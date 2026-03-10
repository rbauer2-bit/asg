#!/usr/bin/env python3
"""
Authority Systems Group™
Roger Bauer — Weekly Focus Brief
March 8, 2026 | Internal Document
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT     = RGBColor(0x1A, 0x25, 0x33)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
LIGHT_BG      = 'F5F5F5'
TINT_BG       = 'EAF6FC'
DIVIDER_HEX   = 'DDDDDD'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'E67E22'
RED_HEX       = 'C0392B'
SLATE_HEX     = '2C3E50'

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_weekly-focus_20260308.docx')


# ── XML / Layout Helpers ───────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_pb(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)
    return p


def section_hdr(doc, text, bg=ASG_DARK_HEX, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    set_para_bg(p, bg)
    r = p.add_run(f'  {text}  ')
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = WHITE; r.font.name = 'Calibri'
    return p


def sub_hdr(doc, text, color=ASG_BLUE, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = color; r.font.name = 'Calibri'
    return p


def body(doc, text, size=10.5, color=BODY_COLOR, bold=False, italic=False, indent=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Calibri'; r.bold = bold; r.italic = italic
    return p


def blt(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = BODY_COLOR; r.font.name = 'Calibri'
    return p


def divider(doc, color=DIVIDER_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1'); bt.set(qn('w:color'), color)
    pBdr.append(bt); pPr.append(pBdr)
    return p


def add_page_number(para):
    run = para.add_run()
    run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin', None), (None, ' PAGE '), ('end', None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0, 0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0, 1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('WEEKLY FOCUS BRIEF  |  MARCH 8, 2026')
    rr.font.size = Pt(8); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True

    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0, 0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Internal Document  —  Confidential')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0, 1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover ──────────────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin', 1.0), ('right_margin', 1.0), ('top_margin', 0.75), ('bottom_margin', 0.75)]:
        setattr(section, attr, Inches(val))

    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(40); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))

    sp(doc, 20)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('INTERNAL BRIEFING — MARCH 8, 2026')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)
    tr = tp.add_run('Weekly Focus Brief')
    tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = DARK_TEXT; tr.font.name = 'Calibri'

    stp = doc.add_paragraph(); stp.paragraph_format.space_before = Pt(0); stp.paragraph_format.space_after = Pt(4)
    str_ = stp.add_run('What to Focus on This Week — Cross-Referenced with Board Meeting + Portfolio Strategy')
    str_.font.size = Pt(13); str_.font.color.rgb = ASG_CHARCOAL; str_.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Director: Roger Bauer     |     Classification: Internal / Confidential')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 20)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(4)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  Source documents: ASG Board Meeting — Weekly Priority Session (March 8, 2026) and '
        'ASG Portfolio Strategy Brief (March 8, 2026). This brief synthesizes both into a single '
        'actionable focus plan for the week, updated to account for the LouisvilleTreePros.com '
        'opportunity and the domain portfolio review.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    add_pb(doc)


# ── Section 1: The Board's Core Message ───────────────────────────────────────

def build_core_message(doc):
    section_hdr(doc, '01 — THE BOARD\'S CORE MESSAGE  |  Still Stands')
    sp(doc, 8)

    body(doc,
        'Before reviewing the specific actions for this week, the board\'s unanimous verdict from this '
        'morning\'s session is worth repeating — because nothing that has happened since changes it:',
        after=10
    )

    vp = doc.add_paragraph(); set_para_bg(vp, SLATE_HEX)
    vp.paragraph_format.space_before = Pt(0); vp.paragraph_format.space_after = Pt(8)
    vp.paragraph_format.left_indent = Inches(0.15); vp.paragraph_format.right_indent = Inches(0.15)
    vr = vp.add_run(
        '  "Revenue this week does not come from a website. It comes from a phone call."  '
    )
    vr.bold = True; vr.font.size = Pt(12); vr.font.color.rgb = WHITE; vr.font.name = 'Calibri'

    body(doc,
        '— Tanya Blackwood, CRO',
        size=9.5, italic=True, color=ASG_CHARCOAL, indent=0.15, after=12
    )

    body(doc,
        'The domain list, the built-out sites, the content assets — these are tools. They are not the '
        'business. The business starts when a real person is on the other end of a conversation, and '
        'Roger is the only one who can make that happen. That is the constraint. Everything else is '
        'downstream of it.',
        after=6
    )

    add_pb(doc)


# ── Section 2: Priority #1 — LouisvilleTreePros Call ─────────────────────────

def build_louisville_tree(doc):
    section_hdr(doc, '02 — PRIORITY #1 THIS WEEK  |  The LouisvilleTreePros.com Call', bg=GREEN_HEX)
    sp(doc, 8)

    # Verdict box
    vp = doc.add_paragraph(); set_para_bg(vp, GREEN_HEX)
    vp.paragraph_format.space_before = Pt(0); vp.paragraph_format.space_after = Pt(8)
    vp.paragraph_format.left_indent = Inches(0.15); vp.paragraph_format.right_indent = Inches(0.15)
    vr = vp.add_run(
        '  ACTION: Call the LouisvilleTreePros.com business owner. Monday morning.  '
    )
    vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = WHITE; vr.font.name = 'Calibri'

    sp(doc, 8)

    sub_hdr(doc, 'Why This Is the #1 Priority')
    body(doc,
        'The board recommended a warm outreach sprint as the top revenue-generating activity this week. '
        'LouisvilleTreePros.com is not a warm outreach opportunity — it is better than that. '
        'The entire value-first framework the team recommends is already complete:',
        after=6
    )
    blt(doc, 'Roger built a site for this business owner without being asked')
    blt(doc, 'Roger installed a tracking phone number and the site is generating real inbound calls')
    blt(doc, 'Roger has recorded call evidence that the site is producing revenue for this business')
    blt(doc, 'The business owner does not yet know any of this exists')

    sp(doc, 8)

    sub_hdr(doc, 'What Makes This Different from Cold Outreach')
    body(doc,
        'The board\'s recommended cold opener is: "I\'m running a limited number of free digital authority '
        'assessments this month for [niche] practices. Want one?" That is a solid opener. '
        'But it requires the prospect to take a leap of faith on the promise of future value.',
        after=6
    )
    body(doc,
        'This situation requires no leap. The value is documented. The calls are recorded. '
        'Roger can open with: "I\'ve been running a pilot with your business in mind and have '
        'some call recordings I\'d like to share with you — do you have 15 minutes this week?" '
        'That is not a sales call. That is a business owner getting a phone call that his phone '
        'is ringing because of work someone else did on his behalf.',
        after=6
    )

    sp(doc, 6)

    sub_hdr(doc, 'The Conversation Path')
    steps = [
        ('Opening', 'Share the evidence — recording count, call volume, the site that\'s driving it. No pitch. Just the facts.'),
        ('The question', '"This is what\'s already happening with a basic setup. Would you like to know what a fully built-out authority site could do for your business?"'),
        ('The offer', 'A Digital Authority Assessment on his current online presence — delivered free, same as the board recommends for LGD and DMA prospects.'),
        ('The close', 'The walkthrough call is the sales call. Let the recorded calls and the assessment do the work.'),
    ]
    t = doc.add_table(len(steps), 2); remove_table_borders(t)
    t.columns[0].width = Inches(1.2); t.columns[1].width = Inches(5.3)
    for i, (label, text) in enumerate(steps):
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        lc = t.cell(i, 0); set_cell_bg(lc, bg)
        lp = lc.paragraphs[0]; lp.paragraph_format.space_before = Pt(4); lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(f'  {label}'); lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = ASG_CHARCOAL; lr.font.name = 'Calibri'
        tc_ = t.cell(i, 1); set_cell_bg(tc_, bg)
        tp = tc_.paragraphs[0]; tp.paragraph_format.space_before = Pt(4); tp.paragraph_format.space_after = Pt(4)
        tr_ = tp.add_run(text); tr_.font.size = Pt(9.5); tr_.font.color.rgb = BODY_COLOR; tr_.font.name = 'Calibri'

    sp(doc, 10)

    body(doc,
        'This is the clearest proof-of-concept in the entire portfolio. A tree service company in '
        'Louisville is receiving calls because of a site Roger built speculatively. If that does not '
        'convert to a paying client, nothing will. Make the call Monday.',
        italic=True, color=ASG_CHARCOAL, after=6
    )

    add_pb(doc)


# ── Section 3: The Expiring Lawyer Domains ────────────────────────────────────

def build_domain_decisions(doc):
    section_hdr(doc, '03 — DOMAIN DECISIONS DUE THIS WEEK  |  Iris Nolan, CTO')
    sp(doc, 8)

    body(doc,
        'Three built-out domains serving the attorney market are expiring in 41 days with auto-renew OFF. '
        'These are not parked domains — they were built out for potential clients. A decision is required '
        'this week. Not next week. The renewal window closes before you will have another planning session.',
        after=10
    )

    # Urgency callout
    up = doc.add_paragraph(); set_para_bg(up, AMBER_HEX)
    up.paragraph_format.space_before = Pt(0); up.paragraph_format.space_after = Pt(8)
    up.paragraph_format.left_indent = Inches(0.15); up.paragraph_format.right_indent = Inches(0.15)
    ur = up.add_run(
        '  All three expire April 18, 2026 — 41 days from today. Auto-renew is OFF. '
        'Decision required this week.  '
    )
    ur.bold = True; ur.font.size = Pt(10); ur.font.color.rgb = WHITE; ur.font.name = 'Calibri'

    sub_hdr(doc, 'The Three Domains')
    domains = [
        ('MyOwensboroLawyer.com', 'Apr 18 2026', 'OFF', 'Built out for a potential attorney client in Owensboro, KY'),
        ('YourLexingtonAttorney.com', 'Apr 18 2026', 'OFF', 'Built out for a potential attorney client in Lexington, KY'),
        ('YourLouisvilleLawyer.com', 'Apr 18 2026', 'OFF', 'Built out for a potential attorney client in Louisville, KY'),
    ]
    dt = doc.add_table(len(domains) + 1, 4); remove_table_borders(dt)
    dt.columns[0].width = Inches(2.2); dt.columns[1].width = Inches(1.0)
    dt.columns[2].width = Inches(0.8); dt.columns[3].width = Inches(2.5)
    hr = dt.rows[0]
    for ci, lbl in enumerate(['Domain', 'Expires', 'Auto-Renew', 'Status']):
        c = hr.cells[ci]; set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
        cr = cp.add_run(lbl); cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = WHITE; cr.font.name = 'Calibri'
    for i, (domain, exp, ar, status) in enumerate(domains):
        row = dt.rows[i + 1]; bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([domain, exp, ar, status]):
            c = row.cells[ci]; set_cell_bg(c, bg)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(val)
            cr.font.size = Pt(9.5); cr.font.color.rgb = BODY_COLOR; cr.font.name = 'Calibri'
            if ci == 0: cr.bold = True

    sp(doc, 10)

    sub_hdr(doc, 'The Decision Framework (Apply to Each Domain)')
    body(doc,
        'Per the Portfolio Strategy Brief, one question determines what to do with each of these:',
        after=6
    )

    qp = doc.add_paragraph(); set_para_bg(qp, SLATE_HEX)
    qp.paragraph_format.space_before = Pt(0); qp.paragraph_format.space_after = Pt(8)
    qp.paragraph_format.left_indent = Inches(0.15); qp.paragraph_format.right_indent = Inches(0.15)
    qr = qp.add_run(
        '  "Do you have a specific attorney in mind for this domain, and are you willing to '
        'pitch them this week or next?"  '
    )
    qr.bold = True; qr.font.size = Pt(10.5); qr.font.color.rgb = WHITE; qr.font.name = 'Calibri'

    body(doc, 'If YES — keep the domain. Make the call this week or next. Use the same LouisvilleTreePros playbook: site is live, it\'s ready for them, you want to show them what\'s possible.', after=5)
    body(doc, 'If NO specific prospect exists — let it expire. $15/year is not the issue. Three "someday" obligations sitting on your list drain focus every time you look at them. Optionality without action is just weight.', after=5)

    sp(doc, 8)

    body(doc,
        'Note: LouisvilleTreePros.com demonstrates the model these domains are designed for. '
        'The attorney domains have the same potential — but only if Roger activates them with '
        'a real prospect conversation. A beautiful site with no conversation attached is just hosting cost.',
        italic=True, color=ASG_CHARCOAL, after=6
    )

    add_pb(doc)


# ── Section 4: Board Priority Summary ─────────────────────────────────────────

def build_priority_summary(doc):
    section_hdr(doc, '04 — FULL PRIORITY SUMMARY  |  Cross-Referenced from Both Documents')
    sp(doc, 8)

    body(doc,
        'The following table consolidates the Board Meeting priorities and Portfolio Strategy actions '
        'into a single weekly execution view. Sequence matters — read the table top to bottom.',
        after=12
    )

    priorities = [
        {
            'num': '01',
            'title': 'Call the LouisvilleTreePros.com business owner',
            'tier': 'DO MONDAY',
            'tier_hex': GREEN_HEX,
            'when': 'Monday AM',
            'action': 'Call with evidence in hand — recording count, call volume, site URL. Offer a Digital Authority Assessment walkthrough. Let the evidence do the work.',
        },
        {
            'num': '02',
            'title': 'Push LGD Services Pages Live',
            'tier': 'DO MONDAY',
            'tier_hex': GREEN_HEX,
            'when': 'Monday',
            'action': 'The copy is written. Log in to LGD WordPress and publish the revised services pages. 1–2 hours max. Every day this sits unpublished is a day a prospect sees the old version.',
        },
        {
            'num': '03',
            'title': 'Decide on the Three Expiring Lawyer Domains',
            'tier': 'BY WEDNESDAY',
            'tier_hex': AMBER_HEX,
            'when': 'Mon–Wed',
            'action': 'Name a specific prospect for each domain or mark for non-renewal. MyOwensboroLawyer.com, YourLexingtonAttorney.com, YourLouisvilleLawyer.com — all expire April 18.',
        },
        {
            'num': '04',
            'title': 'Thomas Banks — Digital Authority Assessment',
            'tier': 'THIS WEEK',
            'tier_hex': AMBER_HEX,
            'when': 'Mid-week',
            'action': 'Run a full Digital Authority Assessment on KentuckianaDivorceAttorney.com and divorceky-in.com. Produce a positioning brief for Thomas. This is the highest-leverage LGD play: one relationship, one firm, six attorneys, multiple future engagements.',
        },
        {
            'num': '05',
            'title': 'Warm Outreach — 10 Implant Dentists + 10 Family Law Attorneys',
            'tier': 'THIS WEEK',
            'tier_hex': AMBER_HEX,
            'when': 'Ongoing',
            'action': 'Identify from Roger\'s existing network. Send a direct, low-friction message: "I\'m running a limited number of free digital authority assessments this month for [niche] practices. Interested?" — no funnel, no pitch, just a door opener.',
        },
        {
            'num': '06',
            'title': 'Post a Upwork Job for a WordPress/Divi 5 Builder',
            'tier': 'THIS WEEK',
            'tier_hex': AMBER_HEX,
            'when': 'Monday–Tuesday',
            'action': 'Roger sets up domains, hosting, Divi 5 install, and blank page structure. Everything after that is builder work. A competent Divi freelancer costs $25–45/hr. At 5 sites, conservative build time is 40–60 hours — that is $1,000–$2,700 in labor Roger does not need to do himself.',
        },
        {
            'num': '07',
            'title': 'Brief the Five Pending Site Builds',
            'tier': 'THIS WEEK',
            'tier_hex': CHARCOAL_HEX,
            'when': 'After builder is hired',
            'action': 'Write one-page brief per site: TopLouisvilleDentist.com, HedgesDental.com, KentuckianaDivorceAttorney.com, ASG site, DMA site. Cover: avatar, positioning, CTA structure, page structure. Hand off to builder with Divi 5 skeleton.',
        },
        {
            'num': '08',
            'title': 'Build a Simple Prospect Pipeline Tracker',
            'tier': 'THIS WEEK',
            'tier_hex': CHARCOAL_HEX,
            'when': '15–30 min',
            'action': 'A Google Sheet with five columns: Name / Niche / Contact Date / Status / Next Step. Set it up before the first outreach call goes out. One missed follow-up is a missed client.',
        },
        {
            'num': '09',
            'title': 'Draft the Direct Mail Packet Copy',
            'tier': 'MEDIUM',
            'tier_hex': CHARCOAL_HEX,
            'when': 'This week — hold for launch',
            'action': 'Draft the letter and envelope teaser. Do not print. Three gates must be open first: (1) live DMA website, (2) clear offer with response mechanism, (3) GHL pipeline to handle responses. Draft this week; mail in Week 3–4.',
        },
    ]

    for p in priorities:
        nt = doc.add_table(1, 2); remove_table_borders(nt)
        nt.columns[0].width = Inches(0.45); nt.columns[1].width = Inches(6.05)

        nc = nt.cell(0, 0); set_cell_bg(nc, p['tier_hex'])
        np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(5); np_.paragraph_format.space_after = Pt(5)
        nr = np_.add_run(p['num'])
        nr.bold = True; nr.font.size = Pt(12); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

        tc_ = nt.cell(0, 1); set_cell_bg(tc_, SLATE_HEX)
        tp = tc_.paragraphs[0]
        tp.paragraph_format.space_before = Pt(5); tp.paragraph_format.space_after = Pt(5)
        tr_ = tp.add_run(p['title'])
        tr_.bold = True; tr_.font.size = Pt(10.5); tr_.font.color.rgb = WHITE; tr_.font.name = 'Calibri'

        meta_t = doc.add_table(1, 3); remove_table_borders(meta_t)
        meta_t.columns[0].width = Inches(1.5); meta_t.columns[1].width = Inches(1.5); meta_t.columns[2].width = Inches(3.5)
        meta_items = [('Tier', p['tier']), ('When', p['when'])]
        for ci, (label, val) in enumerate(meta_items):
            c = meta_t.cell(0, ci); set_cell_bg(c, LIGHT_BG)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
            r1 = cp.add_run(f'{label}: ')
            r1.bold = True; r1.font.size = Pt(8); r1.font.color.rgb = ASG_CHARCOAL; r1.font.name = 'Calibri'
            r2 = cp.add_run(val)
            r2.font.size = Pt(8); r2.font.color.rgb = BODY_COLOR; r2.font.name = 'Calibri'
        # empty third cell for spacing
        set_cell_bg(meta_t.cell(0, 2), LIGHT_BG)

        act_p = doc.add_paragraph(); act_p.paragraph_format.space_before = Pt(0); act_p.paragraph_format.space_after = Pt(0)
        act_p.paragraph_format.left_indent = Inches(0.1); act_p.paragraph_format.right_indent = Inches(0.1)
        a1 = act_p.add_run('  Action: ')
        a1.bold = True; a1.font.size = Pt(9.5); a1.font.color.rgb = RGBColor(0x27, 0xAE, 0x60); a1.font.name = 'Calibri'
        a2 = act_p.add_run(p['action'])
        a2.font.size = Pt(9.5); a2.font.color.rgb = BODY_COLOR; a2.font.name = 'Calibri'

        sp(doc, 10)

    add_pb(doc)


# ── Section 5: The Honest Summary ─────────────────────────────────────────────

def build_closing(doc):
    section_hdr(doc, '05 — THE HONEST SUMMARY  |  Roger Bauer, Director')
    sp(doc, 8)

    body(doc,
        'You are not all over the place. You have real assets with real proof. '
        'The problem is a sequencing issue — not a capability issue.',
        size=11, bold=True, color=DARK_TEXT, after=10
    )

    body(doc,
        'LouisvilleTreePros.com is already doing its job. A business owner\'s phone is ringing '
        'because of a site you built. That is not a portfolio asset sitting dormant — that is '
        'a live demonstration of the entire value proposition. The site is already working. '
        'Now you have to do your part.',
        after=8
    )

    body(doc,
        'The sites built out for potential clients — LouisvilleTreePros, MyOwensboroLawyer, '
        'YourLexingtonAttorney, YourLouisvilleLawyer — are all waiting for the same thing: '
        'a conversation with a real business owner. The site cannot close the client. '
        'Roger closes the client. The site just makes the close easier.',
        after=10
    )

    sub_hdr(doc, 'The Non-Negotiable List for This Week')
    checklist = [
        'Call LouisvilleTreePros.com business owner — Monday',
        'Push LGD services pages live — Monday',
        'Decide on three expiring lawyer domains — by Wednesday',
        'Run Digital Authority Assessment on Thomas Banks\' site — produce the brief',
        'Send 10 warm outreach messages (implant dentists or family law attorneys from existing network)',
        'Post Upwork job for WordPress/Divi 5 builder — stop doing what a $30/hr hire can do',
    ]
    for item in checklist:
        blt(doc, item)

    sp(doc, 16)
    divider(doc)
    sp(doc, 8)

    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(0); ep.paragraph_format.space_after = Pt(4)
    er = ep.add_run(
        'Authority Systems Group™  —  Weekly Focus Brief  —  March 8, 2026\n'
        'Internal Document  |  Confidential  |  Not for Distribution'
    )
    er.font.size = Pt(8); er.font.color.rgb = FOOTER_COLOR; er.italic = True; er.font.name = 'Calibri'


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin', 1.0), ('right_margin', 1.0), ('top_margin', 0.75), ('bottom_margin', 0.75)]:
        setattr(section, attr, Inches(val))

    add_header_footer(doc)
    build_cover(doc)
    build_core_message(doc)
    build_louisville_tree(doc)
    build_domain_decisions(doc)
    build_priority_summary(doc)
    build_closing(doc)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
