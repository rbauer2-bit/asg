#!/usr/bin/env python3
"""
Authority Systems Group™ — Niche Intelligence Brief
Insulation Contractor: Avatar Persona + Pain Points + Insider Intelligence
ASG Internal Asset — Niche Intelligence Library
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK_BG   = '1a1a1a'
ASG_BLUE_HEX  = '25aae1'
ASG_DARK_HEX  = '1A2533'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_TINT      = 'EAF6FC'
SLATE_HEX     = '2C3E50'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'E67E22'
RED_HEX       = '7D2400'
GOLD_HEX      = 'C9A02C'
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_insulation-contractor_avatar-persona_20260309.docx')


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_pb(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)


def set_margins(doc, top=0.75, bottom=0.75, left=0.9, right=0.9):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


# ── Typography Helpers ─────────────────────────────────────────────────────────

def heading(doc, text, level=1, color=None, bg=None, size=None, bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT, caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    if bg:
        set_para_bg(p, bg)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.font.name = 'Calibri'
    r.font.color.rgb = color if color else (WHITE if bg else ASG_CHARCOAL)
    if size:
        r.font.size = Pt(size)
    elif level == 1:
        r.font.size = Pt(20)
    elif level == 2:
        r.font.size = Pt(14)
    elif level == 3:
        r.font.size = Pt(11)
    return p


def body(doc, text, size=10, color=None, bold=False, italic=False, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def bullet(doc, text, size=10, indent=0.2, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(3)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def divider(doc, hex_color=ASG_BLUE_HEX, height_pt=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    set_para_bg(p, hex_color)
    p.paragraph_format.space_before = Pt(height_pt)
    p.paragraph_format.space_after  = Pt(height_pt)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(cover, ASG_DARK_HEX)
    cover.paragraph_format.space_before = Pt(0)
    cover.paragraph_format.space_after  = Pt(0)

    if os.path.exists(LOGO_PATH):
        run = cover.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))

    sp(doc, 4)

    accent = doc.add_paragraph()
    set_para_bg(accent, ASG_BLUE_HEX)
    accent.paragraph_format.space_before = Pt(3)
    accent.paragraph_format.space_after  = Pt(3)

    sp(doc, 10)

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run('NICHE INTELLIGENCE BRIEF')
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = ASG_BLUE
    r1.font.all_caps = True

    sp(doc, 4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Insulation Contractor')
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(32)
    r2.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

    sp(doc, 6)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Avatar Persona  ·  Pain Points  ·  Insider Intelligence  ·  Customer Avatars')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(13)
    r3.font.color.rgb = ASG_CHARCOAL

    sp(doc, 30)

    cred_tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(cred_tbl)
    cred_tbl.columns[0].width = Inches(2.5)
    cred_tbl.columns[1].width = Inches(2.5)
    cred_tbl.columns[2].width = Inches(2.5)

    labels = ['PREPARED BY', 'DOCUMENT TYPE', 'DATE']
    values = ['Authority Systems Group™', 'Internal Niche Intelligence', 'March 9, 2026']

    for i, (lbl, val) in enumerate(zip(labels, values)):
        cell = cred_tbl.rows[0].cells[i]
        set_cell_bg(cell, ASG_TINT)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = cp.add_run(lbl + '\n')
        rl.bold = True
        rl.font.name = 'Calibri'
        rl.font.size = Pt(8)
        rl.font.color.rgb = ASG_BLUE
        rv = cp.add_run(val)
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9)
        rv.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

    sp(doc, 20)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run('Authority Systems Group™  —  Internal Asset. Do not distribute to clients.')
    footer_r.font.name = 'Calibri'
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = FOOTER_COLOR

    add_pb(doc)


# ── Section 1: The Avatar ──────────────────────────────────────────────────────

def build_avatar(doc):
    heading(doc, 'SECTION 1', level=3, color=ASG_BLUE, size=9)
    heading(doc, '"Spray Foam Sam"', level=1, color=RGBColor(0x1A, 0x25, 0x33), size=24, bold=True)
    heading(doc, 'The Insulation Contractor Owner-Operator — ASG Prospect Avatar', level=2,
            color=ASG_CHARCOAL, size=13, bold=False)
    divider(doc)
    sp(doc, 6)

    body(doc,
         'Sam didn\'t go to business school. He started as an installer — pulling batts in new '
         'construction, running hose on a spray rig for someone else\'s company. He got good at it, '
         'realized the margins were real, bought his first rig, and started calling on builders. Fifteen '
         'years later he\'s running a legitimate operation — six to eighteen people, $600K to $3M in '
         'revenue, spray equipment that costs more than most people\'s houses.',
         size=10.5)

    body(doc,
         'The problem is his entire business runs through three or four builder relationships. When '
         'housing starts slow down, his revenue craters — even though there are thousands of homeowners '
         'within twenty miles who would hire him tomorrow if they knew he existed. He has no direct-to-'
         'homeowner engine. He has no referral program. He\'s never used his BPI certification in a '
         'marketing piece. And he has no idea what the IRA 25C tax credit could do for his close rate.',
         size=10.5)

    body(doc,
         'He\'s been pitched by marketing agencies who couldn\'t define spray foam, let alone explain '
         'why closed cell beats open cell in a crawlspace. He\'s skeptical, he\'s smart, and he\'ll '
         'respect the first person who demonstrates real knowledge of his world before trying to sell '
         'him something.',
         size=10.5)

    sp(doc, 4)

    # Profile table
    heading(doc, 'Business Profile', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    profile_data = [
        ('Revenue Range',        '$400,000 – $4,000,000/year'),
        ('Employees',            '2–18 (including owner)'),
        ('Years in Business',    '4–16 years'),
        ('Equipment Investment', '$80,000 – $400,000 (rigs, proportioners, trucks, lifts)'),
        ('Avg Job Value',        '$1,800 – $12,000 (residential); $8,000 – $150,000+ (commercial)'),
        ('CRM / Scheduling',     'None | Jobber | ServiceTitan | Housecall Pro'),
        ('Billing Model',        'Per-job flat fee; square-footage pricing; T&M for commercial'),
        ('Builder Dependency',   '60–90% of revenue from 3–6 GC/builder relationships'),
    ]

    tbl = doc.add_table(rows=len(profile_data), cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(2.4)
    tbl.columns[1].width = Inches(5.1)

    for i, (lbl, val) in enumerate(profile_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(vc, bg)
        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent = Inches(0.1)
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = ASG_CHARCOAL
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.name = 'Calibri'
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = BODY_COLOR

    sp(doc, 8)

    # Revenue Mix
    heading(doc, 'Typical Revenue Mix', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    rev_data = [
        ('Spray Foam Insulation (open & closed cell)',    '45–60%'),
        ('Blown-In / Loose Fill (attic, wall retrofit)',  '20–30%'),
        ('Batt / Roll Installation (new construction)',   '10–15%'),
        ('Air Sealing',                                   '5–10%'),
        ('Vapor / Moisture / Radiant Barriers',           '3–8%'),
        ('Insulation Removal & Replacement',              '3–7%'),
    ]

    rev_tbl = doc.add_table(rows=len(rev_data), cols=2)
    remove_table_borders(rev_tbl)
    rev_tbl.columns[0].width = Inches(5.0)
    rev_tbl.columns[1].width = Inches(2.5)

    for i, (svc, pct) in enumerate(rev_data):
        bg = ASG_LIGHT_BG if i % 2 == 0 else 'FFFFFF'
        sc = rev_tbl.rows[i].cells[0]
        pc = rev_tbl.rows[i].cells[1]
        set_cell_bg(sc, bg)
        set_cell_bg(pc, bg)
        sp2 = sc.paragraphs[0]
        sp2.paragraph_format.left_indent = Inches(0.1)
        sr = sp2.add_run(svc)
        sr.font.name = 'Calibri'
        sr.font.size = Pt(9.5)
        sr.font.color.rgb = BODY_COLOR
        pp = pc.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = pp.add_run(pct)
        pr.bold = True
        pr.font.name = 'Calibri'
        pr.font.size = Pt(9.5)
        pr.font.color.rgb = ASG_BLUE

    sp(doc, 8)

    # Credentials
    heading(doc, 'Credentials That Matter', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    creds = [
        ('BPI Building Analyst / Envelope Professional',
         'Requires documented field training and a proctored exam. The most credible homeowner-facing '
         'credential in building science. Almost no insulation contractor leads with it. '
         'Most homeowners have never heard of it — which is the opportunity.'),
        ('RESNET HERS Rater',
         'Required for new construction energy code compliance in many jurisdictions. '
         'Positions the contractor as the technical authority, not just the installer.'),
        ('Manufacturer Certifications (Icynene, Lapolla, Demilec, Huntsman)',
         'Required to honor product warranties on spray foam jobs. Differentiates from '
         'uncertified applicators using off-brand or expired material.'),
        ('General Liability $1M+ and Workers\' Compensation',
         'Critical differentiator vs. uninsured subs. Most homeowners never ask for the '
         'certificate — an operator who proactively provides it wins authority immediately.'),
        ('Energy Star Partner / DOE Home Performance',
         'Strong positioning lever for the retrofit market. Connects the contractor to '
         'IRA credit eligibility, utility rebate programs, and energy-conscious buyers.'),
    ]

    for cred_name, cred_desc in creds:
        cred_tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(cred_tbl)
        cred_tbl.columns[0].width = Inches(7.5)
        cell = cred_tbl.rows[0].cells[0]
        set_cell_bg(cell, ASG_TINT)
        cp = cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.12)
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(4)
        rn = cp.add_run(cred_name + '  ')
        rn.bold = True
        rn.font.name = 'Calibri'
        rn.font.size = Pt(9.5)
        rn.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)
        rd = cp.add_run(cred_desc)
        rd.font.name = 'Calibri'
        rd.font.size = Pt(9.5)
        rd.font.color.rgb = BODY_COLOR
        sp(doc, 3)

    sp(doc, 6)

    # Belief barriers
    heading(doc, 'Why He Hesitates to Hire ASG', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    barriers = [
        ('Primary Enemy Belief',
         '"Marketing agencies don\'t understand insulation — they\'ll write generic '
         '\'save on energy bills\' ads that make me look like every other guy."',
         RED_HEX),
        ('Secondary Enemy Belief',
         '"My builder relationships are solid. The work comes through them. '
         'I don\'t need to market direct to homeowners."',
         AMBER_HEX),
        ('Normalized Pain',
         '"Slow periods are just part of the construction cycle. You learn to manage through them."',
         SLATE_HEX),
        ('Urgency Blocker',
         '"I\'d rather put that money into a second spray rig than spend it on marketing '
         'I can\'t measure."',
         SLATE_HEX),
        ('Trust Trigger',
         'Reference spray foam chemistry specifics, IRA 25C tax credit mechanics, utility rebate '
         'stacking, or energy code changes by name. Watch the conversation shift immediately.',
         GREEN_HEX),
    ]

    for btype, btext, bhex in barriers:
        btbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(btbl)
        btbl.columns[0].width = Inches(0.12)
        btbl.columns[1].width = Inches(7.38)
        set_cell_bg(btbl.rows[0].cells[0], bhex)
        content_cell = btbl.rows[0].cells[1]
        set_cell_bg(content_cell, ASG_LIGHT_BG)
        cp = content_cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.12)
        cp.paragraph_format.space_before = Pt(5)
        cp.paragraph_format.space_after  = Pt(5)
        rl = cp.add_run(btype + '  ')
        rl.bold = True
        rl.font.name = 'Calibri'
        rl.font.size = Pt(9)
        rl.font.color.rgb = ASG_CHARCOAL
        rv = cp.add_run(btext)
        rv.italic = True
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = BODY_COLOR
        sp(doc, 3)

    add_pb(doc)


# ── Section 2: 10 Pain Points ──────────────────────────────────────────────────

PAIN_POINTS = [
    {
        'number': '01',
        'title': 'The Builder Dependency Trap',
        'body': (
            'He built his business on three or four general contractor relationships. When they\'re '
            'building, he\'s busy. When new construction slows — rate cycle, permitting freeze, '
            'material costs — his revenue craters. He has no direct-to-homeowner pipeline, no '
            'backup source of work, and no marketing system that operates independently of whoever '
            'is handing him the next job.'
        ),
        'agency_insight': (
            'This is the single most dangerous structural vulnerability in the insulation contractor '
            'business model and almost no agency identifies it in the first meeting. Sam doesn\'t '
            'see it as a marketing problem — he sees it as business as usual. The framing that opens '
            'the conversation: "What happens to your revenue when one of your top two builders takes '
            'on 40% fewer starts?" Watch him go quiet.'
        ),
        'color': RED_HEX,
    },
    {
        'number': '02',
        'title': 'The Commodity Comparison Problem',
        'body': (
            'Homeowners have no framework for evaluating an insulation quote. They see three numbers '
            'on three pieces of paper and pick the middle one. They have no idea what closed cell vs. '
            'open cell means, why the ratio of chemicals matters, why BPI certification exists, or '
            'why a spray foam job done wrong smells like cat urine for six months. Sam does better '
            'work than his competitors. Nobody knows it.'
        ),
        'agency_insight': (
            'The insulation industry has a massive education deficit at the consumer level. This is '
            'a positioning opportunity, not a price problem. The contractor who publishes a single '
            'clear piece of content — "Why two insulation quotes can be $3,000 apart for the same '
            'job" — immediately becomes the authority in a market full of operators who hide behind '
            'their quote sheet.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'number': '03',
        'title': 'The IRA Credit He\'s Not Using as a Close Tool',
        'body': (
            'The Inflation Reduction Act 25C tax credit covers 30% of the cost of qualified '
            'insulation and air sealing upgrades, up to $1,200 per year through 2032. Sam probably '
            'knows it exists. He almost certainly doesn\'t have it on his estimate sheet, doesn\'t '
            'mention it in his marketing, and doesn\'t walk homeowners through how to stack it with '
            'their utility rebate. He is leaving a $400–$1,800 per-job closing tool on the table '
            'on every residential retrofit job.'
        ),
        'agency_insight': (
            'A $4,500 attic job becomes $3,100–$3,500 net after stacking the 25C credit and a '
            'typical utility rebate. That reframe, built into the estimate conversation and the '
            'website, changes close rates on mid-ticket jobs immediately. No generalist agency knows '
            'this. Dropping it in the first meeting is an instant authority signal.'
        ),
        'color': GREEN_HEX,
    },
    {
        'number': '04',
        'title': 'The Estimate That Dies Without a Follow-Up',
        'body': (
            'He quotes 40–60% of his residential leads. He closes maybe 35–45% of those. The rest '
            'just disappear — some went with a cheaper guy, some are still thinking about it, some '
            'decided to wait until next year. He has no follow-up sequence, no automated email, no '
            'system to send the tax credit information three days after the estimate. Every dead '
            'estimate is $0 recovered from a prospect who already raised their hand.'
        ),
        'agency_insight': (
            'A structured 3-touch follow-up sequence for unconverted estimates — day 3 (rebate info), '
            'day 10 (social proof), day 21 (seasonal urgency) — can recover 10–20% of dead estimates '
            'with zero additional ad spend. Almost no insulation contractor has built this. It is '
            'also the highest-quality lead you can work because the person has already expressed '
            'interest and given you their contact information.'
        ),
        'color': AMBER_HEX,
    },
    {
        'number': '05',
        'title': 'The Air Sealing Story He Never Tells',
        'body': (
            'Up to 40% of heating and cooling loss in a typical home is through air infiltration — '
            'gaps, penetrations, recessed lights, unsealed chases — not through insufficient '
            'R-value. Sam knows this. He seals before he sprays. But he has never once written a '
            'piece of content about it, put it on his estimate, or explained it to a homeowner in '
            'a way that differentiated his process from the guy who just rolls in and blows '
            'cellulose without sealing anything.'
        ),
        'agency_insight': (
            '"We seal first, then insulate" is a positioning statement, not a technical footnote. '
            'Contractors who explain this upfront in their marketing deliver demonstrably better '
            'results, can justify a price premium, and close against lower-price competitors on '
            'value rather than number. The phrase "air sealing + insulation" as a named process — '
            'rather than just "insulation" — is almost entirely absent from competitor websites '
            'in every market we\'ve examined.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'number': '06',
        'title': 'The Referral Network That\'s Sitting Unused',
        'body': (
            'HVAC technicians are in homes constantly and regularly diagnose comfort problems that '
            'are actually insulation failures. Home inspectors flag inadequate insulation on every '
            'third report. Real estate agents want energy-efficient listings to move faster. All of '
            'these people are encountering Sam\'s ideal customers every single day. He\'s never '
            'sent a formal letter to one of them, offered a referral arrangement, or created a '
            'simple card they could hand to homeowners.'
        ),
        'agency_insight': (
            'A single HVAC company referral relationship — one company whose techs understand '
            'what to say and know they\'ll get a referral back — can generate $120,000–$250,000/year '
            'in warm leads. Three such relationships built and systematized is more powerful than '
            'any digital campaign. The cost is near zero. Almost no insulation contractor has '
            'formalized a single one of these relationships professionally.'
        ),
        'color': SLATE_HEX,
    },
    {
        'number': '07',
        'title': 'The Spray Rig Sitting Idle Between Commercial Jobs',
        'body': (
            'His equipment costs are fixed whether the rig is running or not. Chemical prices have '
            'spiked. Payroll doesn\'t stop. In the shoulder season — late fall, winter, early '
            'spring — there\'s a two-week gap between commercial jobs and the residential phone '
            'isn\'t ringing because no one knows to call him. He fills this with anything he can '
            'find. He should be filling it with a systematic homeowner pipeline.'
        ),
        'agency_insight': (
            'The equipment cost structure of spray foam is punishing when idle. A rig that costs '
            '$150,000+ and $80–$120/board-foot in material has to run. Seasonal promotional '
            'campaigns — "winter crawlspace encapsulation," "pre-summer attic package" — positioned '
            'correctly fill the calendar at full rate. The operators who market these windows early '
            '(6–8 weeks before the season) fill their slow calendar before the phone stops ringing.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'number': '08',
        'title': 'The Crawlspace Encapsulation He Undersells',
        'body': (
            'A full crawlspace encapsulation — liner, closed cell rim joist, conditioning or '
            'dehumidification — is a $5,000–$18,000 project that solves moisture, pest exclusion, '
            'air quality, and comfort simultaneously. The emotional urgency is high (mold, musty '
            'smell, cold floors). The conversion rate is strong. Sam does this work. He doesn\'t '
            'market it as a distinct premium service with its own page, its own content, or its '
            'own sales process. It\'s buried under "insulation services."'
        ),
        'agency_insight': (
            'Crawlspace encapsulation is the highest-ticket, highest-margin residential service in '
            'the insulation contractor\'s portfolio and it is almost universally under-positioned. '
            'Homeowners with moisture problems, cold floors, or musty odors are emotionally motivated '
            'buyers — they are not price shopping, they are problem-solving. A dedicated crawlspace '
            'page, thermal imaging before-and-after content, and a named service process ("Complete '
            'Crawlspace Solution" or similar) converts these buyers at a significantly higher rate.'
        ),
        'color': AMBER_HEX,
    },
    {
        'number': '09',
        'title': 'The BPI Credential He Never Mentions',
        'body': (
            'Sam may hold a BPI Building Analyst certification. If he does, it is almost certainly '
            'not on his website homepage, not in his estimate template, not in his Google Business '
            'Profile description, and not in any conversation with a homeowner. He earned it. He '
            'paid for it. He doesn\'t use it. Most homeowners have never heard of BPI. That is '
            'not a reason to hide the credential — it\'s a reason to explain it.'
        ),
        'agency_insight': (
            'BPI Building Analyst certification requires documented training, a proctored written '
            'exam, and a field performance test. It is a meaningful credential in building science. '
            'The first insulation contractor in most markets who explains this in plain terms — '
            '"here\'s what it means, here\'s why a certified analyst sees your home differently '
            'than a standard installer" — owns the authority position. The search term "BPI '
            'certified insulation [city]" is almost entirely uncontested.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'number': '10',
        'title': 'The "Just Insulation" Identity Tax',
        'body': (
            'Sam describes himself as "an insulation contractor." He does not describe himself as '
            'a certified building performance specialist who reduces energy consumption, eliminates '
            'moisture problems, and permanently solves comfort issues that HVAC alone can\'t fix. '
            'He undersells the scope of what a quality spray foam and air sealing job actually '
            'delivers — and every piece of marketing he has is built around the material, not the outcome.'
        ),
        'agency_insight': (
            'The positioning shift from "insulation contractor" to "home performance specialist" '
            'is not cosmetic — it changes the price conversation, the comparison set, and the '
            'close rate. Homeowners don\'t want insulation. They want a comfortable house, lower '
            'bills, and air they can breathe. The contractor who leads with outcomes rather than '
            'materials occupies a different category and cannot be directly compared to the guy '
            'with a cheaper quote. Almost nobody in this niche has made this shift deliberately.'
        ),
        'color': RED_HEX,
    },
]


def build_pain_points(doc):
    heading(doc, 'SECTION 2', level=3, color=ASG_BLUE, size=9)
    heading(doc, '10 Pain Points for Growing the Business', level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'What keeps Sam up at night — and what every agency misses', level=2,
            color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 8)

    for pp in PAIN_POINTS:
        num_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(num_tbl)
        num_tbl.columns[0].width = Inches(0.65)
        num_tbl.columns[1].width = Inches(6.85)

        num_cell = num_tbl.rows[0].cells[0]
        ttl_cell = num_tbl.rows[0].cells[1]
        set_cell_bg(num_cell, pp['color'])
        set_cell_bg(ttl_cell, pp['color'])

        np_ = num_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(6)
        np_.paragraph_format.space_after  = Pt(6)
        nr = np_.add_run(pp['number'])
        nr.bold = True
        nr.font.name = 'Calibri'
        nr.font.size = Pt(15)
        nr.font.color.rgb = WHITE

        tp = ttl_cell.paragraphs[0]
        tp.paragraph_format.left_indent  = Inches(0.1)
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(6)
        tr = tp.add_run(pp['title'])
        tr.bold = True
        tr.font.name = 'Calibri'
        tr.font.size = Pt(13)
        tr.font.color.rgb = WHITE

        body_tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(body_tbl)
        body_tbl.columns[0].width = Inches(7.5)
        bc = body_tbl.rows[0].cells[0]
        set_cell_bg(bc, 'FFFFFF')
        bp = bc.paragraphs[0]
        bp.paragraph_format.left_indent  = Inches(0.15)
        bp.paragraph_format.right_indent = Inches(0.15)
        bp.paragraph_format.space_before = Pt(7)
        bp.paragraph_format.space_after  = Pt(4)
        br_ = bp.add_run(pp['body'])
        br_.font.name = 'Calibri'
        br_.font.size = Pt(10)
        br_.font.color.rgb = BODY_COLOR

        ai_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(ai_tbl)
        ai_tbl.columns[0].width = Inches(0.08)
        ai_tbl.columns[1].width = Inches(7.42)
        set_cell_bg(ai_tbl.rows[0].cells[0], ASG_BLUE_HEX)
        ai_cell = ai_tbl.rows[0].cells[1]
        set_cell_bg(ai_cell, ASG_TINT)
        ap = ai_cell.paragraphs[0]
        ap.paragraph_format.left_indent  = Inches(0.15)
        ap.paragraph_format.right_indent = Inches(0.15)
        ap.paragraph_format.space_before = Pt(6)
        ap.paragraph_format.space_after  = Pt(7)
        al = ap.add_run('ASG Intelligence  ')
        al.bold = True
        al.font.name = 'Calibri'
        al.font.size = Pt(8.5)
        al.font.color.rgb = ASG_BLUE
        ai = ap.add_run(pp['agency_insight'])
        ai.font.name = 'Calibri'
        ai.font.size = Pt(9.5)
        ai.italic = True
        ai.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

        sp(doc, 10)

    add_pb(doc)


# ── Section 3: Insider Intelligence ───────────────────────────────────────────

INSIDER_FACTS = [
    {
        'topic': 'IRA 25C Tax Credit — The Closing Tool Nobody Is Using',
        'insight': (
            'The Inflation Reduction Act extended and expanded the 25C energy efficiency tax credit '
            'through 2032. Homeowners can claim 30% of the cost of insulation and air sealing '
            'upgrades up to $1,200/year. A $4,500 attic job nets to $3,150 after the credit. Stacked '
            'with a utility rebate ($200–$800 in most territories), the effective cost drops further. '
            'Almost no insulation contractor mentions this on their website or during the estimate. '
            'This single fact, deployed in a first conversation with Sam, immediately separates ASG '
            'from every generalist agency he has ever spoken with.'
        ),
        'color': GREEN_HEX,
    },
    {
        'topic': 'Open Cell vs. Closed Cell — The Explanation Gap',
        'insight': (
            'Closed cell spray foam is denser (2 lb/ft³), achieves R-6.5 per inch, acts as a vapor '
            'barrier, and is required for crawlspaces, rim joists, and conditioned envelope applications. '
            'Open cell is less expensive, achieves R-3.7 per inch, and is appropriate for interior '
            'walls and attics with managed ventilation. Contractors who explain this during the '
            'estimate earn immediate authority over competitors who just quote a number. The '
            'explanation itself is the differentiation. Most competitor websites contain neither term.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'topic': 'The Air Sealing Multiplier — The 40% Nobody Talks About',
        'insight': (
            'Up to 40% of heating and cooling loss in a typical home is through air infiltration — '
            'unsealed penetrations, recessed lights, HVAC chases, rim joists, attic bypasses — not '
            'through insufficient R-value. A contractor who seals first and insulates second delivers '
            'demonstrably better performance. "We seal before we spray" is a positioning statement, '
            'not a technical footnote. It justifies a price premium and makes direct price comparison '
            'with unsealing competitors structurally impossible.'
        ),
        'color': AMBER_HEX,
    },
    {
        'topic': 'Thermal Imaging as a Conversion Tool',
        'insight': (
            'A $300–$600 FLIR thermal camera is one of the highest-ROI tools in an insulation '
            'contractor\'s sales process. Showing a homeowner a thermal image of their attic '
            'floor or crawlspace — with red/orange heat loss clearly visible — converts skeptics '
            'into buyers faster than any verbal explanation. Before-and-after thermal imaging in '
            'content (website, social, YouTube) is nearly nonexistent among competitors in any '
            'market. The contractor who publishes this content owns the visual authority.'
        ),
        'color': RED_HEX,
    },
    {
        'topic': 'Crawlspace Encapsulation — The Underdeveloped Premium Service',
        'insight': (
            'A full crawlspace encapsulation — closed cell rim joist, vapor liner, conditioning '
            'or dehumidification — is a $5,000–$18,000 project that solves moisture, pest exclusion, '
            'air quality, and comfort simultaneously. The buyer is emotionally motivated (mold, '
            'cold floors, musty smell) and not price-shopping. Conversion rate is high. Almost no '
            'insulation contractor positions this as a distinct named service with its own page and '
            'content strategy. It is typically buried under a generic "services" list.'
        ),
        'color': SLATE_HEX,
    },
    {
        'topic': 'HVAC Contractor Referral Economics — One Relationship vs. Thousands in Ad Spend',
        'insight': (
            'HVAC technicians diagnose comfort problems daily that are actually insulation and air '
            'sealing failures. A formal referral arrangement with 3–5 HVAC companies — where their '
            'techs understand what to say and know they will receive referrals back — can generate '
            '$120,000–$250,000/year in warm insulation leads with near-zero acquisition cost. '
            'Almost no insulation contractor has formalized a single HVAC referral relationship '
            'professionally. This is the highest-ROI growth lever in the business.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'topic': 'Home Inspector Pipeline — Pre-Qualified, Inspection-Triggered Leads',
        'insight': (
            'Home inspectors flag inadequate insulation on inspection reports that go directly to '
            'the buyer and seller. A homeowner calling from an inspection report has already been '
            'told by a trusted third party that they have a problem. Close rates are significantly '
            'higher than cold digital leads. A relationship with 5 active home inspectors generates '
            'a predictable stream of these pre-qualified, urgency-triggered leads. Almost no '
            'insulation contractor in any market has systematically approached a home inspector '
            'with a formal referral offer.'
        ),
        'color': GREEN_HEX,
    },
    {
        'topic': 'Energy Code Changes — The Authority Positioning Most Contractors Ignore',
        'insight': (
            'Most jurisdictions have adopted IECC 2021 or similar codes requiring R-49–R-60 attic '
            'insulation and specific air leakage testing for new construction. Renovations requiring '
            'permits often trigger updated code compliance. An insulation contractor who understands '
            'local energy code requirements and can explain compliance implications to builders and '
            'homeowners owns a positioning advantage that no generalist competitor can replicate '
            'without doing the homework. Reference the specific code version for the client\'s '
            'jurisdiction by name in any prospect meeting.'
        ),
        'color': AMBER_HEX,
    },
    {
        'topic': 'Real Estate Pre-Listing Positioning — A Revenue Channel Nobody Has Built',
        'insight': (
            'Homes with documented energy efficiency upgrades achieve higher appraised values and '
            'sell faster in energy-conscious markets. A pre-listing home performance report from a '
            'BPI-certified contractor is a legitimate, differentiating listing tool for real estate '
            'agents. A formal relationship with 5–10 active real estate agents generates pre-sale '
            'leads and post-inspection leads simultaneously. This is an entirely unclaimed channel '
            'in virtually every market — no competitor has built it.'
        ),
        'color': RED_HEX,
    },
    {
        'topic': 'Utility Rebate Stacking — The Calculator That Changes Close Rate',
        'insight': (
            'In most utility territories, attic insulation and air sealing qualify for rebates of '
            '$200–$800 on top of the 25C federal credit. A contractor who builds a simple '
            'calculation into their estimate — "here\'s the job cost, here\'s the rebate, here\'s '
            'the credit, here\'s your net cost and payback period" — changes the conversation from '
            '"is this worth it?" to "when do you want to schedule?" Almost no competitor presents '
            'this calculation. It is the most immediately deployable close-rate improvement in the '
            'residential insulation sales process.'
        ),
        'color': SLATE_HEX,
    },
]


def build_insider_intel(doc):
    heading(doc, 'SECTION 3', level=3, color=ASG_BLUE, size=9)
    heading(doc, 'Insider Intelligence', level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'What separates ASG from every generalist agency that has ever pitched an insulation contractor',
            level=2, color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 4)

    body(doc,
         'Drop any one of the following facts into a first conversation with Sam and you immediately '
         'separate yourself from every agency that has ever pitched him. These are the things a '
         'consultant who has worked with home performance contractors for years knows cold. Use them. '
         'They are your authority signal.',
         size=10.5, italic=True, color=ASG_CHARCOAL)

    sp(doc, 8)

    for i, fact in enumerate(INSIDER_FACTS):
        fact_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(fact_tbl)
        fact_tbl.columns[0].width = Inches(0.18)
        fact_tbl.columns[1].width = Inches(7.32)

        set_cell_bg(fact_tbl.rows[0].cells[0], fact['color'])
        content_cell = fact_tbl.rows[0].cells[1]
        set_cell_bg(content_cell, ASG_LIGHT_BG if i % 2 == 0 else 'FFFFFF')

        cp = content_cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.15)
        cp.paragraph_format.right_indent = Inches(0.15)
        cp.paragraph_format.space_before = Pt(7)
        cp.paragraph_format.space_after  = Pt(8)

        rt = cp.add_run(fact['topic'] + '\n')
        rt.bold = True
        rt.font.name = 'Calibri'
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

        ri = cp.add_run(fact['insight'])
        ri.font.name = 'Calibri'
        ri.font.size = Pt(9.5)
        ri.font.color.rgb = BODY_COLOR

        sp(doc, 5)

    add_pb(doc)


# ── Section 4: Customer Avatars ─────────────────────────────────────────────────

def build_customer_avatars(doc):
    heading(doc, 'SECTION 4', level=3, color=ASG_BLUE, size=9)
    heading(doc, "Sam's Customer Avatars", level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'Who Sam sells to — and what actually moves them to buy',
            level=2, color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 8)

    avatars = [
        {
            'name': 'High-Bill Harold',
            'service': 'Attic Air Sealing + Blown-In Retrofit',
            'color': AMBER_HEX,
            'profile': (
                'Male skew, 38–62. He opens the utility bill. He\'s been absorbing the cost for '
                'years. A neighbor\'s experience, an HVAC tech\'s comment, or a tax preparer '
                'mentioning the 25C credit finally triggers action. He wants ROI he can explain '
                'to his wife and a contractor who treats him like an intelligent adult.'
            ),
            'triggers': [
                'Record utility bill arrives — "we can\'t keep doing this"',
                'Neighbor cut their bill by $200/month after insulation work',
                'HVAC tech says the system is fine — it\'s the insulation and air sealing',
                'Home energy audit flags attic as the primary loss point',
                'IRS preparer mentions the 25C credit — Harold didn\'t know it existed',
            ],
            'fears': [
                'Spending thousands and not seeing a meaningful bill reduction',
                'Being sold more than he needs by a contractor chasing a ticket target',
                'Job done wrong and having to redo it in five years',
                'Disrupting the house for a project that\'s invisible and hard to verify',
            ],
            'hooks': [
                '"Why your HVAC runs constantly but your house is still uncomfortable"',
                '"How to stack the 25C tax credit + your utility rebate to cut cost by 30–40%"',
                '"Why two insulation quotes can be $3,000 apart for the same job"',
                '"What a good insulation job actually looks like — and what a bad one looks like later"',
            ],
        },
        {
            'name': 'Comfort Problem Carol',
            'service': 'Closed Cell Spray Foam — Retrofit / Crawlspace Encapsulation',
            'color': RED_HEX,
            'profile': (
                'Female primary, 36–58. She\'s the one most affected by the uncomfortable room, '
                'the cold floor, the musty smell. The spouse is dismissing it. She\'s not. HVAC '
                'tech found no mechanical explanation. A friend got spray foam in their crawlspace '
                'and raves about it. She wants a permanently comfortable home — not another band-aid.'
            ),
            'triggers': [
                'One room always 10 degrees hotter or colder than the rest',
                'Cold floors in winter — kids refuse to play on the floor',
                'Musty smell or visible moisture under the house',
                'Allergies or respiratory symptoms she connects to the house\'s air quality',
                'HVAC tech finds no mechanical explanation — suggests the building envelope',
            ],
            'fears': [
                'Mold under the house spreading into a $50,000 remediation problem',
                'Spray foam installed incorrectly, off-ratio, causing odor issues',
                'Contractor who doesn\'t explain what they\'re doing or why',
                'Disturbing the crawlspace and finding something worse than expected',
            ],
            'hooks': [
                '"Why that one room is always hot (or cold) — and it\'s not your HVAC system"',
                '"Open cell vs. closed cell spray foam: which one do you actually need?"',
                '"The 3 signs your crawlspace is hurting your indoor air quality right now"',
                '"Crawlspace encapsulation vs. vented crawlspace: the complete homeowner\'s guide"',
            ],
        },
        {
            'name': 'Builder Brad',
            'service': 'New Construction Insulation Package',
            'color': SLATE_HEX,
            'profile': (
                'Heavily male, 34–54. Running 15–80+ starts per year. His current insulation '
                'sub is unreliable, can\'t scale, or doesn\'t do spray foam. He needs a vendor '
                'who shows up, does it right, invoices correctly, and doesn\'t require babysitting. '
                'He\'ll negotiate pricing hard — the relationship value must be clear beyond the number.'
            ),
            'triggers': [
                'Current insulation sub no-showing and holding up dry-in schedule',
                'Scaling build volume and current sub can\'t handle it',
                'Buyer demanded spray foam; current sub can\'t deliver it',
                'Energy code change forcing tighter envelope specs',
                'Callback complaint about comfort in a completed home',
            ],
            'fears': [
                'Sub no-showing and holding up the entire construction schedule',
                'Callback liability if insulation is wrong in completed homes',
                'A sub who can\'t scale with his growth',
                'Getting burned on a new relationship mid-season',
            ],
            'hooks': [
                '"What builders need from their insulation sub — that most subs don\'t deliver"',
                '"How the 2025 energy code changes affect your insulation spec"',
                '"The insulation mistake that generates the most builder callbacks"',
                '"Spray foam on new construction: is the premium worth it for your buyers?"',
            ],
        },
        {
            'name': 'Investment Property Ivan',
            'service': 'Attic / Wall Retrofit + Air Sealing',
            'color': ASG_DARK_HEX,
            'profile': (
                'Male primary, 36–60. Owns rental units, small multifamily, or investment '
                'properties. Tenant complaints about utility bills, a HUD/Section 8 inspection '
                'flag, or a renovation with open walls creates the trigger. He needs ROI math, '
                'written documentation, and a contractor who can handle multiple units '
                'without being supervised. The 179D commercial deduction is his attention-getter.'
            ),
            'triggers': [
                'Tenant complaints about high utility bills reducing ability to raise rent',
                'Section 8 / HUD inspection flagging inadequate insulation',
                'Renovation underway — wants to insulate while walls are open',
                'IRA 179D commercial deduction or 25C credit creates ROI justification',
                'Insurance or energy audit on a commercial-to-residential conversion',
            ],
            'fears': [
                'Capital outlay not recovered through rent premium or sale price increase',
                'Tenant disruption during the project on occupied units',
                'Contractor delays extending beyond planned vacancy window',
                'Being oversold on spray foam when blown-in achieves 90% at 60% of cost',
            ],
            'hooks': [
                '"How insulation upgrades affect rental income, vacancy rates, and property value"',
                '"The 179D deduction and 25C credit: what property owners need to know"',
                '"Why HUD and Section 8 inspections increasingly flag insulation"',
                '"The ROI calculation for energy efficiency upgrades in rental properties"',
            ],
        },
    ]

    for avatar in avatars:
        av_header = doc.add_table(rows=1, cols=1)
        remove_table_borders(av_header)
        av_header.columns[0].width = Inches(7.5)
        hc = av_header.rows[0].cells[0]
        set_cell_bg(hc, avatar['color'])
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.18)
        hp.paragraph_format.space_before = Pt(7)
        hp.paragraph_format.space_after  = Pt(7)
        hn = hp.add_run(avatar['name'])
        hn.bold = True
        hn.font.name = 'Calibri'
        hn.font.size = Pt(13)
        hn.font.color.rgb = WHITE
        hs = hp.add_run('   ·   Primary Service: ' + avatar['service'])
        hs.font.name = 'Calibri'
        hs.font.size = Pt(10)
        hs.font.color.rgb = RGBColor(0xCC, 0xEE, 0xF8)

        body_tbl = doc.add_table(rows=1, cols=3)
        remove_table_borders(body_tbl)
        body_tbl.columns[0].width = Inches(2.8)
        body_tbl.columns[1].width = Inches(2.35)
        body_tbl.columns[2].width = Inches(2.35)

        col_bgs     = [ASG_TINT, ASG_LIGHT_BG, 'FFFFFF']
        col_labels  = ['WHO THEY ARE', 'WHAT THEY FEAR', 'HOOKS THAT WORK']
        col_content = [
            [avatar['profile']] + ['• ' + t for t in avatar['triggers']],
            ['• ' + f for f in avatar['fears']],
            [h for h in avatar['hooks']],
        ]

        for ci in range(3):
            cell = body_tbl.rows[0].cells[ci]
            set_cell_bg(cell, col_bgs[ci])
            p0 = cell.paragraphs[0]
            p0.paragraph_format.left_indent  = Inches(0.1)
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after  = Pt(2)
            rl = p0.add_run(col_labels[ci] + '\n')
            rl.bold = True
            rl.font.name = 'Calibri'
            rl.font.size = Pt(8)
            rl.font.color.rgb = ASG_BLUE

            for line in col_content[ci]:
                lp = cell.add_paragraph()
                lp.paragraph_format.left_indent  = Inches(0.1)
                lp.paragraph_format.space_before = Pt(0)
                lp.paragraph_format.space_after  = Pt(3)
                lr = lp.add_run(line)
                lr.font.name = 'Calibri'
                lr.font.size = Pt(9)
                lr.font.color.rgb = BODY_COLOR

        sp(doc, 12)

    divider(doc, ASG_BLUE_HEX)
    sp(doc, 4)

    close_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(close_tbl)
    close_tbl.columns[0].width = Inches(7.5)
    cc = close_tbl.rows[0].cells[0]
    set_cell_bg(cc, ASG_TINT)
    cp2 = cc.paragraphs[0]
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.paragraph_format.space_before = Pt(10)
    cp2.paragraph_format.space_after  = Pt(10)
    cr = cp2.add_run(
        'Sam isn\'t hard to impress. He\'s been pitched by generalists his whole career.\n'
        'Walk in knowing the difference between open cell and closed cell, why air sealing matters '
        'as much as R-value,\nhow the 25C credit stacks with utility rebates, and what one HVAC '
        'referral relationship is actually worth —\nand he\'ll tell his wife that night he finally '
        'talked to someone who gets it.\n\nThat\'s authority.'
    )
    cr.bold = True
    cr.font.name = 'Calibri'
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)


# ── Footer ─────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        fr = fp.add_run(
            'Authority Systems Group™  —  Internal Asset  ·  Niche Intelligence Library  '
            '·  insulation-contractor  ·  2026-03-09'
        )
        fr.font.name = 'Calibri'
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = FOOTER_COLOR


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    build_cover(doc)
    build_avatar(doc)
    build_pain_points(doc)
    build_insider_intel(doc)
    build_customer_avatars(doc)

    doc.save(OUT_DOCX)
    print(f'✓  Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
