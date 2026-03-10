"""
ASG Internal Strategic Assessment — ColdOutreachEngine.com
Build script: generates branded DOCX
Output: /outputs/asg_coldoutreachengine_internal-assessment_20260308.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Constants ─────────────────────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
ASG_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ASG_LIGHT    = RGBColor(0xF5, 0xF5, 0xF5)
ASG_DARK     = RGBColor(0x1A, 0x1A, 0x1A)

LOGO_PATH = "/Users/Roger/Dropbox/code/authority-systems-group/brand-standards/assets/asgFull.png"
OUTPUT_PATH = "/Users/Roger/Dropbox/code/authority-systems-group/outputs/asg_coldoutreachengine_internal-assessment_20260308.docx"

DOC_TITLE = "ColdOutreachEngine.com — Internal Strategic Assessment"
DOC_DATE  = "March 8, 2026"

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_width(table, width_inches):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def add_bottom_border(paragraph, color='25aae1', sz='6'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_top_border(paragraph, color='25aae1', sz='4'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), sz)
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)

def add_left_border(paragraph, color='25aae1', sz='12'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_run(paragraph, text, bold=False, italic=False, color=None, size=11):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p, '25aae1', '6')
    add_run(p, text, bold=True, color=ASG_BLUE, size=18)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, color=ASG_CHARCOAL, size=14)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, bold=True, color=ASG_BLUE, size=12)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, color=ASG_CHARCOAL, size=11)
    return p

def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    add_left_border(p, '25aae1', '12')
    add_run(p, text, bold=True, color=ASG_BLUE, size=11)
    return p

def add_signoff(doc, name, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    add_run(p, f'— {name}, {title}  |  Authority Systems Group™', italic=True, color=ASG_CHARCOAL, size=10)
    return p

def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, text, color=ASG_CHARCOAL, size=11)
    return p

# ── Header ───────────────────────────────────────────────────────────────────
def add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Inches(1.4))
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)
    title_run = para.add_run('\t' + DOC_TITLE)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = ASG_CHARCOAL
    add_bottom_border(para, '25aae1', '6')

# ── Footer ───────────────────────────────────────────────────────────────────
def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    add_top_border(para, '25aae1', '4')
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)
    left_run = para.add_run('Authority Systems Group™  |  AuthoritySystemsGroup.com  |  Internal Use Only')
    left_run.font.name = 'Calibri'
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = ASG_CHARCOAL
    page_run = para.add_run('\tPage ')
    page_run.font.name = 'Calibri'
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = ASG_CHARCOAL
    # fldChar begin
    run_begin = para.add_run()
    run_begin.font.name = 'Calibri'
    run_begin.font.size = Pt(9)
    run_begin.font.color.rgb = ASG_CHARCOAL
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_begin._r.append(fldChar1)
    # instrText
    run_instr = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' PAGE '
    run_instr._r.append(instrText)
    # fldChar end
    run_end = para.add_run()
    run_end.font.name = 'Calibri'
    run_end.font.size = Pt(9)
    run_end.font.color.rgb = ASG_CHARCOAL
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fldChar2)

# ── Cover Page ───────────────────────────────────────────────────────────────
def add_cover(doc):
    # Dark header block — paragraph background (no table)
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(20)
    logo_p.paragraph_format.space_after = Pt(0)
    set_para_bg(logo_p, '1a1a1a')
    if os.path.exists(LOGO_PATH):
        logo_run = logo_p.add_run()
        logo_run.add_picture(LOGO_PATH, width=Inches(2.2))

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(20)
    set_para_bg(tag_p, '1a1a1a')
    tag_run = tag_p.add_run('Positioning You As The Authority')
    tag_run.font.name = 'Calibri'
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = ASG_BLUE
    tag_run.font.italic = True

    # Blue accent bar
    accent_p = doc.add_paragraph()
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    set_para_bg(accent_p, '25aae1')
    accent_p.paragraph_format.space_before = Pt(3)
    accent_p.paragraph_format.space_after = Pt(3)

    add_spacer(doc, 24)

    # Document title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    t_run = title_p.add_run('ColdOutreachEngine.com')
    t_run.font.name = 'Calibri'
    t_run.font.size = Pt(28)
    t_run.font.bold = True
    t_run.font.color.rgb = ASG_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run('Internal Strategic Assessment')
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(16)
    s_run.font.color.rgb = ASG_CHARCOAL

    add_spacer(doc, 12)

    # Divider
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(16)
    div_p.paragraph_format.space_after = Pt(16)
    add_bottom_border(div_p, '58585a', '8')

    # Attribution
    attr_p = doc.add_paragraph()
    attr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a_run = attr_p.add_run('Directed by Roger Bauer  |  Authority Systems Group™')
    a_run.font.name = 'Calibri'
    a_run.font.size = Pt(11)
    a_run.font.italic = True
    a_run.font.color.rgb = ASG_CHARCOAL

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = date_p.add_run(DOC_DATE)
    d_run.font.name = 'Calibri'
    d_run.font.size = Pt(10)
    d_run.font.color.rgb = ASG_CHARCOAL

    add_spacer(doc, 12)

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = conf_p.add_run('INTERNAL USE ONLY — NOT FOR CLIENT DISTRIBUTION')
    c_run.font.name = 'Calibri'
    c_run.font.size = Pt(9)
    c_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

# ── Main Document ─────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.875)
    section.right_margin  = Inches(0.875)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    add_header(doc)
    add_footer(doc)
    add_cover(doc)

    # ── SECTION: BACKGROUND ──────────────────────────────────────────────────
    add_h1(doc, 'Background & Context')
    add_body(doc,
        'ColdOutreachEngine.com is an internal ASG business concept under evaluation by Roger Bauer. '
        'The following assessment was produced by the ASG leadership team following a review of the domain, '
        'the service model, and the existing one-page website built by Manus.')
    add_spacer(doc, 8)
    add_h3(doc, 'The Service Model')
    add_body(doc,
        'ColdOutreachEngine.com is a done-for-you cold email infrastructure and managed outreach service '
        'for B2B consultants, agencies, and service businesses. The model operates as follows:')
    add_bullet(doc, 'ASG builds dedicated sending infrastructure for each client (15+ domains, 30+ inboxes, private SMTP server)')
    add_bullet(doc, 'ASG manages ongoing campaign operations — outreach, response handling, meeting booking — on the client\'s behalf')
    add_bullet(doc, 'Approximately 80% of the process is automated through n8n or Make; human oversight handles the remaining 20%')
    add_bullet(doc, 'AI voice agents and outbound calling represent a Phase 2 automation layer — not in scope for the initial launch')
    add_bullet(doc, 'Affiliate revenue from platform tools (Instantly, Smartlead, etc.) provides a secondary passive income stream on top of client retainers')
    add_spacer(doc, 8)
    add_h3(doc, 'The Website')
    add_body(doc,
        'A one-page site was produced at predeploy-9b557dfd-coldoutreach-vn9kqlg8-jqsdyamlhrxq2ts8nbapu4.manus.space. '
        'The site positions Cold Outreach Engine as a premium done-for-you client acquisition system with a '
        '"5 qualified sales meetings in 90 days" pipeline guarantee. Pricing is not disclosed on the page; '
        'all traffic is routed to a free strategy call CTA.')
    add_spacer(doc, 8)
    add_callout(doc,
        'STATUS: Backburner — Project is on hold pending development of case studies from LGD and DMA engagements. '
        'No active client acquisition for this service until proof-of-concept is established.')

    doc.add_page_break()

    # ── SECTION: LEADERSHIP ASSESSMENTS ─────────────────────────────────────
    add_h1(doc, 'Leadership Team Assessment')

    # Daniel Frost
    add_h2(doc, 'Daniel Frost — Chief Strategy Officer')
    add_h3(doc, 'Strategic Positioning')
    add_body(doc,
        'This is a strong complementary play, not a standalone business. The strategic value is not in the '
        'service itself — it is in what cold email funds and feeds. ASG is creating a client acquisition engine, '
        'then offering it as a productized service. That is a defensible position.')
    add_body(doc,
        'The positioning risk: "cold email infrastructure" sounds commodity. Every freelancer on Upwork does '
        'setup. The differentiation must be that ASG clients receive a revenue-producing system, not a configured '
        'inbox. The current website copy is closer to the commodity frame.')
    add_body(doc,
        'The real moat is the affiliate stack and the proprietary sending infrastructure. If ASG goes deep on '
        '2-3 tool relationships rather than remaining platform-agnostic, this creates a technical moat most '
        'competitors cannot replicate. The infrastructure story — private SMTP server, dedicated domain pools, '
        'warmup protocols — is the differentiator. Lean into it.')
    add_callout(doc,
        'Strategic priority: Rename "cold email infrastructure" in all positioning. The offer is a revenue-producing '
        'client acquisition system. The infrastructure is proof, not the product.')
    add_signoff(doc, 'Daniel Frost', 'CSO')
    add_spacer(doc, 10)

    # Vivienne Carr
    add_h2(doc, 'Vivienne Carr — Chief Marketing Officer')
    add_h3(doc, 'Copy & Messaging Review')
    add_body(doc,
        'The headline lands: "Double Your Qualified Sales Meetings in 90 Days Without Wasting a Single Minute '
        'on Prospecting" is outcome-first, time-bound, and specific. The problem section is well-written — the '
        '"trapped in the prospecting grind" frame resonates with the B2B consultant avatar.')
    add_body(doc,
        'Where the copy loses traction: by the middle of the page it defaults to generic agency language. '
        '"Battle-tested process." "B2B growth architects." "Value-driven email sequences designed to start '
        'conversations, not just send messages." These phrases appear on a thousand competitor pages. '
        'There is no proprietary language, no distinctive point of view, nothing that makes Cold Outreach '
        'Engine sound built by people with a specific philosophy.')
    add_body(doc,
        '"Trusted by B2B Growth Leaders" with no testimonials, names, or logos underneath is a credibility '
        'hole, not a credibility builder. The animated stat counters displaying "0+" in the live site are '
        'a technical bug that must be resolved before any marketing drives traffic to the page.')
    add_bullet(doc, 'Name the frameworks — give the process proprietary language')
    add_bullet(doc, 'Replace generic phrases with position-taking statements')
    add_bullet(doc, 'Fix the broken counter animations immediately')
    add_bullet(doc, 'Remove or replace social proof claims until real client results exist')
    add_signoff(doc, 'Vivienne Carr', 'CMO')
    add_spacer(doc, 10)

    # Tanya Blackwood
    add_h2(doc, 'Tanya Blackwood — Chief Revenue Officer')
    add_h3(doc, 'Revenue Model & Offer Structure')
    add_body(doc,
        'The model architecture is sound. The setup fee converts the initial sale. Affiliate commissions '
        'build a passive income layer that compounds as the client base grows. The value stack on the pricing '
        'page is well-constructed — breaking out component values totaling $11,000/month is a classic '
        'anchoring move, and the pipeline guarantee is strong.')
    add_body(doc,
        'The missing piece is a retention mechanism. Currently, once setup is complete, client continuity '
        'depends on the retainer relationship holding. A light-touch managed service tier — monthly '
        'deliverability monitoring, domain rotation, list hygiene consulting — creates a reason to remain '
        'engaged and an insurance policy on the affiliate trail.')
    add_body(doc,
        'The CTA is weak. "Book Your Free Strategy Call" is generic. Given the scarcity angle referenced in '
        'the FAQ ("limited new clients per month"), the CTA should reinforce that constraint: '
        '"Claim Your Spot — We Only Take 8 New Clients Per Month" converts urgency into action.')
    add_callout(doc,
        'Fast Cash Opportunity: If Roger has 5 existing ASG clients not currently running cold outreach, '
        'that is a $5,000–$15,000 setup week with minimal overhead and zero new marketing required.')
    add_signoff(doc, 'Tanya Blackwood', 'CRO')
    add_spacer(doc, 10)

    # Iris Nolan
    add_h2(doc, 'Iris Nolan — Chief Technology Officer')
    add_h3(doc, 'Automation Architecture')
    add_body(doc,
        '80% automation is achievable and realistic for this service model. The split between automated '
        'and human-handled tasks is well-defined once the client intake SOP is standardized.')
    add_h3(doc, 'Automatable via n8n or Make:')
    add_bullet(doc, 'Infrastructure provisioning triggers (domain purchase workflow, DNS config, mailbox setup notifications)')
    add_bullet(doc, 'Contact list processing and verification')
    add_bullet(doc, 'Campaign launch sequencing and scheduling')
    add_bullet(doc, 'Response classification: interested / not interested / out of office / unsubscribe')
    add_bullet(doc, 'Meeting booking handoff once a positive reply is detected')
    add_bullet(doc, 'Monthly reporting compilation and client Slack notifications')
    add_h3(doc, 'Requires human review (the 20%):')
    add_bullet(doc, 'Initial ICP definition and strategy per client')
    add_bullet(doc, 'Email sequence writing (AI drafts, human approves)')
    add_bullet(doc, 'Positive replies that do not fit the standard booking flow')
    add_bullet(doc, 'Edge cases: angry responses, legal threats, domain blacklisting incidents')
    add_body(doc,
        'The AI voice agent and outbound calling layer is correctly parked for Phase 2. Get the email '
        'automation running cleanly first — that is a complete build on its own. n8n is the preferred '
        'platform for self-hosted control; Make is faster if the goal is getting something live within a week.')
    add_body(doc,
        'Critical dependency: a client portal or structured Airtable tracker is required to maintain '
        'affiliate attribution per client. That audit trail is the revenue record.')
    add_callout(doc,
        'Architecture note: Standardize the client intake form tightly before building automation. '
        'The quality of automation output is directly proportional to the quality of intake data.')
    add_signoff(doc, 'Iris Nolan', 'CTO')
    add_spacer(doc, 10)

    # James Whitfield
    add_h2(doc, 'James Whitfield III — Board, Revenue Architecture')
    add_h3(doc, 'Pricing Strategy & Unit Economics')
    add_body(doc,
        'With 80% automation, cost-to-serve drops significantly. The unit economics at standard pricing '
        'are attractive:')
    add_bullet(doc, 'Setup fee: $1,500–$2,500 (one-time, covers infrastructure build and initial sequence writing)')
    add_bullet(doc, 'Monthly retainer: $1,500–$2,500/month (sending infrastructure, list procurement, automated management, human oversight)')
    add_bullet(doc, 'Hard costs per client: ~$300–$500/month (tools, SMTP, list data, contractor time for the 20%)')
    add_bullet(doc, 'Gross margin per active client: 70–80%')
    add_body(doc,
        'At 20 active clients: $30,000–$50,000/month in retainer revenue with $24,000–$40,000 gross margin. '
        'At 50 clients on a light managed plan, this becomes a mid-five-figure annual passive revenue stream '
        'requiring almost no new selling.')
    add_body(doc,
        'Pricing guidance: Do not undercharge the setup fee to drive volume. The affiliate tail is only '
        'valuable if the client is well-configured and succeeds — which requires margin in the setup to '
        'execute correctly. Below $1,500 and corners get cut or money gets lost on labor.')
    add_body(doc,
        'The affiliate passive income stream is now secondary. The retainer is the primary revenue engine. '
        'Build pricing around the retainer and treat affiliate commissions as found money.')
    add_signoff(doc, 'James Whitfield III', 'Board — Revenue Architecture')
    add_spacer(doc, 10)

    # Sofia Vega
    add_h2(doc, 'Sofia Vega — Board, Operational Feasibility')
    add_h3(doc, 'Operational Readiness')
    add_body(doc,
        'The model is feasible at scale. The primary concern is the transition period: the first 3–5 clients '
        'before automation is fully operational will be onboarded manually or semi-manually. Roger is the '
        'system during that window. That is not a disqualifier — it is the standard path for this type of '
        'build — but it should be planned for explicitly.')
    add_body(doc,
        'Recommended sequencing:')
    add_bullet(doc, 'Sell 2–3 clients manually — use the process to document every step in granular detail')
    add_bullet(doc, 'Iris builds the automation from that SOP — not from a theoretical process')
    add_bullet(doc, 'Client 4+ onboards through the automated system')
    add_bullet(doc, 'Do not sell beyond manual handling capacity until automation is live and proven')
    add_body(doc,
        'Three questions must be answered in writing before the first client is signed: '
        '(1) What does the client receive? (2) What is the delivery timeline? '
        '(3) Where does ASG\'s scope end — setup only, or ongoing management? '
        'Answered clearly in the contract, these prevent the first five clients from becoming support nightmares.')
    add_callout(doc,
        'The website copy is correct for the managed model. "Done-for-you" is accurate whether the process '
        'is automated or manual. The client does not need to know the delivery mechanism — only the outcome.')
    add_signoff(doc, 'Sofia Vega', 'Board — Operational Feasibility')

    doc.add_page_break()

    # ── SECTION: CONSENSUS ───────────────────────────────────────────────────
    add_h1(doc, 'Leadership Consensus')
    add_body(doc,
        'The Cold Outreach Engine concept is strategically sound. The managed + heavily automated model '
        'is the optimal configuration: it commands premium pricing while maintaining commodity-level '
        'labor cost, and the pipeline guarantee becomes a sales tool rather than a genuine liability '
        'once the automation is refined.')
    add_h3(doc, 'What Works')
    add_bullet(doc, 'The business model architecture — setup fee + monthly retainer + affiliate passive income')
    add_bullet(doc, 'The pipeline guarantee as a conversion mechanism')
    add_bullet(doc, 'The technical differentiation of proprietary sending infrastructure')
    add_bullet(doc, '80% automation target reducing cost-to-serve to 70–80% gross margin')
    add_bullet(doc, 'The one-page website structure — bones are right, copy needs sharpening')
    add_h3(doc, 'What Needs Work Before Launch')
    add_bullet(doc, 'Fix broken animated stat counters on the website immediately')
    add_bullet(doc, 'Rewrite copy to eliminate generic agency language and introduce proprietary framework names')
    add_bullet(doc, 'Replace or remove unsubstantiated social proof claims')
    add_bullet(doc, 'Strengthen the CTA with scarcity-anchored language')
    add_bullet(doc, 'Define "qualified meeting" in writing before the first client contract is signed')
    add_bullet(doc, 'Document the full manual onboarding SOP across 2–3 initial clients before building automation')
    add_bullet(doc, 'Build Airtable or equivalent tracker for affiliate attribution per client')
    add_h3(doc, 'Project Status')
    add_callout(doc,
        'BACKBURNER — No active sales or marketing for ColdOutreachEngine.com until case studies are '
        'developed from LGD and DMA engagements. Case study availability is the trigger for re-activating '
        'this project. At that point, Tanya\'s fast cash sprint angle (closing existing ASG clients on '
        'setup service first) is the recommended Day 1 move.')

    add_spacer(doc, 16)

    # Footer sign-off
    final_p = doc.add_paragraph()
    final_p.paragraph_format.space_before = Pt(12)
    add_top_border(final_p, '25aae1', '4')
    fr = final_p.add_run('Authority Systems Group™ — Confidential. Internal Use Only.')
    fr.font.name = 'Calibri'
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = ASG_CHARCOAL

    doc.save(OUTPUT_PATH)
    print(f'✓ Saved: {OUTPUT_PATH}')

if __name__ == '__main__':
    build()
