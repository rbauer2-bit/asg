#!/usr/bin/env python3
"""
Authority Systems Group™ — Niche Intelligence Brief
Tree Service Business: Avatar Persona + Pain Points + Insider Intelligence
ASG Internal Asset — Niche Intelligence Library
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK_BG   = '1a1a1a'
ASG_BLUE_HEX  = '25aae1'
ASG_DARK_HEX  = '1A2533'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_TINT      = 'EAF6FC'
SLATE_HEX     = '2C3E50'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'E67E22'
RED_HEX       = '7D2400'
GOLD_HEX      = 'C9A02C'
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_tree-service_avatar-persona_20260308.docx')


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_pb(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)


def set_margins(doc, top=0.75, bottom=0.75, left=0.9, right=0.9):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


# ── Typography Helpers ─────────────────────────────────────────────────────────

def heading(doc, text, level=1, color=None, bg=None, size=None, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    if bg:
        set_para_bg(p, bg)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.font.name = 'Calibri'
    r.font.color.rgb = color if color else (WHITE if bg else ASG_CHARCOAL)
    if size:
        r.font.size = Pt(size)
    elif level == 1:
        r.font.size = Pt(20)
    elif level == 2:
        r.font.size = Pt(14)
    elif level == 3:
        r.font.size = Pt(11)
    return p


def body(doc, text, size=10, color=None, bold=False, italic=False, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def bullet(doc, text, size=10, indent=0.2, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(3)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def label_value(doc, label, value, label_color=None, value_color=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    rl = p.add_run(label + '  ')
    rl.bold = True
    rl.font.name = 'Calibri'
    rl.font.size = Pt(size)
    rl.font.color.rgb = label_color if label_color else ASG_CHARCOAL
    rv = p.add_run(value)
    rv.font.name = 'Calibri'
    rv.font.size = Pt(size)
    rv.font.color.rgb = value_color if value_color else BODY_COLOR
    return p


def divider(doc, hex_color=ASG_BLUE_HEX, height_pt=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    set_para_bg(p, hex_color)
    p.paragraph_format.space_before = Pt(height_pt)
    p.paragraph_format.space_after  = Pt(height_pt)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Dark full-width header bar
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(cover, ASG_DARK_HEX)
    cover.paragraph_format.space_before = Pt(0)
    cover.paragraph_format.space_after  = Pt(0)

    if os.path.exists(LOGO_PATH):
        run = cover.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))

    sp(doc, 4)

    # Blue accent bar
    accent = doc.add_paragraph()
    set_para_bg(accent, ASG_BLUE_HEX)
    accent.paragraph_format.space_before = Pt(3)
    accent.paragraph_format.space_after  = Pt(3)

    sp(doc, 10)

    # Title block
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run('NICHE INTELLIGENCE BRIEF')
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = ASG_BLUE
    r1.font.all_caps = True

    sp(doc, 4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Tree Service Business')
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(32)
    r2.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

    sp(doc, 6)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Avatar Persona  ·  Pain Points  ·  Insider Intelligence')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(13)
    r3.font.color.rgb = ASG_CHARCOAL

    sp(doc, 30)

    # Credential block
    cred_tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(cred_tbl)
    cred_tbl.columns[0].width = Inches(2.5)
    cred_tbl.columns[1].width = Inches(2.5)
    cred_tbl.columns[2].width = Inches(2.5)

    labels = ['PREPARED BY', 'DOCUMENT TYPE', 'DATE']
    values = ['Authority Systems Group™', 'Internal Niche Intelligence', 'March 8, 2026']

    for i, (lbl, val) in enumerate(zip(labels, values)):
        cell = cred_tbl.rows[0].cells[i]
        set_cell_bg(cell, ASG_TINT)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = cp.add_run(lbl + '\n')
        rl.bold = True
        rl.font.name = 'Calibri'
        rl.font.size = Pt(8)
        rl.font.color.rgb = ASG_BLUE
        rv = cp.add_run(val)
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9)
        rv.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

    sp(doc, 20)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run('Authority Systems Group™  —  Internal Asset. Do not distribute to clients.')
    footer_r.font.name = 'Calibri'
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = FOOTER_COLOR

    add_pb(doc)


# ── Section 1: The Avatar ──────────────────────────────────────────────────────

def build_avatar(doc):
    heading(doc, 'SECTION 1', level=3, color=ASG_BLUE, size=9)
    heading(doc, '"Cutting Edge Charlie"', level=1, color=RGBColor(0x1A, 0x25, 0x33), size=24, bold=True)
    heading(doc, 'The Tree Service Owner-Operator — ASG Prospect Avatar', level=2,
            color=ASG_CHARCOAL, size=13, bold=False)
    divider(doc)
    sp(doc, 6)

    body(doc,
         'Charlie didn\'t go to business school. He went up trees. Started as a climber or groundsman for '
         'someone else, got good enough to know he was making his boss rich, bought a saw and a pickup, and '
         'built from there. Fifteen years later he\'s running a real operation — 4 to 12 people, $600K to '
         '$2M in revenue, $200,000 worth of equipment with payments to prove it — and he\'s better at his '
         'trade than almost anyone in his market.',
         size=10.5)

    body(doc,
         'The problem is nobody knows it. Online, he\'s invisible or indistinguishable from the guy who '
         'started last year. He\'s been pitched by half a dozen marketing agencies who didn\'t know what '
         'a chipper drum is, let alone what PHC stands for. He\'s skeptical, he\'s smart, and he\'ll '
         'respect the first person who demonstrates they actually understand his world before trying to '
         'sell him something.',
         size=10.5)

    sp(doc, 4)

    # Profile table
    heading(doc, 'Business Profile', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    profile_data = [
        ('Revenue Range', '$350,000 – $2,500,000/year'),
        ('Employees', '2–15 (including owner)'),
        ('Years in Business', '5–18 years'),
        ('Equipment Investment', '$150,000 – $600,000'),
        ('Avg Job Value', '$800 – $8,500'),
        ('CRM / Scheduling', 'None | Jobber | ServiceTitan | Housecall Pro'),
        ('Billing Model', 'Per-job flat fee; T&M for emergency work'),
        ('Emergency Rate Premium', '1.5x – 2.5x standard rate (usually not charged)'),
    ]

    tbl = doc.add_table(rows=len(profile_data), cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(2.4)
    tbl.columns[1].width = Inches(5.1)

    for i, (lbl, val) in enumerate(profile_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(vc, bg)
        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent = Inches(0.1)
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = ASG_CHARCOAL
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.name = 'Calibri'
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = BODY_COLOR

    sp(doc, 8)

    # Revenue Mix
    heading(doc, 'Typical Revenue Mix', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    rev_data = [
        ('Tree Removal (hazardous + standard)',   '50–60%'),
        ('Trimming and Pruning',                   '20–25%'),
        ('Emergency / Storm Work',                 '10–15%'),
        ('Stump Grinding',                         '5–8%'),
        ('Plant Health Care (PHC)',                '2–8%'),
        ('Firewood / Secondary Byproduct',         '0–5%'),
    ]

    rev_tbl = doc.add_table(rows=len(rev_data), cols=2)
    remove_table_borders(rev_tbl)
    rev_tbl.columns[0].width = Inches(5.0)
    rev_tbl.columns[1].width = Inches(2.5)

    for i, (svc, pct) in enumerate(rev_data):
        bg = ASG_LIGHT_BG if i % 2 == 0 else 'FFFFFF'
        sc = rev_tbl.rows[i].cells[0]
        pc = rev_tbl.rows[i].cells[1]
        set_cell_bg(sc, bg)
        set_cell_bg(pc, bg)
        sp2 = sc.paragraphs[0]
        sp2.paragraph_format.left_indent = Inches(0.1)
        sr = sp2.add_run(svc)
        sr.font.name = 'Calibri'
        sr.font.size = Pt(9.5)
        sr.font.color.rgb = BODY_COLOR
        pp = pc.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = pp.add_run(pct)
        pr.bold = True
        pr.font.name = 'Calibri'
        pr.font.size = Pt(9.5)
        pr.font.color.rgb = ASG_BLUE

    sp(doc, 8)

    # Credentials
    heading(doc, 'Credentials That Matter', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    creds = [
        ('ISA Certified Arborist',
         'Requires documented field experience + proctored exam + ongoing CEUs. Most operators have it. Almost none lead with it.'),
        ('ISA Board Certified Master Arborist',
         'Fewer than 400 in the continental US. If Charlie has it, it\'s the single most underutilized marketing asset in his business.'),
        ('TCIA Accreditation',
         'Company-level credential. Requires current membership. Signals professionalism to commercial clients.'),
        ('General Liability $1M–$2M+',
         'Critical differentiator vs. uninsured lowballers. Most homeowners don\'t know to ask for the certificate.'),
        ('Workers\' Compensation',
         'Required in most states. Uninsured operators are common. A legitimate COI eliminates this liability for the homeowner.'),
    ]

    for cred_name, cred_desc in creds:
        cred_tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(cred_tbl)
        cred_tbl.columns[0].width = Inches(7.5)
        cell = cred_tbl.rows[0].cells[0]
        set_cell_bg(cell, ASG_TINT)
        cp = cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.12)
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(4)
        rn = cp.add_run(cred_name + '  ')
        rn.bold = True
        rn.font.name = 'Calibri'
        rn.font.size = Pt(9.5)
        rn.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)
        rd = cp.add_run(cred_desc)
        rd.font.name = 'Calibri'
        rd.font.size = Pt(9.5)
        rd.font.color.rgb = BODY_COLOR
        sp(doc, 3)

    sp(doc, 6)

    # Belief barriers
    heading(doc, 'Why He Hesitates to Hire ASG', level=3, color=RGBColor(0x1A, 0x25, 0x33), size=11)
    sp(doc, 2)

    barriers = [
        ('Primary Enemy Belief',
         '"Marketing agencies don\'t understand how my business actually works — they\'ll waste my money on generic stuff that gets me more of the wrong clients."',
         RED_HEX),
        ('Secondary Enemy Belief',
         '"I built this on referrals. If I just keep doing good work, the business will grow."',
         AMBER_HEX),
        ('Normalized Pain',
         '"Slow seasons are just part of the tree business. Everyone goes through them."',
         SLATE_HEX),
        ('Urgency Blocker',
         '"The money I\'d spend on marketing is money I could put into equipment that actually earns."',
         SLATE_HEX),
        ('Trust Trigger',
         'Demonstrate specific knowledge of his industry before the pitch. PHC, ISA cert, referral network economics. Drop any one of these and he leans in.',
         GREEN_HEX),
    ]

    for btype, btext, bhex in barriers:
        btbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(btbl)
        btbl.columns[0].width = Inches(0.12)
        btbl.columns[1].width = Inches(7.38)
        accent_cell = btbl.rows[0].cells[0]
        set_cell_bg(accent_cell, bhex)
        content_cell = btbl.rows[0].cells[1]
        set_cell_bg(content_cell, ASG_LIGHT_BG)
        cp = content_cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.12)
        cp.paragraph_format.space_before = Pt(5)
        cp.paragraph_format.space_after  = Pt(5)
        rl = cp.add_run(btype + '  ')
        rl.bold = True
        rl.font.name = 'Calibri'
        rl.font.size = Pt(9)
        rl.font.color.rgb = ASG_CHARCOAL
        rv = cp.add_run(btext)
        rv.italic = True
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = BODY_COLOR
        sp(doc, 3)

    add_pb(doc)


# ── Section 2: 10 Pain Points ──────────────────────────────────────────────────

PAIN_POINTS = [
    {
        'number': '01',
        'title': 'The Estimate Graveyard',
        'body': (
            'He estimates 40–60% of his leads. He lands maybe 30–40% of those. The rest? Gone. '
            'Price shoppers who gave the job to the guy with no insurance and a borrowed chipper. '
            'He has zero system for following up. The lead just dies.'
        ),
        'agency_insight': (
            'Tree estimates are time-intensive site visits — not a button click. He\'s driving to the property, '
            'walking the lot, assessing risk, calculating crane or climb, checking utility lines, calculating '
            'disposal, and quoting a legal scope. He loses $50–$150 of time on every estimate that doesn\'t close. '
            'A structured follow-up sequence for unconverted estimates is almost universally missing in this niche.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'number': '02',
        'title': 'The Insurance Cliff',
        'body': (
            'General liability and workers\' comp together run $30,000–$80,000+ per year. One bad year — '
            'one claim — and his premium jumps 40%. He\'s essentially self-funding a risk pool on top of '
            'running a business. And every lowballer who steals his job is carrying a $400 policy that '
            'will deny the claim the moment a branch hits a house.'
        ),
        'agency_insight': (
            'He can\'t just "raise prices to cover overhead." His market has a ceiling because uninsured '
            'operators exist and homeowners can\'t tell the difference until it\'s too late. The marketing '
            'opportunity: make the insurance differential visible, tangible, and emotionally relevant to the '
            'homeowner before the lowballer gets the job.'
        ),
        'color': RED_HEX,
    },
    {
        'number': '03',
        'title': 'The Seasonal Spike-and-Crater',
        'body': (
            'Spring and fall are slammed. Summer is okay. December through February? In half the country, '
            'he\'s staring at a calendar with 40% of his normal bookings and 100% of his fixed costs — '
            'equipment payments, insurance, a couple of guys he wants to keep.'
        ),
        'agency_insight': (
            'Storm season is both his best and worst time. A derecho can flood him with calls he can\'t take '
            'and won\'t get paid on fast enough to cover the surge in fuel, overtime, and equipment wear. '
            'The winter dormancy pitch — educating clients that winter removal is often BETTER for the '
            'remaining trees — is almost never used. The operators who use it fill their slow season at full rate.'
        ),
        'color': SLATE_HEX,
    },
    {
        'number': '04',
        'title': 'The Skilled Labor Drought',
        'body': (
            'A good climber who can read a tree, set rigging, and make a cut under pressure is rare, '
            'underpaid everywhere else, and constantly being recruited by utilities and larger contractors. '
            'Finding one is luck. Keeping one means paying him what he\'s worth, which eats the margin. '
            'Greenhorns run equipment into the ground and have to be supervised constantly.'
        ),
        'agency_insight': (
            'This is not a "post a job on Indeed" problem. This trade is learned in the field over years. '
            'There is no school you can send someone to and get back a good climber in six months. The labor '
            'shortage is structural, not cyclical. This affects his growth ceiling more than any marketing gap.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'number': '05',
        'title': 'The Plant Health Care Blind Spot',
        'body': (
            'He knows how to remove trees. He probably knows how to treat them — deep root fertilization, '
            'disease management, cable bracing, air spading. But most of his revenue is removal. PHC is a '
            'recurring revenue model with 60–80% gross margins, and most tree service guys have never built '
            'it out or gave it away as a loss leader without realizing what they had.'
        ),
        'agency_insight': (
            'PHC is the closest thing this industry has to an annuity. A property with 20 mature trees is '
            'worth $800–$2,000 per year in treatments — annually, indefinitely. Nobody in marketing talks '
            'about this because they don\'t know it exists. Mentioning it in your first conversation with '
            'Charlie is an instant credibility signal.'
        ),
        'color': GREEN_HEX,
    },
    {
        'number': '06',
        'title': 'The Referral Network He Never Built Formally',
        'body': (
            'Landscapers, roofers, real estate agents, home inspectors, HOA managers, general contractors — '
            'these people see tree problems constantly and they send work somewhere. He probably has loose '
            'relationships with a few of them. He has no formal program, no systematic outreach, and no one '
            'source he could point to and say "that guy sends me $40,000 a year."'
        ),
        'agency_insight': (
            'A single HOA property manager can control 200–400 properties. One relationship = a whole revenue '
            'channel. Most tree operators have never sent a professional letter or had a real conversation '
            'with a property management company. This is the highest-ROI fast cash lever in the business.'
        ),
        'color': SLATE_HEX,
    },
    {
        'number': '07',
        'title': 'The Website That Was Good Enough Five Years Ago',
        'body': (
            'He\'s got a website. It has stock photos of guys in hard hats, mentions "licensed and insured," '
            'and hasn\'t been touched since 2019. His Google Business Profile has some reviews — or he\'s '
            'never asked for a single one. A national franchise just moved in and is spending $15K/month on '
            'ads and taking his calls.'
        ),
        'agency_insight': (
            'His actual authority — 15 years in the market, his insurance cert is impeccable, every tree '
            'company in three counties knows his name — is completely invisible online. He is '
            'indistinguishable digitally from the guy who started last year. "ISA certified arborist [city]" '
            'is almost uncontested in most local markets. The first operator to own that term wins.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'number': '08',
        'title': 'The Firewood and Chip Disposal Albatross',
        'body': (
            'Every removal job creates a debris problem. Wood chips have to go somewhere — landfill tipping '
            'fees, give-away to municipalities, or he hauls it. Logs either get bucked for firewood (labor + '
            'storage + sales time) or they\'re a cost center. Most guys have no monetization strategy for '
            'their own byproduct.'
        ),
        'agency_insight': (
            'A tree service doing $800K/year in removals generates hundreds of cords of firewood that either '
            'make money or cost money. Some operators have built $60,000–$150,000 secondary revenue streams '
            'in firewood alone. Nobody in generalist marketing talks about the wood because they\'ve never '
            'thought about what happens after the tree hits the ground.'
        ),
        'color': AMBER_HEX,
    },
    {
        'number': '09',
        'title': 'The Estimating Multiplier He\'s Leaving on the Table',
        'body': (
            'He prices by the job — look at the tree, guess the time, add equipment, add risk, quote a '
            'number. He quotes the job the homeowner asked about rather than walking the whole property. '
            'Average job ticket could be 30–50% higher with a structured walkthrough and upsell sequence.'
        ),
        'agency_insight': (
            'The estimate visit IS the sales opportunity. The problem is he treats it like an obligation to '
            'quote, not a consultation to guide. This is a trained behavior, not a personality flaw — and '
            'it\'s fixable with a one-page checklist. This single change increases revenue without a single '
            'additional lead.'
        ),
        'color': SLATE_HEX,
    },
    {
        'number': '10',
        'title': 'The "Just a Tree Guy" Identity Tax',
        'body': (
            'He undersells himself. He says "we just do tree work" when he\'s a trained arborist with an '
            'ISA credential, $2M in liability coverage, a crew with 40 combined years of field experience, '
            'and a safety record that would get him approved to work for a utility company tomorrow. He '
            'doesn\'t talk about what his certification means. He doesn\'t mention what happens when an '
            'uninsured crew drops a branch on a house.'
        ),
        'agency_insight': (
            'The ISA Certified Arborist credential is genuinely meaningful — documented field experience, '
            'a proctored exam, ongoing CEU requirements. Most homeowners have never heard of it. The first '
            'tree service company that explains it well in their marketing wins. Nobody has. This is the '
            'fastest single authority gap to close in the entire niche.'
        ),
        'color': ASG_DARK_HEX,
    },
]


def build_pain_points(doc):
    heading(doc, 'SECTION 2', level=3, color=ASG_BLUE, size=9)
    heading(doc, '10 Pain Points for Growing the Business', level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'What keeps Charlie up at night — and what every agency misses', level=2,
            color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 8)

    for pp in PAIN_POINTS:
        # Number + title bar
        num_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(num_tbl)
        num_tbl.columns[0].width = Inches(0.65)
        num_tbl.columns[1].width = Inches(6.85)

        num_cell = num_tbl.rows[0].cells[0]
        ttl_cell = num_tbl.rows[0].cells[1]
        set_cell_bg(num_cell, pp['color'])
        set_cell_bg(ttl_cell, pp['color'])

        np_ = num_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(6)
        np_.paragraph_format.space_after  = Pt(6)
        nr = np_.add_run(pp['number'])
        nr.bold = True
        nr.font.name = 'Calibri'
        nr.font.size = Pt(15)
        nr.font.color.rgb = WHITE

        tp = ttl_cell.paragraphs[0]
        tp.paragraph_format.left_indent  = Inches(0.1)
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(6)
        tr = tp.add_run(pp['title'])
        tr.bold = True
        tr.font.name = 'Calibri'
        tr.font.size = Pt(13)
        tr.font.color.rgb = WHITE

        # Body text
        body_tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(body_tbl)
        body_tbl.columns[0].width = Inches(7.5)
        bc = body_tbl.rows[0].cells[0]
        set_cell_bg(bc, 'FFFFFF')
        bp = bc.paragraphs[0]
        bp.paragraph_format.left_indent  = Inches(0.15)
        bp.paragraph_format.right_indent = Inches(0.15)
        bp.paragraph_format.space_before = Pt(7)
        bp.paragraph_format.space_after  = Pt(4)
        br = bp.add_run(pp['body'])
        br.font.name = 'Calibri'
        br.font.size = Pt(10)
        br.font.color.rgb = BODY_COLOR

        # Agency insight
        ai_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(ai_tbl)
        ai_tbl.columns[0].width = Inches(0.08)
        ai_tbl.columns[1].width = Inches(7.42)
        set_cell_bg(ai_tbl.rows[0].cells[0], ASG_BLUE_HEX)
        ai_cell = ai_tbl.rows[0].cells[1]
        set_cell_bg(ai_cell, ASG_TINT)
        aip = ai_cell.paragraphs[0]
        aip.paragraph_format.left_indent  = Inches(0.15)
        aip.paragraph_format.right_indent = Inches(0.15)
        aip.paragraph_format.space_before = Pt(6)
        aip.paragraph_format.space_after  = Pt(6)
        ai_label = aip.add_run('What the typical agency doesn\'t know:  ')
        ai_label.bold = True
        ai_label.font.name = 'Calibri'
        ai_label.font.size = Pt(9)
        ai_label.font.color.rgb = ASG_BLUE
        ai_text = aip.add_run(pp['agency_insight'])
        ai_text.font.name = 'Calibri'
        ai_text.font.size = Pt(9.5)
        ai_text.font.color.rgb = BODY_COLOR

        sp(doc, 10)

    add_pb(doc)


# ── Section 3: Insider Intelligence ───────────────────────────────────────────

INSIDER_FACTS = [
    {
        'topic': 'PHC — The Recurring Revenue Nobody Talks About',
        'insight': (
            'Plant Health Care (deep root fertilization, disease management, cable bracing, air spading) is '
            'the closest thing tree service has to an annuity. A property with 20 mature trees can generate '
            '$800–$2,000/year in treatments — indefinitely. Most operators have never built it out. Mention '
            '"PHC as a recurring revenue program" in your first conversation with Charlie and you\'ve already '
            'separated from every agency that has ever pitched him.'
        ),
        'color': GREEN_HEX,
    },
    {
        'topic': 'The ISA Credential — Most Valuable, Least Used',
        'insight': (
            'ISA Certified Arborist requires documented field experience, a proctored exam, and ongoing '
            'continuing education. ISA Board Certified Master Arborist is held by fewer than 400 people in '
            'the continental US. Almost no tree service operator leads with this in their marketing — it\'s '
            'always buried in fine print. This is the single fastest authority gap to close in the entire niche.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'topic': 'Storm Response Premium Rates — The Conversation Most Operators Won\'t Have',
        'insight': (
            'Emergency storm work commands 1.5–2.5x standard rates. Most operators charge standard rates '
            'because they feel uncomfortable quoting higher in someone\'s moment of crisis. The operators '
            'who define emergency pricing in advance — with a rationale that\'s built into their communications '
            '— close without resistance. This is a marketing and positioning problem, not a pricing problem.'
        ),
        'color': AMBER_HEX,
    },
    {
        'topic': 'The Neighbor Notification Liability Gap',
        'insight': (
            'Before a major removal on a property line, smart operators notify adjacent neighbors in writing. '
            'Most don\'t. When something goes sideways and the neighbor\'s property is damaged, the written '
            'notice is the difference between an insured claim and a lawsuit. An arborist who educates '
            'clients on this protocol positions himself as the professional advisor, not just the crew with a saw.'
        ),
        'color': RED_HEX,
    },
    {
        'topic': 'Crane vs. Climb Economics — The Counterintuitive Sale',
        'insight': (
            'A crane job that takes 3 hours often costs the same or less total job time as a difficult '
            'climbing job that takes 8 hours with significantly more risk. Crane rental ($800–$1,500/day) '
            'is often the MORE profitable choice. Homeowners think crane = expensive. The operator who can '
            'flip that script — explaining faster, safer, and often equivalent cost — closes more crane jobs '
            'and runs a better operation.'
        ),
        'color': SLATE_HEX,
    },
    {
        'topic': 'Winter Dormancy — The Pitch That Fills the Slow Season',
        'insight': (
            'Winter removal is often BETTER for remaining trees: lower disease transfer risk, visible '
            'structure, less site damage when ground is frozen or dry. Most operators don\'t make this '
            'argument and either discount to fill the calendar or simply go quiet. The operator who '
            'educates homeowners on winter removal timing in October and November content fills his slow '
            'season at full rate. Almost nobody does this.'
        ),
        'color': ASG_DARK_HEX,
    },
    {
        'topic': 'Heritage Tree Ordinances — The Professional\'s Differentiator',
        'insight': (
            'Many municipalities have heritage tree ordinances — remove a designated tree without a permit '
            'and the fine is $5,000–$50,000 with a replacement order. An arborist who mentions this during '
            'the estimate provides genuine professional protection and positions the entire firm as advisors, '
            'not just cutters. Most operators never mention it. The first one who does changes the conversation.'
        ),
        'color': SLATE_HEX,
    },
    {
        'topic': 'Stump Grinding — The Standalone Profit Center Almost Nobody Markets',
        'insight': (
            'Many operators include stump grinding as an add-on or give it away to close removals. Marketed '
            'as a standalone service — "that stump you\'ve been looking at for three years" — it\'s a '
            'high-margin, low-competition campaign with almost no customer resistance. The homeowner already '
            'wants it gone. They just need someone to ask.'
        ),
        'color': ASG_BLUE_HEX,
    },
    {
        'topic': 'The Estimate Visit IS the Sales Visit',
        'insight': (
            'Most operators treat the estimate visit as an obligation to quote what the homeowner called '
            'about. A structured walkthrough checklist — designed to identify every tree on the property, '
            'not just the one in question — increases average job ticket 30–50% with zero additional '
            'marketing spend. It is a behavior change, not a marketing campaign. And it is trainable.'
        ),
        'color': GREEN_HEX,
    },
    {
        'topic': 'The Referral Network Economics — One Relationship vs. One Thousand Ads',
        'insight': (
            'A single HOA property manager can control 200–400 properties. A single landscaper sees tree '
            'problems on every job. A real estate agent\'s buyers need pre-sale tree assessments. One '
            'formalized referral relationship with a property manager is worth $30,000–$80,000/year. '
            'Most tree operators have never sent a professional outreach letter to a property management '
            'company. The first one who does owns the channel.'
        ),
        'color': AMBER_HEX,
    },
    {
        'topic': 'Firewood as Secondary Revenue — The Byproduct Nobody Is Monetizing',
        'insight': (
            'A tree service doing $800K/year in removals generates hundreds of cords of firewood annually '
            'that either make money or cost money. Some operators have built $60,000–$150,000 secondary '
            'revenue streams selling seasoned firewood. Nobody in generalist marketing talks about the wood '
            'because they\'ve never thought about what happens after the tree hits the ground. Dropping this '
            'fact in conversation signals deep industry knowledge instantly.'
        ),
        'color': RED_HEX,
    },
    {
        'topic': 'Commercial Contracts and Bid Bonds — The Revenue Game Charlie May Not Know Exists',
        'insight': (
            'The tree service that can submit a bid bond and respond to commercial RFPs — municipalities, '
            'school districts, commercial property management companies — plays a completely different '
            'revenue game. Most small operators have never done it. A single municipal contract can be '
            'worth $50,000–$200,000/year. The operators who learn the process can 10x contract value per '
            'job. Almost no marketing agency knows to ask about this opportunity.'
        ),
        'color': ASG_DARK_HEX,
    },
]


def build_insider_intel(doc):
    heading(doc, 'SECTION 3', level=3, color=ASG_BLUE, size=9)
    heading(doc, 'Insider Intelligence', level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'What separates ASG from every generalist agency that has ever pitched a tree service company',
            level=2, color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 4)

    body(doc,
         'Drop any one of the following facts into a first conversation with Charlie and you immediately '
         'separate yourself from every agency that has ever pitched him. These are the things a consultant '
         'who has worked with tree service clients for years knows cold. Use them. They are your authority signal.',
         size=10.5, italic=True, color=ASG_CHARCOAL)

    sp(doc, 8)

    for i, fact in enumerate(INSIDER_FACTS):
        # Colored side-bar layout
        fact_tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(fact_tbl)
        fact_tbl.columns[0].width = Inches(0.18)
        fact_tbl.columns[1].width = Inches(7.32)

        set_cell_bg(fact_tbl.rows[0].cells[0], fact['color'])
        content_cell = fact_tbl.rows[0].cells[1]
        set_cell_bg(content_cell, ASG_LIGHT_BG if i % 2 == 0 else 'FFFFFF')

        cp = content_cell.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.15)
        cp.paragraph_format.right_indent = Inches(0.15)
        cp.paragraph_format.space_before = Pt(7)
        cp.paragraph_format.space_after  = Pt(8)

        rt = cp.add_run(fact['topic'] + '\n')
        rt.bold = True
        rt.font.name = 'Calibri'
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)

        ri = cp.add_run(fact['insight'])
        ri.font.name = 'Calibri'
        ri.font.size = Pt(9.5)
        ri.font.color.rgb = BODY_COLOR

        sp(doc, 5)

    add_pb(doc)


# ── Section 4: Charlie's Customer Avatars ─────────────────────────────────────

def build_customer_avatars(doc):
    heading(doc, 'SECTION 4', level=3, color=ASG_BLUE, size=9)
    heading(doc, 'Charlie\'s Customer Avatars', level=1,
            color=RGBColor(0x1A, 0x25, 0x33), size=22, bold=True)
    heading(doc, 'Who Charlie sells to — and what actually moves them to buy',
            level=2, color=ASG_CHARCOAL, size=12, bold=False)
    divider(doc)
    sp(doc, 8)

    avatars = [
        {
            'name': 'Liability-Worried Linda',
            'service': 'Hazardous Tree Removal',
            'color': RED_HEX,
            'profile': 'Female skew, 42–62. The household decision-maker on property issues. She\'s been looking at the dead oak for two seasons and her husband keeps saying it\'s fine.',
            'triggers': [
                'Tree noticed leaning or clearly dead after a storm',
                'Neighbor complaining about overhanging limbs',
                'Insurance renewal or home inspection flags the tree',
                'A tree fell in the neighborhood and she saw it on Nextdoor',
            ],
            'fears': [
                'Tree falls on house, fence, or car while she waited',
                'Hiring someone unqualified who damages the property',
                'Job being more expensive than quoted mid-project',
                'Being responsible for neighbor property damage',
            ],
            'hooks': [
                '"The 5 signs a tree needs to come down before it comes down on you"',
                '"What happens when a neighbor\'s tree falls on your house — and who\'s liable"',
                '"Why the cheapest tree quote could cost you the most"',
            ],
        },
        {
            'name': 'Storm Season Steve',
            'service': 'Emergency Removal',
            'color': AMBER_HEX,
            'profile': 'Gender-neutral, 35–60. A tree is on his house or blocking his driveway. He needs someone here today — not in four days. Price is almost no object.',
            'triggers': [
                'Tree down on house, car, fence, or across driveway',
                'Hanging limbs creating immediate hazard',
                'Power line contact situation',
                'Insurance adjuster coming and needs debris cleared first',
            ],
            'fears': [
                'Additional damage while waiting — tree continues to shift',
                'Getting scammed by a storm-chasing fly-by-night crew',
                'Insurance claim being complicated by who he chose',
                'Not being able to get anyone out quickly',
            ],
            'hooks': [
                '"Storm hit your property? Here\'s what to do in the first 24 hours"',
                '"How to document tree damage for your insurance claim"',
                '"Why emergency tree removal rates are different — and why that\'s fair"',
            ],
        },
        {
            'name': 'Property Manager Paula',
            'service': 'Commercial Contract / PHC Program',
            'color': SLATE_HEX,
            'profile': 'Slight female skew, 30–55. Manages 50–400 properties for an HOA or commercial owner. She needs a vendor who shows up, documents everything, and doesn\'t make her look bad to the board.',
            'triggers': [
                'Annual property maintenance budget approval season (Q4)',
                'HOA board complaint about tree condition or incident',
                'Previous contractor did poor work or became unreliable',
                'Insurance audit flagging tree liability across properties',
            ],
            'fears': [
                'Liability from a tree incident on a property she manages',
                'Contractor no-shows that put her in a bad position with the board',
                'Cost overruns outside approved budget',
                'A contractor who doesn\'t communicate professionally',
            ],
            'hooks': [
                '"What every property manager should know about tree liability"',
                '"How to evaluate a commercial tree service vendor before you sign"',
                '"Why annual PHC programs save HOAs money over emergency removals"',
            ],
        },
        {
            'name': 'Estate Owner Edward',
            'service': 'PHC Program + Premium Removal',
            'color': ASG_DARK_HEX,
            'profile': 'Male skew, 50–70. Owns a large residential property (1+ acre) with significant mature tree canopy. He bought the property for the trees. He wants an arborist who treats them like assets.',
            'triggers': [
                'Noticed disease or dieback in a tree he values',
                'Landscaper flagged a tree concern during their visit',
                'Tree on neighboring property fell — now aware of his own risk',
                'Had a removal done and wants the remaining trees professionally managed',
            ],
            'fears': [
                'Losing mature trees he can\'t replace in his lifetime',
                'Hiring someone who treats every problem as a removal opportunity',
                'Not knowing what\'s actually happening in his tree canopy',
            ],
            'hooks': [
                '"Mature trees are assets — here\'s how to protect them"',
                '"The arborist difference: treatment vs. removal"',
                '"What a Plant Health Care program means for your property\'s value"',
            ],
        },
    ]

    for avatar in avatars:
        # Avatar header
        av_header = doc.add_table(rows=1, cols=1)
        remove_table_borders(av_header)
        av_header.columns[0].width = Inches(7.5)
        hc = av_header.rows[0].cells[0]
        set_cell_bg(hc, avatar['color'])
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.18)
        hp.paragraph_format.space_before = Pt(7)
        hp.paragraph_format.space_after  = Pt(7)
        hn = hp.add_run(avatar['name'])
        hn.bold = True
        hn.font.name = 'Calibri'
        hn.font.size = Pt(13)
        hn.font.color.rgb = WHITE
        hs = hp.add_run('   ·   Primary Service: ' + avatar['service'])
        hs.font.name = 'Calibri'
        hs.font.size = Pt(10)
        hs.font.color.rgb = RGBColor(0xCC, 0xEE, 0xF8)

        # Avatar body in 3-col layout
        body_tbl = doc.add_table(rows=1, cols=3)
        remove_table_borders(body_tbl)
        body_tbl.columns[0].width = Inches(2.8)
        body_tbl.columns[1].width = Inches(2.35)
        body_tbl.columns[2].width = Inches(2.35)

        col_bgs = [ASG_TINT, ASG_LIGHT_BG, 'FFFFFF']
        col_labels = ['WHO THEY ARE', 'WHAT THEY FEAR', 'HOOKS THAT WORK']
        col_content = [
            [avatar['profile']] + ['• ' + t for t in avatar['triggers']],
            ['• ' + f for f in avatar['fears']],
            [h for h in avatar['hooks']],
        ]

        for ci in range(3):
            cell = body_tbl.rows[0].cells[ci]
            set_cell_bg(cell, col_bgs[ci])
            p0 = cell.paragraphs[0]
            p0.paragraph_format.left_indent  = Inches(0.1)
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after  = Pt(2)
            rl = p0.add_run(col_labels[ci] + '\n')
            rl.bold = True
            rl.font.name = 'Calibri'
            rl.font.size = Pt(8)
            rl.font.color.rgb = ASG_BLUE

            for line in col_content[ci]:
                lp = cell.add_paragraph()
                lp.paragraph_format.left_indent  = Inches(0.1)
                lp.paragraph_format.space_before = Pt(0)
                lp.paragraph_format.space_after  = Pt(3)
                lr = lp.add_run(line)
                lr.font.name = 'Calibri'
                lr.font.size = Pt(9)
                lr.font.color.rgb = BODY_COLOR

        sp(doc, 12)

    # Closing note
    divider(doc, ASG_BLUE_HEX)
    sp(doc, 4)

    close_tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(close_tbl)
    close_tbl.columns[0].width = Inches(7.5)
    cc = close_tbl.rows[0].cells[0]
    set_cell_bg(cc, ASG_TINT)
    cp2 = cc.paragraphs[0]
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp2.paragraph_format.space_before = Pt(10)
    cp2.paragraph_format.space_after  = Pt(10)
    cr = cp2.add_run(
        'Charlie isn\'t hard to impress. He\'s been pitched by generalists his whole career.\n'
        'Walk in knowing what PHC is, why his ISA credential matters, why he\'s hemorrhaging on estimate '
        'conversions,\nand what a property manager relationship is actually worth — '
        'and he\'ll tell his wife that night he finally talked to someone who gets it.\n\n'
        'That\'s authority.'
    )
    cr.bold = True
    cr.font.name = 'Calibri'
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0x1A, 0x25, 0x33)


# ── Footer ─────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        fr = fp.add_run('Authority Systems Group™  —  Internal Asset  ·  Niche Intelligence Library  ·  tree-service  ·  2026-03-08')
        fr.font.name = 'Calibri'
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = FOOTER_COLOR


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    build_cover(doc)
    build_avatar(doc)
    build_pain_points(doc)
    build_insider_intel(doc)
    build_customer_avatars(doc)

    doc.save(OUT_DOCX)
    print(f'✓  Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
