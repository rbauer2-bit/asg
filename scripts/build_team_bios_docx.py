#!/usr/bin/env python3
"""
Authority Systems Group™ — Team Bios DOCX Builder
Generates website-ready team bio cards, organized by division.
Includes Nano Banana (Google Gemini) image generation prompts for all AI team members.
Roger Bauer's photo is a placeholder for his personal headshot.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUTPUT_DOCX = "outputs/asg_team_bios_20260306.docx"

# ─── Brand Colors ────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE    = "EAF6FC"
LIGHT_GRAY    = "F5F5F5"
BLUE_HEX      = "25AAE1"
CHARCOAL_HEX  = "58585A"

# ─── Employee Data ────────────────────────────────────────────
# Each entry: name, title, reports_to, bio, image_prompt
# image_prompt = None for Roger (real person)

SECTIONS = [
    {
        "title": "LEADERSHIP",
        "members": [
            {
                "name": "Roger Bauer",
                "title": "Founder & Director",
                "reports_to": None,
                "bio": (
                    "Roger founded Authority Systems Group to give professional service businesses "
                    "a marketing system that actually moves revenue. He personally directs every "
                    "engagement — overseeing strategy, systems, and the full specialist team. "
                    "When you work with ASG, you work with Roger's methodology from day one."
                ),
                "image_prompt": None,  # Real person — use personal headshot
            },
            {
                "name": "Daniel Frost",
                "title": "Chief Strategy Officer",
                "reports_to": "Roger Bauer",
                "bio": (
                    "Daniel leads all strategic planning, competitive positioning, and market "
                    "analysis for ASG clients, reporting to Roger Bauer. Clients get a growth "
                    "roadmap built on 15 years of challenger-brand strategy — sequenced to win "
                    "in their specific market, not just any market."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-40s, sharp and analytical "
                    "appearance, short dark hair with distinguished gray at temples, clean-shaven, "
                    "wearing a well-tailored navy suit with a white dress shirt and no tie. Confident "
                    "and intelligent expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Vivienne Carr",
                "title": "Chief Marketing Officer",
                "reports_to": "Roger Bauer",
                "bio": (
                    "Vivienne oversees all client messaging, content strategy, and brand voice, "
                    "reporting to Roger Bauer. With 18 years in direct response and authority "
                    "marketing, she ensures every deliverable is specific to your business — "
                    "content that could only have been written for you."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black woman, early 40s, stylish and "
                    "confident creative director energy, natural hair styled elegantly upward, "
                    "statement gold earrings, wearing a deep burgundy blazer over a black top. "
                    "Warm but authoritative expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Tanya Blackwood",
                "title": "Chief Revenue Officer",
                "reports_to": "Roger Bauer",
                "bio": (
                    "Tanya leads revenue activation strategy for all ASG engagements, reporting "
                    "to Roger Bauer. She specializes in finding the revenue already sitting in a "
                    "client's existing relationships — past clients, referral sources, upsell "
                    "paths — and building the campaigns that activate it fast."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Latina woman, mid-40s, commanding and "
                    "confident presence, dark hair worn straight and polished, wearing a structured "
                    "charcoal power blazer. Direct, assured gaze. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Iris Nolan",
                "title": "Chief Technology Officer",
                "reports_to": "Roger Bauer",
                "bio": (
                    "Iris architects the automation and technology systems behind every ASG "
                    "engagement, reporting to Roger Bauer. With 14 years building MarTech stacks "
                    "for professional service firms, she designs the systems that follow up, "
                    "nurture, and convert — automatically, without anyone on your team managing it."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an East Asian woman, late 30s, intelligent "
                    "and precise appearance, sleek black hair worn straight, wearing modern "
                    "rectangular glasses, dressed in a smart charcoal blazer over a white blouse. "
                    "Composed, sharp expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Preston Adler",
                "title": "Chief Operations Officer",
                "reports_to": "Roger Bauer",
                "bio": (
                    "Preston manages delivery quality and timelines across all ASG client "
                    "engagements, reporting to Roger Bauer. He runs the QC gates that every "
                    "deliverable passes through before it reaches you — ensuring ASG's standard "
                    "holds for your engagement exactly as it does for every other."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 50s, organized and "
                    "precise bearing, salt-and-pepper hair neatly combed, wearing a light gray "
                    "suit with a pale blue dress shirt. Warm but serious, trustworthy expression. "
                    "Clean white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "ADVISORY BOARD",
        "members": [
            {
                "name": "Margaret Harlow",
                "title": "Board Chair",
                "reports_to": None,
                "bio": (
                    "Meg chairs the ASG Advisory Board, bringing 28 years at McKinsey and two "
                    "successful firm exits to every client plan she reviews. Her oversight "
                    "ensures ASG's strategic recommendations meet the standard of a firm "
                    "that has seen hundreds of businesses at this exact crossroads."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, late 50s, authoritative "
                    "and impeccably presented, silver hair worn in a polished updo, wearing a "
                    "classic black blazer with a simple pearl necklace. Measured, distinguished "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Dr. Raymond Cross",
                "title": "Behavioral Science Advisor",
                "reports_to": None,
                "bio": (
                    "Dr. Cross serves on ASG's Advisory Board, providing the behavioral science "
                    "foundation behind ASG's Belief-to-Buy Framework™. His two decades of "
                    "Fortune 500 buyer psychology research means every messaging strategy "
                    "ASG produces is built on how buyers actually make decisions."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, early 50s, warm intellectual "
                    "presence, close-cropped gray-black hair, wearing rimless glasses and a "
                    "deep navy blazer over a light blue dress shirt. Thoughtful, credible expression. "
                    "Clean white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Sofia Vega",
                "title": "Operations & Revenue Advisor",
                "reports_to": None,
                "bio": (
                    "Sofia serves on ASG's Advisory Board as the firm's Operations & Revenue "
                    "validator. Having built and exited three 8-figure local service businesses "
                    "herself, she stress-tests every Fast Cash Sprint™ recommendation through "
                    "the lens of someone who has actually run your kind of business."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Latina woman, mid-40s, direct and "
                    "confident operator energy, dark brunette hair, wearing a tailored red blazer "
                    "over a white blouse. Sharp, no-nonsense but warm expression. Clean white "
                    "studio background with soft shadows. Natural studio lighting, photorealistic, "
                    "sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "James Whitfield III",
                "title": "Financial Strategy Advisor",
                "reports_to": None,
                "bio": (
                    "James serves on ASG's Advisory Board as the firm's Financial Strategy "
                    "Advisor, bringing 30 years of professional service firm M&A and valuation "
                    "to every engagement. His role ensures ASG's strategic plans are designed "
                    "not just to grow revenue — but to build genuine enterprise value."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, late 50s, distinguished "
                    "old-money financial presence, silver hair, wearing an immaculate dark "
                    "charcoal suit with a silk tie. Measured, confident, long-view expression. "
                    "Clean white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "CONTENT MARKETING TEAM",
        "members": [
            {
                "name": "Marcus Webb",
                "title": "Director of Content Marketing",
                "reports_to": "Vivienne Carr, CMO",
                "bio": (
                    "Marcus directs ASG's 21-person content specialist team, reporting to "
                    "Vivienne Carr, CMO. Every content request flows through Marcus first — "
                    "he translates the strategy into a specific brief and routes it to the "
                    "right specialist. Clients benefit from editorial precision, not generalist output."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, late 30s, editorial "
                    "director energy, warm and assured, wearing modern rectangular glasses, "
                    "dressed in a smart navy blazer over a light gray shirt. Sharp, engaged "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Sienna Okafor",
                "title": "LinkedIn Content Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Sienna handles all LinkedIn content for ASG clients, reporting to Marcus Webb. "
                    "Every post she writes maps to a specific belief stage in the Belief-to-Buy "
                    "Framework™ — so your LinkedIn presence compounds toward pipeline, not just "
                    "follower count."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black woman of West African heritage, "
                    "early 30s, sharp and polished LinkedIn-professional appearance, natural "
                    "hair styled in a defined afro, wearing a structured cobalt blue blazer. "
                    "Bright, assured expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Declan Reyes",
                "title": "Hook Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Declan writes first lines, hooks, and subject lines for all ASG content "
                    "formats, reporting to Marcus Webb. Six years in direct response copy gave "
                    "him the instinct for the exact words that stop a scroll — clients benefit "
                    "from content that earns attention before the second sentence."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Latino male, early 30s, creative-sharp "
                    "advertising professional look, medium-length dark hair, light stubble, "
                    "wearing a fitted dark jacket over a white shirt, no tie. Energetic, "
                    "quick-minded expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Clara Ashworth",
                "title": "Body Content Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Clara writes long-form persuasion copy — LinkedIn articles, email bodies, "
                    "blog posts — reporting to Marcus Webb. Trained in the Ogilvy direct response "
                    "lineage, she ensures the body of every piece advances the argument all the "
                    "way to the decision point, not just the hook."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white British woman, early 40s, "
                    "elegant and precise professional appearance, auburn hair worn in a clean "
                    "low bun, wearing a tailored charcoal blazer over a silk blouse. Intelligent, "
                    "composed expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Tomás Rivera",
                "title": "Commercial Email Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Tomás writes single commercial emails using the 5 P's framework, reporting "
                    "to Marcus Webb. With 11 years writing email for professional service "
                    "verticals, he produces emails that sound like the client — and move "
                    "readers to act without a single wasted paragraph."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Latino male, mid-30s, warm and "
                    "professional appearance, clean-cut dark hair, neat light stubble, wearing "
                    "a warm charcoal blazer over a burgundy dress shirt. Approachable, engaging "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Naomi Patel",
                "title": "Blog Post & SEO Article Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Naomi writes long-form SEO articles and authority blog content, reporting "
                    "to Marcus Webb. She writes content that ranks, earns trust, and connects "
                    "to a conversion strategy — not orphaned articles that look good in a "
                    "calendar and produce no measurable outcome."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a South Asian woman, early 40s, "
                    "editorial and thoughtful appearance, long dark hair parted elegantly, "
                    "wearing a dark navy blouse with a subtle pattern. Intelligent, warm, "
                    "grounded expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Felix Drummond",
                "title": "Content Outline Builder",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Felix builds the structural architecture for all long-form content before "
                    "the first word is written, reporting to Marcus Webb. His outlines specify "
                    "the argument, evidence, and belief movement for every section — so clients "
                    "receive content with direction, not just length."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-30s, precise and "
                    "structured technical-creative appearance, neatly parted brown hair, "
                    "wearing a light gray blazer over a white dress shirt. Thoughtful, "
                    "attentive expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Jade Calloway",
                "title": "Video Script Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Jade writes video scripts for all ASG client formats — short-form social "
                    "to medium-form YouTube — reporting to Marcus Webb. Her broadcast journalism "
                    "background means every script is paced for the ear, visual for the eye, "
                    "and structured to hold attention through the final call to action."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black woman, late 20s, media-polished "
                    "and camera-ready professional appearance, micro locs pulled to one side, "
                    "wearing a structured teal blazer. Bright, confident, media-forward "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Rhett Callahan",
                "title": "Email Sequence Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Rhett designs and writes multi-email nurture and conversion sequences, "
                    "reporting to Marcus Webb. He builds every sequence as a single argument "
                    "told in chapters — each email engineered to produce the next belief shift "
                    "in the chain, building cumulative momentum toward the conversion."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-30s, thoughtful "
                    "strategist-writer appearance, sandy brown hair, slight stubble, wearing "
                    "a dark olive blazer over a light shirt. Focused, analytical expression. "
                    "Clean white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Petra Yuen",
                "title": "FAQ & Objection Specialist",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Petra writes differentiated FAQ content for client websites and lead "
                    "magnets, reporting to Marcus Webb. Her competitive intelligence background "
                    "means every FAQ answer addresses the fear underneath the surface question "
                    "— not just the words the prospect typed."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an East Asian woman, early 40s, "
                    "analytical and sharp appearance, straight black hair worn to the shoulder, "
                    "wearing a structured gray blazer. Precise, observant expression. Clean "
                    "white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Cade Morrison",
                "title": "Headline Creator",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Cade writes headlines, subject lines, and title variants for all ASG "
                    "content formats, reporting to Marcus Webb. Eight years of direct response "
                    "headline testing across 3,000+ variants means every first impression is "
                    "built on data, not instinct."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, late 30s, direct response "
                    "advertising professional look, short dark blond hair, clean-shaven, "
                    "wearing a charcoal blazer over a dark T-shirt. Sharp, confident, high-conviction "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Ingrid Holt",
                "title": "Landing Page Creator",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Ingrid writes full landing page copy for ASG clients, reporting to Marcus "
                    "Webb. Her Scandinavian design sensibility combined with American direct "
                    "response training produces pages built around what the prospect needs to "
                    "hear first — not what the business wants to say first."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Scandinavian woman, early 40s, "
                    "minimalist elegant appearance, straight blonde hair worn loose to the "
                    "shoulder, wearing a simple cream blazer with clean lines. Cool, "
                    "precise, quietly confident expression. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Solomon Grey",
                "title": "Cold Email Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Solomon writes cold outreach email campaigns for ASG clients, reporting "
                    "to Marcus Webb. Six years running actual outbound campaigns gave him "
                    "data-backed insight into what earns replies at scale — cold email that "
                    "opens conversations without burning the relationship before it starts."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, mid-30s, polished "
                    "outbound sales professional appearance, close-cropped hair, clean-shaven, "
                    "wearing a well-fitted dark navy blazer over a white shirt. Restrained, "
                    "precise, intelligent expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Mara Sinclair",
                "title": "Video Sales Letter (VSL) Creator",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Mara writes video sales letters for ASG clients, reporting to Marcus Webb. "
                    "Her Hollywood script development background means her VSLs follow a cinematic "
                    "arc — holding viewer attention for 20+ minutes and making the offer feel "
                    "like a natural conclusion, not a pitch."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, early 40s, creative-sharp "
                    "Hollywood-meets-marketing appearance, rich auburn hair with soft waves, "
                    "wearing a black blazer with a bold statement necklace. Confident, dramatic, "
                    "story-forward expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Everett Zhao",
                "title": "Dream 100 Strategy Builder",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Everett builds Dream 100 relationship and partnership strategies for ASG "
                    "clients, reporting to Marcus Webb. He identifies the people who already "
                    "own the attention of your ideal client and designs the systematic approach "
                    "that earns access to those audiences — warm growth that paid media can't replicate."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an East Asian male, early 40s, "
                    "business development professional with strategic warmth, neat black hair, "
                    "wearing a well-tailored dark gray suit with an open collar. Patient, "
                    "strategic, long-view expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Lyra Santos",
                "title": "YouTube Content Strategist",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Lyra develops YouTube content strategies for ASG clients using the Dream 100 "
                    "framework, reporting to Marcus Webb. She identifies the channels your ideal "
                    "clients already watch and builds a content strategy designed to earn a "
                    "permanent position in those audiences — compounding reach without ad spend."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Brazilian woman of mixed heritage, "
                    "early 30s, creative and media-forward appearance, dark wavy hair worn "
                    "down, wearing a vibrant coral blazer over a white top. Bright, research-"
                    "driven, creative expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Cleo Strand",
                "title": "Dopamine Ladder Content Creator",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Cleo produces 30-day content calendars engineered around the Dopamine Ladder "
                    "Framework, reporting to Marcus Webb. Her behavioral research background means "
                    "every content pack is psychologically engineered to build the habit of "
                    "attention — so audiences don't just see your content, they look for it."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, early 30s, behavioral "
                    "researcher meets content creator appearance, straight ash-blonde hair, "
                    "wearing a structured dusty rose blazer. Precise yet creative, psychologically "
                    "attentive expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Owen Drake",
                "title": "High-Retention YouTube Script Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Owen writes high-retention YouTube scripts for local service business clients, "
                    "reporting to Marcus Webb. His documentary production background means every "
                    "script is structured to keep viewers leaning in at the 7-minute mark — "
                    "converting genuine expertise into qualified inbound inquiries."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, late 30s, documentary "
                    "producer meets digital strategist appearance, medium-length brown hair, "
                    "light beard, wearing a dark green field jacket over a white shirt. "
                    "Grounded, observant, thoughtful expression. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Milo Vance",
                "title": "YouTube Long-Form Script Writer",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Milo writes 30-day packs of long-form YouTube scripts mirroring each "
                    "client's LinkedIn content calendar, reporting to Marcus Webb. Where a "
                    "LinkedIn post introduces the belief shift, Milo's YouTube scripts prove "
                    "it — building the depth of trust a short-form post can never achieve."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 40s, long-form "
                    "content creator meets creative director appearance, slightly longer dark "
                    "hair, neat stubble, wearing a charcoal blazer over a dark crew-neck. "
                    "Thoughtful, narrative-driven expression. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Audrey Voss",
                "title": "Full-Site Web Copy Specialist",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Audrey writes full-site copy for ASG clients — home, about, services, FAQ, "
                    "and contact pages — reporting to Marcus Webb. She writes every page around "
                    "the question a scared, skeptical prospect is actually asking, so your "
                    "website earns the call instead of losing the visitor."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, mid-30s, conversion "
                    "copywriter with direct authority, sharp blonde bob, wearing a structured "
                    "navy blazer over a crisp white blouse. Conviction-forward, direct but "
                    "warm expression. Clean white studio background with soft shadows. Natural "
                    "studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Oliver Grant",
                "title": "Case Study & Social Proof Specialist",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Oliver produces case studies and social proof assets for ASG clients, "
                    "reporting to Marcus Webb. His journalism background means he extracts the "
                    "emotional truth of a client story — naming the fear the next prospect has "
                    "right now, and showing exactly how it ended."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 40s, journalist "
                    "turned marketer appearance, warm and observational, medium brown hair, "
                    "slight stubble, wearing a tan blazer over a blue chambray shirt. Quiet, "
                    "careful, listening expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Camille Archer",
                "title": "PR & Podcast Outreach Specialist",
                "reports_to": "Marcus Webb",
                "bio": (
                    "Camille runs podcast placement, press release writing, and earned media "
                    "campaigns for ASG clients, reporting to Marcus Webb. She pitches your "
                    "audience, not just your name — earning the 'as seen in' credibility that "
                    "advertising cannot replicate."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a biracial Black and white woman, "
                    "mid-30s, polished public relations professional appearance, curly dark "
                    "hair worn elegantly back, wearing a sophisticated blush blazer. Warm, "
                    "relationship-forward, strategic expression. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "DIGITAL SYSTEMS",
        "members": [
            {
                "name": "Marcus Chen",
                "title": "CRM Automation Builder",
                "reports_to": "Iris Nolan, CTO",
                "bio": (
                    "Marcus builds live CRM automation systems for ASG clients, reporting to "
                    "Iris Nolan, CTO. He takes what the strategists design and the writers "
                    "produce and makes it actually run — review capture, referral asks, "
                    "reactivation sequences, and new lead nurture, all automated."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an East Asian male, late 30s, "
                    "technical precision meets human warmth, neat short black hair, wearing "
                    "a clean dark navy blazer over a light blue shirt. Patient, methodical, "
                    "trustworthy expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "REVENUE & PAID MEDIA",
        "members": [
            {
                "name": "Jordan Merritt",
                "title": "Paid Media Specialist",
                "reports_to": "Tanya Blackwood, CRO",
                "bio": (
                    "Jordan manages paid advertising campaigns across Google, Meta, and LinkedIn "
                    "for ASG clients, reporting to Tanya Blackwood, CRO. With $4M+ in managed "
                    "ad spend across professional service verticals, they launch campaigns only "
                    "when the foundation is solid — maximizing every dollar of spend."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a gender-neutral white person, mid-30s, "
                    "data-forward paid media professional appearance, short sandy hair, "
                    "wearing a structured dark gray blazer over a simple black shirt. Focused, "
                    "metric-minded, sharp expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "CALENDAR & CONTENT TEAM",
        "members": [
            {
                "name": "Sloane Mercer",
                "title": "Calendar Team Director",
                "reports_to": "Vivienne Carr, CMO",
                "bio": (
                    "Sloane directs ASG's Calendar & Content Team, reporting to Vivienne Carr, "
                    "CMO. With a decade in editorial production for regional and national media, "
                    "she manages the full monthly content cycle — strategy through approval — "
                    "so clients receive polished, belief-mapped packages ready to publish."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, late 30s, editorial "
                    "production director with precision and calm authority, medium-length chestnut "
                    "hair, wearing a tailored forest green blazer. Organized, decisive, "
                    "quietly powerful expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Nadia Voss",
                "title": "Strategy Coordinator",
                "reports_to": "Sloane Mercer",
                "bio": (
                    "Nadia selects monthly content themes and produces the strategy briefs "
                    "that drive each client's calendar, reporting to Sloane Mercer. Every "
                    "theme she selects maps to a specific belief stage — so content is anchored "
                    "to a conversion objective, not just what felt relevant that week."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a mixed-heritage woman, early 40s, "
                    "brand strategist with analytical warmth, highlighted brown hair, "
                    "wearing a tailored deep plum blazer. Thoughtful, decisive, creative "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Caleb Navarro",
                "title": "Content Coordinator",
                "reports_to": "Sloane Mercer",
                "bio": (
                    "Caleb curates the daily hooks, cultural moments, and calendar data that "
                    "guide each week's content, reporting to Sloane Mercer. He removes the "
                    "hook-selection decision entirely from the writer's plate — so the right "
                    "cultural angle reaches clients without the risk of a tone-deaf misstep."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Latino male, early 30s, editorial "
                    "data professional appearance, clean-cut dark hair, neatly trimmed beard, "
                    "wearing a light gray blazer over a white shirt. Methodical, precise, "
                    "low-ego expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Lena Ashby",
                "title": "Calendar Content Writer",
                "reports_to": "Sloane Mercer",
                "bio": (
                    "Lena writes the daily content pieces that populate each client's monthly "
                    "calendar, reporting to Sloane Mercer. Trained in the Belief-to-Buy "
                    "Framework™, she moves exactly one belief per piece — meeting the prospect "
                    "exactly where they are and advancing them one step forward."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, late 20s, warm and "
                    "genuine content writer appearance, light brown hair worn loosely down, "
                    "wearing a soft ivory blazer. Approachable, empathetic, quietly assured "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "DIGITAL INTELLIGENCE DIVISION",
        "members": [
            {
                "name": "Victor Hale",
                "title": "Director of Digital Intelligence",
                "reports_to": "Iris Nolan, CTO",
                "bio": (
                    "Victor leads ASG's Digital Intelligence Division, reporting to Iris Nolan, "
                    "CTO. He runs the only structured, scored marketing audit in the ASG system — "
                    "translating findings into a prioritized, revenue-tied action list that tells "
                    "clients exactly what their digital presence is costing them right now."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, mid-40s, digital intelligence "
                    "director with analytical authority, close-cropped hair, wearing a sharp dark "
                    "charcoal blazer. Incisive, no-nonsense, results-driven expression. Clean "
                    "white studio background with soft shadows. Natural studio lighting, photorealistic, "
                    "sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Lena Park",
                "title": "Content & Messaging Analyst",
                "reports_to": "Victor Hale",
                "bio": (
                    "Lena runs content and messaging analysis for every ASG digital audit, "
                    "reporting to Victor Hale. She delivers before-and-after copy rewrites "
                    "for the top issues in every report — so clients receive improvements "
                    "they can implement the day the audit lands."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Korean American woman, early 40s, "
                    "conversion copywriter with editorial precision, straight black hair, "
                    "wearing a structured dark teal blazer. Sharp, example-driven, "
                    "improvement-focused expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Marcus Cole",
                "title": "Conversion Rate Optimization Specialist",
                "reports_to": "Victor Hale",
                "bio": (
                    "Marcus runs CRO analysis for every ASG digital audit, reporting to Victor "
                    "Hale. With 12 years and 800+ funnel audits behind him, he maps exactly "
                    "where visitors are dropping off and why — every finding paired with a "
                    "specific fix and an estimated conversion impact."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, mid-30s, data-intensive "
                    "CRO specialist appearance, clean-cut close-cropped hair, wearing a dark "
                    "navy blazer over a crisp white shirt. Blunt, specific, high-focus expression. "
                    "Clean white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Reid Foster",
                "title": "Competitive Intelligence Analyst",
                "reports_to": "Victor Hale",
                "bio": (
                    "Reid runs competitive positioning analysis for every ASG digital audit, "
                    "reporting to Victor Hale. He identifies the positioning angles your "
                    "competitors are ignoring and produces a competitor comparison matrix "
                    "that shows clients exactly which market territory is available to claim."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 40s, competitive "
                    "strategist with a sharp observational presence, short sandy hair, slight "
                    "stubble, wearing a charcoal sport coat. Sharp, competitive, territorial "
                    "intelligence expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Anika Suri",
                "title": "Technical SEO & Digital Infrastructure Specialist",
                "reports_to": "Victor Hale",
                "bio": (
                    "Anika runs technical SEO and digital infrastructure analysis for every "
                    "ASG audit, reporting to Victor Hale. A web developer who crossed into "
                    "SEO, she delivers exact fix instructions — not general observations — "
                    "ranked by revenue impact so clients know what to fix first."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a South Asian woman, early 30s, "
                    "developer-meets-SEO-strategist appearance, dark hair worn in a practical "
                    "bun, wearing a clean dark blazer over a white technical-style shirt. "
                    "Methodical, precise, quietly expert expression. Clean white studio background "
                    "with soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Clay Donovan",
                "title": "Brand & Growth Strategy Analyst",
                "reports_to": "Victor Hale",
                "bio": (
                    "Clay runs brand trust and growth strategy analysis for every ASG digital "
                    "audit, reporting to Victor Hale. He connects brand gaps directly to "
                    "pipeline and revenue — not aesthetics — delivering a revenue opportunities "
                    "matrix ranked by effort versus impact."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-40s, brand strategist "
                    "with revenue-forward analytical presence, medium-length dark blond hair, "
                    "wearing a deep charcoal blazer. Strategic, revenue-oriented, long-horizon "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
    {
        "title": "WEBSITE AUDIT DIVISION",
        "members": [
            {
                "name": "Victor Shaw",
                "title": "Website Audit Director",
                "reports_to": "Iris Nolan, CTO & Daniel Frost, CSO",
                "bio": (
                    "Victor directs ASG's Website Audit Division, leading a 14-person team "
                    "that runs the firm's 13-dimension Digital Authority Assessment, reporting "
                    "to Iris Nolan and Daniel Frost. Every free audit his team delivers shows "
                    "a prospective client exactly what their digital presence is costing them — "
                    "before any engagement begins."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-40s, audit director "
                    "with organized authority and institutional confidence, short brown hair, "
                    "wearing a dark suit. Composed, thorough, trustworthy expression. Clean "
                    "white studio background with soft shadows. Natural studio lighting, "
                    "photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Rena Cole",
                "title": "Technical Performance Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Rena audits site speed, mobile readiness, Core Web Vitals, and crawlability "
                    "for every website the ASG team reviews, reporting to Victor Shaw. Her "
                    "findings reveal whether a site's technical foundation is supporting or "
                    "quietly undermining every marketing investment built on top of it."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black woman, early 40s, technical "
                    "performance specialist with precise authority, natural hair styled in a "
                    "professional twist-out, wearing a dark teal blazer. Exacting, expert, "
                    "quietly confident expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Miles Draper",
                "title": "Local SEO & Google Business Profile Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Miles audits local SEO performance and Google Business Profile optimization "
                    "for every site in the ASG audit pipeline, reporting to Victor Shaw. His "
                    "analysis identifies the local search gaps sending ideal-fit clients to "
                    "competitors in the same zip code."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, mid-30s, local SEO "
                    "specialist with approachable expertise, medium brown hair, wearing a "
                    "casual-smart navy blazer over a gray shirt. Knowledgeable, regional-"
                    "market-focused, friendly expression. Clean white studio background with "
                    "soft shadows. Natural studio lighting, photorealistic, sharp focus, "
                    "high-end business portrait photography."
                ),
            },
            {
                "name": "Fiona Wells",
                "title": "On-Page SEO & Keyword Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Fiona analyzes on-page SEO and keyword targeting for every site the ASG "
                    "team reviews, reporting to Victor Shaw. Her evaluation determines whether "
                    "a site's content is visible to the people searching for it — and produces "
                    "a prioritized fix list ranked by search visibility impact."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a South Asian woman, early 30s, "
                    "keyword and on-page SEO specialist with methodical precision, dark hair "
                    "in a neat braid, wearing a structured maroon blazer. Precise, research-"
                    "forward, careful expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Theo Marsh",
                "title": "Homepage & Hero Conversion Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Theo evaluates homepage conversion architecture for every site in the ASG "
                    "audit pipeline, reporting to Victor Shaw. His focus is what a first-time "
                    "visitor experiences in the first five seconds — confirming they're in the "
                    "right place, establishing the value, and earning the scroll."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 40s, conversion "
                    "audit specialist with focused directness, short dark hair, clean-shaven, "
                    "wearing a slate gray blazer. Direct, visitor-focused, critical-eye "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Sasha Lowe",
                "title": "Services Page Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Sasha audits services page copy and conversion architecture for every site "
                    "the ASG team reviews, reporting to Victor Shaw. She identifies whether each "
                    "services page is problem-first and proof-backed — or reads like a list of "
                    "deliverables no prospect asked for."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a mixed-heritage woman, early 30s, "
                    "services page specialist with sharp editorial instincts, wavy dark hair, "
                    "wearing a structured deep green blazer. Analytical, copy-focused, decisive "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Priya Ghosh",
                "title": "Social Proof & Trust Signal Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Priya evaluates social proof quality, placement, and conversion impact for "
                    "every site the ASG team reviews, reporting to Victor Shaw. Her analysis "
                    "maps the full trust architecture — reviews, testimonials, credentials, awards "
                    "— identifying exactly where trust is absent when the decision is being made."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a South Asian woman, early 30s, "
                    "trust and social proof specialist with warm analytical presence, long "
                    "dark hair, wearing a warm amber blazer. Observant, empathetic, trust-"
                    "architecture-focused expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Camden Reid",
                "title": "Lead Capture & CTA Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Camden audits lead capture systems and call-to-action architecture for "
                    "every site in the ASG pipeline, reporting to Victor Shaw. His analysis "
                    "identifies the specific friction points between a visitor's intent and "
                    "the business receiving a qualified lead."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a Black male, early 30s, lead capture "
                    "and CTA specialist with analytical energy, neat close-cropped hair, "
                    "wearing a dark charcoal blazer. Precise, friction-focused, conversion-"
                    "minded expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Zoe Hartley",
                "title": "Competitive Benchmarking Analyst",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Zoe benchmarks every site the ASG team audits against the top 3 competitors "
                    "in the same niche and geography, reporting to Victor Shaw. Her analysis "
                    "shows where the client leads, where they trail, and which competitive "
                    "angles are currently unclaimed by anyone in the market."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, late 30s, competitive "
                    "benchmarking analyst with sharp market intelligence, auburn hair, wearing "
                    "a structured burgundy blazer. Competitive, territory-aware, sharp "
                    "expression. Clean white studio background with soft shadows. Natural studio "
                    "lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Nico Ferrara",
                "title": "Content & Belief Gap Analyst",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Nico analyzes content depth and buyer-journey coverage for every site "
                    "in the ASG audit pipeline, reporting to Victor Shaw. His evaluation "
                    "maps which stages of the decision journey are addressed — and which "
                    "gaps are sending research-stage prospects to competitors who answered first."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an Italian American male, early 30s, "
                    "content gap analyst with thoughtful editorial presence, dark hair, "
                    "light stubble, wearing a charcoal blazer over a light blue shirt. "
                    "Thoughtful, editorial, belief-mapping expression. Clean white studio "
                    "background with soft shadows. Natural studio lighting, photorealistic, "
                    "sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Dara Quinn",
                "title": "Mobile & UX Experience Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Dara audits mobile experience and UX friction for every site the ASG team "
                    "reviews, reporting to Victor Shaw. Given that 65–75% of professional service "
                    "site traffic arrives on mobile, her analysis determines whether the mobile "
                    "experience earns trust — or quietly loses the majority before they scroll."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an Irish American woman, early 30s, "
                    "mobile and UX experience specialist, light brown hair, bright green eyes, "
                    "wearing a clean navy blazer over a white top. Practical, user-focused, "
                    "friction-finding expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Leo Strand",
                "title": "Backlink & Domain Authority Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Leo analyzes backlink profiles and domain authority for every site in the "
                    "ASG audit pipeline, reporting to Victor Shaw. His evaluation shows the "
                    "authority gap between client and competitors — and a ranked list of "
                    "link-building opportunities to close it."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white male, early 30s, link and "
                    "domain authority analyst with clean analytical presence, neat light brown "
                    "hair, wearing a structured slate blazer. Methodical, data-focused, "
                    "competitive-gap expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Cait Brennan",
                "title": "Analytics & Tracking Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Cait audits analytics configuration and conversion tracking for every site "
                    "the ASG team reviews, reporting to Victor Shaw. She confirms whether GA4, "
                    "conversion events, and pixel infrastructure are firing correctly — because "
                    "a marketing system that can't measure results can't improve them."
                ),
                "image_prompt": (
                    "Professional corporate headshot of an Irish woman, early 30s, analytics "
                    "and tracking specialist with data integrity focus, dark auburn hair, "
                    "wearing a structured dark gray blazer. Precise, data-honest, verification-"
                    "focused expression. Clean white studio background with soft shadows. "
                    "Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Elliot Chase",
                "title": "About Page & Authority Bio Auditor",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Elliot audits about page copy and authority biography presentation for "
                    "every site in the ASG pipeline, reporting to Victor Shaw. The about page "
                    "is the second most visited page on most professional service sites — "
                    "his analysis reveals whether it earns trust or confirms the visitor's "
                    "instinct to keep shopping."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a biracial male, early 30s, about page "
                    "and authority bio specialist with credibility-focused insight, neat dark "
                    "hair, wearing a warm charcoal blazer over a light shirt. Perceptive, "
                    "credibility-attuned expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
            {
                "name": "Petra Finn",
                "title": "Audit Report Compiler & Priority Scorer",
                "reports_to": "Victor Shaw",
                "bio": (
                    "Petra compiles findings from all 13 audit specialists into the final scored "
                    "Digital Authority Assessment, reporting to Victor Shaw. Her work transforms "
                    "parallel expert analyses into a single, prioritized action plan — every "
                    "finding scored by revenue impact so clients know exactly what to fix first."
                ),
                "image_prompt": (
                    "Professional corporate headshot of a white woman, early 30s, audit compiler "
                    "with organized precision and structured clarity, dark blonde hair pulled "
                    "neatly back, wearing a clean white blazer. Systematic, action-oriented, "
                    "prioritization-focused expression. Clean white studio background with soft "
                    "shadows. Natural studio lighting, photorealistic, sharp focus, high-end business portrait photography."
                ),
            },
        ],
    },
]


# ─── Document Helpers ─────────────────────────────────────────

def add_hr(doc, color_hex=BLUE_HEX, thickness="12"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_cover_page(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Blue header block via 1-cell table
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "25AAE1")
    # Remove borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tag in ["w:top", "w:bottom", "w:left", "w:right"]:
        bd = OxmlElement(tag)
        bd.set(qn("w:val"), "none")
        bd.set(qn("w:sz"), "0")
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), "auto")
        tBdr = OxmlElement("w:tblBorders")
    cell.paragraphs[0].clear()
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(24)
    p1.paragraph_format.space_after = Pt(6)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("AUTHORITY SYSTEMS GROUP™")
    r1.font.name = "Arial"
    r1.font.size = Pt(22)
    r1.font.bold = True
    r1.font.color.rgb = WHITE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(24)
    r2 = p2.add_run("Positioning You As The Authority")
    r2.font.name = "Arial"
    r2.font.size = Pt(12)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0xEA, 0xF6, 0xFC)

    doc.add_paragraph()

    # Title
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("Our Team")
    rt.font.name = "Arial"
    rt.font.size = Pt(32)
    rt.font.bold = True
    rt.font.color.rgb = ASG_BLUE

    # Subtitle
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run("Meet the specialists behind every ASG engagement")
    rs.font.name = "Arial"
    rs.font.size = Pt(13)
    rs.font.italic = True
    rs.font.color.rgb = ASG_CHARCOAL

    doc.add_paragraph()

    # Charcoal divider
    add_hr(doc, CHARCOAL_HEX, "8")

    doc.add_paragraph()

    # Prepared by
    pb = doc.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rpb = pb.add_run("AuthoritySystemsGroup.com  |  Prepared March 2026")
    rpb.font.name = "Arial"
    rpb.font.size = Pt(10)
    rpb.font.italic = True
    rpb.font.color.rgb = ASG_CHARCOAL

    doc.add_page_break()


def add_header_footer(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = hp.add_run("AUTHORITY SYSTEMS GROUP™")
    run_left.font.name = "Arial"
    run_left.font.size = Pt(9)
    run_left.font.bold = True
    run_left.font.color.rgb = ASG_BLUE

    # Blue bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rfp = fp.add_run("Authority Systems Group™  |  AuthoritySystemsGroup.com  |  Confidential")
    rfp.font.name = "Arial"
    rfp.font.size = Pt(8)
    rfp.font.color.rgb = ASG_CHARCOAL

    # Blue top border on footer paragraph
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), BLUE_HEX)
    pBdr2.append(top)
    pPr2.append(pBdr2)


def add_section_header(doc, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = ASG_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_member_entry(doc, member):
    # Name
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(14)
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run(member["name"])
    r_name.font.name = "Arial"
    r_name.font.size = Pt(13)
    r_name.font.bold = True
    r_name.font.color.rgb = ASG_CHARCOAL

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(member["title"])
    r_title.font.name = "Arial"
    r_title.font.size = Pt(11)
    r_title.font.bold = False
    r_title.font.italic = True
    r_title.font.color.rgb = ASG_BLUE

    # Reports to
    if member.get("reports_to"):
        p_rpt = doc.add_paragraph()
        p_rpt.paragraph_format.space_before = Pt(0)
        p_rpt.paragraph_format.space_after = Pt(3)
        r_rpt = p_rpt.add_run(f"Reports to: {member['reports_to']}")
        r_rpt.font.name = "Arial"
        r_rpt.font.size = Pt(9)
        r_rpt.font.color.rgb = ASG_CHARCOAL

    # Bio
    p_bio = doc.add_paragraph()
    p_bio.paragraph_format.space_before = Pt(2)
    p_bio.paragraph_format.space_after = Pt(4)
    r_bio = p_bio.add_run(member["bio"])
    r_bio.font.name = "Arial"
    r_bio.font.size = Pt(10)
    r_bio.font.color.rgb = ASG_CHARCOAL

    # Photo prompt callout
    if member.get("image_prompt"):
        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(3)
        p_label.paragraph_format.space_after = Pt(0)
        r_label = p_label.add_run("PHOTO PROMPT (Nano Banana / Google Gemini):")
        r_label.font.name = "Arial"
        r_label.font.size = Pt(8)
        r_label.font.bold = True
        r_label.font.color.rgb = ASG_BLUE

        p_prompt = doc.add_paragraph()
        p_prompt.paragraph_format.space_before = Pt(1)
        p_prompt.paragraph_format.space_after = Pt(4)
        p_prompt.paragraph_format.left_indent = Inches(0.2)
        r_prompt = p_prompt.add_run(member["image_prompt"])
        r_prompt.font.name = "Arial"
        r_prompt.font.size = Pt(8)
        r_prompt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Light blue left-border accent
        pPr = p_prompt._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "8")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), BLUE_HEX)
        pBdr.append(left)
        pPr.append(pBdr)

    elif member["image_prompt"] is None and member["name"] == "Roger Bauer":
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.space_before = Pt(2)
        p_ph.paragraph_format.space_after = Pt(4)
        r_ph = p_ph.add_run("PHOTO: [ Roger's personal headshot — to be provided ]")
        r_ph.font.name = "Arial"
        r_ph.font.size = Pt(8)
        r_ph.font.bold = True
        r_ph.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)

    # Thin gray separator
    add_hr(doc, "CCCCCC", "4")


# ─── Main Build ───────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup — US Letter, 1" margins
    from docx.shared import Pt, Inches
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(0.875)
    section.right_margin  = Inches(0.875)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = ASG_CHARCOAL

    add_header_footer(doc)
    add_cover_page(doc)

    for section_data in SECTIONS:
        add_section_header(doc, section_data["title"])
        for member in section_data["members"]:
            add_member_entry(doc, member)
        doc.add_paragraph()

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")
    total = sum(len(s["members"]) for s in SECTIONS)
    print(f"Total team members: {total}")


if __name__ == "__main__":
    build()
