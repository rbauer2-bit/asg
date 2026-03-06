from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
from datetime import date

INPUT_MD = sys.argv[1]
OUTPUT_DOCX = sys.argv[2]
CLIENT_NAME = sys.argv[3] if len(sys.argv) > 3 else "Client"

ASG_BLUE = "25aae1"
ASG_CHARCOAL = "58585a"


def hex_to_rgb(hexstr):
    h = hexstr.strip().lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def add_cover(doc, title, client_name):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = hex_to_rgb(ASG_BLUE)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run(f"Prepared for: {client_name}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = hex_to_rgb(ASG_CHARCOAL)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = p3.add_run(f"Date: {date.today().isoformat()}")
    r3.font.size = Pt(11)
    r3.font.color.rgb = hex_to_rgb(ASG_CHARCOAL)

    doc.add_page_break()


def apply_footer(doc, client_name):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = f"Authority Systems Group™ — Confidential. Prepared exclusively for {client_name}."
        para.runs[0].font.size = Pt(9)
        para.runs[0].font.color.rgb = hex_to_rgb(ASG_CHARCOAL)


def write_styled_docx(md_path, docx_path, client_name):
    doc = Document()
    add_cover(doc, "Authority Blueprint", client_name)
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    block = []
    def flush_block(b):
        if not b:
            return
        text = '\n'.join(b).strip()
        if not text:
            return
        stripped = text.lstrip()
        if stripped.startswith('#'):
            hashes = 0
            for ch in stripped:
                if ch == '#':
                    hashes += 1
                else:
                    break
            heading_text = stripped[hashes:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            if hashes == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = hex_to_rgb(ASG_BLUE)
            elif hashes == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = hex_to_rgb(ASG_CHARCOAL)
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = hex_to_rgb(ASG_CHARCOAL)
        else:
            for line in text.split('\n'):
                if line.strip().startswith(('-', '*')):
                    p = doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(11)

    for ln in lines:
        if ln.strip() == '':
            flush_block(block)
            block = []
        else:
            block.append(ln)
    flush_block(block)

    # QC sign-off block
    doc.add_page_break()
    doc.add_paragraph('QC Sign-offs', style='Heading 2')
    doc.add_paragraph(f'Strategic Alignment (Daniel Frost): ____________________   Date: ________')
    doc.add_paragraph(f'Messaging Integrity (Vivienne Carr): ____________________   Date: ________')
    doc.add_paragraph(f'Revenue Logic (Tanya Blackwood): ____________________   Date: ________')
    doc.add_paragraph(f'Delivery Standard (Preston Adler): ____________________   Date: ________')

    apply_footer(doc, client_name)
    doc.save(docx_path)
    print('STYLED_DOCX_CREATED')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: style_docx.py input_md output_docx [client_name]')
        sys.exit(1)
    write_styled_docx(INPUT_MD, OUTPUT_DOCX, CLIENT_NAME)
