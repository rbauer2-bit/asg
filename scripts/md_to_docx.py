import sys
from docx import Document
from docx.shared import Pt

INPUT = sys.argv[1]
OUTPUT = sys.argv[2]

def write_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    block = []
    def flush_block(b):
        if not b:
            return
        text = '\n'.join(b).strip()
        if not text:
            return
        # heading detection
        stripped = text.lstrip()
        if stripped.startswith('#'):
            # count hashes
            hashes = 0
            for ch in stripped:
                if ch == '#':
                    hashes += 1
                else:
                    break
            heading_text = stripped[hashes:].strip()
            if hashes == 1:
                p = doc.add_paragraph()
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(18)
            elif hashes == 2:
                p = doc.add_paragraph()
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(14)
            else:
                p = doc.add_paragraph()
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(12)
        else:
            # simple list handling: keep lines as separate paragraphs when they start with - or *
            for line in text.split('\n'):
                if line.strip().startswith(('-', '*')):
                    p = doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)

    for ln in lines:
        if ln.strip() == '':
            flush_block(block)
            block = []
        else:
            block.append(ln)
    flush_block(block)
    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: md_to_docx.py input.md output.docx')
        sys.exit(1)
    write_docx(INPUT, OUTPUT)
    print('DOCX_CREATED')
