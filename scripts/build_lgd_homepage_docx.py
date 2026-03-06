#!/usr/bin/env python3
"""
build_lgd_homepage_docx.py
Authority Systems Group™

Generates the Legal Growth Dynamics Homepage Copy deliverable as a branded DOCX.
Usage: python3 scripts/build_lgd_homepage_docx.py
Output: legalgrowth-dynamics_homepage-copy_20260306.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT      = "legalgrowth-dynamics_homepage-copy_20260306.docx"
CLIENT      = "Legal Growth Dynamics"
DOC_TITLE   = "Homepage Copy"
PREP_DATE   = "March 6, 2026"
PREP_BY     = "Audrey Voss, Full-Site Web Copy Specialist  |  Authority Systems Group\u2122"

# ── ASG brand colors ──────────────────────────────────────────────────────────
BLUE        = RGBColor(0x25, 0xaa, 0xe1)
CHARCOAL    = RGBColor(0x58, 0x58, 0x5a)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xEA, 0xF6, 0xFC)
GRAY_NOTE   = RGBColor(0x99, 0x99, 0x99)
FONT        = "Arial"


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _spacing(para, before=0, after=100):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    pPr.append(sp)


def _indent(para, left=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    pPr.append(ind)


def _border_bottom(para, color='25aae1', sz=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b    = OxmlElement('w:bottom')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    str(sz))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def _border_left(para, color='25aae1', sz=18):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    l    = OxmlElement('w:left')
    l.set(qn('w:val'),   'single')
    l.set(qn('w:sz'),    str(sz))
    l.set(qn('w:space'), '6')
    l.set(qn('w:color'), color)
    pBdr.append(l)
    pPr.append(pBdr)


def _cell_bg(cell, hex6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)


def _cell_no_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tcB.append(el)
    tcPr.append(tcB)


def _cell_padding(cell, twips=720):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side in ('top', 'bottom', 'left', 'right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(twips))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)


def _set_table_width(table, twips):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'),    str(twips))
    w.set(qn('w:type'), 'dxa')
    tblPr.append(w)


# ── Paragraph factory helpers ─────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size=11, color=None):
    r = para.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    r.font.color.rgb = color or CHARCOAL
    return r


def h1(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=300, after=100)
    _border_bottom(p, '25aae1', sz=6)
    _run(p, text, bold=True, size=18, color=BLUE)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=220, after=60)
    _run(p, text, bold=True, size=13, color=CHARCOAL)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=160, after=40)
    _run(p, text, bold=True, size=11, color=BLUE)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=40, after=80)
    _run(p, text, size=11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    _spacing(p, before=30, after=30)
    _run(p, text, size=11)
    return p


def arrow(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=30, after=30)
    _indent(p, left=360)
    _run(p, f'\u2192  {text}', size=11)
    return p


def callout(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=120, after=120)
    _indent(p, left=360)
    _border_left(p, '25aae1', sz=18)
    _run(p, text, bold=True, size=11, color=BLUE)
    return p


def dev_note(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=60, after=60)
    _indent(p, left=360)
    _run(p, f'\u25b6 DEVELOPER NOTE: {text}', italic=True, size=9, color=GRAY_NOTE)
    return p


def stage_label(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=20, after=20)
    _run(p, text, italic=True, size=9, color=GRAY_NOTE)
    return p


def tag(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=80, after=10)
    _run(p, f'[ {text} ]', bold=True, italic=True, size=9, color=BLUE)
    return p


def cta(doc, primary, secondary=None):
    p = doc.add_paragraph()
    _spacing(p, before=160, after=60)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _run(p, f'\u25ba  {primary}  \u25c4', bold=True, size=12, color=BLUE)
    if secondary:
        p2 = doc.add_paragraph()
        _spacing(p2, before=20, after=100)
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _run(p2, secondary, italic=True, size=10)


def rule(doc):
    p = doc.add_paragraph()
    _spacing(p, before=160, after=160)
    _border_bottom(p, '25aae1', sz=4)


def gap(doc):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=0)


# ── Header / footer ───────────────────────────────────────────────────────────

def apply_header_footer(doc):
    for section in doc.sections:
        # ── header ──
        hdr  = section.header
        hp   = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        _run(hp, 'AUTHORITY SYSTEMS GROUP\u2122', bold=True, size=9, color=BLUE)
        _run(hp, f'\t{DOC_TITLE}  |  {CLIENT}', size=9, color=CHARCOAL)
        _border_bottom(hp, '25aae1', sz=4)

        # Right-align the tab stop at content width (9072 twips = 6.3")
        pPr  = hp._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)

        # ── footer ──
        ftr = section.footer
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        _border_bottom(fp, '25aae1', sz=4)
        _run(fp,
             f'Authority Systems Group\u2122  |  AuthoritySystemsGroup.com  |  '
             f'Confidential. Prepared exclusively for {CLIENT}.',
             size=8, color=CHARCOAL)


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Full-width blue header table — width in twips (1" = 1440; content = 9072)
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_width(tbl, 9072)
    tbl.style = 'Table Grid'

    cell = tbl.cell(0, 0)
    _cell_bg(cell, '25aae1')
    _cell_no_borders(cell)
    _cell_padding(cell, 800)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _run(p1, 'AUTHORITY SYSTEMS GROUP\u2122', bold=True, size=20, color=WHITE)

    p2 = cell.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _run(p2, 'Positioning You As The Authority', italic=True, size=11, color=LIGHT_BLUE)

    gap(doc)

    # Title
    p_t = doc.add_paragraph()
    p_t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _spacing(p_t, before=200, after=60)
    _run(p_t, DOC_TITLE, bold=True, size=32, color=BLUE)

    # Client
    p_c = doc.add_paragraph()
    p_c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _spacing(p_c, before=40, after=40)
    _run(p_c, CLIENT, size=16, color=CHARCOAL)

    # Date
    p_d = doc.add_paragraph()
    p_d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _spacing(p_d, before=0, after=300)
    _run(p_d, PREP_DATE, italic=True, size=11, color=CHARCOAL)

    # Divider
    p_hr = doc.add_paragraph()
    _spacing(p_hr, before=200, after=200)
    _border_bottom(p_hr, '25aae1', sz=8)

    # Prepared by
    p_b = doc.add_paragraph()
    p_b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _spacing(p_b, before=240, after=0)
    _run(p_b, 'Prepared by', size=10, color=CHARCOAL)

    p_b2 = doc.add_paragraph()
    p_b2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _spacing(p_b2, before=30, after=0)
    _run(p_b2, PREP_BY, bold=True, size=10, color=CHARCOAL)

    doc.add_page_break()


# ── Main document build ───────────────────────────────────────────────────────

def build():
    doc = Document()

    # Margins only — do NOT override page_width/height (default template is Letter)
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(0.875)
        section.right_margin  = Inches(0.875)

    apply_header_footer(doc)
    add_cover(doc)

    # ── DOC OVERVIEW ─────────────────────────────────────────────────────────
    h1(doc, 'Document Overview')
    stage_label(doc, 'Page: Homepage  |  Primary Goal: Trust + Discovery Call Booking')
    stage_label(doc, 'Belief Track: B1 Enemy Belief \u2192 B2 Pre-Frame \u2192 B3 New Worldview')
    stage_label(doc, 'Emotion Track: E1 Dissatisfaction \u2192 E2 Contrast \u2192 E3 Desire \u2192 E5 Relief')
    gap(doc)
    body(doc,
         'NOTE ON NO CASE STUDIES: The Russell Lloyd founding story serves as the primary emotional '
         'anchor in place of client case studies. Process proof (The Triad, DFY/DWY architecture) '
         'replaces outcome-based social proof. Once 2\u20133 case studies exist, Section 5 converts '
         'to a results section and the founder story migrates to the About page.')

    # ── HEADLINE VARIANTS ─────────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Headline Variants')
    stage_label(doc, 'Produced by Cade Morrison, Hook Writer  |  Authority Systems Group\u2122')
    gap(doc)

    h3(doc, 'Variant 1 \u2014 RECOMMENDED  (Enemy Belief + Niche claim)')
    callout(doc,
            'Most Solo Attorneys Don\u2019t Have a Marketing Problem.\n'
            'They Have Three Missing Systems.')

    h3(doc, 'Variant 2 \u2014 Outcome-first')
    body(doc,
         'A Predictable Caseload for Solo Attorneys in Family Law, Criminal Defense, and '
         'Personal Injury \u2014 Built on Systems, Not Ad Spend')

    h3(doc, 'Variant 3 \u2014 Niche exclusivity  (recommended A/B test)')
    body(doc,
         'We Only Work With Solo and Small Law Firm Attorneys in Three Practice Areas.\n'
         'If That\u2019s You, Keep Reading.')

    h3(doc, 'Variant 4 \u2014 Pattern interrupt')
    body(doc,
         'Your Last Marketing Agency Wasn\u2019t Wrong.\nThey Were Just Using the Wrong Playbook.')

    h3(doc, 'Variant 5 \u2014 Direct outcome')
    body(doc, 'Stop Chasing Cases. Start Running a Practice That Attracts Them.')

    gap(doc)
    callout(doc,
            'Cade\u2019s recommendation: Variant 1. Opens at the Enemy Belief stage \u2014 the '
            'attorney who has tried marketing before and is skeptical \u2014 and reframes the '
            'problem without being combative. \u201cThree missing systems\u201d creates '
            'intellectual curiosity that drives the scroll. Variant 3 is the recommended A/B '
            'test for cold traffic.')

    # ── SECTION 1: ABOVE THE FOLD ─────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 1 \u2014 Above the Fold')
    stage_label(doc, 'Belief: B1 Enemy Belief  |  Emotion: E1 Dissatisfaction')
    gap(doc)

    h2(doc, 'H1 \u2014 Primary Headline')
    callout(doc,
            'Most Solo Attorneys Don\u2019t Have a Marketing Problem.\n'
            'They Have Three Missing Systems.')

    h2(doc, 'Subhead')
    body(doc,
         'Legal Growth Dynamics builds predictable growth infrastructure for solo practitioners '
         'and small law firms in family law, criminal defense, and personal injury. Not ad '
         'campaigns. Not a pretty website. A complete system \u2014 tailored to how your specific '
         'practice type actually generates cases.')

    h2(doc, 'Authority Anchor')
    body(doc,
         'Exclusively serving solo and small firm attorneys  |  '
         'Family Law \u00b7 Criminal Defense \u00b7 Personal Injury  |  Jeffersonville, IN')

    h2(doc, 'Primary CTA')
    cta(doc,
        'Book Your Free 30-Minute Strategy Session',
        '[Phone number]  \u2014  or  \u2014  See How It Works \u2193')

    dev_note(doc,
             'Hero section. H1 first line in ASG charcoal (#58585a), second line in ASG blue '
             '(#25aae1) to emphasize "three missing systems." Authority anchor in smaller muted '
             'text. CTA button full-width on mobile. No stock photography of courtrooms or gavels.')

    # ── SECTION 2: THE PROBLEM ────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 2 \u2014 The Problem: Naming Their Reality')
    stage_label(doc, 'Belief: B1 Enemy Belief  |  Emotion: E1 Dissatisfaction')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc, 'If This Sounds Familiar, You\u2019re Not Alone')

    h2(doc, 'Body Copy')
    body(doc,
         'You built your practice the way most attorneys do \u2014 referrals, reputation, and '
         'hustle. For a while, it worked. Some months are genuinely great. Cases come in, '
         'you\u2019re billing, you feel like you\u2019ve figured it out.')
    body(doc,
         'Then there\u2019s March. Or August. Or whenever your calendar goes quiet and you '
         'remember that the whole thing is held together by relationships you can\u2019t fully '
         'control and timing you can\u2019t predict.')
    body(doc,
         'You\u2019ve probably tried to fix it. Maybe Google Ads that burned through budget '
         'without producing the right clients. Maybe an SEO company that showed you rankings '
         'but not revenue. Maybe a website redesign that looked better but converted the same.')
    body(doc,
         'Here\u2019s what nobody told you: those tactics weren\u2019t wrong. They were just '
         'applied without the infrastructure that makes them work. A family law attorney in solo '
         'practice doesn\u2019t need the same marketing system as a 40-attorney firm. The '
         'playbook is different. And most agencies only know one playbook.')

    dev_note(doc,
             'Two-column layout on desktop: left copy, right subtle icon set representing the '
             'three systems. Single column on mobile. No imagery \u2014 attorney is reading, '
             'not scanning.')

    # ── SECTION 3: AUTHORITY ──────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 3 \u2014 Authority Positioning: Why Roger, Why This Firm')
    stage_label(doc, 'Belief: B2 Pre-Frame  |  Emotion: E2 Contrast')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc, 'Why Legal Growth Dynamics Exists \u2014 And Who It\u2019s Built For')

    h2(doc, 'Body Copy')
    body(doc,
         'Roger Bauer started Legal Growth Dynamics because of a family law attorney named '
         'Russell Lloyd.')
    body(doc,
         'Russell was exceptional at his work. He cared about his clients in a way that showed '
         'in every case he took. His results were real. His reputation among the people he '
         'served was spotless. And he was grinding \u2014 constantly stressed about where the '
         'next case was coming from, working too many hours, unable to build the kind of '
         'practice that matched his skill level.')
    body(doc,
         'The problem wasn\u2019t Russell\u2019s ability as an attorney. The problem was that '
         'nobody had ever built him the right marketing infrastructure for a solo family law '
         'practice specifically. He was using generic advice designed for businesses nothing '
         'like his.')
    body(doc,
         'Roger left a full-time career in marketing consulting to solve that problem \u2014 '
         'not generically, but for the specific type of law firm he understood best: solo '
         'practitioners and small firms in family law, criminal defense, and personal injury, '
         'who deserve a practice as strong as their legal ability.')

    h2(doc, 'Three Key Commitments')
    arrow(doc,
          'We only work in three practice areas. Family law, criminal defense, and personal '
          'injury. Not estate planning, not business law, not whatever walked in the door. '
          'Three areas, deep expertise, no dilution.')
    arrow(doc,
          'We match the engagement model to where you actually are. Not every attorney needs '
          'full-service management. Some want their marketing handled completely (DFY). Others '
          'want to build the capability in-house with expert guidance (DWY). We offer both '
          'because the right answer depends on your practice, not our revenue preference.')
    arrow(doc,
          'We measure what moves the needle \u2014 cases, not clicks. Traffic, impressions, '
          'and rankings are inputs. Retained clients are the output. Every recommendation '
          'connects to case volume and revenue.')

    dev_note(doc,
             'Roger\u2019s professional photo \u2014 approachable, not stock \u2014 alongside '
             'this section. Do not use a stock photo of a person. Confirm with Roger whether '
             'the Russell Lloyd mention is appropriate to publish before launch.')

    # ── SECTION 4: SERVICES ───────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 4 \u2014 Services: Two Ways to Work Together')
    stage_label(doc, 'Belief: B3 New Worldview  |  Emotion: E3 Desire')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc,
         'Two Engagement Models. One Goal: A Practice That Grows on a System, Not on Luck.')
    body(doc,
         'Every attorney\u2019s situation is different. Legal Growth Dynamics offers two ways '
         'to work together \u2014 depending on how much you want to own the work versus have '
         'it handled for you.')

    gap(doc)
    tag(doc, 'Done For You (DFY)')
    h3(doc, 'Full-Service Marketing \u2014 We Handle It')
    body(doc,
         'You focus on practicing law. We handle everything else \u2014 your local search '
         'presence, lead generation, website performance, advertising, and the automation '
         'that turns inquiries into booked consultations.')
    body(doc,
         'Built specifically for attorneys who know they need to invest in growth but '
         'don\u2019t have the time or interest to become a marketing expert alongside '
         'running a law firm.')
    h3(doc, 'Services include:')
    for svc in [
        'Local SEO and Google Maps visibility',
        'Lead generation systems',
        'Website design and conversion optimization',
        'Digital advertising (Google and Meta)',
        'Marketing automation and CRM setup',
        'The Triad\u2122 \u2014 our referral and revenue sustainability system',
    ]:
        bullet(doc, svc)
    body(doc, '\u25b8  Learn about Full-Service Marketing \u2192')

    gap(doc)
    tag(doc, 'Do With You (DWY)')
    h3(doc, 'Strategy and Coaching \u2014 Build It With Expert Guidance')
    body(doc,
         'You understand your practice better than anyone. Our coaching and strategy services '
         'give you the framework, the implementation support, and the accountability to build '
         'a marketing system you own \u2014 without guessing at what actually works for your '
         'practice type.')
    body(doc,
         'Built for attorneys who want the capability in-house, prefer to stay closely '
         'involved in their marketing decisions, or are in an earlier stage of building '
         'their practice.')
    h3(doc, 'Services include:')
    for svc in [
        'Strategic planning and marketing roadmap',
        'Content strategy and execution support',
        'Intake optimization and conversion coaching',
        'Videography and content production',
        'Ongoing implementation support',
    ]:
        bullet(doc, svc)
    body(doc, '\u25b8  Learn about Coaching and Strategy \u2192')

    gap(doc)
    tag(doc, 'Cohort-Based')
    h3(doc, 'Group Coaching and Training')
    body(doc,
         'A structured program for solo attorneys who want expert marketing guidance at a '
         'price point built for a practice that\u2019s still scaling. Small cohorts, live '
         'sessions, real implementation \u2014 not a course you\u2019ll never finish.')
    body(doc, '\u25b8  See upcoming cohorts \u2192')

    dev_note(doc,
             'Three cards in a row on desktop, stacked on mobile. Label tag in ASG blue as '
             'secondary identifier only \u2014 primary card headline is plain English. '
             'Cards function as a buyer decision tool, not a sales pitch.')

    # ── SECTION 5: THREE SYSTEMS ──────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 5 \u2014 The System: Process Proof')
    stage_label(doc, 'Belief: B4 Internal Alignment  |  Emotion: E3 \u2192 E4')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc,
         'The Three Systems Every Solo Practice Needs (And Most Are Missing at Least Two)')
    h2(doc, 'Subhead')
    body(doc,
         'This is why tactics without infrastructure fail \u2014 and what changes when all '
         'three are in place.')

    gap(doc)
    h3(doc, 'System 1: Visibility')
    body(doc, 'So the right clients can find you when they\u2019re ready')
    body(doc,
         'A family law attorney in Jeffersonville who can\u2019t be found on Google by '
         'someone searching \u201cdivorce attorney near me\u201d is invisible to their '
         'highest-intent potential clients. Visibility isn\u2019t just about ranking \u2014 '
         'it\u2019s about appearing in the right searches, with the right first impression, '
         'in the right geography.')
    body(doc,
         'Most attorneys either have no visibility strategy or they\u2019re paying for one '
         'that isn\u2019t calibrated to solo practice economics. We build this from the '
         'ground up \u2014 local SEO, Google Business Profile, website architecture \u2014 '
         'designed to produce qualified inquiries, not traffic reports.')

    gap(doc)
    h3(doc, 'System 2: The Triad\u2122 \u2014 Referrals and Revenue Stability')
    body(doc, 'So growth doesn\u2019t depend on luck or peak months')
    body(doc,
         'Most solo attorneys\u2019 referral system is their personality. They\u2019re '
         'likable, they do good work, and people sometimes send them business. This isn\u2019t '
         'a system \u2014 it\u2019s a variable you can\u2019t control.')
    body(doc,
         'The Triad\u2122 is a structured referral and revenue infrastructure that creates '
         'predictability in your caseload. It identifies and cultivates your referral network '
         'systematically, reactivates past clients at the right moments, and creates consistent '
         'revenue touch points that smooth out the feast-or-famine cycle. It\u2019s not '
         'complicated. Most attorneys have never had anyone build it for them.')

    gap(doc)
    h3(doc, 'System 3: Intake and Conversion')
    body(doc, 'So the inquiries you generate actually become retained clients')
    body(doc,
         'Getting a potential client to contact your firm is the beginning of the process '
         '\u2014 not the end. What happens in the first 90 minutes of an inquiry matters '
         'more than most attorneys realize. Response speed, what you say, how you say it, '
         'and what the follow-up looks like determines whether a qualified prospect becomes '
         'a client or calls the next attorney on the list.')
    body(doc,
         'We build and optimize the intake and conversion process so that the inquiries your '
         'visibility system generates don\u2019t leak out the other side.')

    dev_note(doc,
             'Strongest belief-shifting section on the page. Icon or illustration per system. '
             'Consider a subtle "before state" line under each heading in lighter font to create '
             'the contrast the Belief Track needs. Stack vertically on mobile.')

    # ── SECTION 6: RISK REVERSAL ──────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 6 \u2014 The Discovery Call: Risk Reversal')
    stage_label(doc, 'Belief: B5 Certainty  |  Emotion: E5 Relief')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc,
         'The Discovery Call. No Pitch. No Pressure. '
         '30 Minutes That Are 100% About Your Firm.')

    h2(doc, 'Body Copy')
    body(doc, 'Before we take on any new engagement, we have one conversation.')
    body(doc,
         'Not a sales call. Not a capabilities presentation where we tell you about ourselves '
         'for 25 minutes and ask for a proposal at the end.')
    body(doc,
         'A genuine 30-minute session focused entirely on your practice \u2014 where you are, '
         'what you\u2019re trying to build, what you\u2019ve tried before, and what\u2019s '
         'gotten in the way. You\u2019ll leave the call with a clear picture of what\u2019s '
         'actually holding your growth back, whether we\u2019re the right fit to help, and '
         'what the first 90 days would look like if we worked together.')
    body(doc,
         'If we\u2019re not the right fit, you\u2019ll know that too. We\u2019d rather tell '
         'you than waste both our time.')

    h2(doc, 'What the Call Covers')
    for item in [
        'Where your practice stands right now \u2014 visibility, referral flow, intake performance',
        'What the highest-leverage marketing investment is for your specific situation',
        'Which of our engagement models (if any) makes sense for where you are',
        'A clear, honest assessment of what results look like and on what timeline',
    ]:
        arrow(doc, item)

    h2(doc, 'This Call Is Right for You If:')
    for item in [
        'You\u2019re a solo practitioner or small firm attorney in family law, criminal defense, or personal injury',
        'You\u2019re generating some cases but want the growth to be more consistent and predictable',
        'You\u2019ve tried marketing before and want to understand why it didn\u2019t produce what you expected',
    ]:
        arrow(doc, item)

    h2(doc, 'This Call Is Not Right for You If:')
    for item in [
        'You\u2019re looking for a quick fix or want results in 30 days',
        'You\u2019re a large, multi-practice firm with an existing marketing team',
        'You want someone to run marketing without being willing to invest in building the infrastructure',
    ]:
        arrow(doc, item)

    gap(doc)
    cta(doc, 'Book Your Free 30-Minute Strategy Session')
    body(doc,
         'No contract required to have the conversation. No obligation to continue. '
         'Just 30 minutes, focused on your practice.')

    dev_note(doc,
             'Most important conversion section. Calm, not urgent. White background. Button '
             'centered, full-width on mobile. "Right for you / not right for you" bullets '
             'pre-qualify visitors and make the matching reader feel specifically invited. '
             'This intentional exclusion improves close rate.')

    # ── SECTION 7: FINAL CTA ──────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Section 7 \u2014 Final CTA')
    stage_label(doc, 'Emotion: E5 Relief complete')
    gap(doc)

    h2(doc, 'Section Header')
    body(doc, 'The Practice You Want to Run Is a System Problem. Systems Can Be Built.')

    h2(doc, 'Body Copy')
    body(doc,
         'Every attorney who is great at law and struggling with growth has the same core '
         'issue: the legal skill is there, the work ethic is there, and the right marketing '
         'infrastructure isn\u2019t. That\u2019s not a character flaw \u2014 it\u2019s a '
         'resource gap. It\u2019s exactly what Legal Growth Dynamics is here to close.')
    body(doc, 'The first step costs nothing except 30 minutes.')
    cta(doc,
        'Book Your Free Strategy Session',
        '[Phone number] \u2014 prefer to call? We answer.')

    # ── FOOTER COPY ───────────────────────────────────────────────────────────
    rule(doc)
    h1(doc, 'Footer Copy')
    gap(doc)

    h2(doc, 'Business Info')
    body(doc, 'Legal Growth Dynamics')
    body(doc, 'Law Firm Marketing for Solo and Small Firm Attorneys')
    body(doc, 'Jeffersonville, Indiana  |  [Phone]  |  [Email]')

    h2(doc, 'Navigation Links')
    body(doc,
         'Services (Full-Service)  |  Services (Coaching)  |  Coaching & Training  '
         '|  About  |  Blog  |  Schedule a Call')

    h2(doc, 'Legal Disclaimer')
    body(doc,
         'Results referenced are typical of what clients may experience when all recommended '
         'systems are fully implemented. Individual results vary based on practice area, '
         'geography, competition, and implementation. Legal Growth Dynamics is a marketing '
         'consulting firm and does not provide legal advice.')

    # ── AUDREY'S NOTES ────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, 'Audrey\u2019s Notes')
    gap(doc)

    h2(doc, 'Why These Structural Choices')
    body(doc,
         '1.  H1 opens at the Enemy Belief \u2014 not with a value claim but with a problem '
         'reframe. Breaking the pattern of every generic "we\u2019ll grow your firm" headline '
         'earns the scroll. "Three missing systems" creates intellectual curiosity.')
    body(doc,
         '2.  The Russell Lloyd story does three jobs in the absence of case studies: humanizes '
         'the founder, explains why this firm exists, and creates emotional contrast '
         '(exceptional attorney / struggling practice) that functions as a before/after. '
         'Replace with case studies as soon as 2\u20133 are available.')
    body(doc,
         '3.  "Not right for you if" bullets in Section 6 are intentional exclusion language. '
         'They increase conversion by making the included prospect feel specifically invited '
         'and reduce low-quality discovery calls.')
    body(doc,
         '4.  The Three Systems section is the most important belief-shifting section on the '
         'page. It gives the skeptical attorney (tried marketing, got burned) a new explanation '
         'for why past efforts failed \u2014 Internal Alignment stage. Write this before '
         'testimonials are added so testimonials validate a framework the reader already holds.')
    body(doc,
         '5.  DFY/DWY labels are preserved as secondary identifiers inside service cards only '
         '\u2014 not as primary section headers. Keeps the URL structure intact while removing '
         'jargon from the conversion path.')

    h2(doc, 'Competitor Differentiation Flagged')
    bullet(doc,
           'No competitor explicitly names a "three missing systems" framework. This is '
           'ownable positioning if developed into a lead magnet, blog series, and YouTube '
           'content strategy.')
    bullet(doc,
           '"We only work in three practice areas" is a stronger differentiator than it '
           'appears \u2014 most agencies refuse to limit themselves. Exclusivity stated '
           'confidently reads as expertise.')
    bullet(doc,
           'The Triad\u2122 name should be trademarked if it isn\u2019t already. It\u2019s '
           'the most distinctive IP on the service side.')

    h2(doc, 'Recommended A/B Test')
    body(doc,
         'Test Variant 1 ("Most Solo Attorneys Don\u2019t Have a Marketing Problem...") '
         'against Variant 3 ("We Only Work With Solo and Small Firm Attorneys in Three '
         'Practice Areas..."). Variant 3 may outperform for cold traffic because niche '
         'exclusivity functions as a trust shortcut for the right prospect.')

    h2(doc, 'Compliance Check')
    for item in [
        'No specific outcome guarantees in copy. All results language is conditional.  \u2714',
        'Disclaimer added to footer covering results variability.  \u2714',
        'No claims that could be construed as legal advice.  \u2714',
        '"Free" call offer is accurate \u2014 no hidden fee or required next step.  \u2714',
        'Russell Lloyd mention: confirm with Roger before publishing.',
    ]:
        bullet(doc, item)

    # ── QC SIGN-OFFS ──────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, 'QC Sign-Offs')
    gap(doc)
    for line in [
        'Strategic Alignment (Daniel Frost, CSO):    ________________________________   Date: ____________',
        'Messaging Integrity (Vivienne Carr, CMO):   ________________________________   Date: ____________',
        'Revenue Logic (Tanya Blackwood, CRO):       ________________________________   Date: ____________',
        'Delivery Standard (Preston Adler, COO):     ________________________________   Date: ____________',
    ]:
        body(doc, line)

    gap(doc)
    gap(doc)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _run(p_end,
         'Authority Systems Group\u2122  \u2014  Confidential. '
         'Prepared exclusively for Legal Growth Dynamics.',
         italic=True, size=9, color=CHARCOAL)

    doc.save(OUTPUT)
    print(f'DOCX created: {OUTPUT}')


if __name__ == '__main__':
    build()
