#!/usr/bin/env python3
"""
Authority Blueprint™ DOCX Builder
Client: Melissa Doss Law
Authority Systems Group™ — March 2026

Usage:
    python3 scripts/build_melissa_doss_blueprint.py

Output:
    client-onboarding/clients/melissa-doss-law/outputs/
    melissa-doss-law_authority-blueprint_20260302.docx
"""

import re
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────
# BRAND COLORS
# ─────────────────────────────────────────
ASG_BLUE       = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL   = RGBColor(0x58, 0x58, 0x5A)
ASG_BODY       = RGBColor(0x33, 0x33, 0x33)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_RGB = RGBColor(0xF5, 0xF5, 0xF5)
GRAY           = RGBColor(0x88, 0x88, 0x88)
MID_GRAY       = RGBColor(0xAA, 0xAA, 0xAA)
DARK_GRAY      = RGBColor(0x66, 0x66, 0x66)
DARK_BG        = RGBColor(0x1A, 0x1A, 0x1A)

HEX_BLUE       = "25AAE1"
HEX_CHARCOAL   = "58585A"
HEX_DARK_BG    = "1A1A1A"
HEX_BODY       = "333333"
HEX_LIGHT      = "F5F5F5"
HEX_WHITE      = "FFFFFF"
HEX_GRAY       = "888888"
HEX_CCCC       = "CCCCCC"

# ─────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────
BASE_DIR    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CLIENT_SLUG = "melissa-doss-law"
CLIENT_NAME = "Melissa Doss Law"
DOC_DATE    = "March 2026"
INPUT_MD    = os.path.join(
    BASE_DIR,
    "client-onboarding", "clients", CLIENT_SLUG,
    "melissa-doss-law_authority-blueprint_20260302.md"
)
OUTPUT_DOCX = os.path.join(
    BASE_DIR,
    "client-onboarding", "clients", CLIENT_SLUG,
    "outputs", "melissa-doss-law_authority-blueprint_20260302.docx"
)


# ─────────────────────────────────────────
# XML UTILITIES
# ─────────────────────────────────────────
def set_cell_shading(cell, fill_hex):
    """Apply background fill to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper().lstrip("#"))
    # Remove existing shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def set_para_shading(para, fill_hex):
    """Apply background shading to a paragraph (for callout boxes)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper().lstrip("#"))
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(shd)


def add_left_border(para, size_eighths=36, color_hex=HEX_BLUE):
    """Add a colored left border to a paragraph (callout effect)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement("w:left")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    str(size_eighths))
    bdr.set(qn("w:space"), "6")
    bdr.set(qn("w:color"), color_hex.upper().lstrip("#"))
    pBdr.append(bdr)
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(pBdr)


def add_bottom_rule(para, size_eighths=16, color_hex=HEX_BLUE):
    """Add a bottom border to simulate a horizontal rule."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement("w:bottom")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    str(size_eighths))
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), color_hex.upper().lstrip("#"))
    pBdr.append(bdr)
    pPr.append(pBdr)


def set_spacing(para, before_pt=0, after_pt=8, line_multiple=None):
    """Set paragraph spacing."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"),  str(int(after_pt  * 20)))
    if line_multiple is not None:
        sp.set(qn("w:line"),     str(int(line_multiple * 240)))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def set_indent(para, left_inches=0.25):
    """Set paragraph left indent."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(left_inches * 1440)))
    pPr.append(ind)


def set_row_height(row, height_inches, exact=True):
    """Set a table row to an exact or minimum height."""
    tr    = row._tr
    trPr  = tr.get_or_add_trPr()
    trH   = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_inches * 1440)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBdr)


def style_data_table(table):
    """Apply ASG brand styling to a data table."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Light border on all cells
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), HEX_CCCC)
        tblBdr.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBdr)

    # Style rows
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                # Header row: blue background, white text
                set_cell_shading(cell, HEX_BLUE)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = WHITE
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
            else:
                # Alternating rows
                fill = HEX_WHITE if i % 2 == 1 else HEX_LIGHT
                set_cell_shading(cell, fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"
                        run.font.color.rgb = ASG_BODY


HEX_WHITE = "FFFFFF"  # re-declare after function defs use it


# ─────────────────────────────────────────
# INLINE MARKDOWN PARSER
# ─────────────────────────────────────────
def parse_inline(para, text, size=11, color=None, font="Georgia"):
    """Parse **bold** and *italic* inline markdown into runs."""
    if color is None:
        color = ASG_BODY
    # Tokenise: **bold**, *italic*, or plain
    tokens = re.split(r"(\*\*[^*]+?\*\*|\*[^*]+?\*)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            r = para.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = para.add_run(tok[1:-1])
            r.italic = True
        else:
            r = para.add_run(tok)
        r.font.size  = Pt(size)
        r.font.color.rgb = color
        r.font.name  = font


# ─────────────────────────────────────────
# HEADER / FOOTER
# ─────────────────────────────────────────
def apply_header_footer(section, client_name):
    """Set header and footer on a section."""
    section.different_first_page_header_footer = False

    # ── HEADER ──────────────────────────
    hdr  = section.header
    hpara = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hpara.clear()
    hpara.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Left: "ASG™" placeholder for logo
    left_run = hpara.add_run("Authority Systems Group™")
    left_run.font.size  = Pt(9)
    left_run.font.color.rgb = ASG_CHARCOAL
    left_run.font.name  = "Calibri"
    left_run.font.bold  = True

    tab_run = hpara.add_run("\t")
    tab_run.font.size = Pt(9)

    right_run = hpara.add_run(f"Authority Blueprint™  —  {client_name}")
    right_run.font.size  = Pt(9)
    right_run.font.color.rgb = ASG_CHARCOAL
    right_run.font.name  = "Calibri"

    # Bottom border on header
    add_bottom_rule(hpara, size_eighths=4, color_hex=HEX_CCCC)
    set_spacing(hpara, before_pt=0, after_pt=4)

    # ── FOOTER ──────────────────────────
    ftr   = section.footer
    fpara = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER

    conf_run = fpara.add_run(
        f"Authority Systems Group™  —  Confidential. Prepared exclusively for {client_name}."
    )
    conf_run.font.size  = Pt(9)
    conf_run.font.color.rgb = GRAY
    conf_run.font.name  = "Calibri"

    # Page number on right via tab + field
    tab2 = fpara.add_run("\t")
    tab2.font.size = Pt(9)

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fpara.runs[-1]._r.append(fld)

    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    instrR = OxmlElement("w:r")
    instrR.append(instr)
    fpara._p.append(instrR)

    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    endR = OxmlElement("w:r")
    endR.append(fld2)
    fpara._p.append(endR)

    # Top border on footer
    pPr  = fpara._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement("w:top")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    "4")
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), HEX_CCCC)
    pBdr.append(bdr)
    pPr.append(pBdr)
    set_spacing(fpara, before_pt=4, after_pt=0)


# ─────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────
def build_cover_page(doc, client_name, doc_date):
    """Build a dark-background cover page using a full-page table."""
    # Zero out cover section margins
    cover_sec = doc.sections[0]
    cover_sec.top_margin    = Inches(0)
    cover_sec.bottom_margin = Inches(0)
    cover_sec.left_margin   = Inches(0)
    cover_sec.right_margin  = Inches(0)

    # Suppress header/footer on cover
    cover_sec.header.is_linked_to_previous = False
    cover_sec.footer.is_linked_to_previous = False
    for p in cover_sec.header.paragraphs:
        p.clear()
    for p in cover_sec.footer.paragraphs:
        p.clear()

    # Full-page 2-column table: blue stripe (0.15") + dark content (8.35")
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)

    # Column widths
    tbl.columns[0].width = Inches(0.15)
    tbl.columns[1].width = Inches(8.35)

    # Set row to full page height (11")
    set_row_height(tbl.rows[0], 11.0)

    # Left stripe: ASG Blue
    left_cell = tbl.cell(0, 0)
    set_cell_shading(left_cell, HEX_BLUE)
    for p in left_cell.paragraphs:
        p.clear()

    # Right cell: dark background
    right_cell = tbl.cell(0, 1)
    set_cell_shading(right_cell, HEX_DARK_BG)

    # Clear default paragraph
    for p in list(right_cell.paragraphs):
        p._p.getparent().remove(p._p)

    def cover_para(text, size, bold=False, italic=False,
                   color_rgb=None, align=WD_ALIGN_PARAGRAPH.CENTER,
                   before=0, after=12, font="Calibri"):
        p = right_cell.add_paragraph()
        p.alignment = align
        set_spacing(p, before_pt=before, after_pt=after)
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size  = Pt(size)
        r.font.name  = font
        r.font.color.rgb = color_rgb if color_rgb else WHITE
        return p

    # Top spacer (~2.5 inches worth of empty space)
    sp = right_cell.add_paragraph()
    set_spacing(sp, before_pt=145, after_pt=0)

    # ── TITLE ────────────────────────────────
    cover_para("YOUR AUTHORITY BLUEPRINT™",
               size=34, bold=True, before=0, after=4)

    # Blue accent rule (em-dash line)
    line_p = right_cell.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(line_p, before_pt=0, after_pt=20)
    lr = line_p.add_run("─" * 38)
    lr.font.size  = Pt(10)
    lr.font.color.rgb = ASG_BLUE
    lr.font.name  = "Calibri"

    # ── SUBTITLE ─────────────────────────────
    cover_para("A Strategic Growth System Built Exclusively for:",
               size=14, color_rgb=ASG_BLUE, before=0, after=8)

    # ── CLIENT NAME ──────────────────────────
    cover_para(client_name,
               size=26, bold=True, before=0, after=28)

    # ── ATTRIBUTION ──────────────────────────
    cover_para("Directed by Roger  |  Authority Systems Group™",
               size=12, italic=True, color_rgb=MID_GRAY, before=0, after=8)

    # ── DATE ─────────────────────────────────
    cover_para(doc_date,
               size=11, color_rgb=GRAY, before=0, after=200)

    # ── CONFIDENTIAL NOTICE ───────────────────
    cover_para(f"CONFIDENTIAL  —  PREPARED EXCLUSIVELY FOR {client_name.upper()}",
               size=9, color_rgb=DARK_GRAY, before=0, after=24)


# ─────────────────────────────────────────
# CONFIDENTIALITY PAGE
# ─────────────────────────────────────────
def build_confidentiality_page(doc):
    """Add the standard confidentiality notice interior page."""
    doc.add_page_break()

    p = doc.add_paragraph()
    set_spacing(p, before_pt=220, after_pt=0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("CONFIDENTIALITY NOTICE")
    r.bold = True
    r.font.size  = Pt(14)
    r.font.color.rgb = ASG_CHARCOAL
    r.font.name  = "Calibri"
    set_spacing(p2, before_pt=0, after_pt=18)

    notice = (
        "This document was prepared exclusively for Melissa Doss Law by "
        "Authority Systems Group™. It contains proprietary strategic analysis "
        "and recommendations developed specifically for this firm. "
        "It may not be shared, reproduced, or distributed without written "
        "consent from Authority Systems Group™.\n\n"
        "All competitive intelligence, market analysis, and strategic "
        "recommendations contained herein are confidential work product "
        "and constitute trade secrets of Authority Systems Group™."
    )
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(notice)
    r3.font.size  = Pt(11)
    r3.font.color.rgb = ASG_BODY
    r3.font.name  = "Georgia"
    set_spacing(p3, before_pt=0, after_pt=0, line_multiple=1.5)
    set_indent(p3, left_inches=0.75)
    pPr = p3._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        ind.set(qn("w:right"), str(int(0.75 * 1440)))


# ─────────────────────────────────────────
# SECTION HEADING HELPERS
# ─────────────────────────────────────────
def add_h1(doc, text):
    """Major section heading — ASG Blue, 24pt."""
    p = doc.add_paragraph()
    add_bottom_rule(p, size_eighths=16, color_hex=HEX_BLUE)
    set_spacing(p, before_pt=0, after_pt=0)

    p2 = doc.add_paragraph()
    r  = p2.add_run(text)
    r.bold = True
    r.font.size  = Pt(22)
    r.font.color.rgb = ASG_BLUE
    r.font.name  = "Calibri"
    set_spacing(p2, before_pt=6, after_pt=10)
    return p2


def add_h2(doc, text):
    """Subsection heading — ASG Charcoal, 16pt."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(16)
    r.font.color.rgb = ASG_CHARCOAL
    r.font.name  = "Calibri"
    set_spacing(p, before_pt=14, after_pt=6)
    return p


def add_h3(doc, text):
    """Sub-subsection heading — Charcoal, 13pt."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(13)
    r.font.color.rgb = ASG_CHARCOAL
    r.font.name  = "Calibri"
    set_spacing(p, before_pt=10, after_pt=4)
    return p


def add_body(doc, text):
    """Standard body paragraph."""
    p = doc.add_paragraph()
    parse_inline(p, text, size=11, color=ASG_BODY, font="Georgia")
    set_spacing(p, before_pt=0, after_pt=6, line_multiple=1.4)
    return p


def add_bullet(doc, text):
    """Bullet list item."""
    p = doc.add_paragraph(style="List Bullet")
    # Strip bullet prefix (- or *), then any checkbox marker ([ ] / [x])
    clean = re.sub(r"^[-*]\s+", "", text.strip())
    clean = re.sub(r"^\[[ xX✓]\]\s*", "", clean)
    parse_inline(p, clean, size=11, color=ASG_BODY, font="Georgia")
    set_spacing(p, before_pt=0, after_pt=3, line_multiple=1.3)
    return p


def add_callout(doc, text):
    """Callout / blockquote box with left blue border and light gray background."""
    p = doc.add_paragraph()
    parse_inline(p, text, size=11, color=ASG_CHARCOAL, font="Calibri")
    p.runs[0].italic = True if p.runs else None
    add_left_border(p, size_eighths=36, color_hex=HEX_BLUE)
    set_para_shading(p, HEX_LIGHT)
    set_indent(p, left_inches=0.25)
    set_spacing(p, before_pt=10, after_pt=10, line_multiple=1.4)
    return p


def add_hr(doc):
    """Horizontal rule — blue bottom border."""
    p = doc.add_paragraph()
    add_bottom_rule(p, size_eighths=12, color_hex=HEX_BLUE)
    set_spacing(p, before_pt=10, after_pt=10)
    return p


def add_signoff(doc, text):
    """Sign-off italic line."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size  = Pt(11)
    r.font.color.rgb = ASG_CHARCOAL
    r.font.name  = "Calibri"
    set_spacing(p, before_pt=16, after_pt=4)
    set_indent(p, left_inches=0.5)
    return p


# ─────────────────────────────────────────
# MARKDOWN TABLE PARSER
# ─────────────────────────────────────────
def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    stripped = line.strip()
    return stripped.startswith("|") and re.fullmatch(r"[\|\-\:\s]+", stripped)


def parse_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_md_table(doc, table_lines):
    """Convert a block of markdown table lines into a styled Word table."""
    # Extract header and data rows (skip separator rows)
    rows = []
    for ln in table_lines:
        if is_separator_row(ln):
            continue
        rows.append(parse_table_row(ln))

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [""] * (num_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline(p, cell_text, size=10, color=ASG_BODY, font="Calibri")
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = WHITE
                set_cell_shading(cell, HEX_BLUE)
            else:
                fill = HEX_WHITE if r_idx % 2 == 1 else HEX_LIGHT
                set_cell_shading(cell, fill)

    # Apply table border style
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), HEX_CCCC)
        tblBdr.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBdr)

    # Auto-fit
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(tblW)

    set_spacing(doc.add_paragraph(), before_pt=6, after_pt=6)


# ─────────────────────────────────────────
# QC SIGN-OFF PAGE
# ─────────────────────────────────────────
def build_qc_page(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("QC Sign-Offs — Internal")
    r.bold = True
    r.font.size  = Pt(16)
    r.font.color.rgb = ASG_CHARCOAL
    r.font.name  = "Calibri"
    set_spacing(p, before_pt=18, after_pt=16)
    add_bottom_rule(p, size_eighths=8, color_hex=HEX_BLUE)

    gates = [
        ("Strategic Alignment",   "Daniel Frost, CSO"),
        ("Messaging Integrity",   "Vivienne Carr, CMO"),
        ("Revenue Logic",         "Tanya Blackwood, CRO"),
        ("Delivery Standard",     "Preston Adler, COO"),
    ]
    for label, name in gates:
        p2 = doc.add_paragraph()
        r1 = p2.add_run(f"{label}  ({name}):  ")
        r1.bold = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = ASG_BODY
        r1.font.name  = "Calibri"
        r2 = p2.add_run("_" * 32 + "   Date: " + "_" * 14)
        r2.font.size  = Pt(11)
        r2.font.color.rgb = ASG_BODY
        r2.font.name  = "Calibri"
        set_spacing(p2, before_pt=10, after_pt=10)

    # Footer attribution
    spacer = doc.add_paragraph()
    set_spacing(spacer, before_pt=40, after_pt=0)
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = pa.add_run(
        "Authority Systems Group™  —  Confidential. "
        "Prepared exclusively for Melissa Doss Law.\n"
        "7289 Burlington Pike, Florence, KY 41042  |  "
        "(859) 322-5550  |  MelissaDossLaw.com"
    )
    ra.font.size  = Pt(9)
    ra.font.color.rgb = GRAY
    ra.font.name  = "Calibri"
    ra.italic = True
    set_spacing(pa, before_pt=0, after_pt=0, line_multiple=1.5)


# ─────────────────────────────────────────
# MAIN MARKDOWN RENDERER
# ─────────────────────────────────────────
def render_markdown(doc, md_text):
    """
    Parse the Blueprint markdown and render into the document.
    Handles: # headings, **bold**, *italic*, > blockquotes,
    | tables |, - bullets, --- hr, and regular paragraphs.
    H2 sections always start on a new page.
    """
    lines = md_text.splitlines()
    i     = 0
    in_callout_block = False
    callout_lines    = []

    def flush_callout():
        nonlocal in_callout_block, callout_lines
        if callout_lines:
            text = " ".join(callout_lines).strip()
            if text:
                add_callout(doc, text)
        callout_lines    = []
        in_callout_block = False

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        # ── BLANK LINE ────────────────────────
        if not line:
            if in_callout_block:
                flush_callout()
            i += 1
            continue

        # ── MARKDOWN TABLE ────────────────────
        if is_table_row(raw):
            if in_callout_block:
                flush_callout()
            table_block = []
            while i < len(lines) and is_table_row(lines[i]):
                table_block.append(lines[i])
                i += 1
            add_md_table(doc, table_block)
            continue

        # ── BLOCKQUOTE / CALLOUT ──────────────
        if line.startswith(">"):
            content = line.lstrip("> ").strip()
            callout_lines.append(content)
            in_callout_block = True
            i += 1
            continue
        elif in_callout_block:
            flush_callout()

        # ── HORIZONTAL RULE ──────────────────
        if line in ("---", "***", "___"):
            add_hr(doc)
            i += 1
            continue

        # ── H1 (document title — skip rendering as big heading; already on cover) ──
        if line.startswith("# ") and not line.startswith("## "):
            # Skip the document title line (it's on the cover page)
            i += 1
            continue

        # ── H2 (major section) ────────────────
        if line.startswith("## "):
            text = line[3:].strip()
            # Page break before each major section
            doc.add_page_break()
            add_h1(doc, text)
            i += 1
            continue

        # ── H3 ────────────────────────────────
        if line.startswith("### "):
            text = line[4:].strip()
            add_h2(doc, text)
            i += 1
            continue

        # ── H4 ────────────────────────────────
        if line.startswith("#### "):
            text = line[5:].strip()
            add_h3(doc, text)
            i += 1
            continue

        # ── BULLET LIST ───────────────────────
        if line.startswith("- ") or line.startswith("* "):
            add_bullet(doc, line)
            i += 1
            continue

        # ── SIGN-OFF LINE (— Name, Title) ─────
        if line.startswith("— ") or line.startswith("*— ") or line.startswith("*\""):
            # Italic sign-off lines
            cleaned = line.strip("*").strip()
            add_signoff(doc, cleaned)
            i += 1
            continue

        # ── META / FRONT MATTER (skip lines like **Prepared by...**) ──
        if line.startswith("**Prepared by") or line.startswith("**Confidential") or \
           line.startswith("**Date:") or line.startswith("*Authority Systems"):
            i += 1
            continue

        # ── PLAIN PARAGRAPH ───────────────────
        if line:
            add_body(doc, line)

        i += 1

    if in_callout_block:
        flush_callout()


# ─────────────────────────────────────────
# SET DOCUMENT PROPERTIES
# ─────────────────────────────────────────
def set_doc_properties(doc):
    core = doc.core_properties
    core.title   = f"Authority Blueprint™ — {CLIENT_NAME}"
    core.subject = "Strategic Growth System"
    core.author  = "Authority Systems Group™"
    core.company = "Authority Systems Group™"
    core.keywords = f"authority blueprint, strategic plan, {CLIENT_SLUG}"
    core.comments = "Prepared under the direction of Roger, Director — Authority Systems Group™"


# ─────────────────────────────────────────
# MAIN BUILD FUNCTION
# ─────────────────────────────────────────
def build():
    print(f"Reading: {INPUT_MD}")
    with open(INPUT_MD, "r", encoding="utf-8") as f:
        md_text = f.read()

    print("Building DOCX...")
    doc = Document()

    # ── 1. COVER PAGE (Section 1 — zero margins) ──────────
    build_cover_page(doc, CLIENT_NAME, DOC_DATE)

    # ── 2. BODY SECTION BREAK ─────────────────────────────
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin    = Inches(1.0)
    body_section.bottom_margin = Inches(1.0)
    body_section.left_margin   = Inches(1.25)
    body_section.right_margin  = Inches(1.0)
    body_section.header_distance = Inches(0.5)
    body_section.footer_distance = Inches(0.5)

    # Apply header/footer to body section
    apply_header_footer(body_section, CLIENT_NAME)

    # ── 3. CONFIDENTIALITY PAGE ───────────────────────────
    build_confidentiality_page(doc)

    # ── 4. BODY CONTENT (markdown) ────────────────────────
    render_markdown(doc, md_text)

    # ── 5. QC SIGN-OFF PAGE ───────────────────────────────
    build_qc_page(doc)

    # ── 6. DOCUMENT PROPERTIES ───────────────────────────
    set_doc_properties(doc)

    # ── 7. SAVE ───────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"✓  DOCX saved: {OUTPUT_DOCX}")
    print(f"   File size:  {os.path.getsize(OUTPUT_DOCX):,} bytes")


if __name__ == "__main__":
    build()
