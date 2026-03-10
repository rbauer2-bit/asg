#!/usr/bin/env python3
"""
Authority Systems Group™
Roger Bauer — Portfolio Strategy Brief
March 2026 | Internal Strategic Document
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT     = RGBColor(0x1A, 0x25, 0x33)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
LIGHT_BG      = 'F5F5F5'
TINT_BG       = 'EAF6FC'
DIVIDER_HEX   = 'DDDDDD'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'F39C12'
RED_HEX       = 'C0392B'
SLATE_HEX     = '2C3E50'

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_portfolio-strategy-brief_20260308.docx')


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE('w:br')
    br.set(_qn('w:type'), 'page')
    return br


def add_pb(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)
    return p


def add_section_header(doc, text, bg_hex=ASG_DARK_HEX, text_color=WHITE, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    set_para_bg(p, bg_hex)
    p.paragraph_format.left_indent = Inches(0)
    run = p.add_run(f'  {text}  ')
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = text_color
    run.font.name = 'Calibri'
    return p


def add_sub_header(doc, text, color=ASG_BLUE, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p


def add_body(doc, text, size=10.5, color=BODY_COLOR, bold=False, italic=False, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=10.5, indent=0.2):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_COLOR
    run.font.name = 'Calibri'
    return p


def add_divider(doc, color_hex=DIVIDER_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Header / Footer ────────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 2, Inches(6.5))
    remove_table_borders(htable)
    htable.columns[0].width = Inches(2.5)
    htable.columns[1].width = Inches(4.0)

    lc = htable.cell(0, 0)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        lrun = lp.add_run()
        lrun.add_picture(LOGO_PATH, width=Inches(1.8))

    rc = htable.cell(0, 1)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rrun = rp.add_run('PORTFOLIO STRATEGY BRIEF  |  MARCH 2026')
    rrun.font.size = Pt(8)
    rrun.font.color.rgb = FOOTER_COLOR
    rrun.font.name = 'Calibri'
    rrun.bold = True

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    ftable = footer.add_table(1, 2, Inches(6.5))
    remove_table_borders(ftable)
    ftable.columns[0].width = Inches(4.5)
    ftable.columns[1].width = Inches(2.0)

    flc = ftable.cell(0, 0)
    flp = flc.paragraphs[0]
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flrun = flp.add_run('Authority Systems Group™  —  Internal Strategic Document  —  Confidential')
    flrun.font.size = Pt(8)
    flrun.font.color.rgb = FOOTER_COLOR
    flrun.font.name = 'Calibri'

    frc = ftable.cell(0, 1)
    frp = frc.paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


def add_page_number(para):
    run = para.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = FOOTER_COLOR
    run.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Dark background block
    cover_bg = doc.add_paragraph()
    cover_bg.paragraph_format.space_before = Pt(0)
    cover_bg.paragraph_format.space_after = Pt(0)
    set_para_bg(cover_bg, ASG_DARK_HEX)
    cover_bg.add_run('  ')

    # Logo row
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_before = Pt(40)
    logo_p.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        logo_run = logo_p.add_run()
        logo_run.add_picture(LOGO_PATH, width=Inches(2.4))

    add_spacer(doc, 24)

    # Eyebrow
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    eyebrow.paragraph_format.space_before = Pt(0)
    eyebrow.paragraph_format.space_after = Pt(6)
    er = eyebrow.add_run('INTERNAL STRATEGIC DOCUMENT')
    er.font.size = Pt(9)
    er.font.color.rgb = ASG_BLUE
    er.bold = True
    er.font.name = 'Calibri'

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)
    tr = title_p.add_run('Portfolio Strategy Brief')
    tr.font.size = Pt(30)
    tr.bold = True
    tr.font.color.rgb = DARK_TEXT
    tr.font.name = 'Calibri'

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(4)
    sr = sub_p.add_run('Business Inventory · Revenue Path · Short-Term Focus Plan')
    sr.font.size = Pt(14)
    sr.font.color.rgb = ASG_CHARCOAL
    sr.font.name = 'Calibri'

    add_divider(doc, ASG_BLUE_HEX)

    # Meta row
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(8)
    meta_p.paragraph_format.space_after = Pt(4)
    mr = meta_p.add_run('Director: Roger Bauer     |     Date: March 8, 2026     |     Classification: Confidential')
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = ASG_CHARCOAL
    mr.font.name = 'Calibri'

    add_spacer(doc, 20)

    # Framing note
    framing = doc.add_paragraph()
    framing.paragraph_format.space_before = Pt(0)
    framing.paragraph_format.space_after = Pt(4)
    set_para_bg(framing, TINT_BG)
    framing.paragraph_format.left_indent = Inches(0.15)
    framing.paragraph_format.right_indent = Inches(0.15)
    fr = framing.add_run(
        '  This brief inventories Roger\'s active and dormant business portfolio, clarifies the '
        'strategic rationale behind each, and provides team recommendations for the fastest path '
        'to monetization. It is intended as a living reference document — to be updated as '
        'priorities shift and proof points accumulate.  '
    )
    fr.font.size = Pt(10)
    fr.font.color.rgb = BODY_COLOR
    fr.font.name = 'Calibri'
    fr.italic = True

    add_pb(doc)


# ── Section 1: Portfolio Inventory ────────────────────────────────────────────

def build_portfolio_inventory(doc):
    add_section_header(doc, '01 — BUSINESS PORTFOLIO INVENTORY', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'Six businesses currently exist in Roger\'s portfolio, ranging from active client '
        'engagements to dormant assets with deferred potential. Below is the full inventory '
        'with focus, status, and strategic notes from the team.',
        space_after=12
    )

    companies = [
        {
            'name': 'Dental Marketing Authority (DMA)',
            'status': 'ACTIVE',
            'status_hex': GREEN_HEX,
            'tier': 'Tier 1 — Primary Focus',
            'focus': (
                'Implant dentistry (primary), orthodontists offering Invisalign (secondary), '
                'general dentists offering Invisalign (tertiary).'
            ),
            'rationale': (
                'Implants are the highest-ticket, highest-satisfaction procedure in dentistry. '
                'They allow a dentist to build the lifestyle practice they originally envisioned — '
                'fewer patients, higher revenue, technical mastery. The Invisalign add-on creates '
                'a second revenue channel while enabling a giveaway/raffle strategy that attracts '
                'a mature, less price-sensitive patient profile. This is not a commodity market — '
                'it is an authority market. DMA is built for it.'
            ),
            'assets': 'Audience persona (implant dentist), 30-day blog pack (April 2026), DMA brand system, article DOCX generator.',
            'next': 'First paying client. Case study. Referral system.',
        },
        {
            'name': 'Legal Growth Dynamics (LGD)',
            'status': 'ACTIVE',
            'status_hex': GREEN_HEX,
            'tier': 'Tier 1 — Primary Focus',
            'focus': 'Family law attorneys seeking more high-net-worth cases.',
            'rationale': (
                'Family law is a recurring, high-stakes practice area where trust and authority '
                'drive client selection. HNW clients do not choose attorneys from ads — they '
                'choose based on reputation, visibility, and referred confidence. The authority '
                'positioning model fits perfectly. LGD has established infrastructure, client '
                'context, and initial deliverables in place.'
            ),
            'assets': 'Multiple delivered documents including Authority Blueprint, Content Authority Strategy, Dream 100, LinkedIn 30-Day Pack, April/May marketing calendars, YouTube scripts.',
            'next': 'Close first retainer-level client. Build case study. Activate referral network.',
        },
        {
            'name': 'x2 Sales',
            'status': 'PLANNING',
            'status_hex': AMBER_HEX,
            'tier': 'Tier 2 — Deferred',
            'focus': 'B2B sales consulting. Technology companies and outside-sales-oriented businesses.',
            'rationale': (
                'B2B sales consulting is a strong authority play — sales leaders, VPs, and founders '
                'who need to build or fix a revenue function will pay premium prices for a credible '
                'expert. The challenge: the sales cycle for B2B consulting itself is longer and '
                'relationship-driven. This is not a fast-cash vehicle — it is a strategic build. '
                'Defer until DMA and LGD have proven the ASG model.'
            ),
            'assets': 'Brand name. Domain (assumed). No active client infrastructure.',
            'next': 'Activate after first case studies from DMA + LGD are secured.',
        },
        {
            'name': 'Local Consultant Secrets (LCS)',
            'status': 'PLANNING',
            'status_hex': AMBER_HEX,
            'tier': 'Tier 2 — Deferred',
            'focus': (
                'Local consultants offering white-label GHL SaaS or serving local businesses. '
                'Template, snapshot, industry insight deliverables plus a training program component.'
            ),
            'rationale': (
                'This model has structural advantages: productized delivery, recurring revenue '
                'potential, and a community/cohort model that scales without adding proportional '
                'service labor. However, the audience (local consultants) is fragmented, price-'
                'sensitive in some segments, and requires a strong lead magnet funnel to convert. '
                'The training + template model works — but needs dedicated positioning and a '
                'functional GHL demonstration environment. Viable at Phase 2.'
            ),
            'assets': 'Brand concept. GHL platform access (Legal Growth Pilot). Training program concept.',
            'next': 'Define core curriculum. Build pilot cohort offer. Activate post-DMA/LGD.',
        },
        {
            'name': 'SMB Consulting',
            'status': 'DORMANT',
            'status_hex': RED_HEX,
            'tier': 'Tier 3 — No Action Required',
            'focus': 'Local service businesses. Original company, dormant for approximately 10 years.',
            'rationale': (
                'Roger\'s own assessment is accurate: ASG serves this market better, with fewer '
                'constraints and higher price points. Small business clients at the local service '
                'level rarely have the budget for the engagement style that generates meaningful '
                'ROI for either party. The ASG model — authority positioning for businesses that '
                'can command premium prices — is a direct upgrade to the SMB Consulting model. '
                'No action recommended.'
            ),
            'assets': 'Brand name. Domain. Historical client experience (valuable as proof of concept).',
            'next': 'None. Let it remain dormant. Do not actively wind down — no action needed.',
        },
        {
            'name': 'Legal Growth Pilot',
            'status': 'INFRASTRUCTURE',
            'status_hex': AMBER_HEX,
            'tier': 'Tier 2 — Support Role',
            'focus': (
                'White-label GHL platform for running automations and CRM installations for clients. '
                'Originally conceived as a legal-industry CRM; now recognized as broader infrastructure.'
            ),
            'rationale': (
                'This is not a client-facing business in its current form — it is the operational '
                'backbone. GHL is the delivery layer for ASG\'s digital systems work (Marcus Chen\'s '
                'domain). The question is not whether to use it — it will be used for every client '
                'who needs CRM automation. The question is whether to rebrand it as an ASG-aligned '
                'product (e.g., "Authority OS" or similar). That rebrand is a Phase 2 consideration. '
                'For now: operate it as infrastructure, not a product.'
            ),
            'assets': 'Active GHL account. Existing automation and CRM build capacity.',
            'next': 'Use for DMA and LGD client CRM installations. Rebrand consideration deferred.',
        },
    ]

    for co in companies:
        add_spacer(doc, 6)
        # Company name bar
        name_table = doc.add_table(1, 2)
        remove_table_borders(name_table)
        name_table.columns[0].width = Inches(4.8)
        name_table.columns[1].width = Inches(1.7)

        nc = name_table.cell(0, 0)
        set_cell_bg(nc, SLATE_HEX)
        np_ = nc.paragraphs[0]
        np_.paragraph_format.space_before = Pt(4)
        np_.paragraph_format.space_after = Pt(4)
        nr = np_.add_run(f'  {co["name"]}')
        nr.bold = True
        nr.font.size = Pt(11)
        nr.font.color.rgb = WHITE
        nr.font.name = 'Calibri'

        sc = name_table.cell(0, 1)
        set_cell_bg(sc, co['status_hex'])
        sp_ = sc.paragraphs[0]
        sp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp_.paragraph_format.space_before = Pt(4)
        sp_.paragraph_format.space_after = Pt(4)
        sr = sp_.add_run(co['status'])
        sr.bold = True
        sr.font.size = Pt(9)
        sr.font.color.rgb = WHITE
        sr.font.name = 'Calibri'

        # Tier label
        tier_p = doc.add_paragraph()
        tier_p.paragraph_format.space_before = Pt(2)
        tier_p.paragraph_format.space_after = Pt(4)
        set_para_bg(tier_p, LIGHT_BG)
        tier_run = tier_p.add_run(f'  {co["tier"]}')
        tier_run.bold = True
        tier_run.font.size = Pt(9)
        tier_run.font.color.rgb = ASG_CHARCOAL
        tier_run.font.name = 'Calibri'

        # Content table
        content_table = doc.add_table(4, 2)
        remove_table_borders(content_table)
        content_table.columns[0].width = Inches(1.3)
        content_table.columns[1].width = Inches(5.2)

        rows_data = [
            ('Focus', co['focus']),
            ('Rationale', co['rationale']),
            ('Assets In Place', co['assets']),
            ('Next Action', co['next']),
        ]

        for i, (label, text) in enumerate(rows_data):
            row = content_table.rows[i]
            bg = TINT_BG if i % 2 == 0 else 'FFFFFF'

            lc = row.cells[0]
            set_cell_bg(lc, bg)
            lp = lc.paragraphs[0]
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after = Pt(4)
            lr = lp.add_run(f'  {label}')
            lr.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = ASG_CHARCOAL
            lr.font.name = 'Calibri'

            tc_ = row.cells[1]
            set_cell_bg(tc_, bg)
            tp = tc_.paragraphs[0]
            tp.paragraph_format.space_before = Pt(4)
            tp.paragraph_format.space_after = Pt(4)
            tr_ = tp.add_run(text)
            tr_.font.size = Pt(9.5)
            tr_.font.color.rgb = BODY_COLOR
            tr_.font.name = 'Calibri'

        add_spacer(doc, 8)

    add_pb(doc)


# ── Section 2: Strategic Assessment ───────────────────────────────────────────

def build_strategic_assessment(doc):
    add_section_header(doc, '02 — STRATEGIC ASSESSMENT  |  Daniel Frost, CSO', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'From a market positioning and competitive moat perspective, the portfolio concentration '
        'Roger is describing is exactly right. Here is how I see it:',
        space_after=10
    )

    add_sub_header(doc, 'Why Two Niches First — Not Six')
    add_body(doc,
        'Every authority business follows the same pattern: depth precedes width. The consultants '
        'who try to serve everyone serve no one at a price point worth serving. Roger\'s instinct '
        'to go deep on implant dentistry and family law before activating anything else is the '
        'correct strategic call — and here is the specific reason it matters:',
        space_after=6
    )
    add_bullet(doc, 'Case studies are the only real sales asset in a premium service business. No case study = no social proof = longer sales cycle = lower close rate = lower price ceiling.')
    add_bullet(doc, 'Implant dentists and family law attorneys both operate in authority-driven buying environments. The client is making a high-trust, high-stakes decision. They need to see proof that you have done this for someone like them.')
    add_bullet(doc, 'Two niches is not spreading thin — it is building parallel proof tracks. If one underperforms in the short term, the other carries momentum.')
    add_spacer(doc, 6)

    add_sub_header(doc, 'The Authority Pricing Thesis — Confirmed')
    add_body(doc,
        'Roger\'s stated reasoning — authority positioning enables higher price points because the '
        'client can charge more for the same outcomes — is not just intuition. It is documented '
        'market behavior. Here is the chain:',
        space_after=6
    )
    add_bullet(doc, 'Authority positioning increases the client\'s perceived value in their own market.')
    add_bullet(doc, 'Higher perceived value = higher prices the client can charge.')
    add_bullet(doc, 'Higher-priced clients attract better-fit patients/customers = lower churn, fewer complaints, more referrals.')
    add_bullet(doc, 'A client who sees revenue growth attributes it to ASG. That is the case study. That is the renewal. That is the referral.')
    add_bullet(doc, 'This flywheel is not accessible through generic SMB consulting at commodity price points.')
    add_spacer(doc, 6)

    add_sub_header(doc, 'The Sequencing I Recommend')

    seq_table = doc.add_table(5, 3)
    remove_table_borders(seq_table)
    seq_table.columns[0].width = Inches(0.6)
    seq_table.columns[1].width = Inches(2.0)
    seq_table.columns[2].width = Inches(3.9)

    headers = ['Phase', 'Timeline', 'Priority Action']
    header_row = seq_table.rows[0]
    for i, h in enumerate(headers):
        c = header_row.cells[i]
        set_cell_bg(c, ASG_DARK_HEX)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'  {h}' if i == 0 else h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'

    data_rows = [
        ('1', 'Now — 90 days', 'Land 1–2 clients each in DMA (implant) and LGD (family law). Deliver Authority Blueprint for each. Begin content delivery.'),
        ('2', '90–180 days', 'Document first case studies. Build referral system within each niche. Raise price floor on new engagements.'),
        ('3', '180–365 days', 'Evaluate which niche is producing faster/easier revenue. Begin hiring one person (integrator or content support) in that niche.'),
        ('4', 'Year 2', 'Activate x2 Sales and/or LCS using proof from Phase 1–3. GHL rebrand as ASG delivery infrastructure. Systematize.'),
    ]

    for i, (phase, timeline, action) in enumerate(data_rows):
        row = seq_table.rows[i + 1]
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'

        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(f'  {phase}')
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ASG_BLUE
        r0.font.name = 'Calibri'

        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(timeline)
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = BODY_COLOR
        r1.font.name = 'Calibri'

        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(action)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = BODY_COLOR
        r2.font.name = 'Calibri'

    add_spacer(doc, 12)
    add_body(doc,
        'The businesses that are deferred are not abandoned — they are protected. Every hour Roger '
        'spends on a Tier 2 business before Tier 1 has paying clients is an hour that delays the '
        'proof that makes everything else easier to sell.',
        italic=True, color=ASG_CHARCOAL, space_after=6
    )

    add_pb(doc)


# ── Section 3: Revenue Path Recommendations ───────────────────────────────────

def build_revenue_path(doc):
    add_section_header(doc, '03 — FASTEST PATH TO REVENUE  |  Tanya Blackwood, CRO', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'Roger asked where to monetize fastest. Here is the honest answer: the fastest path to '
        'revenue in a professional services authority model is not a funnel. It is a conversation. '
        'Below are the three revenue levers I recommend activating immediately, in priority order.',
        space_after=12
    )

    levers = [
        {
            'number': '01',
            'title': 'The Warm Outreach Sprint (This Week)',
            'hex': GREEN_HEX,
            'body': (
                'Roger has an existing network — former clients, colleagues, referral sources, '
                'and professionals he has interacted with across his career. This is the lowest-'
                'cost, highest-probability first-revenue source available.'
            ),
            'actions': [
                'Identify 10–15 implant-focused dentists and 10–15 family law attorneys in Roger\'s existing network or connections.',
                'Send a direct, non-pitchy outreach message: "I\'m doing a limited number of free digital authority assessments for [niche] practices this month. Interested?" — no sales page, no funnel, just a direct conversation.',
                'From those conversations, close 1–2 paid Authority Blueprint engagements. One from DMA, one from LGD.',
                'Price the Authority Blueprint at no less than $2,500 for the initial engagement. Do not undervalue the entry point.',
            ],
            'timeline': '7–14 days to first conversation. 14–30 days to first paid engagement.',
        },
        {
            'number': '02',
            'title': 'The Free Audit Lead Generator (30–60 Days)',
            'hex': ASG_BLUE_HEX,
            'body': (
                'The Digital Authority Assessment is the most powerful prospecting tool in the ASG '
                'system. It delivers real value before any money changes hands, positions Roger as '
                'the expert, and creates a natural transition to a paid engagement conversation. '
                'Run one free audit per week for implant dentists and family law attorneys.'
            ),
            'actions': [
                'Select one target practice per week from LinkedIn, Google search, or a Dream 100 list.',
                'Run the full 13-dimension Digital Authority Assessment (Victor Shaw\'s team).',
                'Deliver the report with a brief cover note from Roger — no hard pitch, just: "Here\'s what I found. Happy to walk you through it."',
                'The walkthrough call IS the sales call. Let the report do the work.',
                'Convert at least 1 in 5 walkthrough calls to a paid Authority Blueprint engagement.',
            ],
            'timeline': 'First audit delivered in Week 2. First conversion expected by Week 6.',
        },
        {
            'number': '03',
            'title': 'Content Authority Engine (60–90 Days)',
            'hex': CHARCOAL_HEX,
            'body': (
                'While warm outreach and free audits generate near-term revenue, the content system '
                'builds the inbound pipeline that makes the business self-sustaining. Blog content, '
                'LinkedIn, and YouTube scripts are already being produced. The goal now is to get '
                'them published and indexed.'
            ),
            'actions': [
                'Publish the April 2026 blog articles to DMA\'s website (or a landing page if site is not yet live). Prioritize the 4 pillar pieces first.',
                'Post LinkedIn content 3–5x per week targeting each niche. Use the existing content packs for LGD; produce DMA LinkedIn pack next.',
                'Begin building the LGD and DMA YouTube channels. Short-form to start — 60–90 second clips cut from pillar blog content.',
                'Track which content drives profile views and inbound inquiries. Double down on what moves.',
            ],
            'timeline': 'First content live in Week 1–2. Inbound leads begin appearing by Month 2–3.',
        },
    ]

    for lev in levers:
        # Lever header
        lev_table = doc.add_table(1, 2)
        remove_table_borders(lev_table)
        lev_table.columns[0].width = Inches(0.5)
        lev_table.columns[1].width = Inches(6.0)

        nc = lev_table.cell(0, 0)
        set_cell_bg(nc, lev['hex'])
        np_ = nc.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(6)
        np_.paragraph_format.space_after = Pt(6)
        nr = np_.add_run(lev['number'])
        nr.bold = True
        nr.font.size = Pt(14)
        nr.font.color.rgb = WHITE
        nr.font.name = 'Calibri'

        tc_ = lev_table.cell(0, 1)
        set_cell_bg(tc_, SLATE_HEX)
        tp = tc_.paragraphs[0]
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(6)
        tr_ = tp.add_run(lev['title'])
        tr_.bold = True
        tr_.font.size = Pt(11)
        tr_.font.color.rgb = WHITE
        tr_.font.name = 'Calibri'

        add_body(doc, lev['body'], space_after=6)
        for action in lev['actions']:
            add_bullet(doc, action)

        tl_p = doc.add_paragraph()
        tl_p.paragraph_format.space_before = Pt(4)
        tl_p.paragraph_format.space_after = Pt(4)
        set_para_bg(tl_p, TINT_BG)
        tl_r1 = tl_p.add_run('  Timeline: ')
        tl_r1.bold = True
        tl_r1.font.size = Pt(9.5)
        tl_r1.font.color.rgb = ASG_CHARCOAL
        tl_r1.font.name = 'Calibri'
        tl_r2 = tl_p.add_run(lev['timeline'])
        tl_r2.font.size = Pt(9.5)
        tl_r2.font.color.rgb = BODY_COLOR
        tl_r2.font.name = 'Calibri'

        add_spacer(doc, 10)

    add_pb(doc)


# ── Section 4: Short-Term Focus Plan ──────────────────────────────────────────

def build_focus_plan(doc):
    add_section_header(doc, '04 — SHORT-TERM FOCUS PLAN  |  Preston Adler, COO', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'Roger has correctly identified the two-niche focus as the right short-term operating model. '
        'What follows is the operational execution plan: what the team works on, in what order, '
        'and how Roger will know it is working.',
        space_after=12
    )

    add_sub_header(doc, 'The Rule: One Deliverable at a Time Per Niche')
    add_body(doc,
        'The team\'s capacity is not unlimited, and Roger\'s attention is the true constraint. '
        'The operational rule is simple: one active production task per niche at a time. '
        'Parallel workstreams within a niche are fine (e.g., blog + LinkedIn simultaneously for DMA). '
        'But do not open a third niche until DMA and LGD each have at least one case study in hand.',
        space_after=12
    )

    add_sub_header(doc, 'DMA (Dental Marketing Authority) — 90-Day Operational Plan')
    dma_rows = [
        ('Week 1–2', 'Warm outreach to 10–15 implant dentists. Run 1 free Digital Authority Assessment. Publish 4 pillar blog articles.'),
        ('Week 2–4', 'Deliver first Authority Blueprint engagement. Begin DMA LinkedIn content pack (30 days).'),
        ('Month 2', 'Deliver Content Authority Strategy to first client. Begin YouTube content (Milo Vance). Close second client.'),
        ('Month 3', 'First case study documented. Referral ask to first client. Raise price floor by 20%.'),
    ]
    _build_plan_table(doc, dma_rows, ASG_BLUE_HEX)

    add_spacer(doc, 10)

    add_sub_header(doc, 'LGD (Legal Growth Dynamics) — 90-Day Operational Plan')
    lgd_rows = [
        ('Week 1–2', 'Warm outreach to 10–15 family law attorneys. April content calendar goes live. Run 1 free audit.'),
        ('Week 2–4', 'Deliver Authority Blueprint to first LGD client. Begin email sequence build (Rhett Callahan).'),
        ('Month 2', 'May content calendar delivered. Website copy audit (Audrey Voss). Close second client.'),
        ('Month 3', 'First case study documented. Dream 100 cold outreach activated (Solomon Grey). Referral program live.'),
    ]
    _build_plan_table(doc, lgd_rows, CHARCOAL_HEX)

    add_spacer(doc, 12)

    add_sub_header(doc, 'The Hiring Trigger')
    add_body(doc,
        'Roger indicated he will consider hiring one person once a business takes off. Here is the '
        'specific trigger we recommend:',
        space_after=6
    )

    triggers = [
        ('Revenue threshold', '$8,000–$10,000/month in recurring retainer revenue from a single niche. At that level, one hire pays for itself without straining cash flow.'),
        ('Capacity signal', 'Roger is turning down work or consistently behind on deliverable timelines. That is the clearest operational signal.'),
        ('Role to hire first', 'A delivery coordinator or content operations person — not a salesperson. Roger closes deals. The first hire protects delivery quality as volume grows.'),
        ('Niche to hire in', 'Whichever niche hits the revenue threshold first. If both hit simultaneously, hire for DMA first — the content system is more process-driven and easier to hand off than LGD\'s relationship-intensive work.'),
    ]

    for label, text in triggers:
        row_p = doc.add_paragraph()
        row_p.paragraph_format.space_before = Pt(3)
        row_p.paragraph_format.space_after = Pt(3)
        r1 = row_p.add_run(f'{label}: ')
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK_TEXT
        r1.font.name = 'Calibri'
        r2 = row_p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = BODY_COLOR
        r2.font.name = 'Calibri'

    add_spacer(doc, 12)
    add_pb(doc)


def _build_plan_table(doc, rows, accent_hex):
    table = doc.add_table(len(rows) + 1, 2)
    remove_table_borders(table)
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(5.2)

    # Header row
    hr = table.rows[0]
    for ci, label in enumerate(['Period', 'Focus']):
        c = hr.cells[ci]
        set_cell_bg(c, accent_hex)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f'  {label}' if ci == 0 else label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'

    for i, (period, focus) in enumerate(rows):
        row = table.rows[i + 1]
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'

        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(f'  {period}')
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = ASG_CHARCOAL
        r0.font.name = 'Calibri'

        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(focus)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = BODY_COLOR
        r1.font.name = 'Calibri'


# ── Section 5: Domain & Asset Inventory ───────────────────────────────────────

def build_asset_inventory(doc):
    add_section_header(doc, '05 — DORMANT ASSETS & DOMAIN INVENTORY  |  Iris Nolan, CTO', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'Roger mentioned several unused domains and deferred projects. From an infrastructure '
        'perspective, here is how I recommend treating the full asset inventory:',
        space_after=10
    )

    add_sub_header(doc, 'Domains — Decision Framework')
    add_body(doc,
        'For every domain in Roger\'s portfolio, apply this three-question test:',
        space_after=6
    )
    add_bullet(doc, 'Is there an active revenue path attached to this domain within the next 12 months? If yes — prioritize. If no — hold.')
    add_bullet(doc, 'Does it serve a niche that Roger is actively building proof in? If yes — build the site as content infrastructure. If no — park it with a redirect or placeholder.')
    add_bullet(doc, 'Is this domain\'s value tied to Roger being the face of it, or could it function as a productized system? If the latter — it belongs in the Phase 2 pipeline (x2 Sales, LCS) and should not absorb energy now.')
    add_spacer(doc, 8)

    add_sub_header(doc, 'Active Infrastructure — Recommended Priority Build Order')
    infra_rows = [
        ('DMA Website', 'HIGH', 'Dental Marketing Authority needs a content-publishing home. Minimum viable: landing page + blog. Publish April articles now.'),
        ('LGD Website', 'HIGH', 'Already in use. Priority: website copy audit + conversion optimization (Audrey Voss). Ongoing: blog publishing.'),
        ('ASG Website', 'MEDIUM', 'Authority Systems Group needs a credential-building web presence for prospects who research Roger. Not urgent — but build before scaling outreach.'),
        ('Legal Growth Pilot (GHL)', 'MEDIUM', 'Operational backbone. Use for client CRM builds. Rebrand consideration deferred to Phase 2.'),
        ('x2 Sales domain', 'LOW', 'Parked. Activate in Phase 2 after DMA and LGD case studies are documented.'),
        ('LCS domain', 'LOW', 'Parked. Activate with training program build in Phase 2.'),
        ('SMB Consulting domain', 'ARCHIVE', 'No action. Let expire or redirect to ASG.'),
    ]

    infra_table = doc.add_table(len(infra_rows) + 1, 3)
    remove_table_borders(infra_table)
    infra_table.columns[0].width = Inches(1.8)
    infra_table.columns[1].width = Inches(0.8)
    infra_table.columns[2].width = Inches(3.9)

    hr = infra_table.rows[0]
    for ci, label in enumerate(['Asset', 'Priority', 'Recommendation']):
        c = hr.cells[ci]
        set_cell_bg(c, ASG_DARK_HEX)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'

    priority_colors = {
        'HIGH': GREEN_HEX, 'MEDIUM': AMBER_HEX, 'LOW': CHARCOAL_HEX, 'ARCHIVE': 'AAAAAA'
    }

    for i, (asset, priority, rec) in enumerate(infra_rows):
        row = infra_table.rows[i + 1]
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'

        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(asset)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = BODY_COLOR
        r0.font.name = 'Calibri'

        c1 = row.cells[1]
        set_cell_bg(c1, priority_colors.get(priority, CHARCOAL_HEX))
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(priority)
        r1.bold = True
        r1.font.size = Pt(8)
        r1.font.color.rgb = WHITE
        r1.font.name = 'Calibri'

        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(rec)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = BODY_COLOR
        r2.font.name = 'Calibri'

    add_spacer(doc, 12)
    add_body(doc,
        'The unused domains are not a problem — they are latent leverage. The risk is treating them '
        'as obligations rather than options. Options have value. Obligations drain focus. '
        'Until Roger has two active revenue streams generating $5K+/month each, the domains '
        'sit on the shelf — and that is the right call.',
        italic=True, color=ASG_CHARCOAL, space_after=6
    )

    add_pb(doc)


# ── Section 6: Director's Summary ─────────────────────────────────────────────

def build_directors_summary(doc):
    add_section_header(doc, '06 — DIRECTOR\'S SUMMARY  |  Roger Bauer', bg_hex=ASG_DARK_HEX)
    add_spacer(doc, 8)

    add_body(doc,
        'The clarity in this document took longer to get to than it should have. That is fine. '
        'Most people never get there at all.',
        size=11, bold=True, color=DARK_TEXT, space_after=8
    )

    add_body(doc,
        'Here is what I know to be true:',
        space_after=6
    )
    add_bullet(doc, 'The authority model is correct. Small business clients without the budget for premium positioning are a ceiling, not a market. Authority clients — professionals who can raise their own prices when their authority rises — are an unlimited market.')
    add_bullet(doc, 'Two niches first. Everything else is a distraction until DMA and LGD each have a client, a result, and a story to tell.')
    add_bullet(doc, 'The team produces the work. The work becomes the proof. The proof closes the next client. That is the system.')
    add_bullet(doc, 'Dormant assets are not failures. They are investments in future optionality. When the time is right, everything has a place.')
    add_spacer(doc, 10)

    add_body(doc,
        'The focus is set. The team is ready. The work starts now.',
        size=11, bold=True, color=ASG_BLUE, space_after=12
    )

    add_divider(doc)
    add_spacer(doc, 8)

    # Closing block
    close_table = doc.add_table(1, 2)
    remove_table_borders(close_table)
    close_table.columns[0].width = Inches(3.25)
    close_table.columns[1].width = Inches(3.25)

    # Left: team sign-offs
    lc = close_table.cell(0, 0)
    set_cell_bg(lc, LIGHT_BG)
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after = Pt(8)
    lr = lp.add_run(
        '  Reviewed and approved by:\n\n'
        '  Daniel Frost, CSO\n'
        '  Tanya Blackwood, CRO\n'
        '  Preston Adler, COO\n'
        '  Iris Nolan, CTO'
    )
    lr.font.size = Pt(9.5)
    lr.font.color.rgb = ASG_CHARCOAL
    lr.font.name = 'Calibri'

    # Right: date / version
    rc = close_table.cell(0, 1)
    set_cell_bg(rc, TINT_BG)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(8)
    rp.paragraph_format.space_after = Pt(8)
    rr = rp.add_run(
        'Authority Systems Group™\n'
        'Portfolio Strategy Brief\n'
        'Version 1.0  |  March 8, 2026\n'
        'Director: Roger Bauer  '
    )
    rr.font.size = Pt(9)
    rr.font.color.rgb = ASG_CHARCOAL
    rr.font.name = 'Calibri'

    add_spacer(doc, 20)

    # Footer note
    fn = doc.add_paragraph()
    fn.paragraph_format.space_before = Pt(4)
    fn.paragraph_format.space_after = Pt(0)
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fn_r = fn.add_run(
        'Authority Systems Group™ — Confidential. Internal document. Not for distribution.\n'
        'This document will be updated as portfolio priorities shift and case studies are documented.'
    )
    fn_r.font.size = Pt(8)
    fn_r.font.color.rgb = FOOTER_COLOR
    fn_r.font.name = 'Calibri'
    fn_r.italic = True


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    add_header_footer(doc)

    build_cover(doc)
    build_portfolio_inventory(doc)
    build_strategic_assessment(doc)
    build_revenue_path(doc)
    build_focus_plan(doc)
    build_asset_inventory(doc)
    build_directors_summary(doc)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
