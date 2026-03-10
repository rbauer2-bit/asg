"""
ASG Internal Reference: Target Niche Analysis — Top 20 Recommendations
Build script: authority-systems-group_target-niche-analysis_20260309.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Brand constants ────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK      = RGBColor(0x1A, 0x1A, 0x1A)
ASG_BODY      = RGBColor(0x33, 0x33, 0x33)
ASG_LIGHT     = RGBColor(0xF5, 0xF5, 0xF5)
ASG_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ASG_ALT_BLUE  = RGBColor(0xEA, 0xF6, 0xFC)

BASE_DIR  = "/Users/Roger/Dropbox/code/authority-systems-group"
LOGO_PATH = os.path.join(BASE_DIR, "brand-standards/assets/asgFull.png")
OUT_PATH  = os.path.join(BASE_DIR,
    "client-onboarding/clients/authority-systems-group/outputs/docx/"
    "authority-systems-group_target-niche-analysis_20260309.docx")

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_bottom_border_para(para, color="25aae1", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = ASG_BLUE
    add_bottom_border_para(p, color="25aae1", sz="6")
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = ASG_CHARCOAL
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = ASG_BLUE
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_BODY
    return p

def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_BLUE
    run.bold = True
    # left blue border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "25aae1")
    pBdr.append(left)
    pPr.append(pBdr)
    # light bg
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EAF6FC")
    pPr.append(shd)
    return p

def add_section_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    add_bottom_border_para(p, color="25aae1", sz="4")

def add_header(doc, title):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.5))
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)
    title_run = para.add_run("\t" + title)
    title_run.font.name  = "Arial"
    title_run.font.size  = Pt(9)
    title_run.font.color.rgb = ASG_CHARCOAL
    add_bottom_border_para(para, color="25aae1", sz="6")

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = para.add_run("Authority Systems Group\u2122  |  AuthoritySystemsGroup.com  |  Internal Reference")
    left_run.font.name = "Arial"
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = ASG_CHARCOAL
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "25aae1")
    pBdr.append(top)
    pPr.append(pBdr)

def build_niche_table(doc, headers, rows):
    col_widths = [4200, 1300, 1200, 1200, 1100, 820]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr_row = table.rows[0]
    for i, (cell, w, txt) in enumerate(zip(hdr_row.cells, col_widths, headers)):
        cell.width = w
        set_cell_bg(cell, "25aae1")
        set_cell_borders(cell, "CCCCCC")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = ASG_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
    # data rows
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = "F5F5F5" if idx % 2 == 1 else "FFFFFF"
        for cell, w, txt in zip(row.cells, col_widths, row_data):
            cell.width = w
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(txt)
            run.font.name  = "Arial"
            run.font.size  = Pt(9)
            run.font.color.rgb = ASG_BODY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()  # spacer

def build_summary_table(doc, headers, rows):
    col_widths = [2600, 1400, 1300, 2600, 1820]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for cell, w, txt in zip(hdr_row.cells, col_widths, headers):
        cell.width = w
        set_cell_bg(cell, "25aae1")
        set_cell_borders(cell, "CCCCCC")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = ASG_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = "F5F5F5" if idx % 2 == 1 else "FFFFFF"
        for cell, w, txt in zip(row.cells, col_widths, row_data):
            cell.width = w
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(txt)
            run.font.name  = "Arial"
            run.font.size  = Pt(9)
            run.font.color.rgb = ASG_BODY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

# ── Document data ──────────────────────────────────────────────────────────────
CLUSTERS = [
    {
        "title": "Cluster 1 — Financial & Accounting Advisory",
        "subtitle": "High-trust practitioners. The buyer's certainty about the person IS the sale.",
        "callout": "These practitioners ARE the product. Trust precedes every transaction. Authority content is the dominant acquisition lever — and almost none of these consultants have one.",
        "niches": [
            ("Controller Services (Fractional)", "9.60", "$144K", "$48K", "Retainer", "2/5"),
            ("ERISA Compliance Services",         "9.60", "$144K", "$48K", "Retainer", "2/5"),
            ("Pension Plan Administration",        "9.54", "$144K", "$48K", "Retainer", "2/5"),
            ("401(k) Plan Advisory",               "9.20", "$128K", "$32K", "Retainer", "2/5"),
            ("Retirement Planning Services",        "9.20", "$128K", "$32K", "Retainer", "2/5"),
        ],
        "rationale": (
            "CFOs, fractional controllers, and retirement advisors sell expertise packaged as a person. "
            "The buyer's belief about the practitioner's credibility is the entire sale. "
            "LinkedIn authority content combined with a structured belief sequence is the highest-leverage "
            "acquisition system in these markets. Most competitors are invisible on content. The gap is wide open. "
            "The Belief-to-Buy Framework maps perfectly: 'My current advisor is reactive' is the enemy belief, "
            "and a well-positioned fractional CFO with consistent authority content rewrites that belief before "
            "the first conversation happens."
        ),
    },
    {
        "title": "Cluster 2 — HR & People Strategy Consulting",
        "subtitle": "High LTV, short-to-medium sales cycles, LinkedIn-native buyers, underserved by quality marketing.",
        "callout": "HR consultants are fighting for the same CHROs and business owners on LinkedIn with generic content. The Belief-to-Buy Framework maps directly — and the enemy belief is begging to be disrupted.",
        "niches": [
            ("Talent Management Consulting",      "9.55", "$135K", "$45K", "Retainer", "2/5"),
            ("Performance Management Consulting", "9.60", "$135K", "$45K", "Retainer", "2/5"),
            ("HR Consulting Services",            "9.19", "$126K", "$42K", "Retainer", "2/5"),
            ("Labor Relations Consulting",        "9.33", "$126K", "$42K", "Retainer", "2/5"),
            ("Employee Relations Consulting",     "9.33", "$126K", "$42K", "Retainer", "2/5"),
        ],
        "rationale": (
            "HR consultants are fighting for the same CHROs and business owners on LinkedIn with nearly "
            "identical positioning. The enemy belief — 'HR consulting is reactive, compliance-focused, not "
            "strategic' — is universal and almost never challenged in their marketing. "
            "A well-positioned HR consultant with a strong content engine rewrites that belief before "
            "any competitor gets the meeting. These clients also become long-term retainers because HR "
            "challenges don't resolve — they evolve. Churn is low, LTV compounds."
        ),
    },
    {
        "title": "Cluster 3 — Compliance & Risk Consulting",
        "subtitle": "Regulatory complexity creates a credibility crisis for buyers — authority is the only safe harbor.",
        "callout": "When the stakes are regulatory — fines, shutdowns, liability — buyers don't shop on price. They buy certainty. That's the deepest convergence point on the Belief Track.",
        "niches": [
            ("Healthcare Compliance Consulting", "9.33", "$126K", "$42K", "Retainer", "2/5"),
            ("Food Safety Consulting",           "9.33", "$126K", "$42K", "Retainer", "2/5"),
            ("Energy Procurement Consulting",    "9.60", "$144K", "$48K", "Retainer", "2/5"),
        ],
        "rationale": (
            "When regulatory consequences are on the table — OCR audits, FDA violations, energy contract "
            "penalties — buyers make decisions based on perceived safety, not price comparison. "
            "A compliance consultant with a visible, expert content presence is the obvious safe choice. "
            "A consultant with no content presence is invisible and indistinguishable from the next firm. "
            "These markets respond exceptionally well to the Belief-to-Buy Framework because the emotional "
            "track — Dissatisfaction → Contrast → Desire → Urgency → Relief — is already primed. "
            "We just have to show up as the relief."
        ),
    },
    {
        "title": "Cluster 4 — Specialty Consulting",
        "subtitle": "Niche expertise commanding premium retainers. Authority is the only meaningful differentiator.",
        "callout": "Buyers in these markets have almost no framework for evaluating providers — they default to whoever looks most credible. The practitioner who owns the content space owns the category.",
        "niches": [
            ("Hotel Revenue Management",    "9.60", "$144K", "$48K", "Retainer", "2/5"),
            ("Emergency Management Consulting", "9.60", "$144K", "$48K", "Retainer", "2/5"),
            ("Rights Management Services",  "9.60", "$135K", "$45K", "Retainer", "2/5"),
            ("Supply Chain Financing",      "9.32", "$144K", "$48K", "Retainer", "2/5"),
        ],
        "rationale": (
            "These are hyper-specialized fields where buyers have almost no framework for evaluating providers. "
            "The hotel GM evaluating a revenue management consultant doesn't know what great looks like. "
            "The municipal director considering emergency management consulting is equally unsure. "
            "They default to whoever appears most authoritative, most credible, most certain. "
            "Almost none of these consultants have a functioning content or authority system. "
            "First-mover advantage in these niches is significant — the practitioner who builds the "
            "content presence first owns the category perception for years."
        ),
    },
    {
        "title": "Cluster 5 — High-Trust Vertical Marketing",
        "subtitle": "These businesses hire external marketing firms. ASG IS the firm they need.",
        "callout": "These verticals spend heavily on marketing but are flooded with generalist agencies. Niche specificity IS credibility — and the highest deal sizes on this list reflect that.",
        "niches": [
            ("Fertility Clinic Marketing",       "9.60", "$135K", "$45K", "Retainer", "2/5"),
            ("Custom Home Builder Marketing",    "9.60", "$156K", "$52K", "Retainer", "2/5"),
            ("Luxury Auto Dealership Marketing", "9.60", "$156K", "$52K", "Retainer", "2/5"),
        ],
        "rationale": (
            "Fertility clinics, custom home builders, and luxury auto dealerships all share a common "
            "challenge: they spend heavily on marketing but rarely find an agency that speaks their "
            "specific language. A generalist digital agency pitching a fertility clinic sounds exactly "
            "like every other agency. An ASG engagement that leads with deep niche knowledge, "
            "patient acquisition expertise, and authority positioning wins the room immediately. "
            "These also carry the highest deal sizes and LTVs in this entire analysis. "
            "Cluster 5 should be treated as a dedicated prospecting lane with tailored positioning for each vertical."
        ),
    },
]

SUMMARY_ROWS = [
    ("Financial & Accounting", "9.51", "$136K", "LinkedIn / Cold Email / Content", "45–90 days"),
    ("HR & People Strategy",   "9.40", "$130K", "LinkedIn / Cold Email / Content", "60–75 days"),
    ("Compliance & Risk",      "9.42", "$132K", "LinkedIn / Cold Email / Content", "60–75 days"),
    ("Specialty Consulting",   "9.53", "$141K", "LinkedIn / Cold Email / Content", "75–90 days"),
    ("Vertical Marketing",     "9.60", "$149K", "PPC / SEO / Content",             "75–90 days"),
]

# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

# Page setup
section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin   = Inches(1.0)
section.bottom_margin = Inches(1.0)

add_header(doc, "Target Niche Analysis — Internal Reference")
add_footer(doc)

# ── Cover page ─────────────────────────────────────────────────────────────────
from docx.oxml import OxmlElement as OE

# Dark header block via table
cov_table = doc.add_table(rows=1, cols=1)
cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cov_cell = cov_table.rows[0].cells[0]
set_cell_bg(cov_cell, "1a1a1a")
for side in ("top","left","bottom","right"):
    tcPr = cov_cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "none")
    tcBdr.append(b)
    tcPr.append(tcBdr)

logo_para = cov_cell.paragraphs[0]
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para.paragraph_format.space_before = Pt(24)
logo_run = logo_para.add_run()
logo_run.add_picture(LOGO_PATH, width=Inches(2.2))

tag_para = cov_cell.add_paragraph("Positioning You As The Authority")
tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_para.paragraph_format.space_before = Pt(4)
tag_para.paragraph_format.space_after  = Pt(24)
tag_run = tag_para.runs[0]
tag_run.font.name    = "Arial"
tag_run.font.size    = Pt(11)
tag_run.font.italic  = True
tag_run.font.color.rgb = ASG_BLUE

# Spacer
doc.add_paragraph()

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(32)
t_run = title_p.add_run("Target Niche Analysis")
t_run.font.name  = "Arial"
t_run.font.size  = Pt(32)
t_run.font.bold  = True
t_run.font.color.rgb = ASG_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = sub_p.add_run("Top 20 Recommended Markets for ASG Expansion")
s_run.font.name  = "Arial"
s_run.font.size  = Pt(16)
s_run.font.color.rgb = ASG_CHARCOAL

# Divider
rule_p = doc.add_paragraph()
rule_p.paragraph_format.space_before = Pt(24)
rule_p.paragraph_format.space_after  = Pt(24)
add_bottom_border_para(rule_p, color="58585a", sz="8")

# Prepared by block
prep_p = doc.add_paragraph()
prep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
prep_p.paragraph_format.space_before = Pt(16)
prep_run = prep_p.add_run("Internal Reference  |  Authority Systems Group\u2122  |  March 2026")
prep_run.font.name   = "Arial"
prep_run.font.size   = Pt(11)
prep_run.font.italic = True
prep_run.font.color.rgb = ASG_CHARCOAL

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf_run = conf_p.add_run("DIRECTOR USE ONLY \u2014 Roger Bauer / Authority Systems Group\u2122")
conf_run.font.name  = "Arial"
conf_run.font.size  = Pt(9)
conf_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_page_break(doc)

# ── Introduction ───────────────────────────────────────────────────────────────
add_h1(doc, "Overview & Strategic Context")

add_body(doc,
    "This document is a permanent internal reference for Authority Systems Group\u2122. It captures the "
    "20 highest-priority target niches drawn from a database of 1,077 evaluated markets, filtered and "
    "ranked against ASG's acquisition model: composite score, LTV, retainer structure, serviceability, "
    "cashflow risk, and channel alignment."
)
add_body(doc,
    "All 20 niches meet the following minimum criteria:"
)

criteria = [
    "Composite score \u2265 9.2 / 10.0",
    "Lifetime value (LTV) \u2265 $126,000",
    "Retainer payment structure",
    "Serviceability rating 2/5 (easy to deliver at scale)",
    "LOW cashflow risk under all four stress scenarios",
    "Primary acquisition channels: LinkedIn / Cold Email / Content (or PPC/SEO/Content for vertical marketing)",
]
for c in criteria:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(c)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_BODY

doc.add_paragraph()
add_body(doc,
    "Niches already in active client engagements (law / attorney practices, business / executive coaching) "
    "are excluded from this analysis. The 20 niches below represent ASG's highest-confidence expansion "
    "targets as of March 2026."
)

add_callout(doc,
    "Selection criteria prioritize professional service practitioners where the owner IS the product, "
    "trust precedes every purchasing decision, and authority content creates an asymmetric competitive "
    "advantage. These are the exact conditions where the Belief-to-Buy Framework\u2122 dominates."
)

add_section_rule(doc)

# ── Summary scorecard ──────────────────────────────────────────────────────────
add_h2(doc, "Summary Scorecard by Cluster")

build_summary_table(doc,
    ["Cluster", "Avg Composite", "Avg LTV", "Primary Channels", "Avg Sales Cycle"],
    SUMMARY_ROWS
)

add_page_break(doc)

# ── Clusters ───────────────────────────────────────────────────────────────────
NICHE_TABLE_HEADERS = ["Niche", "Composite", "LTV", "Deal Size", "Payment", "Srvc"]

for cluster in CLUSTERS:
    add_h1(doc, cluster["title"])
    add_body(doc, cluster["subtitle"])
    add_callout(doc, cluster["callout"])

    build_niche_table(doc, NICHE_TABLE_HEADERS, cluster["niches"])

    add_h3(doc, "Strategic Rationale")
    add_body(doc, cluster["rationale"])

    add_section_rule(doc)
    add_page_break(doc)

# ── Roger's Read ───────────────────────────────────────────────────────────────
add_h1(doc, "Roger's Read — Priority & Next Steps")

add_body(doc,
    "The highest-confidence immediate targets are Clusters 1 and 2 — Financial & Accounting Advisory "
    "and HR & People Strategy Consulting. Both are LinkedIn-native, authority-starved, and the full "
    "ASG deliverable stack maps directly to how they acquire clients. Prospecting campaigns can launch "
    "with existing infrastructure."
)
add_body(doc,
    "The compliance cluster (Cluster 3) is slightly longer cycle but commands premium retainers and "
    "near-zero churn once embedded. Healthcare Compliance and Energy Procurement are the two strongest "
    "entries in this group — both have clear enemy beliefs and regulatory urgency that accelerates the "
    "emotional track."
)
add_body(doc,
    "Cluster 5 (Vertical Marketing) deserves a dedicated prospecting lane. Deal sizes are the highest "
    "across all 20 niches, buyers are active marketing spenders, and the niche-specialization story "
    "writes itself. Custom Home Builder Marketing and Luxury Auto Dealership Marketing both carry "
    "$156K LTV with 9.60 composite scores."
)

add_callout(doc,
    "Recommended first move: Activate a Cold Email + LinkedIn outreach sequence targeting Fractional "
    "Controller and Performance Management Consulting prospects. Both have 60-day sales cycles, "
    "$135K\u2013$144K LTV, and retainer structures that feed ASG's recurring revenue base immediately."
)

doc.add_paragraph()
add_body(doc,
    "This document should be consulted whenever:"
)
next_steps = [
    "Evaluating a new prospecting campaign target",
    "Assessing whether an inbound inquiry fits the ASG expansion model",
    "Selecting the next vertical for a Dopamine Ladder or Dream 100 build-out",
    "Briefing Daniel Frost or Tanya Blackwood on new market opportunities",
    "Comparing ROI potential across expansion options",
]
for s in next_steps:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(s)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_BODY

doc.add_paragraph()
signoff = doc.add_paragraph()
signoff.paragraph_format.space_before = Pt(24)
signoff.paragraph_format.left_indent  = Inches(0.5)
sr = signoff.add_run(
    "Reviewed and approved for use as an ongoing reference document.\n"
    "Roger Bauer, Director \u2014 Authority Systems Group\u2122"
)
sr.font.name   = "Arial"
sr.font.size   = Pt(11)
sr.font.italic = True
sr.font.color.rgb = ASG_CHARCOAL

# ── Save ───────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
