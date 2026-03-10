#!/usr/bin/env python3
"""
GISCO Corp — Authority Blueprint™ DOCX Builder
Authority Systems Group™
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ─────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
ASG_FOOTER    = RGBColor(0x88, 0x88, 0x88)
ASG_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ASG_DARK_BG   = '1a1a1a'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_BLUE_HEX  = '25aae1'
ASG_GRAY_HEX  = 'CCCCCC'

# ── File Paths ────────────────────────────────────────────────────────────────
BASE = '/Users/Roger/Dropbox/code/authority-systems-group'
SLUG = 'gisco-corp'
DATE = '20260309'
ROOT = f'{BASE}/client-onboarding/clients/{SLUG}'

DOC_CFG = {
    'md':           f'{ROOT}/outputs/md/{SLUG}_authority-blueprint_{DATE}.md',
    'out':          f'{ROOT}/outputs/docx/{SLUG}_authority-blueprint_{DATE}.docx',
    'label':        'AUTHORITY BLUEPRINT\u2122',
    'sub':          'A Strategic Growth System Built Exclusively for Tony Long',
    'client_name':  'GISCO Corporation',
    'client_city':  'Louisville, KY',
    'props_title':  'GISCO Corp — Authority Blueprint\u2122',
    'props_subject':'Strategic Growth System | Commercial, Industrial & Residential Insulation',
}


# ── XML Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def add_left_border(para, color=ASG_BLUE_HEX, sz='24', space='144'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_bottom_border(para, color=ASG_BLUE_HEX, sz='6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_line_spacing(para, multiplier=1.4):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), str(int(multiplier * 240)))
    sp.set(qn('w:lineRule'), 'auto')


def set_indent(para, left_inches=0.0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(left_inches * 1440)))
    pPr.append(ind)


def add_page_number_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ── Inline Formatting ─────────────────────────────────────────────────────────

def add_inline_runs(para, text, base_font='Calibri', base_size=11, base_color=None):
    if base_color is None:
        base_color = ASG_BODY_TEXT
    parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = ASG_CHARCOAL
        else:
            run = para.add_run(part)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color


def strip_inline(text):
    text = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+?)\*', r'\1', text)
    text = re.sub(r'`([^`]+?)`', r'\1', text)
    return text


# ── Cover Page ────────────────────────────────────────────────────────────────

def add_cover_page(doc, cfg):
    def cover_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        set_para_bg(p, ASG_DARK_BG)
        return p

    def cover_run(para, text, size, bold=False, italic=False, color=None):
        if color is None:
            color = ASG_WHITE
        r = para.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        return r

    for _ in range(3):
        cover_para(space_before=0, space_after=0)

    p = cover_para(space_before=36, space_after=6)
    cover_run(p, 'AUTHORITY SYSTEMS GROUP\u2122', 11, bold=True, color=RGBColor(0x88, 0x88, 0x88))

    p = cover_para(space_before=6, space_after=6)
    cover_run(p, 'Positioning You As The Authority', 10, italic=True, color=ASG_BLUE)

    cover_para(space_before=12, space_after=12)

    p = cover_para(space_before=6, space_after=4)
    cover_run(p, 'AUTHORITY BLUEPRINT\u2122', 32, bold=True, color=ASG_BLUE)

    p = cover_para(space_before=4, space_after=4)
    cover_run(p, cfg['sub'], 12, italic=True, color=RGBColor(0xCC, 0xCC, 0xCC))

    cover_para(space_before=16, space_after=16)

    p = cover_para(space_before=4, space_after=8)
    cover_run(p, cfg['client_name'], 26, bold=True)

    p = cover_para(space_before=4, space_after=4)
    cover_run(p, cfg['client_city'], 13, color=RGBColor(0xAA, 0xAA, 0xAA))

    cover_para(space_before=8, space_after=8)

    p = cover_para(space_before=4, space_after=4)
    cover_run(p, 'Directed by Roger Bauer \u2014 Authority Systems Group\u2122', 11,
              italic=True, color=RGBColor(0x88, 0x88, 0x88))

    p = cover_para(space_before=4, space_after=4)
    cover_run(p, 'March 9, 2026 \u2014 Version 1.0', 10, color=RGBColor(0x77, 0x77, 0x77))

    for _ in range(6):
        cover_para(space_before=0, space_after=0)

    p = cover_para(space_before=24, space_after=6)
    cover_run(p, 'CONFIDENTIAL \u2014 PREPARED EXCLUSIVELY FOR TONY LONG / GISCO CORPORATION',
              9, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()


# ── Footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run(
            'Authority Systems Group\u2122 \u2014 Confidential. '
            'Prepared exclusively for GISCO Corporation / Tony Long.'
        )
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9)
        r1.font.color.rgb = ASG_FOOTER
        fp.add_run('\t')
        r2 = fp.add_run()
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.color.rgb = ASG_FOOTER
        add_page_number_field(r2)


# ── Table Builder ─────────────────────────────────────────────────────────────

def build_table(doc, rows_raw):
    rows = []
    for line in rows_raw:
        if re.match(r'^\|[\s\-|:]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row_cells in enumerate(rows):
        tbl_row = tbl.rows[ri]
        is_header = (ri == 0)
        for ci in range(ncols):
            cell_text = row_cells[ci] if ci < len(row_cells) else ''
            cell = tbl_row.cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            plain = strip_inline(cell_text)
            if is_header:
                r = p.add_run(plain)
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = ASG_WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(cell, ASG_BLUE_HEX)
            else:
                add_inline_runs(p, cell_text, base_font='Calibri', base_size=10)
                if ri % 2 == 0:
                    set_cell_bg(cell, ASG_LIGHT_BG)
    doc.add_paragraph()


# ── Code / Pre-formatted block ─────────────────────────────────────────────────

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        set_para_bg(p, ASG_LIGHT_BG)
        set_indent(p, left_inches=0.25)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = ASG_CHARCOAL


# ── Checkbox line ─────────────────────────────────────────────────────────────

def add_checkbox_line(doc, text):
    """Render lines starting with - [ ], - [x], □, ☐ as checklist items."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_line_spacing(p, 1.3)
    add_inline_runs(p, text, base_font='Calibri', base_size=11)


# ── Main Parser & Builder ─────────────────────────────────────────────────────

def build_document(cfg):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    add_cover_page(doc, cfg)

    with open(cfg['md'], 'r', encoding='utf-8') as f:
        lines = [ln.rstrip('\n') for ln in f.readlines()]

    table_buf = []
    in_table  = False
    code_buf  = []
    in_code   = False
    first_h1  = True

    i = 0
    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        # ── Code block ────────────────────────────────────────────────────
        if s.startswith('```'):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code  = False
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────
        if s.startswith('|') and s.endswith('|'):
            table_buf.append(s)
            in_table = True
            i += 1
            continue
        if in_table:
            build_table(doc, table_buf)
            table_buf = []
            in_table  = False
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if s == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            add_bottom_border(p, color=ASG_BLUE_HEX, sz='6')
            i += 1
            continue

        # ── H1 ────────────────────────────────────────────────────────────
        if re.match(r'^# (?!#)', s):
            text = strip_inline(s[2:])
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(12)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(22)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            add_bottom_border(p, color=ASG_BLUE_HEX, sz='10')
            i += 1
            continue

        # ── H2 ────────────────────────────────────────────────────────────
        if re.match(r'^## (?!#)', s):
            text = strip_inline(s[3:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(17)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            add_bottom_border(p, color=ASG_BLUE_HEX, sz='6')
            i += 1
            continue

        # ── H3 ────────────────────────────────────────────────────────────
        if re.match(r'^### (?!#)', s):
            text = strip_inline(s[4:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(13)
            r.font.bold  = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── H4 ────────────────────────────────────────────────────────────
        if re.match(r'^#### (?!#)', s):
            text = strip_inline(s[5:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.font.name   = 'Calibri'
            r.font.size   = Pt(12)
            r.font.bold   = True
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────
        if s.startswith('> '):
            text  = s[2:]
            plain = strip_inline(text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(10)
            set_indent(p, left_inches=0.25)
            set_para_bg(p, ASG_LIGHT_BG)
            add_left_border(p, color=ASG_BLUE_HEX, sz='24', space='144')
            r = p.add_run(plain)
            r.font.name   = 'Calibri'
            r.font.size   = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Checkbox items ────────────────────────────────────────────────
        if re.match(r'^[-*]\s+\[[ x]\]', s) or re.match(r'^[□☐✓☑]', s):
            text = re.sub(r'^[-*]\s+\[[ x]\]\s*', '', s)
            text = re.sub(r'^[□☐✓☑]\s*', '', s)
            add_checkbox_line(doc, text)
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────
        if re.match(r'^[-*]\s+', s):
            text = re.sub(r'^[-*]\s+', '', s)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_line_spacing(p, 1.3)
            add_inline_runs(p, text, base_font='Calibri', base_size=11)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────
        if re.match(r'^\d+\.\s+', s):
            text = re.sub(r'^\d+\.\s+', '', s)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_line_spacing(p, 1.3)
            add_inline_runs(p, text, base_font='Calibri', base_size=11)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── All-bold standalone line ───────────────────────────────────────
        if re.match(r'^\*\*[^*]+\*\*$', s):
            text = s[2:-2]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(12)
            r.font.bold  = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Italic-only standalone line (attribution / sign-off) ──────────
        if re.match(r'^\*[^*].+[^*]\*$', s):
            text = s[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            set_indent(p, left_inches=0.25)
            r = p.add_run(strip_inline('*' + text + '*'))
            r.font.name   = 'Calibri'
            r.font.size   = Pt(10)
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Default body paragraph ────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        set_line_spacing(p, 1.4)
        add_inline_runs(p, s, base_font='Calibri', base_size=11)
        i += 1

    if table_buf:
        build_table(doc, table_buf)
    if code_buf:
        add_code_block(doc, code_buf)

    add_footer(doc)

    doc.core_properties.title   = cfg['props_title']
    doc.core_properties.subject = cfg['props_subject']
    doc.core_properties.author  = 'Authority Systems Group\u2122'

    os.makedirs(os.path.dirname(cfg['out']), exist_ok=True)
    doc.save(cfg['out'])
    size_kb = os.path.getsize(cfg['out']) // 1024
    print(f'Saved: {cfg["out"]}')
    print(f'Size:  {size_kb} KB')


if __name__ == '__main__':
    print(f'Building: {DOC_CFG["label"]} — {DOC_CFG["client_name"]} ...')
    build_document(DOC_CFG)
    print('Done.')
