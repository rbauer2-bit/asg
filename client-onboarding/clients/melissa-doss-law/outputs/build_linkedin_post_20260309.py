"""
Melissa Doss Law — LinkedIn Post Deliverable
"Get Over It Day" — March 9, 2026
Build: melissa-doss-law_linkedin-post_get-over-it-day_20260309.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand constants ────────────────────────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
ASG_BODY     = RGBColor(0x33, 0x33, 0x33)
ASG_LIGHT    = RGBColor(0xF5, 0xF5, 0xF5)
ASG_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ASG_MID_GRAY = RGBColor(0x88, 0x88, 0x88)

BASE_DIR  = "/Users/Roger/Dropbox/code/authority-systems-group"
LOGO_PATH = os.path.join(BASE_DIR, "brand-standards/assets/asgFull.png")
OUT_PATH  = os.path.join(BASE_DIR,
    "client-onboarding/clients/melissa-doss-law/outputs/docx/"
    "melissa-doss-law_linkedin-post_get-over-it-day_20260309.docx")

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_bottom_border(para, color="25aae1", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_left_border(para, color="25aae1", sz="12"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EAF6FC")
    pPr.append(shd)

def add_header(doc):
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.5))
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)
    tr = para.add_run("\tContent Deliverable — Melissa Doss Law")
    tr.font.name = "Arial"; tr.font.size = Pt(9); tr.font.color.rgb = ASG_CHARCOAL
    add_bottom_border(para, color="25aae1", sz="6")

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        "Authority Systems Group\u2122  |  AuthoritySystemsGroup.com  |  "
        "Confidential. Prepared exclusively for Melissa Doss Law."
    )
    run.font.name = "Arial"; run.font.size = Pt(9); run.font.color.rgb = ASG_CHARCOAL
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1"); top.set(qn("w:color"), "25aae1")
    pBdr.append(top); pPr.append(pBdr)

def body(doc, text, size=11, color=None, italic=False, bold=False, space_before=0, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color or ASG_BODY
    run.font.italic = italic
    run.font.bold   = bold
    return p

def label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.font.size = Pt(8)
    r.font.bold = True; r.font.color.rgb = ASG_CHARCOAL
    return p

def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    add_bottom_border(p, color="25aae1", sz="4")

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after  = Pt(0)

# ── Build ──────────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

add_header(doc)
add_footer(doc)

# ── Cover block ────────────────────────────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.alignment = 1  # center
cell = cover.rows[0].cells[0]
set_cell_bg(cell, "1a1a1a")
for side in ("top","left","bottom","right"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none")
    tcB.append(b); tcPr.append(tcB)

lp = cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.paragraph_format.space_before = Pt(18)
lr = lp.add_run(); lr.add_picture(LOGO_PATH, width=Inches(2.0))

tp = cell.add_paragraph("Positioning You As The Authority")
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(4)
tp.paragraph_format.space_after  = Pt(18)
tr = tp.runs[0]
tr.font.name = "Arial"; tr.font.size = Pt(10)
tr.font.italic = True; tr.font.color.rgb = ASG_BLUE

spacer(doc, 16)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run("LinkedIn Post — Ready to Publish")
t.font.name = "Arial"; t.font.size = Pt(22)
t.font.bold = True; t.font.color.rgb = ASG_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run("Melissa Doss Law  |  March 9, 2026  |  Get Over It Day")
s.font.name = "Arial"; s.font.size = Pt(12); s.font.color.rgb = ASG_CHARCOAL

rule(doc)

# ── Usage notes ───────────────────────────────────────────────────────────────
usage = doc.add_paragraph()
usage.paragraph_format.space_before = Pt(6)
usage.paragraph_format.space_after  = Pt(6)
usage.paragraph_format.left_indent  = Inches(0.2)
add_left_border(usage)
ur = usage.add_run(
    "Copy and paste this post directly into LinkedIn. Post today — March 9 is Get Over It Day. "
    "Use the headline as your first line. No hashtags required, but #NorthernKentucky #FamilyLaw "
    "#CollaborativeDivorce may be added at your discretion. LinkedIn profile must include your "
    "firm name and Florence, KY office address before publishing (see compliance note at end)."
)
ur.font.name = "Arial"; ur.font.size = Pt(10)
ur.font.color.rgb = ASG_BLUE; ur.font.bold = True

spacer(doc, 6)

# ── Recommended headline ──────────────────────────────────────────────────────
label(doc, "Recommended Headline (first line of post)")

hl = doc.add_paragraph()
hl.paragraph_format.space_before = Pt(4)
hl.paragraph_format.space_after  = Pt(8)
hr = hl.add_run(
    "If someone is telling you to fight for everything in your divorce,\n"
    "read this before you listen."
)
hr.font.name = "Arial"; hr.font.size = Pt(13)
hr.font.bold = True; hr.font.color.rgb = ASG_CHARCOAL

rule(doc)

# ── The post ──────────────────────────────────────────────────────────────────
label(doc, "Post Copy — Copy from here")

spacer(doc, 4)

post_lines = [
    ("Today is Get Over It Day.\n\nThe irony is not lost on me.", True, False, ASG_CHARCOAL, 11),
    ("", False, False, ASG_BODY, 6),
    ("Get Over It Day exists on one premise: holding on costs more than letting go.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("I've practiced family law in Northern Kentucky for 11 years.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("And I can tell you — the adversarial divorce process was not built to help you get over it.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("Here's what I've watched happen, more times than I can count:", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("A woman enters the process being told she needs to \"protect herself.\"\nShe's advised to document everything, disclose nothing, and outlast.", False, True, ASG_CHARCOAL, 11),
    ("", False, False, ASG_BODY, 4),
    ("Eighteen months later, she's still in it.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("Financially drained. Emotionally exhausted. And in a worse co-parenting relationship with the other parent of her children than when she started.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("She didn't end up here because she had a bad attorney.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("She ended up here because she was given a strategy designed for conflict — when what she actually needed was a path through it.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 6),
    ("Extended proceedings benefit a lot of people in the legal system.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("The person paying for them is not usually one of them.", False, True, ASG_CHARCOAL, 11),
    ("", False, False, ASG_BODY, 6),
    ("This is not an accusation. It's an observation.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("Some cases require aggressive advocacy. Some situations genuinely leave you no other choice.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("But many women I've spoken with over the years were never told that another option existed — one designed to resolve, rather than outlast.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("One that measures success by how quickly and cleanly you can get to the other side — with your finances, your children's stability, and your ability to co-parent still intact.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 6),
    ("Today is Get Over It Day.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("If you're somewhere in the middle of this process, or still deciding how to begin — the most important question you can ask isn't \"how do I win?\"", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("It's \"what process gives me the best chance of actually moving forward?\"", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("Those are very different questions.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 4),
    ("And they lead to very different outcomes.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 6),
    ("Save this for someone who needs to hear it. Or follow along — I share what I've learned over 11 years of family law every week.", False, False, ASG_BODY, 11),
    ("", False, False, ASG_BODY, 6),
    ("—", False, False, ASG_MID_GRAY, 11),
    ("", False, False, ASG_BODY, 2),
    ("This content is for informational purposes only and does not constitute legal advice. Every situation is different — consult a licensed attorney in your state before making legal decisions.", False, True, ASG_MID_GRAY, 9),
]

for text, bold, italic, color, size in post_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if text:
        run = p.add_run(text)
        run.font.name   = "Arial"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        run.font.color.rgb = color

rule(doc)

# ── Compliance note ───────────────────────────────────────────────────────────
label(doc, "Compliance Note — Action Required Before Publishing")

comp = doc.add_paragraph()
comp.paragraph_format.space_before = Pt(4)
comp.paragraph_format.space_after  = Pt(4)
comp.paragraph_format.left_indent  = Inches(0.2)
add_left_border(comp, color="58585a", sz="8")
cr = comp.add_run(
    "Kentucky SCR 3.130 requires all attorney advertising to include at least one "
    "attorney name and Kentucky office address. Your name is displayed on your LinkedIn "
    "profile automatically. Before publishing, confirm your LinkedIn profile bio or "
    "contact section displays: Melissa Doss Law | 7289 Burlington Pike, Florence, KY 41042. "
    "This is a one-time setup item."
)
cr.font.name = "Arial"; cr.font.size = Pt(10); cr.font.color.rgb = ASG_CHARCOAL
cr.font.italic = True

spacer(doc, 4)
body(doc,
    "ABA Compliance Review completed by Diana Voss, Legal Compliance Analyst — "
    "Authority Systems Group\u2122. Status: PASS WITH NOTES. No violations found in copy.",
    size=9, color=ASG_MID_GRAY, italic=True
)

# ── Save ───────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
