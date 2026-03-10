#!/usr/bin/env python3
"""
30-Day YouTube Script Pack PDF Generator — Hedges Dental
Authority Systems Group™ | Roger Bauer, Director
"""

import re
import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
ASG_FOOTER    = RGBColor(0x88, 0x88, 0x88)
ASG_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ASG_DARK_BG   = '1a1a1a'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_BLUE_HEX  = '25aae1'
ASG_NAVY_HEX  = '1F3A5F'

BASE = '/Users/Roger/Dropbox/code/authority-systems-group'
SLUG = 'hedges-dental'
DATE = '20260306'
ROOT = f'{BASE}/client-onboarding/clients/{SLUG}'
MD   = f'{ROOT}/outputs/{SLUG}_april-2026-youtube-scripts_{DATE}.md'
DOCX = f'{ROOT}/outputs/{SLUG}_april-2026-youtube-scripts_{DATE}.docx'
PDF  = f'{ROOT}/outputs/{SLUG}_april-2026-youtube-scripts_{DATE}.pdf'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def add_left_border(para, color=ASG_BLUE_HEX, sz='24', space='144'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_bottom_border(para, color=ASG_BLUE_HEX, sz='6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        'Authority Systems Group\u2122 \u2014 Confidential. '
        'Prepared exclusively for Dr. Jeffrey Hedges / Hedges Dental.'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = ASG_FOOTER
    run.font.name = 'Calibri'
    para.add_run('   |   Page ')
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    para.runs[-1]._r.append(fld)
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    r2 = OxmlElement('w:r')
    r2.append(instr)
    para._p.append(r2)
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    r3 = OxmlElement('w:r')
    r3.append(fld2)
    para._p.append(r3)

def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def add_cover_page(doc):
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0)
    section.left_margin = section.right_margin = Inches(0)

    def dark_line(text, size, bold=False, color=None, space_before=0, space_after=4):
        p = doc.add_paragraph()
        set_para_bg(p, ASG_DARK_BG)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color or ASG_WHITE
        return p

    for _ in range(8): dark_line('', 6)

    # YouTube red accent bar
    p_bar = doc.add_paragraph()
    set_para_bg(p_bar, ASG_DARK_BG)
    p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bar.paragraph_format.space_before = Pt(0)
    p_bar.paragraph_format.space_after  = Pt(0)
    r_bar = p_bar.add_run('\u2588' * 44)
    r_bar.font.color.rgb = RGBColor(0xCC, 0x22, 0x22)
    r_bar.font.size = Pt(6)

    dark_line('', 4)
    dark_line('30-DAY YOUTUBE SCRIPT PACK', 30, bold=True, space_before=0, space_after=2)

    p_rule = doc.add_paragraph()
    set_para_bg(p_rule, ASG_DARK_BG)
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after  = Pt(0)
    r = p_rule.add_run('\u2500' * 42)
    r.font.color.rgb = ASG_BLUE
    r.font.size = Pt(12)

    dark_line('', 8)
    dark_line('Hedges Dental', 28, bold=True)
    dark_line('Dr. Jeffrey Hedges, DMD', 16, color=RGBColor(0xCC, 0xCC, 0xCC))
    dark_line('Louisville, Kentucky', 14, color=RGBColor(0xAA, 0xAA, 0xAA))
    dark_line('', 8)

    p_accent = doc.add_paragraph()
    set_para_bg(p_accent, ASG_DARK_BG)
    p_accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_accent.add_run('\u25cf  \u25cf  \u25cf')
    r2.font.color.rgb = RGBColor(0xCC, 0x22, 0x22)
    r2.font.size = Pt(14)

    dark_line('', 8)
    dark_line('April 2026 \u2014 National Oral Health Month', 14, color=RGBColor(0xCC, 0xCC, 0xCC))
    dark_line('30 Long-Form Scripts  \u2022  B-Roll Cues  \u2022  Production Notes', 11, color=RGBColor(0xAA, 0xAA, 0xAA))
    dark_line('Value Bomb + Contrarian Variants Included', 11, color=RGBColor(0xAA, 0xAA, 0xAA))
    dark_line('', 8)
    dark_line('Prepared exclusively for Dr. Jeffrey Hedges', 11, color=RGBColor(0x88, 0x88, 0x88))
    dark_line('Confidential  |  Authority Systems Group\u2122', 10, color=RGBColor(0x66, 0x66, 0x66))
    dark_line('', 16)

    p_brand = doc.add_paragraph()
    set_para_bg(p_brand, ASG_DARK_BG)
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p_brand.add_run('Authority Systems Group\u2122')
    rb.font.name = 'Calibri'
    rb.font.size = Pt(13)
    rb.font.bold = True
    rb.font.color.rgb = ASG_BLUE

    dark_line('authoritysystemsgroup.com', 9, color=RGBColor(0x66, 0x66, 0x66))

    page_break(doc)
    s2 = doc.add_section()
    s2.top_margin = s2.bottom_margin = Inches(1.0)
    s2.left_margin = s2.right_margin = Inches(1.25)


def add_inline_runs(para, text):
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(1):
            r = para.add_run(m.group(1)); r.bold = True
        elif m.group(2):
            r = para.add_run(m.group(2)); r.italic = True
        elif m.group(3):
            r = para.add_run(m.group(3)); r.font.name = 'Courier New'; r.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def build_table(doc, lines):
    rows = [l for l in lines if l.startswith('|')]
    if not rows: return
    cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
    cells = [r for r in cells if not all(re.match(r'^[-:]+$', c) for c in r if c)]
    if not cells: return
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = 'Table Grid'
    for ri, row_data in enumerate(cells):
        is_header = ri == 0
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row_data[ci] if ci < len(row_data) else ''
            para = cell.paragraphs[0]
            para.clear()
            add_inline_runs(para, text)
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = ASG_WHITE
            if is_header:
                set_cell_bg(cell, ASG_NAVY_HEX)
            elif ri % 2 == 0:
                set_cell_bg(cell, ASG_LIGHT_BG)
    doc.add_paragraph()


def build_document(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    table_buffer = []

    def flush_table():
        if table_buffer:
            build_table(doc, table_buffer)
            table_buffer.clear()

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        # Code blocks
        if line.startswith('```'):
            i += 1
            block_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block_lines.append(lines[i].rstrip())
                i += 1
            if block_lines:
                p = doc.add_paragraph()
                set_para_bg(p, 'F0F0F0')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                for bl in block_lines:
                    run = p.add_run(bl + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = ASG_BODY_TEXT
            i += 1
            continue

        # Tables
        if line.startswith('|'):
            table_buffer.append(line); i += 1; continue
        else:
            flush_table()

        # Horizontal rules — page break before next VIDEO section
        if re.match(r'^-{3,}$', line):
            if i + 1 < len(lines) and lines[i+1].strip().startswith('## VIDEO'):
                page_break(doc)
            i += 1; continue

        # H1 skip
        if line.startswith('# ') and not line.startswith('## '):
            i += 1; continue

        # H2 — Video / section headers
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(14); r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            add_bottom_border(p)
            i += 1; continue

        # H3 — video title
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(12); r.font.bold = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1; continue

        # H4
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            i += 1; continue

        # H5
        if line.startswith('##### '):
            text = line[6:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.bold = True; r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1; continue

        # Italic-only lines — context/intro notes
        if line.startswith('*') and line.endswith('*') and not line.startswith('**') and len(line) > 2:
            text = line[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.left_indent  = Inches(0.2)
            set_para_bg(p, 'F8F8F8')
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1; continue

        # B-ROLL / PACING / CHAPTER inline cues
        if '[B-ROLL:' in line or '[PACING:' in line or '[CHAPTER:' in line:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_para_bg(p, 'E8F4F8')
            add_left_border(p, color='25aae1', sz='12', space='80')
            add_inline_runs(p, line)
            for run in p.runs:
                run.font.name = 'Calibri'; run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1F, 0x6E, 0x9C)
            i += 1; continue

        # PRODUCTION NOTES / VALUE BOMB / CONTRARIAN section headers
        if line.startswith('**PRODUCTION NOTES') or line.startswith('**VALUE BOMB') or line.startswith('**CONTRARIAN'):
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            set_para_bg(p, 'EFEFEF')
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(10.5); r.font.bold = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1; continue

        # Blockquotes
        if line.startswith('> '):
            bq_lines = [line[2:]]
            while i + 1 < len(lines) and lines[i+1].strip().startswith('> '):
                i += 1; bq_lines.append(lines[i].strip()[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            set_para_bg(p, 'EAF6FC')
            add_left_border(p, color=ASG_BLUE_HEX, sz='18', space='100')
            add_inline_runs(p, ' '.join(bq_lines))
            for run in p.runs:
                run.font.name = 'Georgia'; run.font.size = Pt(9.5); run.font.italic = True
                run.font.color.rgb = ASG_CHARCOAL
            i += 1; continue

        # Bullets
        if re.match(r'^[-*]\s+', line):
            text = re.sub(r'^[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline_runs(p, text)
            for run in p.runs:
                run.font.name = 'Calibri'; run.font.size = Pt(10.5)
                run.font.color.rgb = ASG_BODY_TEXT
            i += 1; continue

        # Numbered lists
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline_runs(p, text)
            for run in p.runs:
                run.font.name = 'Calibri'; run.font.size = Pt(10.5)
                run.font.color.rgb = ASG_BODY_TEXT
            i += 1; continue

        # Empty lines
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(4)
            i += 1; continue

        # Body paragraph (script spoken text)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        add_inline_runs(p, line)
        for run in p.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)
            run.font.color.rgb = ASG_BODY_TEXT
        i += 1

    flush_table()


def convert_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert
        print('Converting to PDF via docx2pdf...')
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            print(f'PDF created: {pdf_path}')
            return True
    except Exception as e:
        print(f'docx2pdf failed: {e}')
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', os.path.dirname(pdf_path), docx_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            print(f'PDF created via LibreOffice: {pdf_path}'); return True
    except Exception as e:
        print(f'LibreOffice failed: {e}')
    print(f'\nManual step: open {docx_path} in Word and export as PDF.')
    return False


def main():
    print(f'Reading markdown: {MD}')
    with open(MD, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    add_cover_page(doc)
    add_footer(doc)
    build_document(doc, md_text)

    doc.save(DOCX)
    print(f'DOCX saved: {DOCX}')
    convert_to_pdf(DOCX, PDF)


if __name__ == '__main__':
    main()
