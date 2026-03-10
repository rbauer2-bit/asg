#!/usr/bin/env python3
"""
Marketing Audit PDF Generator — Hedges Dental
Authority Systems Group™ | Victor Hale, Director of Digital Intelligence
Produces a branded DOCX then converts to PDF.
"""

import re
import os
import subprocess
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ──────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
ASG_FOOTER    = RGBColor(0x88, 0x88, 0x88)
ASG_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ASG_RED       = RGBColor(0xCC, 0x33, 0x33)   # critical score colour
ASG_DARK_BG   = '1a1a1a'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_BLUE_HEX  = '25aae1'
ASG_GRAY_HEX  = 'CCCCCC'
ASG_SCORE_BG  = '1F3A5F'   # deep navy for score banner

# ── File Paths ────────────────────────────────────────────────────────────────
BASE = '/Users/Roger/Dropbox/code/authority-systems-group'
SLUG = 'hedges-dental'
DATE = '20260306'
ROOT = f'{BASE}/client-onboarding/clients/{SLUG}'
MD   = f'{ROOT}/outputs/{SLUG}_marketing-audit_{DATE}.md'
DOCX = f'{ROOT}/outputs/{SLUG}_marketing-audit_{DATE}.docx'
PDF  = f'{ROOT}/outputs/{SLUG}_marketing-audit_{DATE}.pdf'


# ── XML Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def add_left_border(para, color=ASG_BLUE_HEX, sz='24', space='144'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_bottom_border(para, color=ASG_BLUE_HEX, sz='6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_line_spacing(para, multiplier=1.4):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), str(int(multiplier * 240)))
    sp.set(qn('w:lineRule'), 'auto')


def set_indent(para, left_inches=0.0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(left_inches * 1440)))
    pPr.append(ind)


def add_page_number_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ── Inline Formatting ─────────────────────────────────────────────────────────

def add_inline_runs(para, text, base_font='Georgia', base_size=11, base_color=None):
    """Parse **bold** and *italic* markdown and add styled runs to para."""
    if base_color is None:
        base_color = ASG_BODY_TEXT
    parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
            run.font.color.rgb = ASG_CHARCOAL
        else:
            run = para.add_run(part)
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color


def strip_inline(text):
    """Strip markdown inline markers — return plain text."""
    text = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+?)\*', r'\1', text)
    text = re.sub(r'`([^`]+?)`', r'\1', text)
    return text


# ── Cover Page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    """Marketing Audit branded cover — dark with score callout."""

    def cover_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        set_para_bg(p, ASG_DARK_BG)
        return p

    def cover_run(para, text, size, bold=False, italic=False, color=None):
        if color is None:
            color = ASG_WHITE
        r = para.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        return r

    # Top spacers
    for _ in range(3):
        cover_para(space_before=0, space_after=0)

    # ASG wordmark
    p = cover_para(space_before=36, space_after=4)
    cover_run(p, 'AUTHORITY SYSTEMS GROUP\u2122', 13, bold=True, color=ASG_BLUE)

    # Tagline
    p = cover_para(space_before=2, space_after=8)
    cover_run(p, 'Digital Intelligence Division', 10,
              italic=True, color=RGBColor(0x88, 0x88, 0x88))

    # Report title
    p = cover_para(space_before=24, space_after=4)
    cover_run(p, 'MARKETING AUDIT', 34, bold=True)

    # Subtitle
    p = cover_para(space_before=4, space_after=4)
    cover_run(p, 'Digital Authority Assessment', 15, color=ASG_BLUE)

    # Divider spacer
    p = cover_para(space_before=16, space_after=16)
    cover_run(p, '\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014',
              12, color=RGBColor(0x44, 0x44, 0x44))

    # Client name
    p = cover_para(space_before=10, space_after=4)
    cover_run(p, 'Hedges Dental', 26, bold=True)

    # Dr. Hedges name
    p = cover_para(space_before=4, space_after=4)
    cover_run(p, 'Dr. Jeffrey Hedges, DMD', 14, color=RGBColor(0xCC, 0xCC, 0xCC))

    # Location
    p = cover_para(space_before=4, space_after=16)
    cover_run(p, '2015 Herr Lane Suite D  \u00b7  Louisville, KY 40222',
              11, color=RGBColor(0x88, 0x88, 0x88))

    # Score banner
    p = cover_para(space_before=18, space_after=4)
    cover_run(p, 'OVERALL MARKETING SCORE', 10, bold=True,
              color=RGBColor(0xAA, 0xAA, 0xAA))

    p = cover_para(space_before=2, space_after=2)
    r = cover_run(p, '22', 52, bold=True, color=RGBColor(0xFF, 0x55, 0x55))
    p2 = cover_para(space_before=0, space_after=2)
    cover_run(p2, 'out of 100  \u2014  Grade: F  \u2014  Critical',
              12, color=RGBColor(0xFF, 0x88, 0x88))

    # Spacers to push attribution to bottom
    for _ in range(6):
        cover_para(space_before=0, space_after=0)

    # Prepared by
    p = cover_para(space_before=12, space_after=4)
    cover_run(p, 'Prepared by: Victor Hale, Director of Digital Intelligence',
              10, italic=True, color=RGBColor(0x77, 0x77, 0x77))

    p = cover_para(space_before=2, space_after=4)
    cover_run(p, 'Directed by Roger Bauer  \u00b7  Authority Systems Group\u2122',
              10, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    # Date + confidential
    p = cover_para(space_before=8, space_after=6)
    cover_run(p, 'March 2026  \u00b7  CONFIDENTIAL \u2014 PREPARED EXCLUSIVELY FOR DR. JEFFREY HEDGES',
              9, color=RGBColor(0x44, 0x44, 0x44))

    doc.add_page_break()


# ── Footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run(
            'Authority Systems Group\u2122 \u2014 Confidential. '
            'Prepared exclusively for Dr. Jeffrey Hedges / Hedges Dental.'
        )
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9)
        r1.font.color.rgb = ASG_FOOTER
        fp.add_run('\t')
        r2 = fp.add_run()
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.color.rgb = ASG_FOOTER
        add_page_number_field(r2)


# ── Table Builder ─────────────────────────────────────────────────────────────

def build_table(doc, rows_raw):
    """Convert markdown table lines into a styled Word table."""
    rows = []
    for line in rows_raw:
        if re.match(r'^\|[\s\-|:]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    for ri, row_cells in enumerate(rows):
        tbl_row = tbl.rows[ri]
        is_header = (ri == 0)
        is_total_row = (
            len(row_cells) > 0 and
            '**TOTAL**' in row_cells[0] or
            ('Total' in row_cells[0] and row_cells[0].startswith('**'))
        )

        for ci in range(ncols):
            cell_text = row_cells[ci] if ci < len(row_cells) else ''
            cell = tbl_row.cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            plain = strip_inline(cell_text)

            if is_header:
                r = p.add_run(plain)
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = ASG_WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(cell, ASG_BLUE_HEX)
            elif is_total_row:
                r = p.add_run(plain)
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = ASG_WHITE
                set_cell_bg(cell, '2D6A8A')
            else:
                add_inline_runs(p, cell_text, base_font='Calibri', base_size=9.5)
                if ri % 2 == 0:
                    set_cell_bg(cell, ASG_LIGHT_BG)

    doc.add_paragraph()


# ── Metadata Header Block ─────────────────────────────────────────────────────

def add_meta_block(doc, lines):
    """Render the **Key:** Value metadata lines at top of report."""
    for line in lines:
        s = line.strip()
        if not s or s == '---':
            break
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        set_line_spacing(p, 1.2)
        add_inline_runs(p, s, base_font='Calibri', base_size=10)


# ── Score Badge ───────────────────────────────────────────────────────────────

def add_score_badge(doc, score_text):
    """Render the overall score as a highlighted banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    set_para_bg(p, ASG_SCORE_BG)
    r = p.add_run(f'  {score_text}  ')
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = ASG_WHITE


# ── Main Parser & Builder ─────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    add_cover_page(doc)

    with open(MD, 'r', encoding='utf-8') as f:
        lines = [ln.rstrip('\n') for ln in f.readlines()]

    table_buf  = []
    in_table   = False
    in_code    = False
    skip_first = True   # skip the opening # H1 line (already on cover)

    i = 0
    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        # ── Skip the H1 title line ─────────────────────────────────────────
        if skip_first and re.match(r'^# (?!#)', s):
            skip_first = False
            i += 1
            continue

        # ── Code block (```) — skip entire block ──────────────────────────
        if s.startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        # ── Table collection ───────────────────────────────────────────────
        if s.startswith('|') and s.endswith('|'):
            table_buf.append(s)
            in_table = True
            i += 1
            continue
        if in_table:
            build_table(doc, table_buf)
            table_buf = []
            in_table  = False
            continue  # reprocess current line

        # ── Horizontal rule ────────────────────────────────────────────────
        if s == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            add_bottom_border(p, color=ASG_BLUE_HEX, sz='6')
            i += 1
            continue

        # ── H2 (##) ────────────────────────────────────────────────────────
        if re.match(r'^## (?!#)', s):
            text = strip_inline(s[3:])
            doc.add_page_break()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(20)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            add_bottom_border(p, color=ASG_BLUE_HEX, sz='8')
            i += 1
            continue

        # ── H3 (###) ───────────────────────────────────────────────────────
        if re.match(r'^### (?!#)', s):
            text = strip_inline(s[4:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(13)
            r.font.bold  = True
            r.font.color.rgb = ASG_CHARCOAL
            add_bottom_border(p, color='CCCCCC', sz='4')
            i += 1
            continue

        # ── H4 (####) ──────────────────────────────────────────────────────
        if re.match(r'^#### (?!#)', s):
            text = strip_inline(s[5:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(11)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            i += 1
            continue

        # ── Blockquote (>) ─────────────────────────────────────────────────
        if s.startswith('> '):
            text  = s[2:]
            plain = strip_inline(text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(10)
            set_indent(p, left_inches=0.25)
            set_para_bg(p, ASG_LIGHT_BG)
            add_left_border(p, color=ASG_BLUE_HEX, sz='24', space='144')
            r = p.add_run(plain)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────────────
        if re.match(r'^[-*]\s+', s):
            text = re.sub(r'^[-*]\s+', '', s)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_line_spacing(p, 1.3)
            add_inline_runs(p, text, base_font='Georgia', base_size=11)
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────────
        if re.match(r'^\d+\.\s+', s):
            text = re.sub(r'^\d+\.\s+', '', s)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            set_line_spacing(p, 1.3)
            add_inline_runs(p, text, base_font='Georgia', base_size=11)
            i += 1
            continue

        # ── Overall score line — special banner treatment ──────────────────
        if s.startswith('**Overall Marketing Score:'):
            plain = strip_inline(s)
            add_score_badge(doc, plain)
            i += 1
            continue

        # ── Bold metadata lines at top (**Key:** Value) ────────────────────
        if re.match(r'^\*\*[A-Za-z].*\*\*:', s) or re.match(r'^\*\*[A-Za-z].*:\*\*', s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            set_line_spacing(p, 1.2)
            add_inline_runs(p, s, base_font='Calibri', base_size=10,
                           base_color=ASG_CHARCOAL)
            i += 1
            continue

        # ── All-bold standalone label (**LABEL**) ──────────────────────────
        if re.match(r'^\*\*[^*]+\*\*$', s):
            text = s[2:-2]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(12)
            r.font.bold  = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── All-italic standalone line ─────────────────────────────────────
        if (re.match(r'^\*[^*].+[^*]\*$', s) or
                (s.startswith('*') and s.endswith('*') and not s.startswith('**'))):
            text = s[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            set_indent(p, left_inches=0.25)
            r = p.add_run(text)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(10)
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            i += 1
            continue

        # ── Empty line ─────────────────────────────────────────────────────
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ── Default body paragraph ─────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        set_line_spacing(p, 1.4)
        add_inline_runs(p, s, base_font='Georgia', base_size=11)
        i += 1

    # Flush any trailing table
    if table_buf:
        build_table(doc, table_buf)

    add_footer(doc)

    doc.core_properties.title   = 'Marketing Audit \u2014 Hedges Dental'
    doc.core_properties.subject = 'Digital Authority Assessment'
    doc.core_properties.author  = 'Authority Systems Group\u2122'

    os.makedirs(os.path.dirname(DOCX), exist_ok=True)
    doc.save(DOCX)
    size_kb = os.path.getsize(DOCX) // 1024
    print(f'DOCX saved: {DOCX}  ({size_kb} KB)')
    return DOCX


# ── PDF Conversion ────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF. Tries docx2pdf first, then LibreOffice."""

    # ── Option 1: docx2pdf (macOS native Word) ─────────────────────────────
    try:
        from docx2pdf import convert
        print('Converting to PDF via docx2pdf...')
        convert(docx_path, PDF)
        if os.path.exists(PDF):
            size_kb = os.path.getsize(PDF) // 1024
            print(f'PDF saved:  {PDF}  ({size_kb} KB)')
            return True
    except ImportError:
        print('docx2pdf not installed — trying LibreOffice...')
    except Exception as e:
        print(f'docx2pdf failed ({e}) — trying LibreOffice...')

    # ── Option 2: LibreOffice ──────────────────────────────────────────────
    libreoffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/bin/soffice',
        '/usr/local/bin/soffice',
    ]
    soffice = None
    for path in libreoffice_paths:
        if os.path.exists(path):
            soffice = path
            break

    if soffice:
        output_dir = os.path.dirname(docx_path)
        print(f'Converting to PDF via LibreOffice...')
        result = subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf',
             '--outdir', output_dir, docx_path],
            capture_output=True, text=True
        )
        # LibreOffice names the output file with .pdf extension
        lo_pdf = docx_path.replace('.docx', '.pdf')
        if os.path.exists(lo_pdf):
            size_kb = os.path.getsize(lo_pdf) // 1024
            print(f'PDF saved:  {lo_pdf}  ({size_kb} KB)')
            return True
        else:
            print(f'LibreOffice stdout: {result.stdout}')
            print(f'LibreOffice stderr: {result.stderr}')

    print('\nNeither docx2pdf nor LibreOffice is available.')
    print('To install docx2pdf:  pip install docx2pdf')
    print('To install LibreOffice: https://www.libreoffice.org/download/')
    print(f'\nThe DOCX is ready at: {docx_path}')
    print('You can open it in Word and export as PDF manually.')
    return False


# ── Entry Point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('=== Hedges Dental — Marketing Audit PDF Generator ===')
    print('Authority Systems Group™ | Digital Intelligence Division\n')

    docx_path = build_document()
    success   = convert_to_pdf(docx_path)

    print('\n=== COMPLETE ===')
    if success:
        print(f'DOCX: {DOCX}')
        print(f'PDF:  {PDF}')
    else:
        print(f'DOCX: {DOCX}')
        print('PDF:  conversion failed — see notes above')
