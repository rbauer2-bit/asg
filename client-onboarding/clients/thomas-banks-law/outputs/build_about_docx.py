#!/usr/bin/env python3
"""
Authority Systems Group™
Thomas E. Banks II — About Pages (KY + IN)
KentuckianaDivorceLawyer.com + KentuckianaFamilyLawyer.com
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT    = RGBColor(0x1A, 0x25, 0x33)
GOLD         = RGBColor(0xC9, 0xA0, 0x2C)
META_COLOR   = RGBColor(0x55, 0x55, 0x55)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
TINT_BG       = 'EAF6FC'
GOLD_HEX      = 'C9A02C'
GOLD_TINT     = 'FDF8EC'
KY_COLOR_HEX  = '154360'   # deep Kentucky blue
IN_COLOR_HEX  = '1B5E3B'   # deep Indiana green
CRED_BG       = 'F7F9FC'
COMPLIANCE_BG = 'FFF8E7'
IMPL_BG       = 'F0F4F8'

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
CLIENT    = os.path.join(BASE, 'client-onboarding/clients/thomas-banks-law')
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
MD_PATH   = os.path.join(CLIENT, 'outputs/md/thomas-banks-law_about-pages_20260308.md')
OUT_DOCX  = os.path.join(CLIENT, 'outputs/docx/thomas-banks-law_about-pages_20260308.docx')


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tblBorders.append(el)
    tblPr.append(tblBorders)

def add_pb(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)

def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(size)

def divider(doc, color='CCCCCC'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'), color); pBdr.append(bt); pPr.append(pBdr)

def add_page_number(para):
    run = para.add_run(); run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin',None),(None,' PAGE '),('end',None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)

def add_inline(para, text, base_size=10.5, base_color=BODY_COLOR, base_bold=False, base_italic=False):
    """Render **bold** and *italic* inline markers within a paragraph."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            r = para.add_run(inner)
            r.bold = True; r.italic = base_italic
            r.font.size = Pt(base_size); r.font.color.rgb = base_color; r.font.name = 'Calibri'
        elif part.startswith('*') and part.endswith('*'):
            inner = part[1:-1]
            r = para.add_run(inner)
            r.bold = base_bold; r.italic = True
            r.font.size = Pt(base_size); r.font.color.rgb = base_color; r.font.name = 'Calibri'
        else:
            if part:
                r = para.add_run(part)
                r.bold = base_bold; r.italic = base_italic
                r.font.size = Pt(base_size); r.font.color.rgb = base_color; r.font.name = 'Calibri'


# ── Header / Footer ────────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0,0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0,1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('ABOUT PAGES  |  THOMAS E. BANKS II  |  MARCH 2026')
    rr.font.size = Pt(7.5); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True
    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0,0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Confidential. Prepared exclusively for Thomas E. Banks II.')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0,1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover ──────────────────────────────────────────────────────────────────────

def build_cover(doc):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(40); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))
    sp(doc, 24)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('AUTHORITY SYSTEMS GROUP™  —  CONFIDENTIAL')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)
    tr_ = tp.add_run('About Pages')
    tr_.font.size = Pt(30); tr_.bold = True; tr_.font.color.rgb = DARK_TEXT; tr_.font.name = 'Calibri'

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(0); sp2.paragraph_format.space_after = Pt(4)
    sr = sp2.add_run('Thomas E. Banks II, AAML Fellow  |  Family Law Attorney  |  Kentucky & Indiana')
    sr.font.size = Pt(13); sr.font.color.rgb = ASG_CHARCOAL; sr.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Prepared: March 8, 2026     |     KentuckianaDivorceLawyer.com  +  KentuckianaFamilyLawyer.com')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 16)

    # Credential bar
    cred_table = doc.add_table(1, 4); remove_table_borders(cred_table)
    for i in range(4): cred_table.columns[i].width = Inches(1.625)
    creds = [
        ('AAML Fellow', 'One of Fewer Than 1,600 Worldwide'),
        ('Super Lawyers', '13 Consecutive Years of Recognition'),
        ('AV Preeminent', 'Martindale-Hubbell Highest Rating'),
        ('KY + IN', 'Dual-State Licensed'),
    ]
    for ci, (title, sub) in enumerate(creds):
        c = cred_table.cell(0, ci); set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8); cp.paragraph_format.space_after = Pt(8)
        r1 = cp.add_run(title + '\n'); r1.bold = True; r1.font.size = Pt(10)
        r1.font.color.rgb = GOLD; r1.font.name = 'Calibri'
        r2 = cp.add_run(sub); r2.font.size = Pt(8); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'

    sp(doc, 16)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(4)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  Two fully differentiated About pages — one for each personal site. '
        'The Kentucky page is built around Jefferson County Family Court presence and the Louisville '
        'competitive landscape. The Indiana page leads with the dual-licensure gap and Southern Indiana '
        'market access. Both apply the Authority Conversion Protocol™. ABA compliant throughout.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    sp(doc, 12)

    # Page index
    pages_idx = doc.add_table(2, 3); remove_table_borders(pages_idx)
    pages_idx.columns[0].width = Inches(0.5); pages_idx.columns[1].width = Inches(4.0)
    pages_idx.columns[2].width = Inches(2.0)
    rows_data = [
        ('KY', 'About Page — KentuckianaDivorceLawyer.com', 'Kentucky Market  |  ~1,050 words'),
        ('IN', 'About Page — KentuckianaFamilyLawyer.com',  'Southern Indiana Market  |  ~1,050 words'),
    ]
    for i, (label, title, note) in enumerate(rows_data):
        bg_hex = 'EAF6FC' if i == 0 else 'E8F5E9'
        c0 = pages_idx.cell(i,0); set_cell_bg(c0, KY_COLOR_HEX if i==0 else IN_COLOR_HEX)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(8); p0.paragraph_format.space_after = Pt(8)
        r0 = p0.add_run(label); r0.bold = True; r0.font.size = Pt(11)
        r0.font.color.rgb = WHITE; r0.font.name = 'Calibri'
        c1 = pages_idx.cell(i,1); set_cell_bg(c1, bg_hex)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_before = Pt(8); p1.paragraph_format.space_after = Pt(8)
        r1 = p1.add_run(title); r1.font.size = Pt(11); r1.bold = True
        r1.font.color.rgb = DARK_TEXT; r1.font.name = 'Calibri'
        c2 = pages_idx.cell(i,2); set_cell_bg(c2, bg_hex)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(8); p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(note); r2.font.size = Pt(8.5); r2.font.color.rgb = ASG_CHARCOAL; r2.font.name = 'Calibri'

    add_pb(doc)


# ── Document Parser ─────────────────────────────────────────────────────────────

def parse_about_doc(md_path):
    """
    Returns:
      ky_page   — dict: {site_url, content_lines}
      in_page   — dict: {site_url, content_lines}
      impl_lines — list of strings
      aba_lines  — list of strings
    """
    with open(md_path, 'r') as f:
        raw = f.read()

    # Split off KY section
    ky_split = re.split(r'\n# ABOUT PAGE — KENTUCKY\b', raw, maxsplit=1)
    rest_after_ky = ky_split[1] if len(ky_split) > 1 else ''

    # Split KY from IN
    in_split = re.split(r'\n# ABOUT PAGE — INDIANA\b', rest_after_ky, maxsplit=1)
    ky_raw = in_split[0]
    rest_after_in = in_split[1] if len(in_split) > 1 else ''

    # Split IN from Implementation Notes
    impl_split = re.split(r'\n## IMPLEMENTATION NOTES\b', rest_after_in, maxsplit=1)
    in_raw = impl_split[0]
    rest_after_impl = impl_split[1] if len(impl_split) > 1 else ''

    # Split Implementation Notes from ABA section
    aba_split = re.split(r'\n## ABA COMPLIANCE NOTES\b', rest_after_impl, maxsplit=1)
    impl_raw = aba_split[0]
    aba_raw  = aba_split[1] if len(aba_split) > 1 else ''

    def parse_page(raw_text):
        lines = raw_text.strip().split('\n')
        # First line should be ## URL
        site_url = ''
        start = 0
        for i, line in enumerate(lines):
            if line.startswith('## ') and ('KentuckianaD' in line or 'KentuckianaF' in line):
                site_url = line[3:].strip()
                start = i + 1
                break
        return {'site_url': site_url, 'content_lines': lines[start:]}

    ky_page  = parse_page(ky_raw)
    in_page  = parse_page(in_raw)
    impl_lines = [l for l in impl_raw.strip().split('\n') if l.strip() and not l.strip() == '---']
    aba_lines  = [l for l in aba_raw.strip().split('\n')  if l.strip() and not l.strip() == '---']

    return ky_page, in_page, impl_lines, aba_lines


# ── Page Renderer ──────────────────────────────────────────────────────────────

def render_page_banner(doc, state_label, site_url, color_hex):
    """Render the top banner for an About page section."""
    bt = doc.add_table(1, 2); remove_table_borders(bt)
    bt.columns[0].width = Inches(0.55); bt.columns[1].width = Inches(5.95)

    nc = bt.cell(0, 0); set_cell_bg(nc, GOLD_HEX)
    np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(8); np_.paragraph_format.space_after = Pt(8)
    nr = np_.add_run(state_label)
    nr.bold = True; nr.font.size = Pt(12); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

    tc_ = bt.cell(0, 1); set_cell_bg(tc_, color_hex)
    tp = tc_.paragraphs[0]
    tp.paragraph_format.space_before = Pt(8); tp.paragraph_format.space_after = Pt(6)
    tr_ = tp.add_run(f'About Page  —  {site_url}')
    tr_.bold = True; tr_.font.size = Pt(11); tr_.font.name = 'Calibri'
    tr_.font.color.rgb = WHITE

    sp(doc, 8)


def is_credential_block_start(line):
    """Detect the start of the credential block section."""
    return '### The Complete Credential Record' in line


def render_credential_block(doc, lines):
    """Render the credential bullet list as a styled two-column table."""
    # Collect credentials: lines starting with -
    creds = []
    name_line = ''
    firm_line = ''
    address_line = ''

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('**') and stripped.endswith('**') and 'Thomas E. Banks' in stripped:
            name_line = stripped.strip('*')
        elif stripped.startswith('*') and stripped.endswith('*') and 'Straw' in stripped:
            firm_line = stripped.strip('*')
        elif stripped.startswith('*') and stripped.endswith('*') and 'West Market' in stripped:
            address_line = stripped.strip('*')
        elif stripped.startswith('- '):
            creds.append(stripped[2:])

    # Name / firm header
    sp(doc, 4)
    hp = doc.add_paragraph(); set_para_bg(hp, ASG_DARK_HEX)
    hp.paragraph_format.space_before = Pt(6); hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.left_indent = Inches(0.1)
    if name_line:
        nr = hp.add_run(name_line)
        nr.bold = True; nr.font.size = Pt(11); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'
    if firm_line:
        hp.add_run('  |  ')._._element if False else None
        fr = hp.add_run(f'  |  {firm_line}')
        fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0xBB, 0xCC, 0xDD); fr.font.name = 'Calibri'

    # Address sub-line
    if address_line:
        ap = doc.add_paragraph(); set_para_bg(ap, CHARCOAL_HEX)
        ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
        ap.paragraph_format.left_indent = Inches(0.1)
        ar = ap.add_run(address_line)
        ar.font.size = Pt(8.5); ar.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC); ar.font.name = 'Calibri'

    # Credentials table (2-column)
    if creds:
        mid = (len(creds) + 1) // 2
        col1 = creds[:mid]; col2 = creds[mid:]
        rows = max(len(col1), len(col2))
        tbl = doc.add_table(rows, 2)
        tbl.style = 'Table Grid'
        tbl.columns[0].width = Inches(3.25); tbl.columns[1].width = Inches(3.25)

        for ri in range(rows):
            bg = CRED_BG if ri % 2 == 0 else 'FFFFFF'
            for ci, col_list in enumerate([col1, col2]):
                cell = tbl.rows[ri].cells[ci]; set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
                if ri < len(col_list):
                    add_inline(p, col_list[ri], base_size=9.5, base_color=BODY_COLOR)
                else:
                    p.add_run('')

    sp(doc, 8)


def render_about_page(doc, page_data, state_label, color_hex):
    render_page_banner(doc, state_label, page_data['site_url'], color_hex)

    lines = page_data['content_lines']
    in_cred_block = False
    cred_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect credential block
        if is_credential_block_start(stripped):
            in_cred_block = True
            cred_lines = []
            # Render the H3 header for the credential block
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(14); hp.paragraph_format.space_after = Pt(2)
            hr = hp.add_run('The Complete Credential Record')
            hr.bold = True; hr.font.size = Pt(11.5)
            hr.font.color.rgb = DARK_TEXT; hr.font.name = 'Calibri'
            i += 1
            continue

        if in_cred_block:
            if stripped.startswith('### ') or stripped.startswith('## ') or stripped.startswith('# '):
                # End of cred block — render it, then process this line
                render_credential_block(doc, cred_lines)
                in_cred_block = False
                cred_lines = []
                # fall through to normal rendering below
            else:
                cred_lines.append(line)
                i += 1
                continue

        # Flush cred block if we reach end of content
        if in_cred_block and i == len(lines) - 1:
            cred_lines.append(line)
            render_credential_block(doc, cred_lines)
            break

        # Normal line rendering
        if not stripped:
            sp(doc, 4)
        elif stripped == '---':
            divider(doc, 'DDDDDD')
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            # H2 = page opening headline — large, prominent
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(8)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
        elif stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 120:
            # Bold-only short line = bold action line (phone/email)
            text = stripped.strip('*')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ASG_BLUE; r.font.name = 'Calibri'
        elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            # Italic-only line = attribution or sub-note
            text = stripped.strip('*')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = META_COLOR; r.font.name = 'Calibri'
        elif stripped.startswith('- '):
            # Bullet
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            bullet_r = p.add_run('• ')
            bullet_r.bold = True; bullet_r.font.size = Pt(10); bullet_r.font.color.rgb = ASG_BLUE
            bullet_r.font.name = 'Calibri'
            add_inline(p, text, base_size=10.5, base_color=BODY_COLOR)
        else:
            # Body paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            add_inline(p, stripped, base_size=10.5, base_color=BODY_COLOR)

        i += 1

    add_pb(doc)


# ── Implementation Notes ──────────────────────────────────────────────────────

def render_impl_section(doc, impl_lines):
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(8)
    set_para_bg(bp, ASG_DARK_HEX)
    br_ = bp.add_run('  IMPLEMENTATION NOTES')
    br_.bold = True; br_.font.size = Pt(11); br_.font.color.rgb = WHITE; br_.font.name = 'Calibri'

    for line in impl_lines:
        stripped = line.strip()
        if not stripped:
            sp(doc, 4)
            continue
        if stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped.strip('*'))
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
        elif stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            set_para_bg(p, IMPL_BG)
            add_inline(p, f'  • {text}', base_size=10, base_color=BODY_COLOR)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            add_inline(p, stripped, base_size=10, base_color=BODY_COLOR)


# ── ABA Compliance Section ─────────────────────────────────────────────────────

def render_aba_section(doc, aba_lines):
    add_pb(doc)
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(8)
    set_para_bg(bp, GOLD_HEX)
    br_ = bp.add_run('  ABA / STATE BAR COMPLIANCE NOTES')
    br_.bold = True; br_.font.size = Pt(11); br_.font.color.rgb = WHITE; br_.font.name = 'Calibri'

    for line in aba_lines:
        stripped = line.strip()
        if not stripped:
            sp(doc, 4)
            continue
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph(); set_para_bg(p, COMPLIANCE_BG)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, text, base_size=9.5, base_color=BODY_COLOR, base_italic=True)
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped.strip('*'))
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
        elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(stripped.strip('*'))
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = META_COLOR; r.font.name = 'Calibri'
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            add_inline(p, stripped, base_size=10, base_color=BODY_COLOR)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_COLOR

    add_header_footer(doc)
    build_cover(doc)

    ky_page, in_page, impl_lines, aba_lines = parse_about_doc(MD_PATH)

    render_about_page(doc, ky_page, 'KY', KY_COLOR_HEX)
    render_about_page(doc, in_page, 'IN', IN_COLOR_HEX)

    render_impl_section(doc, impl_lines)
    render_aba_section(doc, aba_lines)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
