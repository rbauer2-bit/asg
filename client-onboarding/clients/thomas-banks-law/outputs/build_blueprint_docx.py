#!/usr/bin/env python3
"""
Authority Systems Group™
Thomas E. Banks II — Authority Blueprint™
All 13 Sections
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ─────────────────────────────────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT    = RGBColor(0x1A, 0x25, 0x33)
GOLD         = RGBColor(0xC9, 0xA0, 0x2C)

ASG_BLUE_HEX = '25AAE1'
ASG_DARK_HEX = '1A2533'
LIGHT_BG     = 'F5F5F5'
TINT_BG      = 'EAF6FC'
GOLD_HEX     = 'C9A02C'
GOLD_TINT    = 'FDF8EC'
SLATE_HEX    = '2C3E50'
GREEN_HEX    = '27AE60'
AMBER_HEX    = 'E67E22'

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
CLIENT    = os.path.join(BASE, 'client-onboarding/clients/thomas-banks-law')
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
MD_PATH   = os.path.join(CLIENT, 'outputs/md/thomas-banks-law_authority-blueprint_20260308.md')
OUT_DOCX  = os.path.join(CLIENT, 'outputs/docx/thomas-banks-law_authority-blueprint_20260308.docx')

SECTION_COLORS = {
    1: '1A2533', 2: '1A2533', 3: '2C3E50', 4: '4A235A',
    5: '1A4A2C', 6: '7D2400', 7: '1A2533', 8: '25438A',
    9: '005A5A', 10: '2C3E50', 11: '5A3A00', 12: '1A2533', 13: '27AE60',
}


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tblBorders.append(el)
    tblPr.append(tblBorders)

def add_pb(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)

def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(size)

def add_page_number(para):
    run = para.add_run(); run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin',None),(None,' PAGE '),('end',None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)

def divider(doc, color='CCCCCC'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'), color); pBdr.append(bt); pPr.append(pBdr)


# ── Inline Text Renderer ───────────────────────────────────────────────────────

def add_inline(para, text, size=10.5, color=BODY_COLOR, base_bold=False, base_italic=False):
    """Render text with **bold** and *italic* markers inline."""
    # Split on bold first, then italic within segments
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            inner = part[2:-2]
            sub = re.split(r'(\*[^*]+\*)', inner)
            for sp_ in sub:
                if sp_.startswith('*') and sp_.endswith('*') and len(sp_) > 2:
                    r = para.add_run(sp_[1:-1]); r.bold = True; r.italic = True
                else:
                    r = para.add_run(sp_); r.bold = True; r.italic = base_italic
                r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = 'Calibri'
        else:
            # Check for italic *text*
            sub = re.split(r'(\*[^*]+\*)', part)
            for sp_ in sub:
                if sp_.startswith('*') and sp_.endswith('*') and len(sp_) > 2:
                    r = para.add_run(sp_[1:-1]); r.bold = base_bold; r.italic = True
                else:
                    r = para.add_run(sp_); r.bold = base_bold; r.italic = base_italic
                r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = 'Calibri'


# ── Header / Footer ────────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0,0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0,1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('AUTHORITY BLUEPRINT™  |  THOMAS E. BANKS II  |  MARCH 2026')
    rr.font.size = Pt(7.5); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True

    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0,0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Confidential. Prepared exclusively for Thomas E. Banks II.')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0,1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Dark background strip
    bg = doc.add_paragraph(); set_para_bg(bg, ASG_DARK_HEX)
    bg.paragraph_format.space_before = Pt(0); bg.paragraph_format.space_after = Pt(0)
    bg.add_run('  ')

    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(30); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))

    sp(doc, 20)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('AUTHORITY SYSTEMS GROUP™  —  CONFIDENTIAL STRATEGIC DELIVERABLE')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(6)
    tr = tp.add_run('Authority Blueprint™')
    tr.font.size = Pt(32); tr.bold = True; tr.font.color.rgb = DARK_TEXT; tr.font.name = 'Calibri'

    np_ = doc.add_paragraph(); np_.paragraph_format.space_before = Pt(0); np_.paragraph_format.space_after = Pt(4)
    nr = np_.add_run('Thomas E. Banks II, AAML Fellow')
    nr.font.size = Pt(16); nr.bold = True; nr.font.color.rgb = ASG_CHARCOAL; nr.font.name = 'Calibri'

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(0); sp2.paragraph_format.space_after = Pt(4)
    sr = sp2.add_run('Family Law Attorney  |  Kentucky & Indiana  |  KentuckianaDivorceLawyer.com')
    sr.font.size = Pt(11); sr.font.color.rgb = ASG_CHARCOAL; sr.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Prepared: March 8, 2026     |     Director: Roger Bauer     |     Version 1.0')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 16)

    # Credential strip
    ct = doc.add_table(1, 4); remove_table_borders(ct)
    for i in range(4): ct.columns[i].width = Inches(1.625)
    creds = [
        ('AAML Fellow', 'Top 2% Globally\n<1,600 Worldwide'),
        ('Super Lawyers', '2015–2026\n11 Consecutive Years'),
        ('AV Preeminent', 'Martindale-Hubbell\nHighest Possible Rating'),
        ('17+ Years', 'Exclusively KY & IN\nFamily Law'),
    ]
    for ci, (title, sub) in enumerate(creds):
        c = ct.cell(0, ci); set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8); cp.paragraph_format.space_after = Pt(8)
        r1 = cp.add_run(title + '\n'); r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = GOLD; r1.font.name = 'Calibri'
        r2 = cp.add_run(sub); r2.font.size = Pt(8); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'

    sp(doc, 16)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(6)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  A complete 13-section strategic plan for building Thomas Banks\' digital authority to match his '
        'credential standing in the Louisville and Southern Indiana family law market. Covers market '
        'analysis, Belief-to-Buy Framework mapping, 1-year strategic plan, Fast Cash Sprint, KPI '
        'dashboard, content strategy, niche website strategy, email sequences, lead magnets, and '
        'priority next steps.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    sp(doc, 10)

    # TOC table
    toc_items = [
        ('01', 'Executive Summary'), ('02', "Director's Briefing"),
        ('03', 'Market Analysis'), ('04', 'Belief-to-Buy Framework™ Map'),
        ('05', '1-Year Strategic Plan'), ('06', 'Fast Cash Sprint™'),
        ('07', 'KPI Dashboard'), ('08', 'Content Authority Strategy™'),
        ('09', 'Niche Website Strategy'), ('10', 'Email Sequences'),
        ('11', 'Lead Magnet & Referral Program'), ('12', 'Priority Services & Next Steps'),
        ('13', 'QC Gate Sign-Offs'),
    ]
    toc = doc.add_table(len(toc_items), 2); remove_table_borders(toc)
    toc.columns[0].width = Inches(0.5); toc.columns[1].width = Inches(6.0)
    for i, (num, title) in enumerate(toc_items):
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        c0 = toc.cell(i, 0); set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(3); p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(num); r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = ASG_BLUE; r0.font.name = 'Calibri'
        c1 = toc.cell(i, 1); set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(title); r1.font.size = Pt(10); r1.font.color.rgb = BODY_COLOR; r1.font.name = 'Calibri'

    add_pb(doc)


# ── Markdown Line Renderer ─────────────────────────────────────────────────────

def is_table_row(line):
    return line.strip().startswith('|') and line.strip().endswith('|')

def is_separator_row(line):
    stripped = line.strip()
    if not (stripped.startswith('|') and stripped.endswith('|')):
        return False
    inner = stripped[1:-1]
    return all(c in '-|: ' for c in inner)

def parse_table_rows(lines, start):
    """Parse a markdown table starting at start index. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not is_table_row(line):
            break
        if is_separator_row(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i

def render_md_table(doc, rows):
    """Render a parsed markdown table as a Word table."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    t = doc.add_table(len(rows), num_cols); remove_table_borders(t)
    col_width = Inches(6.5 / num_cols)
    for ci in range(num_cols):
        t.columns[ci].width = col_width

    for ri, row in enumerate(rows):
        is_header = (ri == 0)
        bg = ASG_DARK_HEX if is_header else (TINT_BG if ri % 2 == 0 else 'FFFFFF')
        for ci, cell_text in enumerate(row):
            c = t.cell(ri, ci); set_cell_bg(c, bg)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
            text_color = WHITE if is_header else BODY_COLOR
            cr = cp.add_run(cell_text); cr.bold = is_header
            cr.font.size = Pt(9); cr.font.color.rgb = text_color; cr.font.name = 'Calibri'
    sp(doc, 6)


def render_lines(doc, lines):
    """Render a list of markdown lines into the document."""
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip pure markdown separators
        if line.strip() == '---':
            i += 1
            continue

        # H3 subsection header
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(); set_para_bg(p, LIGHT_BG)
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f'  {text}')
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ASG_CHARCOAL; r.font.name = 'Calibri'
            i += 1
            continue

        # H4 minor header
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = ASG_BLUE; r.font.name = 'Calibri'
            i += 1
            continue

        # Bold-only line (standalone **text** as a mini-header)
        if line.strip().startswith('**') and line.strip().endswith('**') and line.strip().count('**') == 2:
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
            i += 1
            continue

        # Italic attribution line *— Name*
        if re.match(r'^\*—\s', line.strip()) or re.match(r'^\*—\s', line):
            text = line.strip().strip('*').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text); r.italic = True
            r.font.size = Pt(9); r.font.color.rgb = FOOTER_COLOR; r.font.name = 'Calibri'
            i += 1
            continue

        # Italic line *text*
        if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            text = line.strip()[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text); r.italic = True
            r.font.size = Pt(9.5); r.font.color.rgb = ASG_CHARCOAL; r.font.name = 'Calibri'
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph(); set_para_bg(p, GOLD_TINT)
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.15)
            r = p.add_run(f'  {text}')
            r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x8B, 0x6A, 0x0A); r.font.name = 'Calibri'
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.2)
            add_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.2)
            add_inline(p, text)
            i += 1
            continue

        # Table
        if is_table_row(line):
            rows, next_i = parse_table_rows(lines, i)
            render_md_table(doc, rows)
            i = next_i
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)
        add_inline(p, line.strip())
        i += 1


# ── Section Renderer ───────────────────────────────────────────────────────────

def render_section(doc, section_num, section_title, section_lines):
    # Section banner
    color = SECTION_COLORS.get(section_num, ASG_DARK_HEX)
    bt = doc.add_table(1, 2); remove_table_borders(bt)
    bt.columns[0].width = Inches(0.45); bt.columns[1].width = Inches(6.05)

    nc = bt.cell(0, 0); set_cell_bg(nc, GOLD_HEX)
    np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(6); np_.paragraph_format.space_after = Pt(6)
    nr = np_.add_run(f'{section_num:02d}')
    nr.bold = True; nr.font.size = Pt(14); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

    tc_ = bt.cell(0, 1); set_cell_bg(tc_, color)
    tp = tc_.paragraphs[0]
    tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(6)
    # Set cell text color via run (bg already set via set_cell_bg)
    tr_ = tp.add_run(f'Section {section_num}: {section_title}')
    tr_.bold = True; tr_.font.size = Pt(11); tr_.font.name = 'Calibri'
    # Parse hex color to set white text
    r_int = int(color[0:2], 16); g_int = int(color[2:4], 16); b_int = int(color[4:6], 16)
    _ = r_int  # ensure vars used
    tr_.font.color.rgb = WHITE

    sp(doc, 6)
    render_lines(doc, section_lines)
    add_pb(doc)


# ── Document Parser ────────────────────────────────────────────────────────────

def parse_document(md_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    lines = [l.rstrip('\n') for l in lines]

    # Find section boundaries
    section_starts = []
    for idx, line in enumerate(lines):
        m = re.match(r'^## Section (\d+):\s*(.+)$', line)
        if m:
            section_starts.append((idx, int(m.group(1)), m.group(2).strip()))

    sections = []
    for i, (start_idx, num, title) in enumerate(section_starts):
        end_idx = section_starts[i+1][0] if i+1 < len(section_starts) else len(lines)
        # Skip the header line itself
        section_lines = lines[start_idx+1:end_idx]
        sections.append({'num': num, 'title': title, 'lines': section_lines})

    return sections


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin',1.0),('right_margin',1.0),('top_margin',0.75),('bottom_margin',0.75)]:
        setattr(section, attr, Inches(val))

    add_header_footer(doc)
    build_cover(doc)

    sections = parse_document(MD_PATH)
    print(f'Parsed {len(sections)} sections')

    for s in sections:
        print(f'  Rendering Section {s["num"]}: {s["title"]} ({len(s["lines"])} lines)')
        render_section(doc, s['num'], s['title'], s['lines'])

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'\nSaved: {OUT_DOCX}')
    size_kb = os.path.getsize(OUT_DOCX) // 1024
    print(f'File size: {size_kb} KB')

if __name__ == '__main__':
    main()
