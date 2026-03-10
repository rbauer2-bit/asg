#!/usr/bin/env python3
"""
Authority Systems Group™
Thomas E. Banks II — Service Pages FAQ
KentuckianaDivorceLawyer.com + KentuckianaFamilyLawyer.com
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT    = RGBColor(0x1A, 0x25, 0x33)
GOLD         = RGBColor(0xC9, 0xA0, 0x2C)
META_COLOR   = RGBColor(0x55, 0x55, 0x55)
INTERNAL_BG  = RGBColor(0xF0, 0xF0, 0xF0)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
LIGHT_BG      = 'F5F5F5'
TINT_BG       = 'EAF6FC'
GOLD_HEX      = 'C9A02C'
INTERNAL_HEX  = 'F0F0F0'
INTERNAL_DARK = '444444'
COMPLIANCE_BG = 'FFF8E7'

# Service banner colors (one per service — distinct palette)
SERVICE_COLORS = {
    1: '2C3E50',   # Divorce — deep navy
    2: '1A5276',   # High-Asset — dark blue
    3: '117A65',   # Child Custody — teal
    4: '6E2F8E',   # Father's Rights — deep purple
    5: 'B7410E',   # Property Division — burnt sienna
    6: '1B4F72',   # Prenuptial — steel blue
}

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
CLIENT    = os.path.join(BASE, 'client-onboarding/clients/thomas-banks-law')
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
MD_PATH   = os.path.join(CLIENT, 'outputs/md/thomas-banks-law_service-pages-faq_20260308.md')
OUT_DOCX  = os.path.join(CLIENT, 'outputs/docx/thomas-banks-law_service-pages-faq_20260308.docx')


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tblBorders.append(el)
    tblPr.append(tblBorders)

def add_pb(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)

def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(size)

def divider(doc, color='CCCCCC'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'), color); pBdr.append(bt); pPr.append(pBdr)

def add_page_number(para):
    run = para.add_run(); run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin',None),(None,' PAGE '),('end',None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)

def add_inline(para, text, base_size=10.5, base_color=BODY_COLOR, base_bold=False):
    """Render **bold** and *italic* inline markers within a paragraph."""
    # Split on **bold** and *italic* markers
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            r = para.add_run(inner)
            r.bold = True; r.font.size = Pt(base_size)
            r.font.color.rgb = base_color; r.font.name = 'Calibri'
        elif part.startswith('*') and part.endswith('*'):
            inner = part[1:-1]
            r = para.add_run(inner)
            r.italic = True; r.font.size = Pt(base_size)
            r.font.color.rgb = base_color; r.font.name = 'Calibri'
        else:
            if part:
                r = para.add_run(part)
                r.bold = base_bold; r.font.size = Pt(base_size)
                r.font.color.rgb = base_color; r.font.name = 'Calibri'


# ── Header / Footer ────────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0,0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0,1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('SERVICE PAGE FAQs  |  THOMAS E. BANKS II  |  MARCH 2026')
    rr.font.size = Pt(7.5); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True
    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0,0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Confidential. Prepared exclusively for Thomas E. Banks II.')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0,1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover ──────────────────────────────────────────────────────────────────────

def build_cover(doc):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(40); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))
    sp(doc, 24)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('AUTHORITY SYSTEMS GROUP™  —  CONFIDENTIAL')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)
    tr_ = tp.add_run('Service Page FAQs')
    tr_.font.size = Pt(30); tr_.bold = True; tr_.font.color.rgb = DARK_TEXT; tr_.font.name = 'Calibri'

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(0); sp2.paragraph_format.space_after = Pt(4)
    sr = sp2.add_run('Thomas E. Banks II, AAML Fellow  |  Family Law Attorney  |  Kentucky & Indiana')
    sr.font.size = Pt(13); sr.font.color.rgb = ASG_CHARCOAL; sr.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Prepared: March 8, 2026     |     KentuckianaDivorceLawyer.com  +  KentuckianaFamilyLawyer.com')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 16)

    # Credential bar
    cred_table = doc.add_table(1, 4); remove_table_borders(cred_table)
    for i in range(4): cred_table.columns[i].width = Inches(1.625)
    creds = [
        ('AAML Fellow', 'One of Fewer Than 1,600 Worldwide'),
        ('Super Lawyers', '2015–2026 (11 Consecutive Years)'),
        ('AV Preeminent', 'Martindale-Hubbell Highest Rating'),
        ('17+ Years', 'Exclusively in KY & IN Family Law'),
    ]
    for ci, (title, sub) in enumerate(creds):
        c = cred_table.cell(0, ci); set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8); cp.paragraph_format.space_after = Pt(8)
        r1 = cp.add_run(title + '\n'); r1.bold = True; r1.font.size = Pt(10)
        r1.font.color.rgb = GOLD; r1.font.name = 'Calibri'
        r2 = cp.add_run(sub); r2.font.size = Pt(8); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'

    sp(doc, 16)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(4)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  60 differentiated FAQs (10 per service) applying the Differentiated FAQ Framework™. '
        'Questions are calibrated to prospects at different belief and emotional stages — '
        'Foundational, Situational, Decision, and Pre-Objection. All content is ABA compliant. '
        'Internal metadata and compliance notes are included at the end of this document.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    sp(doc, 12)

    # Service index
    services = [
        ('01', 'Divorce (Contested and Uncontested)',     '10 FAQs'),
        ('02', 'High-Asset Divorce & Business Valuation', '10 FAQs'),
        ('03', 'Child Custody and Support',               '10 FAQs'),
        ("04", "Father's Rights",                         '10 FAQs'),
        ('05', 'Property Division',                       '10 FAQs'),
        ('06', 'Prenuptial Agreements',                   '10 FAQs'),
    ]
    idx = doc.add_table(len(services), 3); remove_table_borders(idx)
    idx.columns[0].width = Inches(0.5); idx.columns[1].width = Inches(5.0); idx.columns[2].width = Inches(1.0)
    for i, (num, title, count) in enumerate(services):
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        c0 = idx.cell(i,0); set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(4); p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(num); r0.bold = True; r0.font.size = Pt(10)
        r0.font.color.rgb = ASG_BLUE; r0.font.name = 'Calibri'
        c1 = idx.cell(i,1); set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(title); r1.font.size = Pt(10.5); r1.font.color.rgb = BODY_COLOR; r1.font.name = 'Calibri'
        c2 = idx.cell(i,2); set_cell_bg(c2, bg)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(count); r2.font.size = Pt(9); r2.font.color.rgb = ASG_CHARCOAL; r2.font.name = 'Calibri'

    add_pb(doc)


# ── Document Parser ─────────────────────────────────────────────────────────────

def parse_faq_doc(md_path):
    """
    Returns:
      faq_sets    — list of dicts: {num, title, categories: [{cat_name, qas: [{q, a}]}]}
      metadata    — list of dicts: {num, title, rows: [{q_num, real_q, stage, generic, superior}]}
      aba_section — list of lines for ABA compliance section
    """
    with open(md_path, 'r') as f:
        content = f.read()

    # ── Split main body from Internal Metadata ──
    meta_split = re.split(r'\n## INTERNAL METADATA\b', content, maxsplit=1)
    main_body = meta_split[0]
    rest = meta_split[1] if len(meta_split) > 1 else ''

    # ── Split metadata from ABA section ──
    aba_split = re.split(r'\n## ABA COMPLIANCE NOTES\b', rest, maxsplit=1)
    metadata_body = aba_split[0]
    aba_body = aba_split[1] if len(aba_split) > 1 else ''

    # ── Parse FAQ sets ──
    set_blocks = re.split(r'\n## FAQ SET \d+:', main_body)
    faq_sets = []
    for i, block in enumerate(set_blocks[1:], start=1):
        lines = block.strip().split('\n')
        title = lines[0].strip()
        rest_block = '\n'.join(lines[1:])

        # Split into category sections
        cat_blocks = re.split(r'\n### (Foundational Questions|Situational Questions|Decision Questions|Pre-Objection Questions)\b', rest_block)
        categories = []
        j = 1
        while j < len(cat_blocks):
            cat_name = cat_blocks[j].strip()
            cat_content = cat_blocks[j+1] if j+1 < len(cat_blocks) else ''
            qas = parse_qas(cat_content)
            categories.append({'cat_name': cat_name, 'qas': qas})
            j += 2

        faq_sets.append({'num': i, 'title': title, 'categories': categories})

    # ── Parse metadata ──
    meta_sets = []
    meta_set_blocks = re.split(r'\n### Metadata — FAQ Set \d+:', metadata_body)
    for i, block in enumerate(meta_set_blocks[1:], start=1):
        lines = block.strip().split('\n')
        meta_title = lines[0].strip()
        rows = []
        in_table = False
        for line in lines[1:]:
            line = line.strip()
            if not line or line.startswith('|---') or line.startswith('| Q |') or line.startswith('| --- |'):
                in_table = True
                continue
            if in_table and line.startswith('|'):
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if len(cells) >= 5:
                    rows.append({
                        'q_num': cells[0],
                        'real_q': cells[1],
                        'stage': cells[2],
                        'generic': cells[3],
                        'superior': cells[4],
                    })
        meta_sets.append({'num': i, 'title': meta_title, 'rows': rows})

    # ── Parse ABA section ──
    aba_lines = [l for l in aba_body.strip().split('\n') if l.strip() and not l.strip().startswith('---')]

    return faq_sets, meta_sets, aba_lines


def parse_qas(content):
    """Extract Q/A pairs from a category block."""
    qas = []
    # Match **Q1: question text** followed by answer paragraph(s)
    qa_pattern = re.split(r'\n\*\*Q\d+:', content)
    for block in qa_pattern[1:]:
        # First line up to ** is the question
        end_bold = block.find('**')
        if end_bold == -1:
            continue
        question = block[:end_bold].strip()
        answer_raw = block[end_bold+2:].strip()
        # Answer is the next non-empty paragraph(s)
        answer_lines = []
        for line in answer_raw.split('\n'):
            stripped = line.strip()
            if stripped.startswith('**Q') or stripped.startswith('### ') or stripped.startswith('## '):
                break
            answer_lines.append(stripped)
        answer = ' '.join([l for l in answer_lines if l]).strip()
        if question:
            qas.append({'q': question, 'a': answer})
    return qas


# ── Renderers ──────────────────────────────────────────────────────────────────

CATEGORY_LABELS = {
    'Foundational Questions': ('FOUNDATIONAL', ASG_BLUE_HEX),
    'Situational Questions':  ('SITUATIONAL',  '1A7E4F'),
    'Decision Questions':     ('DECISION',     '7D3C98'),
    'Pre-Objection Questions':('PRE-OBJECTION','B7410E'),
}

def render_category_header(doc, cat_name):
    label, color_hex = CATEGORY_LABELS.get(cat_name, (cat_name.upper(), CHARCOAL_HEX))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    set_para_bg(p, color_hex)
    r = p.add_run(f'  {label}')
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
    r.font.name = 'Calibri'

def render_qa(doc, q_num, q_text, a_text):
    """Render a single Q/A pair."""
    # Question line
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(10); qp.paragraph_format.space_after = Pt(2)
    qp.paragraph_format.left_indent = Inches(0.1)
    qr = qp.add_run(f'Q{q_num}:  ')
    qr.bold = True; qr.font.size = Pt(10); qr.font.color.rgb = ASG_BLUE; qr.font.name = 'Calibri'
    qt = qp.add_run(q_text)
    qt.bold = True; qt.font.size = Pt(10.5); qt.font.color.rgb = DARK_TEXT; qt.font.name = 'Calibri'

    # Answer paragraph
    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(2); ap.paragraph_format.space_after = Pt(6)
    ap.paragraph_format.left_indent = Inches(0.3)
    add_inline(ap, a_text, base_size=10.5, base_color=BODY_COLOR)

def render_service_banner(doc, service_num, service_title, color_hex):
    """Render the colored banner for a service section."""
    bt = doc.add_table(1, 2); remove_table_borders(bt)
    bt.columns[0].width = Inches(0.45); bt.columns[1].width = Inches(6.05)

    nc = bt.cell(0, 0); set_cell_bg(nc, GOLD_HEX)
    np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(6); np_.paragraph_format.space_after = Pt(6)
    nr = np_.add_run(f'{service_num:02d}')
    nr.bold = True; nr.font.size = Pt(14); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

    tc_ = bt.cell(0, 1); set_cell_bg(tc_, color_hex)
    tp = tc_.paragraphs[0]
    tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(6)
    tr_ = tp.add_run(f'FAQ Set {service_num:02d}:  {service_title}')
    tr_.bold = True; tr_.font.size = Pt(11); tr_.font.name = 'Calibri'
    tr_.font.color.rgb = WHITE

    sp(doc, 6)

def render_service_faq(doc, faq_set):
    color_hex = SERVICE_COLORS.get(faq_set['num'], ASG_DARK_HEX)
    render_service_banner(doc, faq_set['num'], faq_set['title'], color_hex)

    q_counter = 1
    for cat in faq_set['categories']:
        render_category_header(doc, cat['cat_name'])
        for qa in cat['qas']:
            render_qa(doc, q_counter, qa['q'], qa['a'])
            q_counter += 1

    add_pb(doc)


# ── Internal Metadata Section ──────────────────────────────────────────────────

def render_metadata_banner(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
    set_para_bg(p, INTERNAL_DARK)
    r = p.add_run('  INTERNAL METADATA — NOT FOR CLIENT DELIVERY')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

def render_metadata_set(doc, meta_set):
    # Service label
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(10); lp.paragraph_format.space_after = Pt(4)
    set_para_bg(lp, INTERNAL_HEX)
    lr = lp.add_run(f'  FAQ Set {meta_set["num"]:02d}: {meta_set["title"]}')
    lr.bold = True; lr.font.size = Pt(10); lr.font.color.rgb = META_COLOR; lr.font.name = 'Calibri'

    if not meta_set['rows']:
        p = doc.add_paragraph()
        p.add_run('[Metadata not parsed — see source markdown]').font.size = Pt(9)
        return

    # Table
    headers = ['Q', 'Real Question Underneath', 'B/E Stage', 'Generic Competitor Answer', 'Why Ours Is Superior']
    col_widths = [Inches(0.35), Inches(1.35), Inches(0.65), Inches(2.0), Inches(2.15)]

    tbl = doc.add_table(1 + len(meta_set['rows']), len(headers))
    tbl.style = 'Table Grid'
    for ci, w in enumerate(col_widths):
        tbl.columns[ci].width = w

    # Header row
    hr = tbl.rows[0]
    for ci, hdr in enumerate(headers):
        cell = hr.cells[ci]; set_cell_bg(cell, ASG_DARK_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(hdr)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE; r.font.name = 'Calibri'

    # Data rows
    for ri, row_data in enumerate(meta_set['rows'], start=1):
        row = tbl.rows[ri]
        bg = 'FAFAFA' if ri % 2 == 0 else 'FFFFFF'
        values = [row_data.get('q_num',''), row_data.get('real_q',''), row_data.get('stage',''),
                  row_data.get('generic',''), row_data.get('superior','')]
        for ci, val in enumerate(values):
            cell = row.cells[ci]; set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(8); r.font.color.rgb = BODY_COLOR; r.font.name = 'Calibri'
            if ci == 0: r.bold = True

    sp(doc, 6)


# ── ABA Compliance Section ─────────────────────────────────────────────────────

def render_aba_section(doc, aba_lines):
    # Banner
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(8)
    set_para_bg(bp, GOLD_HEX)
    br_ = bp.add_run('  ABA / STATE BAR COMPLIANCE NOTES')
    br_.bold = True; br_.font.size = Pt(11); br_.font.color.rgb = WHITE; br_.font.name = 'Calibri'

    for line in aba_lines:
        line = line.strip()
        if not line:
            continue
        # Blockquotes → indented italic
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            p = doc.add_paragraph()
            set_para_bg(p, COMPLIANCE_BG)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(text)
            r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = BODY_COLOR; r.font.name = 'Calibri'
        elif line.startswith('**') and line.endswith('**'):
            # Bold-only line = section label
            text = line.strip('*')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'
        elif line.startswith('*') and line.endswith('*'):
            # Attribution / italic line
            text = line.strip('*')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = META_COLOR; r.font.name = 'Calibri'
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            add_inline(p, line, base_size=10, base_color=BODY_COLOR)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_COLOR

    add_header_footer(doc)
    build_cover(doc)

    faq_sets, meta_sets, aba_lines = parse_faq_doc(MD_PATH)

    for faq_set in faq_sets:
        render_service_faq(doc, faq_set)

    # Internal Metadata
    add_pb(doc)
    render_metadata_banner(doc)
    sp(doc, 8)
    for meta_set in meta_sets:
        render_metadata_set(doc, meta_set)

    # ABA Compliance Section
    add_pb(doc)
    render_aba_section(doc, aba_lines)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
