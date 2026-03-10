#!/usr/bin/env python3
"""
Authority Systems Group™
Thomas E. Banks II — Service Pages Web Copy
KentuckianaDivorceLawyer.com + KentuckianaFamilyLawyer.com
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT     = RGBColor(0x1A, 0x25, 0x33)
GOLD          = RGBColor(0xC9, 0xA0, 0x2C)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
LIGHT_BG      = 'F5F5F5'
TINT_BG       = 'EAF6FC'
GOLD_HEX      = 'C9A02C'
GOLD_TINT     = 'FDF8EC'
SLATE_HEX     = '2C3E50'

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
CLIENT    = os.path.join(BASE, 'client-onboarding/clients/thomas-banks-law')
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
MD_PATH   = os.path.join(CLIENT, 'outputs/md/thomas-banks-law_service-pages-web-copy_20260308.md')
OUT_DOCX  = os.path.join(CLIENT, 'outputs/docx/thomas-banks-law_service-pages-web-copy_20260308.docx')


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tblBorders.append(el)
    tblPr.append(tblBorders)

def add_pb(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)

def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(size)

def divider(doc, color='CCCCCC'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'), color); pBdr.append(bt); pPr.append(pBdr)

def add_page_number(para):
    run = para.add_run(); run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin',None),(None,' PAGE '),('end',None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)

def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0,0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0,1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('SERVICE PAGES WEB COPY  |  THOMAS E. BANKS II  |  MARCH 2026')
    rr.font.size = Pt(7.5); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True
    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0,0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Confidential. Prepared exclusively for Thomas E. Banks II.')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0,1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover ──────────────────────────────────────────────────────────────────────

def build_cover(doc):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(40); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH): lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))
    sp(doc, 24)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('AUTHORITY SYSTEMS GROUP™  —  CONFIDENTIAL')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)
    tr = tp.add_run('Service Pages Web Copy')
    tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = DARK_TEXT; tr.font.name = 'Calibri'

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(0); sp2.paragraph_format.space_after = Pt(4)
    sr = sp2.add_run('Thomas E. Banks II, AAML Fellow  |  Family Law Attorney  |  Kentucky & Indiana')
    sr.font.size = Pt(13); sr.font.color.rgb = ASG_CHARCOAL; sr.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Prepared: March 8, 2026     |     KentuckianaDivorceLawyer.com  +  KentuckianaFamilyLawyer.com')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 16)

    # Credential bar
    cred_table = doc.add_table(1, 4); remove_table_borders(cred_table)
    for i in range(4): cred_table.columns[i].width = Inches(1.625)
    creds = [
        ('AAML Fellow', 'Top 2% of Family Lawyers Globally'),
        ('Super Lawyers', '2015–2026 (11 Consecutive Years)'),
        ('AV Preeminent', 'Martindale-Hubbell Highest Rating'),
        ('17+ Years', 'Exclusively in KY & IN Family Law'),
    ]
    for ci, (title, sub) in enumerate(creds):
        c = cred_table.cell(0, ci); set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8); cp.paragraph_format.space_after = Pt(8)
        r1 = cp.add_run(title + '\n'); r1.bold = True; r1.font.size = Pt(10)
        r1.font.color.rgb = GOLD; r1.font.name = 'Calibri'
        r2 = cp.add_run(sub); r2.font.size = Pt(8); r2.font.color.rgb = WHITE; r2.font.name = 'Calibri'

    sp(doc, 16)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(4)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  Six fully written service pages applying the Authority Conversion Protocol™ (Belief-to-Buy '
        'Framework). Each page moves a prospect through both the Belief Track and the Emotion Track '
        'simultaneously. State-specific [KY] and [IN] tags are marked throughout for clean dual-site '
        'implementation.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    sp(doc, 12)

    # Page index
    pages = [
        ('01', 'Divorce (Contested and Uncontested)'),
        ('02', 'High-Asset Divorce & Business Valuation'),
        ('03', 'Child Custody and Support'),
        ('04', 'Father\'s Rights'),
        ('05', 'Property Division'),
        ('06', 'Prenuptial Agreements'),
    ]
    idx = doc.add_table(len(pages), 2); remove_table_borders(idx)
    idx.columns[0].width = Inches(0.5); idx.columns[1].width = Inches(6.0)
    for i, (num, title) in enumerate(pages):
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        c0 = idx.cell(i,0); set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(4); p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(num); r0.bold = True; r0.font.size = Pt(10)
        r0.font.color.rgb = ASG_BLUE; r0.font.name = 'Calibri'
        c1 = idx.cell(i,1); set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(title); r1.font.size = Pt(10.5); r1.font.color.rgb = BODY_COLOR; r1.font.name = 'Calibri'

    add_pb(doc)


# ── Page Parser + Renderer ─────────────────────────────────────────────────────

def parse_pages(md_path):
    with open(md_path, 'r') as f:
        content = f.read()
    # Split on PAGE headers
    blocks = re.split(r'\n(?=# PAGE \d+:)', content)
    pages = []
    for block in blocks:
        if not block.strip().startswith('# PAGE'):
            continue
        lines = block.split('\n')
        # Title line
        title_match = re.match(r'^# PAGE \d+:\s*(.+)$', lines[0])
        page_title = title_match.group(1).strip() if title_match else ''
        # Rest of content
        body_lines = lines[1:]
        pages.append({'title': page_title, 'lines': body_lines})
    return pages


def render_page(doc, page_num, page):
    # Page number + title banner
    nt = doc.add_table(1, 2); remove_table_borders(nt)
    nt.columns[0].width = Inches(0.45); nt.columns[1].width = Inches(6.05)
    nc = nt.cell(0,0); set_cell_bg(nc, GOLD_HEX)
    np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np_.paragraph_format.space_before = Pt(6); np_.paragraph_format.space_after = Pt(6)
    nr = np_.add_run(f'{page_num:02d}'); nr.bold = True; nr.font.size = Pt(14)
    nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'
    tc_ = nt.cell(0,1); set_cell_bg(tc_, ASG_DARK_HEX)
    tp = tc_.paragraphs[0]; tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(6)
    tr_ = tp.add_run(page['title']); tr_.bold = True; tr_.font.size = Pt(11)
    tr_.font.color.rgb = WHITE; tr_.font.name = 'Calibri'

    sp(doc, 4)

    # ACP stage label
    stage_labels = {
        1: 'B1/E1 — Enemy Belief + Dissatisfaction',
        2: 'B2-B3/E2-E3 — Pre-Frame + New Worldview',
        3: 'B3-B4/E3-E4 — New Worldview + Desire',
        4: 'B4/E4 — Internal Alignment + Urgency',
        5: 'B5/E5 — Certainty + Relief',
        6: 'CTA — No-Cost Discovery Call',
    }

    lines = page['lines']
    i = 0
    section_count = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('## '):
            # H2 — page headline (hook)
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = DARK_TEXT; r.font.name = 'Calibri'

        elif line.startswith('### '):
            # H3 — subsection header
            section_count += 1
            text = line[4:].strip()
            # Stage label strip
            stage_p = doc.add_paragraph(); set_para_bg(stage_p, LIGHT_BG)
            stage_p.paragraph_format.space_before = Pt(10); stage_p.paragraph_format.space_after = Pt(0)
            sr = stage_p.add_run(f'  {text}')
            sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = ASG_CHARCOAL; sr.font.name = 'Calibri'

        elif line.startswith('---'):
            # Separator — skip visual dividers between main sections
            pass

        elif line.startswith('**[KY:') or line.startswith('**[IN:') or line.startswith('**Or contact'):
            # CTA contact block
            text = line.strip().strip('*').strip()
            p = doc.add_paragraph(); set_para_bg(p, TINT_BG)
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.1)
            r = p.add_run(f'  {text}'); r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = ASG_BLUE; r.font.name = 'Calibri'

        elif line.startswith('*—') or line.startswith('*— '):
            # Attribution line
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text); r.italic = True; r.font.size = Pt(8.5)
            r.font.color.rgb = FOOTER_COLOR; r.font.name = 'Calibri'

        elif line.startswith('### Start With') or line.startswith('### Let\'s Talk') or \
             line.startswith('### Schedule') or line.startswith('### You Should') or \
             line.startswith('### Your Children') or line.startswith('### Get the Analysis') or \
             line.startswith('### The Conversation'):
            # CTA subsection header
            p = doc.add_paragraph(); set_para_bg(p, ASG_DARK_HEX)
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
            text = line[4:].strip()
            r = p.add_run(f'  {text}'); r.bold = True; r.font.size = Pt(10.5)
            r.font.color.rgb = WHITE; r.font.name = 'Calibri'

        elif line.strip():
            # Body paragraph
            text = line.strip()
            # Bold inline: **text**
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)
            # Simple bold/italic parsing
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2]); r.bold = True
                else:
                    r = p.add_run(part)
                r.font.size = Pt(10.5); r.font.color.rgb = BODY_COLOR; r.font.name = 'Calibri'

        i += 1

    add_pb(doc)


def build_implementation_notes(doc):
    doc.add_paragraph()  # ensure we're on a fresh para

    # Title
    ht = doc.add_table(1, 1); remove_table_borders(ht)
    ht.columns[0].width = Inches(6.5)
    hc = ht.cell(0,0); set_cell_bg(hc, ASG_DARK_HEX)
    hp = hc.paragraphs[0]; hp.paragraph_format.space_before = Pt(6); hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run('  IMPLEMENTATION NOTES — DUAL-SITE DEPLOYMENT')
    hr.bold = True; hr.font.size = Pt(11); hr.font.color.rgb = WHITE; hr.font.name = 'Calibri'

    sp(doc, 8)

    # State swap table
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(0); p_intro.paragraph_format.space_after = Pt(6)
    ri = p_intro.add_run('State-specific reference swaps — apply before publishing to each site:')
    ri.bold = True; ri.font.size = Pt(10); ri.font.color.rgb = DARK_TEXT; ri.font.name = 'Calibri'

    swap_data = [
        ('[KY] / [IN] tags', 'Use KY version', 'Use IN version'),
        ('Court references', 'Jefferson County Family Court', 'Clark County / Floyd County Circuit Court'),
        ('Geographic refs', 'Louisville, KY', 'Jeffersonville, New Albany, Clarksville (Southern Indiana)'),
        ('Phone number', '(502) 568-4700', '(812) 920-5603'),
    ]
    t = doc.add_table(len(swap_data)+1, 3); remove_table_borders(t)
    t.columns[0].width = Inches(1.8); t.columns[1].width = Inches(2.35); t.columns[2].width = Inches(2.35)
    hdrs = ['Tag / Reference', 'KentuckianaDivorceLawyer.com (KY)', 'KentuckianaFamilyLawyer.com (IN)']
    hr_ = t.rows[0]
    for ci, lbl in enumerate(hdrs):
        c = hr_.cells[ci]; set_cell_bg(c, SLATE_HEX)
        cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
        cr = cp.add_run(lbl); cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = WHITE; cr.font.name = 'Calibri'
    for i, row_data in enumerate(swap_data):
        row = t.rows[i+1]; bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = row.cells[ci]; set_cell_bg(c, bg)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(val); cr.font.size = Pt(9.5); cr.font.color.rgb = BODY_COLOR; cr.font.name = 'Calibri'

    sp(doc, 10)

    # Credential pull-quote suggestions
    pq_label = doc.add_paragraph()
    pq_label.paragraph_format.space_before = Pt(0); pq_label.paragraph_format.space_after = Pt(6)
    pqr = pq_label.add_run('Recommended credential callouts for sidebar / pull-quote blocks (web design layer):')
    pqr.bold = True; pqr.font.size = Pt(10); pqr.font.color.rgb = DARK_TEXT; pqr.font.name = 'Calibri'

    callouts = [
        '"AAML Fellow — Top 2% of Family Lawyers Globally. Fewer Than 1,600 Fellows Worldwide."',
        '"Super Lawyers 2021–2026 | Rising Stars 2015–2020 — Recognized Every Year Since 2015"',
        '"Martindale-Hubbell AV Preeminent — Highest Possible Rating for Legal Ability and Ethics"',
        '"17+ Years Exclusively in Family Law in Kentucky and Indiana"',
        '"Cornell University — Graduated #1 in Class | U of L Law School — Scholarship Recipient"',
    ]
    for c_text in callouts:
        cp = doc.add_paragraph(); set_para_bg(cp, GOLD_TINT)
        cp.paragraph_format.space_before = Pt(2); cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.left_indent = Inches(0.1)
        cr = cp.add_run(f'  {c_text}')
        cr.font.size = Pt(9.5); cr.font.color.rgb = RGBColor(0x8B, 0x6A, 0x0A); cr.font.name = 'Calibri'; cr.italic = True

    sp(doc, 12)

    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(0); ep.paragraph_format.space_after = Pt(0)
    er = ep.add_run('Authority Systems Group™ — Confidential. Prepared exclusively for Thomas E. Banks II. | March 8, 2026')
    er.font.size = Pt(8); er.font.color.rgb = FOOTER_COLOR; er.italic = True; er.font.name = 'Calibri'


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin',1.0),('right_margin',1.0),('top_margin',0.75),('bottom_margin',0.75)]:
        setattr(section, attr, Inches(val))

    add_header_footer(doc)
    build_cover(doc)

    pages = parse_pages(MD_PATH)
    for i, page in enumerate(pages, 1):
        render_page(doc, i, page)

    build_implementation_notes(doc)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')

if __name__ == '__main__':
    main()
