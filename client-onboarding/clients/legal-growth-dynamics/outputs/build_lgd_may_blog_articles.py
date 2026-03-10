#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | May 2026 Blog Articles (30-Day Series)
Produced: 2026-03-08
"""

import sys
import os
import re
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)

BASE_DIR    = os.path.dirname(__file__)
INPUT_FILES = [
    os.path.join(BASE_DIR, 'lgd_blog_articles_hnw_familylaw_20260308_may-01-10.md'),
    os.path.join(BASE_DIR, 'lgd_blog_articles_hnw_familylaw_20260308_may-11-20.md'),
    os.path.join(BASE_DIR, 'lgd_blog_articles_hnw_familylaw_20260308_may-21-30.md'),
]
OUTPUT_DOCX = os.path.join(BASE_DIR, 'lgd_blog_articles_hnw_familylaw_20260308_may-complete.docx')


# ── helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color='25AAE1'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name  = 'Georgia'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.line_spacing = Pt(15)

    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(22)
    h1.font.bold  = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)

    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(16)
    h2.font.bold  = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after  = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(13)
    h3.font.bold  = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(4)


# ── cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('30-DAY BLOG ARTICLE SERIES')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(32)
    run.font.bold  = True
    run.font.color.rgb = ASG_BLUE
    cover_title.paragraph_format.space_before = Pt(72)
    cover_title.paragraph_format.space_after  = Pt(8)

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub1.add_run('May 2026')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(24)
    r.font.bold  = True
    r.font.color.rgb = ASG_CHARCOAL
    sub1.paragraph_format.space_after = Pt(4)

    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = firm.add_run('Legal Growth Dynamics')
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(22)
    r2.font.bold  = True
    r2.font.color.rgb = ASG_CHARCOAL
    firm.paragraph_format.space_after = Pt(4)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run('Marketing & Retainer Pain Points | Family Law Attorneys')
    r3.font.name  = 'Calibri'
    r3.font.size  = Pt(14)
    r3.font.color.rgb = ASG_CHARCOAL
    sub2.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    for line in [
        'Prepared exclusively for Legal Growth Dynamics',
        'Roger Bauer, Founder',
        'Jeffersonville, IN',
        '',
        'March 2026  |  Confidential',
        '5 Thematic Hurdles  |  30 Articles  |  4 Pillar Pieces',
        'Belief-to-Buy\u2122 / Authority Conversion Protocol\u2122',
        '',
        'Under the direction of Roger, Director',
        'Authority Systems Group\u2122',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        if 'Authority Systems Group' in line or 'Roger, Director' in line:
            r.font.color.rgb = ASG_BLUE
            r.font.bold = True
        elif 'Confidential' in line or 'March' in line or 'Thematic' in line or 'Belief' in line:
            r.font.color.rgb = ASG_CHARCOAL
            r.font.bold = True
        else:
            r.font.color.rgb = ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Hurdle overview page
    intro = doc.add_paragraph()
    r = intro.add_run('About This Article Series')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(18)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after  = Pt(12)

    overview_text = (
        'This 30-article series addresses five high-priority marketing and retainer pain points '
        'facing family law managing partners. Each article is mapped to a specific stage pair '
        'within the Belief-to-Buy\u2122 / Authority Conversion Protocol\u2122 (ACP) framework, '
        'ensuring that the full content set guides prospects across both the Belief Track and '
        'Emotion Track simultaneously.\n\n'
        'All articles are written for publication on the Legal Growth Dynamics blog in May 2026. '
        'Four pillar articles (~1,250+ words) are designed for weekly anchor publication. '
        'The remaining 26 short articles (400\u2013750 words) publish daily across the remaining '
        'days of the month.'
    )
    p = doc.add_paragraph(overview_text)
    p.paragraph_format.space_after = Pt(16)

    hurdles = [
        ('Hurdle 5', 'Retainer Sticker Shock', 'Articles 1\u20136',
         'Prospects who freeze, stall, or ghost after the retainer amount is named. '
         'Addresses the consultation structure, fee framing, and intake questions that '
         'convert uncertain prospects into signed clients.'),
        ('Hurdle 6', 'The Commoditization Trap', 'Articles 7\u201312',
         'Being evaluated on hourly rates rather than positioned as the obvious specialist. '
         'Addresses authority positioning, website differentiation, and the communication '
         'of value that ends price comparisons.'),
        ('Hurdle 7', 'Google Dependency & Fragile Lead Sources', 'Articles 13\u201318',
         'Over-reliance on a single platform for client acquisition. Addresses multi-channel '
         'lead architecture, referral network development, and the order-of-operations for '
         'marketing investment.'),
        ('Hurdle 8', 'The Marketing Agency Betrayal', 'Articles 19\u201324',
         'The repeating cycle of hiring and firing agencies that don\u2019t understand the '
         'family law business model. Addresses what agencies actually deliver vs. what '
         'practices actually need, and how to evaluate marketing partners correctly.'),
        ('Hurdle 9', 'Feast-or-Famine Revenue Cycle', 'Articles 25\u201330',
         'Unpredictable month-to-month revenue driven by absent pipeline management. '
         'Addresses the complete revenue system: consistent intake, pipeline visibility, '
         'retainer reactivation, and case extension.'),
    ]

    for hurdle_num, hurdle_name, articles, description in hurdles:
        ph = doc.add_paragraph()
        r1 = ph.add_run(f'{hurdle_num}: {hurdle_name}')
        r1.font.name  = 'Calibri'
        r1.font.size  = Pt(13)
        r1.font.bold  = True
        r1.font.color.rgb = ASG_BLUE
        ph.paragraph_format.space_before = Pt(10)
        ph.paragraph_format.space_after  = Pt(2)

        pa = doc.add_paragraph()
        ra = pa.add_run(f'{articles}  |  ')
        ra.font.name  = 'Calibri'
        ra.font.size  = Pt(10)
        ra.font.bold  = True
        ra.font.color.rgb = ASG_CHARCOAL
        rd = pa.add_run(description)
        rd.font.name  = 'Georgia'
        rd.font.size  = Pt(10)
        rd.font.color.rgb = BODY_COLOR
        pa.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


# ── footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '25AAE1')
    pBdr.append(top)
    pPr.append(pBdr)

    run = fp.add_run(
        'Authority Systems Group\u2122 \u2014 Confidential. Prepared exclusively for Legal Growth Dynamics.'
    )
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── article metadata renderer ─────────────────────────────────────────────────

def render_metadata_block(doc, text):
    """Render **Key:** Value metadata lines as a shaded info block."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(9)
    r.font.color.rgb = ASG_CHARCOAL
    set_paragraph_shading(p, 'EEF7FC')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.15)


# ── markdown parser ───────────────────────────────────────────────────────────

def is_metadata_line(line):
    """Detect article metadata lines like **Hurdle:** 5 — ..."""
    return re.match(r'^\*\*[A-Za-z ]+:\*\*', line.strip())


def parse_and_render(doc, md_path, skip_header_lines=5):
    """Parse a markdown file and render it into the document."""
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block   = False
    in_metadata     = False
    header_skipped  = 0
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip the file header block (first N comment lines + blank + ---)
        if header_skipped < skip_header_lines:
            if line.startswith('#') or line.strip() == '' or line.strip() == '---':
                header_skipped += 1
                i += 1
                continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name  = 'Courier New'
            r.font.size  = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Article metadata lines (**Hurdle:** ..., **ACP Stage Pair:** ..., etc.)
        if is_metadata_line(line):
            render_metadata_block(doc, re.sub(r'\*\*(.*?)\*\*', r'\1', line.strip()))
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-\|:]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            cell.text = clean
                            for para in cell.paragraphs:
                                para.paragraph_format.space_after = Pt(2)
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = WHITE
                            if ri == 0:
                                for para in cell.paragraphs:
                                    tc   = cell._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    shd  = OxmlElement('w:shd')
                                    shd.set(qn('w:val'),   'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'),  '25AAE1')
                                    tcPr.append(shd)
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            # H1 used for file headers — render as article separator with page break
            text = line[2:].strip()
            if text and not text.startswith('Legal Growth Dynamics'):
                doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            # Article title — add page break before each article
            text = line[3:].strip()
            if text.startswith('Article'):
                doc.add_page_break()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('### ') and not line.startswith('#### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            r.font.name  = 'Calibri'
            r.font.size  = Pt(12)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip().replace('*', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name   = 'Calibri'
            r.font.size   = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.left_indent = Inches(0.3)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'),   'single')
            left.set(qn('w:sz'),    '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            parts = re.split(r'\*\*(.*?)\*\*', line[2:].strip())
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph — inline bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            italic_parts = re.split(r'\*(.*?)\*', part)
            for k, ipart in enumerate(italic_parts):
                if ipart:
                    r = p.add_run(ipart)
                    r.font.name  = 'Georgia'
                    r.font.size  = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
                    if k % 2 == 1:
                        r.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        i += 1

    return doc


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)

    for md_file in INPUT_FILES:
        print(f'  Processing: {os.path.basename(md_file)}')
        parse_and_render(doc, md_file)

    doc.save(OUTPUT_DOCX)
    print(f'\u2713 Document saved: {OUTPUT_DOCX}')


if __name__ == '__main__':
    main()
