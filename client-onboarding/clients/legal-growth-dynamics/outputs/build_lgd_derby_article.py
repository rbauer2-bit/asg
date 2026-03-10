#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | Kentucky Derby Seasonal Article — May 2026
Produced: 2026-03-08
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)

BASE_DIR    = os.path.dirname(__file__)
INPUT_MD    = os.path.join(BASE_DIR, 'lgd-family-law-attorneys_kentucky-derby-article_20260308.md')
OUTPUT_DOCX = os.path.join(BASE_DIR, 'lgd-family-law-attorneys_kentucky-derby-article_20260308.docx')

CLIENT_NAME       = "Legal Growth Dynamics"
CLIENT_SUBTITLE   = "Family Law Attorney Market"
DOCUMENT_TITLE    = "THE DERBY DOESN'T START ON RACE DAY"
DOCUMENT_SUBTITLE = "What the Kentucky Derby reveals about the only marketing decision that matters this May"
MONTH_YEAR        = "May 2026"


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '25AAE1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = Pt(16)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(6)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(4)


def add_cover_page(doc):
    # Top accent bar
    accent = doc.add_paragraph()
    set_paragraph_shading(accent, '25AAE1')
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(0)
    r = accent.add_run(' ')
    r.font.size = Pt(6)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(DOCUMENT_TITLE)
    r.font.name = 'Calibri'
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = ASG_BLUE
    title_p.paragraph_format.space_before = Pt(48)
    title_p.paragraph_format.space_after  = Pt(12)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run(DOCUMENT_SUBTITLE)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    sub_p.paragraph_format.space_after = Pt(36)

    add_horizontal_rule(doc)

    for line, bold, color in [
        (f'Prepared exclusively for {CLIENT_NAME}', False, ASG_CHARCOAL),
        (CLIENT_SUBTITLE, False, ASG_CHARCOAL),
        ('', False, ASG_CHARCOAL),
        (f'{MONTH_YEAR}  |  Confidential', True, ASG_CHARCOAL),
        ('Produced by Naomi Patel, Blog Post Writer', False, ASG_CHARCOAL),
        ('Directed by Marcus Webb, Director of Content Marketing', False, ASG_CHARCOAL),
        ('Under the direction of Roger Bauer, Director', False, ASG_CHARCOAL),
        ('Authority Systems Group™', True, ASG_BLUE),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        f'Authority Systems Group™ \u2014 Confidential. Prepared exclusively for {CLIENT_NAME}.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def parse_and_render(doc, md_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block = False
    cover_done = False
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if not cover_done:
            if line.strip() == '---':
                cover_done = True
            i += 1
            continue

        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-\|:]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.paragraph_format.space_after = Pt(2)
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = WHITE
                            if ri == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '25AAE1')
                                tcPr.append(shd)
            continue

        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith('### ') and not line.startswith('#### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        if line.startswith('> '):
            text = line[2:].strip().replace('*', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'F0F8FD')
            p.paragraph_format.left_indent = Inches(0.3)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            parts = re.split(r'\*\*(.*?)\*\*', line[2:].strip())
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = 'Georgia'
                r.font.size = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = 'Georgia'
                r.font.size = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph — inline bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            italic_parts = re.split(r'\*(.*?)\*', part)
            for k, ipart in enumerate(italic_parts):
                if ipart:
                    r = p.add_run(ipart)
                    r.font.name = 'Georgia'
                    r.font.size = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
                    if k % 2 == 1:
                        r.font.italic = True
        p.paragraph_format.space_after = Pt(7)
        i += 1

    return doc


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    parse_and_render(doc, INPUT_MD)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
