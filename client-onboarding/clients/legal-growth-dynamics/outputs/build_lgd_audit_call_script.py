#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | Free Intake Audit Call — Working Script
March 2026
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
DARK_BG      = RGBColor(0x1A, 0x1A, 0x2E)
WARN_COLOR   = RGBColor(0xD4, 0x45, 0x00)   # amber-red for "listen for" notes

BASE_DIR    = os.path.dirname(__file__)
OUTPUT_DOCX = os.path.join(BASE_DIR, "docx", "lgd_intake-audit-call-script_20260309.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color="25AAE1", weight="12"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), weight)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body_paragraph(doc, text, space_after=9, indent=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        italic_parts = re.split(r'\*(.*?)\*', part)
        for k, ipart in enumerate(italic_parts):
            if ipart:
                r = p.add_run(ipart)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = color if color else BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
                if k % 2 == 1:
                    r.font.italic = True
    return p


def add_script_block(doc, text):
    """Shaded block for Roger's spoken script — light blue background."""
    lines = text.strip().split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        r = p.add_run(line if line else ' ')
        r.font.name   = 'Georgia'
        r.font.size   = Pt(11)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        set_paragraph_shading(p, 'E8F4FB')
        p.paragraph_format.left_indent  = Inches(0.35)
        p.paragraph_format.right_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(2) if i > 0 else Pt(10)
        p.paragraph_format.space_after  = Pt(2) if i < len(lines) - 1 else Pt(10)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '20')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '25AAE1')
        pBdr.append(left)
        pPr.append(pBdr)


def add_listen_for_block(doc, lines):
    """Orange left-border callout — diagnostic signals to listen for."""
    label_p = doc.add_paragraph()
    r = label_p.add_run('DIAGNOSTIC SIGNAL — LISTEN FOR:')
    r.font.name     = 'Calibri'
    r.font.size     = Pt(9)
    r.font.bold     = True
    r.font.all_caps = True
    r.font.color.rgb = WARN_COLOR
    set_paragraph_shading(label_p, 'FFF4EE')
    label_p.paragraph_format.left_indent  = Inches(0.3)
    label_p.paragraph_format.right_indent = Inches(0.3)
    label_p.paragraph_format.space_before = Pt(10)
    label_p.paragraph_format.space_after  = Pt(2)
    pPr  = label_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '20')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'D44500')
    pBdr.append(left)
    pPr.append(pBdr)

    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            if part:
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(10)
                r.font.color.rgb = RGBColor(0x5C, 0x1E, 0x00)
                if j % 2 == 1:
                    r.font.bold = True
        set_paragraph_shading(p, 'FFF4EE')
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2) if i < len(lines) - 1 else Pt(10)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '20')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), 'D44500')
        pBdr.append(left)
        pPr.append(pBdr)


def add_question(doc, text, sub=None):
    """Formatted interview question with optional follow-up note."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    r_dash = p.add_run('\u2014  ')
    r_dash.font.name  = 'Calibri'
    r_dash.font.size  = Pt(11)
    r_dash.font.bold  = True
    r_dash.font.color.rgb = ASG_BLUE
    r_text = p.add_run(text)
    r_text.font.name  = 'Georgia'
    r_text.font.size  = Pt(11)
    r_text.font.color.rgb = BODY_COLOR
    if sub:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.left_indent = Inches(0.65)
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after  = Pt(4)
        r_s = p_sub.add_run(sub)
        r_s.font.name   = 'Georgia'
        r_s.font.size   = Pt(9.5)
        r_s.font.italic = True
        r_s.font.color.rgb = ASG_CHARCOAL


def add_math_block(doc, rows):
    """Monospace-style calculator block for revenue gap math."""
    for i, (label, formula) in enumerate(rows):
        p = doc.add_paragraph()
        set_paragraph_shading(p, 'F0F8FF')
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(2) if i > 0 else Pt(8)
        p.paragraph_format.space_after  = Pt(2) if i < len(rows) - 1 else Pt(8)
        r_label = p.add_run(label.ljust(44))
        r_label.font.name  = 'Courier New'
        r_label.font.size  = Pt(10)
        r_label.font.color.rgb = BODY_COLOR
        r_formula = p.add_run(formula)
        r_formula.font.name  = 'Courier New'
        r_formula.font.size  = Pt(10)
        r_formula.font.bold  = True
        r_formula.font.color.rgb = ASG_BLUE
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '16')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '25AAE1')
        pBdr.append(left)
        pPr.append(pBdr)


def add_total_row(doc, label, formula):
    """Bold total row for Revenue Gap Calculator."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, 'D6F0FA')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    r_label = p.add_run(label.ljust(44))
    r_label.font.name  = 'Courier New'
    r_label.font.size  = Pt(10.5)
    r_label.font.bold  = True
    r_label.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    r_formula = p.add_run(formula)
    r_formula.font.name  = 'Courier New'
    r_formula.font.size  = Pt(10.5)
    r_formula.font.bold  = True
    r_formula.font.color.rgb = ASG_BLUE
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '25AAE1')
    pBdr.append(left)
    pPr.append(pBdr)


def add_timeline_table(doc):
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    widths = [Inches(1.0), Inches(1.5), Inches(3.5)]

    headers = ['BLOCK', 'TIME', 'PURPOSE']
    hdr_row = table.rows[0]
    for i, (cell, header, width) in enumerate(zip(hdr_row.cells, headers, widths)):
        cell.width = width
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        run.font.name     = 'Calibri'
        run.font.size     = Pt(9)
        run.font.bold     = True
        run.font.all_caps = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '25AAE1')
        tcPr.append(shd)

    data = [
        ('Open & Frame',              '0:00–3:00',   'Set tone; confirm it\'s a working session, not a pitch'),
        ('Practice Context',          '3:00–8:00',   'Understand where they are; confirm avatar fit'),
        ('Diagnostic — Speed to Lead','8:00–14:00',  'Uncover response infrastructure and contact rates'),
        ('Diagnostic — Intake Conv.', '14:00–21:00', 'Uncover consultation framework and conversion rate'),
        ('Diagnostic — Pipeline',     '21:00–25:00', 'Quantify the dormant asset they\'re sitting on'),
        ('Live Revenue Math',         '25:00–27:00', 'Run their numbers out loud — this is the over-deliver'),
        ('Synthesis & Next Step',     '27:00–30:00', 'Tell them what you heard; name the priority; offer the path'),
    ]

    for row_idx, (block, time, purpose) in enumerate(data, start=1):
        row = table.rows[row_idx]
        fill = 'F5FCFF' if row_idx % 2 == 0 else 'FFFFFF'
        for cell_idx, (cell, text) in enumerate(zip(row.cells, [block, time, purpose])):
            cell.width = widths[cell_idx]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.name  = 'Calibri' if cell_idx < 2 else 'Georgia'
            run.font.size  = Pt(10)
            run.font.bold  = cell_idx == 0
            run.font.color.rgb = ASG_BLUE if cell_idx == 0 else BODY_COLOR
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)


def add_priority_table(doc):
    data = [
        ('Speed to Lead',
         '"The cases are coming in. They\'re just not being captured before your competitors get to them."'),
        ('Intake Conversion',
         '"The leads are reaching you. The consultation is where the case is being won — and right now, it\'s being decided against you more often than it should be."'),
        ('Pipeline Nurturing',
         '"You have a database full of people who already trust you — past clients, warm leads, referral sources. They\'re just not hearing from you."'),
    ]
    table = doc.add_table(rows=len(data) + 1, cols=2)
    table.style = 'Table Grid'
    widths = [Inches(2.0), Inches(4.0)]

    hdr_row = table.rows[0]
    for cell, text, width in zip(hdr_row.cells, ['PRIORITY SYSTEM', 'OPENING LINE IN SYNTHESIS'], widths):
        cell.width = width
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(text)
        r.font.name     = 'Calibri'
        r.font.size     = Pt(9)
        r.font.bold     = True
        r.font.all_caps = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '58585A')
        tcPr.append(shd)

    for row_idx, (system, line) in enumerate(data, start=1):
        row = table.rows[row_idx]
        fill = 'F9F9F9' if row_idx % 2 == 0 else 'FFFFFF'
        for cell_idx, (cell, text) in enumerate(zip(row.cells, [system, line])):
            cell.width = widths[cell_idx]
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(text)
            r.font.name  = 'Georgia'
            r.font.size  = Pt(10)
            r.font.bold  = cell_idx == 0
            r.font.italic = cell_idx == 1
            r.font.color.rgb = ASG_BLUE if cell_idx == 0 else BODY_COLOR
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name  = 'Georgia'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.line_spacing = Pt(16)

    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(22)
    h1.font.bold  = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)

    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(16)
    h2.font.bold  = True
    h2.font.color.rgb = ASG_BLUE
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after  = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(4)


# ── cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    p_eye = doc.add_paragraph()
    p_eye.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_eye.add_run('INTERNAL — NOT CLIENT-FACING')
    r.font.name     = 'Calibri'
    r.font.size     = Pt(10)
    r.font.bold     = True
    r.font.all_caps = True
    r.font.color.rgb = WARN_COLOR
    p_eye.paragraph_format.space_after = Pt(10)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('Free Intake Audit Call\nWorking Script')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(30)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    p_title.paragraph_format.space_after = Pt(16)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run(
        'A 30-minute diagnostic framework for identifying where a family law practice\n'
        'is leaking HNW cases \u2014 and which system to fix first.'
    )
    r.font.name   = 'Calibri'
    r.font.size   = Pt(13)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sub.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    meta_lines = [
        ('Roger Bauer, Founder', False),
        ('Legal Growth Dynamics \u2014 Family Law Attorney Market', False),
        ('March 2026', False),
        ('', False),
        ('Produced by Authority Systems Group\u2122', True),
    ]
    for text, is_blue in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(11)
        r.font.bold  = is_blue
        r.font.color.rgb = ASG_BLUE if is_blue else ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ── footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '25AAE1')
    pBdr.append(top)
    pPr.append(pBdr)
    run = fp.add_run(
        'Authority Systems Group\u2122 \u2014 Internal Use Only. Legal Growth Dynamics | Audit Call Script v1.0'
    )
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── document body ─────────────────────────────────────────────────────────────

def add_pre_call(doc):
    doc.add_heading('Pre-Call Prep', level=1)
    add_horizontal_rule(doc)

    add_body_paragraph(doc, 'Before joining the call, confirm the following:')

    for item in [
        'Review any info the prospect submitted on the booking form',
        'Pull up their firm website \u2014 scan intake process, contact form, response language',
        'Check Google reviews \u2014 volume, recency, response pattern',
        'Open the **Metrics Capture Sheet** (Appendix A) \u2014 have it ready to fill in real time',
        'Have the **Revenue Gap Calculator** (Appendix B) open \u2014 you\u2019ll run the math live on the call',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        r_box = p.add_run('\u2610  ')
        r_box.font.name  = 'Calibri'
        r_box.font.size  = Pt(12)
        r_box.font.color.rgb = ASG_BLUE
        parts = re.split(r'\*\*(.*?)\*\*', item)
        for j, part in enumerate(parts):
            if part:
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.bold  = j % 2 == 1
                r.font.color.rgb = BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    doc.add_heading('What to Look for Before the Call Starts', level=3)
    for item in [
        'How long does their website contact form take to respond? Test it if time allows.',
        'Is there a chat widget? Does it display live hours or go silent after 5pm?',
        'Are their Google reviews recent, or did volume drop off? A gap in review recency signals a referral system that\u2019s been left to chance.',
        'Is their site written for HNW prospects or general family law? This tells you whether they\u2019ve identified their avatar or are marketing to everyone.',
    ]:
        add_question(doc, item)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_timeline(doc):
    doc.add_heading('Call Timeline \u2014 30 Minutes', level=1)
    add_horizontal_rule(doc)
    add_timeline_table(doc)


def add_section_1(doc):
    doc.add_heading('Section 1: Open & Frame', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('0:00\u20133:00  |  Tone: Confident, direct, peer-level. Not salesy.')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)

    add_horizontal_rule(doc, color="58585A", weight="6")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_script_block(doc,
        '"Thanks for making the time. I want to be straight with you about how I run these calls, '
        'because I think it\u2019s the opposite of what you\u2019ve probably experienced with agencies.\n'
        '\n'
        'I\u2019m not going to pitch you. I\u2019m not going to walk you through a deck. What I\u2019m '
        'going to do is ask you some direct questions about your intake process and your numbers \u2014 '
        'and at the end of 30 minutes, I\u2019ll tell you exactly what I see.\n'
        '\n'
        'You\u2019ll leave this call knowing more about where your practice is leaking than when you '
        'walked in. Whether or not we ever work together. That\u2019s the deal.\n'
        '\n'
        'Fair?"'
    )

    add_body_paragraph(doc,
        '*Wait for confirmation. Let the silence land.*',
        color=ASG_CHARCOAL)

    add_script_block(doc,
        '"Good. I read your site before this call, so I have some context. Let me ask you a few quick '
        'questions to make sure I\u2019m oriented correctly \u2014 then we\u2019re going to get into '
        'the actual numbers."'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_2(doc):
    doc.add_heading('Section 2: Practice Context', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('3:00\u20138:00  |  Confirm avatar fit. Establish baseline numbers.')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    doc.add_heading('Practice Overview', level=3)
    for q in [
        'How long have you been practicing, and how long have you been focusing on family law?',
        'Is it just you, or do you have associates? Walk me through your team structure real quick.',
        'How many new client consultations are you doing in a typical month \u2014 ballpark?',
    ]:
        add_question(doc, q)

    doc.add_heading('HNW Focus', level=3)
    for q, sub in [
        ('When you say high-net-worth cases \u2014 what does that look like for your practice? What\u2019s a typical matter value on your better cases?', None),
        ('What percentage of your current docket would you call genuinely HNW \u2014 meaning assets, complexity, and retainer size you\u2019d hold up as the work you want more of?', None),
        ('What does the rest of your docket look like? Is it by design or by default?', 'This reveals whether they\u2019re actively filtering or just taking what comes in.'),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('Current Marketing', level=3)
    for q in [
        'What are you doing right now to bring in new cases? PPC, referrals, organic \u2014 walk me through it.',
        'What are you spending monthly on paid marketing, roughly?',
        'Who\u2019s running it \u2014 in-house or an agency?',
    ]:
        add_question(doc, q)

    doc.add_heading('The Frustration', level=3)
    add_question(doc, 'What made you book this call? What\u2019s the specific thing that\u2019s not working the way it should?')
    add_body_paragraph(doc,
        '*Listen closely. This is the first signal of which system is most broken. '
        'Note their exact language \u2014 you\u2019ll use it verbatim in the synthesis.*',
        color=ASG_CHARCOAL, indent=0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_3(doc):
    doc.add_heading('Section 3: Diagnostic \u2014 Speed to Lead', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('8:00\u201314:00  |  Uncover the actual response infrastructure and contact rate.')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    doc.add_heading('The Response Process', level=3)
    for q, sub in [
        ('When someone submits a form on your website right now \u2014 what happens next? Walk me through the exact sequence.', None),
        ('Who gets the notification? A paralegal, you, a shared inbox, something else?', None),
        ('What does that first response look like \u2014 an auto-reply, a phone call, a personal email?', None),
        ('What\u2019s your target response time? And be honest \u2014 what does it actually average out to?', 'Target vs. reality gap is where the problem usually lives.'),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('After-Hours Reality', level=3)
    for q in [
        'What happens to a prospect who reaches out at 9pm on a Tuesday? Or Saturday morning?',
        'Is there anything automated that fires, or do they sit until someone gets in on Monday?',
        'What percentage of your inquiries come in outside business hours \u2014 do you know?',
    ]:
        add_question(doc, q)

    doc.add_heading('Contact Rate', level=3)
    add_question(doc,
        'Of all the leads that come in \u2014 forms, calls, referrals \u2014 what percentage do you actually end up speaking to?',
        sub='This is the single most important number in this section. If they don\u2019t know it, say: "That\u2019s common. We\u2019ll estimate it together."')
    add_question(doc, 'Do you track that number, or is that a guess?')

    doc.add_heading('Intake Staffing', level=3)
    for q in [
        'Who\u2019s the first human voice a prospect hears from \u2014 is it you, a paralegal, a dedicated intake person?',
        'Have they been trained specifically on HNW prospect conversations, or is intake handled the same way regardless of case type?',
    ]:
        add_question(doc, q)

    add_listen_for_block(doc, [
        '**No after-hours coverage** \u2014 automatic Speed to Lead gap; name it explicitly.',
        '**Response time >1 hour on average** \u2014 likely losing cases to faster competitors regardless of quality.',
        '**Contact rate <65%** \u2014 they\u2019re paying for leads that never become a conversation.',
        '**Intake handled by a paralegal with no HNW training** \u2014 Speed to Lead may not be the primary issue; the handoff moment may be where the case is lost.',
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_4(doc):
    doc.add_heading('Section 4: Diagnostic \u2014 Intake Conversion', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('14:00\u201321:00  |  Uncover the consultation framework and the actual conversion rate.')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    doc.add_heading('Consultation Volume and Structure', level=3)
    for q in [
        'How many consultations are you doing per month right now?',
        'Is it a paid consult or free? What does the format look like \u2014 phone, video, in person?',
        'Are you running all the consultations yourself, or does your team handle some?',
    ]:
        add_question(doc, q)

    doc.add_heading('The Consultation Framework', level=3)
    for q, sub in [
        ('When you\u2019re in a consultation with a HNW prospect \u2014 is there a structured framework your team follows, or does it depend on the attorney?',
         'Listen for: "it depends on the attorney" = no framework = conversion rate is personality-dependent.'),
        ('What\u2019s the objective of the consultation in your mind \u2014 is it primarily diagnostic, evaluative, or a close?',
         'Most attorneys say "diagnostic." High-converting firms treat it as all three simultaneously.'),
        ('Where do most consultations break down? Where do prospects go cold?', None),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('The Conversion Rate \u2014 The Most Important Number on This Call', level=3)

    add_script_block(doc,
        '"Here\u2019s a number most attorneys have never actually calculated. Of every ten people who sit '
        'across from you in a consultation \u2014 how many sign a retainer? Not eventually. That day or '
        'within a few days."'
    )

    for q, sub in [
        ('What\u2019s your current consultation-to-retention rate, as best you know it?',
         'Industry average for HNW family law: 35\u201340%. High-performing firms: 50%+. Record this number.'),
        ('If you don\u2019t know the exact number \u2014 of your last 10 consultations, how many became clients?', None),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('Post-Consultation Follow-Up', level=3)
    for q, sub in [
        ('What happens to someone who comes in for a consultation and doesn\u2019t sign that day?', None),
        ('Do you have an automated sequence, or does someone make manual calls?', None),
        ('How long do you follow up before a prospect goes cold in your system?', None),
        ('Of the consultations that don\u2019t sign immediately \u2014 what percentage eventually do? Do you track that?',
         'Most firms have no idea. This gap alone often represents 10\u201315% of closeable revenue.'),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('The Consultation Pre-Frame', level=3)
    for q in [
        'What do prospects receive before they come in? Anything that sets expectations or positions your firm?',
        'Is there any communication between the time they book and the time they show up \u2014 besides a calendar reminder?',
    ]:
        add_question(doc, q)

    add_listen_for_block(doc, [
        '**No pre-consult communication** \u2014 missed trust-building window; the prospect arrives cold.',
        '**No structured framework** \u2014 conversion depends entirely on the attorney\u2019s instincts that day.',
        '**Single manual follow-up** \u2014 likely losing 15\u201320% of consultable prospects who needed more time.',
        '**Consultation-to-retention rate below 40%** \u2014 Intake Conversion is the primary system gap.',
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_5(doc):
    doc.add_heading('Section 5: Diagnostic \u2014 Pipeline Nurturing', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('21:00\u201325:00  |  Quantify the dormant asset they\u2019re sitting on.')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    doc.add_heading('Cold Leads', level=3)
    for q, sub in [
        ('Leads that come in and never schedule a consultation \u2014 what happens to them in your system? CRM, spreadsheet, or do they fall off?', None),
        ('Roughly how many cold leads do you accumulate in a typical month?', None),
        ('Has anyone on your team ever run a reactivation campaign to those cold leads \u2014 a sequence, a call campaign, anything?',
         'The answer is almost always no. Don\u2019t be surprised by it.'),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('Post-Consultation Non-Converts', level=3)
    for q in [
        'The prospects who consulted but didn\u2019t sign \u2014 do they stay active in your system, or do they go dark after a few calls?',
        'How many of those would you estimate are sitting in your database right now?',
    ]:
        add_question(doc, q)

    doc.add_heading('Past Clients', level=3)
    for q, sub in [
        ('How many clients have you served in the last three to five years \u2014 rough estimate?', None),
        ('How are you staying in contact with them? Email newsletter, personal outreach, anything structured?', None),
        ('In the last 12 months, how many past client referrals have you received? Is that number growing, flat, or declining?',
         'Declining referrals from past clients = relationship maintenance gap, not a reputation problem.'),
    ]:
        add_question(doc, q, sub)

    doc.add_heading('Referral Network', level=3)
    for q, sub in [
        ('Do you have active relationships with estate attorneys, CPAs, financial advisors \u2014 the professionals who work with the same client base you do?', None),
        ('Is there a systematic approach to those relationships, or is it more organic?', None),
        ('How many active referral source relationships do you have right now?',
         'Under 5 = completely organic. 5\u201315 = partially developed. 15+ = a real referral infrastructure exists.'),
    ]:
        add_question(doc, q, sub)

    add_listen_for_block(doc, [
        '**No CRM or disorganized database** \u2014 the pipeline leaks at every seam; volume is unknowable.',
        '**No reactivation history** \u2014 high-value dormant asset sitting untouched; typically the fastest cash opportunity.',
        '**Past clients not hearing from them** \u2014 referral revenue is being forfeited every month by inaction.',
        '**Organic referral relationships only** \u2014 Pipeline Nurturing is the highest-leverage long-term system gap.',
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_6(doc):
    doc.add_heading('Section 6: Live Revenue Math', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run(
        '25:00\u201327:00  |  Run this out loud, with the prospect on the call. '
        'This is the over-deliver. Use their numbers.'
    )
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    add_script_block(doc,
        '"Okay. Let me show you what I\u2019m seeing in the numbers. This is rough \u2014 we\u2019re using '
        'estimates \u2014 but this is how I want you to think about what these systems are worth to '
        'your practice."'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_heading('Revenue Gap 1 \u2014 Speed to Lead', level=3)
    add_math_block(doc, [
        ('Monthly inquiries:',              '[ A ] = ______'),
        ('Current contact rate (%):',       '[ B ] = ______'),
        ('Target contact rate:',            '[ C ] = 80%'),
        ('Additional contacts/month:',      '[ D ] = A \u00d7 (C \u2212 B) = ______'),
        ('Consult-to-hire rate (%):',       '[ E ] = ______'),
        ('Additional retentions/month:',    '[ F ] = D \u00d7 E = ______'),
        ('Average case value:',             '[ G ] = $______'),
    ])
    add_total_row(doc, 'SPEED TO LEAD MONTHLY GAP:', '[ H ] = F \u00d7 G = $______')

    add_script_block(doc,
        '"You told me you\u2019re getting roughly [A] inquiries per month and your contact rate is around [B]%. '
        'If we move that contact rate to 80%, that\u2019s [D] additional conversations per month. '
        'At your conversion rate of [E]% and an average case value of $[G], that\u2019s $[H] in '
        'additional monthly revenue \u2014 from leads you\u2019re already paying for."'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.add_heading('Revenue Gap 2 \u2014 Intake Conversion', level=3)
    add_math_block(doc, [
        ('Monthly consultations:',              '[ A ] = ______'),
        ('Current consult-to-hire rate (%):',   '[ B ] = ______'),
        ('Benchmark rate:',                     '[ C ] = 50%'),
        ('Conversion gap (%):',                 '[ D ] = C \u2212 B = ______'),
        ('Additional retentions/month:',        '[ E ] = A \u00d7 D = ______'),
        ('Average case value:',                 '[ F ] = $______'),
    ])
    add_total_row(doc, 'INTAKE CONVERSION MONTHLY GAP:', '[ G ] = E \u00d7 F = $______')

    add_script_block(doc,
        '"Your consultation-to-retention rate is around [B]%. The benchmark for a well-run HNW intake '
        'system is 50%. You\u2019re doing [A] consultations per month. The gap between [B]% and 50% is '
        '[D]% \u2014 that\u2019s [E] additional retained clients per month on the same consultation volume. '
        'At $[F] average case value, that\u2019s $[G] every month."'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.add_heading('Revenue Gap 3 \u2014 Pipeline Nurturing', level=3)
    add_math_block(doc, [
        ('Cold leads + non-converts:',      '[ A ] = ______'),
        ('Past clients in database:',       '[ B ] = ______'),
        ('Total addressable database:',     '[ C ] = A + B = ______'),
        ('Conservative reactivation rate:', '[ D ] = 4%'),
        ('Reactivation cases:',             '[ E ] = C \u00d7 D = ______'),
        ('Average case value:',             '[ F ] = $______'),
        ('One-time opportunity:',           '[ G ] = E \u00d7 F = $______'),
    ])
    add_total_row(doc, 'PIPELINE MONTHLY EQUIVALENT:', '[ H ] = G \u00f7 12 = $______/mo')

    add_script_block(doc,
        '"You have roughly [C] people in your database who already know you. A structured reactivation '
        'campaign on a warm list converts at 3 to 7%. At 4%, that\u2019s [E] cases from assets you '
        'already have. At $[F] per case, that\u2019s $[G]. '
        'That\u2019s not a marketing number. That\u2019s a database number."'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.add_heading('Total Revenue Gap', level=3)
    add_math_block(doc, [
        ('Speed to Lead monthly gap:',      '$______ (Gap 1)'),
        ('Intake Conversion monthly gap:',  '$______ (Gap 2)'),
        ('Pipeline monthly equivalent:',    '$______ (Gap 3)'),
    ])
    add_total_row(doc, 'TOTAL ESTIMATED MONTHLY GAP:', '$______')
    add_math_block(doc, [
        ('Conservative capture (50%):',     '$______/month'),
        ('Annualized at 50% capture:',      '$______/year'),
    ])

    add_script_block(doc,
        '"Add those three together. That\u2019s the range of what these systems are worth to your '
        'practice every month if they\u2019re built correctly. I\u2019m not telling you we can capture '
        'all of it. But if we capture half \u2014 that\u2019s [total \u00d7 50%] in monthly revenue. '
        'And none of it requires you to spend more on ads."\n'
        '\n'
        '*Pause. Let the math settle.*'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_section_7(doc):
    doc.add_heading('Section 7: Synthesis & Recommendation', level=1)
    p_time = doc.add_paragraph()
    r = p_time.add_run('27:00\u201330:00  |  Restate what you heard. Name the priority. Don\u2019t hedge.')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_time.paragraph_format.space_after = Pt(10)
    add_horizontal_rule(doc, color="58585A", weight="6")

    add_body_paragraph(doc,
        'Use the **Prioritization Decision Tree** (Appendix C) to identify the lead system. '
        'Then run this script:')

    add_script_block(doc,
        '"Here\u2019s what I heard on this call.\n'
        '\n'
        '[Restate 2\u20133 specific things they said \u2014 use their exact words from the '
        '\u2018The Frustration\u2019 section.]\n'
        '\n'
        'Based on what you\u2019ve told me, the system that\u2019s costing you the most right now '
        'is [Speed to Lead / Intake Conversion / Pipeline Nurturing].\n'
        '\n'
        'The reason I say that is [specific evidence: their response time / their conversion rate / '
        'their database situation].\n'
        '\n'
        'Before anything else \u2014 before we talk about ads, before we talk about SEO, before we '
        'change anything else \u2014 that system needs to be built. Everything else amplifies a '
        'broken foundation.\n'
        '\n'
        'What would change in your practice if [describe the fixed system in one sentence]?"'
    )

    add_body_paragraph(doc,
        '*Let them answer. That answer is your buy signal.*',
        color=ASG_CHARCOAL, indent=0.3)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_heading('The Honest Close', level=3)

    add_script_block(doc,
        '"Here\u2019s where I land.\n'
        '\n'
        'What you\u2019ve described today is very fixable. The gap isn\u2019t your marketing, '
        'it isn\u2019t your skills, and it isn\u2019t your market. It\u2019s the infrastructure '
        'between the lead and the retainer.\n'
        '\n'
        'If you want to understand what building that looks like for your specific practice \u2014 '
        'what the scope is, what the sequence is, what the investment looks like \u2014 the next '
        'conversation is a strategy session where we go deeper.\n'
        '\n'
        'If you\u2019re not there yet, I understand that completely. But I want to leave you with '
        'this: the [primary gap system] we talked about today is costing you $[revenue gap number] '
        'every month it goes unfixed. That math doesn\u2019t change.\n'
        '\n'
        'Whatever you decide to do with that, today was worth it."'
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_heading('If They Want to Move Forward', level=3)

    add_script_block(doc,
        '"Then let\u2019s schedule a strategy session. It\u2019s typically 60 minutes \u2014 '
        'no charge. I\u2019ll do some prep work on your market before we talk, and we\u2019ll '
        'map out exactly what a system-first approach looks like for your firm.\n'
        '\n'
        'Does [specific date/time] work for you?"'
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_appendix_a(doc):
    doc.add_page_break()
    doc.add_heading('Appendix A \u2014 Metrics Capture Sheet', level=1)
    add_horizontal_rule(doc)
    add_body_paragraph(doc, '*Fill in real time during the call. These numbers feed the Revenue Gap Calculator.*', color=ASG_CHARCOAL)

    sections = [
        ('PRACTICE BASICS', [
            'Firm name:', 'Years in practice:', 'Team size (attorneys + staff):',
            'Monthly consultations:', 'Average HNW case value: $', 'Monthly marketing spend: $',
            'Primary marketing channels:',
        ]),
        ('SPEED TO LEAD', [
            'Monthly total inquiries:', 'Estimated contact rate: ____%',
            'Avg response time (honest):', 'After-hours coverage:  Yes / No / Partial',
            'Intake handled by:', 'HNW-specific intake training:  Yes / No',
        ]),
        ('INTAKE CONVERSION', [
            'Monthly paid consultations:', 'Consultation-to-retention rate: ____%  (est: ___ of last 10)',
            'Consultation framework:  Structured / Ad hoc / Depends on attorney',
            'Pre-consult communication:  Yes / No',
            'Post-consult follow-up:  Automated / Manual / None',
            'Follow-up duration (days):',
        ]),
        ('PIPELINE NURTURING', [
            'Cold leads in database:', 'Post-consult non-converts:', 'Past clients (3\u20135 year window):',
            'Past client contact cadence:  Regular / Occasional / None',
            'Referral source relationships:  Systematic / Organic / None',
            'Active referral source count:',
            'Reactivation campaigns run:  Yes / No / Never',
        ]),
    ]

    for section_title, fields in sections:
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(12)
        p_hdr.paragraph_format.space_after  = Pt(4)
        r = p_hdr.add_run(section_title)
        r.font.name     = 'Calibri'
        r.font.size     = Pt(10)
        r.font.bold     = True
        r.font.all_caps = True
        r.font.color.rgb = ASG_BLUE
        add_horizontal_rule(doc, color="25AAE1", weight="6")

        for field in fields:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            r_f = p.add_run(field + '  ')
            r_f.font.name  = 'Georgia'
            r_f.font.size  = Pt(10)
            r_f.font.color.rgb = BODY_COLOR
            r_line = p.add_run('_' * 30)
            r_line.font.name  = 'Georgia'
            r_line.font.size  = Pt(10)
            r_line.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # The Frustration box
    p_hdr2 = doc.add_paragraph()
    p_hdr2.paragraph_format.space_before = Pt(12)
    p_hdr2.paragraph_format.space_after  = Pt(4)
    r = p_hdr2.add_run('THE FRUSTRATION \u2014 Quote Their Exact Words')
    r.font.name     = 'Calibri'
    r.font.size     = Pt(10)
    r.font.bold     = True
    r.font.all_caps = True
    r.font.color.rgb = WARN_COLOR
    add_horizontal_rule(doc, color="D44500", weight="6")

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(16)
        r_line = p.add_run('_' * 72)
        r_line.font.name  = 'Georgia'
        r_line.font.size  = Pt(10)
        r_line.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)


def add_appendix_b(doc):
    doc.add_page_break()
    doc.add_heading('Appendix B \u2014 Revenue Gap Calculator', level=1)
    add_horizontal_rule(doc)
    add_body_paragraph(doc, '*Use the Metrics Capture Sheet inputs. Run this live on the call.*', color=ASG_CHARCOAL)

    doc.add_heading('Gap 1: Speed to Lead', level=3)
    add_math_block(doc, [
        ('Monthly inquiries:',              '[ A ] = ______'),
        ('Current contact rate (%):',       '[ B ] = ______'),
        ('Target contact rate:',            '[ C ] = 80%'),
        ('Additional contacts/month:',      '[ D ] = A \u00d7 (C \u2212 B) = ______'),
        ('Consult-to-hire rate (%):',       '[ E ] = ______'),
        ('Additional retentions/month:',    '[ F ] = D \u00d7 E = ______'),
        ('Average case value:',             '[ G ] = $______'),
    ])
    add_total_row(doc, 'SPEED TO LEAD MONTHLY GAP:', '[ H ] = F \u00d7 G = $______')

    doc.add_heading('Gap 2: Intake Conversion', level=3)
    add_math_block(doc, [
        ('Monthly consultations:',              '[ A ] = ______'),
        ('Current consult-to-hire rate (%):',   '[ B ] = ______'),
        ('Benchmark rate:',                     '[ C ] = 50%'),
        ('Conversion gap (%):',                 '[ D ] = C \u2212 B = ______'),
        ('Additional retentions/month:',        '[ E ] = A \u00d7 D = ______'),
        ('Average case value:',                 '[ F ] = $______'),
    ])
    add_total_row(doc, 'INTAKE CONVERSION MONTHLY GAP:', '[ G ] = E \u00d7 F = $______')

    doc.add_heading('Gap 3: Pipeline Nurturing', level=3)
    add_math_block(doc, [
        ('Cold leads + non-converts:',      '[ A ] = ______'),
        ('Past clients in database:',       '[ B ] = ______'),
        ('Total addressable database:',     '[ C ] = A + B = ______'),
        ('Conservative reactivation rate:', '[ D ] = 4%'),
        ('Reactivation cases:',             '[ E ] = C \u00d7 D = ______'),
        ('Average case value:',             '[ F ] = $______'),
        ('One-time opportunity:',           '[ G ] = E \u00d7 F = $______'),
    ])
    add_total_row(doc, 'PIPELINE MONTHLY EQUIVALENT:', '[ H ] = G \u00f7 12 = $______/mo')

    doc.add_heading('Total Revenue Gap', level=3)
    add_math_block(doc, [
        ('Speed to Lead monthly gap:',      '$______ (H, Gap 1)'),
        ('Intake Conversion monthly gap:',  '$______ (G, Gap 2)'),
        ('Pipeline monthly equivalent:',    '$______ (H, Gap 3)'),
    ])
    add_total_row(doc, 'TOTAL ESTIMATED MONTHLY GAP:', '$______')
    add_math_block(doc, [
        ('Conservative capture (50%):',     '$______/month'),
        ('Annualized at 50% capture:',      '$______/year'),
    ])


def add_appendix_c(doc):
    doc.add_page_break()
    doc.add_heading('Appendix C \u2014 Prioritization Decision Tree', level=1)
    add_horizontal_rule(doc)
    add_body_paragraph(doc, 'Use this during the synthesis to identify which system to lead with.')

    tree_items = [
        ('Contact rate below 65%?',
         'YES \u2192 **Speed to Lead is Priority 1.** They\u2019re losing leads before any conversation happens. Fix the infrastructure first \u2014 everything else is moot.'),
        ('Response time over 1 hour on average?',
         'YES \u2192 **Speed to Lead is Priority 1.** HNW prospects are gone within hours. Speed is the first fix.'),
        ('Consultation-to-hire rate below 40%?',
         'YES \u2192 **Intake Conversion is Priority 1.** High contact rate, but cases are walking. The leak is in the room.'),
        ('Post-consultation follow-up minimal or nonexistent?',
         'YES \u2192 **Intake Conversion is Priority 1.** They\u2019re abandoning cases that needed one more touchpoint.'),
        ('Database of 100+ contacts with no systematic nurture?',
         'YES + contact rate and conversion are acceptable \u2192 **Pipeline Nurturing is Priority 1.** Fastest cash from existing assets.'),
        ('All three gaps are roughly equal?',
         'DEFAULT sequence: Speed to Lead \u2192 Intake Conversion \u2192 Pipeline Nurturing. Build in this order. Each system supports the next.'),
    ]

    for i, (trigger, outcome) in enumerate(tree_items):
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(10)
        p_q.paragraph_format.space_after  = Pt(3)
        r_num = p_q.add_run(f'{i+1}.  ')
        r_num.font.name  = 'Calibri'
        r_num.font.size  = Pt(11)
        r_num.font.bold  = True
        r_num.font.color.rgb = ASG_BLUE
        r_q = p_q.add_run(trigger)
        r_q.font.name  = 'Calibri'
        r_q.font.size  = Pt(11)
        r_q.font.bold  = True
        r_q.font.color.rgb = BODY_COLOR

        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.45)
        p_a.paragraph_format.space_after = Pt(6)
        parts = re.split(r'\*\*(.*?)\*\*', outcome)
        for j, part in enumerate(parts):
            if part:
                r = p_a.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.bold  = j % 2 == 1
                r.font.color.rgb = ASG_BLUE if j % 2 == 1 else BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    doc.add_heading('Priority Framing by System', level=3)
    add_priority_table(doc)


def add_appendix_d(doc):
    doc.add_heading('Appendix D \u2014 Post-Call Protocol', level=1)
    add_horizontal_rule(doc)

    add_body_paragraph(doc, '**Immediately after the call:**')
    for item in [
        'Complete Metrics Capture Sheet if any fields are blank',
        'Run final Revenue Gap Calculator totals and record the numbers',
        'Note primary system gap and the specific evidence that points to it',
        'Record their exact words from the \u201cThe Frustration\u201d section',
        'If moving to strategy session: prep the market brief before the 60-min call',
        'If not ready: add to 30-day follow-up sequence \u2014 do not let them go cold',
        'If not the right avatar: close graciously and do not pursue',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        r_box = p.add_run('\u2610  ')
        r_box.font.name  = 'Calibri'
        r_box.font.size  = Pt(12)
        r_box.font.color.rgb = ASG_BLUE
        r_t = p.add_run(item)
        r_t.font.name  = 'Georgia'
        r_t.font.size  = Pt(11)
        r_t.font.color.rgb = BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    doc.add_heading('Qualification Gate', level=3)
    add_body_paragraph(doc, 'Before booking a strategy session, confirm all of the following:')

    for item in [
        '3+ years in practice',
        'At least some HNW family law in current docket, or explicit intent to shift',
        'Monthly marketing spend or referral volume sufficient to justify system investment',
        'Practice owner or primary decision-maker was on this call',
        'Their stated frustration aligns with a system gap (not a fundamentally different problem)',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        r_box = p.add_run('\u2610  ')
        r_box.font.name  = 'Calibri'
        r_box.font.size  = Pt(12)
        r_box.font.color.rgb = ASG_BLUE
        r_t = p.add_run(item)
        r_t.font.name  = 'Georgia'
        r_t.font.size  = Pt(11)
        r_t.font.color.rgb = BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_horizontal_rule(doc)
    p_sig = doc.add_paragraph()
    r = p_sig.add_run(
        'Authority Systems Group\u2122 \u2014 Internal Operations Document\n'
        'Legal Growth Dynamics | Roger Bauer, Director\n'
        'Version 1.0 \u2014 March 2026'
    )
    r.font.name   = 'Calibri'
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sig.paragraph_format.space_after = Pt(8)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    add_pre_call(doc)
    add_timeline(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_appendix_a(doc)
    add_appendix_b(doc)
    add_appendix_c(doc)
    add_appendix_d(doc)

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f'\u2713 Document saved: {OUTPUT_DOCX}')


if __name__ == '__main__':
    main()
