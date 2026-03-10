#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | Mid-Funnel Report
"The Three Systems Your Family Law Practice Is Missing"
March 2026
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK     = RGBColor(0x1A, 0x1A, 0x1A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)

BASE_DIR    = os.path.dirname(__file__)
OUTPUT_DOCX = os.path.join(BASE_DIR, "docx", "lgd_three-missing-systems-report_20260309.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color="25AAE1"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body_paragraph(doc, text, space_after=10, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        italic_parts = re.split(r'\*(.*?)\*', part)
        for k, ipart in enumerate(italic_parts):
            if ipart:
                r = p.add_run(ipart)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
                if k % 2 == 1:
                    r.font.italic = True
    return p


def add_callout_box(doc, text):
    """Blue left-border callout — used for 'What Now' action items."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = 'Georgia'
    r.font.size  = Pt(11)
    r.font.color.rgb = BODY_COLOR
    set_paragraph_shading(p, 'EAF6FC')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '25AAE1')
    pBdr.append(left)
    pPr.append(pBdr)


def add_callout_lines(doc, lines):
    """Multi-line callout box — accepts a list of strings."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            if part:
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
        set_paragraph_shading(p, 'EAF6FC')
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(2) if i > 0 else Pt(10)
        p.paragraph_format.space_after  = Pt(2) if i < len(lines) - 1 else Pt(10)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '24')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '25AAE1')
        pBdr.append(left)
        pPr.append(pBdr)


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name  = 'Georgia'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.line_spacing = Pt(16)

    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(26)
    h1.font.bold  = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after  = Pt(10)

    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(18)
    h2.font.bold  = True
    h2.font.color.rgb = ASG_BLUE
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after  = Pt(4)

    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(13)
    h3.font.bold  = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(4)


# ── cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    add_horizontal_rule(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(56)

    p_eye = doc.add_paragraph()
    p_eye.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_eye.add_run('MID-FUNNEL REPORT')
    r.font.name     = 'Calibri'
    r.font.size     = Pt(11)
    r.font.bold     = True
    r.font.all_caps = True
    r.font.color.rgb = ASG_CHARCOAL
    p_eye.paragraph_format.space_after = Pt(14)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(
        'The Three Systems Your\nFamily Law Practice Is Missing'
    )
    r.font.name  = 'Calibri'
    r.font.size  = Pt(32)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    p_title.paragraph_format.space_after = Pt(20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run(
        'Why you\u2019re generating leads and still not signing\n'
        'the cases you want \u2014 and what to do about it.'
    )
    r.font.name   = 'Calibri'
    r.font.size   = Pt(14)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sub.paragraph_format.space_after = Pt(56)

    add_horizontal_rule(doc)

    meta_lines = [
        ('Roger Bauer, Founder', False),
        ('Legal Growth Dynamics \u2014 Family Law Attorney Market', False),
        ('March 2026  |  Confidential', False),
        ('', False),
        ('Produced by Authority Systems Group\u2122', True),
    ]
    for text, is_blue in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(11)
        r.font.bold  = is_blue
        r.font.color.rgb = ASG_BLUE if is_blue else ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ── footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '25AAE1')
    pBdr.append(top)
    pPr.append(pBdr)

    run = fp.add_run(
        'Authority Systems Group\u2122 \u2014 Confidential. Prepared exclusively for Legal Growth Dynamics.'
    )
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── report body ───────────────────────────────────────────────────────────────

def add_report_body(doc):

    # ── Report H1 + standfirst ──
    doc.add_heading('The Three Systems Your Family Law Practice Is Missing', level=1)

    p_sub = doc.add_paragraph()
    r = p_sub.add_run(
        'Why you\u2019re generating leads and still not signing the cases you want \u2014 '
        'and what to do about it.'
    )
    r.font.name   = 'Georgia'
    r.font.size   = Pt(12)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sub.paragraph_format.space_after = Pt(16)

    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Opening ──
    add_body_paragraph(doc,
        'If you\u2019ve been practicing family law for more than three years, you\u2019ve probably already '
        'spent money on marketing. Maybe a lot of it. Maybe you\u2019re running PPC right now and getting '
        'a decent volume of inquiries. The phone rings. Forms come in. The dashboard your agency sends '
        'every Friday looks fine.')

    add_body_paragraph(doc,
        'And yet \u2014 the HNW cases aren\u2019t converting the way they should.')

    add_body_paragraph(doc,
        'The high-asset divorces. The complex custody matters. The cases that actually move your revenue '
        'number. Those prospects are calling. Some are calling multiple firms. And a meaningful number of '
        'them are ending up somewhere else.')

    add_body_paragraph(doc,
        'The standard explanation from most agencies is that you need better targeting, more impressions, '
        'or a bigger budget. More leads will fix it.')

    add_body_paragraph(doc,
        'Here\u2019s the thing \u2014 they\u2019re wrong. And they\u2019re not wrong because they\u2019re '
        'stupid. They\u2019re wrong because their incentives point in a completely different direction from '
        'yours. They get paid when you run ads. They don\u2019t get paid when you sign cases.')

    add_body_paragraph(doc,
        'What I\u2019ve seen \u2014 from ten years sitting across the table from attorneys in the legal '
        'services industry \u2014 is that the gap between firms that consistently attract and convert HNW '
        'family law clients, and firms that don\u2019t, almost never comes down to marketing spend.')

    add_body_paragraph(doc, 'It comes down to three systems.')

    add_body_paragraph(doc, 'Most family law firms have none of them.', space_after=20)

    add_horizontal_rule(doc, color="58585A")

    # ═══════════════════════════════════════════════════════════════════════════
    # SYSTEM 1
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_heading('System 1: Speed to Lead', level=2)

    # Why
    doc.add_heading('The Race You Don\u2019t Know You\u2019re In', level=3)

    add_body_paragraph(doc,
        'HNW prospects don\u2019t behave like general family law clients.')

    add_body_paragraph(doc,
        'When someone facing a high-asset divorce starts looking for an attorney, they\u2019re not '
        'comparison-shopping on price. They\u2019re evaluating trust, competence, and \u2014 whether '
        'they\u2019d consciously say this or not \u2014 **who made them feel most confident, most quickly.**')

    add_body_paragraph(doc,
        'The research on this is uncomfortable but consistent. A prospect who reaches out and hears back '
        'within five minutes is dramatically more likely to schedule a consultation than one who hears back '
        'in two hours. The firms that win these cases are not always the best. They\u2019re **first.**')

    add_body_paragraph(doc,
        'Here\u2019s what most family law firms are operating with instead: voicemail, callback windows, '
        'and a paralegal who gets back to people when they can. After hours? Nothing until morning. '
        'Weekend? Monday at best.')

    add_body_paragraph(doc,
        'HNW clients contact two to four attorneys before making a decision. You know what every '
        'competitor is doing too? The same slow intake process. Which means the first firm to reach out '
        'with any kind of confidence and structure has a significant advantage that has nothing to do with '
        'reputation, skill, or ad spend.')

    add_body_paragraph(doc,
        'You\u2019re probably losing cases right now \u2014 not because you\u2019re not good enough, '
        'but because you\u2019re not fast enough.')

    # What
    doc.add_heading('What an Actual First Response Looks Like', level=3)

    add_body_paragraph(doc,
        'Speed-to-lead is an automated first-response system that ensures every prospect who reaches out '
        '\u2014 through a form, a call, a chat \u2014 gets an immediate, intelligent, branded response. '
        'Not an auto-reply that says \u201cthanks for reaching out, someone will contact you shortly.\u201d')

    add_body_paragraph(doc,
        'A real response. One that establishes authority, sets expectations, and begins building the case '
        'for why your firm is the right choice \u2014 before a human being ever picks up the phone.')

    # How
    doc.add_heading('The Four Things That Fire When a Prospect Reaches Out', level=3)

    add_body_paragraph(doc, 'When a prospect submits a form or leaves a voicemail, the system triggers:')

    for item in [
        'An immediate personalized text and email acknowledging their inquiry \u2014 within minutes, '
        'any hour of the day',
        'A brief, authoritative message that reflects your positioning (not a generic acknowledgment)',
        'A follow-up sequence if they don\u2019t respond, spaced over the next 24\u201348 hours',
        'A notification to your intake team with full prospect context so the first live conversation '
        'is warm, not cold',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(5)
        r_dash = p.add_run('\u2014  ')
        r_dash.font.name  = 'Georgia'
        r_dash.font.size  = Pt(11)
        r_dash.font.bold  = True
        r_dash.font.color.rgb = ASG_BLUE
        r_text = p.add_run(item)
        r_text.font.name  = 'Georgia'
        r_text.font.size  = Pt(11)
        r_text.font.color.rgb = BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_body_paragraph(doc,
        'The system doesn\u2019t replace your intake team. It makes sure no prospect ever falls into a '
        'gap between inquiry and contact. It levels a playing field that most of your competitors '
        'don\u2019t even know they\u2019re losing on.')

    # What Now
    doc.add_heading('Look at Your Last 30 Inquiries', level=3)

    add_callout_lines(doc, [
        'Pull up your last 30 inquiries. For each one, answer honestly:',
        '\u2014  What was the actual response time? Not your policy \u2014 the real number.',
        '\u2014  What did that first response say?',
        '\u2014  Did every inquiry get followed up, or did some fall through the cracks?',
        '',
        'If you can\u2019t answer those questions with certainty, your speed-to-lead system is '
        'costing you cases. Most firms, when they actually look, find they\u2019re running at '
        '40\u201360% contact rates. That means four out of ten prospects who reach out are never '
        'having a real conversation with anyone at your firm.',
    ])

    add_horizontal_rule(doc, color="58585A")

    # ═══════════════════════════════════════════════════════════════════════════
    # SYSTEM 2
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_heading('System 2: Intake Conversion', level=2)

    # Why
    doc.add_heading('The First Call Is an Audition', level=3)

    add_body_paragraph(doc, 'Getting someone on the phone is only the beginning.')

    add_body_paragraph(doc,
        'The first conversation your firm has with a HNW prospect is not a fact-gathering exercise. '
        'It\u2019s an audition. They\u2019re evaluating your competence, your confidence, and whether '
        'they believe you understand the stakes of their situation before they hand you a $25,000 retainer.')

    add_body_paragraph(doc,
        'Most family law intake processes are built for a different client entirely. They\u2019re designed '
        'to gather case information \u2014 what are the assets involved, are there children, has a petition '
        'been filed. That\u2019s useful data. But it is completely the wrong frame for the first '
        'conversation with a high-net-worth prospect.')

    add_body_paragraph(doc,
        'HNW clients don\u2019t want to be processed. They want to be understood. They want to feel like '
        'whoever they\u2019re talking to has seen this situation before, knows what it means, and is '
        'capable of navigating it. When that doesn\u2019t happen \u2014 when the intake feels '
        'transactional, rushed, or generic \u2014 they leave. They go back to their list and '
        'call the next name.')

    add_body_paragraph(doc,
        'Your intake team may be excellent. Warm, professional, knowledgeable. But if they haven\u2019t '
        'been trained specifically on high-net-worth consultation frameworks, they\u2019re having '
        'the wrong conversation.')

    # What
    doc.add_heading('A Framework, Not a Script', level=3)

    add_body_paragraph(doc,
        'Intake conversion is the structured process your firm uses to move a prospect from first contact '
        'to signed retainer \u2014 with a framework designed specifically for HNW clients.')

    add_body_paragraph(doc,
        'It\u2019s not a script. It\u2019s a conversation architecture. A way of structuring the '
        'engagement so the prospect feels heard, the stakes are acknowledged, and your firm\u2019s '
        'expertise is communicated before the prospect ever has to ask about it.')

    # How
    doc.add_heading('The Four Components of a Conversion System', level=3)

    add_body_paragraph(doc, 'A functioning intake conversion system includes:')

    for bold_lead, body_text in [
        ('A pre-consultation touchpoint',
         ' that sets the frame before the first conversation \u2014 what to expect, what you\u2019ll '
         'cover, and a subtle positioning statement that establishes your firm\u2019s approach '
         'to HNW matters.'),
        ('A structured consultation framework',
         ' your intake team follows consistently \u2014 not verbatim, but architecturally. How to open, '
         'how to surface the prospect\u2019s real concerns, how to communicate competence without '
         'pitching, and how to move toward a decision naturally.'),
        ('A post-consultation follow-up sequence',
         ' for prospects who don\u2019t sign on the spot. Most HNW prospects don\u2019t sign in the '
         'first meeting. If your follow-up process is a single email and then nothing, you\u2019re '
         'leaving cases on the table that were very nearly yours.'),
        ('KPIs that tell you where it\u2019s breaking.',
         ' Consultation-to-retention rate, follow-up contact rate, average time from consultation to '
         'signed retainer. Most firms have no idea what these numbers are. When you know them, you know '
         'exactly where to fix it.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(7)
        r_b = p.add_run(bold_lead)
        r_b.font.name  = 'Georgia'
        r_b.font.size  = Pt(11)
        r_b.font.bold  = True
        r_b.font.color.rgb = BODY_COLOR
        r_t = p.add_run(body_text)
        r_t.font.name  = 'Georgia'
        r_t.font.size  = Pt(11)
        r_t.font.color.rgb = BODY_COLOR

    # What Now
    doc.add_heading('Sit in on Three Calls This Week', level=3)

    add_callout_lines(doc, [
        'Sit in on three intake calls this week \u2014 or listen to recordings if you have them. '
        'Don\u2019t evaluate legal knowledge. Evaluate this:',
        '',
        '\u2014  Did the conversation make the prospect feel understood, or processed?',
        '\u2014  Was the firm\u2019s expertise communicated, or just assumed?',
        '\u2014  Did the prospect leave with a clear next step and a reason to choose your firm?',
        '',
        'If you\u2019re uncomfortable with what you hear, that\u2019s the intake '
        'conversion gap.',
    ])

    add_horizontal_rule(doc, color="58585A")

    # ═══════════════════════════════════════════════════════════════════════════
    # SYSTEM 3
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_heading('System 3: Pipeline Nurturing', level=2)

    # Why
    doc.add_heading('The Quietest Revenue Leak in Your Practice', level=3)

    add_body_paragraph(doc, 'Not every HNW prospect is ready this week.')

    add_body_paragraph(doc,
        'Some are gathering information. Some are in the early stages of a situation that will be a case '
        'in 60 or 90 days. Some received your name as a referral and are quietly evaluating you before '
        'they ever reach out. Some called, had a consultation, didn\u2019t sign \u2014 and are '
        'still deciding.')

    add_body_paragraph(doc, 'Without a nurture system, all of those people disappear.')

    add_body_paragraph(doc,
        'They don\u2019t disappear because they chose someone else. They disappear because there was no '
        'thread connecting them back to you. No reason to remember your name when the decision point '
        'arrived. No content they\u2019d seen from you that reinforced confidence in your expertise. '
        'Just silence.')

    add_body_paragraph(doc, 'And then they show up as a client at another firm.')

    add_body_paragraph(doc,
        'This is the quietest revenue leak in family law practices. You\u2019ll never see it in your '
        'metrics because you never knew those prospects existed. They reached out or were referred, went '
        'cold, and became invisible. The cases they represent don\u2019t show up as lost \u2014 '
        'they just never show up at all.')

    # What
    doc.add_heading('The Thread That Keeps You in the Room', level=3)

    add_body_paragraph(doc,
        'Pipeline nurturing is the infrastructure that keeps your firm present and credible with '
        'prospects who aren\u2019t ready yet \u2014 without adding anything to your workload.')

    add_body_paragraph(doc,
        'It\u2019s not a newsletter they\u2019ll ignore. It\u2019s a sequenced, relevant, automated '
        'follow-up system built around the actual decision timeline of a HNW family law prospect. '
        'Content that earns trust. Touchpoints that demonstrate expertise. A consistent signal that your '
        'firm is still there, still capable, and still the right choice \u2014 whenever they\u2019re ready.')

    # How
    doc.add_heading('Three Groups. One System', level=3)

    add_body_paragraph(doc, 'A functioning nurture system operates across three groups:')

    for bold_lead, body_text in [
        ('Active prospects who didn\u2019t convert.',
         ' Consultations that ended without a signed retainer get placed into a follow-up sequence \u2014 '
         'not aggressive, not salesy, but substantive. Something useful about the process they\u2019re '
         'facing. An article that addresses a concern they raised. A case study that reflects their '
         'situation. Content that moves the conversation forward without requiring them to respond.'),
        ('Cold leads.',
         ' Prospects who inquired but never scheduled, or who went quiet after initial contact. These '
         'aren\u2019t dead leads \u2014 they\u2019re pending leads. A structured reactivation sequence '
         'can bring a meaningful percentage of these back into conversation.'),
        ('Past clients and referral sources.',
         ' The highest-value pipeline any family law firm has is the one it already built. Past clients '
         'who had a positive experience are your most efficient source of referrals \u2014 if you stay '
         'in contact. Most firms don\u2019t. A simple, automated sequence keeps you front-of-mind with '
         'the people most likely to send you their next case.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        r_b = p.add_run(bold_lead)
        r_b.font.name  = 'Georgia'
        r_b.font.size  = Pt(11)
        r_b.font.bold  = True
        r_b.font.color.rgb = BODY_COLOR
        r_t = p.add_run(body_text)
        r_t.font.name  = 'Georgia'
        r_t.font.size  = Pt(11)
        r_t.font.color.rgb = BODY_COLOR

    add_body_paragraph(doc,
        'The system runs in the background. You set it up once, connect it to your CRM, and the '
        'infrastructure does the work. Prospects receive relevant, valuable communication from your firm '
        'at intervals that reflect how they actually make decisions \u2014 not how fast you want '
        'them to decide.')

    # What Now
    doc.add_heading('Three Questions That Tell You Where You Stand', level=3)

    add_callout_lines(doc, [
        'Answer these three questions:',
        '',
        '\u2014  What happens to a prospect who attends a consultation and doesn\u2019t sign? '
        'Is there a follow-up sequence, or does someone make a couple of calls and let it go?',
        '\u2014  What happened to your last 20 inquiries that didn\u2019t convert immediately? '
        'Where are they right now?',
        '\u2014  When did you last communicate with a past client who hasn\u2019t sent you '
        'a referral this year?',
        '',
        'If the honest answer is \u201cnot much\u201d or \u201cI don\u2019t know,\u201d your '
        'pipeline nurturing system is missing. And the revenue it\u2019s costing you is real \u2014 '
        'it\u2019s just invisible.',
    ])

    add_horizontal_rule(doc, color="58585A")

    # ═══════════════════════════════════════════════════════════════════════════
    # WHAT ALL THREE HAVE IN COMMON
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_heading('What All Three Systems Have in Common', level=2)

    add_body_paragraph(doc, 'None of these are marketing problems.')

    add_body_paragraph(doc,
        'Your agency can\u2019t fix them. More ad spend won\u2019t solve them. A new website won\u2019t '
        'address them. These are operations problems \u2014 systems that live between \u201cprospect '
        'reaches out\u201d and \u201cprospect signs\u201d \u2014 and they\u2019re the actual reason '
        'high-quality leads don\u2019t become high-value cases.')

    add_body_paragraph(doc,
        'The firms winning HNW family law cases aren\u2019t outspending you on ads. They\u2019re '
        'outoperating you on intake. They respond faster. They convert better. They stay in front of '
        'the people who aren\u2019t ready yet \u2014 until they are.')

    add_body_paragraph(doc,
        'That\u2019s the difference between a practice that generates revenue and a practice that '
        'generates leads.')

    add_horizontal_rule(doc, color="58585A")

    # ═══════════════════════════════════════════════════════════════════════════
    # WHERE TO START
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_heading('Where to Start', level=2)

    add_body_paragraph(doc,
        'Before you fix any of these systems, you need to know which one is costing you the most.')

    add_body_paragraph(doc,
        '**Speed to lead** is usually the fastest win \u2014 it\u2019s the one with the most immediate, '
        'measurable impact and the clearest implementation path. **Intake conversion** is typically the '
        'highest-leverage fix \u2014 it raises your consultation-to-retention rate, which makes every '
        'lead more valuable. **Pipeline nurturing** is the longest-term compound investment \u2014 the '
        'one that keeps producing returns long after the initial build.')

    add_body_paragraph(doc,
        'The right sequence depends on what your firm actually looks like right now. Some practices need '
        'speed-to-lead first. Others need to fix the consultation framework before anything else makes '
        'sense. Some are losing the most revenue in the nurture gap.')

    add_body_paragraph(doc, 'The starting point is an honest look at where the cases are going cold.')

    add_body_paragraph(doc,
        'If you want a structured framework for that audit \u2014 a way to look at your intake process '
        'and identify exactly which system is your biggest gap \u2014 that\u2019s the conversation '
        'we\u2019re built for.')

    add_body_paragraph(doc,
        'It\u2019s not a pitch call. It\u2019s a diagnostic. Thirty minutes, your actual numbers, '
        'a direct read on what to fix first.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # CTA line
    p_cta = doc.add_paragraph()
    r = p_cta.add_run(
        'When you\u2019re ready: legalgrowthdynamics.com/intake-audit'
    )
    r.font.name  = 'Calibri'
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    p_cta.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Author sign-off
    p_sig = doc.add_paragraph()
    r = p_sig.add_run(
        'Roger Bauer is the founder of Legal Growth Dynamics, a firm that helps family law attorneys '
        'build predictable HNW client pipelines. This report is a starting framework \u2014 not a '
        'complete prescription. Every firm\u2019s gap is in a different place. The audit call '
        'is where we find yours.'
    )
    r.font.name   = 'Georgia'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sig.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    add_report_body(doc)

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"\u2713 Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
