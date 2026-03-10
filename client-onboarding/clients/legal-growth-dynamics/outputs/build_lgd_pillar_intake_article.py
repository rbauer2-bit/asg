#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | Pillar Article: Intake & Conversion
"You Don't Have a Lead Problem. You Have a Leak Problem."
March 2026
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK     = RGBColor(0x1A, 0x1A, 0x1A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY   = "F5F5F5"

BASE_DIR    = os.path.dirname(__file__)
OUTPUT_DOCX = os.path.join(BASE_DIR, "lgd-family-law_pillar-article_intake-conversion_20260306.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color="25AAE1"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_inline_paragraph(doc, text, bold=False, italic=False, color=None,
                          font_size=11, alignment=None, space_before=None,
                          space_after=8):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # Parse **bold** and *italic* inline
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        italic_parts = re.split(r'\*(.*?)\*', part)
        for k, ipart in enumerate(italic_parts):
            if ipart:
                run = p.add_run(ipart)
                run.font.name = 'Georgia'
                run.font.size = Pt(font_size)
                run.font.color.rgb = color if color else BODY_COLOR
                run.font.bold   = bold or (j % 2 == 1)
                run.font.italic = italic or (k % 2 == 1)
    return p


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name  = 'Georgia'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.line_spacing = Pt(16)

    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(24)
    h1.font.bold  = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(28)
    h1.paragraph_format.space_after  = Pt(10)

    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(17)
    h2.font.bold  = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(13)
    h3.font.bold  = True
    h3.font.color.rgb = ASG_BLUE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(4)


# ── cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Top blue rule
    add_horizontal_rule(doc)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # Eyebrow
    p_eye = doc.add_paragraph()
    p_eye.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_eye.add_run('PILLAR CONTENT ARTICLE')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = ASG_CHARCOAL
    r.font.all_caps = True
    p_eye.paragraph_format.space_after = Pt(12)

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('You Don\u2019t Have a Lead Problem.\nYou Have a Leak Problem.')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(30)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    p_title.paragraph_format.space_after = Pt(16)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run(
        'Why family law attorneys keep losing money on marketing that should be working\n'
        '\u2014 and the one audit you need to run before you spend another dollar.'
    )
    r.font.name   = 'Calibri'
    r.font.size   = Pt(13)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sub.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    # Metadata block
    meta_lines = [
        ('Author: Roger Bauer', False),
        ('Legal Growth Dynamics \u2014 Family Law Attorney Market', False),
        ('March 2026  |  Confidential', False),
        ('', False),
        ('Produced by Authority Systems Group\u2122', True),
    ]
    for text, is_blue in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(11)
        r.font.color.rgb = ASG_BLUE if is_blue else ASG_CHARCOAL
        r.font.bold  = is_blue
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ── footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Blue top border on footer
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '25AAE1')
    pBdr.append(top)
    pPr.append(pBdr)

    run = fp.add_run(
        'Authority Systems Group\u2122 \u2014 Confidential. Prepared exclusively for Legal Growth Dynamics.'
    )
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── outline section ───────────────────────────────────────────────────────────

def add_outline_section(doc):
    doc.add_heading('ARTICLE OUTLINE', level=1)
    add_horizontal_rule(doc)

    outline_items = [
        ('I.', 'The Hook — "You\'ve Been Given the Wrong Diagnosis"',
         'Most attorneys who\'ve been burned by marketing were told the same thing: you need more leads. '
         'That diagnosis is almost always wrong. The real problem is not volume — it\'s conversion. '
         'More traffic into a broken system produces more waste, not more revenue.'),

        ('II.', 'The Numbers Most Attorneys Have Never Actually Looked At',
         '60% of law firms miss incoming calls entirely (Clio 2024 Legal Trends Report) \u2022 '
         'Only 33% of law firms respond to email inquiries \u2022 '
         '80% of prospects will contact another attorney within 48 hours of no response \u2022 '
         '79% of leads never make it to the conversion stage industry-wide'),

        ('III.', 'What "Intake Conversion" Actually Means',
         'Define the two rates every attorney must know: lead-to-consultation and consultation-to-hire. '
         'Industry benchmark context: high-performing firms close 50%+ of consultations; most family law '
         'firms operate at 30% or lower. The math of the gap at 20 consultations per month.'),

        ('IV.', 'Why More Marketing Makes a Broken System Worse',
         'Paying for leads you can\'t capture is not a marketing ROI problem \u2014 it\'s a systems failure. '
         'Speed reality: firms responding within 5 minutes see 400% higher conversion. '
         '97% of legal PPC users say ROI is poor \u2014 the reason is usually intake, not ads.'),

        ('V.', 'The Foundation Every Attorney Skips',
         'What a functioning intake system actually looks like: Speed. Process. Consistency. '
         'Clio data: firms with intake technology earn 50% more revenue and attract 50% more potential clients.'),

        ('VI.', 'The Math That Should Stop You Cold',
         'Running the conversion calculation at 30% close vs. 50% close on the same consultation volume. '
         'At 20 consultations/month and $8,000 average matter value, the difference is $32,000/month. '
         'No additional marketing spend required.'),

        ('VII.', 'The Right Sequence (Foundation First, Marketing Second)',
         'What to audit before you write another check to an agency. '
         'The four numbers every family law firm should know by the end of this week.'),

        ('VIII.', 'What to Do Right Now',
         'Run a 30-day intake audit: track response time, lead-to-consult rate, consultation-to-hire rate, '
         'drop-off points. Fix the leak before you amplify the flow. CTA: Assessment.'),
    ]

    for num, heading, desc in outline_items:
        # Section number + heading
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(12)
        p_head.paragraph_format.space_after  = Pt(3)
        r_num = p_head.add_run(num + '  ')
        r_num.font.name  = 'Calibri'
        r_num.font.size  = Pt(12)
        r_num.font.bold  = True
        r_num.font.color.rgb = ASG_BLUE
        r_text = p_head.add_run(heading)
        r_text.font.name  = 'Calibri'
        r_text.font.size  = Pt(12)
        r_text.font.bold  = True
        r_text.font.color.rgb = ASG_CHARCOAL

        # Description
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.35)
        p_desc.paragraph_format.space_after = Pt(6)
        r_d = p_desc.add_run(desc)
        r_d.font.name  = 'Georgia'
        r_d.font.size  = Pt(10)
        r_d.font.color.rgb = BODY_COLOR

    doc.add_page_break()


# ── article body ──────────────────────────────────────────────────────────────

def add_stat_callout(doc, text):
    """Blue left-border callout box for key statistics."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    set_paragraph_shading(p, 'EAF6FC')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    # Left blue border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '25AAE1')
    pBdr.append(left)
    pPr.append(pBdr)


def add_body_paragraph(doc, text, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    # Handle **bold** and *italic*
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        italic_parts = re.split(r'\*(.*?)\*', part)
        for k, ipart in enumerate(italic_parts):
            if ipart:
                r = p.add_run(ipart)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
                if k % 2 == 1:
                    r.font.italic = True
    return p


def add_article_section(doc):
    # Article eyebrow
    p_eye = doc.add_paragraph()
    r = p_eye.add_run('FULL ARTICLE')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = ASG_CHARCOAL
    r.font.all_caps = True
    p_eye.paragraph_format.space_after = Pt(4)

    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Article H1
    doc.add_heading('You Don\u2019t Have a Lead Problem. You Have a Leak Problem.', level=1)

    # Subtitle / standfirst
    p_sub = doc.add_paragraph()
    r = p_sub.add_run(
        'Why family law attorneys keep losing money on marketing that should be working \u2014 '
        'and the one audit you need to run before you spend another dollar.'
    )
    r.font.name   = 'Georgia'
    r.font.size   = Pt(12)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_sub.paragraph_format.space_after = Pt(20)

    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Opening ──
    add_body_paragraph(doc,
        'Most family law attorneys I speak with have been burned by marketing at least once. '
        'An agency promised more cases. They paid for a new website, or ran Google Ads for a quarter, '
        'or hired someone to handle their SEO. The results were disappointing. The clients didn\u2019t '
        'materialize. The money walked out the door.')

    add_body_paragraph(doc, 'They drew the obvious conclusion: marketing doesn\u2019t work for family law.')

    add_body_paragraph(doc, 'That conclusion is wrong.')

    add_body_paragraph(doc,
        'But here\u2019s what\u2019s worse \u2014 the advice they got next was also wrong. The prescription '
        'was always the same: *you need more leads.* More ad spend. More traffic. More content. More visibility.')

    add_body_paragraph(doc,
        'What they needed was a leak repaired. And nobody told them.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section 1 ──
    doc.add_heading('The Numbers Most Family Law Attorneys Have Never Looked At', level=2)

    add_body_paragraph(doc,
        'In 2024, Clio conducted a secret shopper study of 500 law firms across the United States. '
        'What they found should make every family law attorney stop cold.')

    add_stat_callout(doc,
        '60% of law firms did not answer their phone when a prospective client called. '
        'That number has gotten worse over time \u2014 in 2019, the answer rate was 56%. '
        'By 2024, it had dropped to 40%. Of the firms that missed the call, only 20% bothered to return it.')

    add_body_paragraph(doc,
        'Email wasn\u2019t better. **Only 33% of firms responded to an emailed inquiry at all.** '
        'Of those that did, fewer than one in five provided any clear next steps or information about fees.')

    add_stat_callout(doc,
        '80% of prospective legal clients will contact another attorney '
        'if they don\u2019t receive a response within 48 hours.')

    add_body_paragraph(doc, 'Not 48 days. 48 hours.')

    add_body_paragraph(doc,
        'Let that sit for a moment.')

    add_body_paragraph(doc,
        'You can have a well-designed website, a competitive Google Ads campaign, and a strong referral '
        'network \u2014 and still lose more than half your potential clients before they ever speak to anyone '
        'at your firm. Not because your marketing failed. Because your intake system did.')

    add_body_paragraph(doc, 'This is the leak.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section 2 ──
    doc.add_heading('What \u201cConversion\u201d Actually Means for Your Practice', level=2)

    add_body_paragraph(doc,
        'Every family law attorney needs to know two numbers. Most know neither.')

    add_body_paragraph(doc,
        'The first is your **lead-to-consultation rate** \u2014 of all the people who express interest in '
        'your firm (calls, web forms, referrals, emails), what percentage actually schedule and show up '
        'for a consultation?')

    add_body_paragraph(doc,
        'The second is your **consultation-to-hire rate** \u2014 of the prospects who sit across from you, '
        'what percentage retain you?')

    add_body_paragraph(doc,
        'Industry data on consultation close rates shows a significant spread. High-performing firms close '
        '**50% or more** of their consultations. Many family law firms are operating at **30% or below** '
        '\u2014 some considerably lower.')

    add_body_paragraph(doc, 'Here is what that difference means at practical scale.')

    add_stat_callout(doc,
        'At 20 consultations per month with a 30% close rate: 6 retained clients.\n'
        'At a 50% close rate \u2014 same 20 consultations, same marketing spend: 10 retained clients.\n'
        'At an average matter value of $8,000, that is $32,000 in additional monthly revenue. '
        'No additional advertising required.')

    add_body_paragraph(doc, 'That is not a marginal improvement. That is a structural one.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section 3 ──
    doc.add_heading('Why More Marketing Makes a Broken System Worse', level=2)

    add_body_paragraph(doc,
        'When an attorney with a leaking intake system hires an agency to drive more traffic, something '
        'predictable happens. More leads arrive. More leads fall through the cracks. The agency points to '
        'impressions and clicks and call volume. The attorney sees the same disappointing conversion to '
        'retained clients. Everyone blames the leads.')

    add_body_paragraph(doc, 'The leads were not the problem.')

    add_body_paragraph(doc,
        'The research on response speed alone is enough to make the case. Firms that respond to an incoming '
        'inquiry **within five minutes** see conversion rates **400% higher** than firms that respond within '
        '30 minutes or more. Four hundred percent. Not from better copywriting. Not from a higher ad budget. '
        'From speed.')

    add_body_paragraph(doc,
        'The Clio data also shows that firms with functioning intake systems \u2014 online scheduling, '
        'digital intake forms, structured follow-up \u2014 earn **50% more revenue on average** and '
        'attract **50% more potential clients** than firms operating on ad-hoc processes.')

    add_body_paragraph(doc,
        'That is the leverage sitting on the table for almost every family law firm I have encountered. '
        'Not more traffic. Better capture.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section 4 ──
    doc.add_heading('The Foundation Every Firm Skips', level=2)

    add_body_paragraph(doc,
        'Most attorneys, when they decide to grow their practice, jump straight to marketing. They think '
        'about their brand, their website, their SEO rankings, their Google Ads account. These are all '
        'downstream of a problem they haven\u2019t addressed.')

    add_body_paragraph(doc,
        'The foundation of a growing family law firm is a predictable intake system. Not sophisticated. '
        'Not expensive. Predictable.')

    add_body_paragraph(doc, 'That means three things.')

    # Three-item list with bold lead-ins
    for item in [
        ('Speed.', 'Someone needs to respond to every inquiry within a time window your firm has actually '
                   'defined and committed to. Not \u201cwhen someone gets to it.\u201d A defined window.'),
        ('Process.', 'Every prospect should move through a consistent sequence \u2014 initial contact, '
                     'information gathering, consultation scheduling, follow-up. Inconsistency is where '
                     'revenue leaks.'),
        ('Measurement.', 'You need to know your numbers. If you cannot tell me your lead-to-consultation '
                         'rate or your consultation-to-hire rate today, you are running your practice on '
                         'instinct. Instinct is not a growth strategy.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_after  = Pt(8)
        r_b = p.add_run(item[0] + '  ')
        r_b.font.name  = 'Georgia'
        r_b.font.size  = Pt(11)
        r_b.font.bold  = True
        r_b.font.color.rgb = BODY_COLOR
        r_t = p.add_run(item[1])
        r_t.font.name  = 'Georgia'
        r_t.font.size  = Pt(11)
        r_t.font.color.rgb = BODY_COLOR

    add_body_paragraph(doc,
        'When these three elements are in place, marketing amplifies revenue. When they are not in place, '
        'marketing amplifies waste.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section 5 ──
    doc.add_heading('The Right Sequence', level=2)

    add_body_paragraph(doc,
        'I talk to family law attorneys every week who are planning to increase their marketing budget '
        'this year. Most of them should not be doing that yet.')

    add_body_paragraph(doc,
        'Before you spend money on paid search, content, or any visibility campaign, you need to answer '
        'four questions honestly:')

    for q in [
        'What is your current lead-to-consultation rate?',
        'What is your current consultation-to-hire rate?',
        'What is your average response time from first point of contact to first conversation with a prospect?',
        'What happens to a prospect who doesn\u2019t book on the first contact?',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(5)
        r_num = p.add_run('\u2014  ')
        r_num.font.name = 'Georgia'
        r_num.font.size = Pt(11)
        r_num.font.color.rgb = ASG_BLUE
        r_num.font.bold = True
        r_q = p.add_run(q)
        r_q.font.name  = 'Georgia'
        r_q.font.size  = Pt(11)
        r_q.font.color.rgb = BODY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_body_paragraph(doc,
        'If you don\u2019t know the answers \u2014 or if the answers reveal a conversion problem \u2014 '
        'fix the leak first.')

    add_body_paragraph(doc,
        'Run a 30-day audit. Track every inbound inquiry. Log the source, the response time, whether they '
        'scheduled, whether they showed, whether they hired. This is not complicated. It requires '
        'discipline, not technology.')

    add_body_paragraph(doc,
        'What you will find in 30 days will be more valuable than any agency presentation you have '
        'ever sat through.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Closing ──
    doc.add_heading('Stop Pouring Water Into a Bucket With a Hole in It', level=2)

    add_body_paragraph(doc,
        'The family law attorneys who grow predictably are not the ones with the biggest marketing budgets. '
        'They are the ones who know their numbers, close at a high rate, and respond fast. When they invest '
        'in marketing, they invest with confidence \u2014 because they know their system captures '
        'what they attract.')

    add_body_paragraph(doc,
        'If you are frustrated with marketing results, I want you to ask a different question before you '
        'call another agency. Don\u2019t ask *how do I get more leads?* Ask: *what happens to the leads '
        'I already have?*')

    add_body_paragraph(doc, 'The answer to that question is where the real growth is hiding.')

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Author bio / CTA
    p_bio = doc.add_paragraph()
    r = p_bio.add_run(
        'Roger Bauer is the founder of Legal Growth Dynamics, a firm that helps family law attorneys build '
        'predictable client pipelines without referral dependency. If you want to identify exactly where '
        'your firm is losing revenue in the intake and conversion process, start with our free five-minute '
        'practice assessment at legalgrowthdynamics.com.'
    )
    r.font.name   = 'Georgia'
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ASG_CHARCOAL
    p_bio.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    add_outline_section(doc)
    add_article_section(doc)

    doc.save(OUTPUT_DOCX)
    print(f"\u2713 Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
