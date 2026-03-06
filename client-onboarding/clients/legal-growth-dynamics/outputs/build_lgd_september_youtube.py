#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics | September 2026 YouTube Long-Form Script Pack
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../..'))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ASG_BLUE    = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK    = RGBColor(0x1A, 0x1A, 0x1A)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR  = RGBColor(0x33, 0x33, 0x33)

BASE_DIR   = os.path.dirname(__file__)
INPUT_MD   = os.path.join(BASE_DIR, "lgd-family-law-attorneys_september-2026-youtube-scripts_20260305.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "lgd-family-law-attorneys_september-2026-youtube-scripts_20260305.docx")


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '25AAE1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name  = 'Georgia'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.line_spacing = Pt(15)

    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(22)
    h1.font.bold  = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)

    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(16)
    h2.font.bold  = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after  = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(13)
    h3.font.bold  = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(4)


def add_cover_page(doc):
    # Main title
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('20-DAY YOUTUBE LONG-FORM SCRIPT PACK')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(32)
    run.font.bold  = True
    run.font.color.rgb = ASG_BLUE
    cover_title.paragraph_format.space_before = Pt(72)
    cover_title.paragraph_format.space_after  = Pt(8)

    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = firm.add_run('Legal Growth Dynamics')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(26)
    r.font.bold  = True
    r.font.color.rgb = ASG_CHARCOAL
    firm.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('Family Law Attorneys | September 2026')
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(18)
    r2.font.color.rgb = ASG_CHARCOAL
    sub.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    for line in [
        '20 Complete Long-Form Scripts  |  5-Week Belief-to-Buy Arc',
        'Full Production Notes + Value Bomb + Contrarian Variants',
        '',
        'Prepared exclusively for Legal Growth Dynamics',
        'March 2026  |  Confidential',
        '',
        'Under the direction of Roger Bauer, Director',
        'Authority Systems Group™',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        if 'Authority Systems Group' in line or 'Roger Bauer' in line:
            r.font.color.rgb = ASG_BLUE
            r.font.bold = True
        elif 'Confidential' in line or 'March' in line:
            r.font.color.rgb = ASG_CHARCOAL
            r.font.bold = True
        else:
            r.font.color.rgb = ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        'Authority Systems Group™ — Confidential. Prepared exclusively for Legal Growth Dynamics.'
    )
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def parse_and_render(doc, md_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block = False
    cover_done = False
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip YAML frontmatter (--- to ---)
        if not cover_done:
            if line.strip() == '---' and i > 0:
                cover_done = True
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name  = 'Courier New'
            r.font.size  = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-\|:]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.paragraph_format.space_after = Pt(2)
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = WHITE
                            if ri == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '25AAE1')
                                tcPr.append(shd)
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith('### ') and not line.startswith('#### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            r.font.name  = 'Calibri'
            r.font.size  = Pt(12)
            r.font.bold  = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip().replace('*', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name   = 'Calibri'
            r.font.size   = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.left_indent = Inches(0.3)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'),   'single')
            left.set(qn('w:sz'),    '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', line[2:].strip())
            p.clear()
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name  = 'Georgia'
                r.font.size  = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Checkbox items
        if line.strip().startswith('- [ ]'):
            text = line.strip()[5:].strip()
            p = doc.add_paragraph()
            r = p.add_run('\u2610  ' + text)
            r.font.name  = 'Georgia'
            r.font.size  = Pt(11)
            r.font.color.rgb = BODY_COLOR
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph — inline bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            italic_parts = re.split(r'\*(.*?)\*', part)
            for k, ipart in enumerate(italic_parts):
                if ipart:
                    r = p.add_run(ipart)
                    r.font.name  = 'Georgia'
                    r.font.size  = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
                    if k % 2 == 1:
                        r.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        i += 1


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin   = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    parse_and_render(doc, INPUT_MD)

    doc.save(OUTPUT_DOCX)
    print(f"Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
