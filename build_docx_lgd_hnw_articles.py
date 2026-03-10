#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics: 30-Day HNW Blog Article Pack
Family Law Attorneys | High-Net-Worth Case Attraction
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Brand Colors ────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)   # #25aae1
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)   # #58585a
ASG_DARK     = RGBColor(0x1A, 0x1A, 0x1A)   # #1a1a1a
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)

INPUT_FILES = [
    "outputs/lgd_blog_articles_hnw_familylaw_20260307_01-10.md",
    "outputs/lgd_blog_articles_hnw_familylaw_20260307_11-20.md",
    "outputs/lgd_blog_articles_hnw_familylaw_20260307_21-30.md",
]
OUTPUT_DOCX = "outputs/lgd_hnw_blog_articles_30day_20260307.docx"


# ─── XML Helpers ─────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color_hex='25AAE1', weight='8'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), weight)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_blue_left_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '25AAE1')
    pBdr.append(left)
    pPr.append(pBdr)


# ─── Styles ──────────────────────────────────────────────

def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = Pt(15)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(5)

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ASG_BLUE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


# ─── Cover Page ──────────────────────────────────────────

def add_cover_page(doc):
    # Top brand bar (blue shaded paragraph)
    bar = doc.add_paragraph()
    bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bar.add_run('LEGAL GROWTH DYNAMICS™')
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = WHITE
    set_paragraph_shading(bar, '25AAE1')
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(0)

    sub_bar = doc.add_paragraph()
    sub_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_bar.add_run('A Division of Authority Systems Group™   |   AuthoritySystemsGroup.com')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = WHITE
    set_paragraph_shading(sub_bar, '1D88B4')
    sub_bar.paragraph_format.space_before = Pt(0)
    sub_bar.paragraph_format.space_after = Pt(48)

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run('30-Day HNW Blog Article Pack')
    rt.font.name = 'Calibri'
    rt.font.size = Pt(34)
    rt.font.bold = True
    rt.font.color.rgb = ASG_BLUE
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run('Targeting Family Law Attorneys')
    rs.font.name = 'Calibri'
    rs.font.size = Pt(20)
    rs.font.bold = False
    rs.font.color.rgb = ASG_CHARCOAL
    subtitle.paragraph_format.space_after = Pt(4)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = subtitle2.add_run('Attracting & Converting High-Net-Worth Cases')
    rs2.font.name = 'Calibri'
    rs2.font.size = Pt(16)
    rs2.font.color.rgb = ASG_CHARCOAL
    subtitle2.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    # Description block
    for line, size, bold, color in [
        ('30 fully written blog articles',          12, True,  ASG_CHARCOAL),
        ('5 hurdles × 6 articles each',             11, False, ASG_CHARCOAL),
        ('Belief-to-Buy™ Framework integrated throughout', 11, False, ASG_CHARCOAL),
        ('Authority Conversion Protocol™ mapped',   11, False, ASG_CHARCOAL),
        ('',                                        10, False, ASG_CHARCOAL),
        ('Hurdle 1 — Invisible Positioning',        11, True,  ASG_BLUE),
        ('Hurdle 2 — Referral Network Gap',         11, True,  ASG_BLUE),
        ('Hurdle 3 — Pricing & Retainer Confidence',11, True,  ASG_BLUE),
        ('Hurdle 4 — Content Authority Gap',        11, True,  ASG_BLUE),
        ('Hurdle 5 — Pipeline Nurturing & Follow-Up',11, True, ASG_BLUE),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(3)

    add_horizontal_rule(doc)

    # Footer block
    for line, bold in [
        ('Produced: March 2026  |  Confidential', True),
        ('Under the direction of Roger, Director', False),
        ('Authority Systems Group™', True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.color.rgb = ASG_BLUE if 'Authority Systems' in line else ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()


# ─── Table of Contents ───────────────────────────────────

def add_toc(doc, all_lines):
    """Build a table of contents from ## Article headings."""
    p = doc.add_heading('Table of Contents', level=1)

    article_titles = []
    for line in all_lines:
        m = re.match(r'^## (Article \d+: .+)$', line.strip())
        if m:
            article_titles.append(m.group(1))

    for title in article_titles:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()


# ─── Footer ──────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        'Legal Growth Dynamics™  |  Authority Systems Group™  |  Confidential'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ─── Metadata Box ────────────────────────────────────────

def add_metadata_box(doc, meta_lines):
    """Render article metadata (Hurdle, ACP Stage, Angle, Hook) as a shaded callout."""
    for line in meta_lines:
        # Strip bold markers to get clean key: value
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line).strip()
        if not clean:
            continue
        p = doc.add_paragraph()
        set_paragraph_shading(p, 'EAF6FC')
        # Split key: value
        if ':' in clean:
            key, _, val = clean.partition(':')
            rk = p.add_run(key.strip() + ':  ')
            rk.font.name = 'Calibri'
            rk.font.size = Pt(9)
            rk.font.bold = True
            rk.font.color.rgb = ASG_BLUE
            rv = p.add_run(val.strip())
            rv.font.name = 'Calibri'
            rv.font.size = Pt(9)
            rv.font.color.rgb = ASG_CHARCOAL
        else:
            r = p.add_run(clean)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)

    # Small spacer after metadata block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.space_before = Pt(0)


# ─── Inline Bold/Italic Renderer ────────────────────────

def add_rich_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    # Split by bold markers first
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        italic_parts = re.split(r'\*(.*?)\*', part)
        for k, ipart in enumerate(italic_parts):
            if ipart:
                r = p.add_run(ipart)
                r.font.name = 'Georgia'
                r.font.size = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
                if k % 2 == 1:
                    r.font.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


# ─── Main Content Renderer ───────────────────────────────

def load_all_lines(files):
    """Read all files and concatenate, stripping per-file headers."""
    all_lines = []
    for filepath in files:
        with open(filepath, 'r') as f:
            raw = f.readlines()
        # Skip the file-level header block (comment lines + first --- separator)
        past_header = False
        for line in raw:
            if not past_header:
                stripped = line.strip()
                # The file header ends at the first blank line after the first ---
                if stripped == '---' and not past_header:
                    past_header = True
                continue
            all_lines.append(line.rstrip('\n'))
    return all_lines


def is_metadata_line(line):
    """Return True if the line is an article metadata annotation."""
    stripped = line.strip()
    return (
        stripped.startswith('**Hurdle:**') or
        stripped.startswith('**ACP Stage Pair:**') or
        stripped.startswith('**Angle:**') or
        stripped.startswith('**Hook Type:**')
    )


def render_content(doc, all_lines):
    """Parse and render all article content."""
    first_article = True
    i = 0
    n = len(all_lines)

    while i < n:
        line = all_lines[i]

        # ── Skip blank lines at file boundaries
        if not line.strip() and i < 3:
            i += 1
            continue

        # ── Article heading (## Article N: ...)
        if re.match(r'^## Article \d+:', line.strip()):
            if not first_article:
                doc.add_page_break()
            first_article = False

            title = line.strip()[3:].strip()  # strip '## '
            p = doc.add_heading(title, level=1)

            # Collect following metadata lines
            i += 1
            meta_lines = []
            # Skip blank line(s) immediately after heading
            while i < n and not all_lines[i].strip():
                i += 1
            # Collect metadata lines
            while i < n and is_metadata_line(all_lines[i]):
                meta_lines.append(all_lines[i].strip())
                i += 1
            if meta_lines:
                add_metadata_box(doc, meta_lines)
            continue

        # ── Section headings (### ...)
        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # ── H4 level (#### ...)
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Horizontal rules (---) — light divider within article
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc, color_hex='CCCCCC', weight='4')
            i += 1
            continue

        # ── Blockquotes
        if line.startswith('> '):
            text = line[2:].strip().replace('*', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'EAF6FC')
            p.paragraph_format.left_indent = Inches(0.3)
            add_blue_left_border(p)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                italic_parts = re.split(r'\*(.*?)\*', part)
                for k, ipart in enumerate(italic_parts):
                    if ipart:
                        r = p.add_run(ipart)
                        r.font.name = 'Georgia'
                        r.font.size = Pt(11)
                        r.font.color.rgb = BODY_COLOR
                        if j % 2 == 1:
                            r.font.bold = True
                        if k % 2 == 1:
                            r.font.italic = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if part:
                    r = p.add_run(part)
                    r.font.name = 'Georgia'
                    r.font.size = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Disclaimer lines (*This content is...)
        if line.strip().startswith('*This content is') or line.strip().startswith('*_This content'):
            text = line.strip().strip('*').strip('_').strip()
            p = doc.add_paragraph()
            set_paragraph_shading(p, 'F5F5F5')
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = ASG_CHARCOAL
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.1)
            i += 1
            continue

        # ── Empty lines
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # ── Regular paragraph (with inline bold/italic)
        add_rich_paragraph(doc, line)
        i += 1


# ─── Main ────────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)

    all_lines = load_all_lines(INPUT_FILES)
    add_toc(doc, all_lines)
    render_content(doc, all_lines)

    doc.save(OUTPUT_DOCX)
    print(f"✓ Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
