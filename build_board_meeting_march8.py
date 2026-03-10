#!/usr/bin/env python3
"""
Authority Systems Group™
Board Meeting — March 8, 2026
Weekly Priorities, Site Build Strategy, Thomas Banks Analysis
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT     = RGBColor(0x1A, 0x25, 0x33)

ASG_BLUE_HEX  = '25AAE1'
ASG_DARK_HEX  = '1A2533'
CHARCOAL_HEX  = '58585A'
LIGHT_BG      = 'F5F5F5'
TINT_BG       = 'EAF6FC'
DIVIDER_HEX   = 'DDDDDD'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'E67E22'
RED_HEX       = 'C0392B'
SLATE_HEX     = '2C3E50'
PURPLE_HEX    = '6C3483'

BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_board-meeting_weekly-priorities_20260308.docx')


# ── XML / Layout Helpers ───────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_pb(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)
    return p


def section_hdr(doc, text, bg=ASG_DARK_HEX, color=WHITE, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    set_para_bg(p, bg)
    r = p.add_run(f'  {text}  ')
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = color; r.font.name = 'Calibri'
    return p


def sub_hdr(doc, text, color=ASG_BLUE, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = color; r.font.name = 'Calibri'
    return p


def body(doc, text, size=10.5, color=BODY_COLOR, bold=False, italic=False, indent=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Calibri'; r.bold = bold; r.italic = italic
    return p


def blt(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = BODY_COLOR; r.font.name = 'Calibri'
    return p


def divider(doc, color=DIVIDER_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '4')
    bt.set(qn('w:space'), '1'); bt.set(qn('w:color'), color)
    pBdr.append(bt); pPr.append(pBdr)
    return p


def add_page_number(para):
    run = para.add_run()
    run.font.size = Pt(8); run.font.color.rgb = FOOTER_COLOR; run.font.name = 'Calibri'
    for tag, txt in [('begin', None), (None, ' PAGE '), ('end', None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = txt; run._r.append(it)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    ht = header.add_table(1, 2, Inches(6.5)); remove_table_borders(ht)
    ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    lp = ht.cell(0, 0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    rp = ht.cell(0, 1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('BOARD MEETING  |  WEEKLY PRIORITIES  |  MARCH 8, 2026')
    rr.font.size = Pt(8); rr.font.color.rgb = FOOTER_COLOR; rr.font.name = 'Calibri'; rr.bold = True

    footer = section.footer; footer.is_linked_to_previous = False
    ft = footer.add_table(1, 2, Inches(6.5)); remove_table_borders(ft)
    ft.columns[0].width = Inches(4.5); ft.columns[1].width = Inches(2.0)
    flp = ft.cell(0, 0).paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run('Authority Systems Group™  —  Internal Document  —  Confidential')
    flr.font.size = Pt(8); flr.font.color.rgb = FOOTER_COLOR; flr.font.name = 'Calibri'
    frp = ft.cell(0, 1).paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(frp)


# ── Cover ──────────────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin', 1.0), ('right_margin', 1.0), ('top_margin', 0.75), ('bottom_margin', 0.75)]:
        setattr(section, attr, Inches(val))

    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(40); lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(2.4))

    sp(doc, 20)

    ey = doc.add_paragraph(); ey.paragraph_format.space_before = Pt(0); ey.paragraph_format.space_after = Pt(6)
    er = ey.add_run('INTERNAL BOARD MEETING — MARCH 8, 2026')
    er.font.size = Pt(9); er.font.color.rgb = ASG_BLUE; er.bold = True; er.font.name = 'Calibri'

    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(0); tp.paragraph_format.space_after = Pt(8)
    tr = tp.add_run('Weekly Priority Session')
    tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = DARK_TEXT; tr.font.name = 'Calibri'

    stp = doc.add_paragraph(); stp.paragraph_format.space_before = Pt(0); stp.paragraph_format.space_after = Pt(4)
    str_ = stp.add_run('10 Priorities · Site Build Strategy · Thomas Banks Opportunity · Honest Verdicts')
    str_.font.size = Pt(13); str_.font.color.rgb = ASG_CHARCOAL; str_.font.name = 'Calibri'

    divider(doc, ASG_BLUE_HEX)

    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8); mp.paragraph_format.space_after = Pt(4)
    mr = mp.add_run('Director: Roger Bauer     |     Board: Full Team     |     Classification: Confidential')
    mr.font.size = Pt(9.5); mr.font.color.rgb = ASG_CHARCOAL; mr.font.name = 'Calibri'

    sp(doc, 20)

    fp = doc.add_paragraph(); set_para_bg(fp, TINT_BG)
    fp.paragraph_format.space_before = Pt(0); fp.paragraph_format.space_after = Pt(4)
    fp.paragraph_format.left_indent = Inches(0.15); fp.paragraph_format.right_indent = Inches(0.15)
    fr = fp.add_run(
        '  Ground rules for this session: No sugar-coating. No validating bad sequencing. '
        'No pretending a five-site build week is compatible with closing new clients. '
        'The team\'s job is to tell Roger what is true — and what will move the needle fastest.  '
    )
    fr.font.size = Pt(10); fr.font.color.rgb = BODY_COLOR; fr.italic = True; fr.font.name = 'Calibri'

    add_pb(doc)


# ── Section 1: The 10 Priorities ──────────────────────────────────────────────

def build_priorities(doc):
    section_hdr(doc, '01 — THE 10 PRIORITIES THIS WEEK  |  Preston Adler, COO')
    sp(doc, 8)

    body(doc,
        'Roger came in with five items. The board reviewed them and added five more. Below is the '
        'full list, sequenced by strategic weight — not by the order Roger listed them. '
        'The sequence matters. Read the rationale.',
        after=12
    )

    priorities = [
        {
            'num': '01',
            'title': 'Activate Warm Outreach — First Paying Client (DMA or LGD)',
            'origin': 'Board Addition',
            'tier_hex': RED_HEX,
            'tier': 'CRITICAL',
            'owner': 'Roger',
            'time': '2–4 hrs total this week',
            'rationale': (
                'This is the most important item on the list. Not a website. Not a blog post. '
                'A conversation with a real human who might pay money. Roger has a network. '
                'Ten outreach messages to implant dentists and ten to family law attorneys — this week. '
                'Not after the sites are built. Not after the DOCX is printed. Now. '
                'Every website built before the first client is booked is infrastructure for a business '
                'that does not yet have revenue. Sequence this first.'
            ),
            'action': (
                'Identify 10 implant dentists + 10 family law attorneys in Roger\'s existing network. '
                'Send a direct, low-friction message: "I\'m running a limited number of free digital '
                'authority assessments this month for [niche] practices. Want one?" — no pitch, no funnel, '
                'just a door opener.'
            ),
        },
        {
            'num': '02',
            'title': 'Thomas Banks — Build Personal Authority Site (KentuckianaDivorceAttorney.com)',
            'origin': 'Roger (Item 3)',
            'tier_hex': RED_HEX,
            'tier': 'CRITICAL',
            'owner': 'Roger + Team',
            'time': '4–6 hrs strategy + delegation',
            'rationale': (
                'The board reviewed both sites. KentuckianaDivorceAttorney.com is a generic placeholder — '
                'no authority claims, no trust signals, no differentiation. The firm site (divorceky-in.com) '
                'is professionally redesigned and positions well, but the individual attorney pages are weak. '
                'This is the play: make Thomas Banks\' personal site dramatically better than his colleagues\' '
                'pages on the firm site. When the other attorneys see what Roger did for Thomas, they will '
                'come to Roger. One relationship, multiple clients, multiple case studies, referral flywheel. '
                'This is the highest-leverage item in the entire list.'
            ),
            'action': (
                'Step 1 (this week): Run a full Digital Authority Assessment on KentuckianaDivorceAttorney.com '
                'and divorceky-in.com. Produce a positioning brief for Thomas. Step 2: Build his personal '
                'authority site on the existing domain with copy that outranks the generic attorney bio on '
                'the firm site. Step 3: After delivery, ask Thomas who else at the firm should hear about this.'
            ),
        },
        {
            'num': '03',
            'title': 'LGD Services Pages — Push Last Night\'s Rewrites Live',
            'origin': 'Roger (Item 5)',
            'tier_hex': GREEN_HEX,
            'tier': 'DO TODAY',
            'owner': 'Roger',
            'time': '1–2 hrs',
            'rationale': (
                'The content is written. The copy is done. Every day that improved copy sits in a document '
                'instead of on the live site is a day a prospect lands on the old version and bounces. '
                'This is the fastest item on the list. Do it first thing Monday.'
            ),
            'action': 'Log into LGD WordPress. Paste revised copy into services pages. Publish. Done.',
        },
        {
            'num': '04',
            'title': 'Domain + Divi 5 Theme Setup — All Pending Sites',
            'origin': 'Roger (Board Reframe)',
            'tier_hex': AMBER_HEX,
            'tier': 'HIGH',
            'owner': 'Roger (setup only — delegate build)',
            'time': '2–3 hrs',
            'rationale': (
                'Roger\'s instinct to install Divi 5 and hand off the content build is exactly right. '
                'The strategic framing: Roger\'s job is domain registration, hosting setup, theme install, '
                'and basic page structure creation. The content-to-page build is a $25–40/hr task. '
                'Do not spend 20 hours in Divi doing what a part-time VA can do in 8. '
                'This week\'s goal: get all five pending domains pointed to hosting, Divi 5 installed, '
                'blank page structure created. Then hand the brief and content files to a builder.'
            ),
            'action': (
                'Register/verify domains. Assign hosting. Install Divi 5. Create page structure (Home, '
                'Services, About, Blog, Contact) for each site. Export content briefs and hand off. '
                'Sites covered: TopLouisvilleDentist.com, HedgesDental.com, KentuckianaDivorceAttorney.com, '
                'ASG site, DMA site.'
            ),
        },
        {
            'num': '05',
            'title': 'Jeff Hedges — Implant-Focused Site (TopLouisvilleDentist.com)',
            'origin': 'Roger (Item 1)',
            'tier_hex': AMBER_HEX,
            'tier': 'HIGH',
            'owner': 'Roger (brief) → Builder (execution)',
            'time': '1–2 hrs brief + handoff',
            'rationale': (
                'This is a real client site that represents DMA\'s first case study candidate. '
                'Jeff Hedges is the proof-of-concept for the implant dentist model. '
                'The content assets (persona, April blog articles) are in place. The strategy is clear: '
                'position implants as the practice\'s flagship service, use the Belief-to-Buy architecture '
                'to move a hesitant patient from "implants are too expensive" to "this is the only right choice." '
                'Roger should write the brief and hand the build to a developer. '
                'Do not build this manually — it is too time-consuming relative to the leverage it creates.'
            ),
            'action': (
                'Write a one-page site brief covering: target avatar, primary service (implants), key '
                'messaging hooks, page structure, CTA strategy. Attach relevant DMA content assets. '
                'Hand off to builder with Divi 5 skeleton already in place.'
            ),
        },
        {
            'num': '06',
            'title': 'Jeff Hedges — General Dentistry Site (HedgesDental.com)',
            'origin': 'Roger (Item 2)',
            'tier_hex': AMBER_HEX,
            'tier': 'HIGH',
            'owner': 'Roger (brief) → Builder (execution)',
            'time': '1 hr brief + handoff',
            'rationale': (
                'This is the secondary Hedges site — general dentistry with implants featured prominently. '
                'The positioning challenge: do not make this look like "a dentist who also does implants." '
                'Make it look like "an implant specialist who also serves your family\'s general dental needs." '
                'The distinction matters for the avatar: a patient researching implants should land on '
                'TopLouisvilleDentist.com (pure implant authority). A patient researching a family dentist '
                'who offers implants should land on HedgesDental.com.'
            ),
            'action': 'Brief + handoff same as Item 05. Keep content distinct between both Hedges properties.',
        },
        {
            'num': '07',
            'title': 'Hire a WordPress / Divi Builder This Week',
            'origin': 'Board Addition',
            'tier_hex': AMBER_HEX,
            'tier': 'HIGH',
            'owner': 'Roger',
            'time': '1–2 hrs sourcing + vetting',
            'rationale': (
                'This is not a "should I someday hire someone" question. It is a "do I want to spend '
                '25–40 hours this week in WordPress, or do I want to spend that time closing clients?" '
                'A competent Divi freelancer costs $25–45/hr on Upwork or similar platforms. '
                'At five sites, conservative estimate is 8–12 hours of build work per site = 40–60 hours total. '
                'That is $1,000–$2,700 in labor. One Authority Blueprint engagement pays for all of it '
                'and leaves margin. The math is not close.'
            ),
            'action': (
                'Post a Upwork or Contra job posting today. Requirements: WordPress expert, Divi 5, '
                'ability to follow a content brief, experience with professional service businesses. '
                'Budget: $25–40/hr or flat-rate per site. Provide Divi access + brief. Review work in stages.'
            ),
        },
        {
            'num': '08',
            'title': 'ASG + DMA Sites — Minimum Viable Credential Pages',
            'origin': 'Roger (Item 4)',
            'tier_hex': CHARCOAL_HEX,
            'tier': 'MEDIUM',
            'owner': 'Roger (brief) → Builder (execution)',
            'time': '1 hr brief per site',
            'rationale': (
                'These sites do not need to be elaborate this week. They need to exist and pass the '
                '"Google them after a cold outreach" test. A prospect who receives Roger\'s direct mail '
                'packet will search ASG and DMA before calling. They need to find a professional '
                'presence — not a blank page. Minimum viable: Home (who we are, who we serve), '
                'Services (what we do), About (Roger\'s credibility + team framing), Contact.'
            ),
            'action': (
                'Write a one-page brief for each. Install Divi 5. Hand to builder with copy '
                'direction. DMA can repurpose the audience persona positioning and implant dentist '
                'content already produced. ASG can repurpose the portfolio brief narrative.'
            ),
        },
        {
            'num': '09',
            'title': 'Draft the Direct Mail Packet — Implant Dentists',
            'origin': 'Board Addition',
            'tier_hex': CHARCOAL_HEX,
            'tier': 'MEDIUM',
            'owner': 'Team (Vivienne Carr, CMO)',
            'time': '2–3 hrs production',
            'rationale': (
                'Direct mail is on Roger\'s radar as a client acquisition channel. Before printing anything, '
                'three conditions must be true: (1) DMA website exists and is live, (2) the offer is clear '
                '(free Digital Authority Assessment), and (3) the landing page is built to receive the response. '
                'This week\'s goal is not to mail anything — it is to draft the packet copy so it is ready '
                'to deploy the moment the site is live. Do not print before you have a destination.'
            ),
            'action': (
                'Vivienne Carr drafts a two-component direct mail packet: (1) outer envelope teaser line, '
                '(2) one-page letter with offer (free authority assessment) + response mechanism '
                '(URL + phone). Roger reviews and approves. Hold for site launch.'
            ),
        },
        {
            'num': '10',
            'title': 'Create a Simple Prospect Pipeline Tracker',
            'origin': 'Board Addition',
            'tier_hex': CHARCOAL_HEX,
            'tier': 'MEDIUM',
            'owner': 'Roger (5 min setup)',
            'time': '15–30 min',
            'rationale': (
                'When outreach starts producing responses — and it will — Roger needs a way to track '
                'who he contacted, their status, and the next action. A spreadsheet or a basic GHL '
                'pipeline is fine. The risk without this: leads fall through the cracks while Roger is '
                'building websites. One missed follow-up is a missed client. At this stage, a manual '
                'Google Sheet with columns (Name / Niche / Contact Date / Status / Next Step) takes '
                'fifteen minutes to build and prevents a $2,500 engagement from disappearing.'
            ),
            'action': 'Build a simple Google Sheet or GHL pipeline. Log every outreach contact made this week.',
        },
    ]

    for p in priorities:
        # Number + title bar
        nt = doc.add_table(1, 2); remove_table_borders(nt)
        nt.columns[0].width = Inches(0.45); nt.columns[1].width = Inches(6.05)

        nc = nt.cell(0, 0); set_cell_bg(nc, p['tier_hex'])
        np_ = nc.paragraphs[0]; np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(5); np_.paragraph_format.space_after = Pt(5)
        nr = np_.add_run(p['num'])
        nr.bold = True; nr.font.size = Pt(12); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

        tc_ = nt.cell(0, 1); set_cell_bg(tc_, SLATE_HEX)
        tp = tc_.paragraphs[0]
        tp.paragraph_format.space_before = Pt(5); tp.paragraph_format.space_after = Pt(5)
        tr_ = tp.add_run(p['title'])
        tr_.bold = True; tr_.font.size = Pt(10.5); tr_.font.color.rgb = WHITE; tr_.font.name = 'Calibri'

        # Meta row
        meta_t = doc.add_table(1, 4); remove_table_borders(meta_t)
        meta_t.columns[0].width = Inches(1.3); meta_t.columns[1].width = Inches(1.1)
        meta_t.columns[2].width = Inches(1.5); meta_t.columns[3].width = Inches(2.6)
        meta_items = [
            ('Origin', p['origin']),
            ('Tier', p['tier']),
            ('Owner', p['owner']),
            ('Time Estimate', p['time']),
        ]
        for ci, (label, val) in enumerate(meta_items):
            c = meta_t.cell(0, ci); set_cell_bg(c, LIGHT_BG)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
            r1 = cp.add_run(f'{label}: ')
            r1.bold = True; r1.font.size = Pt(8); r1.font.color.rgb = ASG_CHARCOAL; r1.font.name = 'Calibri'
            r2 = cp.add_run(val)
            r2.font.size = Pt(8); r2.font.color.rgb = BODY_COLOR; r2.font.name = 'Calibri'

        # Rationale
        rat_p = doc.add_paragraph(); set_para_bg(rat_p, TINT_BG)
        rat_p.paragraph_format.space_before = Pt(0); rat_p.paragraph_format.space_after = Pt(0)
        rat_p.paragraph_format.left_indent = Inches(0.1); rat_p.paragraph_format.right_indent = Inches(0.1)
        r1 = rat_p.add_run('  Rationale: ')
        r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = ASG_CHARCOAL; r1.font.name = 'Calibri'
        r2 = rat_p.add_run(p['rationale'])
        r2.font.size = Pt(9.5); r2.font.color.rgb = BODY_COLOR; r2.font.name = 'Calibri'

        # Action
        act_p = doc.add_paragraph(); act_p.paragraph_format.space_before = Pt(0); act_p.paragraph_format.space_after = Pt(0)
        act_p.paragraph_format.left_indent = Inches(0.1); act_p.paragraph_format.right_indent = Inches(0.1)
        a1 = act_p.add_run('  Action: ')
        a1.bold = True; a1.font.size = Pt(9.5); a1.font.color.rgb = RGBColor(0x27, 0xAE, 0x60); a1.font.name = 'Calibri'
        a2 = act_p.add_run(p['action'])
        a2.font.size = Pt(9.5); a2.font.color.rgb = BODY_COLOR; a2.font.name = 'Calibri'

        sp(doc, 10)

    add_pb(doc)


# ── Section 2: Site Build — Hire or DIY ───────────────────────────────────────

def build_hire_verdict(doc):
    section_hdr(doc, '02 — SHOULD ROGER HIRE A SITE BUILDER?  |  The Honest Answer')
    sp(doc, 8)

    body(doc,
        'The board does not vote on this one. This is unanimous.',
        size=11, bold=True, color=DARK_TEXT, after=10
    )

    # Verdict box
    vp = doc.add_paragraph(); set_para_bg(vp, GREEN_HEX)
    vp.paragraph_format.space_before = Pt(0); vp.paragraph_format.space_after = Pt(0)
    vp.paragraph_format.left_indent = Inches(0.15); vp.paragraph_format.right_indent = Inches(0.15)
    vr = vp.add_run(
        '  VERDICT: YES. Hire a WordPress / Divi 5 developer. This week. '
        'Roger\'s time is worth more spent on sales and strategy than inside a page builder.  '
    )
    vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = WHITE; vr.font.name = 'Calibri'

    sp(doc, 10)

    sub_hdr(doc, 'The Math (Preston Adler, COO)')
    rows = [
        ('Roger builds all 5 sites himself', '40–60 hrs in Divi', '$0 labor cost', 'Zero time for sales, zero new clients booked this week'),
        ('Roger hires a Divi builder', '~$1,200–$2,000 total', 'One client closes = $2,500+', 'Roger uses freed time to close 1–2 clients; net positive before sites launch'),
    ]
    t = doc.add_table(len(rows) + 1, 4); remove_table_borders(t)
    t.columns[0].width = Inches(1.8); t.columns[1].width = Inches(1.3)
    t.columns[2].width = Inches(1.4); t.columns[3].width = Inches(2.0)
    hdr_row = t.rows[0]
    for ci, lbl in enumerate(['Scenario', 'Time Cost', 'Revenue Potential', 'Net Impact']):
        c = hdr_row.cells[ci]; set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
        cr = cp.add_run(lbl); cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = WHITE; cr.font.name = 'Calibri'
    for i, row_data in enumerate(rows):
        row = t.rows[i + 1]; bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = row.cells[ci]; set_cell_bg(c, bg)
            cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(val); cr.font.size = Pt(9.5); cr.font.color.rgb = BODY_COLOR; cr.font.name = 'Calibri'

    sp(doc, 10)

    sub_hdr(doc, 'What Roger Does. What the Builder Does.')
    body(doc, "Roger's role — cannot be delegated:", bold=True, color=DARK_TEXT, after=4)
    blt(doc, 'Writes or approves the site brief (avatar, positioning, CTA structure, services) — 30–60 min per site')
    blt(doc, 'Reviews each page before it goes live — 20–30 min per site')
    blt(doc, 'Has the sales conversations that produce the clients the sites are meant to convert')
    blt(doc, 'Domains, hosting, Divi license — Roger sets this up once, then hands keys to builder')

    sp(doc, 6)
    body(doc, "Builder's role — fully delegatable:", bold=True, color=DARK_TEXT, after=4)
    blt(doc, 'Takes the Divi 5 skeleton Roger set up and builds out all pages')
    blt(doc, 'Drops in the copy from the brief and content assets the team has produced')
    blt(doc, 'Formats headers, service sections, team bios, CTA buttons, mobile responsiveness')
    blt(doc, 'Makes revision cycles based on Roger\'s feedback')

    sp(doc, 8)
    sub_hdr(doc, 'The Divi 5 Setup-First Idea — Confirmed Smart')
    body(doc,
        'Roger\'s instinct to get the domains set up with Divi 5 installed before handing off is correct. '
        'A builder who hits a blank host with no structure wastes time on setup decisions that Roger '
        'should control. Install Divi 5, create the page shell (Home / Services / About / Contact / Blog), '
        'set the color palette and typography to match each brand — then hand off. That is the efficient split.',
        after=6
    )
    body(doc,
        'Estimated time for Roger to do this for all five sites: 3–4 hours total. '
        'That is the right investment. Everything after that is builder work.',
        italic=True, color=ASG_CHARCOAL, after=6
    )

    add_pb(doc)


# ── Section 3: Thomas Banks Strategic Analysis ────────────────────────────────

def build_thomas_banks(doc):
    section_hdr(doc, '03 — THE THOMAS BANKS OPPORTUNITY  |  Daniel Frost, CSO + Tanya Blackwood, CRO')
    sp(doc, 8)

    body(doc,
        'The team reviewed both sites before this meeting. Here is the full read:',
        after=10
    )

    sub_hdr(doc, 'Site 1: KentuckianaDivorceAttorney.com (Thomas Banks — Southern Indiana)')
    rows1 = [
        ('Positioning', 'Generic location + specialty claim. No differentiation. No authority markers.'),
        ('Trust Signals', 'None visible. No reviews, no credentials, no years of experience, no outcomes.'),
        ('Messaging Quality', 'Template-default. Could apply to any family law attorney in any state.'),
        ('Conversion Elements', 'Basic contact form. No urgency, no offer, no reason to call today vs. tomorrow.'),
        ('Verdict', 'This site is doing no work for Thomas. A prospect who Googles him after a referral sees nothing that confirms the referral was right.'),
    ]
    _build_analysis_table(doc, rows1, RED_HEX)

    sp(doc, 10)
    sub_hdr(doc, 'Site 2: divorceky-in.com (Firm Site — Straw-Boone Doheny Banks et al.)')
    rows2 = [
        ('Positioning', '"Elite Family Law Advocacy Standing Beside You Every Step of the Way." Stronger than most. Authority + compassion framing.'),
        ('Trust Signals', 'Super Lawyers designation, 4.7-star aggregate, multiple attorney bios, third-party directory presence. Real social proof.'),
        ('Messaging Quality', 'Professionally written. Consistent voice. Acknowledges complexity and emotional burden.'),
        ('Weakness', 'The "elite + compassionate" pairing is common in family law. Not uniquely differentiated. Blog/thought leadership absent.'),
        ('Verdict', 'This is the benchmark Thomas needs to beat on his personal site — not just match. If his site outperforms the firm\'s, the partners will notice.'),
    ]
    _build_analysis_table(doc, rows2, AMBER_HEX)

    sp(doc, 12)
    sub_hdr(doc, 'The Strategic Play (Daniel Frost)')
    body(doc,
        'Thomas Banks is a multi-opportunity account. Here is how to sequence it:',
        after=6
    )

    steps = [
        ('Step 1 — Deliver exceptional work for Thomas personally.',
         'Redesign KentuckianaDivorceAttorney.com to outperform the firm site on every dimension: '
         'stronger positioning, documented trust signals, emotional resonance, clear CTA. '
         'Make it look like Thomas has his own PR firm behind him.'),
        ('Step 2 — Let the result do the selling.',
         'When the other attorneys at Straw-Boone Doheny Banks see Thomas\'s personal site, they will '
         'ask who built it. Thomas becomes the referral source without Roger having to pitch anyone. '
         'That is the highest-leverage sales motion available: let the work close the room.'),
        ('Step 3 — Approach the firm after Thomas\'s site launches.',
         'Roger follows up with Thomas: "Is there an attorney at your firm who could use the same '
         'treatment?" One introduction. One meeting. Multiple potential engagements. The firm has '
         'at least 6 named attorneys visible on their site. Even converting two of them to '
         'Authority Blueprint clients at $2,500 each represents more revenue than two cold-sourced clients.'),
        ('Step 4 — Document Thomas as Case Study #1 (LGD).',
         'Before/after screenshots, traffic or inquiry metrics (if Thomas will share them), '
         'a brief written testimonial. This becomes the proof that sells every subsequent family law attorney.'),
    ]

    for label, text in steps:
        row_p = doc.add_paragraph(); row_p.paragraph_format.space_before = Pt(4); row_p.paragraph_format.space_after = Pt(4)
        r1 = row_p.add_run(f'{label}  ')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = DARK_TEXT; r1.font.name = 'Calibri'
        r2 = row_p.add_run(text)
        r2.font.size = Pt(10); r2.font.color.rgb = BODY_COLOR; r2.font.name = 'Calibri'

    sp(doc, 8)
    body(doc,
        'Tanya Blackwood adds: the Thomas Banks engagement is the single fastest path to an LGD case study '
        'in the portfolio. Roger already has the relationship. The site brief is straightforward. '
        'The deliverable is clear. Prioritize this above any cold-sourced LGD prospect.',
        italic=True, color=ASG_CHARCOAL, after=6
    )

    add_pb(doc)


def _build_analysis_table(doc, rows, accent_hex):
    t = doc.add_table(len(rows), 2); remove_table_borders(t)
    t.columns[0].width = Inches(1.4); t.columns[1].width = Inches(5.1)
    for i, (label, text) in enumerate(rows):
        bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        lc = t.cell(i, 0); set_cell_bg(lc, bg)
        lp = lc.paragraphs[0]; lp.paragraph_format.space_before = Pt(4); lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = ASG_CHARCOAL; lr.font.name = 'Calibri'
        tc_ = t.cell(i, 1); set_cell_bg(tc_, bg)
        tp = tc_.paragraphs[0]; tp.paragraph_format.space_before = Pt(4); tp.paragraph_format.space_after = Pt(4)
        tr_ = tp.add_run(text); tr_.font.size = Pt(9.5); tr_.font.color.rgb = BODY_COLOR; tr_.font.name = 'Calibri'


# ── Section 4: Direct Mail Question ───────────────────────────────────────────

def build_direct_mail(doc):
    section_hdr(doc, '04 — THE DIRECT MAIL QUESTION  |  Tanya Blackwood, CRO')
    sp(doc, 8)

    body(doc,
        'Roger mentioned sending direct mail packets to potential DMA and ASG prospects. '
        'The team supports this as a medium-term channel. The honest timing note is below.',
        after=10
    )

    # Timing verdict
    vp = doc.add_paragraph(); set_para_bg(vp, AMBER_HEX)
    vp.paragraph_format.space_before = Pt(0); vp.paragraph_format.space_after = Pt(8)
    vp.paragraph_format.left_indent = Inches(0.15)
    vr = vp.add_run(
        '  Do not mail this week. Three gates must be open first: '
        '(1) live website to send them to, (2) clear offer with response mechanism, '
        '(3) pipeline to handle responses. None of these are ready yet.  '
    )
    vr.bold = True; vr.font.size = Pt(10); vr.font.color.rgb = WHITE; vr.font.name = 'Calibri'

    sub_hdr(doc, 'The Sequence')
    seq_rows = [
        ('This week', 'Draft the direct mail letter copy (team handles this). Hold for printing.'),
        ('Week 2–3', 'DMA and ASG sites go live. Verify response landing pages are functional.'),
        ('Week 3–4', 'Build the GHL pipeline to receive and track inbound mail responses.'),
        ('Week 4+', 'Print and mail first batch. Start with 50–75 implant dentists in one geographic cluster. Measure response rate before scaling.'),
    ]
    t = doc.add_table(len(seq_rows) + 1, 2); remove_table_borders(t)
    t.columns[0].width = Inches(1.3); t.columns[1].width = Inches(5.2)
    hr = t.rows[0]
    for ci, lbl in enumerate(['When', 'Action']):
        c = hr.cells[ci]; set_cell_bg(c, ASG_DARK_HEX)
        cp = c.paragraphs[0]; cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
        cr = cp.add_run(lbl); cr.bold = True; cr.font.size = Pt(9); cr.font.color.rgb = WHITE; cr.font.name = 'Calibri'
    for i, (when, action) in enumerate(seq_rows):
        row = t.rows[i + 1]; bg = TINT_BG if i % 2 == 0 else 'FFFFFF'
        c0 = row.cells[0]; set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]; p0.paragraph_format.space_before = Pt(4); p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(when); r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = ASG_CHARCOAL; r0.font.name = 'Calibri'
        c1 = row.cells[1]; set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(action); r1.font.size = Pt(9.5); r1.font.color.rgb = BODY_COLOR; r1.font.name = 'Calibri'

    sp(doc, 10)
    body(doc,
        'One more note on direct mail economics: it works best when the offer has a clear, low-friction '
        'response path ("scan this QR code to get your free site analysis") and when the list is highly '
        'targeted (implant dentists within a specific MSA, not national broadcast). Start local, '
        'measure, then expand. Do not spend $800 on printing before you know the message converts.',
        italic=True, color=ASG_CHARCOAL, after=6
    )

    add_pb(doc)


# ── Section 5: Closing Verdicts ────────────────────────────────────────────────

def build_verdicts(doc):
    section_hdr(doc, '05 — BOARD CLOSING VERDICTS')
    sp(doc, 8)

    body(doc, 'One sentence per team member. No hedging.', after=12)

    verdicts = [
        ('Daniel Frost, CSO',
         'The Thomas Banks firm play is the most overlooked opportunity on the list — one relationship, '
         'six attorneys, multiple case studies, zero cold outreach required. Make that site outstanding.'),
        ('Tanya Blackwood, CRO',
         'Revenue this week does not come from a website — it comes from a phone call. '
         'Send the warm outreach messages before touching a single page builder.'),
        ('Preston Adler, COO',
         'Five websites in one week built by one person is a plan that produces zero websites '
         'and zero clients. Hire the builder, set the structure, hand off the execution.'),
        ('Iris Nolan, CTO',
         'The Divi 5 setup-first idea is correct — control the architecture, delegate the content build. '
         'Get hosting, DNS, and theme install done for all five domains in one sitting this week.'),
        ('Vivienne Carr, CMO',
         'The copy assets are already built. The only thing stopping them from being on live sites '
         'is sequencing. LGD services pages go live Monday — the content is done, it just needs to be published.'),
        ('Roger Bauer, Director',
         'The list is real. The work will get done. The sites will exist. '
         'What matters is that a paying client is in motion before any of them launch.'),
    ]

    for name, verdict in verdicts:
        vt = doc.add_table(1, 2); remove_table_borders(vt)
        vt.columns[0].width = Inches(1.8); vt.columns[1].width = Inches(4.7)

        nc = vt.cell(0, 0); set_cell_bg(nc, ASG_DARK_HEX)
        np_ = nc.paragraphs[0]; np_.paragraph_format.space_before = Pt(6); np_.paragraph_format.space_after = Pt(6)
        nr = np_.add_run(f'  {name}')
        nr.bold = True; nr.font.size = Pt(9); nr.font.color.rgb = WHITE; nr.font.name = 'Calibri'

        tc_ = vt.cell(0, 1); set_cell_bg(tc_, TINT_BG)
        tp = tc_.paragraphs[0]; tp.paragraph_format.space_before = Pt(6); tp.paragraph_format.space_after = Pt(6)
        tr_ = tp.add_run(verdict)
        tr_.font.size = Pt(9.5); tr_.font.color.rgb = BODY_COLOR; tr_.font.name = 'Calibri'

        sp(doc, 4)

    sp(doc, 20)
    divider(doc)
    sp(doc, 8)

    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(0); ep.paragraph_format.space_after = Pt(4)
    er = ep.add_run(
        'Authority Systems Group™  —  Board Meeting  —  March 8, 2026\n'
        'Internal Document  |  Confidential  |  Not for Distribution'
    )
    er.font.size = Pt(8); er.font.color.rgb = FOOTER_COLOR; er.italic = True; er.font.name = 'Calibri'


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for attr, val in [('left_margin', 1.0), ('right_margin', 1.0), ('top_margin', 0.75), ('bottom_margin', 0.75)]:
        setattr(section, attr, Inches(val))

    add_header_footer(doc)
    build_cover(doc)
    build_priorities(doc)
    build_hire_verdict(doc)
    build_thomas_banks(doc)
    build_direct_mail(doc)
    build_verdicts(doc)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
