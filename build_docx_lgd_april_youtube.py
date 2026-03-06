#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
April 2026 YouTube Long-Form Script Pack
Client: Legal Growth Dynamics — Family Law Attorney Market
Author: Roger Bauer
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ASG_BLUE = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
BROLL_COLOR = RGBColor(0x44, 0x77, 0x99)   # muted blue for B-roll cues
PACING_COLOR = RGBColor(0x77, 0x77, 0x77)  # gray for pacing notes

INPUT_MD = "client-onboarding/clients/authority-systems-group/outputs/lgd-family-law-attorneys_april-youtube-30day-scripts_20260305.md"
OUTPUT_DOCX = "client-onboarding/clients/authority-systems-group/outputs/lgd-family-law-attorneys_april-youtube-30day-scripts_20260305.docx"


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color='25AAE1'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_left_border_paragraph(doc, text, font_color, fill_hex, border_color, font_size=10, italic=False, bold=False):
    """Render a callout-style paragraph with a left border accent."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(font_size)
    r.font.color.rgb = font_color
    r.font.italic = italic
    r.font.bold = bold
    set_paragraph_shading(p, fill_hex)
    p.paragraph_format.left_indent = Inches(0.25)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = Pt(15)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ASG_BLUE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)


def add_cover_page(doc):
    # Main title
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('APRIL 2026')
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = ASG_BLUE
    cover_title.paragraph_format.space_before = Pt(48)
    cover_title.paragraph_format.space_after = Pt(4)

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub1.add_run('YouTube Long-Form Script Pack')
    r.font.name = 'Calibri'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = ASG_CHARCOAL
    sub1.paragraph_format.space_after = Pt(4)

    theme = doc.add_paragraph()
    theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = theme.add_run('The Q2 Reckoning — 22 Complete Scripts')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(14)
    r2.font.italic = True
    r2.font.color.rgb = ASG_CHARCOAL
    theme.paragraph_format.space_after = Pt(20)

    add_horizontal_rule(doc)

    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = firm.add_run('Legal Growth Dynamics')
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = ASG_CHARCOAL
    firm.paragraph_format.space_after = Pt(4)

    market = doc.add_paragraph()
    market.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = market.add_run('Family Law Attorney Market — Continental US')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(13)
    r2.font.color.rgb = ASG_CHARCOAL
    market.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)

    credits = [
        ('Produced by: Milo Vance, YouTube Long-Form Script Writer', True, False),
        ('Directed by: Marcus Webb, Director of Content Marketing', False, False),
        ('Under the direction of Roger Bauer, Director — Authority Systems Group™', True, True),
        ('', False, False),
        ('April 2026  |  Confidential', False, True),
        ('Legal Growth Dynamics — Family Law Attorney Market', False, False),
    ]

    for text, is_blue, is_bold in credits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = is_bold
        r.font.color.rgb = ASG_BLUE if is_blue else ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        'Authority Systems Group™ — Confidential. Prepared exclusively for Legal Growth Dynamics.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def render_inline(p, text, base_bold=False, base_italic=False):
    """Parse **bold**, *italic*, and `code` inline markers and add runs to paragraph."""
    # Split on **bold**, then *italic*, then `code`
    token_pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = token_pattern.split(text)
    for part in parts:
        if not part:
            continue
        r = p.add_run()
        r.font.name = 'Georgia'
        r.font.size = Pt(11)
        r.font.color.rgb = BODY_COLOR
        if part.startswith('**') and part.endswith('**'):
            r.text = part[2:-2]
            r.font.bold = True
            r.font.italic = base_italic
        elif part.startswith('*') and part.endswith('*'):
            r.text = part[1:-1]
            r.font.italic = True
            r.font.bold = base_bold
        elif part.startswith('`') and part.endswith('`'):
            r.text = part[1:-1]
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
            r.font.color.rgb = ASG_CHARCOAL
        else:
            r.text = part
            r.font.bold = base_bold
            r.font.italic = base_italic


def is_script_section_header(line):
    """Detect section headers like 'SECTION 1 — HOOK (0:00–0:45)' etc."""
    patterns = [
        r'^(SECTION \d+ —)',
        r'^(HOOK|SETUP|THE TURN|THE PROOF|THE CLOSE|PRODUCTION NOTES|VALUE BOMB|CONTRARIAN)',
        r'^(5A —|5B —|5C —)',
    ]
    stripped = line.strip().lstrip('#').strip()
    for pat in patterns:
        if re.match(pat, stripped):
            return True
    return False


def parse_and_render(doc, md_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block = False
    skip_frontmatter = True
    frontmatter_dashes = 0
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip YAML front matter (between first two --- lines)
        if skip_frontmatter:
            stripped = line.strip()
            if stripped == '---':
                frontmatter_dashes += 1
                if frontmatter_dashes == 2:
                    skip_frontmatter = False
            i += 1
            continue

        # Code block toggle
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ('---', '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-\|:]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            para = cell.paragraphs[0]
                            para.paragraph_format.space_after = Pt(2)
                            r = para.add_run(cell_text)
                            r.font.name = 'Calibri'
                            r.font.size = Pt(9)
                            if ri == 0:
                                r.font.bold = True
                                r.font.color.rgb = WHITE
                            else:
                                r.font.color.rgb = BODY_COLOR
                            if ri == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '25AAE1')
                                tcPr.append(shd)
                            elif ri % 2 == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), 'F5F5F5')
                                tcPr.append(shd)
            continue

        # Headings
        if line.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # B-roll cues: [B-ROLL: ...]
        broll_match = re.match(r'^\[B-ROLL:\s*(.*?)\]$', line.strip())
        if broll_match:
            add_left_border_paragraph(
                doc,
                '▶ B-ROLL: ' + broll_match.group(1),
                font_color=BROLL_COLOR,
                fill_hex='EAF6FC',
                border_color='25AAE1',
                font_size=9.5,
                italic=True
            )
            i += 1
            continue

        # Chapter markers: [CHAPTER: ...]
        chapter_match = re.match(r'^\[CHAPTER:\s*(.*?)\]$', line.strip())
        if chapter_match:
            p = doc.add_paragraph()
            r = p.add_run('⏱ CHAPTER: ' + chapter_match.group(1))
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F0F0F0')
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.1)
            i += 1
            continue

        # Pacing notes: [PACING: ...]
        pacing_match = re.match(r'^\[PACING:\s*(.*?)\]$', line.strip())
        if pacing_match:
            p = doc.add_paragraph()
            r = p.add_run('PACING: ' + pacing_match.group(1))
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = PACING_COLOR
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.1)
            i += 1
            continue

        # Script section headers (SECTION 1 — HOOK, etc.)
        # These often appear as bold lines or ### headers — catch plain-text versions
        section_patterns = [
            r'^(SECTION \d+ —\s+.*)',
            r'^(HOOK \(|SETUP / PROBLEM|THE TURN \(|THE PROOF|THE CLOSE \(|PRODUCTION NOTES —|VALUE BOMB VARIANT —|CONTRARIAN VARIANT —)',
            r'^(5A —|5B —|5C —)',
            r'^(Identity Anchor|CTA \(|Channel Hook)',
        ]
        is_sect = any(re.match(pat, line.strip()) for pat in section_patterns)
        if is_sect:
            p = doc.add_paragraph()
            r = p.add_run(line.strip())
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.bold = True
            if 'VALUE BOMB' in line or 'CONTRARIAN' in line:
                r.font.color.rgb = ASG_CHARCOAL
                set_paragraph_shading(p, 'F5F5F5')
            elif 'PRODUCTION NOTES' in line:
                r.font.color.rgb = ASG_CHARCOAL
            else:
                r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Blockquotes (spoken script lines)
        if line.startswith('> '):
            text = line[2:].strip()
            # Remove surrounding *italic* markers if present
            text = re.sub(r'^\*(.+)\*$', r'\1', text)
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Georgia'
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x22, 0x22, 0x44)
            set_paragraph_shading(p, 'EAF6FC')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.1)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet points
        if re.match(r'^[-*]\s', line):
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            render_inline(p, line[2:].strip())
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            p.clear()
            render_inline(p, text)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        render_inline(p, line)
        p.paragraph_format.space_after = Pt(6)
        i += 1


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    parse_and_render(doc, INPUT_MD)

    doc.save(OUTPUT_DOCX)
    print(f"✓ Document saved: {OUTPUT_DOCX}")
    print(f"  Input:  {INPUT_MD}")
    print(f"  Output: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
