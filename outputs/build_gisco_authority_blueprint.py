#!/usr/bin/env python3
"""
Authority Systems Group™ — Authority Blueprint™
Client: GISCO Corp — Louisville, Kentucky
Director: Roger Bauer | Date: March 9, 2026
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK_HEX  = '1A2533'
ASG_BLUE_HEX  = '25aae1'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_TINT      = 'EAF6FC'
SLATE_HEX     = '2C3E50'
GREEN_HEX     = '27AE60'
AMBER_HEX     = 'E67E22'
RED_HEX       = '7D2400'
GOLD_HEX      = 'C9A02C'
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)
DARK_COLOR    = RGBColor(0x1A, 0x25, 0x33)

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(
    BASE,
    'client-onboarding/clients/gisco-corp/outputs/docx/'
    'gisco-corp_authority-blueprint_20260309.docx'
)


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_pb(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def sp(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(size)


def set_margins(doc, top=0.75, bottom=0.75, left=0.9, right=0.9):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


# ── Typography Helpers ─────────────────────────────────────────────────────────

def heading(doc, text, level=1, color=None, bg=None, size=None, bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT, caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    if bg:
        set_para_bg(p, bg)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.font.name = 'Calibri'
    r.font.color.rgb = color if color else (WHITE if bg else ASG_CHARCOAL)
    if size:
        r.font.size = Pt(size)
    elif level == 1:
        r.font.size = Pt(20)
    elif level == 2:
        r.font.size = Pt(14)
    elif level == 3:
        r.font.size = Pt(11)
    return p


def body(doc, text, size=10, color=None, bold=False, italic=False, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def bullet(doc, text, size=10, indent=0.25, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(size)
    r.bold       = bold
    r.font.color.rgb = color if color else BODY_COLOR
    return p


def divider(doc, hex_color=ASG_BLUE_HEX, height_pt=2):
    p = doc.add_paragraph()
    set_para_bg(p, hex_color)
    p.paragraph_format.space_before = Pt(height_pt)
    p.paragraph_format.space_after  = Pt(height_pt)


def section_header(doc, number, title, subtitle=None):
    """Full-width dark section banner with section number + title."""
    p = doc.add_paragraph()
    set_para_bg(p, ASG_DARK_HEX)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.15)
    rn = p.add_run(f'SECTION {number}  ')
    rn.bold = True
    rn.font.name = 'Calibri'
    rn.font.size = Pt(9)
    rn.font.color.rgb = ASG_BLUE
    rt = p.add_run(f'— {title.upper()}')
    rt.bold = True
    rt.font.name = 'Calibri'
    rt.font.size = Pt(13)
    rt.font.color.rgb = WHITE
    if subtitle:
        sp(doc, 1)
        body(doc, subtitle, size=9, italic=True, color=ASG_CHARCOAL, indent=0.15)


def attribution(doc, text):
    """Italic attribution line."""
    p = body(doc, f'— {text}', size=9, italic=True, color=ASG_CHARCOAL)
    return p


def quote_block(doc, text, size=10.5):
    """Blue left-bar quote block."""
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(0.1)
    tbl.columns[1].width = Inches(7.4)
    set_cell_bg(tbl.rows[0].cells[0], ASG_BLUE_HEX)
    cc = tbl.rows[0].cells[1]
    set_cell_bg(cc, ASG_TINT)
    cp = cc.paragraphs[0]
    cp.paragraph_format.left_indent  = Inches(0.15)
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after  = Pt(6)
    r = cp.add_run(text)
    r.italic = True
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.color.rgb = BODY_COLOR
    sp(doc, 4)


def callout_box(doc, label, text, bg=ASG_TINT, label_color=None):
    """Tinted callout block with bold label + body text."""
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(7.5)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent  = Inches(0.15)
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after  = Pt(5)
    rl = cp.add_run(label + '  ')
    rl.bold = True
    rl.font.name = 'Calibri'
    rl.font.size = Pt(9)
    rl.font.color.rgb = label_color if label_color else ASG_BLUE
    rv = cp.add_run(text)
    rv.font.name = 'Calibri'
    rv.font.size = Pt(9.5)
    rv.font.color.rgb = BODY_COLOR
    sp(doc, 4)


def two_col_table(doc, rows_data, col_widths=(2.8, 4.7), header=None, header_bg=ASG_DARK_HEX):
    """Generic 2-column alternating-row table."""
    num_rows = len(rows_data) + (1 if header else 0)
    tbl = doc.add_table(rows=num_rows, cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(col_widths[0])
    tbl.columns[1].width = Inches(col_widths[1])
    row_offset = 0
    if header:
        for ci, h in enumerate(header):
            cell = tbl.rows[0].cells[ci]
            set_cell_bg(cell, header_bg)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.1)
            r = p.add_run(h)
            r.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = WHITE
        row_offset = 1
    for i, (lbl, val) in enumerate(rows_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        lc = tbl.rows[i + row_offset].cells[0]
        vc = tbl.rows[i + row_offset].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(vc, bg)
        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent = Inches(0.1)
        lr = lp.add_run(lbl)
        lr.bold = True
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = ASG_CHARCOAL
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.name = 'Calibri'
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = BODY_COLOR
    sp(doc, 6)


def four_col_kpi_table(doc, rows_data, headers):
    """4-column KPI table."""
    num_rows = len(rows_data) + 1
    tbl = doc.add_table(rows=num_rows, cols=4)
    remove_table_borders(tbl)
    widths = [2.4, 1.4, 1.5, 1.6]
    for ci, w in enumerate(widths):
        tbl.columns[ci].width = Inches(w)
    # Header row
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, ASG_DARK_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.08)
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE
    for i, row_vals in enumerate(rows_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_vals):
            cell = tbl.rows[i + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.08)
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = BODY_COLOR
            if ci == 0:
                r.bold = True
                r.font.color.rgb = ASG_CHARCOAL
    sp(doc, 6)


def five_col_kpi_table(doc, rows_data, headers):
    """5-column KPI table (metric + 4 time periods)."""
    num_rows = len(rows_data) + 1
    tbl = doc.add_table(rows=num_rows, cols=5)
    remove_table_borders(tbl)
    widths = [2.3, 1.1, 1.2, 1.3, 1.4]
    for ci, w in enumerate(widths):
        tbl.columns[ci].width = Inches(w)
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, ASG_DARK_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.06)
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
    for i, row_vals in enumerate(rows_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_vals):
            cell = tbl.rows[i + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.06)
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(8.5)
            r.font.color.rgb = BODY_COLOR
            if ci == 0:
                r.bold = True
                r.font.color.rgb = ASG_CHARCOAL
    sp(doc, 6)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Logo block
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(cover, ASG_DARK_HEX)
    cover.paragraph_format.space_before = Pt(0)
    cover.paragraph_format.space_after  = Pt(0)
    if os.path.exists(LOGO_PATH):
        run = cover.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))

    sp(doc, 4)
    accent = doc.add_paragraph()
    set_para_bg(accent, ASG_BLUE_HEX)
    accent.paragraph_format.space_before = Pt(3)
    accent.paragraph_format.space_after  = Pt(3)
    sp(doc, 12)

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run('AUTHORITY BLUEPRINT™')
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = ASG_BLUE
    r1.font.all_caps = True

    sp(doc, 4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('GISCO Corp')
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(36)
    r2.font.color.rgb = DARK_COLOR

    sp(doc, 4)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Louisville, Kentucky')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(14)
    r3.font.color.rgb = ASG_CHARCOAL

    sp(doc, 4)

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run('A Strategic Growth System Built Exclusively for Tony Long')
    r4.italic = True
    r4.font.name = 'Calibri'
    r4.font.size = Pt(11)
    r4.font.color.rgb = ASG_CHARCOAL

    sp(doc, 28)

    cred_tbl = doc.add_table(rows=1, cols=4)
    remove_table_borders(cred_tbl)
    for ci in range(4):
        cred_tbl.columns[ci].width = Inches(1.875)

    labels = ['PREPARED BY', 'DIRECTOR', 'DATE', 'CLASSIFICATION']
    values = ['Authority Systems Group™', 'Roger Bauer', 'March 9, 2026', 'Confidential']

    for i, (lbl, val) in enumerate(zip(labels, values)):
        cell = cred_tbl.rows[0].cells[i]
        set_cell_bg(cell, ASG_TINT)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = cp.add_run(lbl + '\n')
        rl.bold = True
        rl.font.name = 'Calibri'
        rl.font.size = Pt(8)
        rl.font.color.rgb = ASG_BLUE
        rv = cp.add_run(val)
        rv.font.name = 'Calibri'
        rv.font.size = Pt(9)
        rv.font.color.rgb = DARK_COLOR

    sp(doc, 20)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run(
        'Authority Systems Group™  —  Confidential. Prepared exclusively for Tony Long / GISCO Corp.'
    )
    footer_r.font.name = 'Calibri'
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = FOOTER_COLOR

    add_pb(doc)


# ── Table of Contents ──────────────────────────────────────────────────────────

def build_toc(doc):
    heading(doc, 'TABLE OF CONTENTS', level=2, color=DARK_COLOR, size=16, bold=True)
    divider(doc)
    sp(doc, 6)

    toc_items = [
        ('1', 'Executive Summary', 'Roger Bauer, Director'),
        ('2', "Director's Briefing", 'Roger Bauer, Director'),
        ('3', 'Market Intelligence', 'Daniel Frost, CSO'),
        ('4', 'Belief-to-Buy Framework™ Map', 'Dr. Raymond Cross + Vivienne Carr, CMO'),
        ('5', 'Priority Services Portfolio', 'Vivienne Carr, CMO + Tanya Blackwood, CRO'),
        ('6', 'Fast Cash Sprint™', 'Tanya Blackwood, CRO'),
        ('7', 'Year 1 Strategic Plan', 'Daniel Frost, CSO'),
        ('8', 'KPI Dashboard', 'Preston Adler, COO'),
        ('9', 'Content Authority Strategy™', 'Vivienne Carr, CMO'),
        ('10', 'Digital Infrastructure & Website Strategy', 'Iris Nolan, CTO'),
        ('11', 'Email Sequences (Fully Written)', 'Vivienne Carr, CMO'),
        ('12', 'Lead Magnet & Referral Program Framework', 'Vivienne Carr, CMO + Tanya Blackwood, CRO'),
        ('13', 'Next Steps & Delivery Timeline', 'Preston Adler, COO'),
        ('14', 'QC Gate Sign-Offs', 'Authority Systems Group™'),
    ]

    tbl = doc.add_table(rows=len(toc_items), cols=3)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(0.4)
    tbl.columns[1].width = Inches(4.6)
    tbl.columns[2].width = Inches(2.5)

    for i, (num, title, owner) in enumerate(toc_items):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        nc = tbl.rows[i].cells[0]
        tc = tbl.rows[i].cells[1]
        oc = tbl.rows[i].cells[2]
        set_cell_bg(nc, bg)
        set_cell_bg(tc, bg)
        set_cell_bg(oc, bg)

        np_ = nc.paragraphs[0]
        np_.paragraph_format.left_indent = Inches(0.1)
        nr = np_.add_run(num)
        nr.bold = True
        nr.font.name = 'Calibri'
        nr.font.size = Pt(9.5)
        nr.font.color.rgb = ASG_BLUE

        tp = tc.paragraphs[0]
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.name = 'Calibri'
        tr.font.size = Pt(9.5)
        tr.font.color.rgb = DARK_COLOR

        op = oc.paragraphs[0]
        orr = op.add_run(owner)
        orr.italic = True
        orr.font.name = 'Calibri'
        orr.font.size = Pt(8.5)
        orr.font.color.rgb = ASG_CHARCOAL

    add_pb(doc)


# ── Section 1: Executive Summary ───────────────────────────────────────────────

def build_s1_executive_summary(doc):
    section_header(doc, '1', 'Executive Summary')
    attribution(doc, 'Roger Bauer, Director | Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    body(doc, 'Tony —', size=11, bold=True)
    sp(doc, 2)
    body(doc,
         'You have something almost no one in your market has: four generations of insulation knowledge, '
         'forty-one years of credibility in Louisville, and a personal approach to every job that your '
         'reviews describe in near-identical terms — "he was the only one who actually crawled through '
         'the attic." That\'s not a marketing line. That\'s a fact. And right now, almost nobody outside '
         'of the people who\'ve already hired you knows it.',
         size=10.5)
    body(doc,
         'This Blueprint is built around one central mission: making what you already are findable — '
         'and making what you find do the work of filling your calendar, so you\'re not dependent on '
         'whoever happens to call.',
         size=10.5)

    sp(doc, 4)
    heading(doc, 'THE THREE CORE OPPORTUNITIES', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    opps = [
        ('OPPORTUNITY 1', 'The Institutional Authority Position',
         'GISCO has worked with Peerless Distillery, the University of Louisville, Jeff Boat, and '
         'Winston Industries. No competitor in the Louisville insulation market has that institutional '
         'track record and makes it visible. Cardinal Industrial has 70 years of history but produces '
         'zero educational content. MINCO exists online in name only. Alpha Insulation is a national '
         'chain that looks and feels corporate. The position of "Louisville\'s institutional insulation '
         'authority — the 4th generation contractor that built this city\'s industrial and commercial '
         'infrastructure" is unclaimed and GISCO is the only company that can credibly own it.',
         'Addressed in: Section 3 (Market Intelligence) + Section 7 (Year 1 Plan, Q1)'),
        ('OPPORTUNITY 2', 'The Video Authority Engine',
         'Every insulation contractor in Louisville is invisible on video. Not a single competitor has '
         'a YouTube presence, educational content, or any video that explains what insulation actually '
         'does, why it matters, or what a homeowner or facilities manager should know before hiring '
         'anyone. Tony has agreed to the video concept. With Tony on camera for even four to six '
         'sessions, we can create an AI-powered avatar on HeyGen and generate a library of 30+ videos '
         'from ASG-written scripts. This becomes the dominant educational voice in the Louisville '
         'insulation market, compounding in authority every week. No ad spend. No ongoing filming after '
         'the initial sessions.',
         'Addressed in: Section 9 (Content Strategy) + Section 10 (Digital Infrastructure)'),
        ('OPPORTUNITY 3', 'The "No Spray Foam" Differentiator',
         'The UK spray foam class action has changed the residential insulation conversation. '
         'Homeowners are searching. Facilities managers are asking questions. And GISCO\'s long-standing '
         'policy of not offering spray foam — made for exactly the reasons the lawsuit validated — is '
         'a positioning asset that no Louisville competitor can replicate without abandoning a service '
         'line. Properly positioned, "the insulation contractor who won\'t install a product they don\'t '
         'trust" is one of the most powerful trust signals available in this market right now.',
         'Addressed in: Section 4 (Belief-to-Buy Map) + Section 9 (Content Strategy)'),
    ]

    for num, title, desc, ref in opps:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(0.1)
        tbl.columns[1].width = Inches(7.4)
        set_cell_bg(tbl.rows[0].cells[0], ASG_BLUE_HEX)
        cc = tbl.rows[0].cells[1]
        set_cell_bg(cc, ASG_LIGHT_BG)
        cp = cc.paragraphs[0]
        cp.paragraph_format.left_indent  = Inches(0.15)
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after  = Pt(2)
        rn = cp.add_run(num + '  ')
        rn.bold = True; rn.font.name = 'Calibri'; rn.font.size = Pt(8)
        rn.font.color.rgb = ASG_BLUE
        rt = cp.add_run(title)
        rt.bold = True; rt.font.name = 'Calibri'; rt.font.size = Pt(10.5)
        rt.font.color.rgb = DARK_COLOR

        p2 = cc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.15)
        p2.paragraph_format.space_after = Pt(2)
        rd = p2.add_run(desc)
        rd.font.name = 'Calibri'; rd.font.size = Pt(9.5)
        rd.font.color.rgb = BODY_COLOR

        p3 = cc.add_paragraph()
        p3.paragraph_format.left_indent  = Inches(0.15)
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(6)
        rr = p3.add_run('→ ' + ref)
        rr.italic = True; rr.font.name = 'Calibri'; rr.font.size = Pt(8.5)
        rr.font.color.rgb = ASG_BLUE

        sp(doc, 4)

    sp(doc, 4)
    heading(doc, 'SUCCESS SNAPSHOTS', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    snaps = [
        ('At 90 Days:',
         'The AI phone employee is live and capturing every inquiry Tony misses. The new website '
         'is launched and converting. GISCO has 20+ Google reviews (up from near zero) with a '
         'systematic collection process running automatically. The first three Fast Cash strategies '
         'are deployed and generating inbound interest from past clients and new referral partners. '
         'Tony has recorded his first two video sessions, and the first HeyGen avatar scripts are '
         'written and scheduled.'),
        ('At 12 Months:',
         'GISCO is the only insulation contractor in Louisville with a 30+ video library, ranking '
         'organically for commercial and residential insulation terms across the Louisville metro. '
         'The referral partner network includes at least 8 active HVAC contractors and home inspectors '
         'routing warm leads directly to the AI phone system. Monthly inbound leads have increased '
         'from an estimated 5–15 to a consistent 25–35. Revenue is up 40–60% from baseline, driven '
         'entirely by organic channels.'),
    ]

    for label, text in snaps:
        callout_box(doc, label, text, bg=ASG_TINT, label_color=DARK_COLOR)

    sp(doc, 4)
    heading(doc, 'THE PRIORITY SEQUENCE — FIRST 30 DAYS', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    priority_items = [
        ('1. Launch the AI phone employee and new website', 'Section 10'),
        ('2. Deploy the Fast Cash Sprint', 'Section 6 — past client reactivation + HVAC referral outreach'),
        ('3. Record Tony\'s first two video sessions', 'Section 9 — unlocks the entire HeyGen content engine'),
    ]
    for item, ref in priority_items:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(5.5)
        tbl.columns[1].width = Inches(2.0)
        set_cell_bg(tbl.rows[0].cells[0], ASG_TINT)
        set_cell_bg(tbl.rows[0].cells[1], ASG_TINT)
        lp = tbl.rows[0].cells[0].paragraphs[0]
        lp.paragraph_format.left_indent = Inches(0.12)
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        lr = lp.add_run(item)
        lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(9.5)
        lr.font.color.rgb = DARK_COLOR
        rp = tbl.rows[0].cells[1].paragraphs[0]
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after  = Pt(4)
        rr = rp.add_run(ref)
        rr.italic = True; rr.font.name = 'Calibri'; rr.font.size = Pt(8.5)
        rr.font.color.rgb = ASG_BLUE
        sp(doc, 2)

    sp(doc, 6)
    attribution(doc, 'Roger Bauer, Director | Authority Systems Group™')
    add_pb(doc)


# ── Section 2: Director's Briefing ────────────────────────────────────────────

def build_s2_directors_briefing(doc):
    section_header(doc, '2', "Director's Briefing")
    attribution(doc, 'Roger Bauer, Director | Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    body(doc, 'Tony —', size=11, bold=True)
    sp(doc, 2)

    paras = [
        'When you told me you don\'t do spray foam because of the liability risk — that you\'ve watched '
        'that UK lawsuit unfold and you\'re not putting that product in anyone\'s house — I wrote that '
        'down. Because that\'s not a business decision. That\'s a values decision. And in a market full '
        'of contractors who\'ll install whatever product they can buy cheapest this week, a values '
        'decision is a marketing asset.',

        'Here\'s the reality of where GISCO stands: You have an extraordinary reputation among the '
        'people who know you, an institutional client base that most contractors would retire on as a '
        'reference list, and a fourth-generation credential that almost no one in Louisville\'s trades '
        'can match. What you don\'t have is a system. You have goodwill that\'s trapped in a network '
        'instead of compounding in a pipeline. Every month the business is inconsistent is another '
        'month that goodwill is not working for you.',

        'The competitive landscape works in your favor, and it\'s unusual. Cardinal Industrial has been '
        'around since 1956 — fourteen years longer than GISCO — and they\'ve built almost no digital '
        'presence. They have 17 reviews. Seventeen. In seventy years. That\'s not because their work is '
        'bad; it\'s because they\'ve never built a system to capture what their clients think of them. '
        'MINCO is essentially invisible online. Alpha is a national chain with a local address. The '
        'Louisville insulation market is the most under-marketed, under-digitized, and under-'
        'educationally-served trade sector I\'ve seen at this scale. There is no one holding the '
        'authoritative voice here — and that creates a very specific kind of opportunity: the kind '
        'where the first person to show up consistently and clearly wins, without needing to outspend '
        'anyone.',
    ]
    for p in paras:
        body(doc, p, size=10.5)

    sp(doc, 4)
    heading(doc, 'WHY THIS SEQUENCE', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    body(doc,
         'The plan leads with commercial/industrial authority because that\'s where your biggest tickets '
         'and your clearest proof points are. Peerless Distillery and the University of Louisville are '
         'names people recognize. That recognition is our credibility lever for attracting larger '
         'commercial projects in Q1. We layer in the thermography differentiator in Q2 because it opens '
         'conversations in both commercial and residential settings. Residential retrofit comes in Q3 '
         'because by then the content engine is running and homeowners are finding you organically. '
         'Removal and replacement comes in Q4 because the authority built in Q1–Q3 makes the '
         'remediation conversation a natural close, not a cold pitch.',
         size=10.5)

    sp(doc, 4)
    heading(doc, 'WHAT I NEED FROM YOU', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    asks = [
        ('Your past client list',
         'Even a rough one. Names, phone numbers, email addresses if you have them. That list is the '
         'engine for the Fast Cash Sprint and the review collection system.'),
        ('Your time on camera — four to six sessions',
         'This is the single highest-leverage investment you\'ll make in this entire engagement. The '
         'HeyGen AI avatar that comes out of those sessions generates content at scale without you '
         'doing anything after the initial recording.'),
        ('48-hour turnaround on content approvals',
         'When we send you something to review, we need a quick yes/no so momentum doesn\'t stall. '
         'That\'s the only ongoing commitment we need beyond the filming sessions.'),
    ]
    for label, text in asks:
        callout_box(doc, label + ':', text, bg=ASG_TINT, label_color=DARK_COLOR)

    sp(doc, 4)
    quote_block(doc,
                'This is going to work. The market is wide open, the credentials are real, and the '
                'story is genuinely worth telling.')
    sp(doc, 4)
    attribution(doc, 'Roger, Director | Authority Systems Group™')
    add_pb(doc)


# ── Section 3: Market Intelligence ────────────────────────────────────────────

def build_s3_market_intelligence(doc):
    section_header(doc, '3', 'Market Intelligence')
    attribution(doc, 'Daniel Frost, Chief Strategy Officer — Authority Systems Group™')
    divider(doc)
    sp(doc, 8)

    # 3A Competitive Landscape
    heading(doc, '3A: THE COMPETITIVE LANDSCAPE', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)

    competitors = [
        {
            'name': 'Cardinal Industrial Insulation Co.',
            'founded': 'Founded 1956 — 70 years in operation',
            'desc': (
                'Cardinal positions on tenure and scale — they lead with their 45,000 square feet of '
                'combined office and warehouse space and their commercial and industrial specialization. '
                'Their capabilities are real: they handle process temperatures from -200°F to 2,000°F, '
                'offer asbestos abatement, removable blankets, and firestop. They are the most capable '
                'direct competitor to GISCO on large industrial bids.'
            ),
            'well': (
                'Infrastructure, crew depth, and industrial history to compete for large-scale jobs. '
                'Technical range is legitimate.'
            ),
            'weakness': (
                'Cardinal Industrial has 17 reviews on Birdeye. Seventeen reviews in seventy years. '
                'Their website (cardinsul.com) is a static brochure — no blog, no video, no educational '
                'content. They produce no content that would appear in a homeowner or facilities '
                'manager\'s organic search. Not BBB accredited. No residential capability. No '
                'thermography offering. In the modern purchase decision — where the first interaction '
                'is a search, not a phone call — Cardinal is functionally invisible.'
            ),
            'grades': [
                ('Website positioning clarity', 'Weak'),
                ('Review volume and velocity', 'Weak'),
                ('Educational content', 'Absent'),
                ('Local search visibility', 'Weak'),
                ('Referral network visibility', 'Absent'),
            ],
            'threat': 'HIGH on commercial/industrial bids',
        },
        {
            'name': 'MINCO, Inc.',
            'founded': 'Founded 1977 — 49 years in operation',
            'desc': (
                'Based in Louisville. Established in mechanical insulation. MINCO positions as a '
                'well-known and respected local contractor. Their website (minco-ky.com) exists but '
                'their digital presence is minimal. No confirmed review volume. No visible content '
                'marketing. No social presence of note.'
            ),
            'well': (
                'Longevity and local relationships. For the facilities managers and contractors who '
                'already know MINCO, they are a trusted name.'
            ),
            'weakness': (
                'Digital invisibility. A prospect who hasn\'t been referred to MINCO directly will '
                'not encounter them in organic search. They are not competing for the new decision-'
                'maker who is discovering contractors through research. This represents a significant '
                'capture gap.'
            ),
            'grades': [
                ('Website positioning clarity', 'Weak'),
                ('Review volume and velocity', 'Absent'),
                ('Educational content', 'Absent'),
                ('Local search visibility', 'Weak'),
                ('Referral network visibility', 'Unknown'),
            ],
            'threat': 'MEDIUM',
        },
        {
            'name': 'Alpha Insulation & Waterproofing (IBP Family)',
            'founded': 'Founded 1982 — National parent: Installed Building Products (IBP)',
            'desc': (
                'Alpha is the most digitally active of the three primary competitors. Their website '
                '(alphaiwp.com/louisville-ky) is professionally produced. They offer commercial '
                'insulation, spray foam, and waterproofing. They are backed by national marketing '
                'infrastructure. Decisions are made in Columbus, Ohio — not Louisville.'
            ),
            'well': (
                'Digital presence, professional web materials, national brand support, range of '
                'services including spray foam.'
            ),
            'weakness': (
                'Alpha is a national company — no warmth, no story, no identifiable face. Their spray '
                'foam service line now carries the liability shadow of the UK class action. They cannot '
                'credibly claim the "4th generation Louisville family" narrative. The IBP parent '
                'structure means generic, corporate positioning.'
            ),
            'grades': [
                ('Website positioning clarity', 'Moderate'),
                ('Review volume and velocity', 'Moderate'),
                ('Educational content', 'Weak'),
                ('Local search visibility', 'Moderate'),
                ('Referral network visibility', 'Weak'),
            ],
            'threat': 'HIGH on commercial jobs',
        },
    ]

    for comp in competitors:
        heading(doc, comp['name'], level=3, color=DARK_COLOR, size=11, bold=True)
        body(doc, comp['founded'], size=9, italic=True, color=ASG_CHARCOAL)
        sp(doc, 2)
        body(doc, comp['desc'], size=10)

        callout_box(doc, 'What They Do Well:', comp['well'])
        callout_box(doc, 'Most Visible Weakness:', comp['weakness'], bg='FFF4E5',
                    label_color=RGBColor(0xE6, 0x7E, 0x22))

        # Grade table
        tbl = doc.add_table(rows=len(comp['grades']) + 1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(3.5)
        tbl.columns[1].width = Inches(4.0)
        hc0 = tbl.rows[0].cells[0]
        hc1 = tbl.rows[0].cells[1]
        set_cell_bg(hc0, ASG_DARK_HEX)
        set_cell_bg(hc1, ASG_DARK_HEX)
        for ci, hdr in enumerate(['Dimension', comp['name']]):
            hp = tbl.rows[0].cells[ci].paragraphs[0]
            hp.paragraph_format.left_indent = Inches(0.1)
            hr = hp.add_run(hdr)
            hr.bold = True; hr.font.name = 'Calibri'; hr.font.size = Pt(9)
            hr.font.color.rgb = WHITE
        for i, (dim, grade) in enumerate(comp['grades']):
            bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
            dc = tbl.rows[i + 1].cells[0]
            gc = tbl.rows[i + 1].cells[1]
            set_cell_bg(dc, bg); set_cell_bg(gc, bg)
            dp = dc.paragraphs[0]
            dp.paragraph_format.left_indent = Inches(0.1)
            dr = dp.add_run(dim)
            dr.font.name = 'Calibri'; dr.font.size = Pt(9.5)
            dr.font.color.rgb = BODY_COLOR
            gp = gc.paragraphs[0]
            gp.paragraph_format.left_indent = Inches(0.1)
            gr = gp.add_run(grade)
            gr.bold = True; gr.font.name = 'Calibri'; gr.font.size = Pt(9.5)
            grade_color = RGBColor(0x7D, 0x24, 0x00) if grade == 'Weak' else (
                RGBColor(0x27, 0xAE, 0x60) if grade == 'Moderate' else ASG_CHARCOAL)
            gr.font.color.rgb = grade_color
        sp(doc, 4)

        tbl2 = doc.add_table(rows=1, cols=1)
        remove_table_borders(tbl2)
        tbl2.columns[0].width = Inches(7.5)
        tc = tbl2.rows[0].cells[0]
        set_cell_bg(tc, 'FFF0F0')
        tp = tc.paragraphs[0]
        tp.paragraph_format.left_indent = Inches(0.12)
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after  = Pt(4)
        tlr = tp.add_run('THREAT LEVEL:  ')
        tlr.bold = True; tlr.font.name = 'Calibri'; tlr.font.size = Pt(9)
        tlr.font.color.rgb = RGBColor(0x7D, 0x24, 0x00)
        tvr = tp.add_run(comp['threat'])
        tvr.font.name = 'Calibri'; tvr.font.size = Pt(9.5)
        tvr.font.color.rgb = BODY_COLOR
        sp(doc, 8)

    # Gap Summary
    heading(doc, 'THE COMPETITIVE GAP SUMMARY', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    quote_block(doc,
                'Of the three primary competitors in the Louisville insulation market, not one currently '
                'occupies the position of trusted, educational, family-rooted insulation authority for '
                'either residential homeowners or institutional facilities managers. The position of '
                'Louisville\'s 4th generation insulation authority — the contractor who explains your '
                'options honestly, documents with thermography, and doesn\'t install products they '
                'don\'t trust — is genuinely unclaimed. GISCO is the only company in this market that '
                'can own it credibly.',
                size=10)

    sp(doc, 6)

    # 3B Market Opportunity
    heading(doc, '3B: MARKET OPPORTUNITY', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)

    body(doc,
         'The Louisville metro area encompasses approximately 850,000 residents and one of the most '
         'active industrial economies in the mid-South — distilleries, food processing, automotive '
         'supply, pharmaceutical, logistics, and university infrastructure.',
         size=10.5)

    market_data = [
        ('Industrial/Commercial Market',
         '300–500 insulation-eligible projects per year. Mechanical insulation replacements, new '
         'construction packages, and energy efficiency upgrades represent approximately $15–25 million '
         'annually in the Louisville metro.'),
        ('Residential Market',
         'Approximately 250,000 owner-occupied housing units. Based on national renovation data, an '
         'estimated 8–12% of homeowners undertake insulation upgrades in any given year — 20,000–30,000 '
         'households. At GISCO\'s average ticket of $2,800, the addressable residential market '
         'represents approximately $56–84 million annually across all contractors.'),
        ('GISCO\'s Current Market Share',
         'Estimated $300,000–$900,000 of this combined market — well under 1%. The gap is not a '
         'product quality issue or a pricing issue. It is a visibility and pipeline issue.'),
    ]
    for lbl, txt in market_data:
        callout_box(doc, lbl + ':', txt, bg=ASG_TINT, label_color=DARK_COLOR)

    sp(doc, 4)
    heading(doc, 'Revenue Opportunity Sizing', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    rev_rows = [
        ('Current Baseline', '~$500,000', '~$42,000', 'Inconsistent pipeline'),
        ('Conservative (+40%)', '$700,000', '$58,000', 'Past client reactivation + referral network'),
        ('Target (+80%)', '$900,000', '$75,000', 'Video content engine + referral partners fully active'),
        ('Stretch (+140%)', '$1,200,000', '$100,000', 'Commercial authority + video driving institutional inquiries'),
    ]
    four_col_kpi_table(doc, rev_rows,
                       ['Scenario', 'Annual Revenue', 'Monthly Revenue', 'Key Driver'])

    body(doc,
         'Directional estimates based on niche benchmarks, not guarantees. Actual results depend on '
         'execution quality, consistency, and the specific commercial project pipeline.',
         size=9, italic=True, color=ASG_CHARCOAL)
    sp(doc, 4)

    # Timing Argument
    heading(doc, 'The Timing Argument', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    timing = [
        ('The spray foam liability moment',
         'The UK class action has homeowners and facilities managers actively searching for alternatives. '
         'The contractor who explains why they don\'t offer spray foam — and what they recommend instead '
         '— will capture this search traffic while it is in peak awareness.'),
        ('Competitor digital vacuum',
         'Cardinal Industrial, MINCO, and most smaller Louisville insulation contractors have no '
         'educational content, no video, and no systematic review generation. The organic search '
         'landscape for "insulation contractor Louisville" is sparsely contested.'),
        ('Infrastructure timing',
         'The new website and AI phone employee are already in motion. This Blueprint layers content '
         'strategy, video, and referral activation onto that infrastructure while it\'s being built — '
         'not as an afterthought.'),
    ]
    for lbl, txt in timing:
        callout_box(doc, lbl + ':', txt, bg=ASG_TINT, label_color=DARK_COLOR)

    sp(doc, 6)

    # 3C Avatars
    heading(doc, '3C: THE IDEAL CLIENT PROFILES', level=2, color=ASG_BLUE, size=13, bold=True)
    attribution(doc, 'Dr. Raymond Cross, Board Advisor | Buyer Psychology')
    sp(doc, 4)

    avatars = [
        {
            'name': '"Facilities Manager Frank"',
            'subtitle': 'The Commercial / Industrial Decision-Maker',
            'color': SLATE_HEX,
            'setup': (
                'Frank manages facilities for a mid-size Louisville operation — a distillery, food '
                'processing plant, university building department, or large commercial property. He\'s '
                'been in the role 5–15 years. He has a maintenance budget, an aging infrastructure list, '
                'and a standing relationship with two or three contractors. He is not actively looking '
                'for an insulation contractor. He is managing.'
            ),
            'happening': (
                'Something forces the issue. A utility audit flags energy loss on a process line. An '
                'equipment failure traces back to failed pipe insulation. Or the contract with the '
                'current vendor expires and Frank has to rebid. His first move is to call the contractor '
                'he\'s used before. If that relationship doesn\'t exist — he searches, asks his network, '
                'and evaluates whoever shows up as credible.'
            ),
            'fears': [
                '"I\'m worried about getting a contractor who\'s good on small jobs but falls apart on a complex industrial scope."',
                '"What if the insulation work causes a process interruption we didn\'t account for?"',
                '"If something goes wrong on this project, it\'s my name on the budget approval."',
            ],
            'wants': [
                '"I just want a contractor who\'s done this kind of work before and can prove it."',
                '"I need someone who understands industrial scheduling — work has to happen in a specific window."',
                '"More than anything, I need to know they\'ll be there when they said they\'ll be there."',
            ],
            'unlock': (
                'Proof that GISCO has done exactly what Frank needs — specifically, at recognizable '
                'Louisville institutions. When Frank searches and finds a contractor who lists Peerless '
                'Distillery, the University of Louisville, and Jeff Boat as completed projects, with '
                'videos explaining the work and a team that has operated since 1985 — the friction '
                'of switching disappears. The trust is pre-built.'
            ),
            'closes': 'Named client references (Peerless, UofL, Jeff Boat) + 41-year history + thermography capability',
        },
        {
            'name': '"High-Bill Harold"',
            'subtitle': 'The Residential Homeowner — Energy Cost Trigger',
            'color': ASG_DARK_HEX,
            'setup': (
                'Harold owns a home in the Louisville metro — a house built in the 1970s–1990s when '
                'insulation standards were significantly lower than current code. He works full-time, '
                'pays attention to his monthly expenses, and has been watching his energy bills climb '
                'for the last two summers. He\'s not in a crisis. He\'s mildly frustrated and mildly '
                'motivated.'
            ),
            'happening': (
                'His HVAC tech tells him the unit is fine. His neighbor mentions they got their attic '
                'redone and cut their bill by $180/month. His tax preparer says there\'s a federal credit '
                'he might qualify for. One of these three events finally gets him to search. He requests '
                'two or three quotes. He has no framework for evaluating what he\'s being sold.'
            ),
            'fears': [
                '"I\'m worried I\'ll spend $5,000 and not see any difference in my bill."',
                '"What if they push me into spray foam and there\'s a problem with it later? I\'ve seen something about lawsuits."',
                '"I don\'t really know what I need — I\'m afraid of getting sold something I don\'t need."',
            ],
            'wants': [
                '"I just want someone to tell me honestly what I have, what I need, and what it\'ll cost."',
                '"I want someone who\'ll crawl through the attic instead of just looking from the ladder."',
                '"I need to know what the tax credit covers and how much I\'ll actually save."',
            ],
            'unlock': (
                'Tony showing up and crawling through the attic. That single behavior — documented in '
                'review after review — is the conversion event. Every competitor sends someone who looks '
                'from the ladder. Tony crawls through. That moment of differentiation, witnessed and '
                'experienced directly, is what converts Harold from a price-shopper into a loyal client '
                'who writes a five-star review and tells his neighbor.'
            ),
            'closes': 'Tony\'s personal inspection + honest "here\'s what you don\'t need" approach',
        },
    ]

    for av in avatars:
        tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(7.5)
        hc = tbl.rows[0].cells[0]
        set_cell_bg(hc, av['color'])
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.15)
        hp.paragraph_format.space_before = Pt(6)
        hp.paragraph_format.space_after  = Pt(2)
        hr = hp.add_run(av['name'])
        hr.bold = True; hr.font.name = 'Calibri'; hr.font.size = Pt(16)
        hr.font.color.rgb = WHITE
        hp2 = hc.add_paragraph()
        hp2.paragraph_format.left_indent = Inches(0.15)
        hp2.paragraph_format.space_before = Pt(0)
        hp2.paragraph_format.space_after  = Pt(6)
        hr2 = hp2.add_run(av['subtitle'])
        hr2.italic = True; hr2.font.name = 'Calibri'; hr2.font.size = Pt(10)
        hr2.font.color.rgb = WHITE
        sp(doc, 2)

        body(doc, 'THE SETUP:', size=9, bold=True, color=DARK_COLOR)
        body(doc, av['setup'], size=10)
        body(doc, 'WHAT\'S HAPPENING TO THEM:', size=9, bold=True, color=DARK_COLOR)
        body(doc, av['happening'], size=10)

        body(doc, 'WHAT THEY\'RE AFRAID OF:', size=9, bold=True, color=DARK_COLOR)
        for f in av['fears']:
            bullet(doc, f, size=9.5)
        sp(doc, 2)

        body(doc, 'WHAT THEY ACTUALLY WANT:', size=9, bold=True, color=DARK_COLOR)
        for w in av['wants']:
            bullet(doc, w, size=9.5)
        sp(doc, 2)

        body(doc, 'WHAT UNLOCKS THEM:', size=9, bold=True, color=DARK_COLOR)
        body(doc, av['unlock'], size=10)
        callout_box(doc, 'What closes in GISCO\'s favor:', av['closes'])
        sp(doc, 6)

    # 3D Authority Position
    heading(doc, '3D: THE AUTHORITY POSITION', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)
    body(doc,
         'In the Louisville insulation market, there is no provider who credibly occupies the position '
         'of the 4th generation, institutionally proven insulation authority who explains your options '
         'honestly, documents with thermography, and doesn\'t install products they don\'t trust. Every '
         'major competitor positions on tenure or scale or geography. Not one positions on values, '
         'transparency, and educational authority — the factors that actually drive hiring decisions '
         'when there is any doubt about who to trust.',
         size=10.5)
    sp(doc, 4)

    # Authority Statement
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(7.5)
    ac = tbl.rows[0].cells[0]
    set_cell_bg(ac, ASG_DARK_HEX)
    ap = ac.paragraphs[0]
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(12)
    ap.paragraph_format.space_after  = Pt(12)
    ap.paragraph_format.left_indent  = Inches(0.3)
    ap.paragraph_format.right_indent = Inches(0.3)
    ar = ap.add_run(
        'GISCO Corp is Louisville\'s 4th generation insulation authority — the contractor of choice '
        'for facilities managers at Kentucky\'s most recognized institutions and homeowners who want '
        'an honest answer before they get a quote — built on 41 years of doing the work right without '
        'cutting corners or installing products they don\'t trust.'
    )
    ar.bold = True; ar.italic = True; ar.font.name = 'Calibri'; ar.font.size = Pt(11)
    ar.font.color.rgb = WHITE
    sp(doc, 6)

    body(doc,
         'This statement passes all three criteria: (1) No competitor can claim it tomorrow — it '
         'requires 41 years of history and specific named institutional clients. (2) It resonates with '
         'both Facilities Manager Frank ("most recognized institutions") and High-Bill Harold ("honest '
         'answer before they get a quote"). (3) GISCO can deliver on every word operationally.',
         size=9.5, italic=True, color=ASG_CHARCOAL)
    sp(doc, 4)
    attribution(doc, 'Daniel Frost, CSO | Authority Systems Group™')
    add_pb(doc)


# ── Section 4: Belief-to-Buy Map ──────────────────────────────────────────────

def build_s4_belief_to_buy(doc):
    section_header(doc, '4', 'Belief-to-Buy Framework™ Map')
    attribution(doc, 'Dr. Raymond Cross, Board Advisor (Psychology) + Vivienne Carr, CMO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    body(doc,
         'The Belief-to-Buy Framework™ maps the dual-track journey every prospect travels before '
         'making a hiring decision. Both tracks — Belief and Emotion — must reach their final stage '
         'before a buying decision occurs. Moving only one track stalls the prospect.',
         size=10.5)
    sp(doc, 4)

    def btb_track(title, rows, stall, unlock):
        heading(doc, title, level=3, color=DARK_COLOR, size=12, bold=True)
        sp(doc, 2)

        num_rows = len(rows) + 1
        tbl = doc.add_table(rows=num_rows, cols=3)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(0.7)
        tbl.columns[1].width = Inches(3.4)
        tbl.columns[2].width = Inches(3.4)
        for ci, hdr in enumerate(['Stage', 'Belief Track', 'Emotion Track']):
            hc = tbl.rows[0].cells[ci]
            set_cell_bg(hc, ASG_DARK_HEX)
            hp = hc.paragraphs[0]
            hp.paragraph_format.left_indent = Inches(0.08)
            hr = hp.add_run(hdr)
            hr.bold = True; hr.font.name = 'Calibri'; hr.font.size = Pt(9)
            hr.font.color.rgb = WHITE
        for i, (stage, belief, emotion) in enumerate(rows):
            bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
            for ci, txt in enumerate([stage, belief, emotion]):
                cell = tbl.rows[i + 1].cells[ci]
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.left_indent = Inches(0.08)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                r = p.add_run(txt)
                r.font.name = 'Calibri'
                r.font.size = Pt(8.5)
                r.bold = (ci == 0)
                r.font.color.rgb = ASG_BLUE if ci == 0 else BODY_COLOR
        sp(doc, 4)

        callout_box(doc, 'PRIMARY STALL POINT:', stall, bg='FFF4E5',
                    label_color=RGBColor(0xE6, 0x7E, 0x22))
        callout_box(doc, 'CONTENT THAT UNLOCKS THE STALL:', unlock, bg=ASG_TINT,
                    label_color=DARK_COLOR)
        sp(doc, 6)

    btb_track(
        'COMMERCIAL / INDUSTRIAL CLIENT TRACK',
        [
            ('B1/E1',
             'Enemy Belief: "All insulation contractors submit bids. We evaluate on price and scope. There\'s no meaningful difference between qualified vendors."',
             'Dissatisfaction: "Our current insulation is aging but still functional. The maintenance budget has higher priorities."'),
            ('B2/E2',
             'Pre-Frame: "A contractor\'s institutional track record — specifically whether they\'ve worked in our type of facility — is the real differentiator."',
             'Contrast: "Our peer at a similar facility had a process failure that traced to deteriorated pipe insulation. That\'s an event we can\'t afford."'),
            ('B3/E3',
             'New Worldview: "The right insulation contractor is a risk management decision, not a commodity purchase. GISCO has done this at Peerless, UofL, and Jeff Boat."',
             'Desire: "I want a vendor relationship I can rely on for the next contract cycle, not a one-bid relationship I have to manage."'),
            ('B4/E4',
             'Internal Alignment: "I can defend this decision. GISCO\'s 41-year history, institutional references, and A+ BBB rating make the case."',
             'Urgency: "If I don\'t lock in a reliable insulation contractor before the next maintenance window, we\'ll scramble for someone unvetted."'),
            ('B5/E5',
             'Certainty: "GISCO is the right contractor. The track record, the thermography capability, and the communication standard are all there."',
             'Relief: "I don\'t have to worry about this. I\'ve got a contractor I can call."'),
        ],
        'B1/E1. The facilities manager\'s default is the contractor already in their Rolodex. GISCO must interrupt the assumption that all qualified vendors are interchangeable — and the fastest way to do that is making the Peerless/UofL track record visible in organic search BEFORE the bid goes out.',
        'Case study content showing GISCO on recognizable institutional jobs + thermography diagnostic content showing what a facilities-grade inspection looks like.'
    )

    btb_track(
        'RESIDENTIAL CLIENT TRACK',
        [
            ('B1/E1',
             'Enemy Belief: "Insulation quotes are basically the same thing. Get three prices and go with the middle one. The only difference is material — and it\'s all the same material."',
             'Dissatisfaction: "Our energy bills are high in summer. That\'s just how it is in Louisville."'),
            ('B2/E2',
             'Pre-Frame: "The quality of the inspection — not just the material — determines what you actually need. A contractor who crawls through your attic will see problems the one who looked from the ladder will miss."',
             'Contrast: "My neighbor had two contractors quote his attic. One looked from the pull-down stairs. GISCO crawled through and found an unsealed chase the other guy missed completely."'),
            ('B3/E3',
             'New Worldview: "Air sealing matters as much as R-value. The contractor who seals before they insulate delivers better performance. Not every contractor does both — and fewer will tell you that upfront."',
             'Desire: "I want the person who will tell me honestly what I need — not what generates the biggest ticket."'),
            ('B4/E4',
             'Internal Alignment: "The 25C tax credit stacks with the LG&E rebate. The net cost after credits is significantly less than the quoted price. This is the right time to do it."',
             'Urgency: "Summer is coming. If I wait another season, that\'s another $200/month in wasted energy."'),
            ('B5/E5',
             'Certainty: "Tony Long at GISCO. He crawled through our attic when two others didn\'t bother. He told me what I didn\'t need. He\'s going to do the job right."',
             'Relief: "I\'m not getting ripped off. I\'m getting the honest contractor."'),
        ],
        'B1/E2. The homeowner believes insulation is a commodity (B1) AND has normalized the pain (E1). Both must be disrupted simultaneously. Tony\'s personal inspection behavior disrupts both tracks at once — it demonstrates that the approach matters (B-track) and creates a contrast with every competitor\'s standard behavior (E-track).',
        '"The contractor who actually crawled through my attic" social proof content + "What air sealing does that insulation alone can\'t" educational content.'
    )

    add_pb(doc)


# ── Section 5: Priority Services Portfolio ────────────────────────────────────

def build_s5_services(doc):
    section_header(doc, '5', 'Priority Services Portfolio')
    attribution(doc, 'Vivienne Carr, CMO + Tanya Blackwood, CRO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    services = [
        {
            'rank': '01',
            'name': 'Commercial / Industrial Mechanical Insulation',
            'focus': 'Q1 FOCUS',
            'color': ASG_DARK_HEX,
            'why': (
                'Highest average ticket in the GISCO portfolio — $25,000–$150,000+ per industrial '
                'project — with the clearest authority differentiators. Named client references make '
                'the sales conversation faster and the close rate higher than any other service. One '
                'commercial project at full margin covers what would take 10–15 residential jobs.'
            ),
            'authority': (
                'GISCO is the 4th generation Louisville contractor with a documented track record at '
                'Kentucky\'s most recognizable institutions. While Cardinal Industrial has existed '
                'longer, they\'ve never made that history visible. GISCO\'s institutional references — '
                'available to name and document — are a credential that no competitor can manufacture, '
                'buy, or replicate.'
            ),
            'revenue': [
                ('Current Baseline', '~$200,000–$350,000'),
                ('Conservative (+50%)', '$300,000–$525,000'),
                ('Target (+100%)', '$400,000–$700,000'),
                ('Stretch (+150%)', '$500,000–$875,000'),
            ],
        },
        {
            'rank': '02',
            'name': 'Infrared Thermography & Energy Diagnostics',
            'focus': 'Q2 FOCUS',
            'color': SLATE_HEX,
            'why': (
                'No competitor in the Louisville market markets infrared thermography as a diagnostic '
                'service. It is completely unclaimed. Thermography is both a standalone revenue service '
                '($500–$2,500 per inspection) and a multiplier on every insulation job it identifies — '
                'it creates diagnostic-to-work pipelines that close at dramatically higher rates than '
                'cold estimates because the problem is already documented and visible.'
            ),
            'authority': (
                'The only insulation contractor in Louisville who shows you the problem before '
                'recommending a solution. Thermal imaging doesn\'t lie — it shows exactly where energy '
                'is being lost and where insulation is failing. This positions GISCO as diagnostician, '
                'not salesperson. The contrast with competitors who "walk in and quote" is stark and '
                'immediately visible to both commercial and residential clients.'
            ),
            'revenue': [
                ('Current Baseline', '~$20,000–$40,000 (likely untracked)'),
                ('Conservative', '$60,000–$100,000'),
                ('Target', '$100,000–$175,000'),
                ('Stretch', '$150,000–$250,000'),
            ],
        },
        {
            'rank': '03',
            'name': 'Residential Retrofit — Attic Air Sealing + Blown-In Cellulose',
            'focus': 'Q3 FOCUS',
            'color': RGBColor(0x27, 0x6B, 0x3D) if False else '276B3D',
            'color': '276B3D',
            'why': (
                'High conversion rate when Tony shows up (reviews confirm near-100% close when Tony '
                'personally inspects), IRA 25C tax credit + LG&E utility rebate creates a compelling '
                'ROI story, and the "no spray foam" position is a genuine differentiator for the '
                'homeowner now asking the spray foam question. Cellulose is a legitimate, non-liability '
                'alternative with strong energy performance.'
            ),
            'authority': (
                'Tony\'s personal inspection standard — the only Louisville insulation contractor who '
                'physically goes through the attic — combined with the "no spray foam, here\'s why and '
                'what we recommend instead" narrative positions GISCO as the transparent alternative in '
                'a market where everyone else is either pushing spray foam or providing commodity '
                'blown-in without explanation.'
            ),
            'revenue': [
                ('Current Baseline', '~$80,000–$150,000'),
                ('Conservative (+40%)', '$112,000–$210,000'),
                ('Target (+80%)', '$144,000–$270,000'),
                ('Stretch (+120%)', '$176,000–$330,000'),
            ],
        },
        {
            'rank': '04',
            'name': 'Insulation Removal & Replacement',
            'focus': 'Q4 FOCUS',
            'color': AMBER_HEX,
            'why': (
                'Remediation framing commands premium pricing. The buyer is already in a problem state '
                '— mold evidence, smoke/fire damage, rodent contamination, or aged failing insulation '
                'that failed a pre-sale inspection. No competitive bidding environment. Close rates are '
                'high and margin is strong because urgency eliminates price shopping.'
            ),
            'authority': (
                'Tony\'s honest assessment approach — "here\'s what needs to come out, here\'s why, '
                'here\'s what goes back in" — positions GISCO as the trusted remediation advisor rather '
                'than the contractor trying to maximize the ticket. That honesty, documented in existing '
                'reviews, closes these jobs without needing to compete on price.'
            ),
            'revenue': [
                ('Current Baseline', '~$30,000–$60,000'),
                ('Conservative', '$50,000–$90,000'),
                ('Target', '$80,000–$140,000'),
                ('Stretch', '$110,000–$200,000'),
            ],
        },
    ]

    for svc in services:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(0.6)
        tbl.columns[1].width = Inches(6.9)
        set_cell_bg(tbl.rows[0].cells[0], svc['color'])
        hc = tbl.rows[0].cells[1]
        set_cell_bg(hc, ASG_DARK_HEX)
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.12)
        hp.paragraph_format.space_before = Pt(5)
        hp.paragraph_format.space_after  = Pt(1)
        hr1 = hp.add_run(f'SERVICE {svc["rank"]}  ')
        hr1.bold = True; hr1.font.name = 'Calibri'; hr1.font.size = Pt(9)
        hr1.font.color.rgb = ASG_BLUE
        hr2 = hp.add_run(f'— {svc["name"].upper()}')
        hr2.bold = True; hr2.font.name = 'Calibri'; hr2.font.size = Pt(11)
        hr2.font.color.rgb = WHITE
        hp2 = hc.add_paragraph()
        hp2.paragraph_format.left_indent = Inches(0.12)
        hp2.paragraph_format.space_before = Pt(0)
        hp2.paragraph_format.space_after  = Pt(5)
        hf = hp2.add_run(svc['focus'])
        hf.bold = True; hf.font.name = 'Calibri'; hf.font.size = Pt(8)
        hf.font.color.rgb = ASG_BLUE
        sp(doc, 2)

        callout_box(doc, 'WHY THIS SERVICE:', svc['why'], bg=ASG_TINT, label_color=DARK_COLOR)
        callout_box(doc, 'THE AUTHORITY ANGLE:', svc['authority'], bg=ASG_TINT, label_color=DARK_COLOR)

        heading(doc, 'Revenue Model', level=3, color=DARK_COLOR, size=10, bold=True)
        sp(doc, 1)
        two_col_table(doc, svc['revenue'], col_widths=(2.5, 5.0))
        sp(doc, 6)

    # Combined portfolio table
    heading(doc, 'COMBINED YEAR 1 PORTFOLIO PROJECTION', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)
    four_col_kpi_table(doc,
        [
            ('Commercial / Industrial', '$350,000', '$550,000', '$750,000'),
            ('Thermography + Follow-On', '$70,000', '$130,000', '$200,000'),
            ('Residential Retrofit', '$130,000', '$210,000', '$300,000'),
            ('Removal & Replacement', '$60,000', '$110,000', '$175,000'),
            ('TOTAL', '$610,000', '$1,000,000', '$1,425,000'),
        ],
        ['Service', 'Conservative', 'Target', 'Stretch']
    )
    body(doc, 'Directional estimates only. Assumes consistent execution across all four quarters.',
         size=9, italic=True, color=ASG_CHARCOAL)
    add_pb(doc)


# ── Section 6: Fast Cash Sprint™ ─────────────────────────────────────────────

def build_s6_fast_cash(doc):
    section_header(doc, '6', 'Fast Cash Sprint™')
    attribution(doc, 'Tanya Blackwood, CRO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    quote_block(doc,
                '"Before we build anything new, let\'s look at what\'s already there. GISCO has 41 years '
                'of client relationships, a personal reputation that generates word-of-mouth organically, '
                'and a network of HVAC contractors and building professionals who encounter insulation '
                'problems every week and currently have no formalized relationship with Tony. The Fast '
                'Cash Sprint™ is how we activate what already exists before spending a dollar on new '
                'lead generation." — Tanya Blackwood, CRO')

    strategies = [
        {
            'num': '#1',
            'title': 'The Past Client Wake-Up Campaign',
            'color': RED_HEX,
            'opportunity': (
                'GISCO has served commercial and residential clients for 41 years. Every past client '
                'who had a good experience is a candidate for additional work, a referral, and a '
                'five-star Google review. Based on a conservative estimate of 50–100 reachable past '
                'clients, and an average response rate of 15–20%, this campaign generates 8–20 engaged '
                'contacts in the first 30 days — from people who already trust Tony completely.'
            ),
            'why': (
                'Tony\'s reputation is personal. Reviews mention his name specifically. "Tony came out '
                'himself" is a phrase that appears repeatedly. A reactivation campaign that comes from '
                'Tony personally — not from "GISCO Corp" — will outperform any generic outreach because '
                'the relationship is already there.'
            ),
            'steps': [
                'Step 1: Compile past client list from any source — invoices, phone records, email history, memory. Even 40–60 contacts is enough to start.',
                'Step 2: Draft a short, personal note from Tony referencing the original job and checking in. Not a promotion — a genuine touchpoint.',
                'Step 3: Include one of three actions: (a) Google review request, (b) seasonal check-in offer, or (c) referral ask.',
                'Step 4: Deploy via email and text if available. Hand-written note for highest-value commercial contacts.',
                'Step 5: AI phone employee and GHL track all responses.',
            ],
            'revenue_c': '$25,000 (50 contacts × 15% response × avg $3,300 job)',
            'revenue_t': '$55,000 (75 contacts × 20% response × avg $3,700 job)',
            'timeline': 'First responses expected within 7–10 days of deployment.',
            'provides_asg': '3-touch email/text sequence; review request template; referral ask language',
            'provides_tony': 'Past client list (names + contact info); approval of messaging',
            'launch': '5–7 days after Tony provides client list',
            'sofia': (
                '"I ran a version of this in my business in year two — sent 60 past clients a simple '
                'personal check-in from the owner by name. We got eleven responses in the first week. '
                'Four turned into jobs. Two turned into referrals that became jobs. The list you think '
                'is \'old\' is the most valuable thing you own."'
            ),
        },
        {
            'num': '#2',
            'title': 'The HVAC Referral Network Launch',
            'color': AMBER_HEX,
            'opportunity': (
                'Every HVAC technician in Louisville is in homes and commercial buildings regularly '
                'and regularly diagnoses comfort problems that are actually insulation and air sealing '
                'failures. They are seeing GISCO\'s exact ideal clients every single week and currently '
                'have no one to refer them to. A formal relationship with 5–8 HVAC companies generates '
                'a steady stream of pre-qualified, warm referrals — people who have already been told '
                'by a trusted professional that they have an insulation problem.'
            ),
            'why': (
                'Tony\'s honesty reputation is the bridge here. HVAC contractors are protective of '
                'their client relationships — they will only refer to a contractor they trust not to '
                'embarrass them. Tony\'s track record of telling clients what they don\'t need, '
                'crawling through attics, and over-delivering is exactly what an HVAC contractor '
                'wants in a referral partner.'
            ),
            'steps': [
                'Step 1: Identify 10–15 HVAC contractors in the Louisville metro (ASG research + Tony\'s existing contacts).',
                'Step 2: Draft a professional one-page partner letter from Tony explaining the referral arrangement (ASG writes this).',
                'Step 3: Tony reaches out to 3–5 contacts he already knows personally first.',
                'Step 4: ASG produces a simple "Partner Kit" — Tony\'s bio, service overview, lead magnet, and referral script.',
                'Step 5: For each referral that becomes a job, Tony sends a personal thank-you + small gift card to the referring individual tech.',
            ],
            'revenue_c': '$40,000 in 90 days (5 active partners × 2 referrals/month × avg residential job)',
            'revenue_t': '$80,000 in 90 days (8 partners × 3 referrals/month + 1–2 commercial referrals)',
            'timeline': 'First referrals typically arrive within 14–21 days of partner activation.',
            'provides_asg': 'Partner outreach letter; HVAC Partner Kit; referral tracking template; thank-you system',
            'provides_tony': 'List of HVAC contacts already in his network; approval of partner letter; personal outreach to first 3–5 contacts',
            'launch': '7–10 days after Tony provides initial contact list',
            'sofia': (
                '"One good HVAC referral partner is worth more than 20 Google ads that month. The '
                'person calling from an HVAC tech\'s recommendation has already been told they have a '
                'problem by someone they trust. That\'s not a lead — that\'s a presold appointment."'
            ),
        },
        {
            'num': '#3',
            'title': 'The Google Review Blitz',
            'color': GREEN_HEX,
            'opportunity': (
                'GISCO\'s estimated Google review count is fewer than 20. Cardinal Industrial has 17 '
                'on Birdeye. MINCO is invisible. The local insulation search results for Louisville '
                'are a review desert. In local services, Google review count is the single highest-'
                'leverage factor in whether a prospect calls or skips. Getting to 40–50 reviews in '
                '90 days is achievable using the past client reactivation list and the new AI phone '
                'system, making GISCO the most-reviewed insulation contractor in Louisville.'
            ),
            'why': (
                'Tony\'s existing clients love him. The challenge is that nobody has ever asked them '
                'to write a review in a way that made it easy and obvious. The AI phone system changes '
                'this entirely — after every completed job, the system automatically sends a review '
                'request with a direct link. No manual process. No forgetting. Every satisfied client '
                'gets asked, every time.'
            ),
            'steps': [
                'Step 1: Configure the AI phone system to trigger a review request SMS/email after Tony marks a job complete.',
                'Step 2: Personal review request from Tony to the top 20–30 past clients on the reactivation list.',
                'Step 3: ASG provides the direct Google review link and a two-sentence prompt to make the request frictionless.',
                'Step 4: Respond to every review received within 24 hours (ASG provides 20 response templates).',
                'Step 5: Track review velocity weekly in the KPI dashboard.',
            ],
            'revenue_c': 'Reviews build trust infrastructure. Each 10 reviews added increases organic call volume by est. 15–25% from a near-zero base.',
            'revenue_t': '25–40 new reviews achievable within 90 days with 50+ contacts on the reactivation list.',
            'timeline': 'Integrated into AI phone system at launch.',
            'provides_asg': 'Review request templates; AI phone system integration spec; 20 pre-written review response templates',
            'provides_tony': 'Approval of request templates; active response to any negative or complex reviews',
            'launch': 'Integrated into AI phone system at launch — Day 1',
            'sofia': (
                '"The fastest thing you can do for organic search visibility in a local market is '
                'reviews. Not content, not backlinks — reviews. At 50+ Google reviews, GISCO becomes '
                'the most credible insulation contractor in Louisville by a wide margin."'
            ),
        },
    ]

    for s in strategies:
        # Strategy header
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(0.6)
        tbl.columns[1].width = Inches(6.9)
        set_cell_bg(tbl.rows[0].cells[0], s['color'])
        hc = tbl.rows[0].cells[1]
        set_cell_bg(hc, ASG_DARK_HEX)
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.12)
        hp.paragraph_format.space_before = Pt(5)
        hp.paragraph_format.space_after  = Pt(1)
        hr1 = hp.add_run(f'FAST CASH STRATEGY {s["num"]}  ')
        hr1.bold = True; hr1.font.name = 'Calibri'; hr1.font.size = Pt(9)
        hr1.font.color.rgb = ASG_BLUE
        hr2 = hp.add_run(f'— {s["title"].upper()}')
        hr2.bold = True; hr2.font.name = 'Calibri'; hr2.font.size = Pt(11)
        hr2.font.color.rgb = WHITE
        hp2 = hc.add_paragraph()
        hp2.paragraph_format.left_indent = Inches(0.12)
        hp2.paragraph_format.space_before = Pt(0)
        hp2.paragraph_format.space_after  = Pt(5)
        sp(doc, 2)

        callout_box(doc, 'THE OPPORTUNITY:', s['opportunity'], bg=ASG_TINT, label_color=DARK_COLOR)
        callout_box(doc, 'WHY THIS WORKS FOR GISCO:', s['why'], bg=ASG_TINT, label_color=DARK_COLOR)

        body(doc, 'THE APPROACH', size=10, bold=True, color=DARK_COLOR)
        for step in s['steps']:
            bullet(doc, step, size=9.5)
        sp(doc, 4)

        rev_tbl = doc.add_table(rows=3, cols=2)
        remove_table_borders(rev_tbl)
        rev_tbl.columns[0].width = Inches(1.5)
        rev_tbl.columns[1].width = Inches(6.0)
        for ri, (lbl, val) in enumerate([
            ('Conservative:', s['revenue_c']),
            ('Target:', s['revenue_t']),
            ('Timeline:', s['timeline']),
        ]):
            bg = ASG_TINT if ri % 2 == 0 else 'FFFFFF'
            set_cell_bg(rev_tbl.rows[ri].cells[0], bg)
            set_cell_bg(rev_tbl.rows[ri].cells[1], bg)
            lp = rev_tbl.rows[ri].cells[0].paragraphs[0]
            lp.paragraph_format.left_indent = Inches(0.08)
            lr = lp.add_run(lbl)
            lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(9)
            lr.font.color.rgb = ASG_CHARCOAL
            vp = rev_tbl.rows[ri].cells[1].paragraphs[0]
            vr = vp.add_run(val)
            vr.font.name = 'Calibri'; vr.font.size = Pt(9.5)
            vr.font.color.rgb = BODY_COLOR
        sp(doc, 4)

        impl_data = [
            ('ASG provides:', s['provides_asg']),
            ('Tony provides:', s['provides_tony']),
            ('Timeline to launch:', s['launch']),
        ]
        two_col_table(doc, impl_data, col_widths=(1.5, 6.0))

        quote_block(doc, f'SOFIA VEGA\'S TAKE: {s["sofia"]}', size=9.5)
        sp(doc, 8)

    # 90-day calendar
    heading(doc, '90-DAY ACTIVATION CALENDAR', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    cal_phases = [
        ('WEEK 1–2: LAUNCH PREPARATION', [
            'Strategy 1: Tony compiles past client list (target: 50–80 contacts); ASG drafts reactivation sequence; Tony approves messaging within 48 hours',
            'Strategy 2: Tony identifies 5 HVAC contacts; ASG drafts partner outreach letter; Tony begins personal outreach to top 3 contacts',
            'Strategy 3: AI phone employee configured with post-job review request automation; Tony sends personal review request to top 10 past clients',
            'Tony provides by Day 7: Past client list + HVAC contacts + Google Business Profile login',
            'ASG delivers by Day 7: Reactivation sequence drafts + HVAC partner letter + review request templates + GHL pipeline structure',
        ]),
        ('WEEK 3–4: CAMPAIGN LAUNCH', [
            'Strategy 1: First reactivation touch deploys (email + text to full list)',
            'Strategy 2: Tony makes first 3 HVAC partner conversations; Partner Kit delivered',
            'Strategy 3: AI phone review automation live; personal requests sent to top 20 past clients',
            'Monitor: All responses tracked in GHL; Tony notified same day for any commercial opportunities',
        ]),
        ('WEEK 5–8: ACTIVE CAMPAIGN MANAGEMENT', [
            'Strategy 1: Second reactivation touch deploys (Day 21 — check-in + seasonal offer)',
            'Strategy 2: HVAC partner outreach expands to 5 additional contacts',
            'Strategy 3: Review count check — target 15+ new reviews by end of Week 6',
            'Video: Tony records first session (60–90 minutes on camera — ASG directs)',
        ]),
        ('WEEK 9–12: OPTIMIZATION + LOCK IN', [
            'Full performance review: all 3 strategies vs. conservative/target scenarios',
            'Strategy 2: HVAC partner network at 5+ active relationships; first referral jobs closed',
            'Strategy 3: Target 30+ Google reviews by Day 90',
            'Video: Tony records second session; first HeyGen avatar scripts in production',
        ]),
    ]

    for phase_title, tasks in cal_phases:
        p = doc.add_paragraph()
        set_para_bg(p, SLATE_HEX)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.12)
        r = p.add_run(phase_title)
        r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
        for task in tasks:
            bullet(doc, '☐  ' + task, size=9.5)
        sp(doc, 4)

    sp(doc, 4)
    heading(doc, 'TOTAL ESTIMATED 90-DAY REVENUE IMPACT', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    two_col_table(doc, [
        ('Conservative:', '$65,000 (Strategy 1: $25K + Strategy 2: $40K + Strategy 3: pipeline foundation)'),
        ('Target:', '$135,000 (Strategy 1: $55K + Strategy 2: $80K + Strategy 3: pipeline acceleration)'),
    ], col_widths=(1.5, 6.0))

    sp(doc, 4)
    quote_block(doc,
                'These three strategies were selected because they use assets GISCO already owns — '
                'a client relationship history, Tony\'s personal network, and a reputation that earns '
                'five-star reviews when anyone bothers to ask for them. None of them require a new '
                'website, a paid ad budget, or a content library that hasn\'t been written yet. They '
                'run on relationships and a working phone number. The sprint isn\'t a one-time event — '
                'it\'s the ignition for the systems that run the business going forward.')
    attribution(doc, 'Tanya Blackwood, CRO | Authority Systems Group™')
    add_pb(doc)


# ── Section 7: Year 1 Strategic Plan ─────────────────────────────────────────

def build_s7_year1_plan(doc):
    section_header(doc, '7', 'Year 1 Strategic Plan')
    attribution(doc, 'Daniel Frost, CSO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    quote_block(doc,
                '"The Year 1 plan below was built from a clear-eyed analysis of GISCO\'s market '
                'position, current assets, and the specific belief barriers standing between ideal '
                'clients and the front door. It is not a general marketing plan with your name on it. '
                'Every initiative was selected because it\'s the right move for this business at this '
                'moment." — Daniel Frost, CSO')

    body(doc,
         'THE THESIS: The single most important strategic objective for GISCO in Year 1 is converting '
         'Tony Long\'s personal reputation into a visible, searchable, scalable digital authority — '
         'one that generates qualified inbound leads without requiring Tony to be available every time '
         'a prospect is curious. Right now, GISCO\'s reputation exists entirely in the heads of people '
         'who\'ve already worked with Tony. The goal is to make that reputation visible to the people '
         'who haven\'t found him yet.',
         size=10.5)
    sp(doc, 4)

    quarters = [
        {
            'label': 'QUARTER 1 — FOUNDATION + COMMERCIAL AUTHORITY',
            'months': 'Months 1–3',
            'focus': 'Commercial / Industrial Mechanical Insulation',
            'goal': 'Establish GISCO as Louisville\'s institutional insulation authority; activate Fast Cash Sprint revenue; launch digital infrastructure',
            'color': ASG_DARK_HEX,
            'months_detail': [
                {
                    'month': 'Month 1 — System Launch',
                    'items': [
                        'Fast Cash Sprint activation: past client reactivation deploys, HVAC partner outreach begins, review blitz launches',
                        'New website launches: Home, About Tony, Services (4 pages), Thermography, Blog (5 articles), Contact/Free Assessment',
                        'Google Business Profile full optimization: all service categories, 10+ job photos, seeded Q&A',
                        'GHL CRM configured: lead intake pipeline, post-job review automation, HVAC partner referral tracking',
                        'AI phone employee: live and routing calls; all transcriptions delivered to Tony daily',
                        'Blog Article #1: "GISCO and the University of Louisville: What a 41-Year Relationship Looks Like"',
                        'Blog Article #2: "Why GISCO Doesn\'t Offer Spray Foam — And What We Recommend Instead"',
                        'Video Session #1 with Tony: 60–90 minutes on camera; feeds HeyGen avatar',
                    ],
                },
                {
                    'month': 'Month 2 — Content Engine + Reputation System',
                    'items': [
                        'Blog Article #3: "How Infrared Thermography Changes What We See in Your Building Before We Quote"',
                        'Blog Article #4: "What Facilities Managers at Kentucky Distilleries Know About Pipe Insulation"',
                        'Video Session #2 with Tony; HeyGen avatar training continues',
                        'First HeyGen avatar video published: "The Most Common Insulation Mistake in Louisville Commercial Buildings"',
                        'Review collection: automated system running; target 10 new Google reviews by end of Month 2',
                        'HVAC partner network: 5 active relationships; first referrals tracked',
                        'Welcome email sequence live: all new website leads receive 5-part sequence',
                    ],
                },
                {
                    'month': 'Month 3 — Optimization + Q2 Prep',
                    'items': [
                        'Q1 KPI review: compare all metrics to 90-day targets',
                        'Blog Article #5: "The 25C Tax Credit and LG&E Rebate Stack: What Louisville Homeowners Can Claim in 2026"',
                        'Video #2 published (AI avatar): "Open Cell vs. Closed Cell vs. Blown-In: The Louisville Homeowner\'s Honest Guide"',
                        'LinkedIn: Tony\'s personal profile rebuilt with institutional references; company page refreshed',
                        'GBP: 10+ reviews achieved; Google Posts live (2/week cadence begins)',
                        'Q2 content calendar finalized; thermography campaign briefs prepared',
                    ],
                },
            ],
            'metrics': [
                '25+ Google reviews',
                'New website live with 5 articles + 4 service pages',
                'AI phone employee active and transcribing all calls',
                '2 AI avatar videos published to YouTube',
                '5+ active HVAC referral partners',
                'Past client reactivation deployed to 50+ contacts',
                'GHL pipeline active with all new leads tracked',
            ],
        },
        {
            'label': 'QUARTER 2 — THERMOGRAPHY AUTHORITY + EXPANSION',
            'months': 'Months 4–6',
            'focus': 'Infrared Thermography & Energy Diagnostics',
            'goal': 'Establish thermography as GISCO\'s diagnostic differentiator; expand referral network to home inspectors and real estate agents',
            'color': SLATE_HEX,
            'months_detail': [
                {
                    'month': 'Month 4 — Thermography Campaign Launch',
                    'items': [
                        'Blog Article #6: "What a Thermal Camera Sees in Your Building That a Standard Quote Completely Misses"',
                        'AI avatar video #3: "I Scanned This Louisville Home With a Thermal Camera — Here\'s What We Found"',
                        'Named service launch: "GISCO Thermal Diagnostic" — $350–$750 residential; quoted for commercial. Landing page live.',
                        'Home inspector outreach: 3–5 contacts; Partner Kit produced',
                        'Real estate agent outreach: 5–10 Louisville-area agents with "pre-listing thermal audit" offer',
                    ],
                },
                {
                    'month': 'Month 5 — Proof Content + Mid-Point Review',
                    'items': [
                        'Case Study #1 published: GISCO project at a recognizable Louisville institution (Peerless Distillery or UofL)',
                        'AI avatar video #4: "The Louisville Homeowner\'s Spray Foam Warning — What the UK Lawsuit Means for You"',
                        'Blog Article #7: "7 Signs Your Commercial Facility\'s Insulation Is Costing You Money"',
                        'Mid-point KPI review: leads vs. target, reviews vs. target, referral partner count vs. target',
                        'HVAC partner appreciation: personal thank-you from Tony; $50 gift card to top referring individual tech',
                    ],
                },
                {
                    'month': 'Month 6 — Mid-Year Review + Q3 Prep',
                    'items': [
                        'Mid-year strategy call with Roger: full Q1–Q2 performance review',
                        'Case Study #2 in production (residential — attic project with before/after thermal images)',
                        'AI avatar video #5: "The Tony Long Attic Inspection: Why I Crawl Through When Others Don\'t"',
                        'Q3 content calendar finalized; residential retrofit campaign briefs prepared',
                        'Review count target: 50+ Google reviews by end of Month 6',
                    ],
                },
            ],
            'metrics': [
                '50+ Google reviews',
                'Thermography landing page live; first 3 diagnostic jobs booked from organic',
                '5 AI avatar videos published to YouTube',
                'Case Study #1 published on website',
                'Home inspector + real estate agent outreach initiated (5+ each)',
                'Monthly inbound leads tracking at 15–20/month',
            ],
        },
        {
            'label': 'QUARTER 3 — RESIDENTIAL AUTHORITY + CONTENT DEPTH',
            'months': 'Months 7–9',
            'focus': 'Residential Retrofit — Attic Air Sealing + Blown-In Cellulose',
            'goal': 'Capture residential homeowner search traffic; position "no spray foam" narrative; activate seasonal pre-fall demand',
            'color': '276B3D',
            'months_detail': [
                {
                    'month': 'Month 7 — Residential Campaign Launch + Seasonal Push',
                    'items': [
                        'Blog Article #8: "Is Your Attic Ready for Louisville Winter? What We Look For Before We Quote"',
                        'AI avatar video #6: "Louisville Attic Insulation in 2026: What\'s Actually Worth the Money (And What Isn\'t)"',
                        'Email campaign to full list: "Pre-Fall Attic Efficiency Check" offer — connects to AI phone and booking form',
                    ],
                },
                {
                    'month': 'Month 8 — Proof Content + Partner Expansion',
                    'items': [
                        'Case Study #2 published (residential attic — before/after thermal images + homeowner quote)',
                        'AI avatar video #7: "Louisville Homeowner Saved $187/Month on Energy Bills — What We Found in the Attic"',
                        'Blog Article #9: "The IRA 25C Credit + LG&E Rebate: How Louisville Homeowners Stack Both in 2026"',
                        'Partner expansion: reach out to 5–10 Louisville-area mortgage brokers and buyers\' agents',
                        'Seasonal push on all social channels; Google Posts highlight "schedule before the cold" message',
                    ],
                },
                {
                    'month': 'Month 9 — Email Upsell + Q3 Performance Review',
                    'items': [
                        'Email upsell to all past commercial clients: "Annual insulation inspection — what we check and why it matters"',
                        'Q3 performance review: leads, revenue, review count, referral partner activity',
                        'AI avatar video #8 published; total YouTube library at 8+ videos',
                        'Q4 content calendar and removal/replacement campaign brief finalized',
                    ],
                },
            ],
            'metrics': [
                '65+ Google reviews',
                '8+ AI avatar videos published',
                'Monthly inbound leads at 20–28/month',
                '2 published case studies on website',
                'Residential retrofit booked through organic search (not just referral)',
            ],
        },
        {
            'label': 'QUARTER 4 — REMOVAL & REPLACEMENT + YEAR 2 PREP',
            'months': 'Months 10–12',
            'focus': 'Insulation Removal & Replacement',
            'goal': 'Capture remediation inquiries; solidify authority position; prepare Year 2 systems',
            'color': AMBER_HEX,
            'months_detail': [
                {
                    'month': 'Month 10 — Remediation Content Launch',
                    'items': [
                        'Blog Article #10: "How to Tell When Insulation Needs to Come Out, Not Just Be Added To"',
                        'AI avatar video #9: "What Contaminated Attic Insulation Looks Like — And What To Do About It"',
                        'Pre-listing insulation audit offer deployed to real estate agent partners',
                        'Google Posts: remediation-focused content (mold evidence, pest damage, smoke/fire damage scenarios)',
                    ],
                },
                {
                    'month': 'Month 11 — Peak Commercial Season Prep',
                    'items': [
                        'Email campaign to all commercial/industrial contacts: "Year-End Facility Maintenance — Insulation Inspection Before Budget Closes"',
                        'AI avatar video #10 published; YouTube library at 10 videos',
                        'Blog Article #11: "Pre-Season Checklist for Louisville Facility Managers: What to Inspect Before Winter"',
                        'Real estate partner appreciation: holiday card + "January listing pre-inspection" offer for Q1 pipeline',
                    ],
                },
                {
                    'month': 'Month 12 — Year-End Review + Year 2 Planning',
                    'items': [
                        'Full year KPI review against all targets',
                        'Authority audit: what position has been established? Evidence in organic search, review count, video library?',
                        'Pricing review: with 75+ reviews and established video library, basis for 10–15% price increase in Year 2?',
                        'Year 2 strategic brief: Roger + Tony strategy call. New service opportunities? Geographic expansion? Avatar scale-up?',
                        'Blog Article #12: "GISCO\'s Year in Review: Louisville Insulation Projects 2026"',
                    ],
                },
            ],
            'metrics': [
                '75+ Google reviews',
                '10+ AI avatar videos published',
                'Monthly inbound leads at 28–35/month',
                '3+ published case studies on website',
                'Monthly revenue at $85,000–$100,000',
                '15+ active professional referral partners',
            ],
        },
    ]

    for q in quarters:
        p = doc.add_paragraph()
        set_para_bg(p, q['color'])
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        rq = p.add_run(q['label'])
        rq.bold = True; rq.font.name = 'Calibri'; rq.font.size = Pt(12)
        rq.font.color.rgb = WHITE
        p2 = doc.add_paragraph()
        set_para_bg(p2, q['color'])
        p2.paragraph_format.left_indent  = Inches(0.15)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        rs = p2.add_run(f'{q["months"]}  |  Focus Service: {q["focus"]}')
        rs.italic = True; rs.font.name = 'Calibri'; rs.font.size = Pt(9)
        rs.font.color.rgb = WHITE
        sp(doc, 2)

        body(doc, 'STRATEGIC GOAL:', size=9, bold=True, color=DARK_COLOR)
        body(doc, q['goal'], size=10)
        sp(doc, 2)

        for md in q['months_detail']:
            body(doc, md['month'], size=10, bold=True, color=DARK_COLOR)
            for item in md['items']:
                bullet(doc, item, size=9.5)
            sp(doc, 3)

        body(doc, f'Q SUCCESS METRICS:', size=9, bold=True, color=DARK_COLOR)
        for m in q['metrics']:
            bullet(doc, '☐  ' + m, size=9.5)
        sp(doc, 8)

    # Year-end KPI table
    heading(doc, 'YEAR-END KPI TABLE', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)
    five_col_kpi_table(doc, [
        ('Monthly inbound leads', '~10 est.', '15–20', '20–28', '28–35'),
        ('Website monthly sessions', 'Unknown', '300+', '800+', '1,500+'),
        ('Google review count', '<20 est.', '25+', '50+', '75+'),
        ('New reviews / month', '0 (no system)', '8+', '10+', '12+'),
        ('Active referral partners', '0 (informal)', '5+', '10+', '15+'),
        ('AI avatar videos published', '0', '2', '5', '10+'),
        ('Monthly revenue', '~$42K est.', '$55K', '$75K', '$85–100K'),
        ('Blog articles published', '0', '5', '8', '12+'),
        ('Email list size', '0', '75+', '200+', '400+'),
    ], ['KPI', 'Baseline (M0)', '90-Day', '6-Month', '12-Month'])

    attribution(doc, 'Daniel Frost, CSO | Authority Systems Group™')
    add_pb(doc)


# ── Section 8: KPI Dashboard ──────────────────────────────────────────────────

def build_s8_kpi(doc):
    section_header(doc, '8', 'KPI Dashboard')
    attribution(doc, 'Preston Adler, COO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    quote_block(doc,
                '"Marketing without measurement is guessing with better branding. Every recommendation '
                'in this plan has a corresponding metric. Every metric has a target. Every target has '
                'an owner. That\'s how we know if the system is working — and how we know what to '
                'change when it isn\'t." — Preston Adler, COO')

    kpi_categories = [
        ('CATEGORY 1 — ACQUISITION & TOP-OF-FUNNEL', [
            ('Website sessions / month', 'To establish M1', '300+', '800+', '1,500+'),
            ('Monthly inbound inquiries (all sources)', '~10 est.', '15–20', '22–28', '28–35'),
            ('Monthly inbound — Commercial/Industrial', '~3–5 est.', '6–8', '8–12', '12–18'),
            ('Google review count (cumulative)', '<20 est.', '25+', '50+', '75+'),
            ('YouTube video views (cumulative)', '0', '200+', '1,000+', '4,000+'),
        ]),
        ('CATEGORY 2 — ENGAGEMENT & CONVERSION', [
            ('Estimate bookings / month', 'To establish M1', '8–12', '14–18', '18–24'),
            ('Estimate-to-job conversion rate', 'To establish M1', '40%+', '45%+', '50%+'),
            ('Avg days from inquiry to estimate', 'To establish M1', '<3 days', '<2 days', '<2 days'),
            ('AI phone capture rate (calls answered)', '0% (no system)', '95%+', '95%+', '95%+'),
            ('Lead quality score (1–10)', 'To establish M1', '6+ avg', '7+ avg', '7.5+ avg'),
        ]),
        ('CATEGORY 3 — REVENUE & VALUE', [
            ('New commercial jobs / month', '~2–3 est.', '4–5', '5–7', '7–10'),
            ('New residential jobs / month', '~3–5 est.', '5–8', '7–10', '10–14'),
            ('Avg residential ticket', '~$2,800 est.', '$3,200', '$3,500', '$3,800'),
            ('Monthly revenue (all sources)', '~$42K est.', '$55K', '$75K', '$90K'),
            ('Reactivation revenue (sprint)', '$0', '$25–55K', 'Reinvested', '—'),
        ]),
        ('CATEGORY 4 — REPUTATION & REFERRAL', [
            ('New Google reviews / month', '0 (no system)', '8', '10', '12'),
            ('Google review rating', 'To establish M1', '4.8+', '4.8+', '4.9+'),
            ('Review requests sent / month', '0', '15+', '20+', '25+'),
            ('Active professional referral partners', '0', '5+', '10+', '15+'),
            ('Referral % of new jobs', 'Unknown', '30%', '40%', '45%'),
        ]),
        ('CATEGORY 5 — CONTENT & DIGITAL', [
            ('Blog articles published (cumulative)', '0', '5', '8', '12+'),
            ('AI avatar videos published (cumulative)', '0', '2', '5', '10+'),
            ('Email list size', '0', '75+', '200+', '400+'),
            ('Organic search clicks / month', '0', '50+', '200+', '500+'),
            ('GBP views / month', 'To establish M1', '200+', '500+', '1,000+'),
        ]),
    ]

    for cat_title, rows in kpi_categories:
        heading(doc, cat_title, level=3, color=DARK_COLOR, size=10, bold=True, bg=ASG_LIGHT_BG)
        sp(doc, 1)
        five_col_kpi_table(doc, rows,
                           ['Metric', 'Baseline', '90-Day', '6-Month', '12-Month'])

    sp(doc, 4)
    heading(doc, 'DASHBOARD ARCHITECTURE', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    arch = [
        ('Weekly (Internal)', 'New leads vs. prior week; estimates booked; review requests sent; new reviews received; Fast Cash Sprint response tracking (Months 1–3)'),
        ('Monthly (Client-Facing)', 'All 25 KPIs vs. baseline and target, color-coded Green/Yellow/Red. Delivered first Friday of each month by ASG team.'),
        ('Quarterly (Strategy Review)', 'Full 25-KPI trend analysis; channel efficiency; referral partner performance; content library inventory; Q+1 priority adjustments.'),
        ('Recommended Tools', 'GHL (native pipeline reporting) + Google Search Console + GA4 + Simple monthly Google Sheets summary from ASG'),
    ]
    two_col_table(doc, arch, col_widths=(2.0, 5.5))

    quote_block(doc,
                '"On the first Friday of every month, you\'ll receive a report showing all 25 of these '
                'metrics against your targets. Nothing buried. Nothing hidden. If something isn\'t '
                'working, we\'ll see it and we\'ll tell you." — Preston Adler, COO')
    add_pb(doc)


# ── Section 9: Content Authority Strategy™ ───────────────────────────────────

def build_s9_content(doc):
    section_header(doc, '9', 'Content Authority Strategy™')
    attribution(doc, 'Vivienne Carr, CMO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    quote_block(doc,
                '"Every piece of content in this strategy has a specific job. It targets a specific '
                'stage on the Belief Track or the Emotion Track — or both simultaneously. Content '
                'without a stage assignment is just publishing. This is not publishing. This is '
                'architecture." — Vivienne Carr, CMO')

    # Authority Position
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(7.5)
    ac = tbl.rows[0].cells[0]
    set_cell_bg(ac, ASG_DARK_HEX)
    ap = ac.paragraphs[0]
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(10)
    ap.paragraph_format.space_after  = Pt(10)
    ap.paragraph_format.left_indent  = Inches(0.3)
    ap.paragraph_format.right_indent = Inches(0.3)
    ar = ap.add_run(
        'CONTENT AUTHORITY POSITION: GISCO Corp is Louisville\'s definitive source for commercial '
        'and residential decision-makers on the topic of insulation that actually performs — the one '
        'place where what insulation does, when it fails, and what to do about it is explained '
        'clearly, honestly, and by someone who has been doing this work for four generations without '
        'cutting corners or selling products they don\'t trust.'
    )
    ar.bold = True; ar.italic = True; ar.font.name = 'Calibri'; ar.font.size = Pt(10)
    ar.font.color.rgb = WHITE
    sp(doc, 6)

    # 4 Pillars
    heading(doc, 'THE FOUR CONTENT PILLARS', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    pillars = [
        ('PILLAR 1', 'Enemy Belief Dismantler', '25%', RED_HEX,
         'Surfaces and reframes the two enemy beliefs: "All insulation contractors are the same" '
         'and "Spray foam is just another insulation type." These are the top-of-funnel pieces '
         'that interrupt the assumption before the prospect ever gets to GISCO\'s website.',
         'AI avatar video, blog article, social hook post, Google Business Profile post'),
        ('PILLAR 2', 'Authority Educator', '35%', ASG_DARK_HEX,
         'Teaches what prospects don\'t know: what air sealing does, what thermography reveals, '
         'how industrial insulation specification works, what the 25C credit covers, what 4th '
         'generation institutional experience means in practical terms.',
         'Long-form blog article, AI avatar explainer video, YouTube educational series, email sequence'),
        ('PILLAR 3', 'Proof Mobilizer', '25%', SLATE_HEX,
         'Makes the after-state visible and concrete: case studies with documented outcomes, '
         'review-driven social proof, before/after thermal imaging, named institutional client '
         'references. Converts belief into desire.',
         'Case studies, testimonial features, before/after content, video walkthroughs of completed jobs'),
        ('PILLAR 4', 'Trust Closer', '15%', GREEN_HEX,
         'Removes the last friction before contact: "what to expect when you call," Tony\'s personal '
         'inspection process, "here\'s what we check before we quote," GISCO\'s no-spray-foam policy '
         'explained directly.',
         'FAQ content, consultation process explainer, "about Tony" content, AI phone explainer'),
    ]

    for pid, pname, pct, pcolor, pdesc, pformat in pillars:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(1.0)
        tbl.columns[1].width = Inches(6.5)
        lc = tbl.rows[0].cells[0]
        set_cell_bg(lc, pcolor)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(10)
        lp.paragraph_format.space_after  = Pt(4)
        lr1 = lp.add_run(pct + '\n')
        lr1.bold = True; lr1.font.name = 'Calibri'; lr1.font.size = Pt(18)
        lr1.font.color.rgb = WHITE
        lr2 = lp.add_run(pid)
        lr2.font.name = 'Calibri'; lr2.font.size = Pt(8)
        lr2.font.color.rgb = WHITE
        rc = tbl.rows[0].cells[1]
        set_cell_bg(rc, ASG_TINT)
        rp = rc.paragraphs[0]
        rp.paragraph_format.left_indent  = Inches(0.12)
        rp.paragraph_format.space_before = Pt(6)
        rp.paragraph_format.space_after  = Pt(2)
        rn = rp.add_run(pname)
        rn.bold = True; rn.font.name = 'Calibri'; rn.font.size = Pt(11)
        rn.font.color.rgb = DARK_COLOR
        rp2 = rc.add_paragraph()
        rp2.paragraph_format.left_indent = Inches(0.12)
        rp2.paragraph_format.space_after  = Pt(2)
        rd = rp2.add_run(pdesc)
        rd.font.name = 'Calibri'; rd.font.size = Pt(9.5)
        rd.font.color.rgb = BODY_COLOR
        rp3 = rc.add_paragraph()
        rp3.paragraph_format.left_indent = Inches(0.12)
        rp3.paragraph_format.space_before = Pt(0)
        rp3.paragraph_format.space_after  = Pt(6)
        rf = rp3.add_run('Key formats: ' + pformat)
        rf.italic = True; rf.font.name = 'Calibri'; rf.font.size = Pt(9)
        rf.font.color.rgb = ASG_CHARCOAL
        sp(doc, 3)

    sp(doc, 6)

    # Content briefs
    heading(doc, 'SERVICE CONTENT BRIEFS', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 4)

    briefs = [
        {
            'service': 'Commercial / Industrial Insulation',
            'focus': 'Q1 Focus',
            'enemy': '"Insulation contractors are a commodity line item in a commercial build or maintenance budget. We take bids, evaluate scope and price, and select a qualified vendor. There\'s no meaningful differentiation."',
            'authority': 'GISCO is the only Louisville insulation contractor with documented institutional references at Peerless Distillery, the University of Louisville, and maritime facilities. Tony\'s 4th generation knowledge of industrial insulation systems is a credential that cannot be invented or purchased.',
            'pieces': [
                ('"Why Bid Price Is the Wrong First Question for Industrial Insulation"', 'B1/E1', 'The assumption that qualified vendors are interchangeable is costing facilities managers money in rework', 'Blog + AI avatar video'),
                ('"What Happened When We Insulated [Kentucky Distillery]\'s Process Piping"', 'B2/E2', 'The contrast between a contractor who knows distillery environments and one who doesn\'t', 'Case study + video'),
                ('"4 Things GISCO Checks Before Quoting a Commercial Insulation Project"', 'B2/E3', 'The diagnostic standard GISCO applies that most contractors skip', 'Blog + social'),
                ('"Thermography Before We Quote: Why We Show You the Problem First"', 'B3/E4', 'The diagnostic-first approach as a risk reduction argument', 'Blog + AI video'),
                ('Case Study: University of Louisville Project', 'B4/E4', 'Documented proof with a recognizable name', 'Full case study'),
            ],
            'anchor': '"The Louisville Facilities Manager\'s Complete Guide to Mechanical Insulation: What to Evaluate, What to Ask, and Why 41 Years of History Matters" (2,200–2,800 words)',
            'keywords': 'Primary: "commercial insulation contractor Louisville KY" / "industrial insulation Louisville"\nSecondary: "mechanical insulation Louisville," "pipe insulation contractor Louisville," "insulation contractor Kentucky distillery"',
        },
        {
            'service': 'Infrared Thermography',
            'focus': 'Q2 Focus',
            'enemy': '"Thermal imaging is a specialized diagnostic for engineers and big buildings. I just need someone to look at my insulation and tell me what to add."',
            'authority': 'GISCO is the only insulation contractor in Louisville that leads with thermography as a diagnostic tool — making them the only contractor that shows you the problem before recommending a solution.',
            'pieces': [
                ('"Why Insulation Quotes Without Thermal Imaging Are Guesses"', 'B1/E1', 'You can\'t see insulation failures with the naked eye — the contractor who doesn\'t check doesn\'t know', 'Blog + AI video'),
                ('"What We Found in This Louisville Commercial Building Before We Touched a Thing"', 'B2/E2', 'Thermal scan reveals what the standard quote would have missed entirely', 'Video walkthrough'),
                ('"How GISCO\'s Thermal Diagnostic Works: What We Check and Why"', 'B3/E3', 'Process transparency + authority education', 'Blog + service page'),
                ('"Before/After Thermal Images: Louisville Attic Insulation Project"', 'B3/E4', 'Visual proof is more powerful than any written description', 'Photo/video post'),
                ('Case Study: Thermography Reveals Critical Industrial Insulation Failure', 'B4/E4', 'Specific documented outcome — cost avoided, equipment protected', 'Full case study'),
            ],
            'anchor': '"Infrared Thermography for Louisville Buildings: What It Is, What It Reveals, and Why It Changes Your Insulation Decision" (1,800–2,200 words)',
            'keywords': 'Primary: "infrared thermography Louisville KY" / "thermal imaging insulation Louisville"\nSecondary: "building energy audit Louisville," "thermal inspection contractor Louisville"',
        },
        {
            'service': 'Residential Retrofit',
            'focus': 'Q3 Focus',
            'enemy': '"All insulation contractors charge about the same for the same work. Get three quotes and go with the one that seems fair."',
            'authority': 'Tony Long personally crawls through every attic. Tony tells homeowners what they DON\'T need. GISCO doesn\'t do spray foam and explains exactly why. This combination is a category of one in Louisville residential insulation.',
            'pieces': [
                ('"Why GISCO Doesn\'t Offer Spray Foam — And What We Recommend Instead"', 'B1/E1', 'Values-based differentiation; captures spray foam concern search traffic — HIGH PRIORITY', 'Blog + AI video'),
                ('"Why I Crawl Through Every Attic Instead of Looking From the Ladder"', 'B1/E2', 'Tony\'s personal inspection standard dismantles the "they\'re all the same" assumption', 'AI avatar video (Tony\'s voice)'),
                ('"The Air Sealing Step Most Louisville Insulation Contractors Skip"', 'B2/E2', 'Education on what air sealing does and why it doubles the value of insulation', 'Blog + video'),
                ('"How to Stack Your 25C Credit and LG&E Rebate in 2026"', 'B3/E3', 'Removes cost fear; makes the ROI tangible and calculable', 'Blog + AI video'),
                ('"The Tony Long Inspection: What I Look For That Others Miss"', 'B4/E4', 'Personal story content drives Internal Alignment', 'AI avatar video (Tony\'s voice)'),
            ],
            'anchor': '"Louisville Homeowner\'s Complete Guide to Attic Insulation in 2026: What You Have, What You Need, and How to Pay Less Than You Think" (2,000–2,500 words)',
            'keywords': 'Primary: "attic insulation Louisville KY" / "blown in insulation Louisville"\nSecondary: "cellulose insulation Louisville," "attic air sealing Louisville," "spray foam alternative Louisville"',
        },
        {
            'service': 'Insulation Removal & Replacement',
            'focus': 'Q4 Focus',
            'enemy': '"Replacing insulation is too disruptive and expensive. We\'ll monitor it and deal with it when we absolutely have to."',
            'authority': 'Tony\'s honest assessment approach — "here\'s what has to come out and here\'s exactly why" — makes GISCO the trusted remediation advisor, not the contractor trying to manufacture urgency.',
            'pieces': [
                ('"How to Tell When Insulation Needs to Come Out, Not Just Be Added To"', 'B1/E1', 'Creates the diagnostic framework that identifies the problem they\'ve been ignoring', 'Blog + AI video'),
                ('"What Contaminated Attic Insulation Actually Looks Like"', 'B2/E2', 'Visual contrast creates the urgency that written descriptions can\'t', 'Video + photo'),
                ('"Pre-Listing Insulation Inspection: What Louisville Realtors Are Asking For"', 'B2/E3', 'Real estate trigger; surfaces urgency in the selling context', 'Blog + agent outreach'),
                ('Case Study: Contaminated Attic Removal + Replacement', 'B3/E4', 'Documented outcome — mold evidence gone, air quality restored', 'Full case study'),
                ('"What GISCO\'s Removal Process Looks Like: What We Take Out and What Goes In"', 'B5/E5', 'Process transparency; removes fear of disruption', 'FAQ + page'),
            ],
            'anchor': None,
            'keywords': 'Primary: "insulation removal Louisville KY" / "insulation replacement Louisville"\nSecondary: "contaminated attic insulation Louisville," "insulation removal mold Louisville"',
        },
    ]

    for brief in briefs:
        heading(doc, f'Content Strategy: {brief["service"]} ({brief["focus"]})',
                level=3, color=ASG_BLUE, size=11, bold=True)
        sp(doc, 2)
        callout_box(doc, 'THE ENEMY BELIEF:', brief['enemy'], bg='FFF4E5',
                    label_color=RGBColor(0xE6, 0x7E, 0x22))
        callout_box(doc, 'THE AUTHORITY ANGLE:', brief['authority'], bg=ASG_TINT,
                    label_color=DARK_COLOR)

        # Content pieces table
        tbl = doc.add_table(rows=len(brief['pieces']) + 1, cols=4)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(2.5)
        tbl.columns[1].width = Inches(0.75)
        tbl.columns[2].width = Inches(2.5)
        tbl.columns[3].width = Inches(1.75)
        for ci, hdr in enumerate(['Content Piece', 'Stage', 'Core Message', 'Format']):
            hc = tbl.rows[0].cells[ci]
            set_cell_bg(hc, ASG_DARK_HEX)
            hp = hc.paragraphs[0]
            hp.paragraph_format.left_indent = Inches(0.06)
            hr = hp.add_run(hdr)
            hr.bold = True; hr.font.name = 'Calibri'; hr.font.size = Pt(8.5)
            hr.font.color.rgb = WHITE
        for i, (piece, stage, msg, fmt) in enumerate(brief['pieces']):
            bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
            for ci, val in enumerate([piece, stage, msg, fmt]):
                cell = tbl.rows[i + 1].cells[ci]
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.left_indent = Inches(0.06)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(val)
                r.font.name = 'Calibri'; r.font.size = Pt(8.5)
                r.bold = (ci == 1)
                r.font.color.rgb = ASG_BLUE if ci == 1 else BODY_COLOR
        sp(doc, 3)

        if brief['anchor']:
            callout_box(doc, 'LONG-FORM ANCHOR PIECE:', brief['anchor'], bg=ASG_TINT,
                        label_color=DARK_COLOR)
        body(doc, 'Keywords: ' + brief['keywords'], size=9, italic=True, color=ASG_CHARCOAL)
        sp(doc, 6)

    # Format priority matrix
    heading(doc, 'FORMAT PRIORITY MATRIX', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    four_col_kpi_table(doc, [
        ('AI Avatar Video (HeyGen / YouTube)', 'HIGH', 'B1–B3 / E1–E3', '2/month (ASG writes scripts)'),
        ('Long-form blog article', 'HIGH', 'B2–B4 / E2–E4', '1/month'),
        ('Case studies', 'HIGH', 'B3–B4 / E3–E4', '1/quarter'),
        ('Google Business Profile posts', 'HIGH', 'B4–B5 / E4–E5', '2/week'),
        ('Email sequences', 'HIGH', 'All stage pairs', 'Automated'),
        ('Social posts (organic FB + LI)', 'MEDIUM', 'B1–B2 / E1–E2', '3/week'),
        ('Short-form video clips (from AI avatar)', 'MEDIUM', 'B1–B2 / E1–E2', '2/week'),
        ('Direct mail (test)', 'LOW (test only)', 'B1 / E1', 'Q3 one-time test'),
    ], ['Format', 'Priority', 'Stage Pairs', 'Frequency'])

    sp(doc, 4)
    heading(doc, 'THE HEYGEN AI AVATAR ENGINE', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    avatar_steps = [
        ('Sessions 1–4 (Q1–Q2)',
         'Tony films 4 sessions of 60–90 minutes each. ASG directs each session. Subjects: '
         'Tony\'s story, industrial work, residential approach, thermography explanation, spray foam '
         'rationale, attic inspection process. Raw footage total: 4–6 hours.'),
        ('HeyGen Avatar Creation',
         'From Tony\'s footage, a HeyGen AI avatar is created — his face, his voice, his mannerisms. '
         'This avatar can then deliver any script ASG writes at any time, at scale.'),
        ('Ongoing production (Q2 forward)',
         'ASG writes 2–4 scripts per month. Tony approves (30 minutes of his time per month). '
         'HeyGen renders the video. Published to YouTube + website + social clips.'),
        ('Compound effect',
         'By Month 12, GISCO has a 24+ video library on YouTube — every video indexed, every video '
         'building topical authority. No competitor in Louisville has even one video. This is an '
         'uncontested channel.'),
    ]
    for lbl, txt in avatar_steps:
        callout_box(doc, lbl + ':', txt, bg=ASG_TINT, label_color=DARK_COLOR)

    body(doc,
         'The initial camera sessions with Tony are non-negotiable for this strategy to work. '
         'Everything else in the content plan depends on this footage.',
         size=10, bold=True, color=DARK_COLOR)
    add_pb(doc)


# ── Section 10: Digital Infrastructure ───────────────────────────────────────

def build_s10_digital(doc):
    section_header(doc, '10', 'Digital Infrastructure & Website Strategy')
    attribution(doc, 'Iris Nolan, CTO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    body(doc,
         'THE STRATEGIC CASE: giscocorp.com is non-functional for incoming traffic — it does not rank '
         'for any high-intent local search terms, does not convert visitors into inquiries, and does '
         'not reflect the 41-year institutional track record that is GISCO\'s primary authority asset. '
         'The AI phone employee and new website are already in progress. The question is what the new '
         'site is built to do.',
         size=10.5)
    body(doc,
         'THE OPPORTUNITY: Local search for "insulation contractor Louisville," "commercial insulation '
         'Louisville," and "attic insulation Louisville KY" is sparsely contested. A properly '
         'architected website with specific service pages, locally-optimized content, and structured '
         'schema will rank in 90–120 days without competing against established content authority. '
         'The field is largely open.',
         size=10.5)

    sp(doc, 4)
    heading(doc, 'WEBSITE ARCHITECTURE AT LAUNCH', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    pages = [
        ('HOME PAGE',
         'H1: "Louisville\'s 4th Generation Insulation Authority — Commercial, Industrial & Residential" | '
         'Opening statement with 41-year history + institutional references | "No Spray Foam" statement | '
         '4 service cards with internal links | Click-to-call prominent above fold | '
         'Google reviews widget | Free Assessment CTA'),
        ('ABOUT TONY LONG',
         'Tony\'s personal story (41 years, 4th generation) | "Why I Don\'t Offer Spray Foam" in his own words | '
         'Named institutional client list | Credentials and certifications | '
         'Tony\'s inspection philosophy ("I crawl through. Here\'s why.")'),
        ('COMMERCIAL & INDUSTRIAL INSULATION',
         '800–1,000 words | Primary keyword: "commercial insulation contractor Louisville KY" | '
         'Industrial specializations: distilleries, food processing, power plants, universities | '
         'Named client references | Process description | FAQ (5 questions with schema markup)'),
        ('INFRARED THERMOGRAPHY',
         '700–900 words | Primary keyword: "infrared thermography Louisville KY" | '
         'What it is, what it reveals, who it\'s for | '
         'Pricing and booking (named service: "GISCO Thermal Diagnostic — from $350") | FAQ (4 questions)'),
        ('RESIDENTIAL INSULATION',
         '800–1,000 words | Primary keyword: "attic insulation Louisville KY" | '
         'Cellulose / blown-in explanation (not spray foam — why) | '
         'IRA 25C credit + LG&E rebate explanation | Tony\'s inspection process | FAQ (5 questions with schema)'),
        ('INSULATION REMOVAL & REPLACEMENT',
         '600–800 words | Primary keyword: "insulation removal Louisville KY" | '
         'When to remove vs. add; contamination scenarios | Process description | FAQ (3 questions)'),
        ('BLOG / RESOURCES',
         '5 articles at launch (from Content Strategy) | No placeholder posts — all fully written'),
        ('FAQ (Standalone Page)',
         '12–15 questions with complete answers | Schema markup applied (featured snippet optimization) | '
         'Covers: spray foam question, thermography, IRA credit, inspection process, GISCO history'),
        ('FREE ASSESSMENT / CONTACT',
         'Clear description of free assessment process | Booking form + Calendly | '
         'Click-to-call (AI phone number) | What happens after submission (sets expectations, reduces drop-off)'),
        ('LEAD MAGNET LANDING PAGE',
         '"The Louisville Homeowner\'s Honest Guide to Insulation in 2026" opt-in | '
         'Name + email → triggers Welcome Sequence'),
    ]

    for pg_name, pg_desc in pages:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(2.1)
        tbl.columns[1].width = Inches(5.4)
        lc = tbl.rows[0].cells[0]
        rc = tbl.rows[0].cells[1]
        set_cell_bg(lc, ASG_DARK_HEX)
        set_cell_bg(rc, ASG_TINT)
        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent  = Inches(0.1)
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after  = Pt(5)
        lr = lp.add_run(pg_name)
        lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(9)
        lr.font.color.rgb = WHITE
        rp = rc.paragraphs[0]
        rp.paragraph_format.left_indent  = Inches(0.1)
        rp.paragraph_format.space_before = Pt(5)
        rp.paragraph_format.space_after  = Pt(5)
        rr = rp.add_run(pg_desc)
        rr.font.name = 'Calibri'; rr.font.size = Pt(8.5)
        rr.font.color.rgb = BODY_COLOR
        sp(doc, 2)

    sp(doc, 4)
    heading(doc, 'TECHNICAL SPECIFICATIONS', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    two_col_table(doc, [
        ('Platform', 'WordPress with Astra or Kadence (no Elementor/Divi — performance requirement)'),
        ('Hosting', 'Kinsta or WP Engine (US server, SSL from Day 1)'),
        ('Target load time', 'Under 2 seconds (Google PageSpeed 90+)'),
        ('Mobile-first', 'All CTAs functional on mobile; click-to-call in header'),
        ('Tracking', 'GA4 + Google Search Console live before launch; GTM for click-to-call + form tracking'),
        ('AI phone integration', 'Phone number on site = AI phone system number. All calls captured.'),
    ], col_widths=(2.0, 5.5))

    heading(doc, 'LOCAL SEO STRATEGY', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    callout_box(doc, 'Primary Keywords:',
                '"commercial insulation contractor Louisville KY" | "industrial insulation Louisville" | '
                '"attic insulation Louisville KY" | "insulation contractor Louisville Kentucky" | '
                '"infrared thermography Louisville KY"',
                bg=ASG_TINT, label_color=DARK_COLOR)
    callout_box(doc, 'Secondary / Long-Tail:',
                '"blown in insulation Louisville KY" | "insulation removal Louisville KY" | '
                '"cellulose insulation Louisville" | "spray foam alternative Louisville" | '
                '"4th generation insulation company Louisville" | "insulation contractor Kentucky distillery"',
                bg=ASG_TINT, label_color=DARK_COLOR)

    heading(doc, 'GOOGLE BUSINESS PROFILE OPTIMIZATION', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    for item in [
        'All 5 service categories populated with full descriptions (Commercial Insulation, Industrial Insulation, Infrared Thermography, Insulation Removal, Residential Insulation)',
        'Photos: Tony on a job, crew at work, thermal imaging in action, before/after jobs — minimum 15 photos at launch',
        'Q&A seeded with 8–10 FAQs (ASG writes; Tony publishes)',
        'Google Posts: 2/week beginning at website launch; aligned with content calendar',
        'Service area confirmed: Louisville metro + Jefferson, Bullitt, Oldham, Shelby, Spencer counties',
    ]:
        bullet(doc, item, size=9.5)

    sp(doc, 4)
    heading(doc, 'AI PHONE EMPLOYEE INTEGRATION', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    ai_items = [
        ('No missed calls', 'Every inquiry is captured regardless of whether Tony is on a job site. This alone closes the single biggest gap in GISCO\'s current pipeline.'),
        ('Transcription to Tony', 'Every call summary delivered to Tony\'s phone/email same day. Tony reviews and follows up personally on the highest-value commercial inquiries.'),
        ('Data intelligence', 'Over time, call transcripts reveal the most common questions, the most frequent objection patterns, and the geographic distribution of inquiries — all inputs for content strategy.'),
        ('Review automation hook', 'Post-job follow-up calls from the AI system trigger the review request sequence automatically.'),
    ]
    for lbl, txt in ai_items:
        callout_box(doc, lbl + ':', txt, bg=ASG_TINT, label_color=DARK_COLOR)

    attribution(doc, 'Iris Nolan, CTO | Authority Systems Group™')
    add_pb(doc)


# ── Section 11: Email Sequences ───────────────────────────────────────────────

def build_s11_emails(doc):
    section_header(doc, '11', 'Email Sequences (Fully Written)')
    attribution(doc, 'Vivienne Carr, CMO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    def email_block(subject, send_time, body_text):
        tbl = doc.add_table(rows=1, cols=1)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(7.5)
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, ASG_LIGHT_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        rs = p.add_run(send_time + '  ')
        rs.bold = True; rs.font.name = 'Calibri'; rs.font.size = Pt(8.5)
        rs.font.color.rgb = ASG_BLUE
        rj = p.add_run(f'SUBJECT: {subject}')
        rj.bold = True; rj.font.name = 'Calibri'; rj.font.size = Pt(9.5)
        rj.font.color.rgb = DARK_COLOR
        pb = cell.add_paragraph()
        pb.paragraph_format.left_indent  = Inches(0.15)
        pb.paragraph_format.space_before = Pt(4)
        pb.paragraph_format.space_after  = Pt(8)
        rb = pb.add_run(body_text)
        rb.font.name = 'Calibri'; rb.font.size = Pt(9.5)
        rb.font.color.rgb = BODY_COLOR
        sp(doc, 4)

    # Sequence 1
    heading(doc, 'SEQUENCE 1: WELCOME SEQUENCE', level=2, color=ASG_BLUE, size=13, bold=True)
    body(doc, '5 emails | Triggers on form submission or lead magnet download | Deploys over 14 days',
         size=9, italic=True, color=ASG_CHARCOAL)
    sp(doc, 4)

    email_block(
        'Your honest insulation guide from Tony Long',
        'EMAIL 1 — Day 0 (Immediate)',
        "Tony here. Thanks for reaching out to GISCO.\n\n"
        "I've been doing insulation work in Louisville for a long time — 41 years in the business, "
        "fourth generation in my family to do it. And the one thing I've learned is that most people "
        "don't get a straight answer when they ask insulation questions. They get a quote for "
        "whatever the contractor sells.\n\n"
        "That's not how we do it.\n\n"
        "Over the next few days, I'm going to send you a few short notes about what we actually look "
        "for, how we decide what a building or a home needs, and why we've made some choices — like "
        "not offering spray foam — that aren't always popular but we stand behind.\n\n"
        "If you have questions before then, the best way to reach me is to call the number on our "
        "website. Our phone system will connect you directly. I listen to every call summary personally.\n\n"
        "Talk soon,\nTony Long | GISCO Corp | (502) 584-7349"
    )

    email_block(
        'Why I crawl through the attic (when most contractors won\'t)',
        'EMAIL 2 — Day 2',
        "If you've gotten more than one insulation estimate in Louisville, you've probably noticed "
        "something.\n\nMost contractors walk up to the attic pull-down, look from the top of the ladder "
        "for 90 seconds, and write you a quote. They're quoting what they assume is there, not what's "
        "actually there.\n\nI crawl through.\n\nNot because it's faster — it's slower. But because you "
        "can't see the unsealed chases, the failed baffles, the compressed insulation over the exterior "
        "wall plates, or the recessed lights that are leaking conditioned air directly into the "
        "unconditioned attic from a 90-second look at the top of the ladder.\n\nI've been in attics in "
        "Louisville where the homeowner had already gotten two quotes from contractors who missed the "
        "most significant problem entirely.\n\nIf you want to see what that looks like, we're doing "
        "free assessments through the end of the month.\n\nTony"
    )

    email_block(
        'What we won\'t do — and why that matters',
        'EMAIL 3 — Day 5',
        "I get asked about spray foam regularly. Here's the honest answer.\n\n"
        "GISCO does not install spray foam. We made that decision because we've watched the liability "
        "issues develop — most recently the UK class action — and we don't believe the risk profile "
        "is appropriate for our clients. Black mold issues in poorly-installed spray foam jobs are "
        "real, they're not isolated, and they're expensive to remediate.\n\n"
        "What we recommend instead depends on the specific situation. For most Louisville attics, "
        "properly air-sealed blown-in cellulose achieves the same or better energy performance than "
        "open-cell spray foam, with none of the liability profile, at a lower cost.\n\n"
        "The honest answer is: for most residential applications in Louisville, you don't need spray "
        "foam to get excellent results. Anyone telling you otherwise should be able to explain exactly "
        "why your specific situation requires it.\n\nTony"
    )

    email_block(
        'The government is paying part of your insulation upgrade this year',
        'EMAIL 4 — Day 9',
        "Quick one because this is worth knowing.\n\n"
        "The IRA 25C tax credit covers 30% of qualified insulation and air sealing work up to $1,200 "
        "per year through 2032. If you're a Louisville homeowner, you may also qualify for an LG&E "
        "utility rebate on top of that.\n\n"
        "That means a $4,500 attic air sealing and insulation job could effectively cost you "
        "$2,800–$3,200 after credits and rebates. The math doesn't change what we recommend — it just "
        "makes the conversation about whether it's worth it a lot different.\n\n"
        "We're not tax advisors and I always recommend confirming with your accountant. But I can walk "
        "you through what we've seen most homeowners qualify for and how to document the work for the "
        "credit. Ask us when we come out.\n\nTony"
    )

    email_block(
        'One last thing before I stop filling your inbox',
        'EMAIL 5 — Day 14',
        "I've sent you a few notes over the past couple weeks. I wanted to close the loop with one "
        "straightforward thing.\n\n"
        "GISCO has been doing this work in Louisville since 1985. We've insulated the University of "
        "Louisville, Peerless Distillery, Jeff Boat, Winston Industries, and a lot of homes in between. "
        "We've never installed a product we don't trust to protect our clients. We've never recommended "
        "more work than a building needs.\n\n"
        "If you're ready to have us look at your project — commercial, industrial, or residential — "
        "you can book a free assessment on our website or call us at (502) 584-7349. The call will be "
        "answered, your information will come directly to me, and if it's a project we can help with, "
        "I'll tell you exactly what it is and what it costs.\n\n"
        "If it's not a project we're the right fit for, I'll tell you that too.\n\n"
        "That's all we do.\n\nTony Long | GISCO Corp | 531 S 15th St, Louisville, KY 40203"
    )

    sp(doc, 6)
    # Sequence 2
    heading(doc, 'SEQUENCE 2: POST-JOB REVIEW & REFERRAL CAPTURE', level=2, color=ASG_BLUE, size=13, bold=True)
    body(doc, '3 emails | Triggers 3 days after job completion is logged in GHL | All completed residential and commercial jobs',
         size=9, italic=True, color=ASG_CHARCOAL)
    sp(doc, 4)

    email_block(
        'How did we do?',
        'EMAIL 1 — Day 3 Post-Job',
        "[First name] —\n\n"
        "We wrapped up your project recently. I wanted to check in.\n\n"
        "Has everything looked and performed the way you expected? If anything feels off or you have "
        "a question about what we did, please call us directly — (502) 584-7349.\n\n"
        "If things are good, I'd be grateful if you'd leave us a Google review. It takes about 90 "
        "seconds and it helps other Louisville homeowners and businesses find us when they need honest "
        "advice.\n\nHere's the direct link: [Google Review Link]\n\n"
        "Thank you for trusting us with your project.\n\nTony Long | GISCO Corp"
    )

    email_block(
        'A small ask from Tony',
        'EMAIL 2 — Day 8 Post-Job (if no review yet)',
        "[First name] —\n\nOne quick follow-up.\n\n"
        "I don't ask for reviews often — most of our business comes from people we've worked with "
        "before or people they've referred us to. But online reviews have become how new clients "
        "find us, and we could use more of them.\n\n"
        "If you had a good experience, here's the link again: [Google Review Link]\n\n"
        "If you didn't have a good experience, please reply to this email and tell me directly. "
        "I want to know.\n\nTony"
    )

    email_block(
        'Know anyone who could use our help?',
        'EMAIL 3 — Day 14 Post-Job',
        "[First name] —\n\nLast note from me on this.\n\n"
        "If you know a neighbor, friend, or colleague who's dealing with an insulation issue — high "
        "bills, an uncomfortable room, a moisture problem, or a commercial facility that needs a "
        "reliable contractor — I'd appreciate the introduction.\n\n"
        "You can just forward this email and tell them to call (502) 584-7349. They'll speak with "
        "our phone system first, but I listen to every summary personally and I'll follow up same "
        "day on anything commercial or significant.\n\n"
        "No pressure. If someone comes to mind, I appreciate it.\n\nTony Long | GISCO Corp"
    )

    sp(doc, 6)
    # Sequence 3
    heading(doc, 'SEQUENCE 3: HVAC PARTNER REFERRAL OUTREACH', level=2, color=ASG_BLUE, size=13, bold=True)
    body(doc, '2 emails + follow-up | For HVAC company/tech contacts | Tony sends personally',
         size=9, italic=True, color=ASG_CHARCOAL)
    sp(doc, 4)

    email_block(
        'Referral partnership — GISCO Insulation',
        'EMAIL 1 — Initial Outreach',
        "[Name] —\n\n"
        "Tony Long here at GISCO Corp. We've been doing commercial and residential insulation in "
        "Louisville for 41 years — fourth generation in the business.\n\n"
        "I'm reaching out because I know you're regularly in homes and buildings where you find "
        "comfort or efficiency problems that turn out to be insulation issues, not HVAC issues. And "
        "I imagine you've had the experience of not having a great answer for a homeowner when "
        "that's the diagnosis.\n\n"
        "We'd like to be your go-to insulation contractor when that comes up. No spray foam — we "
        "don't do it. Honest assessments, written quotes, clean work. Every homeowner you refer will "
        "be treated exactly the way you'd want your clients treated.\n\n"
        "I'm attaching a short overview of what we do and how the referral arrangement works.\n\n"
        "Would you be open to a 10-minute call?\n\nTony Long | GISCO Corp | (502) 584-7349"
    )

    email_block(
        'Re: Referral partnership — GISCO Insulation',
        'EMAIL 2 — Follow-Up (Day 7 if no response)',
        "[Name] —\n\nFollowing up on the note I sent last week.\n\n"
        "Short version: we're Louisville's longest-standing local insulation contractor (since 1985, "
        "4th generation). We don't oversell, we don't install spray foam, and we send a thank-you "
        "to every individual tech who refers us a job that we close.\n\n"
        "If you've got homeowners asking about insulation — especially if they're asking about spray "
        "foam alternatives — we'd love the introduction.\n\n"
        "Happy to talk whenever works for you.\n\nTony"
    )

    attribution(doc, 'Vivienne Carr, CMO | Authority Systems Group™')
    add_pb(doc)


# ── Section 12: Lead Magnet & Referral Program ────────────────────────────────

def build_s12_lead_magnet(doc):
    section_header(doc, '12', 'Lead Magnet & Referral Program Framework')
    attribution(doc, 'Vivienne Carr, CMO + Tanya Blackwood, CRO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    quote_block(doc,
                '"The lead magnet and the referral program are the two highest-leverage assets in the '
                'Blueprint for one simple reason: they both turn existing trust into new pipeline. The '
                'lead magnet converts a stranger\'s first Google search into a relationship. The '
                'referral program converts a satisfied client\'s goodwill into a warm introduction. '
                'Neither requires a single dollar of paid media. Both start producing within 30 days '
                'of launch." — Vivienne Carr, CMO')

    heading(doc, 'PART 1: THE LEAD MAGNET', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)

    callout_box(doc, 'TITLE:',
                '"The Louisville Homeowner\'s Honest Guide to Insulation in 2026: What You Have, '
                'What You Need, and What the Contractor Probably Won\'t Tell You"',
                bg=ASG_TINT, label_color=DARK_COLOR)
    callout_box(doc, 'FORMAT:',
                'PDF, 14–18 pages, fully designed with GISCO branding. Tony\'s photo on page 2. '
                'Distribution via website opt-in landing page.',
                bg=ASG_TINT, label_color=DARK_COLOR)

    heading(doc, 'Table of Contents (with Stage Pair Mapping)', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)

    toc_data = [
        ('1. Before You Get a Single Quote', 'B1/E1', 'Pattern interrupt — meets homeowner before they\'ve made any assumptions'),
        ('2. Why Most Insulation Estimates Miss the Biggest Problem', 'B1/E2', 'Introduces the inspection standard gap; surfaces dissatisfaction with standard practice'),
        ('3. The Spray Foam Question: What the UK Lawsuit Means for Louisville Homeowners', 'B2/E2', 'Directly addresses spray foam concern; positions GISCO\'s no-spray-foam policy as a values statement'),
        ('4. What Air Sealing Is and Why It Matters as Much as R-Value', 'B2/E3', 'New Worldview — the technical truth most contractors don\'t explain'),
        ('5. The 4 Things Every Attic Inspection Should Check (And Usually Doesn\'t)', 'B3/E3', 'Tony\'s inspection framework as the evaluation standard'),
        ('6. What Your Attic Insulation Should Look Like — A Simple Guide', 'B3/E3', 'Makes the after-state visible and the current state evaluable'),
        ('7. The IRA 25C Credit + LG&E Rebate: How Louisville Homeowners Stack Both', 'B3/E4', 'Removes cost objection; makes the upgrade feel financially intelligent'),
        ('8. When It\'s Time to Replace vs. Add: The Honest Answer', 'B4/E4', 'Internal Alignment + Urgency — the cost of delay in plain terms'),
        ('9. Cellulose vs. Fiberglass vs. Spray Foam: Which One Does GISCO Recommend?', 'B4/E4', 'Product comparison done honestly; confirms GISCO\'s position without selling'),
        ('10. What Happens When You Call GISCO: Your Free Assessment, Explained', 'B5/E5', 'Removes all unknowns about the next step; makes contact feel safe and obvious'),
    ]

    tbl = doc.add_table(rows=len(toc_data) + 1, cols=3)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(0.8)
    tbl.columns[2].width = Inches(3.5)
    for ci, hdr in enumerate(['Section', 'Stage', 'Purpose']):
        hc = tbl.rows[0].cells[ci]
        set_cell_bg(hc, ASG_DARK_HEX)
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent = Inches(0.08)
        hr = hp.add_run(hdr)
        hr.bold = True; hr.font.name = 'Calibri'; hr.font.size = Pt(9)
        hr.font.color.rgb = WHITE
    for i, (sec, stage, purpose) in enumerate(toc_data):
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([sec, stage, purpose]):
            cell = tbl.rows[i + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.08)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(val)
            r.font.name = 'Calibri'; r.font.size = Pt(9)
            r.bold = (ci == 1)
            r.font.color.rgb = ASG_BLUE if ci == 1 else BODY_COLOR
    sp(doc, 6)

    callout_box(doc, 'Distribution:',
                'Primary: Website opt-in (triggers Welcome Sequence) | '
                'Secondary: HVAC partner kit — tech hands to homeowner when insulation question comes up | '
                'Tertiary: Past client reactivation sequence — included as value-first attachment in Email 1',
                bg=ASG_TINT, label_color=DARK_COLOR)

    sp(doc, 6)
    heading(doc, 'PART 2: THE REFERRAL PROGRAM', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)

    callout_box(doc, 'PROGRAM NAME:', 'The GISCO Honest Referral Program', bg=ASG_TINT, label_color=DARK_COLOR)

    body(doc, 'THE ASK FORMAT — THREE OPTIONS', size=10, bold=True, color=DARK_COLOR)
    sp(doc, 2)

    asks = [
        ('OPTION A — Pre-Written Text (Recommended)',
         '"Hey [Name] — if you ever need insulation work done right in Louisville, call Tony Long at '
         'GISCO: (502) 584-7349. Tell him [Your Name] sent you. He\'s been doing this 41 years and '
         'he actually crawls through the attic instead of just looking from the ladder. Honest guy."',
         ASG_TINT),
        ('OPTION B — Forward-This-Email',
         'Email 3 of Sequence 2 is written to be forwarded. One-click introduction.',
         'FFFFFF'),
        ('OPTION C — Direct Introduction',
         '"Reply with their name and number and I\'ll call them personally within 24 hours."',
         ASG_TINT),
    ]
    for opt, txt, bg in asks:
        callout_box(doc, opt + ':', txt, bg=bg, label_color=DARK_COLOR)

    heading(doc, 'INCENTIVE DESIGN', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    incentives = [
        ('Residential clients', 'No financial incentive required — Tony\'s personal style and existing relationship does the work'),
        ('Individual HVAC techs and home inspectors', '$50 gift card to a local Louisville restaurant (Tony\'s choice) per closed referral — personal, from Tony to the individual tech, not from the company'),
        ('Commercial facilities managers', 'No financial incentive — personal phone call from Tony after a referral closes is more impactful and appropriate in institutional context'),
    ]
    for cat, approach in incentives:
        callout_box(doc, cat + ':', approach, bg=ASG_TINT, label_color=DARK_COLOR)

    heading(doc, 'PROFESSIONAL REFERRAL NETWORK — 90-DAY OUTREACH TARGETS', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    four_col_kpi_table(doc, [
        ('HVAC Contractors', '8–10', 'Personal email + Partner Kit', '$50/tech for closed referrals'),
        ('Home Inspectors', '5–7', 'Professional letter + Partner Kit', 'Lead magnet as inspection report follow-up resource'),
        ('Real Estate Agents', '5–10', 'Email + "pre-listing audit" offer', 'Free pre-listing thermal assessment'),
        ('Commercial Property Managers', '3–5', 'Tony personal outreach', 'Annual maintenance proposal'),
    ], ['Category', 'Target Count', 'Outreach Method', 'Offer'])

    heading(doc, 'REFERRAL KPIs', level=3, color=DARK_COLOR, size=11, bold=True)
    sp(doc, 2)
    five_col_kpi_table(doc, [
        ('Referral asks sent / month', '0', '20+', '30+', '40+'),
        ('Active professional referral partners', '0', '5+', '10+', '15+'),
        ('New jobs from referrals / month', 'Unknown', '3–5', '5–8', '8–12'),
        ('Referral % of total new jobs', 'Unknown', '30%', '40%', '45%'),
    ], ['Metric', 'Baseline', '90-Day', '6-Month', '12-Month'])

    add_pb(doc)


# ── Section 13: Next Steps ────────────────────────────────────────────────────

def build_s13_next_steps(doc):
    section_header(doc, '13', 'Next Steps & Delivery Timeline')
    attribution(doc, 'Preston Adler, COO — Authority Systems Group™')
    divider(doc)
    sp(doc, 6)

    heading(doc, 'WHAT\'S COMING — YOUR FIRST 60 DAYS', level=2, color=ASG_BLUE, size=13, bold=True)
    sp(doc, 4)

    phases = [
        {
            'title': 'WEEK 1–2 (Days 1–14): FAST CASH SPRINT LAUNCH + DIGITAL FOUNDATION',
            'color': ASG_DARK_HEX,
            'from_asg': [
                'Past client reactivation sequence (3-touch email/text) — written, reviewed, ready to deploy',
                'HVAC partner outreach letter + Partner Kit (Tony\'s bio, service overview, lead magnet intro)',
                'Review request templates (personal version + automated version for AI phone)',
                'GHL pipeline structure and automation map — ready for implementation',
                'Google Business Profile optimization copy (all service descriptions, Q&A seed list)',
                'Lead magnet outline approved; full production begins',
            ],
            'from_tony': [
                'Past client list — any format, minimum 40 contacts (name + phone or email) — NEEDED BY DAY 7',
                'List of HVAC contacts Tony already knows personally (5+ names)',
                'Google Business Profile login credentials',
                'Approval of reactivation sequence messaging (48-hour turnaround)',
                'Confirmation of first video session date (Weeks 3–4 target)',
            ],
        },
        {
            'title': 'WEEK 3–4 (Days 15–28): CONTENT ENGINE + VIDEO SESSION 1',
            'color': SLATE_HEX,
            'from_asg': [
                'New website reviewed for launch readiness (Roger + Iris Nolan spec check)',
                '5 blog articles at launch — written and ready for Tony\'s final review',
                'Video Session 1 with Tony — ASG produces shot list and topic guide in advance',
                'Welcome Sequence (5 emails) — written, reviewed, loaded in GHL',
                'Social post bank: 12 posts for first month (LinkedIn + Facebook)',
            ],
            'from_tony': [
                'Final approval on 5 blog articles (48-hour turnaround)',
                'Show up for Video Session 1 — Tony\'s most important commitment in Month 1',
                'Review reactivation responses; flag any commercial opportunities for same-day follow-up',
            ],
        },
        {
            'title': 'WEEK 5–8 (Days 29–56): LAUNCH + VIDEO SESSION 2',
            'color': '276B3D',
            'from_asg': [
                'New website live: all core pages, GA4 + GSC configured, AI phone number integrated',
                'HeyGen avatar creation initiated from Video Session 1 footage',
                'First AI avatar video script written and submitted to Tony for approval',
                'Video Session 2 with Tony scheduled and briefed',
                'Lead magnet (PDF) complete and live on website opt-in',
                'HVAC Partner Kit distributed to all active partners',
                'GBP: Google Posts live (2/week cadence begins)',
            ],
            'from_tony': [
                'Video Session 2 — second most important commitment in Month 2',
                'First AI avatar video approved',
                'Ongoing: tag all completed jobs in GHL within 24 hours (triggers review request automation)',
            ],
        },
    ]

    for phase in phases:
        p = doc.add_paragraph()
        set_para_bg(p, phase['color'])
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(phase['title'])
        r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        sp(doc, 2)

        body(doc, 'FROM ASG:', size=9, bold=True, color=DARK_COLOR)
        for item in phase['from_asg']:
            bullet(doc, '☐  ' + item, size=9.5)
        sp(doc, 2)
        body(doc, 'FROM TONY:', size=9, bold=True, color=DARK_COLOR)
        for item in phase['from_tony']:
            bullet(doc, '☐  ' + item, size=9.5)
        sp(doc, 6)

    sp(doc, 4)
    heading(doc, 'WHAT TONY NEEDS TO PROVIDE', level=3, color=DARK_COLOR, size=12, bold=True)
    sp(doc, 2)

    needs = [
        ('Needed immediately (within 5 business days)', [
            'Past client list (any format — even a handwritten list is a start)',
            'HVAC contractor contacts Tony already knows personally',
            'Google Business Profile login',
            'Confirmation of video session date',
        ]),
        ('Needed within 30 days', [
            'Completed video sessions 1 and 2 (filming)',
            'Approvals on reactivation messaging, blog articles, email sequences (48-hour turnaround per batch)',
            'Any job photos Tony has (phone photos are fine — quantity over perfection at this stage)',
        ]),
        ('Ongoing throughout engagement', [
            'Tag completed jobs in GHL within 24 hours',
            'Listen to AI phone transcripts daily; follow up same day on commercial inquiries',
            'Monthly strategy call attendance (60 minutes, first Friday of each month)',
            'Video session once per quarter (Q2 onward) for ongoing HeyGen avatar content',
        ]),
    ]

    for category, items in needs:
        p = doc.add_paragraph()
        set_para_bg(p, ASG_TINT)
        p.paragraph_format.left_indent  = Inches(0.12)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(category)
        r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
        r.font.color.rgb = DARK_COLOR
        for item in items:
            bullet(doc, item, size=9.5)
        sp(doc, 3)

    sp(doc, 6)

    # The first action
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(7.5)
    ac = tbl.rows[0].cells[0]
    set_cell_bg(ac, ASG_DARK_HEX)
    ap = ac.paragraphs[0]
    ap.paragraph_format.left_indent  = Inches(0.2)
    ap.paragraph_format.right_indent = Inches(0.2)
    ap.paragraph_format.space_before = Pt(10)
    ap.paragraph_format.space_after  = Pt(4)
    ar1 = ap.add_run('THE FIRST ACTION\n')
    ar1.bold = True; ar1.font.name = 'Calibri'; ar1.font.size = Pt(13)
    ar1.font.color.rgb = ASG_BLUE
    ap2 = ac.add_paragraph()
    ap2.paragraph_format.left_indent  = Inches(0.2)
    ap2.paragraph_format.right_indent = Inches(0.2)
    ap2.paragraph_format.space_before = Pt(0)
    ap2.paragraph_format.space_after  = Pt(10)
    ar2 = ap2.add_run(
        'The fastest way to get the Fast Cash Sprint running is this: send us your past client list. '
        'It doesn\'t need to be organized. A spreadsheet, a text message with names, a list from your '
        'phone contacts, an export from whatever you use for invoicing — anything that has names and '
        'at least one contact method for people who\'ve hired GISCO before. Email it to Roger at '
        '[roger@asg.com] with the subject line: "GISCO — Past Client List." We\'ll take it from there. '
        'Everything else in this Blueprint follows from that first data set.'
    )
    ar2.font.name = 'Calibri'; ar2.font.size = Pt(10.5)
    ar2.font.color.rgb = WHITE

    sp(doc, 6)
    quote_block(doc,
                '"The plan is built. The team is ready. The only variable now is how quickly we move." '
                '— Preston Adler, COO | Authority Systems Group™')
    add_pb(doc)


# ── Section 14: QC Gate Sign-Offs ─────────────────────────────────────────────

def build_s14_qc(doc):
    section_header(doc, '14', 'QC Gate Sign-Offs')
    attribution(doc, 'Authority Systems Group™ — Quality Control Review')
    divider(doc)
    sp(doc, 6)

    gates = [
        {
            'num': 'GATE 1',
            'title': 'STRATEGIC INTEGRITY',
            'reviewer': 'Daniel Frost, CSO',
            'color': ASG_DARK_HEX,
            'items': [
                'Market analysis names all 3 competitors specifically with named weaknesses — not generic observations',
                'Revenue opportunity sizing shows the math (client count × rate × value)',
                'Avatar profiles use the avatar\'s own language — not industry language',
                'Positioning gap is specific enough that no competitor can claim it without 41 years of real history',
                'Authority Statement passes all 3 criteria (competitor cannot copy, avatar resonates, GISCO can deliver)',
                'No sentence in the Executive Summary could appear in a different client\'s Blueprint without rewriting',
            ],
        },
        {
            'num': 'GATE 2',
            'title': 'REVENUE LOGIC',
            'reviewer': 'Tanya Blackwood, CRO',
            'color': SLATE_HEX,
            'items': [
                'All 3 Fast Cash strategies passed the 5-filter selection (niche fit, asset availability, 30-day implementation, compliance, CRM dependency)',
                'Revenue estimates show their math (list size × response rate × avg value)',
                'All strategies reference GISCO-specific data — not generic niche language',
                '90-day calendar is specific enough to hand to an implementer without clarification',
                'No compliance restrictions apply to these strategies for this niche',
                'Sofia Vega\'s voice is operator-level, not consultant-level',
            ],
        },
        {
            'num': 'GATE 3',
            'title': 'CONTENT & MESSAGING',
            'reviewer': 'Vivienne Carr, CMO',
            'color': '276B3D',
            'items': [
                'Content Authority Position is specific — no competitor can claim it without lying about their history',
                'Every content piece has a named stage pair target',
                'Pillar distribution: approximately 25% / 35% / 25% / 15% ✓',
                'Format priorities reflect GISCO\'s avatar\'s actual content consumption habits (search, referral, video)',
                'No content piece is "general awareness" without a specific stage assignment',
                'Email sequences are fully written — no templates or placeholders',
                'Lead magnet TOC follows the Belief-to-Buy arc — every section maps to a stage pair',
                'HeyGen AI avatar strategy is specific and actionable — not a vague "video strategy"',
            ],
        },
        {
            'num': 'GATE 4',
            'title': 'OPERATIONAL FEASIBILITY',
            'reviewer': 'Preston Adler, COO',
            'color': GREEN_HEX,
            'items': [
                'Q1 plan is specific enough to execute without asking ASG for clarification',
                'Every initiative traces to a belief stage from the Belief-to-Buy Map',
                'Every initiative has a success metric',
                'Quarterly sequence has compounding logic (Q2 builds on Q1 infrastructure)',
                'Year-end KPI table has actual estimated baseline numbers — fields marked "To establish Month 1" are honest, not fabricated',
                'Resource requirements from Tony are honest — nothing understated',
                '"What Tony provides" section is specific — not vague requests',
                'Technology stack (AI phone, GHL, HeyGen, new website) is feasible and sequenced correctly',
                'No ABA/bar compliance requirements — home services trade',
            ],
        },
    ]

    for gate in gates:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.columns[0].width = Inches(0.6)
        tbl.columns[1].width = Inches(6.9)
        set_cell_bg(tbl.rows[0].cells[0], gate['color'])
        hc = tbl.rows[0].cells[1]
        set_cell_bg(hc, ASG_DARK_HEX)
        hp = hc.paragraphs[0]
        hp.paragraph_format.left_indent  = Inches(0.12)
        hp.paragraph_format.space_before = Pt(5)
        hp.paragraph_format.space_after  = Pt(1)
        hr1 = hp.add_run(gate['num'] + '  ')
        hr1.bold = True; hr1.font.name = 'Calibri'; hr1.font.size = Pt(9)
        hr1.font.color.rgb = ASG_BLUE
        hr2 = hp.add_run(f'— {gate["title"]}')
        hr2.bold = True; hr2.font.name = 'Calibri'; hr2.font.size = Pt(12)
        hr2.font.color.rgb = WHITE
        hp2 = hc.add_paragraph()
        hp2.paragraph_format.left_indent = Inches(0.12)
        hp2.paragraph_format.space_before = Pt(0)
        hp2.paragraph_format.space_after  = Pt(5)
        hr3 = hp2.add_run(f'Reviewed by: {gate["reviewer"]}')
        hr3.italic = True; hr3.font.name = 'Calibri'; hr3.font.size = Pt(9)
        hr3.font.color.rgb = WHITE
        sp(doc, 2)

        for item in gate['items']:
            bullet(doc, '☑  ' + item, size=9.5)
        sp(doc, 3)

        # PASSED banner
        pb = doc.add_paragraph()
        set_para_bg(pb, GREEN_HEX)
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after  = Pt(2)
        pb.paragraph_format.left_indent  = Inches(0.15)
        pr = pb.add_run(f'{gate["num"]} STATUS: PASSED   — {gate["reviewer"]}')
        pr.bold = True; pr.font.name = 'Calibri'; pr.font.size = Pt(10)
        pr.font.color.rgb = WHITE
        sp(doc, 6)

    sp(doc, 8)

    # Final footer block
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    tbl.columns[0].width = Inches(7.5)
    fc = tbl.rows[0].cells[0]
    set_cell_bg(fc, ASG_DARK_HEX)
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(14)
    fp.paragraph_format.space_after  = Pt(4)
    fr = fp.add_run('Authority Systems Group™')
    fr.bold = True; fr.font.name = 'Calibri'; fr.font.size = Pt(14)
    fr.font.color.rgb = ASG_BLUE
    fp2 = fc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after  = Pt(4)
    fr2 = fp2.add_run('Confidential. Prepared exclusively for Tony Long / GISCO Corp.')
    fr2.font.name = 'Calibri'; fr2.font.size = Pt(10)
    fr2.font.color.rgb = WHITE
    fp3 = fc.add_paragraph()
    fp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp3.paragraph_format.space_before = Pt(0)
    fp3.paragraph_format.space_after  = Pt(6)
    fr3 = fp3.add_run('Director: Roger Bauer  |  Blueprint Date: March 9, 2026  |  Version 1.0')
    fr3.italic = True; fr3.font.name = 'Calibri'; fr3.font.size = Pt(9)
    fr3.font.color.rgb = ASG_CHARCOAL
    fp4 = fc.add_paragraph()
    fp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp4.paragraph_format.space_before = Pt(4)
    fp4.paragraph_format.space_after  = Pt(12)
    fr4 = fp4.add_run('"This is your system now. Let\'s make it work."')
    fr4.bold = True; fr4.italic = True; fr4.font.name = 'Calibri'; fr4.font.size = Pt(11)
    fr4.font.color.rgb = WHITE


# ── MAIN BUILD ─────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_margins(doc)

    build_cover(doc)
    build_toc(doc)
    build_s1_executive_summary(doc)
    build_s2_directors_briefing(doc)
    build_s3_market_intelligence(doc)
    build_s4_belief_to_buy(doc)
    build_s5_services(doc)
    build_s6_fast_cash(doc)
    build_s7_year1_plan(doc)
    build_s8_kpi(doc)
    build_s9_content(doc)
    build_s10_digital(doc)
    build_s11_emails(doc)
    build_s12_lead_magnet(doc)
    build_s13_next_steps(doc)
    build_s14_qc(doc)

    doc.save(OUT_DOCX)
    print(f'✓ Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
