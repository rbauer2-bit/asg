#!/usr/bin/env python3
"""
Authority Systems Group™ — Niche Audience Persona
Local Optometrist: Selective Sam + Specialty-Pivot Sophia
ASG Internal Asset — Niche Intelligence Library
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ───────────────────────────────────────────────────────────────
ASG_BLUE      = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL  = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK_BG   = '1a1a1a'
ASG_BLUE_HEX  = '25aae1'
ASG_LIGHT_BG  = 'F5F5F5'
ASG_TINT      = 'EAF6FC'
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR    = RGBColor(0x33, 0x33, 0x33)
FOOTER_COLOR  = RGBColor(0x88, 0x88, 0x88)

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/Users/Roger/Dropbox/code/authority-systems-group'
LOGO_PATH = os.path.join(BASE, 'brand-standards/assets/asgFull.png')
OUT_DOCX  = os.path.join(BASE, 'outputs/asg_local-optometrist_avatar-persona_20260308.docx')


# ── XML Helpers ────────────────────────────────────────────────────────────────

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_bottom_border(para, color=ASG_BLUE_HEX, sz='6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_left_border(para, color=ASG_BLUE_HEX, sz='24', space='144'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        'Authority Systems Group\u2122  \u2014  Niche Intelligence Library  \u2014  Internal Use Only'
    )
    run.font.size  = Pt(8)
    run.font.color.rgb = FOOTER_COLOR
    run.font.name  = 'Calibri'
    # Page number
    para.add_run('   |   Page ')
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    para.runs[-1]._r.append(fld)
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    r2 = OxmlElement('w:r')
    r2.append(instr)
    para._p.append(r2)
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    r3 = OxmlElement('w:r')
    r3.append(fld2)
    para._p.append(r3)


# ── Typography Helpers ─────────────────────────────────────────────────────────

def h1(doc, text):
    """Section heading — ASG Blue, 20pt, bold, blue bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    add_bottom_border(p, color=ASG_BLUE_HEX, sz='8')
    r = p.add_run(text.upper())
    r.font.name  = 'Calibri'
    r.font.size  = Pt(20)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    return p


def h2(doc, text):
    """Sub-section heading — charcoal, 14pt, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = ASG_CHARCOAL
    return p


def h3(doc, text):
    """Sub-sub heading — ASG Blue, 11pt, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = ASG_BLUE
    return p


def body(doc, text, space_after=6):
    """Standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10.5)
    r.font.color.rgb = BODY_COLOR
    return p


def bullet(doc, text, level=0):
    """Bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10.5)
    r.font.color.rgb = BODY_COLOR
    return p


def callout(doc, text, bg=ASG_TINT, border_color=ASG_BLUE_HEX):
    """Callout / highlighted block — light blue background, left border."""
    p = doc.add_paragraph()
    set_para_bg(p, bg)
    add_left_border(p, color=border_color, sz='20', space='120')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.15)
    r = p.add_run(text)
    r.font.name  = 'Calibri'
    r.font.size  = Pt(10.5)
    r.font.color.rgb = ASG_CHARCOAL
    return p


def label_value(doc, label, value):
    """Inline label: value line — label in blue bold, value in charcoal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    rl = p.add_run(label + '  ')
    rl.font.name  = 'Calibri'
    rl.font.size  = Pt(10.5)
    rl.font.bold  = True
    rl.font.color.rgb = ASG_BLUE
    rv = p.add_run(value)
    rv.font.name  = 'Calibri'
    rv.font.size  = Pt(10.5)
    rv.font.color.rgb = BODY_COLOR
    return p


def divider(doc):
    """Thin blue horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    add_bottom_border(p, color=ASG_BLUE_HEX, sz='4')
    r = p.add_run('')
    return p


def belief_block(doc, stage, belief, why_it_stalls, reframe, note=None):
    """Formatted Belief Track block."""
    # Stage badge
    p_stage = doc.add_paragraph()
    p_stage.paragraph_format.space_before = Pt(10)
    p_stage.paragraph_format.space_after  = Pt(2)
    rs = p_stage.add_run(f'  {stage}  ')
    rs.font.name  = 'Calibri'
    rs.font.size  = Pt(8)
    rs.font.bold  = True
    rs.font.color.rgb = WHITE
    set_para_bg(p_stage, ASG_BLUE_HEX)

    # Enemy belief
    p_belief = doc.add_paragraph()
    p_belief.paragraph_format.space_before = Pt(2)
    p_belief.paragraph_format.space_after  = Pt(3)
    rb = p_belief.add_run(f'\u201c{belief}\u201d')
    rb.font.name   = 'Calibri'
    rb.font.size   = Pt(11)
    rb.font.bold   = True
    rb.font.italic = True
    rb.font.color.rgb = ASG_CHARCOAL

    # Why it stalls
    p_why = doc.add_paragraph()
    p_why.paragraph_format.space_before = Pt(0)
    p_why.paragraph_format.space_after  = Pt(4)
    rw_label = p_why.add_run('Why it stalls:  ')
    rw_label.font.name  = 'Calibri'
    rw_label.font.size  = Pt(10)
    rw_label.font.bold  = True
    rw_label.font.color.rgb = ASG_CHARCOAL
    rw = p_why.add_run(why_it_stalls)
    rw.font.name  = 'Calibri'
    rw.font.size  = Pt(10)
    rw.font.color.rgb = BODY_COLOR

    # Reframe
    callout(doc, f'Reframe: {reframe}')

    if note:
        p_note = doc.add_paragraph()
        set_para_bg(p_note, 'FFF8E1')
        add_left_border(p_note, color='D4A844', sz='16', space='100')
        p_note.paragraph_format.space_before = Pt(2)
        p_note.paragraph_format.space_after  = Pt(6)
        p_note.paragraph_format.left_indent  = Inches(0.12)
        rn_label = p_note.add_run('ASG Handling Note:  ')
        rn_label.font.name  = 'Calibri'
        rn_label.font.size  = Pt(10)
        rn_label.font.bold  = True
        rn_label.font.color.rgb = RGBColor(0xB8, 0x86, 0x00)
        rn = p_note.add_run(note)
        rn.font.name  = 'Calibri'
        rn.font.size  = Pt(10)
        rn.font.color.rgb = RGBColor(0x5C, 0x45, 0x00)


def objection_block(doc, objection, fear, reframe):
    """Formatted objection block."""
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_before = Pt(10)
    p_obj.paragraph_format.space_after  = Pt(3)
    ro = p_obj.add_run(f'\u201c{objection}\u201d')
    ro.font.name   = 'Calibri'
    ro.font.size   = Pt(11)
    ro.font.bold   = True
    ro.font.italic = True
    ro.font.color.rgb = ASG_CHARCOAL

    p_fear = doc.add_paragraph()
    p_fear.paragraph_format.space_before = Pt(0)
    p_fear.paragraph_format.space_after  = Pt(4)
    rf_label = p_fear.add_run('Underlying fear:  ')
    rf_label.font.name  = 'Calibri'
    rf_label.font.size  = Pt(10)
    rf_label.font.bold  = True
    rf_label.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    rf = p_fear.add_run(fear)
    rf.font.name  = 'Calibri'
    rf.font.size  = Pt(10)
    rf.font.color.rgb = BODY_COLOR

    callout(doc, f'Response: {reframe}')


def two_col_table(doc, rows, col_widths=(3000, 5500)):
    """Simple 2-column table: left label (blue), right value (charcoal)."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (label, value) in enumerate(rows):
        row = tbl.rows[i]
        bg = ASG_TINT if i % 2 == 0 else 'FFFFFF'
        # Left cell
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        c0.width = Inches(col_widths[0] / 1440)
        p0 = c0.paragraphs[0]
        p0.clear()
        r0 = p0.add_run(label)
        r0.font.name  = 'Calibri'
        r0.font.size  = Pt(10)
        r0.font.bold  = True
        r0.font.color.rgb = ASG_BLUE
        c0.paragraphs[0].paragraph_format.space_before = Pt(3)
        c0.paragraphs[0].paragraph_format.space_after  = Pt(3)
        # Right cell
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        c1.width = Inches(col_widths[1] / 1440)
        p1 = c1.paragraphs[0]
        p1.clear()
        r1 = p1.add_run(value)
        r1.font.name  = 'Calibri'
        r1.font.size  = Pt(10)
        r1.font.color.rgb = BODY_COLOR
        c1.paragraphs[0].paragraph_format.space_before = Pt(3)
        c1.paragraphs[0].paragraph_format.space_after  = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Cover Page ─────────────────────────────────────────────────────────────────

def add_cover(doc):
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin   = Inches(0)
    section.right_margin  = Inches(0)

    def dk(text, size, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4):
        p = doc.add_paragraph()
        set_para_bg(p, ASG_DARK_BG)
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(text)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = color or WHITE
        return p

    # Top spacer
    for _ in range(6):
        dk('', 6)

    # Logo or wordmark
    p_logo = doc.add_paragraph()
    set_para_bg(p_logo, ASG_DARK_BG)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(6)
    if os.path.exists(LOGO_PATH):
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))
    else:
        r = p_logo.add_run('AUTHORITY SYSTEMS GROUP\u2122')
        r.font.name  = 'Calibri'
        r.font.size  = Pt(22)
        r.font.bold  = True
        r.font.color.rgb = ASG_BLUE

    dk('Positioning You As The Authority', 11, color=RGBColor(0x80, 0xD0, 0xF0), sa=2)

    # Rule
    p_rule = doc.add_paragraph()
    set_para_bg(p_rule, ASG_DARK_BG)
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_before = Pt(16)
    p_rule.paragraph_format.space_after  = Pt(2)
    r_rule = p_rule.add_run('\u2500' * 50)
    r_rule.font.color.rgb = ASG_BLUE
    r_rule.font.size = Pt(10)

    dk('', 8, sa=2)
    dk('NICHE INTELLIGENCE LIBRARY', 10, color=RGBColor(0x99, 0x99, 0x99), sa=2)
    dk('', 4, sa=2)
    dk('AUDIENCE PERSONA', 30, bold=True, color=WHITE, sa=4)

    p_niche = doc.add_paragraph()
    set_para_bg(p_niche, ASG_DARK_BG)
    p_niche.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_niche.paragraph_format.space_before = Pt(0)
    p_niche.paragraph_format.space_after  = Pt(2)
    rn = p_niche.add_run('Local Optometrist')
    rn.font.name  = 'Calibri'
    rn.font.size  = Pt(20)
    rn.font.bold  = False
    rn.font.color.rgb = ASG_BLUE

    dk('Independent Practice \u2014 Premium Patient Growth', 13, color=RGBColor(0xCC, 0xCC, 0xCC), sa=2)

    dk('', 12, sa=2)

    # Dot accent
    p_dot = doc.add_paragraph()
    set_para_bg(p_dot, ASG_DARK_BG)
    p_dot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dot = p_dot.add_run('\u25cf  \u25cf  \u25cf')
    r_dot.font.color.rgb = ASG_BLUE
    r_dot.font.size = Pt(12)
    p_dot.paragraph_format.space_before = Pt(0)
    p_dot.paragraph_format.space_after  = Pt(12)

    dk('Prepared by Vivienne Carr, CMO + Daniel Frost, CSO', 10, color=RGBColor(0xAA, 0xAA, 0xAA))
    dk('Authority Systems Group\u2122  \u2014  March 2026  \u2014  Internal Use Only', 9, color=RGBColor(0x77, 0x77, 0x77))

    # Bottom filler
    for _ in range(8):
        dk('', 6)

    page_break(doc)


# ── Body Content ───────────────────────────────────────────────────────────────

def add_body(doc):

    # Reset page margins for body
    section2 = doc.add_section()
    section2.page_width    = Inches(8.5)
    section2.page_height   = Inches(11)
    section2.top_margin    = Inches(1.0)
    section2.bottom_margin = Inches(1.0)
    section2.left_margin   = Inches(1.25)
    section2.right_margin  = Inches(1.25)

    # Add header
    hdr = section2.header
    hdr.is_linked_to_previous = False
    hdr_para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hdr_para.clear()
    if os.path.exists(LOGO_PATH):
        run = hdr_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.4))
    else:
        r_hdr = hdr_para.add_run('AUTHORITY SYSTEMS GROUP\u2122')
        r_hdr.font.name  = 'Calibri'
        r_hdr.font.size  = Pt(11)
        r_hdr.font.bold  = True
        r_hdr.font.color.rgb = ASG_BLUE
    # Tab to right for document title
    pPr = hdr_para._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '8280')
    tabs_el.append(tab_el)
    pPr.append(tabs_el)
    r_title = hdr_para.add_run('\tAudience Persona — Local Optometrist')
    r_title.font.name  = 'Calibri'
    r_title.font.size  = Pt(9)
    r_title.font.color.rgb = ASG_CHARCOAL
    add_bottom_border(hdr_para, color=ASG_BLUE_HEX, sz='6')

    # Add footer
    ftr = section2.footer
    ftr.is_linked_to_previous = False
    ftr_para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    ftr_para.clear()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ftr = ftr_para.add_run(
        'Authority Systems Group\u2122  \u2014  Niche Intelligence Library  \u2014  Internal Use Only'
    )
    r_ftr.font.name  = 'Calibri'
    r_ftr.font.size  = Pt(8)
    r_ftr.font.color.rgb = FOOTER_COLOR
    add_bottom_border(ftr_para, color='CCCCCC', sz='2')

    # ── NICHE OVERVIEW ───────────────────────────────────────────────────────────
    h1(doc, 'Niche Overview')

    body(doc,
        'The independent optometrist niche requires precision positioning — not a lead generation '
        'campaign. The OD is not invisible. They are miscategorized. The market sees them as interchangeable '
        'with every LensCrafters, Walmart Vision, and online eyewear retailer in the zip code.',
        space_after=6)
    body(doc,
        'Patients who don\u2019t know better make decisions on price, convenience, and insurance coverage. '
        'The patients this OD wants \u2014 those who value clinical judgment, buy premium optical, and refer '
        'their families \u2014 are already in the market. They\u2019re just not finding this practice.',
        space_after=6)

    callout(doc,
        'The goal: become the recognized eye care authority in their community \u2014 the doctor that quality-minded '
        'patients seek out pre-sold, pre-trusting, and already committed to buying their optical in-office '
        'rather than taking the prescription to Zenni.')

    h2(doc, 'Critical Engagement Notes')
    bullet(doc,
        'This avatar RESISTS content creation and automation. Both are essential \u2014 but must be '
        'reframed and de-jargoned at every touchpoint.')
    bullet(doc,
        '\u201cContent\u201d becomes \u201cpatient education assets.\u201d '
        '\u201cAutomation\u201d becomes \u201cpatient care systems.\u201d')
    bullet(doc,
        'The OD is the approver and expert \u2014 never the content machine. Position ASG as the team '
        'that translates their clinical expertise into the assets that attract quality patients.')
    bullet(doc,
        'This avatar has been pitched by 4\u20136 consultants. Lead with outcomes, not activities. '
        'Never open with clicks, impressions, or new patient volume.')

    page_break(doc)

    # ── PRIMARY AVATAR ───────────────────────────────────────────────────────────
    h1(doc, 'Primary Avatar — \u201cSelective Sam\u201d')

    callout(doc,
        'He built a premium practice. His marketing keeps attracting the wrong patients.',
        bg='1a1a1a')
    # White text on dark callout — rebuild manually
    doc.paragraphs[-1].runs[0].font.color.rgb = ASG_BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    two_col_table(doc, [
        ('Avatar Name',        'Selective Sam'),
        ('Niche',              'Independent Optometrist'),
        ('Age Range',          '35\u201354'),
        ('Gender Skew',        'Majority female in new ODs (\u224867%) \u2014 use gender-neutral language throughout'),
        ('Gross Collections',  '$450,000\u2013$1,200,000'),
        ('Net Take-Home',      '$130,000\u2013$320,000'),
        ('Years in Practice',  '7\u201322 years'),
        ('Geography',          'Suburban markets and mid-sized cities with enough density, not yet chain-dominated'),
    ])

    # Income gap
    h2(doc, 'The Income Gap')
    body(doc,
        'Sam sees 14 patients today. He makes $7,200 gross. After overhead, he takes home about $2,200. '
        'The patient who bought glasses from Warby Parker with his prescription? That was $600 Sam didn\u2019t '
        'collect \u2014 for a patient he already spent 40 minutes with.',
        space_after=6)
    callout(doc,
        'If just 3 more patients per day bought their optical in-office, Sam\u2019s take-home would increase '
        'by approximately $80,000 annually. That math is always in the background.')

    # Current situation
    h2(doc, 'Daily Reality')
    body(doc,
        'Sam\u2019s 9am patient arrives, gets a thorough exam, leaves with a handwritten prescription for '
        'progressive lenses. She smiles and says \u201cthank you, I\u2019ll think about the glasses.\u201d '
        'He knows she\u2019s going home to Zenni or Warby Parker. He also knows the premium Zeiss lenses '
        'he would have fitted her with would have been life-changing for her and $480 for his practice. '
        'He sees this 4\u20136 times a day. He\u2019s stopped trying to figure out how to stop it.',
        space_after=6)

    h2(doc, 'The Core Tension')
    body(doc,
        'Sam\u2019s clinical skills are not in question. His patients trust him in the exam room. The problem '
        'is what happens before they arrive and after they leave.',
        space_after=4)
    body(doc,
        'Before: The patients who find him are often price-driven, insurance-first, and willing to go to whoever '
        'takes their plan. They don\u2019t have a reason to seek him out specifically until they\u2019re already '
        'sitting in his chair \u2014 and by then, the relationship is transactional.',
        space_after=4)
    body(doc,
        'After: Without a system to stay connected, the patient who had a great experience simply doesn\u2019t '
        'think to refer, doesn\u2019t come back until they \u201cneed new glasses,\u201d and doesn\u2019t know '
        'about the dry eye treatment or myopia management protocol that would genuinely improve their child\u2019s life.',
        space_after=6)
    callout(doc,
        'The gap is not clinical. It\u2019s the authority and relationship infrastructure that makes quality '
        'patients choose Sam, stay with Sam, and send everyone they know.')

    # What keeps him stuck
    h2(doc, 'What Keeps Sam Stuck')
    stuck_items = [
        'Insurance contracts pull in price-motivated patients \u2014 but dropping plans feels like financial suicide',
        'Tried Google Ads and got calls from people asking \u201cdo you take [discount plan]?\u201d Not what he wanted',
        'Tried a local SEO agency. Got on page one. For \u201ceye exam near me.\u201d Wrong patient.',
        'His current marketing attracts patients who compare him to Walmart Vision \u2014 a comparison he should never be in',
        'Spent 8 years earning his degree \u2014 he\u2019s not making TikToks to explain it',
        'Knows he should be doing something about patient retention, but doesn\u2019t have time to manually follow up',
        'Has been pitched by 4\u20136 consultants in the last 3 years \u2014 every pitch led with \u201cmore traffic\u201d and \u201cmore clicks\u201d',
        'Word of mouth is his best patient source \u2014 but he has no system to engineer it',
        'His website is functional but indistinguishable from every other optometry site in his market',
    ]
    for item in stuck_items:
        bullet(doc, item)

    page_break(doc)

    # ── GOALS ────────────────────────────────────────────────────────────────────
    h1(doc, 'Goals & Desires')

    h2(doc, 'Primary Goals')
    primary_goals = [
        'Replace 4\u20136 price-shopping patients per day with 2\u20133 quality patients who buy optical in-office',
        'Grow premium optical conversion rate from 40\u201355% to 70\u201380% of exam patients',
        'Build a patient base that comes in pre-trusting him \u2014 not comparing him to LensCrafters',
        'Add 1\u20132 dry eye treatment cases per week without adding chair time',
        'Be the optometrist in his community that patients specifically request \u2014 not \u201cwhoever takes my insurance\u201d',
        'Engineer referrals without having to personally follow up or feel like he\u2019s begging',
    ]
    for g in primary_goals:
        bullet(doc, g)

    h2(doc, 'Secondary Goals')
    secondary_goals = [
        'Develop enough authority to selectively drop the most punishing insurance plans',
        'Attract pediatric patients for myopia management \u2014 the highest-LTV patient type in optometry',
        'Build a practice with enough brand equity that it sells at a premium valuation',
        'Stop being squeezed by online optical retailers and insurance reimbursement schedules simultaneously',
        'Have staff who handle patient communication without it falling on him personally',
    ]
    for g in secondary_goals:
        bullet(doc, g)

    h2(doc, 'The Revenue Math')
    body(doc,
        'Sam doesn\u2019t need more patients. He needs more value per patient. Right now, his practice captures '
        '35\u201345 cents of every dollar his patients spend on vision care. The other 55\u201365 cents goes to '
        'Zenni, 1-800-Contacts, and Costco.',
        space_after=4)
    two_col_table(doc, [
        ('Current revenue/day',      '$5,500\u2013$7,500 gross'),
        ('Target revenue/day',       '$8,500\u2013$12,000 gross (fewer exams, higher per-patient value)'),
        ('Key lever 1',              '+3 premium optical sales/day = +$800\u2013$1,400/day'),
        ('Key lever 2',              '+2 dry eye cases/week @ $1,500 avg = +$156,000/year'),
        ('Key lever 3',              '5 myopia management enrollments = +$7,500\u2013$15,000 recurring/year'),
    ])

    page_break(doc)

    # ── BELIEF-TO-BUY FRAMEWORK ───────────────────────────────────────────────────
    h1(doc, 'Belief-to-Buy Framework\u2122')
    body(doc,
        'Both tracks must be moved simultaneously. Shifting only the Belief Track or only the Emotion Track '
        'stalls the prospect. Every content piece, conversation, and touchpoint maps to a specific stage pair.',
        space_after=10)

    h2(doc, 'BELIEF TRACK \u2014 What Sam Thinks Is True (and Must Shift)')

    belief_block(doc,
        stage='Enemy Belief 1',
        belief='Marketing just attracts price shoppers. I\u2019ve tried it and it brings in the wrong patients.',
        why_it_stalls=(
            'Sam has conflated volume marketing \u2014 designed to fill chairs with any patient \u2014 '
            'with authority positioning, which pre-selects for quality patients. Every agency he hired ran '
            'the former and called it the latter. His conclusion is reasonable. It\u2019s aimed at the '
            'wrong category of marketing.'
        ),
        reframe=(
            'Draw the line between volume marketing (clicks, impressions, new patient volume) and authority '
            'positioning (pre-sold patients who already trust you before they call). The type of patient you '
            'attract is determined by what you say, not how loud you say it. Sam isn\u2019t attracting price '
            'shoppers because he\u2019s marketing. He\u2019s attracting them because his positioning looks '
            'identical to a practice that competes on price.'
        )
    )

    belief_block(doc,
        stage='Enemy Belief 2',
        belief='Good patients come from referrals. I can\u2019t really control that.',
        why_it_stalls=(
            'Sam has correctly identified the channel and incorrectly concluded it\u2019s unengineerable. '
            'He treats word-of-mouth as a weather event: it happens when it happens. The result is '
            'feast-or-famine quality patient flow with no system to accelerate or sustain it.'
        ),
        reframe=(
            'Referrals are engineerable. They happen more often when patients have a specific reason to refer \u2014 '
            'a memorable outcome, a clear story of what makes this practice different. Authority positioning gives '
            'existing patients the language and the prompt to refer. It doesn\u2019t ask them to be salespeople. '
            'It gives them a story worth telling.'
        )
    )

    belief_block(doc,
        stage='Enemy Belief 3',
        belief='I can\u2019t drop insurance plans \u2014 I\u2019d lose too many patients.',
        why_it_stalls=(
            'Insurance dependency is the cage. As long as Sam believes his revenue floor requires insurance '
            'volume, he will keep attracting the patient profile that insurance delivers: price-first, plan-first, '
            'loyalty-last. This belief is partially accurate short-term and completely limiting long-term.'
        ),
        reframe=(
            'Authority positioning is the bridge out of insurance dependency \u2014 not a cliff. When enough '
            'patients choose Sam specifically (not because he\u2019s on their plan, but because he\u2019s the '
            'trusted authority), he can selectively exit the worst-paying plans without revenue impact. '
            'The sequence matters: position first, exit second. No one successfully drops insurance without '
            'first building a reason patients will pay.'
        )
    )

    belief_block(doc,
        stage='Enemy Belief 4',
        belief='I\u2019m not going to record videos or post on social media every day. That\u2019s not who I am.',
        why_it_stalls=(
            'This is the most common authority-building objection in the OD niche \u2014 and the most valid. '
            'Sam is a clinician, not a content creator. If authority building requires him to become a TikTok '
            'personality, it will never happen. He doesn\u2019t have the time, the inclination, or the identity '
            'for it. Forcing a content-creation model on this avatar kills the engagement before it starts.'
        ),
        reframe=(
            'Authority is not the same as visibility. You don\u2019t need to be on social media every day. '
            'You need your expertise to be legible in the moments when quality patients are researching and '
            'deciding. Your clinical knowledge becomes patient education assets \u2014 written, formatted, and '
            'published by ASG\u2019s content team, reviewed and approved by you. You are the expert whose name '
            'is on the content. You are not the content machine.'
        ),
        note=(
            'NEVER pitch this avatar with \u201ccontent calendar,\u201d \u201csocial media management,\u201d '
            'or \u201cposting strategy.\u201d Frame all content as clinical authority assets and patient '
            'education. Lead with outcomes (patients who arrive pre-trusting, pre-qualified) not activities '
            '(posts per week, video frequency). Sam doesn\u2019t care about the output \u2014 he cares about '
            'which patients walk through the door.'
        )
    )

    belief_block(doc,
        stage='Enemy Belief 5',
        belief='Automation is impersonal. My patients know me. I don\u2019t want to feel like a robot sending them emails.',
        why_it_stalls=(
            'Sam\u2019s concern is legitimate. Bad patient automation is impersonal, generic, and erodes '
            'the trust he has built. He has probably received bad examples of it \u2014 generic appointment '
            'reminders, mass-blast emails, \u201cDear Patient\u201d letters. His instinct to protect patient '
            'relationships from that kind of damage is exactly right.'
        ),
        reframe=(
            'The alternative to systematized patient care is not personal care \u2014 it\u2019s no care. '
            'Patients who don\u2019t hear from Sam\u2019s practice after their exam simply drift. They go to '
            'LensCrafters next year because it was convenient. A well-built patient care system \u2014 one that '
            'sounds like it came from Sam\u2019s practice, uses the patient\u2019s name, references their actual '
            'visit, and is timed to moments that matter \u2014 is not a robot. It is the thoughtful practice '
            'Sam runs, extended beyond his chair time.'
        ),
        note=(
            'Never use the word \u201cautomation\u201d with this avatar. Use \u201cpatient care systems,\u201d '
            '\u201crelationship infrastructure,\u201d or \u201cyour practice staying in touch \u2014 without it '
            'falling on you personally.\u201d Lead with the patient experience, not the mechanism.'
        )
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    callout(doc,
        'New Worldview Target: \u201cI don\u2019t need more patients. I need the right patients \u2014 the ones '
        'who value clinical expertise, trust my optical recommendations, refer their family, and come back every year. '
        'Those patients exist in my market. I need them to find me first, not stumble across me while comparing '
        'prices. That\u2019s a positioning problem. And positioning is solvable.\u201d')

    page_break(doc)

    h2(doc, 'EMOTION TRACK \u2014 What Sam Feels Is True (and Must Shift)')

    emotion_stages = [
        ('DISSATISFACTION', ASG_BLUE_HEX, 'FFFFFF',
         'Sam\u2019s 2pm slot just canceled. Again. He knows her \u2014 she came in last year, liked him, but '
         'chose Costco this year because it was \u201cmore convenient.\u201d He didn\u2019t have a reason in her '
         'mind to come back to him specifically. His schedule has 4 empty slots this week. His overhead doesn\u2019t. '
         'He\u2019s quietly frustrated. Not panicked \u2014 his practice is fine. But fine isn\u2019t the goal. '
         'He suspects the patients who would value what he does are out there. They\u2019re just walking past '
         'his practice on the way to somewhere else.'),
        ('CONTRAST', '58585a', 'FFFFFF',
         'A patient comes in and says: \u201cMy colleague told me you\u2019re the person to see for dry eye. '
         'She said you actually fixed it \u2014 not just sold her drops.\u201d The patient books three sessions. '
         'Spends $1,800. Sends her husband the next month. Sam walks out thinking: \u201cThat\u2019s what I want '
         'more of. That patient came in pre-sold. I didn\u2019t have to explain myself at all.\u201d He\u2019s '
         'felt this before. He can\u2019t figure out how to engineer it.'),
        ('DESIRE', ASG_BLUE_HEX, 'FFFFFF',
         'A schedule with 10 patients who chose him specifically. Patients who ask for his recommendation on '
         'frames \u2014 and take it. A front desk not fielding calls from \u201cwhat does my insurance cover.\u201d '
         'Enough referral inflow that he\u2019s selective about which new patients he accepts. Walking out by 4pm '
         'having made $9,500 gross. Being introduced at a community event as \u201cthe eye doctor everyone recommends.\u201d'),
        ('URGENCY', '8B0000', 'FFFFFF',
         'Warby Parker just opened two miles from his office. A PE-backed regional chain is expanding into his '
         'market. His optical revenue is down 11% year-over-year. Three of his best patients moved to the new '
         'ophthalmology group that \u201cdoes everything under one roof.\u201d He is not in a crisis. But he can '
         'see it from here. Doing nothing is a slow erosion, not a stable floor.'),
        ('RELIEF', '1a6e1a', 'FFFFFF',
         'The version of Sam\u2019s practice that authority positioning builds doesn\u2019t compete with Warby '
         'Parker or Walmart Vision \u2014 it doesn\u2019t occupy the same category in patients\u2019 minds. '
         'Premium patients seek it out. The optical conversion rate is 72%. Dry eye patients arrive having already '
         'decided. Myopia management families are on a waitlist. Sam works 4 days a week, sees 10 patients a day, '
         'and nets more than he does now seeing 16. That is not a fantasy. It is a positioning problem with a solution.'),
    ]

    for stage_name, bg, fg, text in emotion_stages:
        # Stage header
        p_s = doc.add_paragraph()
        set_para_bg(p_s, bg)
        p_s.paragraph_format.space_before = Pt(10)
        p_s.paragraph_format.space_after  = Pt(3)
        rs = p_s.add_run(f'  {stage_name}  ')
        rs.font.name  = 'Calibri'
        rs.font.size  = Pt(9)
        rs.font.bold  = True
        rs.font.color.rgb = RGBColor(int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
        body(doc, text, space_after=4)

    page_break(doc)

    # ── WHAT SAM HAS TRIED ────────────────────────────────────────────────────────
    h1(doc, 'What Sam Has Already Tried')
    body(doc,
        'This avatar has a history of failed marketing engagements. Understanding this context is critical '
        'before any sales conversation. Each prior solution left a residual belief that shapes how Sam hears '
        'every new pitch.',
        space_after=8)

    tried = [
        ('Local SEO Agency',
         'Ranked higher for \u201ceye exam near me.\u201d Got more insurance-first calls.',
         '\u201cSEO brings in whoever is searching for the cheapest exam \u2014 not quality patients.\u201d',
         'SEO targeting was wrong. \u201cDry eye treatment [city]\u201d and \u201cbest optometrist [city]\u201d attract '
         'different patients than \u201ceye exam near me.\u201d Authority positioning changes what you rank for, '
         'not just how high.'),
        ('Google Ads / PPC',
         'Spent $600\u2013$1,500/month. Calls from people asking \u201cdo you take [discount vision plan]?\u201d',
         '\u201cAds attract price shoppers \u2014 not the patients I want.\u201d',
         None),
        ('Facebook / Instagram Ads',
         'Ran promotions. Got some new patients. Most bought one pair and never came back.',
         '\u201cSocial media brings in one-timers \u2014 they don\u2019t become loyal patients.\u201d',
         None),
        ('New Website',
         'Looks professional. Still says \u201ccomprehensive eye care for the whole family.\u201d Doesn\u2019t differentiate.',
         '\u201cThe website is fine \u2014 the problem is something else.\u201d',
         'The website is not fine. It is indistinguishable from every other practice in the market.'),
        ('Practice Software Recall System',
         'Sends appointment reminders. Generic. Patients ignore or unsubscribe.',
         '\u201cAutomated messages feel impersonal \u2014 patients know it\u2019s a system, not a person.\u201d',
         None),
        ('Hiring Front Desk to \u201cDo Social Media\u201d',
         'Posted a few times. Generic eye health awareness posts. No discernible patient impact.',
         '\u201cSocial media doesn\u2019t really move the needle for a local practice like mine.\u201d',
         None),
    ]

    for solution, result, residual, note in tried:
        h3(doc, solution)
        label_value(doc, 'Result:', result)
        label_value(doc, 'Residual belief:', residual)
        if note:
            p_note = doc.add_paragraph()
            set_para_bg(p_note, 'FFF8E1')
            add_left_border(p_note, color='D4A844', sz='16', space='100')
            p_note.paragraph_format.space_before = Pt(2)
            p_note.paragraph_format.space_after  = Pt(8)
            p_note.paragraph_format.left_indent  = Inches(0.12)
            r_note = p_note.add_run(f'Note: {note}')
            r_note.font.name  = 'Calibri'
            r_note.font.size  = Pt(10)
            r_note.font.color.rgb = RGBColor(0x5C, 0x45, 0x00)

    page_break(doc)

    # ── OBJECTIONS ───────────────────────────────────────────────────────────────
    h1(doc, 'Objections & Reframes')
    body(doc,
        'These objections appear consistently in the first and second sales conversation. '
        'Each must be handled without dismissing the underlying concern \u2014 because each concern is legitimate. '
        'The goal is reframe, not rebut.',
        space_after=8)

    objections = [
        (
            "I\u2019ve already spent money on marketing and it brought in the wrong patients.",
            "I\u2019ll spend another $20,000 and still be stuck with the same problem.",
            "What didn\u2019t work was positioning you as a commodity competing on price and convenience \u2014 exactly "
            "what SEO-for-exam-volume and generic ad campaigns do. Authority positioning pre-selects for patients making "
            "a trust-based decision, not a price-based one. Before they call, they already know you\u2019re the right choice. "
            "The comparison category changes. The patient type changes with it."
        ),
        (
            "I don\u2019t have time to create content or manage a social media presence.",
            "This is going to become another thing I have to do personally.",
            "You won\u2019t be creating content. You\u2019ll be approving it. Your clinical expertise \u2014 your philosophy "
            "about dry eye, about myopia in children, about why certain lenses matter \u2014 becomes the authority content "
            "that your team writes and publishes under your name. Your role is 20 minutes of review and sign-off. "
            "The content engine runs without requiring you to be in it."
        ),
        (
            "I don\u2019t want to do heavy follow-up with patients. That\u2019s not my style.",
            "I\u2019ll feel like I\u2019m pestering people or being a salesperson.",
            "Patients who matter don\u2019t experience well-timed, relevant communication as pestering \u2014 they experience "
            "it as the practice remembering them. A note that says \u201cYour daughter\u2019s myopia has been stable for "
            "8 months \u2014 here\u2019s what that means\u201d is not a sales pitch. It\u2019s clinical care delivered "
            "outside the exam room. The difference between annoying follow-up and appreciated communication is relevance "
            "and timing. We build the latter."
        ),
        (
            "I don\u2019t want to automate my patient relationships. It feels impersonal.",
            "My patients know me personally. I don\u2019t want them to feel like they\u2019re just a number.",
            "The question isn\u2019t automation vs. personal care. It\u2019s: what happens when there is no system at all? "
            "Right now, 40% of your patients don\u2019t come back the following year \u2014 not because they had a bad "
            "experience, because no one from your practice reached out in a way that mattered to them. A well-built patient "
            "care system doesn\u2019t replace your personal care. It extends it to the 8,760 hours of the year you\u2019re "
            "not in the exam room."
        ),
        (
            "I\u2019ve been pitched by a lot of marketing consultants. What makes you different?",
            "Another agency that promises quality patients and delivers clicks and followers.",
            "Every agency you\u2019ve worked with was optimizing for volume metrics \u2014 traffic, clicks, impressions, "
            "new patients. We don\u2019t measure any of those things. We measure patient quality signals: in-office optical "
            "conversion rate, LTV per patient, specialty case uptake, and referral source concentration. Our work is "
            "successful when the patients who find you are already convinced before they call. That is a different goal, "
            "built with a different system."
        ),
        (
            "My best patients come from referrals. I just need more of that.",
            "I know what works \u2014 I just can\u2019t figure out how to get more of it.",
            "You\u2019re right that referrals are the best source. You\u2019re missing the system that makes them happen "
            "more often and more specifically. Right now, referrals happen because a patient had a great experience and "
            "happened to mention it. The system we build makes those moments predictable: a post-dry-eye-treatment follow-up "
            "that invites patients to share, a pediatric myopia welcome kit parents naturally share with other parents, "
            "specific touchpoints that give satisfied patients a reason and a mechanism to refer. Engineering the referral "
            "doesn\u2019t make it feel less genuine. It makes it happen more often."
        ),
    ]

    for obj, fear, reframe in objections:
        objection_block(doc, obj, fear, reframe)

    page_break(doc)

    # ── LANGUAGE GUIDE ─────────────────────────────────────────────────────────
    h1(doc, 'Language Guide')

    h2(doc, 'Industry Terms Sam Uses')
    industry_terms = [
        ('Chair time', 'The scarce resource \u2014 Sam measures everything in chair time'),
        ('Optical conversion / capture rate', '% of exam patients who buy optical in-office'),
        ('Recall', 'Patient retention / annual re-examination scheduling'),
        ('Medical billing', 'Billing eye disease to medical insurance \u2014 higher reimbursement'),
        ('DED / dry eye disease', 'The specialty diagnosis \u2014 not \u201cdry eyes\u201d'),
        ('Myopia management / Ortho-K / MiSight', 'Specialty pediatric services'),
        ('Premium progressives / free-form / digital surfaced', 'High-margin lens category'),
        ('AR', 'Anti-reflective coating \u2014 primary upgrade conversation'),
        ('Complete pair', 'Frame + lenses as a bundle \u2014 optical retail framing'),
        ('VSP / EyeMed / Davis', 'The major vision insurance plans \u2014 he loves to hate them'),
        ('Frame board', 'The optical display \u2014 his biggest retail variable'),
        ('New patient mix', 'The ratio of quality patients to insurance-first patients'),
        ('LTV / lifetime value', 'Sam thinks in these terms, especially after being burned by volume'),
    ]
    two_col_table(doc, industry_terms, col_widths=(2800, 5700))

    h2(doc, 'How Sam Talks About the Problem')
    how_he_talks = [
        '\u201cI\u2019m tired of seeing patients who just want the cheapest glasses.\u201d',
        '\u201cMy best patients send me referrals \u2014 I just can\u2019t figure out how to get more of them.\u201d',
        '\u201cI keep losing optical sales to Warby Parker and Zenni.\u201d',
        '\u201cI want patients who trust my recommendations \u2014 not patients who just want to use their benefits.\u201d',
        '\u201cMy practice is good. The market just doesn\u2019t see it that way yet.\u201d',
        '\u201cI\u2019ve tried marketing and it just brings in price shoppers.\u201d',
        '\u201cI don\u2019t want to be on social media \u2014 I want my work to speak for itself.\u201d',
    ]
    for item in how_he_talks:
        bullet(doc, item)

    h2(doc, '\u26a0  What Not to Say to Sam')
    do_not_say = [
        'Do NOT lead with \u201cnew patient volume\u201d \u2014 he wants quality, not quantity',
        'Do NOT use the word \u201cautomation\u201d without immediately reframing as patient care systems',
        'Do NOT pitch social media management or content calendars as the lead value',
        'Do NOT lead with clicks, impressions, or website traffic \u2014 he has been burned by these metrics',
        'Do NOT compare him to LensCrafters or corporate chains \u2014 the comparison offends him',
        'Do NOT pitch follow-up sequences as \u201clead nurturing\u201d \u2014 he is not running a sales funnel; he is a doctor',
        'Do NOT be aggressive or high-pressure \u2014 he specifically does not want to feel like he is being sold to',
        'Do NOT use \u201cgrowth hacking,\u201d \u201cscaling,\u201d or startup-culture language',
    ]
    for item in do_not_say:
        p = bullet(doc, item)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)

    page_break(doc)

    # ── DEFINITION OF A WIN ─────────────────────────────────────────────────────
    h1(doc, 'Definition of a Win')

    wins = [
        ('3-Month Win',
         'Google Business Profile generating 2\u20133 new quality patient calls/week from condition-specific '
         'searches (dry eye, myopia management, specialty contacts). In-office optical conversion rate moves '
         'from 52% to 58%+. At least one dry eye case initiated from a patient who read an education asset '
         'before calling.'),
        ('6-Month Win',
         'Referral volume from existing patients up measurably. 1\u20132 myopia management enrollments per month. '
         'At least one insurance plan Sam has been considering dropping \u2014 he now has enough quality patient volume '
         'to make the decision confidently. Front desk reporting that new patients frequently say \u201cI heard '
         'you\u2019re the best for [specific thing].\u201d'),
        ('12-Month Win',
         'Sam works 4 days/week. Gross collections up 15\u201325% on the same or fewer exams. In-office optical '
         'capture rate at 68%+. Dry eye cases represent 12%+ of revenue. His practice has a specific thing it is '
         'known for in his community. New patients arrive having already researched him, already trusting him, '
         'already planning to buy their optical there. He turns away one insurance plan without anxiety.'),
    ]
    for label, text in wins:
        h3(doc, label)
        callout(doc, text)

    h2(doc, 'The Moment Sam Knew It Worked')
    body(doc,
        'A new patient calls and says: \u201cI\u2019ve been dealing with dry eye for two years. Three different '
        'doctors gave me artificial tears. Someone in my neighborhood Facebook group said you actually treat the '
        'cause \u2014 that you\u2019re the person to see for this. I want to come in. I\u2019m ready to do '
        'whatever it takes.\u201d',
        space_after=4)
    callout(doc,
        'That call is the system working. She didn\u2019t compare prices. She didn\u2019t ask about insurance first. '
        'She arrived pre-sold and pre-qualified. Sam did 90 minutes of clinical work. He collected $1,400. '
        'She referred her husband the next month.')

    page_break(doc)

    # ── SECONDARY AVATAR ─────────────────────────────────────────────────────────
    h1(doc, 'Secondary Avatar \u2014 \u201cSpecialty-Pivot Sophia\u201d')

    callout(doc, 'She\u2019s done building a general practice. She wants to build a specialty-driven one.')

    two_col_table(doc, [
        ('Avatar Name',       'Specialty-Pivot Sophia'),
        ('Age Range',         '32\u201348'),
        ('Gender Skew',       'Majority female \u2014 newer OD, specialty-trained or specialty-curious'),
        ('Years in Practice', '5\u201315 years'),
        ('Primary Service',   'Dry eye clinic and/or myopia management launch'),
    ])

    h2(doc, 'Current Situation')
    body(doc,
        'Sophia completed CE training in dry eye disease management or myopia control, invested in the diagnostic '
        'equipment (LipiFlow, OCULUS Keratograph, corneal topographer), and launched a specialty service. '
        'Six months in, she has a functional protocol but her patient pipeline doesn\u2019t know it exists. '
        'She\u2019s seeing 2\u20133 dry eye cases per month when she needs 8\u201310 to cover the equipment '
        'investment and justify the overhead.',
        space_after=6)
    callout(doc,
        'The problem is not clinical readiness \u2014 it\u2019s that her market doesn\u2019t know to seek her out for it.')

    h2(doc, 'Primary Motivations')
    sophia_goals = [
        'Make the specialty equipment investment pay back within 18\u201324 months',
        'Build a reputation as the dry eye or myopia management specialist in her market \u2014 not just an OD who \u201calso does that\u201d',
        'Exit the commodity exam model and build a practice where clinical skill is the differentiator',
        'Create a referral pipeline from ophthalmologists, PCPs, and pediatricians for specialty cases',
        'Attract motivated patients willing to invest in their vision \u2014 not just use their benefits',
    ]
    for g in sophia_goals:
        bullet(doc, g)

    h2(doc, 'Key Differences from Selective Sam')
    differences = [
        ('Digital comfort',     'More comfortable with content and digital channels \u2014 grew up in the social media era'),
        ('Marketing jadedness', 'Less jaded about marketing \u2014 has not been burned as many times yet'),
        ('Focus',               'Specialty positioning rather than general premium differentiation'),
        ('Urgency driver',      'Equipment investment creates urgency \u2014 the ROI math is pressing'),
        ('Network status',      'Often building a referral network from scratch \u2014 less established'),
    ]
    two_col_table(doc, differences)

    h2(doc, 'Belief Barriers')
    belief_block(doc,
        stage='Primary Enemy Belief',
        belief='Patients don\u2019t know what dry eye disease management is. I have to educate my whole market from scratch \u2014 that takes years.',
        why_it_stalls=(
            'This belief makes the task feel impossibly large and the timeline impossibly long. '
            'It causes Sophia to delay positioning decisions while waiting for the market to \u201cbe ready.\u201d'
        ),
        reframe=(
            'The market doesn\u2019t need to understand dry eye management before they seek Sophia out. '
            'They need to understand their symptoms: tired eyes, burning, fluctuating vision, dependence on drops, '
            'discomfort with contact lenses. These symptoms are widely experienced and widely under-treated. '
            'Authority content that describes the symptoms and offers a solution intercepts patients at the moment '
            'of frustration. They don\u2019t need to know the clinical term before they call.'
        )
    )

    belief_block(doc,
        stage='Secondary Enemy Belief',
        belief='I don\u2019t have enough reviews or case studies yet to be seen as a specialist.',
        why_it_stalls=(
            'Sophia is waiting for permission the market will never give her unprompted. '
            'She\u2019s in the same chicken-and-egg trap as every specialist who delayed positioning: '
            'she won\u2019t position as a specialist until she has the cases; she won\u2019t get the cases '
            'until she positions as a specialist.'
        ),
        reframe=(
            'Authority is built forward, not backward. The practices that dominate their niche did not wait until '
            'they had 200 cases before positioning themselves. They positioned early and the cases followed. '
            'The first 10 dry eye cases Sophia treats with genuine outcomes become the foundation of a patient story '
            'library that drives the next 100. Start positioning. Document outcomes. Authority builds as cases build.'
        )
    )

    page_break(doc)

    # ── COMPETITIVE LANDSCAPE ─────────────────────────────────────────────────────
    h1(doc, 'Competitive Landscape')
    body(doc,
        'Understanding the competitive context shapes both the positioning strategy and the sales conversation. '
        'The opportunity is not contested. It is simply unclaimed.',
        space_after=8)

    h2(doc, 'Competitor Weaknesses')
    comp_weaknesses = [
        'Corporate chains (LensCrafters, Visionworks, MyEyeDr, Walmart Vision) market on price, convenience, and insurance acceptance \u2014 not clinical expertise. Premium patients who value expertise actively distrust corporate feel.',
        'Online retailers (Warby Parker, Zenni) have taken optical retail share \u2014 but cannot replicate clinical care, specialty services, or trusted professional relationships.',
        'The typical independent OD has a generic website, no educational content presence, and fewer than 30 Google reviews.',
        'Ophthalmology practices rarely compete for routine vision care and optical retail \u2014 they occupy a different lane and don\u2019t typically build consumer-facing authority content.',
        'Almost no independent OD has positioned themselves as the recognized authority for a specific condition or service type.',
        'Most OD marketing, where it exists, drives exam volume \u2014 not a specific patient profile.',
    ]
    for item in comp_weaknesses:
        bullet(doc, item)

    h2(doc, 'Open Authority Positions (Available in Most Markets)')
    open_positions = [
        '\u201cDry eye specialist [city]\u201d \u2014 almost universally uncontested as an authority position in small and mid-sized markets',
        '\u201cMyopia management [city]\u201d and \u201cchildren\u2019s vision [city]\u201d \u2014 growing search categories with no dominant independent OD voice in most markets',
        'The moment between symptom awareness and care-seeking \u2014 almost entirely unaddressed by any local OD content',
        'Google Business Profile reviews for specialty services \u2014 all reviews are generic; condition-specific social proof is nearly always absent',
        'Medical provider referral relationships \u2014 almost universally underdeveloped',
    ]
    for item in open_positions:
        bullet(doc, item)

    page_break(doc)

    # ── TOP 5 FAST CASH ───────────────────────────────────────────────────────────
    h1(doc, 'Top 5 Fast Cash Priorities')
    body(doc,
        'These five strategies are sequenced for maximum impact with minimum friction. '
        'Strategies 1 and 3 can begin in week one. Strategies 4 and 5 build over 60\u201390 days.',
        space_after=8)

    fast_cash = [
        ('1', 'Reactivate Lapsed Patients + Dry Eye Pre-Diagnosed',
         'Fastest near-zero-cost revenue in the practice. Every independent OD has 300\u2013800 lapsed '
         'patients overdue for recall. Additional high-value segment: patients who received a dry eye '
         'discussion but never booked treatment. These are pre-diagnosed, pre-interested, and the softest '
         'possible sales conversation. A 3-touch reactivation sequence consistently returns 4\u20138 bookings '
         'per 100 contacts. Cost: near zero.'),
        ('2', 'Pareto the Practice \u2014 Sell What Pays',
         'Until Sam sees his per-chair-hour revenue breakdown by service line, he continues allocating '
         'his most valuable resource equally across services with dramatically different margin profiles. '
         'This analysis is the strategic foundation for every other decision. Premium optical, specialty '
         'services, and medical eye care generate 3\u20135x the per-hour revenue of routine insurance exams.'),
        ('3', 'In-Office Optical Conversion System',
         'The optical conversation happens after every exam \u2014 but most practices leave it to chance. '
         'A structured optical consultation process, trained into staff, with clear recommendation language '
         'for each lens upgrade category, moves capture rate from 45\u201355% to 65\u201375%+ within 30\u201360 days. '
         'At 15 exams/day, 2 more captures = $600\u2013$900 more per day = $150,000+ per year.'),
        ('4', 'Local SEO for Condition-Specific Searches',
         '\u201cDry eye specialist [city]\u201d and \u201cmyopia management [city]\u201d are high-intent, '
         'low-competition, and almost always unowned by any local practice. Owning these search positions '
         'produces a fundamentally different patient: someone researching a specific problem who is already '
         'motivated to solve it. Google Business Profile optimization for these conditions consistently '
         'produces condition-specific patient flow within 60\u201390 days.'),
        ('5', 'Specialty Patient Referral + Medical Provider Loop',
         'Every dry eye patient Sam has successfully treated is a referral engine. Most have never been asked. '
         'A post-treatment outcome sequence that invites patients to share their experience activates the '
         'referral channel that produced Sam\u2019s best patients. Medical provider outreach (PCPs, '
         'rheumatologists, endocrinologists, pediatricians) runs in parallel and produces referrals with '
         'multi-year LTV.'),
    ]

    for priority, name, rationale in fast_cash:
        p_hdr = doc.add_paragraph()
        set_para_bg(p_hdr, ASG_BLUE_HEX)
        p_hdr.paragraph_format.space_before = Pt(10)
        p_hdr.paragraph_format.space_after  = Pt(3)
        r_pri = p_hdr.add_run(f'  Priority {priority}  ')
        r_pri.font.name  = 'Calibri'
        r_pri.font.size  = Pt(9)
        r_pri.font.bold  = True
        r_pri.font.color.rgb = WHITE
        r_name = p_hdr.add_run(f'  {name}')
        r_name.font.name  = 'Calibri'
        r_name.font.size  = Pt(11)
        r_name.font.bold  = True
        r_name.font.color.rgb = WHITE
        body(doc, rationale, space_after=6)

    page_break(doc)

    # ── CONTENT INTELLIGENCE ──────────────────────────────────────────────────────
    h1(doc, 'Content Intelligence')

    callout(doc,
        'FRAMING NOTE: All content for this niche is framed as patient education and clinical authority \u2014 '
        'never as marketing. The OD\u2019s name and expertise are on the content. The content team writes and '
        'publishes. The OD approves. Never pitch content as \u201csocial media posts\u201d to this avatar \u2014 '
        'pitch it as \u201cpatient education assets that build trust before patients call.\u201d')

    h2(doc, 'Authority Themes (Evergreen)')
    themes = [
        'Why your chronic dry eye isn\u2019t going to improve with more eye drops \u2014 and what actually treats it',
        'The truth about online eyeglasses: what you\u2019re giving up when you skip the fitting',
        'Why your child\u2019s glasses prescription keeps changing \u2014 and what you can do about it',
        'The difference between an eye exam and a comprehensive vision evaluation',
        '3 things your eye doctor learns from your eyes about your overall health',
        'When \u201cmy eyes are just tired\u201d is actually a diagnosable condition',
        'What premium progressive lenses actually do differently \u2014 a plain-English explanation',
        'Why the cheapest eye exam is often the most expensive decision you\u2019ll make',
    ]
    for t in themes:
        bullet(doc, t)

    h2(doc, 'Social Proof Angles')
    social_proof = [
        'Patient had been using eye drops every 30 minutes for three years. Now she uses them twice a day \u2014 maybe.',
        'Her son\u2019s myopia was progressing fast. Ortho-K stopped it. His prescription barely changed in 18 months.',
        'He bought glasses online twice. Neither pair felt right. One appointment changed that.',
        'She came in expecting another eye exam. She left with a diagnosis and a treatment plan for a condition she\u2019d had for years.',
    ]
    for sp in social_proof:
        bullet(doc, sp)

    h2(doc, 'Content Asset Formats (by Patient Journey Stage)')
    formats = [
        ('Google Business Profile Q&A', 'Condition-specific answered questions built for local search \u2014 dry eye, myopia, specialty lenses'),
        ('Website \u2014 condition pages', 'Dedicated pages for dry eye and myopia management written to convert a researching patient, not describe a service'),
        ('Patient education one-pagers', 'Dry eye self-assessment, myopia risk factors, premium lens guide \u2014 in-office and linked from site'),
        ('Email recall sequences', 'Condition-specific, clinically relevant, timed to annual recall windows \u2014 not generic reminders'),
        ('Google review responses', 'Each response reinforces Sam\u2019s clinical expertise and specialty positioning'),
        ('Medical provider letter', 'Formal introduction to PCPs, pediatricians, rheumatologists \u2014 positions Sam as the referral destination'),
    ]
    two_col_table(doc, formats, col_widths=(2400, 6100))

    # ── COMPLIANCE FLAGS ──────────────────────────────────────────────────────────
    h1(doc, 'Compliance Flags')
    body(doc, 'All content and deliverables for this niche must observe the following:', space_after=6)

    flags = [
        'Cannot guarantee vision outcomes. Use \u201ctypically,\u201d \u201cin most cases,\u201d \u201cresults vary\u201d framing.',
        'Before/after or outcome descriptions require HIPAA-compliant patient consent and de-identification or explicit written release.',
        'Testimonials must be from real patients who have provided written authorization. No fabricated outcomes.',
        'Advertising as a \u201cspecialist\u201d in dry eye or myopia management should be accurate to training. Verify state board rules on specialty designation language.',
        'MiSight (CooperVision) and Ortho-K are FDA-cleared for myopia management \u2014 this claim is accurate and approvable. Other lens types for myopia management should be referenced per current FDA guidance.',
        'Medical billing content must accurately reflect that medical insurance is billed for medical diagnoses \u2014 not as a workaround to vision plan limits.',
        'All promotional materials must include practice identification and comply with state optometry board advertising regulations.',
        'Email marketing must include opt-out mechanism and practice identification per CAN-SPAM.',
        'Content comparing independent ODs favorably to corporate chains must be factual, not disparaging. Focus on what the independent practice offers.',
    ]
    for f in flags:
        p = bullet(doc, f)

    # ── CLOSING ───────────────────────────────────────────────────────────────────
    divider(doc)
    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_close.paragraph_format.space_before = Pt(12)
    r_close = p_close.add_run(
        'Authority Systems Group\u2122  \u2014  Niche Intelligence Library\n'
        'Prepared by Vivienne Carr, CMO + Daniel Frost, CSO\n'
        'Under the direction of Roger Bauer, Director  \u2014  March 2026'
    )
    r_close.font.name  = 'Calibri'
    r_close.font.size  = Pt(9)
    r_close.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r_close.font.italic = True


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_COLOR

    add_cover(doc)
    add_footer(doc)
    add_body(doc)

    doc.save(OUT_DOCX)
    print(f'Saved: {OUT_DOCX}')


if __name__ == '__main__':
    main()
