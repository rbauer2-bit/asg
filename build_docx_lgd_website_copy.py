#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Legal Growth Dynamics: Full Website Copy — HNW Family Law
Target: Family Law Attorneys / Small Multi-Attorney Practices
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Brand Colors ─────────────────────────────────────────
ASG_BLUE     = RGBColor(0x25, 0xAA, 0xE1)   # #25aae1
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)   # #58585a
ASG_DARK     = RGBColor(0x1A, 0x1A, 0x1A)   # #1a1a1a
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR   = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
META_BLUE    = RGBColor(0xEA, 0xF6, 0xFC)

INPUT_FILE  = "outputs/lgd_website_copy_hnw_familylaw_20260307.md"
OUTPUT_DOCX = "outputs/lgd_website_copy_hnw_familylaw_20260307.docx"


# ─── XML Helpers ──────────────────────────────────────────

def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color_hex='25AAE1', weight='6'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), weight)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_blue_left_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '25AAE1')
    pBdr.append(left)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br


def insert_page_break(doc):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el
    p = doc.add_paragraph()
    r = p.add_run()
    br = _el('w:br')
    br.set(_qn('w:type'), 'page')
    r._r.append(br)


# ─── Styles ───────────────────────────────────────────────

def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = Pt(16)

    # H1 — Page title (blue, large)
    try:
        h1 = styles['Heading 1']
    except KeyError:
        h1 = styles.add_style('Heading 1', 1)
    h1.font.name = 'Calibri'
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(10)

    # H2 — Section heading (charcoal)
    try:
        h2 = styles['Heading 2']
    except KeyError:
        h2 = styles.add_style('Heading 2', 1)
    h2.font.name = 'Calibri'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # H3 — Sub-section (blue)
    try:
        h3 = styles['Heading 3']
    except KeyError:
        h3 = styles.add_style('Heading 3', 1)
    h3.font.name = 'Calibri'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = ASG_BLUE
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    # H4 — Sub-sub-section (charcoal bold)
    try:
        h4 = styles['Heading 4']
    except KeyError:
        h4 = styles.add_style('Heading 4', 1)
    h4.font.name = 'Calibri'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = ASG_CHARCOAL
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)


# ─── Cover Page ───────────────────────────────────────────

def add_cover_page(doc):
    # Blue header bar via shaded table
    from docx.oxml import OxmlElement as _el
    from docx.oxml.ns import qn as _qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt as _Pt

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Blue fill
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = _el('w:shd')
    shd.set(_qn('w:val'), 'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'), '25AAE1')
    tcPr.append(shd)

    # Remove cell borders
    tcBorders = _el('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = _el(f'w:{side}')
        border.set(_qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # Cell padding
    tcMar = _el('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = _el(f'w:{side}')
        m.set(_qn('w:w'), '900')
        m.set(_qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("LEGAL GROWTH DYNAMICS™")
    run1.font.name = 'Calibri'
    run1.font.size = _Pt(32)
    run1.font.bold = True
    run1.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Systems-First Marketing for Family Law Attorneys")
    run2.font.name = 'Calibri'
    run2.font.size = _Pt(16)
    run2.font.color.rgb = RGBColor(0xEA, 0xF6, 0xFC)
    run2.font.italic = True

    doc.add_paragraph()

    # Document title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Full Website Copy")
    tr.font.name = 'Calibri'
    tr.font.size = _Pt(36)
    tr.font.bold = True
    tr.font.color.rgb = ASG_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("HNW Family Law | Intake-First Positioning")
    sr.font.name = 'Calibri'
    sr.font.size = _Pt(16)
    sr.font.color.rgb = ASG_CHARCOAL

    doc.add_paragraph()

    # Page list
    pages = [
        "Home  ·  About  ·  Ideal Client",
        "Intake Optimization  ·  Lead Generation",
        "Local SEO  ·  Marketing Automation",
        "The Triad  ·  Web Design & Optimization",
        "Content Strategy  ·  Strategic Planning  ·  Schedule",
    ]
    for pg in pages:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = pp.add_run(pg)
        pr.font.name = 'Calibri'
        pr.font.size = _Pt(12)
        pr.font.color.rgb = ASG_CHARCOAL

    doc.add_paragraph()

    # Date / confidential block
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run("Prepared: March 7, 2026  |  Confidential — Internal Review Copy")
    dr.font.name = 'Calibri'
    dr.font.size = _Pt(10)
    dr.font.color.rgb = ASG_CHARCOAL
    dr.font.italic = True

    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sig_p.add_run("Directed by Roger Bauer  |  Authority Systems Group™")
    sr2.font.name = 'Calibri'
    sr2.font.size = _Pt(10)
    sr2.font.color.rgb = ASG_CHARCOAL

    insert_page_break(doc)


# ─── TOC ──────────────────────────────────────────────────

def add_toc(doc, lines):
    toc_h = doc.add_paragraph()
    r = toc_h.add_run("TABLE OF CONTENTS")
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = ASG_BLUE

    add_horizontal_rule(doc)

    page_num = 3
    for line in lines:
        m = re.match(r'^# PAGE (\d+): (.+)$', line.strip())
        if m:
            pnum, title = m.group(1), m.group(2)
            p = doc.add_paragraph()
            r_title = p.add_run(f"Page {pnum}: {title}")
            r_title.font.name = 'Calibri'
            r_title.font.size = Pt(11)
            r_title.font.color.rgb = ASG_CHARCOAL
            p.paragraph_format.space_after = Pt(4)
            page_num += 1

    insert_page_break(doc)


# ─── Rich Paragraph (inline bold/italic) ──────────────────

def add_rich_paragraph(doc, text, style='Normal', align=None,
                        font_size=None, color=None, bold=False,
                        space_before=None, space_after=None,
                        left_indent=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)

    # Split on **bold** and *italic* markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True

        run.font.name = 'Calibri'
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color

    return p


# ─── Callout Box ──────────────────────────────────────────

def add_callout_box(doc, text):
    """Blue left-border callout for bold insight lines."""
    p = add_rich_paragraph(doc, text, font_size=11.5, color=ASG_BLUE,
                            bold=True, left_indent=0.3,
                            space_before=10, space_after=10)
    add_blue_left_border(p)
    return p


# ─── Page Header Box ──────────────────────────────────────

def add_page_header_box(doc, page_num, page_title, url, seo_title=None):
    """Light blue metadata box at top of each page section."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, 'EAF6FC')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"PAGE {page_num}: {page_title}")
    r1.font.name = 'Calibri'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = ASG_BLUE

    p2 = doc.add_paragraph()
    set_paragraph_shading(p2, 'EAF6FC')
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(f"URL: {url}")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9)
    r2.font.color.rgb = ASG_CHARCOAL

    if seo_title:
        p3 = doc.add_paragraph()
        set_paragraph_shading(p3, 'EAF6FC')
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(6)
        r3 = p3.add_run(f"SEO Title: {seo_title}")
        r3.font.name = 'Calibri'
        r3.font.size = Pt(9)
        r3.font.color.rgb = ASG_CHARCOAL

    doc.add_paragraph()


# ─── CTA Box ──────────────────────────────────────────────

def add_cta_box(doc, headline, cta_text):
    """Blue-bordered CTA section."""
    doc.add_paragraph()
    add_horizontal_rule(doc, color_hex='25AAE1', weight='4')

    p = add_rich_paragraph(doc, headline, font_size=13, color=ASG_BLUE,
                            bold=True, space_before=8, space_after=6)

    cta_p = doc.add_paragraph()
    cta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cta_p.add_run(f"[ {cta_text} ]")
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = WHITE

    # Blue background for CTA button
    set_paragraph_shading(cta_p, '25AAE1')
    cta_p.paragraph_format.space_before = Pt(6)
    cta_p.paragraph_format.space_after = Pt(6)

    add_horizontal_rule(doc, color_hex='25AAE1', weight='4')
    doc.add_paragraph()


# ─── Footer ───────────────────────────────────────────────

def add_footer(doc):
    from docx.oxml import OxmlElement as _el
    from docx.oxml.ns import qn as _qn

    section = doc.sections[0]
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        footer._element.append(_el('w:p'))

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = fp._p.get_or_add_pPr()
    pBdr = _el('w:pBdr')
    top = _el('w:top')
    top.set(_qn('w:val'), 'single')
    top.set(_qn('w:sz'), '4')
    top.set(_qn('w:space'), '1')
    top.set(_qn('w:color'), '25AAE1')
    pBdr.append(top)
    pPr.append(pBdr)

    r = fp.add_run("Legal Growth Dynamics™  |  Authority Systems Group™  |  Confidential")
    r.font.name = 'Calibri'
    r.font.size = Pt(8)
    r.font.color.rgb = ASG_CHARCOAL


# ─── Main Content Renderer ────────────────────────────────

def is_cta_line(line):
    return line.strip().startswith('[') and line.strip().endswith(']') and 'SCHEDULE' in line.strip().upper()


def render_content(doc, lines):
    """Parse markdown and render to DOCX."""
    i = 0
    current_page_num = None
    skip_file_header = True
    in_cta_section = False
    cta_headline = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip file header comment lines and first --- separator
        if skip_file_header:
            if stripped.startswith('#') and ('Legal Growth Dynamics' in stripped or 'Target:' in stripped or 'Framework:' in stripped or 'Date:' in stripped):
                i += 1
                continue
            elif stripped == '---':
                skip_file_header = False
                i += 1
                continue

        # --- separator (section divider or page separator)
        if stripped == '---':
            if current_page_num is not None:
                # Double --- means end of page — insert page break
                if i + 1 < len(lines) and lines[i+1].strip() == '---':
                    insert_page_break(doc)
                    i += 2
                    continue
                else:
                    add_horizontal_rule(doc, color_hex='58585A', weight='4')
            i += 1
            continue

        # PAGE N header
        page_match = re.match(r'^# PAGE (\d+): (.+)$', stripped)
        if page_match:
            pnum = page_match.group(1)
            ptitle = page_match.group(2)
            current_page_num = pnum
            in_cta_section = False

            # Collect metadata lines
            j = i + 1
            page_title_line = None
            url_line = None
            seo_line = None
            while j < len(lines) and not lines[j].strip().startswith('##'):
                meta = lines[j].strip()
                if meta.startswith('**Page Title:**'):
                    page_title_line = meta.replace('**Page Title:**', '').strip()
                elif meta.startswith('**URL:**'):
                    url_line = meta.replace('**URL:**', '').strip()
                elif meta.startswith('**SEO Title:**'):
                    seo_line = meta.replace('**SEO Title:**', '').strip()
                j += 1

            add_page_header_box(doc, pnum, ptitle,
                                 url_line or '',
                                 seo_title=seo_line)
            i += 1
            continue

        # Skip metadata lines (they were consumed above)
        if stripped.startswith('**Page Title:**') or stripped.startswith('**URL:**') or stripped.startswith('**SEO Title:**'):
            i += 1
            continue

        # ## SECTION heading (H2)
        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            if heading_text == 'HERO SECTION':
                add_horizontal_rule(doc, color_hex='25AAE1', weight='6')
                p = doc.add_paragraph()
                r = p.add_run('HERO SECTION')
                r.font.name = 'Calibri'
                r.font.size = Pt(9)
                r.font.color.rgb = ASG_BLUE
                r.font.bold = True
                p.paragraph_format.space_after = Pt(4)
            elif heading_text == 'CTA SECTION' or heading_text == 'CTA':
                in_cta_section = True
                add_horizontal_rule(doc, color_hex='25AAE1', weight='4')
            else:
                in_cta_section = False
                p = doc.add_paragraph(heading_text, style='Heading 2')
            i += 1
            continue

        # ### Sub-section heading (H3 — blue)
        if stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            if heading_text in ('Headline:', 'Subheadline:'):
                # Next line is the actual text
                i += 1
                continue
            p = doc.add_paragraph(heading_text, style='Heading 3')
            i += 1
            continue

        # #### Sub-sub heading (H4 — charcoal bold)
        if stripped.startswith('#### '):
            heading_text = stripped[5:].strip()
            p = doc.add_paragraph(heading_text, style='Heading 4')
            i += 1
            continue

        # Headline: / Subheadline: (hero content)
        if stripped.startswith('Headline:') or stripped == 'Headline:':
            i += 1
            if i < len(lines):
                hl_text = lines[i].strip()
                p = add_rich_paragraph(doc, hl_text, font_size=26, color=ASG_BLUE,
                                        bold=True, space_before=6, space_after=8)
            i += 1
            continue

        if stripped.startswith('Subheadline:') or stripped == 'Subheadline:':
            i += 1
            if i < len(lines):
                sh_text = lines[i].strip()
                p = add_rich_paragraph(doc, sh_text, font_size=13, color=ASG_CHARCOAL,
                                        space_before=0, space_after=12)
                p.paragraph_format.left_indent = Inches(0.15)
            i += 1
            continue

        if stripped == 'CTA:':
            i += 1
            if i < len(lines):
                cta_text = lines[i].strip()
                cta_p = doc.add_paragraph()
                cta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = cta_p.add_run(f"→  {cta_text}")
                r.font.name = 'Calibri'
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = WHITE
                set_paragraph_shading(cta_p, '25AAE1')
                cta_p.paragraph_format.space_before = Pt(6)
                cta_p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # CTA block in CTA section — [TEXT]
        if is_cta_line(stripped) and in_cta_section:
            cta_text = stripped.strip('[').strip(']').strip()
            cta_p = doc.add_paragraph()
            cta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cta_p.add_run(cta_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_paragraph_shading(cta_p, '25AAE1')
            cta_p.paragraph_format.space_before = Pt(8)
            cta_p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Blockquote / callout >
        if stripped.startswith('> '):
            callout_text = stripped[2:].strip()
            add_callout_box(doc, callout_text)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- '):
            bullet_text = stripped[2:].strip()
            p = add_rich_paragraph(doc, bullet_text, font_size=11, color=BODY_COLOR,
                                    space_before=2, space_after=2, left_indent=0.25)
            # Add bullet character
            run0 = p.runs[0] if p.runs else p.add_run('')
            p._p.insert(0, OxmlElement('w:r'))
            # Just prepend bullet character to first run's text
            if p.runs:
                first_run = p.runs[0]
                first_run.text = '•  ' + first_run.text
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\. (.+)$', stripped)
        if num_match:
            num = num_match.group(1)
            item_text = num_match.group(2)
            p = add_rich_paragraph(doc, item_text, font_size=11, color=BODY_COLOR,
                                    space_before=2, space_after=2, left_indent=0.25)
            if p.runs:
                p.runs[0].text = f"{num}.  " + p.runs[0].text
            i += 1
            continue

        # Italic disclaimer line
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            disclaimer = stripped.strip('*')
            p = doc.add_paragraph()
            r = p.add_run(disclaimer)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            r.font.italic = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # CTA in non-CTA context (inline CTA with arrow)
        if stripped.startswith('[') and stripped.endswith(']') and (
                'SCHEDULE' in stripped.upper() or 'CALL' in stripped.upper()):
            cta_text = stripped.strip('[').strip(']')
            cta_p = doc.add_paragraph()
            cta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cta_p.add_run(cta_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_paragraph_shading(cta_p, '25AAE1')
            cta_p.paragraph_format.space_before = Pt(6)
            cta_p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph (with inline bold/italic)
        add_rich_paragraph(doc, stripped, font_size=11, color=BODY_COLOR,
                            space_before=0, space_after=8)
        i += 1


# ─── Page Setup ───────────────────────────────────────────

def setup_page(doc):
    # IMPORTANT: section.page_width / page_height take EMU (English Metric Units).
    # NEVER pass raw DXA values (e.g. 12240) — python-docx will treat them as EMU
    # and produce a page ~0.013" wide. Always use Inches() here.
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)


# ─── Main ─────────────────────────────────────────────────

def main():
    doc = Document()
    setup_page(doc)
    setup_styles(doc)

    # Load input lines
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        all_lines = f.readlines()

    add_cover_page(doc)
    add_toc(doc, all_lines)
    add_footer(doc)
    render_content(doc, all_lines)

    doc.save(OUTPUT_DOCX)
    print(f"✓ Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
