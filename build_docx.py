#!/usr/bin/env python3
"""
Authority Systems Group™ — DOCX Builder
Converts the Authority Blueprint markdown into a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Brand Colors ───────────────────────────────────────────
ASG_BLUE = RGBColor(0x25, 0xAA, 0xE1)      # #25aae1
ASG_CHARCOAL = RGBColor(0x58, 0x58, 0x5A)  # #58585a
ASG_DARK = RGBColor(0x1A, 0x1A, 0x1A)      # #1a1a1a
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)

INPUT_MD = "client-onboarding/clients/kristen-lemastus-coaching/authority_blueprint_20260304.md"
OUTPUT_DOCX = "client-onboarding/clients/kristen-lemastus-coaching/kristen-lemastus-coaching_authority-blueprint_20260304.docx"


def set_paragraph_shading(paragraph, fill_hex):
    """Add background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color=ASG_BLUE):
    """Add a full-width horizontal rule in ASG Blue."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '25AAE1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def setup_styles(doc):
    """Configure document styles."""
    styles = doc.styles

    # Normal (body) style
    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = Pt(15)

    # Heading 1 — Section titles
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2 — Sub-sections
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 — Sub-sub-sections
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)


def add_cover_page(doc):
    """Build the cover page."""
    # Dark background section — simulated with a bordered box paragraph
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('AUTHORITY BLUEPRINT™')
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = ASG_BLUE
    cover_title.paragraph_format.space_before = Pt(72)
    cover_title.paragraph_format.space_after = Pt(8)

    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = firm.add_run('Kristen LeMastus Coaching & Consulting')
    r.font.name = 'Calibri'
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = ASG_CHARCOAL
    firm.paragraph_format.space_after = Pt(4)

    attorney = doc.add_paragraph()
    attorney.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = attorney.add_run('Kristen LeMastus | Louisville, KY')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(18)
    r2.font.color.rgb = ASG_CHARCOAL
    attorney.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    for line in [
        'Prepared exclusively for Kristen LeMastus',
        'Kristen LeMastus Coaching & Consulting',
        'Louisville, KY',
        '',
        'March 2026  |  Confidential',
        'Under the direction of Roger, Director',
        'Authority Systems Group™',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        if 'Authority Systems Group' in line or 'Roger' in line:
            r.font.color.rgb = ASG_BLUE
            r.font.bold = True
        elif 'Confidential' in line or 'March' in line:
            r.font.color.rgb = ASG_CHARCOAL
            r.font.bold = True
        else:
            r.font.color.rgb = ASG_CHARCOAL
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def add_footer(doc, client_name="Kristen LeMastus / Kristen LeMastus Coaching & Consulting"):
    """Add footer to all sections."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f'Authority Systems Group™ — Confidential. Prepared exclusively for {client_name}.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def parse_and_render(doc, md_path):
    """Parse markdown and render to docx."""
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False
    skip_cover = True   # Skip the first section (cover page data already added)
    cover_done = False
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip the original cover page content (first 8 lines until first --- )
        if not cover_done:
            if line.strip() == '---':
                cover_done = True
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Horizontal rules — render as visual divider
        if line.strip() in ('---', '---\n'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # Skip separator rows
                if not re.match(r'^\|[\s\-\|]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.paragraph_format.space_after = Pt(2)
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = WHITE
                            if ri == 0:
                                for para in cell.paragraphs:
                                    # Header row shading
                                    tc = cell._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'), '25AAE1')
                                    tcPr.append(shd)
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if text.startswith('SECTION') or text in ('TABLE OF CONTENTS',):
                p = doc.add_heading(text, level=1)
            else:
                p = doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip()
            # Strip italic markers
            text = text.replace('*', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.left_indent = Inches(0.3)
            # Left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # strip bold markers for now
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold inline
            parts = re.split(r'\*\*(.*?)\*\*', line[2:].strip())
            p.clear()
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = 'Georgia'
                r.font.size = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            text = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = 'Georgia'
                r.font.size = Pt(11)
                r.font.color.rgb = BODY_COLOR
                if j % 2 == 1:
                    r.font.bold = True
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Checkbox items
        if line.strip().startswith('- [ ]') or line.strip().startswith('□'):
            text = line.strip().lstrip('- [ ]').lstrip('□').strip()
            p = doc.add_paragraph()
            r = p.add_run('☐  ' + text)
            r.font.name = 'Georgia'
            r.font.size = Pt(11)
            r.font.color.rgb = BODY_COLOR
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Page break markers
        if 'page_break' in line.lower() or line.strip() == '\\newpage':
            doc.add_page_break()
            i += 1
            continue

        # Subject line headers in email sequences
        if line.strip().startswith('Subject:'):
            p = doc.add_paragraph()
            set_paragraph_shading(p, 'F5F5F5')
            r = p.add_run(line.strip())
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = ASG_CHARCOAL
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Section divider lines (full line of dashes within sections)
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            # Small spacer
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph — handle inline bold/italic
        p = doc.add_paragraph()
        # Split by bold markers
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            # Handle italic within part
            italic_parts = re.split(r'\*(.*?)\*', part)
            for k, ipart in enumerate(italic_parts):
                if ipart:
                    r = p.add_run(ipart)
                    r.font.name = 'Georgia'
                    r.font.size = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
                    if k % 2 == 1:
                        r.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        i += 1

    return doc


def main():
    doc = Document()

    # Page setup — letter size with moderate margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)
    parse_and_render(doc, INPUT_MD)

    doc.save(OUTPUT_DOCX)
    print(f"✓ Document saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
