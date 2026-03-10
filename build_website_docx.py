#!/usr/bin/env python3
"""
Authority Systems Group™ — Website Copy DOCX Builder
Combines all website copy files into a single formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ─── Brand Colors ───────────────────────────────────────────
ASG_BLUE    = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL= RGBColor(0x58, 0x58, 0x5A)
ASG_DARK    = RGBColor(0x1A, 0x1A, 0x1A)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BODY_COLOR  = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY  = RGBColor(0xAA, 0xAA, 0xAA)

WEBSITE_DIR = "client-onboarding/clients/authority-systems-group/website"
OUTPUT_DOCX = "client-onboarding/clients/authority-systems-group/outputs/authority-systems-group_website-copy_20260307.docx"

# Ordered file list
PAGES = [
    ("HOMEPAGE", os.path.join(WEBSITE_DIR, "01-homepage.md")),
    ("ABOUT", os.path.join(WEBSITE_DIR, "08-about.md")),
    ("SERVICE: Authority Blueprint™ (DFY)", os.path.join(WEBSITE_DIR, "services/02-dfy-authority-blueprint.md")),
    ("SERVICE: Full Authority Retainer (DFY)", os.path.join(WEBSITE_DIR, "services/03-dfy-authority-retainer.md")),
    ("SERVICE: Authority Accelerator™ (DWY)", os.path.join(WEBSITE_DIR, "services/04-dwy-authority-accelerator.md")),
    ("SERVICE: Belief Map Intensive™ (DWY)", os.path.join(WEBSITE_DIR, "services/05-dwy-belief-map-intensive.md")),
    ("SERVICE: Authority Intelligence Report™ (DIY)", os.path.join(WEBSITE_DIR, "services/06-diy-intelligence-report.md")),
    ("SERVICE: ACP Self-Guided Mastery Program (DIY)", os.path.join(WEBSITE_DIR, "services/07-diy-acp-mastery-program.md")),
]


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc, color='25AAE1'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_section_label(doc, label_text):
    """Render [HERO], [SECTION 1], etc. as styled section markers."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, 'EEF8FD')
    r = p.add_run(label_text)
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = ASG_BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)


def add_page_divider(doc, page_title):
    """Full-width page title banner between pages."""
    doc.add_page_break()
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(page_title.upper())
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = ASG_BLUE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_horizontal_rule(doc)


def setup_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = Pt(15)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ASG_BLUE
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = ASG_CHARCOAL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(5)

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = ASG_CHARCOAL
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)


def add_cover_page(doc):
    # Top spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('AUTHORITY SYSTEMS GROUP™')
    r.font.name = 'Calibri'
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = ASG_BLUE
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('Website Copy Package')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(20)
    r2.font.color.rgb = ASG_CHARCOAL
    sub.paragraph_format.space_after = Pt(4)

    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = url.add_run('AuthoritySystemsGroup.com')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(14)
    r3.font.italic = True
    r3.font.color.rgb = LIGHT_GRAY
    url.paragraph_format.space_after = Pt(48)

    add_horizontal_rule(doc)

    for line, color, bold in [
        ('Homepage  ·  About  ·  6 Service Pages', ASG_CHARCOAL, False),
        ('DFY  ·  DWY  ·  DIY', ASG_CHARCOAL, False),
        ('', ASG_CHARCOAL, False),
        ('March 2026  |  Internal Reference Document', LIGHT_GRAY, False),
        ('Directed by Roger  |  Authority Systems Group™', ASG_BLUE, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = color
        r.font.bold = bold
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Table of Contents
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = toc_title.add_run('TABLE OF CONTENTS')
    r.font.name = 'Calibri'
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ASG_CHARCOAL
    toc_title.paragraph_format.space_after = Pt(16)

    toc_entries = [
        ('Homepage', 'The primary landing page — belief sequence arc, services overview, CTA'),
        ('About', 'Team structure, personas, "The Honest Answer" section'),
        ('Authority Blueprint™', 'DFY — Flagship 13-section strategy document'),
        ('Full Authority Retainer', 'DFY — Ongoing monthly implementation engagement'),
        ('Authority Accelerator™', 'DWY — 90-day build-alongside program with Roger'),
        ('Belief Map Intensive™', 'DWY — Single-session buyer belief architecture mapping'),
        ('Authority Intelligence Report™', 'DIY — Market research deliverable, standalone'),
        ('ACP Self-Guided Mastery Program', 'DIY — Self-paced training in the full ACP methodology'),
    ]
    for title, desc in toc_entries:
        p = doc.add_paragraph()
        r_title = p.add_run(f'{title}  —  ')
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = ASG_CHARCOAL
        r_desc = p.add_run(desc)
        r_desc.font.name = 'Calibri'
        r_desc.font.size = Pt(11)
        r_desc.font.color.rgb = BODY_COLOR
        p.paragraph_format.space_after = Pt(5)

    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run('Authority Systems Group™  —  Website Copy Package  —  Internal Reference  |  March 2026')
    r.font.name = 'Calibri'
    r.font.size = Pt(8)
    r.font.color.rgb = LIGHT_GRAY


def is_metadata_line(line):
    """Skip comment-style metadata lines at the top of each file."""
    stripped = line.strip()
    # Lines like: # FILENAME.md, # Voice: ..., # Avatar: ..., # Copy Writer: ...
    meta_prefixes = [
        '# Service Page:', '# Tier:', '# Voice:', '# Avatar', '# Avatar Stage',
        '# Copy Writer', '# Framework:', '# Stage Coverage', '# Reviewed:',
        '# AuthoritySystemsGroup', '# This is',
    ]
    if stripped.startswith('#') and any(stripped.startswith(p) for p in meta_prefixes):
        return True
    # Also skip lines that are purely filename headers like "# 01-homepage.md" style
    if re.match(r'^# [A-Z_]+\.md$', stripped):
        return True
    return False


def render_page(doc, md_path):
    """Parse one markdown file and render it into the document."""
    with open(md_path, 'r') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False
    metadata_section = True  # True until we hit the first ---
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip metadata header block (everything before first ---)
        if metadata_section:
            if line.strip() == '---':
                metadata_section = False
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = ASG_CHARCOAL
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-\|:]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < cols:
                            cell = table.cell(ri, ci)
                            # Strip markdown bold from cell text
                            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
                            cell.text = clean_text
                            for para in cell.paragraphs:
                                para.paragraph_format.space_after = Pt(2)
                                for run in para.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = WHITE
                            if ri == 0:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '25AAE1')
                                tcPr.append(shd)
            continue

        # Section labels like [HERO], [SECTION 1 — TITLE], [NAV], [NAV BREADCRUMB]
        bracket_match = re.match(r'^\[([A-Z][^\]]*)\]$', line.strip())
        if bracket_match:
            label = bracket_match.group(1)
            if label in ('NAV', 'NAV BREADCRUMB'):
                # Render nav lines in a subtle muted style
                p = doc.add_paragraph()
                r = p.add_run(f'[ {label} ]')
                r.font.name = 'Calibri'
                r.font.size = Pt(8)
                r.font.color.rgb = LIGHT_GRAY
                r.font.italic = True
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(6)
            else:
                add_section_label(doc, f'[ {label} ]')
            i += 1
            continue

        # Headings
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = ASG_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = ASG_BLUE
            set_paragraph_shading(p, 'F5F5F5')
            p.paragraph_format.left_indent = Inches(0.3)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '25AAE1')
            pBdr.append(left)
            pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Callout boxes: **[CALLOUT BOX]**
        if line.strip() in ('**[CALLOUT BOX]**', '[CALLOUT BOX]'):
            i += 1
            continue  # The content following is blockquoted — handled above

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            raw = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', raw)
            for j, part in enumerate(parts):
                italic_parts = re.split(r'\*(.*?)\*', part)
                for k, ipart in enumerate(italic_parts):
                    if ipart:
                        r = p.add_run(ipart)
                        r.font.name = 'Georgia'
                        r.font.size = Pt(11)
                        r.font.color.rgb = BODY_COLOR
                        if j % 2 == 1:
                            r.font.bold = True
                        if k % 2 == 1:
                            r.font.italic = True
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line):
            raw = re.sub(r'^\d+[\.\)]\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', raw)
            for j, part in enumerate(parts):
                italic_parts = re.split(r'\*(.*?)\*', part)
                for k, ipart in enumerate(italic_parts):
                    if ipart:
                        r = p.add_run(ipart)
                        r.font.name = 'Georgia'
                        r.font.size = Pt(11)
                        r.font.color.rgb = BODY_COLOR
                        if j % 2 == 1:
                            r.font.bold = True
                        if k % 2 == 1:
                            r.font.italic = True
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Page notes (italic comment at end of file)
        if line.strip().startswith('*Page notes') or line.strip().startswith('*page notes'):
            p = doc.add_paragraph()
            r = p.add_run(line.strip().strip('*'))
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = LIGHT_GRAY
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue

        # Regular paragraph — inline bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for j, part in enumerate(parts):
            italic_parts = re.split(r'\*(.*?)\*', part)
            for k, ipart in enumerate(italic_parts):
                if ipart:
                    r = p.add_run(ipart)
                    r.font.name = 'Georgia'
                    r.font.size = Pt(11)
                    r.font.color.rgb = BODY_COLOR
                    if j % 2 == 1:
                        r.font.bold = True
                    if k % 2 == 1:
                        r.font.italic = True
        p.paragraph_format.space_after = Pt(6)
        i += 1


def main():
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin  = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin   = Inches(1.0)
    section.bottom_margin= Inches(1.0)

    setup_styles(doc)
    add_cover_page(doc)
    add_footer(doc)

    for idx, (page_label, md_path) in enumerate(PAGES):
        print(f'  Rendering: {page_label}')
        add_page_divider(doc, page_label)
        render_page(doc, md_path)

    doc.save(OUTPUT_DOCX)
    print(f'\n✓ Document saved: {OUTPUT_DOCX}')


if __name__ == '__main__':
    main()
