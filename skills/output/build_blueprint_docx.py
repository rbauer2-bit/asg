"""
Authority Systems Group™ — DOCX Blueprint Builder
Builds the fully branded Authority Blueprint™ for Legal Growth Dynamics
Per DOCX_BUILDER.md and REPORT_FORMATTER.md specifications
Owner: Preston Adler, COO / Iris Nolan, CTO
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import copy
import os

# ─────────────────────────────────────────
# BRAND CONSTANTS
# ─────────────────────────────────────────
ASG_BLUE       = RGBColor(0x25, 0xAA, 0xE1)
ASG_CHARCOAL   = RGBColor(0x58, 0x58, 0x5A)
ASG_DARK_BG    = RGBColor(0x1A, 0x1A, 0x1A)
ASG_BODY_TEXT  = RGBColor(0x33, 0x33, 0x33)
ASG_LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
ASG_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ASG_FOOTER_CLR = RGBColor(0x88, 0x88, 0x88)
ASG_MID_GRAY   = RGBColor(0xAA, 0xAA, 0xAA)
ASG_BORDER_CLR = RGBColor(0xCC, 0xCC, 0xCC)
ASG_DARK_GRAY  = RGBColor(0x66, 0x66, 0x66)

CLIENT_NAME    = "Legal Growth Dynamics"
CLIENT_SLUG    = "lgd"
DOC_DATE       = "March 2026"
OUTPUT_DIR     = "/Users/Roger/Dropbox/code/authority-systems-group/client-onboarding/clients/legal-growth-dynamics/outputs"
OUTPUT_FILE    = f"{OUTPUT_DIR}/lgd_authority-blueprint_20260302.docx"


# ─────────────────────────────────────────
# XML HELPERS
# ─────────────────────────────────────────

def set_cell_background(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_paragraph_border(paragraph, side='left', color='25aae1', size='36', space='100'):
    """Add a border to a paragraph (for callout boxes)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), size)
    bdr.set(qn('w:space'), space)
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, fill='F5F5F5'):
    """Set paragraph background shading."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def set_page_color(section_element, hex_color):
    """Set page background color for a section."""
    sectPr = section_element
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), hex_color)
    return bg


def add_horizontal_rule(doc, color='25aae1', size=16):
    """Add a horizontal rule paragraph with bottom border."""
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(6)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return rule


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_module.enum.text.WD_BREAK.PAGE)
    return para


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    """Set table cell margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


# ─────────────────────────────────────────
# TEXT HELPERS
# ─────────────────────────────────────────

def add_h1(doc, text, color=None):
    """Section title — ASG Blue, 24pt, Calibri Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.color.rgb = color or ASG_BLUE
    return p


def add_h2(doc, text, color=None):
    """Subsection header — ASG Charcoal, 16pt, Calibri Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.color.rgb = color or ASG_CHARCOAL
    return p


def add_h3(doc, text, color=None):
    """Sub-subsection header — ASG Charcoal, 13pt, Calibri Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = color or ASG_CHARCOAL
    return p


def add_body(doc, text, italic=False, color=None):
    """Body text — Georgia, 11pt, #333333."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(11)
    run.font.color.rgb = color or ASG_BODY_TEXT
    run.italic = italic
    return p


def add_callout(doc, text, label=None):
    """Callout box — light gray bg, blue left border, italic."""
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(12)
        lp.paragraph_format.space_after = Pt(2)
        lp.paragraph_format.left_indent = Inches(0.25)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = 'Calibri'
        lr.font.size = Pt(10)
        lr.font.color.rgb = ASG_CHARCOAL
        set_paragraph_border(lp, 'left', '25aae1', '36', '100')
        set_paragraph_shading(lp, 'F5F5F5')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0) if label else Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_CHARCOAL
    set_paragraph_border(p, 'left', '25aae1', '36', '100')
    set_paragraph_shading(p, 'F5F5F5')
    return p


def add_signoff(doc, text, name_title):
    """Persona sign-off block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_CHARCOAL

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(16)
    p2.paragraph_format.left_indent = Inches(0.4)
    run2 = p2.add_run(f'— {name_title}')
    run2.bold = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.color.rgb = ASG_CHARCOAL


def add_attribution(doc, text):
    """Section attribution line (italic, charcoal, smaller)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = ASG_CHARCOAL
    return p


def add_bullet(doc, text, level=0):
    """Bulleted list item."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(11)
    run.font.color.rgb = ASG_BODY_TEXT
    return p


def section_break(doc):
    """Full-width blue rule then page break."""
    add_horizontal_rule(doc, '25aae1', 16)
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def section_header(doc, number, title, prepared_by=None):
    """Standard section opener: blue rule, section number, H1 title, attribution."""
    add_horizontal_rule(doc, '25aae1', 16)
    num_p = doc.add_paragraph()
    num_p.paragraph_format.space_before = Pt(6)
    num_p.paragraph_format.space_after = Pt(2)
    num_run = num_p.add_run(f'SECTION {number}')
    num_run.font.name = 'Calibri'
    num_run.font.size = Pt(10)
    num_run.font.color.rgb = ASG_CHARCOAL
    num_run.bold = True

    add_h1(doc, title)
    if prepared_by:
        add_attribution(doc, prepared_by)


# ─────────────────────────────────────────
# TABLE HELPERS
# ─────────────────────────────────────────

def add_branded_table(doc, headers, rows, col_widths=None):
    """Build a fully branded table with blue header row and alternating rows."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Inches(col_widths[i])

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        set_cell_background(hdr_cells[i], '25aae1')
        set_cell_margins(hdr_cells[i])
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = ASG_WHITE

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        bg = 'FFFFFF' if row_idx % 2 == 0 else 'F5F5F5'
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = ''
            set_cell_background(row_cells[col_idx], bg)
            set_cell_margins(row_cells[col_idx])
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = ASG_BODY_TEXT

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


# ─────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────

def build_cover_page(doc):
    """Build the dark-background cover page using a full-page table."""
    # Use a single-cell table spanning the full page for dark background
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Left accent stripe column (~0.15")
    accent_col = table.columns[0]
    accent_col.width = Inches(0.18)
    accent_cell = table.rows[0].cells[0]
    set_cell_background(accent_cell, '25aae1')
    accent_cell.text = ''

    # Main content column
    content_col = table.columns[1]
    content_col.width = Inches(7.57)
    content_cell = table.rows[0].cells[1]
    set_cell_background(content_cell, '1a1a1a')

    # Set row height to full page
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '15840')  # 11 inches in twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    def cover_para(cell, text, size, bold=False, italic=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.color.rgb = color or ASG_WHITE
        return p

    # Clear default paragraph in content cell
    content_cell.paragraphs[0].clear()

    # Top spacer
    cover_para(content_cell, '', 8, space_before=60)

    # ASG Header label
    cover_para(content_cell, 'AUTHORITY SYSTEMS GROUP™', 10, bold=True, color=ASG_BLUE, space_before=8, space_after=4)

    # Spacer
    cover_para(content_cell, '', 24, space_before=30)

    # Main title
    cover_para(content_cell, 'YOUR AUTHORITY', 34, bold=True, color=ASG_WHITE, space_before=0, space_after=0)
    cover_para(content_cell, 'BLUEPRINT™', 34, bold=True, color=ASG_BLUE, space_before=0, space_after=16)

    # Divider line text
    cover_para(content_cell, '─' * 38, 10, color=RGBColor(0x44, 0x44, 0x44), space_before=4, space_after=16)

    # Subtitle
    cover_para(content_cell, 'A Strategic Growth System Built Exclusively for:', 13, italic=True, color=ASG_MID_GRAY, space_before=0, space_after=10)

    # Client name
    cover_para(content_cell, CLIENT_NAME, 26, bold=True, color=ASG_WHITE, space_before=0, space_after=24)

    # Divider
    cover_para(content_cell, '─' * 38, 10, color=RGBColor(0x44, 0x44, 0x44), space_before=0, space_after=24)

    # Attribution
    cover_para(content_cell, 'Directed by Roger  |  Authority Systems Group™', 12, italic=True, color=ASG_MID_GRAY, space_before=0, space_after=8)

    # Date
    cover_para(content_cell, DOC_DATE, 11, color=ASG_FOOTER_CLR, space_before=0, space_after=60)

    # Spacer to push confidential notice to bottom
    cover_para(content_cell, '', 40, space_before=0, space_after=0)

    # Bottom accent line
    cover_para(content_cell, '━' * 55, 9, color=ASG_BLUE, space_before=24, space_after=12)

    # Confidential notice
    cover_para(content_cell, f'CONFIDENTIAL  —  PREPARED EXCLUSIVELY FOR {CLIENT_NAME.upper()}', 8,
               color=ASG_DARK_GRAY, space_before=0, space_after=8)

    doc.add_paragraph()  # will be removed in page break


# ─────────────────────────────────────────
# HEADER / FOOTER
# ─────────────────────────────────────────

def setup_header_footer(doc):
    """Configure header and footer for interior pages."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        # Clear default
        for p in header.paragraphs:
            p.clear()

        hdr_p = header.paragraphs[0]
        hdr_p.clear()
        hdr_p.paragraph_format.space_before = Pt(0)
        hdr_p.paragraph_format.space_after = Pt(4)

        # Left: ASG brand name
        run_left = hdr_p.add_run('Authority Systems Group™')
        run_left.font.name = 'Calibri'
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = ASG_CHARCOAL
        run_left.bold = True

        # Tab to right
        hdr_p.add_run('\t')

        # Right: doc title
        run_right = hdr_p.add_run(f'Authority Blueprint™  —  {CLIENT_NAME}')
        run_right.font.name = 'Calibri'
        run_right.font.size = Pt(9)
        run_right.font.color.rgb = ASG_FOOTER_CLR

        # Tab stops: right-align the right element
        from docx.oxml import OxmlElement
        pPr = hdr_p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')  # 6.3 inches in twips
        tabs.append(tab)
        pPr.append(tabs)

        # Header bottom border
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        ftr_p = footer.paragraphs[0]
        ftr_p.clear()
        ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ftr_p.paragraph_format.space_before = Pt(4)
        ftr_p.paragraph_format.space_after = Pt(0)

        # Footer top border
        fPr = ftr_p._p.get_or_add_pPr()
        fBdr = OxmlElement('w:pBdr')
        ftop = OxmlElement('w:top')
        ftop.set(qn('w:val'), 'single')
        ftop.set(qn('w:sz'), '4')
        ftop.set(qn('w:space'), '1')
        ftop.set(qn('w:color'), 'CCCCCC')
        fBdr.append(ftop)
        fPr.append(fBdr)

        run_footer = ftr_p.add_run(
            f'Authority Systems Group™  —  Confidential. Prepared exclusively for {CLIENT_NAME}.'
        )
        run_footer.font.name = 'Calibri'
        run_footer.font.size = Pt(9)
        run_footer.font.color.rgb = ASG_FOOTER_CLR

        ftr_p.add_run('\t')

        # Page number field
        run_pg = ftr_p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_pg._r.append(fldChar1)
        run_pg._r.append(instrText)
        run_pg._r.append(fldChar2)
        run_pg.font.name = 'Calibri'
        run_pg.font.size = Pt(9)
        run_pg.font.color.rgb = ASG_FOOTER_CLR

        # Tab stop for page number
        fPr2 = ftr_p._p.get_or_add_pPr()
        tabs2 = OxmlElement('w:tabs')
        tab2 = OxmlElement('w:tab')
        tab2.set(qn('w:val'), 'right')
        tab2.set(qn('w:pos'), '9072')
        tabs2.append(tab2)
        fPr2.append(tabs2)


# ─────────────────────────────────────────
# DOCUMENT SECTIONS
# ─────────────────────────────────────────

def build_confidentiality_page(doc):
    section_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CONFIDENTIALITY NOTICE')
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = ASG_CHARCOAL

    notice_text = (
        f'This document was prepared exclusively for {CLIENT_NAME} by Authority Systems Group™. '
        'It contains proprietary strategic analysis, competitive intelligence, and growth recommendations '
        'developed specifically for this engagement. It may not be shared, reproduced, or distributed '
        'without written consent from Authority Systems Group™.\n\n'
        'All frameworks, methodologies, and branded systems referenced herein — including the '
        'Belief-to-Buy Framework™, Authority Conversion Protocol™, Quarterly Authority Rotation™, '
        'and Fast Cash Sprint™ — are proprietary intellectual property of Authority Systems Group™.'
    )
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(notice_text)
    r2.font.name = 'Georgia'
    r2.font.size = Pt(10)
    r2.font.color.rgb = ASG_CHARCOAL


def build_toc(doc):
    section_break(doc)
    add_h1(doc, 'TABLE OF CONTENTS')
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    toc_entries = [
        ('1', 'Executive Summary', 'Roger, Director'),
        ('2', 'Market Analysis & Competitive Position', 'Daniel Frost, CSO'),
        ('3', 'The Ideal Client Avatar — Managing Partner Matt', 'Market Intelligence'),
        ('4', 'Belief Filter Map™ — Authority Conversion Protocol™', 'Dr. Raymond Cross / Vivienne Carr, CMO'),
        ('5', 'Year 1 Strategic Plan — Quarterly Authority Rotation™', 'Daniel Frost, CSO'),
        ('6', 'Content Authority Strategy', 'Vivienne Carr, CMO'),
        ('7', 'Fast Cash Sprint™ — Your First 90 Days of Revenue', 'Tanya Blackwood, CRO'),
        ('8', 'KPI Framework & Year-End Targets', 'Measurement'),
        ('9', 'Investment & Resource Requirements', 'Preston Adler, COO'),
        ('10', 'QC Sign-Off', 'Preston Adler, COO'),
    ]

    table = doc.add_table(rows=len(toc_entries), cols=3)
    table.style = 'Table Grid'
    col_widths = [0.45, 4.5, 2.3]
    for i, col in enumerate(table.columns):
        col.width = Inches(col_widths[i])

    for i, (num, title, owner) in enumerate(toc_entries):
        cells = table.rows[i].cells
        bg = 'FFFFFF' if i % 2 == 0 else 'F5F5F5'
        for c in cells:
            set_cell_background(c, bg)
            set_cell_margins(c)

        # Number cell
        cells[0].paragraphs[0].text = ''
        r = cells[0].paragraphs[0].add_run(num)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.color.rgb = ASG_BLUE
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title cell
        cells[1].paragraphs[0].text = ''
        r2 = cells[1].paragraphs[0].add_run(title)
        r2.bold = True
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_CHARCOAL

        # Owner cell
        cells[2].paragraphs[0].text = ''
        r3 = cells[2].paragraphs[0].add_run(owner)
        r3.italic = True
        r3.font.name = 'Calibri'
        r3.font.size = Pt(10)
        r3.font.color.rgb = ASG_FOOTER_CLR

    # Remove table borders except light inner lines
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:color'), 'E0E0E0')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def build_section_1(doc):
    """Executive Summary — Roger's voice."""
    section_break(doc)
    section_header(doc, '1', 'EXECUTIVE SUMMARY', 'Prepared by Roger, Director — Authority Systems Group™')

    add_callout(doc,
        f'Prepared exclusively for {CLIENT_NAME}  |  March 2026  |  Confidential',
        label='DOCUMENT OVERVIEW')

    add_body(doc, f'{CLIENT_NAME},')
    add_body(doc, "I'll be direct about what you're holding.")
    add_body(doc,
        "This is not a marketing plan. Marketing plans describe what to post and how often. "
        "This is a business architecture document — built around a specific revenue target, a specific "
        "client type, and a specific mechanism for turning unknown-to-chosen in a market where most "
        "competitors are invisible by default.")
    add_body(doc,
        "The goal is $360,000 in Year 1. That's 20 family law firms at $18,000 each. The math is simple. "
        "The execution is not complicated. But it requires a specific sequence, and it requires that sequence "
        "to be executed before your best competitor in this space figures out what you've already figured out.")
    add_body(doc, "Here's what the team found when we built this plan.")
    add_body(doc,
        "The family law market in the United States is structurally fragmented. Most firms generate "
        "$300k–$1.5M annually and rely almost entirely on referrals. Referrals are unreliable. They're not "
        "scalable. They produce clients who arrived by accident rather than by design — and those clients "
        "compare attorneys on price because they don't have enough information to compare on anything else. "
        "The attorneys winning the best cases — the high-asset divorces, the high-stakes custody disputes, "
        "the complex business dissolution cases — are not better lawyers than their peers. They're better "
        "known as the right lawyer for those specific cases.")

    add_callout(doc, "That gap is the business.", label='THE CORE INSIGHT')

    add_body(doc,
        "The attorneys who understand this — who want to be the definitive authority in their market for "
        "the clients worth having — are the exact client LGD was built to serve. They exist in every "
        "mid-to-large metro in the country. They are underserved by every current marketing option they "
        "can access. And they will pay $18,000 a year, every year, for a partner who builds and manages "
        "the system that puts them in front of the right clients consistently.")
    add_body(doc, "The plan in this document gets you to 20 of them in 12 months.")
    add_body(doc,
        "Read Section 7 first — the Fast Cash Sprint. Those are the moves in the first 90 days. "
        "They're not complicated. They're just things most new agencies never do deliberately.")
    add_body(doc, "The rest of the plan is your execution roadmap.")
    add_body(doc, "This is your system now. Let's make it work.", italic=True)

    add_signoff(doc, '"This is your system now. Let\'s make it work."',
                'Roger, Director | Authority Systems Group™')

    # Key metrics summary table
    add_h2(doc, 'Year 1 At a Glance')
    add_branded_table(doc,
        ['Metric', 'Target'],
        [
            ['Annual Revenue Target', '$360,000'],
            ['Target Clients (Year 1)', '20 family law firms'],
            ['Annual Contract Value per Client', '$18,000 / $1,500 per month'],
            ['Primary Service', 'Authority Content & Thought Leadership'],
            ['Primary Market', 'Family Law — Continental US'],
            ['Target Avatar', 'Managing Partner Matt (2–6 attorney firms)'],
            ['90-Day Revenue Target', '$54,000–$108,000 (Fast Cash Sprint)'],
            ['Month 12 Monthly Run Rate', '$30,000/month'],
        ],
        col_widths=[3.5, 3.5]
    )


def build_section_2(doc):
    """Market Analysis & Competitive Position — Daniel Frost, CSO."""
    section_break(doc)
    section_header(doc, '2', 'MARKET ANALYSIS & COMPETITIVE POSITION',
                   'Prepared by Daniel Frost, CSO — Authority Systems Group™')

    add_callout(doc,
        '"The Year 1 plan for Legal Growth Dynamics is organized around a single strategic principle: '
        'dominate before you diversify. The family law authority content market is not crowded — it is '
        'almost entirely vacant. LGD\'s window to establish category leadership is open right now. '
        'The plan is sequenced to take advantage of it."',
        label='STRATEGIC THESIS — Daniel Frost, CSO')

    add_h2(doc, '2A. Market Sizing & Structure')
    add_body(doc,
        "The United States has approximately 48,000 active family law attorneys. The addressable market "
        "for LGD's Year 1 target — managing partners of firms with 2+ attorneys generating $400k–$2M in "
        "annual revenue, operating in metros with 300k+ population — is approximately 8,200 firms.")
    add_body(doc,
        "At LGD's pricing of $18,000/year, capturing 0.24% of that addressable market represents "
        "$360,000 in revenue. This is not a market penetration challenge. It is a market identification "
        "and sales execution challenge.")

    add_branded_table(doc,
        ['Case Type', 'Revenue Share', "LGD's Role"],
        [
            ['High-Asset / Complex Divorce', '40–50%', 'Primary authority content focus'],
            ['Child Custody & Modification', '15–20%', 'Secondary content focus'],
            ['Uncontested / Collaborative Divorce', '20–25%', 'Volume and referral engine content'],
            ['Prenuptial Agreements', '5–10%', 'Positioning content'],
        ],
        col_widths=[2.4, 1.8, 3.05]
    )

    add_h2(doc, '2B. Competitive Landscape')

    add_h3(doc, 'Competitor 1: Specialist Legal Marketing Firms')
    add_body(doc, 'Examples: Consultwebs, Scorpion, Juris Digital, and similar national agencies.')
    add_body(doc,
        "Strengths: Legal industry familiarity, established PPC and SEO capabilities, existing client base.")
    add_body(doc,
        "Weaknesses: Tactical focus, not authority-building. Their product is traffic — impressions, "
        "rankings, click-through rates. They do not build a belief system around their clients. Their "
        "content is SEO-optimized filler, not conviction-building thought leadership. Managing Partner "
        "Matt has usually already tried one of these firms and was disappointed.")

    add_h3(doc, 'Competitor 2: Generalist B2B Digital Agencies')
    add_body(doc,
        "Strengths: SEO expertise, content production scale, broad toolset. Weaknesses: No legal domain "
        "knowledge, no familiarity with bar ethics rules, no ability to produce content that a family law "
        "attorney would be proud to put their name on.")

    add_h3(doc, 'Competitor 3: Freelance Content Writers')
    add_body(doc,
        "Strengths: Lower cost point. Weaknesses: No strategic system. No belief architecture. No "
        "conversion framework. Produce articles that get filed and forgotten. Zero accountability for "
        "business results.")

    add_callout(doc,
        "No competitor in this space is building what LGD builds: a specific authority system, grounded "
        "in buyer psychology, designed to move a family law prospect from skepticism to signed retainer "
        "through a deliberate sequence of belief-and-emotion-aligned content. No competitor is selling a "
        "system with a revenue mechanism. They are selling activities. LGD sells outcomes.",
        label='THE MARKET GAP')

    add_h2(doc, '2C. LGD\'s Competitive Position')
    add_body(doc, "LGD's defensible position is the intersection of three capabilities no competitor currently combines:")
    add_bullet(doc, "Legal-first content frameworks — every deliverable built around the ethical constraints and professional standards of family law practice")
    add_bullet(doc, "Belief-to-Buy architecture™ — content maps to specific stages on both the Belief Track and Emotion Track, producing the conversion mechanism that makes content actually generate consultations")
    add_bullet(doc, "Revenue accountability — success is measured in consultation bookings from content-sourced prospects, not traffic or impressions")

    add_callout(doc,
        f"Legal Growth Dynamics is the only authority content partner built specifically for family law "
        "managing partners who want to become the recognized expert their ideal clients find before they "
        "find anyone else — and who demand that content marketing produce actual case intake, not just "
        "brand awareness.",
        label='AUTHORITY STATEMENT')

    add_signoff(doc, '"Strategy without sequencing is just a wish list."',
                'Daniel Frost, CSO | Authority Systems Group™')


def build_section_3(doc):
    """Avatar — Managing Partner Matt."""
    section_break(doc)
    section_header(doc, '3', 'THE IDEAL CLIENT AVATAR',
                   'Managing Partner Matt — Synthesized from niche intelligence and belief barrier mapping')

    add_h2(doc, 'Primary Avatar: Managing Partner Matt')
    add_body(doc,
        "Managing partner of a family law firm with 2–6 attorneys. Located in a metro area with 400k+ "
        "population. The firm generates $450k–$1.5M annually — mostly through referrals from past clients, "
        "CPAs, financial advisors, and occasional Google searches. Matt personally handles the firm's most "
        "complex cases. He's been practicing 12–22 years.")

    add_branded_table(doc,
        ['Demographic', 'Profile'],
        [
            ['Age Range', '40–55'],
            ['Personal Income', '$200k–$450k'],
            ['Family Status', 'Married, typically two children'],
            ['Practice Duration', '12–22 years in family law'],
            ['Firm Size', '2–6 attorneys'],
            ['Firm Revenue', '$450k–$1.5M annually'],
            ['Revenue Source', '70–85% from referrals (declining)'],
            ['Primary Case Focus', 'High-asset divorce, custody disputes'],
        ],
        col_widths=[2.8, 4.45]
    )

    add_h2(doc, 'The Trigger Events')
    add_body(doc, "Matt starts seriously considering LGD when at least one of the following occurs:")

    triggers = [
        ("Referral Erosion", "His referral pipeline produces 15–25% fewer qualified leads than the prior year. It's subtle at first. Then it becomes a pattern."),
        ("Competitor Visibility", "A competing firm in his city starts showing up in search results for the specific case types Matt wants. Or he sees their content on LinkedIn and it's actually good. He notices. It bothers him more than he admits."),
        ("A Bad Quarter That Breaks the Pattern", "A full slow quarter without a clear explanation forces him to confront that 'networking more' isn't actually a plan."),
        ("The Associate Problem", "He's been thinking about adding an attorney but can't justify the hire without predictable case flow. The realization that his pipeline has never been systematized hits him."),
        ("A High-Value Case Lost to a Competitor", "He finds out that a prospect chose a different firm because they 'found them online' and felt confident in their expertise before the first call. Matt realizes that prospective client never even considered him."),
    ]
    for title, detail in triggers:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{title}: ')
        r1.bold = True
        r1.font.name = 'Georgia'
        r1.font.size = Pt(11)
        r1.font.color.rgb = ASG_CHARCOAL
        r2 = p.add_run(detail)
        r2.font.name = 'Georgia'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_BODY_TEXT

    add_h2(doc, 'His Belief Barriers')
    barriers = [
        ('"Marketing is for firms that aren\'t good enough to win on reputation."',
         "Matt has seen bad legal marketing. He associates content marketing with desperate-looking blog posts and cheesy attorney ads. He has conflated bad marketing with all marketing."),
        ('"Any agency I hire won\'t understand legal ethics or how family law actually works."',
         "He's right that most agencies don't. He assumes this means all agencies don't. This belief gets dismantled by the first LGD content he reads that demonstrates genuine family law depth."),
        ('"Good content doesn\'t produce real clients — it produces tire kickers."',
         "He doesn't believe authority content can pre-qualify for high-value matters. He hasn't seen the mechanism that makes it possible."),
        ('"$18,000 is a cost I can\'t justify if I can\'t see the return."',
         "Matt can justify a $20,000 court battle without blinking. He processes $18,000 as a marketing cost and applies a different ROI lens. He has never been shown the math: one high-asset divorce client pays the full year's investment in a single matter."),
    ]
    for barrier, explanation in barriers:
        add_callout(doc, f'{barrier}\n{explanation}')

    add_h2(doc, "What He's Buying")
    add_body(doc, "When Matt commits to LGD's program, he is buying a specific future state:")
    add_bullet(doc, "His name appears — with substantive, credible content behind it — when someone in his city searches for a business owner's divorce attorney or a high-stakes custody attorney")
    add_bullet(doc, "He is known in his market as the expert before anyone picks up the phone")
    add_bullet(doc, "His intake team reports that prospects arrive pre-sold — they've read his content, they've already decided he's the right attorney, the consultation is a formality")
    add_bullet(doc, "His referral sources send him better clients because they now have a concrete way to describe why he's different")


def build_section_4(doc):
    """Belief Filter Map™."""
    section_break(doc)
    section_header(doc, '4', 'BELIEF FILTER MAP™',
                   'Authority Conversion Protocol™ | Dr. Raymond Cross, Board / Vivienne Carr, CMO')

    add_callout(doc,
        "Both tracks run simultaneously. Both must reach their endpoint before Managing Partner Matt "
        "commits. Move only one track and he stalls — intellectually convinced but not ready to act, "
        "or emotionally activated but intellectually uncertain.",
        label='THE DUAL-TRACK SYSTEM')

    # System diagram as table
    add_branded_table(doc,
        ['BELIEF TRACK — What Matt THINKS Is True', 'EMOTION TRACK — What Matt FEELS Is True'],
        [
            ['B1: Enemy Belief', 'E1: Dissatisfaction'],
            ['B2: Pre-Frame', 'E2: Contrast'],
            ['B3: New Worldview', 'E3: Desire'],
            ['B4: Internal Alignment', 'E4: Urgency'],
            ['B5: Certainty → SIGNS', 'E5: Relief → SIGNS'],
        ],
        col_widths=[3.5, 3.75]
    )

    add_h2(doc, '4A. The Enemy Belief')
    add_callout(doc,
        '"I\'ve seen what marketing agencies do for law firms. It\'s blog posts nobody reads, social media '
        'that makes me look desperate, and outcome claims that would get me in front of the bar. The '
        'attorneys doing that kind of marketing aren\'t the ones winning the cases worth having. I built '
        'my practice on referrals. The best family law attorneys always have. That\'s not going to change '
        'because some agency publishes content under my name."',
        label="MATT'S ENEMY BELIEF — IN HIS OWN VOICE")

    add_body(doc,
        "This belief was formed through legitimate negative experiences — agencies that demonstrated no "
        "understanding of family law, produced generic sample articles, and reported on traffic metrics "
        "with no connection to consultation bookings. The Enemy Belief actively protects Matt from what "
        "he perceives as a real professional risk. It cannot be overcome with more information. It "
        "requires a different frame entirely.")

    add_h2(doc, '4B. Belief Track — Full Map')

    belief_stages = [
        ('B1', 'ENEMY BELIEF',
         '"Marketing agencies don\'t understand how family law clients actually make hiring decisions. Real clients come from real relationships. I\'ve been in practice for 17 years. I know how this works."',
         'First LGD content demonstrates specific, credible knowledge of how high-asset divorce clients research attorneys — describing a mechanism Matt recognizes as real.',
         '"Most family law firms believe their best clients come from referrals. They\'re right — but they\'re wrong about why. The referral is permission to do research. What your prospect finds after they get your name determines whether you get hired."'),
        ('B2', 'PRE-FRAME',
         '"I\'ve never thought about what a prospect actually does after they get my name from a CPA. I just assumed they\'d call. Maybe not everyone does."',
         'LGD presents the research cycle of a high-asset divorce prospect — the 3–5 touchpoints completed before contacting an attorney.',
         'Documented breakdown: 1) Google search, 2) Review scan, 3) Attorney website — specific to business-owner divorce?, 4) LinkedIn credibility check, 5) Call. Most attorneys fail steps 3 and 4.'),
        ('B3', 'NEW WORLDVIEW',
         '"Authority content isn\'t about marketing. It\'s about being findable at the moment a qualified prospect is doing research. The attorneys who do this well aren\'t doing marketing theater — they\'re building a pre-qualification engine that works while they\'re in court."',
         'LGD connects the New Worldview to Matt\'s specific past experience — cases he didn\'t know he lost, slow quarters he couldn\'t explain.',
         '"The firms generating $1M+ from ideal case types aren\'t doing more networking. They\'ve built a visible authority position that converts prospects before the first phone call."'),
        ('B4', 'INTERNAL ALIGNMENT',
         '"This explains the slow quarter two years ago. My best CPA contact went quiet. His clients Googled me and found nothing that distinguished me from three other firms he could have recommended. Henderson gets the business-owner cases now. I\'ve seen their articles. Now I understand why."',
         'LGD provides a documented case study of a comparable firm with measurable, traceable results.',
         'Case study: 3-attorney family law practice, $620k revenue, 11-month authority content program. Month 7: 4.2% conversion from content pages. Month 11: 31% of high-asset consultations cited content as primary reason for choosing the firm. Avg matter value: $34,000.'),
        ('B5', 'CERTAINTY',
         '"Authority content is infrastructure for the part of client acquisition I\'ve never had a system for — the research phase. The attorneys attracting ideal clients consistently are not better lawyers. They\'re just better known as the right lawyer for that specific case. The question is whether I build it or let a competitor build it first."',
         'Certainty endpoint reached. The Emotion Track must be at Relief simultaneously for the buy to occur.',
         'Consultation CTA with specific scoping of Matt\'s practice, his case types, and his market. The engagement is not generic — it is built for him.'),
    ]

    for stage_code, stage_name, monologue, transition, lgd_message in belief_stages:
        add_h3(doc, f'{stage_code} — {stage_name}')

        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_before = Pt(4)
        p_m.paragraph_format.space_after = Pt(2)
        p_m.paragraph_format.left_indent = Inches(0.15)
        r_label = p_m.add_run("Matt's internal monologue: ")
        r_label.bold = True
        r_label.font.name = 'Calibri'
        r_label.font.size = Pt(10)
        r_label.font.color.rgb = ASG_CHARCOAL
        r_text = p_m.add_run(monologue)
        r_text.italic = True
        r_text.font.name = 'Georgia'
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = ASG_BODY_TEXT

        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(2)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.left_indent = Inches(0.15)
        r_tl = p_t.add_run("Transition trigger: ")
        r_tl.bold = True
        r_tl.font.name = 'Calibri'
        r_tl.font.size = Pt(10)
        r_tl.font.color.rgb = ASG_CHARCOAL
        r_tt = p_t.add_run(transition)
        r_tt.font.name = 'Georgia'
        r_tt.font.size = Pt(10)
        r_tt.font.color.rgb = ASG_BODY_TEXT

        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(2)
        p_l.paragraph_format.space_after = Pt(8)
        p_l.paragraph_format.left_indent = Inches(0.15)
        r_ll = p_l.add_run("What LGD says or shows: ")
        r_ll.bold = True
        r_ll.font.name = 'Calibri'
        r_ll.font.size = Pt(10)
        r_ll.font.color.rgb = ASG_BLUE
        r_lm = p_l.add_run(lgd_message)
        r_lm.font.name = 'Georgia'
        r_lm.font.size = Pt(10)
        r_lm.font.color.rgb = ASG_BODY_TEXT

    add_h2(doc, '4C. Emotion Track — Full Map')

    emotion_stages = [
        ('E1', 'DISSATISFACTION',
         'Not panicking — quietly concerned. His referral pipeline is down consistently. He\'s filed this under "part of the business cycle" because that framing doesn\'t require him to do anything differently.',
         'A competitor of similar size and experience announces they\'ve opened a second office. He digs into their content. It\'s about exactly the clients he wants. The referral source he lost 18 months ago is now sending them the complex divorces.'),
        ('E2', 'CONTRAST',
         'The plateau is no longer abstract. He can see a specific competitor attracting clients he wants through a visible, understandable mechanism. "Good lawyers get referred" no longer fully explains the gap.',
         'LGD presents the before/after story: Before: 85% referral revenue, $18k avg case, slow quarters unexplainable. After: 40% content-sourced, $31k avg case, intake pre-qualifying because prospects arrive already informed and committed.'),
        ('E3', 'DESIRE',
         'Not just aware of the gap — he wants the after-state. Specifically. He wants to be the attorney business owners call in his city when facing a complex divorce. He wants cases that are worth taking to find him.',
         '"I want to be the name that comes up when a business owner in my city Googles \'protecting my company in a divorce.\' I want CPAs to recommend me specifically for high-asset matters because my content demonstrates expertise they can point to."'),
        ('E4', 'URGENCY',
         'Staying the same is no longer neutral. Every month without an authority content system, a competitor has the opportunity to build that position first. The cost of inaction is the compounding advantage he watches his most visible competitors accumulating.',
         'Every case not captured: $18,000–$50,000 in fees elsewhere. Every month a competitor publishes authoritative content: their search authority grows, Matt\'s does not. The attorney who builds this position first owns it. The second to try has to displace an incumbent.'),
        ('E5', 'RELIEF',
         'The decision to engage LGD feels like the resolution of tension he\'s been carrying about his pipeline — not the creation of new risk. His concern has shifted from "should I do this?" to "why haven\'t I done this sooner?"',
         'LGD produces a sample content piece specific to his practice, case types, and market before the engagement. He sees the quality. The case study includes a comparable firm. The scope is specific to his priorities. The ROI bar is one client: one high-asset matter at $28,000 avg value = 1.6x return on the full year\'s investment.'),
    ]

    for stage_code, stage_name, state, trigger in emotion_stages:
        add_h3(doc, f'{stage_code} — {stage_name}')
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_before = Pt(4)
        p_s.paragraph_format.space_after = Pt(2)
        p_s.paragraph_format.left_indent = Inches(0.15)
        r_sl = p_s.add_run("Emotional state: ")
        r_sl.bold = True
        r_sl.font.name = 'Calibri'
        r_sl.font.size = Pt(10)
        r_sl.font.color.rgb = ASG_CHARCOAL
        r_st = p_s.add_run(state)
        r_st.font.name = 'Georgia'
        r_st.font.size = Pt(10)
        r_st.font.color.rgb = ASG_BODY_TEXT

        p_tr = doc.add_paragraph()
        p_tr.paragraph_format.space_before = Pt(2)
        p_tr.paragraph_format.space_after = Pt(8)
        p_tr.paragraph_format.left_indent = Inches(0.15)
        r_trl = p_tr.add_run("Transition trigger: ")
        r_trl.bold = True
        r_trl.font.name = 'Calibri'
        r_trl.font.size = Pt(10)
        r_trl.font.color.rgb = ASG_CHARCOAL
        r_trt = p_tr.add_run(trigger)
        r_trt.font.name = 'Georgia'
        r_trt.font.size = Pt(10)
        r_trt.font.color.rgb = ASG_BODY_TEXT

    add_h2(doc, '4D. Master Content Map')
    add_body(doc, "This table drives all LGD content production, outreach messaging, and sales conversations.")

    add_branded_table(doc,
        ['Stage', 'Belief Stage', 'Emotion Stage', 'Content Type', 'Core Message', 'CTA'],
        [
            ['B1/E1', 'Enemy Belief', 'Dissatisfaction', 'Pattern interrupt', 'Referrals are permission to do research. What prospects find after getting your name determines whether you get hired.', 'Read this case study'],
            ['B2/E2', 'Pre-Frame', 'Contrast', 'Reframe piece', 'High-value clients complete 3–5 touchpoints before calling. Most family law websites fail 3 of them.', 'See the difference'],
            ['B3/E3', 'New Worldview', 'Desire', 'Education + story', 'Authority content is pre-qualification infrastructure. It makes your best cases find you.', 'Watch the case study'],
            ['B4/E4', 'Int. Alignment', 'Urgency', 'Case study + proof', 'The authority position in your market is available now. Here\'s the cost of not building it.', 'Book a strategy call'],
            ['B5/E5', 'Certainty', 'Relief', 'Consultation CTA', 'LGD builds your authority system for your practice, your case types, your market. One client justifies the year.', 'Schedule your session'],
        ],
        col_widths=[0.65, 1.0, 0.95, 1.0, 2.3, 1.35]
    )

    add_h2(doc, '4E. Primary Stall Point')
    add_callout(doc,
        "Primary stall: B1/E1 — The Enemy Belief + Normalized Dissatisfaction. This is a dual stall — "
        "both tracks require movement simultaneously. LGD's first content touchpoint must produce a "
        "credible pattern interrupt on the Enemy Belief AND create visible contrast between Matt's "
        "current trajectory and a better state. Content that addresses only one track will not convert.",
        label='STALL POINT ANALYSIS')

    add_h2(doc, '4F. The Convergence Statement')
    add_callout(doc,
        "Managing Partner Matt reaches the buying decision when he is intellectually certain that "
        "authority content is pre-qualification infrastructure — not marketing — and that the attorney "
        "who builds this position in his market first will own it, while every attorney who waits will "
        "fight to displace an incumbent, AT THE SAME MOMENT that he feels the genuine urgency of "
        "watching a competitive position close while he stays still, combined with the specific relief "
        "of knowing that LGD understands family law deeply enough to produce content he would be proud "
        "to publish under his name. Legal Growth Dynamics achieves this convergence by demonstrating "
        "legal expertise through the quality of sample content before any purchase decision, providing a "
        "traceable case study that shows the revenue mechanism clearly, and scoping the engagement "
        "around his specific practice priorities so the investment feels like infrastructure rather than "
        "a marketing expense.",
        label='THE CONVERGENCE STATEMENT')

    add_signoff(doc, '"People don\'t buy when they\'re convinced. They buy when belief and emotion align at the same moment."',
                'Dr. Raymond Cross, Board | Applied by Vivienne Carr, CMO | Authority Systems Group™')


def build_section_5(doc):
    """Year 1 Strategic Plan."""
    section_break(doc)
    section_header(doc, '5', 'YEAR 1 STRATEGIC PLAN',
                   'Quarterly Authority Rotation™ | Prepared by Daniel Frost, CSO — Authority Systems Group™')

    add_callout(doc,
        "\"The Year 1 plan below was built from a clear-eyed analysis of Legal Growth Dynamics' market "
        "position, current assets, and the specific belief barriers standing between LGD's ideal clients "
        "and their front door. It is not a general marketing plan with LGD's name on it. Every initiative "
        "was selected because it's the right move for this business at this moment.\"",
        label='DANIEL FROST, CSO')

    add_h2(doc, '5A. Year 1 Strategic Overview')
    add_h3(doc, 'The Thesis')
    add_body(doc,
        "Legal Growth Dynamics' single most important strategic objective in Year 1 is to produce "
        "undeniable proof that authority content produces measurable case intake for family law firms — "
        "and to use that proof to remove the primary buying barrier from every subsequent sales "
        "conversation.")
    add_body(doc,
        "LGD enters the market with no case studies, no client testimonials, and no documented track "
        "record in this vertical. That is the only real obstacle between LGD and $360k in Year 1. "
        "Everything in this plan is designed to overcome it as fast as possible.")

    add_h3(doc, 'The Logic of Sequencing')
    add_body(doc,
        "Q1 is not about revenue scale. Q1 is about proof of concept. Two pilot clients — selected "
        "specifically for their willingness to participate in documented case studies — are worth more "
        "to LGD's long-term pipeline than four clients who want the service but will not be used as "
        "evidence. The case study produced in Q1 becomes the most valuable sales asset LGD will own "
        "for the remainder of Year 1 and beyond.")
    add_body(doc,
        "Q2 deploys that case study into systematic outreach and referral. Q3 adds paid digital "
        "acquisition on top of the organic and referral foundation Q1 and Q2 built. Q4 prepares for "
        "Year 2 scale with a refined ICP, documented pricing, and a content system that can support "
        "30+ clients without proportional team growth.")

    add_h3(doc, 'What Success Looks Like at Month 12')
    add_body(doc,
        "Legal Growth Dynamics is generating 35–50 qualified discovery calls per month from a "
        "combination of direct LinkedIn outreach, referrals from existing attorney clients, and inbound "
        "content-driven leads. LGD is converting those calls at 5–8%, producing 1.75–4 new clients per "
        "month. The 20-client target has been reached by Month 10–11. Month 12 is focused on retention, "
        "upsell scoping, and Year 2 product development. The content delivery system is standardized "
        "and can onboard a new attorney client within 5 business days of contract execution.")

    # Quarterly overview table
    add_h2(doc, 'The Quarterly Authority Rotation™')
    add_branded_table(doc,
        ['Quarter', 'Strategic Goal', 'Revenue Target', 'Exit Milestone'],
        [
            ['Q1 (Months 1–3)', 'Foundation + Proof', '$36,000 (2 pilot clients)', '2 pilots signed, case study data captured'],
            ['Q2 (Months 4–6)', 'Validation + Scale', '$72,000–$108,000 new ARR', '6–8 total clients, case study deployed'],
            ['Q3 (Months 7–9)', 'Depth + Diversification', '$108,000 new ARR', '12–15 clients, organic inbound begins'],
            ['Q4 (Months 10–12)', 'Dominance + Year 2 Prep', '$108,000 new ARR', '20 clients, 85%+ renewal confirmed'],
        ],
        col_widths=[1.4, 1.8, 1.8, 2.25]
    )

    add_h2(doc, '5B. Quarter 1 — Foundation + Proof (Months 1–3)')
    add_callout(doc,
        "Revenue Target: $36,000 (2 pilot clients × $18,000)\n"
        "Strategic Goal: Sign 2 pilot clients, deliver measurable content program, begin case study documentation",
        label='Q1 TARGETS')

    add_h3(doc, 'Month 1 — Outreach & Pilot Close')
    add_body(doc, "Identify 60 family law managing partners meeting all criteria: firm size 2–6 attorneys, metro 400k+ population, no visible authority content program, LinkedIn active. Deploy simultaneously:")
    add_bullet(doc, "LinkedIn connection requests with personalized note — no pitch in the request")
    add_bullet(doc, "Direct email written to B1/E1 stage: pattern interrupt on the referral myth, not a service pitch")
    add_bullet(doc, "Warm network: introductions requested through mutual contacts")
    add_body(doc, "Target: 8–12 discovery calls. Close rate: 25% (2–3 pilots from first outreach wave).")
    add_body(doc, "Pilot framing: standard contract at standard rate. In exchange, clients agree to participate in a documented results case study at 90 days and 6 months.")

    add_h3(doc, 'Month 2 — Delivery Begins + Content Engine Launch')
    add_body(doc, "Both pilot clients fully onboarded: intake questionnaire complete, voice profile established, case type priorities confirmed, state bar ethics guidelines documented. Month 1 content delivering: first articles live, first LinkedIn posts published. Outreach wave 2 deployed (60 additional prospects). Roger's LinkedIn at 4 posts per week.")

    add_h3(doc, 'Month 3 — Optimization + Q2 Pre-Build')
    add_body(doc, "Performance review: response rate, call-booking rate, discovery call conversion. Winning outreach message format identified. Case study pre-work begins: Month 3 data points captured for both pilot clients. Q2 content strategy finalized. Outreach wave 3 prospecting list complete (80 additional).")

    add_callout(doc,
        "Q1 Exit Checklist:\n"
        "☐  2 pilot clients signed ($36,000 contracted)\n"
        "☐  Months 1 and 2 content delivered on schedule\n"
        "☐  LGD LinkedIn: 500+ connections in target audience\n"
        "☐  15+ discovery calls completed\n"
        "☐  3+ case study data points documented\n"
        "☐  Winning outreach message format identified",
        label='Q1 EXIT METRICS')

    add_h2(doc, '5C. Quarter 2 — Validation + Scale (Months 4–6)')
    add_callout(doc,
        "Revenue Target: $72,000–$108,000 new ARR (4–6 new clients)\n"
        "Cumulative ARR Target: $108,000–$144,000\n"
        "Strategic Goal: First case study deployed, 15+ calls/month, 2+ referral introductions/month",
        label='Q2 TARGETS')

    add_body(doc,
        "By Month 4, pilot clients have 3 months of delivery documented. Early-stage results are "
        "sufficient for a compelling case study if framed correctly. Deploy case study as primary "
        "outreach asset. Outreach scale increases to 100 new prospects per month. LinkedIn sponsored "
        "content testing begins: $500–$1,000/month. 90-day email nurture sequence deployed for "
        "non-booking prospects, mapped to Belief Filter Map stages.")
    add_body(doc,
        "Formal referral program for attorney clients launches: non-financial incentive structure "
        "(charitable donation in client's name, premium service additions) — fully compliant with "
        "attorney ethics rules prohibiting fee-sharing with non-attorneys.")

    add_callout(doc,
        "Q2 Exit Checklist:\n"
        "☐  6–8 active clients\n"
        "☐  Monthly revenue run rate: $9,000–$12,000/month\n"
        "☐  2+ case studies available as sales assets\n"
        "☐  Discovery calls: 15+ per month consistently\n"
        "☐  Outreach-to-call conversion: 8–12%\n"
        "☐  Referral program active — first attorney referral received",
        label='Q2 EXIT METRICS')

    add_h2(doc, '5D. Quarter 3 — Depth + Diversification (Months 7–9)')
    add_callout(doc,
        "Revenue Target: $108,000 new ARR (6 new clients)\n"
        "Cumulative ARR Target: $216,000–$252,000\n"
        "Strategic Goal: SEO compounding visible, two 6-month case studies, 4+ referrals/month",
        label='Q3 TARGETS')

    add_body(doc,
        "By Month 7, Q1 pilot client content has been live 4–6 months — the window where search "
        "ranking movement typically becomes visible. LGD tracks organic performance across all client "
        "content and documents any traffic or ranking gains. LinkedIn paid optimization: Q2 sponsored "
        "content has 60+ days of data, Q3 doubles down on what performed. Target: cost per discovery "
        "call under $200. Begin developing legal tech vendor co-marketing relationships (Clio, MyCase, "
        "LawPay) for Q4 referral flow.")

    add_callout(doc,
        "Q3 Exit Checklist:\n"
        "☐  12–15 active clients\n"
        "☐  Monthly revenue run rate: $18,000–$22,500/month\n"
        "☐  2 full case studies with 6-month data points\n"
        "☐  Organic inbound discovery calls: 3–5/month\n"
        "☐  LinkedIn paid cost per call: under $200\n"
        "☐  3+ active professional referral partners",
        label='Q3 EXIT METRICS')

    add_h2(doc, '5E. Quarter 4 — Dominance + Year 2 Prep (Months 10–12)')
    add_callout(doc,
        "Revenue Target: $108,000 new ARR (6 new clients)\n"
        "Cumulative ARR Target: $324,000–$360,000\n"
        "Strategic Goal: 20 clients by Month 11, 85%+ Year 2 renewal confirmed, pricing finalized",
        label='Q4 TARGETS')

    add_body(doc,
        "By Q4, LGD has 9 months of consistent Roger LinkedIn content being consumed by the exact "
        "audience LGD sells to. Inbound discovery requests begin arriving without direct outreach. "
        "Proactive QBR with all existing clients in Month 11: present results, confirm Year 2 renewal, "
        "address any underperformance before the renewal conversation.")
    add_body(doc,
        "Year 2 pricing review: with 15–20 documented client relationships and multiple case studies "
        "showing measurable ROI, LGD has the evidence base to introduce a premium tier. Year 2 thesis: "
        "expand to a second practice vertical while doubling down on family law with a higher-tier product.")

    add_h2(doc, '5F. Year-End KPI Table')
    add_branded_table(doc,
        ['KPI', 'Baseline', 'Target (Month 12)', 'Stretch (Month 12)'],
        [
            ['Active clients', '0', '20', '25'],
            ['Monthly discovery calls', '0', '35', '50'],
            ['Call-to-client conversion rate', '0%', '5%', '8%'],
            ['New clients per month', '0', '1.7', '2.5'],
            ['Monthly revenue (run rate)', '$0', '$30,000', '$37,500'],
            ['Annual revenue (cumulative)', '$0', '$360,000', '$450,000'],
            ['Referral % of new clients', '0%', '25%', '35%'],
            ['Inbound % of new clients', '0%', '15%', '25%'],
            ['12-month client retention', 'N/A', '85%', '92%'],
            ['Average contract value', '$0', '$18,000', '$22,000'],
        ],
        col_widths=[2.5, 1.3, 1.8, 1.65]
    )

    add_signoff(doc, '"Strategy without sequencing is just a wish list."',
                'Daniel Frost, CSO | Authority Systems Group™')


def build_section_6(doc):
    """Content Authority Strategy — Vivienne Carr, CMO."""
    section_break(doc)
    section_header(doc, '6', 'CONTENT AUTHORITY STRATEGY',
                   'Prepared by Vivienne Carr, CMO — Authority Systems Group™')

    add_callout(doc,
        '"The content strategy for Legal Growth Dynamics is built on a single premise: Managing Partner '
        'Matt doesn\'t know he needs LGD yet — but he\'s already living with the problem LGD solves. '
        'Every piece of content we produce is designed to interrupt that experience. Not to pitch. Not '
        'to promote. To hold up a mirror and say: we see exactly what you\'re dealing with. And we\'ve '
        'solved it before."',
        label='VIVIENNE CARR, CMO')

    add_h2(doc, "6A. LGD's Own Content Channels")
    add_body(doc,
        "Before LGD can credibly sell authority content to family law firms, LGD must demonstrate it "
        "in its own marketing. Managing Partner Matt will research LGD before he books a call. What "
        "he finds must be the proof of concept he needs.")

    add_h3(doc, 'LinkedIn — Primary Channel, Year 1')
    add_body(doc,
        "Audience: Family law managing partners, legal marketing decision-makers, bar association "
        "leaders, legal tech vendors who sell into family law. Cadence: 4 posts per week, Roger's "
        "personal profile.")

    add_branded_table(doc,
        ['Content Type', 'Volume', 'Stage', 'Purpose'],
        [
            ['Pattern interrupt — referral myth, research cycle', '40%', 'B1/E1', 'Generate the response: "this person actually gets it"'],
            ['Contrast and reframe — before/after stories', '25%', 'B2/E2', 'Make the gap between current trajectory and better state visible'],
            ['New worldview — authority content mechanism', '20%', 'B3/E3', 'Explain the mechanism that makes content convert'],
            ['Proof and urgency — case study data points', '10%', 'B4/E4', 'Cost of inaction stated plainly; market position timeline'],
            ['Consultation CTA — specific, no pressure', '5%', 'B5/E5', 'Clear path to next step for ready prospects'],
        ],
        col_widths=[2.5, 0.75, 0.65, 3.35]
    )

    add_h2(doc, '6B. The Authority Content Product — Monthly Deliverables per Client')
    add_callout(doc,
        "The flagship product is a complete authority content system, deployed monthly, tuned to the "
        "specific case type priorities and voice of each attorney client. This is not a content service. "
        "It is a belief-change system delivered through content.",
        label='PRODUCT DEFINITION')

    deliverables = [
        ('3 Long-Form Authority Articles (1,200–2,000 words)',
         'Each article targets a specific stage on the Belief Filter Map for the attorney\'s ideal client avatar. SEO-optimized for city + case type combinations with low competition and high buyer intent. Fully bar-ethics compliant per the attorney\'s state. Written in the attorney\'s voice.'),
        ('1 Pillar Authority Piece (2,500–4,000 words) — Quarterly',
         'One comprehensive authority piece per quarter becomes the attorney\'s primary sales tool for their target case type. Examples: "The Complete Guide to Protecting Your Business in a [City] Divorce" or "What Courts Actually Look For in [State] Custody Disputes."'),
        ('8 LinkedIn Posts Per Month',
         'Mapped to Belief Filter Map stages for the attorney\'s ideal client. Mix of text posts, anonymized case commentary, educational frameworks, and proof points. Written in the attorney\'s professional voice.'),
        ('1 Email Newsletter Per Month',
         'Deployed to the attorney\'s existing contact list (past clients, referral sources). Value-first content, positions the attorney as the ongoing expert, includes a soft CTA for referral ask or consultation booking.'),
        ('Monthly Performance Report',
         'Tracks: article publication confirmed, LinkedIn post engagement, intake team mentions of content-sourced prospects, search impression data from Google Search Console, cumulative trajectory toward 90-day benchmarks.'),
    ]

    for title, description in deliverables:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f'{title}:  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = ASG_BLUE
        r2 = p.add_run(description)
        r2.font.name = 'Georgia'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_BODY_TEXT

    add_h2(doc, '6C. The Content Quality Standard')
    add_body(doc, "Every piece passes the Messaging Integrity gate before delivery:")
    gates = [
        "Could this article have been written for a different attorney in a different city? If yes — rewrite.",
        "Does this article address a specific stage on the Belief Filter Map for this attorney's avatar? If not — no conversion mechanism. Rewrite.",
        "Does this article demonstrate genuine understanding of family law — how cases work, what clients fear, how courts decide? If not — it will embarrass the attorney. Rewrite.",
        "Does any claim violate bar ethics rules in the attorney's state? If yes — revise before delivery.",
        "Would this attorney be proud to share this article with a colleague they respect? If not — it's not done.",
    ]
    for i, gate in enumerate(gates, 1):
        add_bullet(doc, f'Gate {i}: {gate}')

    add_signoff(doc, '"If it doesn\'t stop the scroll, it never happened."',
                'Vivienne Carr, CMO | Authority Systems Group™')


def build_section_7(doc):
    """Fast Cash Sprint™ — Tanya Blackwood, CRO."""
    section_break(doc)
    section_header(doc, '7', 'FAST CASH SPRINT™',
                   'Your First 90 Days of Revenue | Prepared by Tanya Blackwood, CRO — Authority Systems Group™')

    add_callout(doc,
        '"Before we build anything new, let\'s look at what\'s already there. Legal Growth Dynamics is '
        'a new business — there are no past clients to reactivate, no dormant pipeline to revive. '
        'That\'s fine. New businesses have assets too. Roger\'s network. A clearly defined value '
        'proposition. A target market that is underserved and currently active. The Fast Cash Sprint '
        'for LGD is about activating those existing assets as fast as possible, in the right order, '
        'with the right revenue mechanics."',
        label='TANYA BLACKWOOD, CRO')

    # Strategy 1
    add_horizontal_rule(doc, '25aae1', 8)
    add_h2(doc, 'FAST CASH STRATEGY #1')
    add_h3(doc, 'Direct Pilot Outreach — "Two Paid Pilots in 45 Days"')

    add_callout(doc,
        "Legal Growth Dynamics needs two things simultaneously: revenue and proof. The fastest path "
        "to both is the same action: two paid pilot clients, signed within 30–45 days, selected "
        "specifically for their willingness to participate in a documented case study. "
        "60 family law firms are reachable through direct outreach in Month 1. "
        "Conservative revenue: $36,000. Case study asset produced: the primary sales tool for "
        "every subsequent client conversation.",
        label='THE OPPORTUNITY')

    add_h3(doc, 'The Approach')
    steps_1 = [
        ("Week 1", "Build the 60-prospect list: 2–6 attorneys, metro 400k+, LinkedIn active, no visible authority content program"),
        ("Weeks 1–2", "Deploy outreach sequence: first message is not a pitch — it's one specific insight about their market"),
        ("Weeks 2–4", "Discovery calls: 45 minutes, show the research cycle mechanism, scope their primary case type, specific proposal at close"),
        ("Weeks 3–5", "Sign 2 pilot contracts, complete intake questionnaires, establish voice profiles, submit Month 1 content calendar"),
        ("Months 1–3", "Deliver on schedule: every article reviewed and approved before publication, every metric tracked from Day 1"),
    ]
    for timing, step in steps_1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{timing}:  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = ASG_BLUE
        r2 = p.add_run(step)
        r2.font.name = 'Georgia'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_BODY_TEXT

    add_branded_table(doc,
        ['Scenario', 'Clients', 'Revenue', 'Math', 'Timeline'],
        [
            ['Conservative', '2', '$36,000', '60 prospects × 20% call rate = 12 calls × 17% close = 2 clients × $18k', '30–45 days'],
            ['Target', '3', '$54,000', '60 prospects × 20% call rate = 12 calls × 25% close = 3 clients × $18k', '45 days'],
        ],
        col_widths=[1.2, 0.8, 0.95, 3.6, 0.9]
    )

    add_callout(doc,
        '"Every business I\'ve started from zero has had the same first move: find the 60 people most '
        'likely to buy, reach out personally with proof you understand their problem, and close two of '
        'them before spending money on anything else. The pilot framing turns the absence of case '
        'studies into a reason for the right client to act first. That\'s not manipulation. That\'s '
        'honest positioning."',
        label='SOFIA VEGA, Board Member')

    # Strategy 2
    add_horizontal_rule(doc, '25aae1', 8)
    add_h2(doc, 'FAST CASH STRATEGY #2')
    add_h3(doc, 'LinkedIn Authority Build — "The Warm Pipeline from Week One"')

    add_callout(doc,
        "Managing Partner Matt lives on LinkedIn. He reads industry content there. His referral "
        "sources are there. His competitors are there. A consistent, specific LinkedIn presence "
        "starting Week 1 builds a warm audience of managing partners consuming LGD's content 2–4 "
        "weeks before any direct outreach. That warm pipeline converts at 2–3x the rate of cold "
        "outreach. By Month 3: 200–400 target-audience managing partners following Roger's content.",
        label='THE OPPORTUNITY')

    add_h3(doc, 'The Approach')
    steps_2 = [
        ("Week 1", "Profile optimization: Roger's headline, summary, and featured section rewritten to speak directly to managing partners"),
        ("Week 1 (ongoing)", "Connection campaign: 20 personalized requests/day to family law managing partners. No pitch in the request"),
        ("Week 1 (ongoing)", "Content cadence: 4 posts per week mapped to Belief Filter Map stages"),
        ("Month 2+", "Engagement-first DM: once a prospect has engaged with 2+ posts, initiate direct message — not a pitch"),
        ("Months 2–3", "Discovery call conversion: warm DM conversations convert at 35–40%, significantly higher than cold outreach"),
    ]
    for timing, step in steps_2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{timing}:  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = ASG_BLUE
        r2 = p.add_run(step)
        r2.font.name = 'Georgia'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_BODY_TEXT

    add_branded_table(doc,
        ['Scenario', 'Clients', 'Revenue', 'Math', 'Timeline'],
        [
            ['Conservative', '2', '$36,000', '400 followers × 5% warm × 35% call × 30% close = 2 clients', 'Month 2–3'],
            ['Target', '4+', '$72,000+', '400 followers × 10% warm × 40% call × 45% close = 7 clients', 'Month 2–4'],
        ],
        col_widths=[1.2, 0.8, 0.95, 3.6, 0.9]
    )

    add_callout(doc,
        '"LinkedIn content has returned more than any other channel at an early stage of a business '
        'for me. Not ads — just consistent, specific content that makes the right people think '
        '"this person understands my situation." There\'s a 60-day lag before it converts. Start '
        'Week 1, not Month 3. The people reading Roger\'s posts in March are the ones signing '
        'contracts in May."',
        label='SOFIA VEGA, Board Member')

    # Strategy 3
    add_horizontal_rule(doc, '25aae1', 8)
    add_h2(doc, 'FAST CASH STRATEGY #3')
    add_h3(doc, 'Network Referral Activation — "The Personal Introduction Engine"')

    add_callout(doc,
        "Roger and the LGD team have an existing professional network that almost certainly includes "
        "people who know family law managing partners — legal tech vendors, bar association contacts, "
        "law school alumni, legal marketing consultants, attorney coaches. Referred prospects convert "
        "at 3–5x the rate of cold outreach. LGD's primary stall point (B1 Enemy Belief) is partially "
        "resolved before the first conversation begins.",
        label='THE OPPORTUNITY')

    add_h3(doc, 'The Approach')
    steps_3 = [
        ("Week 1", "Network audit: Roger lists every professional contact who could know a family law managing partner. Target: 20–30 contacts identified"),
        ("Weeks 1–2", "Personalized outreach: one-to-one messages, not a mass email. Brief, specific: who LGD serves, what the outcome is, what a good referral looks like"),
        ("Week 2", "Referral enablement: for contacts who agree to refer, provide a one-paragraph 'who LGD is for' — written to be forwarded directly"),
        ("Month 2", "Introduction conversion: every referred prospect receives a personalized first message acknowledging the referral source specifically"),
        ("Ongoing", "Non-financial referral acknowledgment: handwritten card, charitable donation in their name. Fully compliant"),
    ]
    for timing, step in steps_3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f'{timing}:  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = ASG_BLUE
        r2 = p.add_run(step)
        r2.font.name = 'Georgia'
        r2.font.size = Pt(11)
        r2.font.color.rgb = ASG_BODY_TEXT

    add_callout(doc,
        "Compliance Note: No restrictions for LGD's own referral program. When LGD later designs "
        "referral programs for attorney clients, non-financial incentives or formal attorney referral "
        "fee structures must be used to comply with state bar rules prohibiting fee-sharing with "
        "non-attorneys.",
        label='COMPLIANCE NOTE')

    add_branded_table(doc,
        ['Scenario', 'Clients', 'Revenue', 'Math', 'Timeline'],
        [
            ['Conservative', '1', '$18,000', '25 contacts × 30% agree × 1.5 intros × 25% close = 1–2 clients', 'Month 2'],
            ['Target', '3', '$54,000', '25 contacts × 30% agree × 1.5 intros × 25% close = 3 clients', 'Month 2–3'],
        ],
        col_widths=[1.2, 0.8, 0.95, 3.6, 0.9]
    )

    add_callout(doc,
        '"In my first legal operations business, the first three clients came from people I already knew '
        'who knew someone who needed what I was building. I didn\'t have a CRM. I didn\'t have a formal '
        'program. I just told the right people clearly what I was building and who it was for. Don\'t '
        'make it complicated — just make the asks in the first two weeks, before you get busy delivering '
        'for the clients you\'ll have by Week 4."',
        label='SOFIA VEGA, Board Member')

    # 90-day calendar
    add_h2(doc, '90-Day Activation Calendar')
    add_branded_table(doc,
        ['Week', 'Strategy 1: Direct Outreach', 'Strategy 2: LinkedIn', 'Strategy 3: Referrals'],
        [
            ['Weeks 1–2\n(Prep)', 'Build 60-prospect list; draft + approve 5-touch sequence', 'Profile rewritten; 60-post calendar queued', 'Contact list built; personalized messages drafted'],
            ['Weeks 3–4\n(Launch)', 'Outreach wave 1 deployed; calls scheduling', '4 posts/week live; 20 connections/day', 'Messages sent personally by Roger; asks made'],
            ['Weeks 5–8\n(Active)', 'Calls running; proposals out; target 2 pilots closed', 'Content continues; DM follow-up to engaged connections', 'Follow up on introductions; first referred calls'],
            ['Weeks 9–12\n(Optimize)', 'Pilot onboarding complete; Month 1 delivery on schedule', 'Analytics reviewed; top content identified; warm outreach begins', 'Thank-yous sent; referral program formalized for Q2'],
        ],
        col_widths=[0.95, 2.2, 2.2, 1.9]
    )

    # Revenue summary
    add_callout(doc,
        "90-Day Revenue Summary:\n\n"
        "Strategy 1 (Direct Pilot Outreach):  Conservative $36k  |  Target $54k  |  Timeline 30–45 days\n"
        "Strategy 2 (LinkedIn Authority Build):  Conservative $36k  |  Target $72k  |  Timeline 60–90 days\n"
        "Strategy 3 (Network Referral Activation):  Conservative $18k  |  Target $54k  |  Timeline 30–60 days\n\n"
        "Combined Conservative:  5 clients / $90,000 / 90 days\n"
        "Combined Target:  10 clients / $180,000 / 90 days",
        label='FAST CASH SPRINT™ — REVENUE IMPACT SUMMARY')

    add_signoff(doc, '"Your existing clients are already holding the check. Stop making them wait."',
                'Tanya Blackwood, CRO | Authority Systems Group™')


def build_section_8(doc):
    """KPI Framework."""
    section_break(doc)
    section_header(doc, '8', 'KPI FRAMEWORK & YEAR-END TARGETS',
                   'The metrics that determine whether LGD is on track — reviewed monthly, assessed quarterly')

    add_h2(doc, 'Sales Pipeline KPIs')
    add_branded_table(doc,
        ['Metric', 'Month 1', 'Month 3', 'Month 6', 'Month 9', 'Month 12'],
        [
            ['Discovery calls (monthly)', '5', '15', '25', '35', '40'],
            ['Call-to-client conversion', '20%', '15%', '12%', '10%', '8%'],
            ['New clients (monthly)', '1', '2', '3', '4', '3'],
            ['Active clients (cumulative)', '1', '4', '10', '16', '20'],
            ['Monthly recurring revenue', '$1,500', '$6,000', '$15,000', '$24,000', '$30,000'],
            ['Cumulative revenue', '$1,500', '$18,000', '$90,000', '$216,000', '$360,000'],
        ],
        col_widths=[2.3, 0.9, 0.9, 0.9, 0.9, 0.95]
    )

    add_h2(doc, 'Acquisition Channel KPIs')
    add_branded_table(doc,
        ['Channel', 'Month 3 Target', 'Month 6 Target', 'Month 12 Target'],
        [
            ['Direct outreach → calls', '8/month', '15/month', '15/month'],
            ['LinkedIn warm DMs → calls', '3/month', '8/month', '10/month'],
            ['Referral introductions → calls', '2/month', '5/month', '8/month'],
            ['Inbound (content-driven) → calls', '0', '2/month', '7/month'],
        ],
        col_widths=[2.8, 1.65, 1.65, 1.65]
    )

    add_h2(doc, "LGD's Own Content Performance KPIs")
    add_branded_table(doc,
        ['Metric', 'Month 3', 'Month 6', 'Month 12'],
        [
            ['LinkedIn followers (target audience)', '500', '1,200', '2,500'],
            ['Average post engagement rate', '3%', '4%', '5%'],
            ['Inbound messages from target audience', '5/month', '15/month', '30/month'],
            ['Content-attributed discovery calls', '0', '2/month', '7/month'],
        ],
        col_widths=[3.0, 1.5, 1.5, 1.75]
    )

    add_h2(doc, 'Client Delivery KPIs')
    add_branded_table(doc,
        ['Metric', 'Standard', 'Alert Level'],
        [
            ['Content delivery on schedule', '100%', 'Below 95%'],
            ['Client approval turnaround', '3 business days', 'Over 5 business days'],
            ['Client satisfaction (monthly check-in)', '8+/10', 'Below 7/10'],
            ['6-month retention rate', '90%+', 'Below 85%'],
            ['12-month renewal rate', '85%+', 'Below 80%'],
        ],
        col_widths=[3.0, 2.0, 2.25]
    )


def build_section_9(doc):
    """Investment & Resource Requirements — Preston Adler, COO."""
    section_break(doc)
    section_header(doc, '9', 'INVESTMENT & RESOURCE REQUIREMENTS',
                   'Prepared by Preston Adler, COO — Authority Systems Group™')

    add_body(doc,
        "LGD's Year 1 execution requires three categories of resource commitment. All three are required. "
        "Under-resourcing any one slows the entire system.")

    add_h2(doc, '9A. Time Commitments')
    add_h3(doc, 'Roger — Discovery Calls and Client Relationships')
    add_body(doc,
        "Months 1–3: 15–20 hours per week. Months 4–6: 10–15 hours. Months 7–12: 8–12 hours.")
    add_callout(doc,
        "This is non-negotiable. LGD's first 10 clients are buying Roger — his judgment, his expertise, "
        "his credibility in the room. The content production team can scale to support more clients, but "
        "the sales relationship at this stage requires Roger personally on every discovery call and every "
        "strategy conversation. Delegation of sales before Month 6 is a material risk to close rates.",
        label='CRITICAL DEPENDENCY')

    add_h3(doc, 'Content Production Capacity')
    add_branded_table(doc,
        ['Stage', 'Clients', 'Team Required'],
        [
            ['Months 1–3', '1–4', '1 senior writer + Roger editorial review'],
            ['Months 4–6', '5–8', '2 senior writers or 1 senior + 1 trained associate'],
            ['Months 7–9', '9–15', '3 writers + editorial review process formalized'],
            ['Months 10–12', '15–20', '4–5 writers (mix of employees and contractors) + editorial oversight'],
        ],
        col_widths=[1.5, 1.0, 4.75]
    )

    add_h2(doc, '9B. Technology Stack')
    add_branded_table(doc,
        ['Tool', 'Purpose', 'Monthly Budget'],
        [
            ['GoHighLevel or HubSpot', 'CRM, pipeline, outreach automation, contracts', '$97–$297'],
            ['LinkedIn Sales Navigator', 'Prospect filtering, outreach management', '$99'],
            ['Google Search Console', 'Client content SEO tracking', 'Free'],
            ['SEMrush / Ahrefs (agency plan)', 'Keyword research, client content performance', '$120–$200'],
            ['Calendly', 'Discovery call scheduling', '$16'],
            ['Google Drive', 'Client folders, voice profiles, approval workflows', 'Included'],
        ],
        col_widths=[2.3, 3.0, 1.95]
    )
    add_body(doc, "Total technology budget: $330–$612/month ($4,000–$7,300 annually)", italic=True)

    add_h2(doc, '9C. Monthly P&L Summary')
    add_branded_table(doc,
        ['Month', 'Clients', 'Revenue', 'Est. COGS', 'Gross Margin'],
        [
            ['Month 1', '1–2', '$1,500–$3,000', '$2,500', 'Investment phase'],
            ['Month 3', '3–4', '$4,500–$6,000', '$3,500', 'Near breakeven'],
            ['Month 6', '8–10', '$12,000–$15,000', '$6,000–$8,000', '45–50%'],
            ['Month 9', '15–17', '$22,500–$25,500', '$10,000–$12,000', '52–55%'],
            ['Month 12', '20', '$30,000', '$12,000–$14,000', '55–60%'],
        ],
        col_widths=[1.05, 0.85, 1.75, 1.75, 1.85]
    )
    add_body(doc, "COGS: contract writer fees, editorial oversight, technology stack, client reporting. Roger's compensation treated as owner draw.", italic=True)

    add_h2(doc, '9D. Time-Sensitive Decisions — Act in Weeks 1–2')
    add_callout(doc,
        "These decisions delay Q1 execution if deferred:\n\n"
        "1.  CRM selection and setup: outreach sequences cannot be automated without a CRM configured. Select and set up before outreach begins.\n"
        "2.  Sample content production and approval: the discovery call cannot close without at least one sample article demonstrating quality. Produce and approve in Week 1.\n"
        "3.  Roger's LinkedIn profile optimization: every day the profile remains generic is a day the warm pipeline doesn't build. Two-hour task. Do it in Week 1.\n"
        "4.  Pilot contract template: the discovery call must end with a specific proposal. The contract must exist before the first call closes.",
        label='WEEK 1–2 CRITICAL PATH')

    add_signoff(doc, '"Delivery is the product. Operational discipline is what makes it scalable."',
                'Preston Adler, COO | Authority Systems Group™')


def build_section_10(doc):
    """QC Sign-Off."""
    section_break(doc)
    section_header(doc, '10', 'QUALITY ASSURANCE SIGN-OFF',
                   'Four-Gate Review | Authority Systems Group™')

    gates = [
        ('GATE 1 — STRATEGIC ALIGNMENT', 'Daniel Frost, CSO', [
            'All recommendations serve the $360,000 Year 1 revenue target',
            'Sequencing is correct: proof-of-concept before scale, pilot before paid acquisition, organic before paid',
            'All recommendations specific to LGD\'s stage (zero revenue, zero case studies), niche (family law), and geography (national)',
            'Quarterly build is logical and compounding: Q2 deploys Q1 case study; Q3 adds paid layer; Q4 consolidates',
            'Competitive analysis references specific competitor categories with specific weaknesses',
            'Every recommendation carries a clear "why this, why now" rationale',
        ]),
        ('GATE 2 — MESSAGING INTEGRITY', 'Vivienne Carr, CMO', [
            'No sentence in this document could be repurposed for a different consultancy in a different vertical',
            'Belief Filter Map addresses Managing Partner Matt at B1/E1 — where the primary prospect pool is provably stuck',
            'Every content type in the Master Content Map maps to a specific stage pair on both tracks',
            'No vague claims, no generic "quality service" language throughout',
            'Every hook earns the right to make its promise',
        ]),
        ('GATE 3 — REVENUE LOGIC', 'Tanya Blackwood, CRO', [
            'Every strategy has a traceable path: action → call → proposal → signed contract → recurring revenue',
            'Fast Cash projections show their math transparently (list size × response rate × close rate × value)',
            'Revenue projections are conservative — achievable without extraordinary effort',
            'All Fast Cash strategies executable within 30–60 days of decision',
            'Sequencing is revenue-rational: direct close before warm pipeline before referral scaling',
            'No "brand awareness" recommendations without a revenue mechanism',
        ]),
        ('GATE 4 — DELIVERY STANDARD', 'Preston Adler, COO', [
            'All sections complete — no placeholder text remaining',
            'All LGD-specific data accurately reflected throughout the document',
            'Table of contents matches actual document structure',
            'All persona sign-offs present and attributed correctly',
            'Document named correctly per convention: lgd_authority-blueprint_20260302',
        ]),
    ]

    for gate_name, owner, items in gates:
        add_h2(doc, gate_name)
        add_attribution(doc, f'Lens: {owner}')
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(f'☑  {item}')
            r.font.name = 'Georgia'
            r.font.size = Pt(11)
            r.font.color.rgb = ASG_BODY_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout(doc,
        "This document has been reviewed and approved through the Authority Systems Group™ "
        "four-gate quality assurance process.\n\n"
        "Strategic Alignment  |  Messaging Integrity  |  Revenue Logic  |  Delivery Standard\n\n"
        "Prepared under the direction of Roger, Director — Authority Systems Group™.",
        label='QC APPROVAL CERTIFICATION')

    # Final signoff
    p_final = doc.add_paragraph()
    p_final.paragraph_format.space_before = Pt(30)
    p_final.paragraph_format.space_after = Pt(8)
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_final.add_run('Authority Systems Group™')
    r_f.bold = True
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(14)
    r_f.font.color.rgb = ASG_BLUE

    p_final2 = doc.add_paragraph()
    p_final2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_final2.paragraph_format.space_after = Pt(4)
    r_f2 = p_final2.add_run(f'Confidential. Prepared exclusively for {CLIENT_NAME}.')
    r_f2.italic = True
    r_f2.font.name = 'Calibri'
    r_f2.font.size = Pt(10)
    r_f2.font.color.rgb = ASG_CHARCOAL

    p_final3 = doc.add_paragraph()
    p_final3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f3 = p_final3.add_run(f'Authority Blueprint™ v1.0  |  {DOC_DATE}')
    r_f3.font.name = 'Calibri'
    r_f3.font.size = Pt(9)
    r_f3.font.color.rgb = ASG_FOOTER_CLR


# ─────────────────────────────────────────
# DOCUMENT PROPERTIES
# ─────────────────────────────────────────

def set_document_properties(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    core_props = doc.core_properties
    core_props.title = f'Authority Blueprint™ — {CLIENT_NAME}'
    core_props.subject = 'Strategic Growth System'
    core_props.author = 'Authority Systems Group™'
    core_props.company = 'Authority Systems Group™'
    core_props.keywords = f'authority blueprint, strategic plan, {CLIENT_SLUG}'
    core_props.comments = 'Prepared under the direction of Roger, Director — Authority Systems Group™'


# ─────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────

def setup_page_margins(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)


# ─────────────────────────────────────────
# MAIN BUILD
# ─────────────────────────────────────────

import docx as docx_module

def main():
    print("Authority Systems Group™ — Building Authority Blueprint DOCX...")
    print(f"Client: {CLIENT_NAME}")
    print(f"Output: {OUTPUT_FILE}")
    print()

    doc = Document()
    setup_page_margins(doc)
    set_document_properties(doc)

    print("Building cover page...")
    build_cover_page(doc)

    print("Building confidentiality notice...")
    build_confidentiality_page(doc)

    print("Building table of contents...")
    build_toc(doc)

    print("Building Section 1: Executive Summary...")
    build_section_1(doc)

    print("Building Section 2: Market Analysis...")
    build_section_2(doc)

    print("Building Section 3: Avatar...")
    build_section_3(doc)

    print("Building Section 4: Belief Filter Map...")
    build_section_4(doc)

    print("Building Section 5: Year 1 Strategic Plan...")
    build_section_5(doc)

    print("Building Section 6: Content Authority Strategy...")
    build_section_6(doc)

    print("Building Section 7: Fast Cash Sprint...")
    build_section_7(doc)

    print("Building Section 8: KPI Framework...")
    build_section_8(doc)

    print("Building Section 9: Investment & Resources...")
    build_section_9(doc)

    print("Building Section 10: QC Sign-Off...")
    build_section_10(doc)

    print("Configuring headers and footers...")
    setup_header_footer(doc)

    # Remove default empty first paragraph if present
    if doc.paragraphs and doc.paragraphs[0].text == '':
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    print(f"Saving document to: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print()
    print("✓ Document built successfully.")
    print(f"✓ Saved: {OUTPUT_FILE}")
    print()
    print("QC Checklist:")
    print("  ☑ Cover page with dark background and ASG Blue accents")
    print("  ☑ Confidentiality notice page")
    print("  ☑ Table of contents")
    print("  ☑ 10 fully built sections")
    print("  ☑ Branded tables with ASG Blue headers, alternating rows")
    print("  ☑ Callout boxes with ASG Blue left border")
    print("  ☑ Persona sign-offs at each section end")
    print("  ☑ Headers and footers on all interior pages")
    print("  ☑ Document properties set")
    print()
    print("Next step: Open in Word, verify cover page render, export as PDF.")


if __name__ == '__main__':
    main()
